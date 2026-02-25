#!/usr/bin/env python3
"""
MLGH Data Viewer - Shows the MLGH datatracker data from test files
This displays the Meta-Layer Task Force data so you can see it working.

⚠️ CRITICAL: THIS IS THE MLGH VERSION - DO NOT REVERT TO IETF ⚠️
If you see "IETF Data Viewer" in the docstring, this file has been reverted incorrectly.
The correct version should say "MLGH Data Viewer" and "Meta-Layer Task Force".

BUILD: 1
Last Updated: 2026-01-23 (Ordinals integration with markdown detection)
"""

# Build number for cache busting and version tracking (increments from file on each run)
def _load_and_increment_build_number():
    import os
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'instance_dev', 'build_number.txt')
    try:
        with open(path, 'r') as f:
            n = int(f.read().strip())
    except (FileNotFoundError, ValueError):
        n = 74
    try:
        with open(path, 'w') as f:
            f.write(str(n + 1))
    except Exception:
        pass
    return n

BUILD_NUMBER = _load_and_increment_build_number()

def create_hypothesis_account(user):
    """Create a Hypothesis account for a Meta-Layer user via API"""
    import requests
    
    # Check if user already has a Hypothesis account
    existing = HypothesisAccount.query.filter_by(user_id=user['id']).first()
    if existing:
        return existing
    
    if not HYPOTHESIS_CONFIG.get('API_TOKEN'):
        app.logger.error("No Hypothesis API token configured")
        return None
    
    # Generate unique username
    base_username = user.get('displayName', user.get('username', f'user{user["id"]}'))
    # Clean username (Hypothesis requirements: alphanumeric + hyphens + underscores)
    clean_username = ''.join(c for c in base_username if c.isalnum() or c in '-_').lower()
    if not clean_username:
        clean_username = f'mluser{user["id"]}'
    
    # Ensure uniqueness by adding timestamp
    import time
    username = f"{clean_username}_{int(time.time())}"
    
    try:
        # Create user via Hypothesis API
        headers = {
            'Authorization': f'Bearer {HYPOTHESIS_CONFIG["API_TOKEN"]}',
            'Content-Type': 'application/json'
        }
        
        # Use the standard hypothes.is authority
        hypothesis_userid = f"acct:{username}@hypothes.is"
        
        # Create user payload
        user_data = {
            'authority': 'hypothes.is',
            'username': username,
            'email': user.get('email', f'{username}@rfc.themetalayer.org'),
            'display_name': user.get('displayName', username)
        }
        
        # Note: The Hypothesis API doesn't have a direct user creation endpoint
        # We'll store the mapping and let users authenticate normally
        
        # Store the account link
        hypothesis_account = HypothesisAccount(
            user_id=user['id'],
            hypothesis_username=username,
            hypothesis_userid=hypothesis_userid
        )
        db.session.add(hypothesis_account)
        db.session.commit()
        
        app.logger.info(f"Created Hypothesis account mapping for user {user['id']}: {username}")
        return hypothesis_account
        
    except Exception as e:
        app.logger.error(f"Failed to create Hypothesis account for user {user['id']}: {e}")
        return None

def get_document_annotations(document_name, document_type='draft'):
    """Fetch existing annotations for a document using Hypothesis API"""
    import requests
    
    if not HYPOTHESIS_CONFIG.get('API_TOKEN'):
        return []
    
    try:
        headers = {
            'Authorization': f'Bearer {HYPOTHESIS_CONFIG["API_TOKEN"]}',
            'Content-Type': 'application/json'
        }
        
        # Search for annotations with document-specific tags
        tag = f"{document_type}:{document_name}"
        url = f"{HYPOTHESIS_CONFIG['API_URL']}/search"
        params = {
            'tag': tag,
            'limit': 200  # Maximum per request
        }
        
        response = requests.get(url, headers=headers, params=params, timeout=10)
        if response.status_code == 200:
            data = response.json()
            return data.get('rows', [])
        else:
            app.logger.warning(f"Failed to fetch annotations: {response.status_code}")
            return []
            
    except Exception as e:
        app.logger.error(f"Error fetching annotations: {e}")
        return []

def create_annotation_via_api(document_name, document_type, text, quote, user):
    """Create annotation via Hypothesis API.
    WARNING: Uses the server's API token, so the annotation would be attributed to the
    token owner (you), NOT the end user. Do NOT use for user-created content.
    Only use for system/bot annotations if ever needed. User annotations are created
    by the Hypothesis client in the browser under each user's own account.
    """
    import requests
    
    if not HYPOTHESIS_CONFIG.get('API_TOKEN'):
        return None
    
    try:
        headers = {
            'Authorization': f'Bearer {HYPOTHESIS_CONFIG["API_TOKEN"]}',
            'Content-Type': 'application/json'
        }
        
        # Create annotation payload
        annotation_data = {
            'uri': f'https://dev.rfc.themetalayer.org/doc/{document_type}/{document_name}/',
            'text': text,
            'tags': [f'{document_type}:{document_name}', f'meta-layer:{document_type}'],
            'target': [{
                'source': f'https://dev.rfc.themetalayer.org/doc/{document_type}/{document_name}/',
                'selector': [{
                    'type': 'TextQuoteSelector',
                    'exact': quote
                }]
            }],
            'permissions': {
                'read': ['group:__world__'],
                'update': [f'acct:{user.get("username", "anonymous")}@hypothes.is'],
                'delete': [f'acct:{user.get("username", "anonymous")}@hypothes.is'],
                'admin': [f'acct:{user.get("username", "anonymous")}@hypothes.is']
            }
        }
        
        response = requests.post(
            f"{HYPOTHESIS_CONFIG['API_URL']}/annotations",
            headers=headers,
            json=annotation_data,
            timeout=10
        )
        
        if response.status_code == 200:
            return response.json()
        else:
            app.logger.warning(f"Failed to create annotation: {response.status_code}")
            return None
            
    except Exception as e:
        app.logger.error(f"Error creating annotation: {e}")
        return None

def generate_hypothesis_config(document_name=None, document_type='draft'):
    """Generate Hypothesis configuration HTML for document pages"""
    if not HYPOTHESIS_ENABLED:
        return ""
    
    # Check if user has annotations enabled (via cookie)
    annotations_enabled = request.cookies.get('annotations', 'off') == 'on'
    if not annotations_enabled:
        return ""
    
    # Get current user
    current_user = get_current_user()
    
    # Generate document-specific tags
    if document_type == 'draft':
        # For drafts, include document name for revision-specific annotations
        tags = f'["draft:{document_name}", "meta-layer:draft"]'
    else:
        # For other document types
        tags = f'["{document_type}:{document_name}", "meta-layer:{document_type}"]'
    
    # For now, use standard Hypothesis (users create their own accounts)
    # TODO: Implement full API integration when we get Hypothesis approval
    auth_config = ""
    
    return f"""
    <script>
    window.hypothesisConfig = function () {{
      return {{
        branding: {{
          appBackgroundColor: '{HYPOTHESIS_CONFIG['BRANDING']['appBackgroundColor']}',
          ctaBackgroundColor: '{HYPOTHESIS_CONFIG['BRANDING']['ctaBackgroundColor']}',
          ctaTextColor: '{HYPOTHESIS_CONFIG['BRANDING']['ctaTextColor']}',
          selectionFontFamily: '{HYPOTHESIS_CONFIG['BRANDING']['selectionFontFamily']}'
        }},
        enableExperimentalNewNoteButton: {str(HYPOTHESIS_CONFIG['ENABLE_EXPERIMENTAL_NEW_NOTE_BUTTON']).lower()},
        showHighlights: '{HYPOTHESIS_CONFIG['SHOW_HIGHLIGHTS']}',
        openSidebar: false,{auth_config}
        // Focus on document-specific annotations
        focus: {{
          user: {{
            filter: {{
              any: {{
                tag: {tags}
              }}
            }}
          }}
        }}
      }};
    }};
    </script>
    <script async src="{HYPOTHESIS_CONFIG['EMBED_URL']}"></script>
    """

from flask import Flask, render_template_string, request, redirect, url_for, flash, session, send_file, send_from_directory, jsonify, g
from flask_sqlalchemy import SQLAlchemy
import os
import re
from dotenv import load_dotenv

# Load environment variables
load_dotenv('/home/ubuntu/xowlz/burned/.env')

# Hypothesis Annotation Configuration
HYPOTHESIS_ENABLED = True  # Set to False to disable annotations globally
HYPOTHESIS_CONFIG = {
    'EMBED_URL': 'https://hypothes.is/embed.js',
    'API_URL': 'https://hypothes.is/api',
    'API_TOKEN': os.getenv('HYPOTHESIS_API_TOKEN'),  # Server-only: read/count; never sent to client
    'AUTHORITY': 'hypothes.is',  # Use hypothes.is authority for now
    'BRANDING': {
        'appBackgroundColor': '#16181c',  # Dark theme background
        'ctaBackgroundColor': '#1d9bf0',  # Meta-Layer accent color
        'ctaTextColor': '#ffffff',
        'selectionFontFamily': '-apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif',
    },
    'ENABLE_EXPERIMENTAL_NEW_NOTE_BUTTON': True,
    'SHOW_HIGHLIGHTS': 'whenSidebarOpen',
}
import json
import uuid
import requests
from datetime import datetime, timedelta
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from collections import defaultdict
import time

# Import file processing libraries
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import markdown2
    import bleach
    MARKDOWN_SUPPORT = True
except ImportError:
    MARKDOWN_SUPPORT = False

# Rate limiting for security
rate_limit_store = defaultdict(list)

def check_rate_limit(identifier, max_requests=5, window_seconds=300):
    """Simple in-memory rate limiting"""
    now = time.time()
    key = f"{identifier}"

    # Clean old entries
    rate_limit_store[key] = [timestamp for timestamp in rate_limit_store[key]
                           if now - timestamp < window_seconds]

    # Check if under limit
    if len(rate_limit_store[key]) >= max_requests:
        return False

    # Add current request
    rate_limit_store[key].append(now)
    return True

# Database initialization
def init_db():
    """Initialize database and create tables"""
    with app.app_context():
        db.create_all()

        # Run database migrations for ordinals support
        migrate_ordinals_support()
        
        # Check if we need to migrate for Hypothesis accounts
        try:
            # Test if HypothesisAccount table exists
            db.session.execute(db.text("SELECT 1 FROM hypothesis_account LIMIT 1"))
            print("✅ HypothesisAccount table exists")
        except:
            print("🔄 Creating HypothesisAccount table...")
            db.create_all()
            print("✅ HypothesisAccount table created")
        
        # Add banner_image, headline, bio, social_links to user table if missing
        try:
            import sqlite3
            db_path = app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute("PRAGMA table_info(user)")
            user_columns = [c[1] for c in cursor.fetchall()]
            if 'banner_image' not in user_columns:
                cursor.execute("ALTER TABLE user ADD COLUMN banner_image VARCHAR(500)")
                conn.commit()
                print("✅ Added banner_image column to user table")
            if 'headline' not in user_columns:
                cursor.execute("ALTER TABLE user ADD COLUMN headline VARCHAR(200)")
                conn.commit()
                print("✅ Added headline column to user table")
            if 'bio' not in user_columns:
                cursor.execute("ALTER TABLE user ADD COLUMN bio TEXT")
                conn.commit()
                print("✅ Added bio column to user table")
            if 'social_links' not in user_columns:
                cursor.execute("ALTER TABLE user ADD COLUMN social_links TEXT")
                conn.commit()
                print("✅ Added social_links column to user table")
            if 'referral_code' not in user_columns:
                cursor.execute("ALTER TABLE user ADD COLUMN referral_code VARCHAR(50)")
                conn.commit()
                print("✅ Added referral_code column to user table")
            conn.close()
        except Exception as e:
            print(f"⚠️  Error adding profile columns: {e}")
        
        # --- public_id migration (MUST run before any User.query calls) ---
        try:
            import sqlite3
            from uuid import uuid4
            db_path = app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            for table_name in ['user', 'project', 'submission', 'badge']:
                try:
                    cursor.execute(f"SELECT public_id FROM {table_name} LIMIT 1")
                    print(f"✅ public_id already exists on {table_name}")
                except sqlite3.OperationalError:
                    print(f"🔄 Adding public_id to {table_name}...")
                    cursor.execute(f"ALTER TABLE {table_name} ADD COLUMN public_id VARCHAR(36)")
                    conn.commit()
                    
                    # Backfill existing rows
                    cursor.execute(f"SELECT id FROM {table_name} WHERE public_id IS NULL")
                    rows = cursor.fetchall()
                    for row in rows:
                        cursor.execute(
                            f"UPDATE {table_name} SET public_id = ? WHERE id = ?",
                            (str(uuid4()), row[0])
                        )
                    conn.commit()
                    print(f"✅ Backfilled {len(rows)} rows in {table_name}")
                    
                    # Add unique index
                    try:
                        cursor.execute(f"CREATE UNIQUE INDEX idx_{table_name}_public_id ON {table_name}(public_id)")
                        conn.commit()
                        print(f"✅ Created unique index on {table_name}.public_id")
                    except sqlite3.OperationalError:
                        pass  # Index already exists
            
            conn.close()
        except Exception as e:
            print(f"⚠️  Error adding public_id columns: {e}")
        
        # Ensure project_member table exists
        try:
            db.session.execute(db.text("SELECT 1 FROM project_member LIMIT 1"))
            print("✅ project_member table exists")
        except Exception:
            db.create_all()
            print("✅ project_member table created")
        
        # Ensure waitlist tables exist
        try:
            db.session.execute(db.text("SELECT 1 FROM waitlist LIMIT 1"))
            print("✅ waitlist tables exist")
        except Exception:
            db.create_all()
            print("✅ waitlist tables created")
        
        # Add image_url columns to project, workgroup, guild, waitlist
        try:
            import sqlite3
            conn = sqlite3.connect(app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', ''))
            cursor = conn.cursor()
            
            # Check and add image_url to project
            cursor.execute("PRAGMA table_info(project)")
            project_columns = [c[1] for c in cursor.fetchall()]
            if 'image_url' not in project_columns:
                cursor.execute("ALTER TABLE project ADD COLUMN image_url VARCHAR(500)")
                conn.commit()
                print("✅ Added image_url column to project table")
            
            # Check and add image_url to waitlist
            cursor.execute("PRAGMA table_info(waitlist)")
            waitlist_columns = [c[1] for c in cursor.fetchall()]
            if 'image_url' not in waitlist_columns:
                cursor.execute("ALTER TABLE waitlist ADD COLUMN image_url VARCHAR(500)")
                conn.commit()
                print("✅ Added image_url column to waitlist table")
            
            # Check and add image_url to working_group
            cursor.execute("PRAGMA table_info(working_group)")
            wg_columns = [c[1] for c in cursor.fetchall()]
            if 'image_url' not in wg_columns:
                cursor.execute("ALTER TABLE working_group ADD COLUMN image_url VARCHAR(500)")
                conn.commit()
                print("✅ Added image_url column to working_group table")
            
            # Check and add image_url to guild
            cursor.execute("PRAGMA table_info(guild)")
            guild_columns = [c[1] for c in cursor.fetchall()]
            if 'image_url' not in guild_columns:
                cursor.execute("ALTER TABLE guild ADD COLUMN image_url VARCHAR(500)")
                conn.commit()
                print("✅ Added image_url column to guild table")
            
            conn.close()
        except Exception as e:
            print(f"⚠️  Error adding image_url columns: {e}")

        # Migrate hardcoded users to database if not already done
        if User.query.count() == 0:
            migrate_hardcoded_users()

        # Note: DRAFTS list is now populated from approved/published submissions dynamically
        # No need to pre-load from a separate table
        migrate_coordinator_and_member_requests()
        
        print(f"Database initialized: {User.query.count()} users")

def migrate_coordinator_and_member_requests():
    """Ensure coordinator_request, workgroup_member_request tables exist; add user_id to working_group_chair and working_group_member."""
    try:
        import sqlite3
        db_path = app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        # Add user_id to working_group_chair if missing
        cursor.execute("PRAGMA table_info(working_group_chair)")
        wgc_columns = [c[1] for c in cursor.fetchall()]
        if 'user_id' not in wgc_columns:
            cursor.execute("ALTER TABLE working_group_chair ADD COLUMN user_id INTEGER REFERENCES user(id)")
            conn.commit()
        # Add user_id to working_group_member if missing (membership by id)
        cursor.execute("PRAGMA table_info(working_group_member)")
        wgm_columns = [c[1] for c in cursor.fetchall()]
        if 'user_id' not in wgm_columns:
            cursor.execute("ALTER TABLE working_group_member ADD COLUMN user_id INTEGER REFERENCES user(id)")
            conn.commit()
            # Backfill user_id from user_name where we can match a user
            cursor.execute("SELECT id, group_acronym, user_name FROM working_group_member")
            for row in cursor.fetchall():
                mid, gac, uname = row
                cursor.execute("SELECT id FROM user WHERE username = ? OR name = ? OR displayName = ? OR oauthName = ? LIMIT 1", (uname, uname, uname, uname))
                urow = cursor.fetchone()
                if urow:
                    cursor.execute("UPDATE working_group_member SET user_id = ? WHERE id = ?", (urow[0], mid))
            conn.commit()
        conn.close()
    except Exception as e:
        print(f"Migration coordinator/member_requests: {e}")

def migrate_ordinals_support():
    """Add ordinals support columns to existing submission table"""
    try:
        import sqlite3
        db_path = app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Check if ordinals columns exist
        cursor.execute("PRAGMA table_info(submission)")
        columns = [col[1] for col in cursor.fetchall()]
        
        ordinals_columns = {
            'sourceType': ('TEXT', 'file'),
            'ordinalId': ('TEXT', None),
            'ordinalContentUrl': ('TEXT', None),
            'ordinalContentType': ('TEXT', None),
            'inscriptionNumber': ('INTEGER', None),
            'blockHeight': ('INTEGER', None),
            'inscriptionTimestamp': ('DATETIME', None),
            'doc_type': ('TEXT', 'draft')
        }
        
        added_columns = []
        for col_name, (col_type, default_value) in ordinals_columns.items():
            if col_name not in columns:
                if default_value:
                    cursor.execute(f"ALTER TABLE submission ADD COLUMN {col_name} {col_type} DEFAULT '{default_value}'")
                else:
                    cursor.execute(f"ALTER TABLE submission ADD COLUMN {col_name} {col_type}")
                added_columns.append(col_name)
        
        # Migrate existing ML numbers to new format (ML-001 -> ML-Draft-001)
        cursor.execute("SELECT id, ml_number FROM submission WHERE ml_number IS NOT NULL")
        submissions = cursor.fetchall()
        migrated_count = 0
        for sub_id, ml_num in submissions:
            if ml_num and not ml_num.startswith('ML-Draft-') and not ml_num.startswith('ML-RFC-'):
                # Old format: ML-001, ML-002, etc.
                # Extract number and convert to ML-Draft-XXX
                try:
                    num_part = ml_num.split('-')[-1]
                    new_ml_num = f"ML-Draft-{num_part}"
                    cursor.execute("UPDATE submission SET ml_number = ? WHERE id = ?", (new_ml_num, sub_id))
                    migrated_count += 1
                except (ValueError, IndexError):
                    pass
        
        conn.commit()
        conn.close()
        
        if added_columns:
            print(f"✅ Added columns to submission table: {', '.join(added_columns)}")
        else:
            print("✅ All columns already exist in submission table")
        
        if migrated_count > 0:
            print(f"✅ Migrated {migrated_count} ML numbers to new format (ML-Draft-XXX)")
    except Exception as e:
        print(f"⚠️  Error migrating ordinals support: {e}")
        # Non-fatal - table might already have columns

def migrate_hardcoded_users():
    """Migrate hardcoded users to database"""
    hardcoded_users = {
        'admin': {'password': 'admin123', 'name': 'Admin User', 'email': 'admin@metalayer.org', 'role': 'admin', 'theme': 'dark'},
        'daveed': {'password': 'admin123', 'name': 'Daveed', 'email': 'daveed@bridgit.io', 'role': 'admin', 'theme': 'dark'},
        'john': {'password': 'password123', 'name': 'John Doe', 'email': 'john@example.com', 'role': 'editor', 'theme': 'dark'},
        'jane': {'password': 'password123', 'name': 'Jane Smith', 'email': 'jane@example.com', 'role': 'user', 'theme': 'dark'},
        'shiftshapr': {'password': 'mynewpassword123', 'name': 'Shift Shapr', 'email': 'shiftshapr@example.com', 'role': 'editor', 'theme': 'dark'}
    }

    for username, user_data in hardcoded_users.items():
        if not User.query.filter_by(username=username).first():
            user = User(
                username=username,
                password_hash=generate_password_hash(user_data['password']),
                name=user_data['name'],
                email=user_data['email'],
                role=user_data.get('role', 'user'),
                theme=user_data.get('theme', 'dark')
            )
            db.session.add(user)

    db.session.commit()
    print(f"Migrated {len(hardcoded_users)} hardcoded users to database")

# Environment configuration
ENV = os.environ.get('FLASK_ENV', 'production').lower()
IS_PRODUCTION = ENV == 'production'
IS_DEVELOPMENT = ENV == 'development'

# Set up paths based on environment
if IS_DEVELOPMENT:
    INSTANCE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'instance_dev')
    DB_NAME = 'datatracker_dev.db'
    PORT = int(os.environ.get('FLASK_PORT', 8001))
    DEBUG = True
else:
    INSTANCE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'instance')
    DB_NAME = 'datatracker.db'
    PORT = int(os.environ.get('FLASK_PORT', 8000))
    DEBUG = False

# Host → Layer middleware configuration
RESERVED_SUBDOMAINS = {
    "dev", "rfc", "www", "api", "static",
    "assets", "admin", "staging", "beta"
}
BASE_DOMAIN = "themetalayer.org"

# DEPLOYMENT SAFETY - Block data modifications during deployment
deployment_flag_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), f".deployment_{'dev' if IS_DEVELOPMENT else 'prod'}")
DEPLOYMENT_MODE = os.path.exists(deployment_flag_file)
if DEPLOYMENT_MODE:
    print("🚨 DEPLOYMENT MODE ENABLED - Data modifications blocked")

def check_deployment_safety(operation="database operation"):
    """Check if operations are allowed during deployment"""
    if DEPLOYMENT_MODE:
        error_msg = f"🚨 BLOCKED: {operation} not allowed during deployment"
        print(error_msg)
        raise RuntimeError(error_msg)

def init_deployment_safety():
    """Initialize deployment safety checks after database is set up"""
    if DEPLOYMENT_MODE:
        # Override SQLAlchemy session methods to check deployment safety
        global original_add, original_commit, original_delete, original_create_all
        original_add = db.session.add
        original_commit = db.session.commit
        original_delete = db.session.delete
        original_create_all = db.create_all

        def safe_add(instance):
            check_deployment_safety("database add operation")
            return original_add(instance)

        def safe_commit():
            check_deployment_safety("database commit operation")
            return original_commit()

        def safe_delete(instance):
            check_deployment_safety("database delete operation")
            return original_delete(instance)

        def safe_create_all(*args, **kwargs):
            check_deployment_safety("database schema creation")
            return original_create_all(*args, **kwargs)

        # Monkey patch the session and DDL methods
        db.session.add = safe_add
        db.session.commit = safe_commit
        db.session.delete = safe_delete
        db.create_all = safe_create_all
        print("🚨 Database operations blocked during deployment")

# Ensure instance directory exists
os.makedirs(INSTANCE_DIR, exist_ok=True)

app = Flask(__name__, instance_path=INSTANCE_DIR, instance_relative_config=True)
app.secret_key = os.environ.get('SECRET_KEY', 'your-secret-key-here-change-in-production')  # For flash messages

# Database setup
DB_PATH = os.path.join(INSTANCE_DIR, DB_NAME)
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{DB_PATH}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['DEBUG'] = DEBUG

# Session security configuration
app.config['SESSION_COOKIE_SECURE'] = not IS_DEVELOPMENT  # HTTPS only in production
app.config['SESSION_COOKIE_HTTPONLY'] = True  # Prevent XSS access to session cookie
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # CSRF protection

db = SQLAlchemy(app)

# Database Models
class Submission(db.Model):
    id = db.Column(db.String(8), primary_key=True)
    public_id = db.Column(db.String(36), unique=True, nullable=False, default=lambda: str(uuid4()))
    title = db.Column(db.String(255))
    authors = db.Column(db.JSON)  # List of author dicts
    abstract = db.Column(db.Text)
    group = db.Column(db.String(50))
    project_id = db.Column(db.String(50), db.ForeignKey('project.id'), nullable=True, index=True)
    filename = db.Column(db.String(255))
    file_path = db.Column(db.String(500))
    draft_name = db.Column(db.String(255))
    status = db.Column(db.String(20), default='submitted')  # submitted, approved, rejected
    submitted_at = db.Column(db.DateTime, default=datetime.utcnow)
    submitted_by = db.Column(db.String(100), default='Anonymous User')
    approved_at = db.Column(db.DateTime, nullable=True)
    rejected_at = db.Column(db.DateTime, nullable=True)
    ml_number = db.Column(db.String(20), nullable=True)  # ML-Draft-001, ML-RFC-001, etc.
    doc_type = db.Column(db.String(10), default='draft')  # 'draft' or 'rfc'
    pages = db.Column(db.Integer, default=1)  # Calculated page count
    words = db.Column(db.Integer, default=0)  # Calculated word count
    # Ordinal integration fields
    sourceType = db.Column(db.String(20), default='file')  # 'file' or 'ordinal'
    ordinalId = db.Column(db.String(255), nullable=True)  # Inscription ID
    ordinalContentUrl = db.Column(db.String(500), nullable=True)  # URL to content
    ordinalContentType = db.Column(db.String(100), nullable=True)  # MIME type
    inscriptionNumber = db.Column(db.Integer, nullable=True)  # Ordinal inscription number
    blockHeight = db.Column(db.Integer, nullable=True)  # Bitcoin block height
    inscriptionTimestamp = db.Column(db.DateTime, nullable=True)  # When inscribed
    # Revision fields
    parent_draft_name = db.Column(db.String(255), nullable=True)  # Link to parent draft for revisions
    revision_number = db.Column(db.String(10), nullable=True)  # e.g., "01", "02"
    what_changed = db.Column(db.Text, nullable=True)  # Description of changes in this revision
    is_revision = db.Column(db.Boolean, default=False)  # Flag to indicate this is a revision
    # RFC publication field
    rfc_number = db.Column(db.Integer, nullable=True)  # RFC number when status='published'


class Comment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    draft_name = db.Column(db.String(255), index=True)
    text = db.Column(db.Text)
    author = db.Column(db.String(100))
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)
    parent_id = db.Column(db.Integer, db.ForeignKey('comment.id'), nullable=True)
    edited_at = db.Column(db.DateTime, nullable=True)
    is_deleted = db.Column(db.Boolean, default=False)
    original_text = db.Column(db.Text, nullable=True)  # Store original text for edit history

    # Relationship for replies
    replies = db.relationship('Comment', backref=db.backref('parent', remote_side=[id]), lazy=True)

class DocumentHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    draft_name = db.Column(db.String(255), index=True)
    action = db.Column(db.String(50))
    user = db.Column(db.String(100))
    details = db.Column(db.Text)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

class WorkingGroupMember(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    group_acronym = db.Column(db.String(50), index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)  # primary lookup
    user_name = db.Column(db.String(100), index=True)  # for display
    joined_at = db.Column(db.DateTime, default=datetime.utcnow)

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    public_id = db.Column(db.String(36), unique=True, nullable=False, default=lambda: str(uuid4()))
    username = db.Column(db.String(50), unique=True, index=True)
    password_hash = db.Column(db.String(255))

    # Web3Auth fields
    web3authVerifierId = db.Column(db.String(255), unique=True, index=True)  # Web3Auth unique identifier
    typeOfLogin = db.Column(db.String(50))  # 'google', 'wallet', 'twitter', 'email_passwordless'

    # Display name system (user-editable)
    displayName = db.Column(db.String(50))  # User's chosen display name
    displayNameSetAt = db.Column(db.DateTime)  # When user first set/changed it

    # OAuth data (read-only reference)
    oauthName = db.Column(db.String(100))  # Original name from OAuth provider
    name = db.Column(db.String(100))  # Legacy field - will be deprecated
    email = db.Column(db.String(100), unique=True, index=True)
    profileImage = db.Column(db.String(500))  # Avatar URL

    # Wallet data (always visible)
    evmAddress = db.Column(db.String(42), unique=True, index=True)  # EVM wallet address
    solanaAddress = db.Column(db.String(44), unique=True, index=True)  # Solana wallet address

    # Handle (unique identifier)
    handle = db.Column(db.String(50), unique=True, index=True)  # Unique handle for user

    # Profile customization
    banner_image = db.Column(db.String(500))  # Banner/header image URL
    headline = db.Column(db.String(200))  # Short headline/tagline
    bio = db.Column(db.Text)  # Longer bio/description
    social_links = db.Column(db.Text)  # JSON string of social media links
    
    # Referral system
    referral_code = db.Column(db.String(50), unique=True, index=True)  # User's unique referral code

    # Other fields
    role = db.Column(db.String(20), default='user')  # admin, editor, user
    theme = db.Column(db.String(10), default='dark')  # light, dark, auto
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    last_login = db.Column(db.DateTime, nullable=True)

class UserFollow(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    draft_name = db.Column(db.String(100), nullable=False)
    followed_at = db.Column(db.DateTime, default=datetime.utcnow)
    notification_level = db.Column(db.String(20), default='all')  # all, significant, major, comments, none

    # Relationship
    user = db.relationship('User', backref=db.backref('follows', lazy=True))

    __table_args__ = (db.UniqueConstraint('user_id', 'draft_name', name='unique_user_draft_follow'),)

class HypothesisAccount(db.Model):
    """Links Meta-Layer users to their Hypothesis accounts"""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, unique=True)
    hypothesis_username = db.Column(db.String(100), nullable=False, unique=True)
    hypothesis_userid = db.Column(db.String(100), nullable=False, unique=True)  # acct:username@authority
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationship
    user = db.relationship('User', backref=db.backref('hypothesis_account', uselist=False))

    NOTIFICATION_LEVELS = {
        'all': 'All changes and comments',
        'significant': 'Only significant changes (state changes, new revisions)',
        'major': 'Only major changes (IESG actions, RFC publication)',
        'comments': 'Only comments',
        'none': 'No notifications (just tracking)'
    }

class CoordinatorRequest(db.Model):
    """User-requested coordinator role; requires approval. Ties coordinator to user id."""
    id = db.Column(db.Integer, primary_key=True)
    group_acronym = db.Column(db.String(50), index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)  # known user
    username = db.Column(db.String(100), index=True)  # always set for lookup
    display_name = db.Column(db.String(200))  # for display in lists
    status = db.Column(db.String(20), default='pending')  # pending, approved, rejected
    requested_at = db.Column(db.DateTime, default=datetime.utcnow)
    reviewed_at = db.Column(db.DateTime, nullable=True)
    reviewed_by = db.Column(db.String(100), nullable=True)

class WorkgroupMemberRequest(db.Model):
    """Pending member join when workgroup has members_require_approval=True."""
    id = db.Column(db.Integer, primary_key=True)
    group_acronym = db.Column(db.String(50), index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    user_name = db.Column(db.String(100), index=True)
    status = db.Column(db.String(20), default='pending')  # pending, approved, rejected
    requested_at = db.Column(db.DateTime, default=datetime.utcnow)
    reviewed_at = db.Column(db.DateTime, nullable=True)
    reviewed_by = db.Column(db.String(100), nullable=True)

class WorkingGroupChair(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    group_acronym = db.Column(db.String(50), index=True)
    chair_name = db.Column(db.String(100))
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)  # set when created from CoordinatorRequest
    approved = db.Column(db.Boolean, default=False)
    set_at = db.Column(db.DateTime, default=datetime.utcnow)

class RoleImage(db.Model):
    """Visual representation proposed for a role"""
    __tablename__ = 'role_image'
    
    id = db.Column(db.String(50), primary_key=True)  # rimg_...
    project_id = db.Column(db.String(50), nullable=True, index=True)  # For future project scoping
    role_slug = db.Column(db.String(100), nullable=False, index=True)  # Role identifier
    
    # Source
    source_type = db.Column(db.String(20), nullable=False)  # 'upload', 'url', 'ordinal'
    image_url = db.Column(db.String(500), nullable=True)  # For upload or URL source
    file_path = db.Column(db.String(500), nullable=True)  # For uploaded files
    
    # Ordinal metadata (optional)
    chain = db.Column(db.String(50), nullable=True)  # 'bitcoin', etc.
    inscription_id = db.Column(db.String(255), nullable=True, index=True)
    content_type = db.Column(db.String(100), nullable=True)  # MIME type
    
    # Status and promotion
    is_primary = db.Column(db.Boolean, default=False, index=True)  # Primary role image
    is_hidden = db.Column(db.Boolean, default=False, index=True)  # Hidden by admin
    
    # Voting (aggregated)
    upvotes = db.Column(db.Integer, default=0)
    downvotes = db.Column(db.Integer, default=0)
    net_score = db.Column(db.Integer, default=0, index=True)  # upvotes - downvotes
    
    # Admin actions
    admin_note = db.Column(db.Text, nullable=True)
    promoted_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    promoted_at = db.Column(db.DateTime, nullable=True)
    
    # Audit
    submitted_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    submitted_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, onupdate=datetime.utcnow)
    
    # Relationships
    submitted_by = db.relationship('User', foreign_keys=[submitted_by_id], backref='submitted_role_images')
    promoted_by = db.relationship('User', foreign_keys=[promoted_by_id], backref='promoted_role_images')
    
    __table_args__ = (
        db.Index('idx_role_image_role_primary', 'role_slug', 'is_primary'),
        db.Index('idx_role_image_role_score', 'role_slug', 'net_score'),
    )
    
    def to_dict(self):
        """Convert to dictionary for API responses"""
        return {
            'id': self.id,
            'project_id': self.project_id,
            'role_slug': self.role_slug,
            'source_type': self.source_type,
            'image_url': self.image_url,
            'file_path': self.file_path,
            'chain': self.chain,
            'inscription_id': self.inscription_id,
            'content_type': self.content_type,
            'is_primary': self.is_primary,
            'is_hidden': self.is_hidden,
            'upvotes': self.upvotes,
            'downvotes': self.downvotes,
            'net_score': self.net_score,
            'admin_note': self.admin_note,
            'submitted_by_id': self.submitted_by_id,
            'submitted_by_name': self.submitted_by.displayName or self.submitted_by.username if self.submitted_by else 'Unknown',
            'submitted_at': self.submitted_at.isoformat() if self.submitted_at else None,
            'promoted_by_id': self.promoted_by_id,
            'promoted_at': self.promoted_at.isoformat() if self.promoted_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class RoleImageVote(db.Model):
    """User vote on a role image proposal"""
    __tablename__ = 'role_image_vote'
    
    id = db.Column(db.Integer, primary_key=True)
    image_id = db.Column(db.String(50), db.ForeignKey('role_image.id'), nullable=False, index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    
    # Vote value: 1 (upvote) or -1 (downvote)
    value = db.Column(db.Integer, nullable=False)  # 1 or -1
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, onupdate=datetime.utcnow)
    
    # Relationships
    image = db.relationship('RoleImage', backref=db.backref('votes', lazy=True))
    user = db.relationship('User', backref=db.backref('role_image_votes', lazy=True))
    
    __table_args__ = (
        db.UniqueConstraint('image_id', 'user_id', name='unique_user_image_vote'),
        db.Index('idx_vote_image', 'image_id'),
        db.Index('idx_vote_user', 'user_id'),
    )

# ============================================================================
# Projects, Workgroups, and Guilds Models
# ============================================================================

class Project(db.Model):
    """Primary organizing entity for submissions, documents, and workgroups"""
    __tablename__ = 'project'
    
    id = db.Column(db.String(50), primary_key=True)  # proj_...
    public_id = db.Column(db.String(36), unique=True, nullable=False, default=lambda: str(uuid4()))
    name = db.Column(db.String(255), unique=True, nullable=False, index=True)
    slug = db.Column(db.String(255), unique=True, nullable=False, index=True)
    
    # Initiator
    initiator_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    # Status (descriptive, not evaluative)
    status = db.Column(db.String(20), default='proposed', index=True)
    # proposed, active, stabilizing, maintaining, dormant, concluded, archived
    status_reason = db.Column(db.Text, nullable=True)
    
    # Image
    image_url = db.Column(db.String(500), nullable=True)
    
    # Admin approval
    approval_status = db.Column(db.String(20), default='pending', index=True)
    # pending, approved, rejected
    approved_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    approved_at = db.Column(db.DateTime, nullable=True)
    
    # Mission and description
    mission = db.Column(db.Text, nullable=True)
    description = db.Column(db.Text, nullable=True)
    
    # Activity tracking
    last_activity = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Succession
    superseded_by_id = db.Column(db.String(50), db.ForeignKey('project.id'), nullable=True)
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, onupdate=datetime.utcnow)
    
    # Relationships
    initiator = db.relationship('User', foreign_keys=[initiator_id], backref='initiated_projects')
    approved_by = db.relationship('User', foreign_keys=[approved_by_id], backref='approved_projects')
    superseded_by = db.relationship('Project', remote_side=[id], backref='supersedes')
    
    def to_dict(self):
        return {
            'id': self.id,
            'public_id': self.public_id,
            'name': self.name,
            'slug': self.slug,
            'initiator_id': self.initiator_id,
            'initiator_name': self.initiator.displayName or self.initiator.username if self.initiator else None,
            'status': self.status,
            'status_reason': self.status_reason,
            'approval_status': self.approval_status,
            'mission': self.mission,
            'description': self.description,
            'image_url': self.image_url,
            'last_activity': self.last_activity.isoformat() if self.last_activity else None,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None,
        }


class ProjectMember(db.Model):
    """Track project membership and referrals"""
    __tablename__ = 'project_member'
    
    id = db.Column(db.Integer, primary_key=True)
    project_id = db.Column(db.String(50), db.ForeignKey('project.id'), nullable=False, index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    
    # Referral tracking
    referred_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True, index=True)
    referral_code = db.Column(db.String(50), nullable=True)  # The referral code used to join
    
    # Role in project (optional)
    role = db.Column(db.String(100), nullable=True)  # contributor, maintainer, etc.
    
    # Status
    status = db.Column(db.String(20), default='active')  # active, inactive, left
    
    # Timestamps
    joined_at = db.Column(db.DateTime, default=datetime.utcnow, index=True)
    left_at = db.Column(db.DateTime, nullable=True)
    
    # Relationships
    project = db.relationship('Project', backref='project_members', foreign_keys=[project_id])
    user = db.relationship('User', backref='project_memberships', foreign_keys=[user_id])
    referred_by = db.relationship('User', backref='referrals_made', foreign_keys=[referred_by_id])
    
    # Unique constraint: one membership per user per project
    __table_args__ = (
        db.UniqueConstraint('project_id', 'user_id', name='unique_project_member'),
        db.Index('idx_project_member_status', 'status'),
    )
    
    def to_dict(self):
        return {
            'id': self.id,
            'project_id': self.project_id,
            'user_id': self.user_id,
            'referred_by_id': self.referred_by_id,
            'role': self.role,
            'status': self.status,
            'joined_at': self.joined_at.isoformat() if self.joined_at else None,
            'left_at': self.left_at.isoformat() if self.left_at else None,
        }


class ProjectAdmin(db.Model):
    """Assigned project admins (in addition to initiator/owner). Owner cannot be removed."""
    __tablename__ = 'project_admin'
    
    id = db.Column(db.Integer, primary_key=True)
    project_id = db.Column(db.String(50), db.ForeignKey('project.id'), nullable=False, index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    added_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    __table_args__ = (
        db.UniqueConstraint('project_id', 'user_id', name='uq_project_admin_project_user'),
    )
    
    project = db.relationship('Project', backref=db.backref('project_admins', lazy='dynamic'))
    user = db.relationship('User', backref=db.backref('project_admin_of', lazy='dynamic'))


class Waitlist(db.Model):
    """Project waitlist: name, description, public/private, referrals, active, dates, max, milestones."""
    __tablename__ = 'waitlist'
    
    id = db.Column(db.Integer, primary_key=True)
    project_id = db.Column(db.String(50), db.ForeignKey('project.id'), nullable=False, index=True)
    name = db.Column(db.String(255), nullable=False)
    description = db.Column(db.Text, nullable=True)
    
    public = db.Column(db.Boolean, default=True)  # False = only project members or link-holders see it
    referrals = db.Column(db.Boolean, default=False)  # If True, joiners get referral link; referrer gets credit
    active = db.Column(db.Boolean, default=True)  # If False, tab not shown
    start_date = db.Column(db.DateTime, nullable=False)  # Join disabled until start
    closing_date = db.Column(db.DateTime, nullable=True)
    max_number = db.Column(db.Integer, nullable=True)  # "Full" when reached
    archived = db.Column(db.Boolean, default=False)  # Soft delete / archive
    
    # Image
    image_url = db.Column(db.String(500), nullable=True)
    
    milestones = db.Column(db.Boolean, default=False)
    show_milestones = db.Column(db.String(20), default='all')  # 'all', 'next', 'future'
    
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    project = db.relationship('Project', backref=db.backref('waitlists', lazy='dynamic'))
    
    def to_dict(self):
        count = WaitlistEntry.query.filter_by(waitlist_id=self.id, left_at=None).count()
        return {
            'id': self.id,
            'project_id': self.project_id,
            'name': self.name,
            'description': self.description,
            'public': self.public,
            'referrals': self.referrals,
            'active': self.active,
            'start_date': self.start_date.isoformat() if self.start_date else None,
            'closing_date': self.closing_date.isoformat() if self.closing_date else None,
            'max_number': self.max_number,
            'archived': self.archived,
            'image_url': self.image_url,
            'milestones': self.milestones,
            'show_milestones': self.show_milestones,
            'count': count,
            'full': self.max_number is not None and count >= self.max_number,
            'closed': self.closing_date is not None and datetime.utcnow() >= self.closing_date,
            'started': self.start_date is not None and datetime.utcnow() >= self.start_date,
            'created_at': self.created_at.isoformat() if self.created_at else None,
        }


class WaitlistEntry(db.Model):
    """One user on a waitlist; can leave (left_at set). Position = order of join."""
    __tablename__ = 'waitlist_entry'
    
    id = db.Column(db.Integer, primary_key=True)
    waitlist_id = db.Column(db.Integer, db.ForeignKey('waitlist.id'), nullable=False, index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    message = db.Column(db.Text, nullable=True)
    position = db.Column(db.Integer, nullable=False)  # 1-based queue order
    referred_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True, index=True)
    referral_code = db.Column(db.String(50), nullable=True)
    source = db.Column(db.String(255), nullable=True)  # Track signup source (e.g., 'embed:example.com', 'direct', 'referral')
    source_url = db.Column(db.String(500), nullable=True)  # Full URL where signup occurred
    joined_at = db.Column(db.DateTime, default=datetime.utcnow)
    left_at = db.Column(db.DateTime, nullable=True)  # If set, user left
    
    waitlist = db.relationship('Waitlist', backref=db.backref('entries', lazy='dynamic'))
    user = db.relationship('User', foreign_keys=[user_id], backref=db.backref('waitlist_entries', lazy='dynamic'))
    referred_by = db.relationship('User', foreign_keys=[referred_by_id])
    
    __table_args__ = (db.UniqueConstraint('waitlist_id', 'user_id', name='uq_waitlist_entry_user'),)


class WaitlistMilestone(db.Model):
    """Milestone: activates at threshold (number on waitlist). Order by threshold."""
    __tablename__ = 'waitlist_milestone'
    
    id = db.Column(db.Integer, primary_key=True)
    waitlist_id = db.Column(db.Integer, db.ForeignKey('waitlist.id'), nullable=False, index=True)
    title = db.Column(db.String(255), nullable=False)
    description = db.Column(db.Text, nullable=True)
    threshold = db.Column(db.Integer, nullable=False)  # Number on waitlist to activate (ordering = by this)
    action_type = db.Column(db.String(50), nullable=True)  # e.g. email, badge, webhook
    action_payload = db.Column(db.Text, nullable=True)  # JSON
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    waitlist = db.relationship('Waitlist', backref=db.backref('milestone_list', lazy='dynamic', order_by='WaitlistMilestone.threshold'))


def is_project_admin(project, user):
    """True if user is project owner (initiator), an assigned project admin, or site admin."""
    if not user:
        return False
    if user.get('role') == 'admin':
        return True
    if not project:
        return False
    initiator_id = project.initiator_id if hasattr(project, 'initiator_id') else project.get('initiator_id')
    if initiator_id == user['id']:
        return True
    if hasattr(project, 'id'):
        pid = project.id
    else:
        pid = project.get('id')
    return ProjectAdmin.query.filter_by(project_id=pid, user_id=user['id']).first() is not None


class Workgroup(db.Model):
    """Task-focused group within a project"""
    __tablename__ = 'working_group'
    
    id = db.Column(db.Integer, primary_key=True)
    acronym = db.Column(db.String(50), unique=True, index=True)  # Legacy field
    name = db.Column(db.String(255), nullable=False)
    slug = db.Column(db.String(255), index=True)
    description = db.Column(db.Text, nullable=True)
    type = db.Column(db.String(50), nullable=True)  # Legacy field
    state = db.Column(db.String(20), nullable=True)  # Legacy field
    
    # Project relationship (required)
    project_id = db.Column(db.String(50), db.ForeignKey('project.id'), nullable=False, index=True)
    
    # Image
    image_url = db.Column(db.String(500), nullable=True)
    
    # Coordinator (formerly "chair")
    coordinator_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    
    # Status
    status = db.Column(db.String(20), default='active', index=True)
    # active, inactive, completed, archived, concluded
    
    # Approval
    approval_status = db.Column(db.String(20), default='pending', index=True)
    # pending, approved, rejected
    approved_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    approved_at = db.Column(db.DateTime, nullable=True)
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, onupdate=datetime.utcnow)
    
    # Relationships
    project = db.relationship('Project', backref=db.backref('workgroups', lazy=True))
    coordinator = db.relationship('User', foreign_keys=[coordinator_id], backref='coordinated_workgroups')
    approved_by = db.relationship('User', foreign_keys=[approved_by_id], backref='approved_workgroups')
    
    def to_dict(self):
        return {
            'id': self.id,
            'acronym': self.acronym,  # Legacy field
            'name': self.name,
            'slug': self.slug or self.acronym,  # Use acronym as fallback for legacy groups
            'project_id': self.project_id,
            'project_name': self.project.name if self.project else None,
            'coordinator_id': self.coordinator_id,
            'coordinator_name': self.coordinator.displayName or self.coordinator.username if self.coordinator else None,
            'status': self.status,
            'approval_status': self.approval_status,
            'description': self.description,
            'image_url': self.image_url,
            'type': self.type,  # Legacy field
            'state': self.state,  # Legacy field
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class Guild(db.Model):
    """Cross-project collaboration group"""
    __tablename__ = 'guild'
    
    id = db.Column(db.String(50), primary_key=True)  # guild_...
    name = db.Column(db.String(255), unique=True, nullable=False, index=True)
    slug = db.Column(db.String(255), unique=True, nullable=False, index=True)
    
    # Initiator (automatically becomes admin)
    initiator_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    # Description
    description = db.Column(db.Text, nullable=True)
    
    # Image
    image_url = db.Column(db.String(500), nullable=True)
    
    # Status (guilds don't require approval - instant registration)
    status = db.Column(db.String(20), default='active', index=True)
    # active, archived
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, onupdate=datetime.utcnow)
    
    # Relationships
    initiator = db.relationship('User', backref='initiated_guilds')
    
    def to_dict(self):
        return {
            'id': self.id,
            'name': self.name,
            'slug': self.slug,
            'initiator_id': self.initiator_id,
            'initiator_name': self.initiator.displayName or self.initiator.username if self.initiator else None,
            'description': self.description,
            'image_url': self.image_url,
            'status': self.status,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class GuildMembership(db.Model):
    """Guild membership with roles"""
    __tablename__ = 'guild_membership'
    
    id = db.Column(db.Integer, primary_key=True)
    guild_id = db.Column(db.String(50), db.ForeignKey('guild.id'), nullable=False, index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    
    # Role: initiator, admin, member
    role = db.Column(db.String(20), default='member', nullable=False)
    
    # Timestamps
    joined_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships
    guild = db.relationship('Guild', backref=db.backref('memberships', lazy=True))
    user = db.relationship('User', backref=db.backref('guild_memberships', lazy=True))
    
    __table_args__ = (
        db.UniqueConstraint('guild_id', 'user_id', name='unique_guild_membership'),
        db.Index('idx_guild_membership_guild', 'guild_id'),
        db.Index('idx_guild_membership_user', 'user_id'),
    )

class GuildInvitation(db.Model):
    """Guild invitation system"""
    __tablename__ = 'guild_invitation'
    
    id = db.Column(db.String(50), primary_key=True)  # ginv_...
    guild_id = db.Column(db.String(50), db.ForeignKey('guild.id'), nullable=False, index=True)
    inviter_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    invitee_email = db.Column(db.String(255), nullable=False, index=True)
    invitee_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)  # If user exists
    
    # Status
    status = db.Column(db.String(20), default='pending', index=True)
    # pending, accepted, declined, expired
    
    # Token for email verification
    token = db.Column(db.String(100), unique=True, nullable=False, index=True)
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    expires_at = db.Column(db.DateTime, nullable=False)  # 7 days from creation
    responded_at = db.Column(db.DateTime, nullable=True)
    
    # Relationships
    guild = db.relationship('Guild', backref=db.backref('invitations', lazy=True))
    inviter = db.relationship('User', foreign_keys=[inviter_id], backref='sent_guild_invitations')
    invitee = db.relationship('User', foreign_keys=[invitee_id], backref='received_guild_invitations')
    
    __table_args__ = (
        db.Index('idx_guild_invitation_guild', 'guild_id'),
        db.Index('idx_guild_invitation_status', 'status'),
        db.Index('idx_guild_invitation_token', 'token'),
    )

# ============================================================================
# Roles, Claims, and Badges Models
# ============================================================================

class Cluster(db.Model):
    """Organizational grouping of roles within a project"""
    __tablename__ = 'cluster'
    
    id = db.Column(db.String(50), primary_key=True)  # clu_...
    project_id = db.Column(db.String(50), db.ForeignKey('project.id'), nullable=False, index=True)
    cluster_slug = db.Column(db.String(100), nullable=False)
    name = db.Column(db.String(255), nullable=False)
    description = db.Column(db.Text, nullable=True)
    order = db.Column(db.Integer, default=0)
    
    # Status
    status = db.Column(db.String(20), default='active')
    # active, archived
    
    # Audit
    created_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, onupdate=datetime.utcnow)
    
    # Relationships
    project = db.relationship('Project', backref=db.backref('clusters', lazy=True))
    created_by = db.relationship('User', backref='created_clusters')
    
    __table_args__ = (
        db.UniqueConstraint('project_id', 'cluster_slug', name='unique_cluster_slug_per_project'),
        db.Index('idx_cluster_project_status', 'project_id', 'status'),
        db.Index('idx_cluster_project_order', 'project_id', 'order'),
    )
    
    def to_dict(self):
        return {
            'id': self.id,
            'project_id': self.project_id,
            'cluster_slug': self.cluster_slug,
            'name': self.name,
            'description': self.description,
            'order': self.order,
            'status': self.status,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class Role(db.Model):
    """Defined unit of responsibility scoped to a project"""
    __tablename__ = 'role'
    
    id = db.Column(db.String(50), primary_key=True)  # rol_...
    project_id = db.Column(db.String(50), db.ForeignKey('project.id'), nullable=False, index=True)
    role_slug = db.Column(db.String(100), nullable=False)
    
    # Titles
    title_guild = db.Column(db.String(255), nullable=False)
    title_operational = db.Column(db.String(255), nullable=True)
    
    # Description and image
    description = db.Column(db.Text, nullable=False)
    image_url = db.Column(db.String(500), nullable=True)
    
    # Organization
    cluster_id = db.Column(db.String(50), db.ForeignKey('cluster.id'), nullable=True)
    order = db.Column(db.Integer, default=0)
    
    # Status
    status = db.Column(db.String(20), default='draft', index=True)
    # draft, approved, deprecated, archived
    
    # Visibility
    public_visible = db.Column(db.Boolean, default=False, index=True)
    
    # Configuration
    claim_requires_approval = db.Column(db.Boolean, default=False)
    badge_enabled = db.Column(db.Boolean, default=True)
    badge_requires_approval = db.Column(db.Boolean, default=True)
    
    # Audit
    created_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    approved_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    approved_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, onupdate=datetime.utcnow)
    
    # Relationships
    project = db.relationship('Project', backref=db.backref('roles', lazy=True))
    cluster = db.relationship('Cluster', backref=db.backref('roles', lazy=True))
    created_by = db.relationship('User', foreign_keys=[created_by_id], backref='created_roles')
    approved_by = db.relationship('User', foreign_keys=[approved_by_id], backref='approved_roles')
    
    __table_args__ = (
        db.UniqueConstraint('project_id', 'role_slug', name='unique_role_slug_per_project'),
        db.Index('idx_role_project_status', 'project_id', 'status'),
        db.Index('idx_role_project_visible', 'project_id', 'public_visible'),
        db.Index('idx_role_status_visible', 'status', 'public_visible'),
    )
    
    def to_dict(self):
        return {
            'id': self.id,
            'project_id': self.project_id,
            'role_slug': self.role_slug,
            'slug': self.role_slug,  # Alias for consistency
            'title_guild': self.title_guild,
            'title_operational': self.title_operational,
            'description': self.description,
            'image_url': self.image_url,
            'cluster_id': self.cluster_id,
            'order': self.order,
            'status': self.status,
            'public_visible': self.public_visible,
            'claim_requires_approval': self.claim_requires_approval,
            'badge_enabled': self.badge_enabled,
            'badge_requires_approval': self.badge_requires_approval,
            'created_at': self.created_at.isoformat() if self.created_at else None
        }

class Claim(db.Model):
    """User's declaration of stewarding a role"""
    __tablename__ = 'claim'
    
    id = db.Column(db.String(50), primary_key=True)  # clm_...
    project_id = db.Column(db.String(50), db.ForeignKey('project.id'), nullable=False, index=True)
    role_id = db.Column(db.String(50), db.ForeignKey('role.id'), nullable=False, index=True)
    claimant_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    
    # Intent and evidence
    intent = db.Column(db.Text, nullable=True)
    evidence_links = db.Column(db.JSON, default=list)
    
    # Status
    status = db.Column(db.String(20), default='active', index=True)
    # active, pending_approval, paused, expired, revoked
    
    # Approval (if required)
    approval_required = db.Column(db.Boolean, default=False)
    approved_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    approved_at = db.Column(db.DateTime, nullable=True)
    
    # Term (optional time-bounding)
    term_start = db.Column(db.Date, nullable=True)
    term_end = db.Column(db.Date, nullable=True)
    term_duration_days = db.Column(db.Integer, nullable=True)
    term_status = db.Column(db.String(20), nullable=True)
    # active, expired, paused, canceled
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, onupdate=datetime.utcnow)
    
    # Relationships
    project = db.relationship('Project', backref=db.backref('claims', lazy=True))
    role = db.relationship('Role', backref=db.backref('claims', lazy=True))
    claimant = db.relationship('User', foreign_keys=[claimant_id], backref='role_claims')
    approved_by = db.relationship('User', foreign_keys=[approved_by_id], backref='approved_claims')
    
    __table_args__ = (
        db.Index('idx_claim_project_status', 'project_id', 'status'),
        db.Index('idx_claim_role_status', 'role_id', 'status'),
        db.Index('idx_claim_claimant_status', 'claimant_id', 'status'),
        db.Index('idx_claim_created', 'created_at'),
    )
    
    def to_dict(self):
        return {
            'id': self.id,
            'project_id': self.project_id,
            'role_id': self.role_id,
            'claimant_id': self.claimant_id,
            'intent': self.intent,
            'evidence_links': self.evidence_links,
            'status': self.status,
            'approval_required': self.approval_required,
            'approved_at': self.approved_at.isoformat() if self.approved_at else None,
            'term_start': self.term_start.isoformat() if self.term_start else None,
            'term_end': self.term_end.isoformat() if self.term_end else None,
            'term_duration_days': self.term_duration_days,
            'term_status': self.term_status,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class Badge(db.Model):
    """Recognition artifact linked to a claim"""
    __tablename__ = 'badge'
    
    id = db.Column(db.String(50), primary_key=True)  # bdg_...
    public_id = db.Column(db.String(36), unique=True, nullable=False, default=lambda: str(uuid4()))
    project_id = db.Column(db.String(50), db.ForeignKey('project.id'), nullable=False, index=True)
    claim_id = db.Column(db.String(50), db.ForeignKey('claim.id'), nullable=False, index=True)
    role_id = db.Column(db.String(50), db.ForeignKey('role.id'), nullable=False)
    claimant_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    requested_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    # Badge type
    badge_type = db.Column(db.String(50), default='role_badge')
    # role_badge, founding_wave_badge, term_renewal_marker
    
    # Status
    status = db.Column(db.String(20), default='requested', index=True)
    # requested, needs_info, approved, issued, denied, canceled
    
    # Evidence
    evidence_links = db.Column(db.JSON, default=list)
    
    # Custody
    custody_mode = db.Column(db.String(20), default='user_wallet')
    # user_wallet, overweb_treasury
    btc_taproot_address = db.Column(db.String(255), nullable=True)
    
    # Approval
    approved_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    approved_at = db.Column(db.DateTime, nullable=True)
    approval_note = db.Column(db.Text, nullable=True)
    
    # Issuance (ordinal metadata)
    issuance_kind = db.Column(db.String(20), default='offchain')
    # offchain, ordinal
    inscription_id = db.Column(db.String(255), nullable=True)
    tx_ref = db.Column(db.String(255), nullable=True)
    chain = db.Column(db.String(50), default='bitcoin', nullable=True)
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, onupdate=datetime.utcnow)
    
    # Relationships
    project = db.relationship('Project', backref=db.backref('badges', lazy=True))
    claim = db.relationship('Claim', backref=db.backref('badges', lazy=True))
    role = db.relationship('Role', backref=db.backref('badges', lazy=True))
    claimant = db.relationship('User', foreign_keys=[claimant_id], backref='badges_received')
    requested_by = db.relationship('User', foreign_keys=[requested_by_id], backref='badges_requested')
    approved_by = db.relationship('User', foreign_keys=[approved_by_id], backref='badges_approved')
    
    __table_args__ = (
        db.Index('idx_badge_project_status', 'project_id', 'status'),
        db.Index('idx_badge_claim_status', 'claim_id', 'status'),
        db.Index('idx_badge_status', 'status'),
        db.Index('idx_badge_created', 'created_at'),
    )
    
    def to_dict(self):
        return {
            'id': self.id,
            'project_id': self.project_id,
            'claim_id': self.claim_id,
            'role_id': self.role_id,
            'claimant_id': self.claimant_id,
            'badge_type': self.badge_type,
            'status': self.status,
            'evidence_links': self.evidence_links,
            'custody_mode': self.custody_mode,
            'btc_taproot_address': self.btc_taproot_address,
            'approved_at': self.approved_at.isoformat() if self.approved_at else None,
            'approval_note': self.approval_note,
            'issuance_kind': self.issuance_kind,
            'inscription_id': self.inscription_id,
            'tx_ref': self.tx_ref,
            'chain': self.chain,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class StatusChange(db.Model):
    """Audit trail for status changes across all entities"""
    __tablename__ = 'status_change'
    
    id = db.Column(db.String(50), primary_key=True)  # sc_...
    
    # Polymorphic reference
    entity_type = db.Column(db.String(20), nullable=False)
    # role, claim, badge, cluster, project, workgroup, guild
    entity_id = db.Column(db.String(50), nullable=False, index=True)
    
    # Change details
    field_name = db.Column(db.String(50), nullable=False)
    from_value = db.Column(db.String(100), nullable=True)
    to_value = db.Column(db.String(100), nullable=False)
    note = db.Column(db.Text, nullable=True)
    
    # Audit
    changed_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    changed_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships
    changed_by = db.relationship('User', backref='status_changes')
    
    __table_args__ = (
        db.Index('idx_status_change_entity', 'entity_type', 'entity_id'),
        db.Index('idx_status_change_changed_at', 'changed_at'),
    )

# ================================================================
# VOTING MODELS
# ================================================================

class Vote(db.Model):
    """A vote/ballot on a draft submission within a project context."""
    __tablename__ = 'vote'
    
    id = db.Column(db.Integer, primary_key=True)
    public_id = db.Column(db.String(36), unique=True, nullable=False, default=lambda: str(uuid4()))
    
    project_id = db.Column(db.String(50), db.ForeignKey('project.id'), nullable=False, index=True)
    submission_id = db.Column(db.String(8), db.ForeignKey('submission.id'), nullable=False, index=True)
    created_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    title = db.Column(db.String(255), nullable=False)
    description = db.Column(db.Text, nullable=True)
    
    start_at = db.Column(db.DateTime, nullable=False)
    end_at = db.Column(db.DateTime, nullable=False)
    
    quorum_count = db.Column(db.Integer, nullable=False)
    win_threshold = db.Column(db.Float, nullable=False, default=0.5)
    
    status = db.Column(db.String(20), nullable=False, default='scheduled', index=True)
    result = db.Column(db.String(20), nullable=True)
    result_summary = db.Column(db.Text, nullable=True)
    
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    closed_at = db.Column(db.DateTime, nullable=True)
    
    # Relationships
    project = db.relationship('Project', backref=db.backref('votes', lazy=True))
    submission = db.relationship('Submission', backref=db.backref('votes', lazy=True))
    created_by = db.relationship('User', backref='created_votes')

class VoteEligibilitySnapshot(db.Model):
    """Snapshot of eligible voters at vote activation time."""
    __tablename__ = 'vote_eligibility_snapshot'
    
    id = db.Column(db.Integer, primary_key=True)
    vote_id = db.Column(db.Integer, db.ForeignKey('vote.id'), nullable=False, index=True)
    person_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    is_eligible = db.Column(db.Boolean, default=True, nullable=False)
    reason = db.Column(db.String(255), nullable=True)
    captured_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    __table_args__ = (
        db.UniqueConstraint('vote_id', 'person_id', name='unique_vote_eligibility'),
    )
    
    vote = db.relationship('Vote', backref=db.backref('eligibility_snapshot', lazy=True))
    person = db.relationship('User', backref=db.backref('vote_eligibility', lazy=True))

class Ballot(db.Model):
    """A single person's ballot cast in a vote."""
    __tablename__ = 'ballot'
    
    id = db.Column(db.Integer, primary_key=True)
    vote_id = db.Column(db.Integer, db.ForeignKey('vote.id'), nullable=False, index=True)
    person_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    choice = db.Column(db.String(10), nullable=False)
    cast_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    __table_args__ = (
        db.UniqueConstraint('vote_id', 'person_id', name='unique_ballot'),
    )
    
    vote = db.relationship('Vote', backref=db.backref('ballots', lazy=True))
    person = db.relationship('User', backref=db.backref('ballots_cast', lazy=True))

# ================================================================
# VOTING LOGIC
# ================================================================

def activate_vote(vote):
    """Activate a scheduled vote: set status, snapshot eligibility."""
    if vote.status != 'scheduled':
        return False, f"Cannot activate vote in status '{vote.status}'"
    
    # Snapshot eligible voters = active ProjectMembers for vote.project_id
    members = ProjectMember.query.filter_by(
        project_id=vote.project_id,
        status='active'
    ).all()
    
    for member in members:
        snapshot = VoteEligibilitySnapshot(
            vote_id=vote.id,
            person_id=member.user_id,
            is_eligible=True,
            reason='active project member at vote activation'
        )
        db.session.add(snapshot)
    
    vote.status = 'active'
    db.session.commit()
    
    print(f"[VOTE] Activated vote {vote.id} ({vote.title}) — {len(members)} eligible voters")
    return True, f"Activated with {len(members)} eligible voters"


def close_vote(vote):
    """Close an active vote: tally ballots, determine result."""
    if vote.status != 'active':
        return False, f"Cannot close vote in status '{vote.status}'"
    
    ballots = Ballot.query.filter_by(vote_id=vote.id).all()
    
    yes_count = sum(1 for b in ballots if b.choice == 'yes')
    no_count = sum(1 for b in ballots if b.choice == 'no')
    abstain_count = sum(1 for b in ballots if b.choice == 'abstain')
    votes_cast = yes_count + no_count + abstain_count
    
    eligible_count = VoteEligibilitySnapshot.query.filter_by(
        vote_id=vote.id, is_eligible=True
    ).count()
    
    quorum_met = votes_cast >= vote.quorum_count
    
    if yes_count + no_count > 0:
        yes_ratio = yes_count / (yes_count + no_count)
    else:
        yes_ratio = 0.0
    
    if not quorum_met:
        vote.result = 'no_quorum'
    elif yes_ratio >= vote.win_threshold:
        vote.result = 'passed'
    else:
        vote.result = 'failed'
    
    vote.status = 'closed'
    vote.closed_at = datetime.utcnow()
    vote.result_summary = json.dumps({
        'eligible': eligible_count,
        'votes_cast': votes_cast,
        'yes': yes_count,
        'no': no_count,
        'abstain': abstain_count,
        'quorum_required': vote.quorum_count,
        'quorum_met': quorum_met,
        'yes_ratio': round(yes_ratio, 4),
        'win_threshold': vote.win_threshold,
        'result': vote.result
    })
    
    db.session.commit()
    
    print(f"[VOTE] Closed vote {vote.id} ({vote.title}) — result: {vote.result}")
    return True, vote.result

# Users are now stored in database - this dict is kept for backward compatibility during migration

# Store document history in memory
DOCUMENT_HISTORY = {}

# Store comments in memory
COMMENTS = {}

# Store comment likes in memory
COMMENT_LIKES = {}

# Store comment replies in memory
COMMENT_REPLIES = {}

# Store workgroup coordinators in memory
WORKING_GROUP_CHAIRS = {}

# Configuration for file uploads
UPLOAD_FOLDER = '/home/ubuntu/data-tracker/uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'xml', 'doc', 'docx'}
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'svg'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
ROLE_IMAGE_UPLOAD_FOLDER = '/home/ubuntu/data-tracker/uploads/role_images'
ENTITY_IMAGE_UPLOAD_FOLDER = '/home/ubuntu/data-tracker/uploads/entity_images'
PROFILE_IMAGE_UPLOAD_FOLDER = '/home/ubuntu/data-tracker/uploads/profile_images'
os.makedirs(ROLE_IMAGE_UPLOAD_FOLDER, exist_ok=True)
os.makedirs(ENTITY_IMAGE_UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROFILE_IMAGE_UPLOAD_FOLDER, exist_ok=True)

# Image dimension constraints
IMAGE_MAX_WIDTH = 600
IMAGE_MAX_HEIGHT = 600
IMAGE_MAX_SIZE_MB = 5

# Entity images (projects, workgroups, guilds, waitlists): same 600×600 limit as role images
ENTITY_IMAGE_MAX_DIMENSION = 600
ENTITY_IMAGE_UPLOAD_FOLDER = '/home/ubuntu/data-tracker/uploads/entity_images'
MAX_IMAGE_FILE_SIZE = 5 * 1024 * 1024  # 5MB for 600×600 images
os.makedirs(ENTITY_IMAGE_UPLOAD_FOLDER, exist_ok=True)

# Comment edit/delete time limit (in minutes)
EDIT_DELETE_TIME_MINUTES = 15

# Create upload directory if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_image_file(filename):
    """Check if image file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS

def upload_image_600x600(file_storage, upload_folder, url_prefix, filename_prefix='img'):
    """
    Validate and save an uploaded image with max dimensions 600×600.
    Returns (image_url, None) on success, or (None, error_message) on failure.
    file_storage: FileStorage from request.files['file']
    upload_folder: absolute path to directory to save the file
    url_prefix: URL path prefix for the saved file (e.g. '/uploads/role_images' or '/uploads/entity_images')
    filename_prefix: prefix for the saved filename (e.g. 'rimg' or 'entity')
    """
    if not file_storage or not file_storage.filename:
        return None, 'No file provided'
    filename = file_storage.filename
    if not allowed_image_file(filename):
        return None, 'Invalid file type. Allowed: PNG, JPG, GIF, WebP, SVG'
    ext = filename.rsplit('.', 1)[1].lower()
    file_storage.seek(0, os.SEEK_END)
    file_size = file_storage.tell()
    file_storage.seek(0)
    if file_size > MAX_IMAGE_FILE_SIZE:
        return None, f'File too large. Maximum size is {MAX_IMAGE_FILE_SIZE // (1024*1024)}MB.'
    # Validate dimensions for raster formats (PIL); SVG is allowed without dimension check
    if ext != 'svg':
        try:
            from PIL import Image
            img = Image.open(file_storage)
            img.load()
            w, h = img.size
            if w > ENTITY_IMAGE_MAX_DIMENSION or h > ENTITY_IMAGE_MAX_DIMENSION:
                return None, f'Image dimensions must be at most {ENTITY_IMAGE_MAX_DIMENSION}×{ENTITY_IMAGE_MAX_DIMENSION} pixels (got {w}×{h}).'
            file_storage.seek(0)
        except Exception as e:
            file_storage.seek(0)
            return None, f'Invalid image or unsupported format: {e}'
    safe_name = f"{filename_prefix}_{os.urandom(8).hex()}.{ext}"
    file_path = os.path.join(upload_folder, safe_name)
    try:
        file_storage.save(file_path)
    except Exception as e:
        return None, f'Failed to save file: {e}'
    return f"{url_prefix}/{safe_name}", None

def generate_draft_name(title, authors):
    """Generate a draft name from title and authors"""
    # Extract first author's last name
    first_author = authors[0] if authors else "unknown"
    author_last = first_author.split()[-1].lower() if first_author else "unknown"
    
    # Create a slug from the title
    title_slug = re.sub(r'[^a-zA-Z0-9\s-]', '', title.lower())
    title_slug = re.sub(r'\s+', '-', title_slug.strip())
    title_slug = title_slug[:30]  # Limit length
    
    return f"draft-{author_last}-{title_slug}"

def generate_role_image_id():
    """Generate unique role image ID with rimg_ prefix"""
    import random
    import string
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=12))
    return f"rimg_{suffix}"

def update_image_vote_counts(image_id):
    """Recalculate vote counts for a role image"""
    image = RoleImage.query.get(image_id)
    if not image:
        return False
    
    votes = RoleImageVote.query.filter_by(image_id=image_id).all()
    upvotes = sum(1 for v in votes if v.value == 1)
    downvotes = sum(1 for v in votes if v.value == -1)
    
    image.upvotes = upvotes
    image.downvotes = downvotes
    image.net_score = upvotes - downvotes
    db.session.commit()
    return True

def generate_project_id():
    """Generate unique project ID with proj_ prefix"""
    import random
    import string
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=12))
    return f"proj_{suffix}"

def generate_workgroup_id():
    """Generate unique workgroup ID with wg_ prefix"""
    import random
    import string
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=12))
    return f"wg_{suffix}"

def generate_guild_id():
    """Generate unique guild ID with guild_ prefix"""
    import random
    import string
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=12))
    return f"guild_{suffix}"

def generate_cluster_id():
    """Generate unique cluster ID with clu_ prefix"""
    import random
    import string
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=12))
    return f"clu_{suffix}"

def generate_role_id():
    """Generate unique role ID with rol_ prefix"""
    import random
    import string
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=12))
    return f"rol_{suffix}"

def generate_claim_id():
    """Generate unique claim ID with clm_ prefix"""
    import random
    import string
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=12))
    return f"clm_{suffix}"

def generate_badge_id():
    """Generate unique badge ID with bdg_ prefix"""
    import random
    import string
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=12))
    return f"bdg_{suffix}"

def generate_status_change_id():
    """Generate unique status change ID with sc_ prefix"""
    import random
    import string
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=12))
    return f"sc_{suffix}"

def generate_guild_invitation_id():
    """Generate unique guild invitation ID with ginv_ prefix"""
    import random
    import string
    suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=12))
    return f"ginv_{suffix}"

def generate_invitation_token():
    """Generate secure token for guild invitations"""
    import secrets
    return secrets.token_urlsafe(32)

def create_slug(text):
    """Create URL-safe slug from text"""
    import re
    slug = text.lower()
    slug = re.sub(r'[^a-z0-9\s-]', '', slug)
    slug = re.sub(r'\s+', '-', slug.strip())
    return slug[:100]

def require_auth(f):
    """Decorator to require authentication"""
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash('Please log in to access this page.', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    decorated_function.__name__ = f.__name__
    return decorated_function

def require_role(required_role):
    """Decorator to require a specific role (admin or editor can access admin features)"""
    def decorator(f):
        def decorated_function(*args, **kwargs):
            if 'user' not in session:
                flash('Please log in to access this page.', 'error')
                return redirect(url_for('login'))
            current_user = get_current_user()
            if not current_user:
                return "Access denied: Not logged in", 403
            
            user_role = current_user.get('role', 'user')
            
            # Admin and editor can access admin pages
            if required_role == 'admin':
                if user_role not in ['admin', 'editor']:
                    return "Access denied: Admin or Editor role required", 403
            # For other roles, exact match required
            elif user_role != required_role:
                return f"Access denied: {required_role} role required", 403
            
            return f(*args, **kwargs)
        decorated_function.__name__ = f.__name__
        return decorated_function
    return decorator

def shorten_inscription_id(inscription_id, chars_each_side=8):
    """
    Shorten an inscription ID to show first N chars...last N chars
    Example: 8e24de515cc0dc305188f3c4a0e563466723bf9cf8d4576184bf3d13e287615bi0
    becomes: 8e24de51....7615bi0 (with chars_each_side=8)
    
    Always includes the 'i0' at the end as it's part of the ordinal identifier.
    """
    if not inscription_id:
        return ''
    
    # Ensure we have enough characters
    if len(inscription_id) <= (chars_each_side * 2 + 4):
        return inscription_id
    
    # Extract parts
    start = inscription_id[:chars_each_side]
    end = inscription_id[-chars_each_side:]
    
    return f'{start}....{end}'

def generate_referral_code(username):
    """Generate a unique referral code for a user"""
    import hashlib
    import time
    # Create a unique code based on username and timestamp
    raw = f"{username}-{time.time()}"
    hash_obj = hashlib.md5(raw.encode())
    return hash_obj.hexdigest()[:8].upper()

def get_or_create_referral_code(user):
    """Get user's referral code or create one if it doesn't exist"""
    if not user.referral_code:
        user.referral_code = generate_referral_code(user.username)
        db.session.commit()
    return user.referral_code

def get_current_user():
    """Get current logged in user"""
    if 'user' in session:
        user = User.query.filter_by(username=session['user']).first()
        if user:
            # Use displayName or oauthName as fallback if name is not set
            user_name = user.name or user.displayName or user.oauthName or user.username
            return {
                'id': user.id,
                'username': user.username,
                'name': user_name,
                'email': user.email,
                'role': user.role,
                'theme': user.theme,
                # Web3Auth fields
                'displayName': user.displayName,
                'oauthName': user.oauthName,
                'profileImage': user.profileImage,
                'typeOfLogin': user.typeOfLogin,
                'evmAddress': user.evmAddress,
                'solanaAddress': user.solanaAddress
            }
    return None

def render_page(title, content, theme=None, user_menu=None):
    """Helper to render a page with BASE_TEMPLATE including build number"""
    if theme is None:
        theme = session.get('theme', get_current_user().get('theme', 'dark') if get_current_user() else 'dark')
    if user_menu is None:
        user_menu = generate_user_menu()
    return BASE_TEMPLATE.format(
        title=title,
        theme=theme,
        user_menu=user_menu,
        content=content,
        build_number=BUILD_NUMBER,
        hypothesis_config=""
    )

def process_ordinal_markdown(markdown_text):
    """
    Shared function to process ordinal markdown content into HTML.
    Used by both the API endpoint and draft_detail page for consistency.
    Returns the processed HTML string.
    """
    if not MARKDOWN_SUPPORT:
        import html
        return html.escape(markdown_text).replace('\n', '<br>')
    
    import re
    import markdown2
    import bleach
    
    # Pre-process markdown to handle figure tags with images
    def replace_figure_image(match):
        alt_text = match.group(1) if match.group(1) else ''
        image_url = match.group(2)
        caption = match.group(3) if len(match.groups()) >= 3 and match.group(3) else ''
        
        html = '<figure class="figure">\n'
        html += f'  <img src="{image_url}" alt="{alt_text}" class="img-fluid figure-img">\n'
        if caption:
            html += f'  <figcaption class="figure-caption"><small>{caption}</small></figcaption>\n'
        html += '</figure>'
        return html
    
    # Pattern to match: <figure>\n![alt](url)\n<figcaption>caption</figcaption>\n</figure>
    # Updated to handle nested HTML tags in figcaption using non-greedy match
    markdown_text = re.sub(
        r'<figure[^>]*>\s*!\[([^\]]*)\]\(([^\)]+)\)\s*(?:<figcaption[^>]*>(.*?)</figcaption>)?\s*</figure>',
        replace_figure_image,
        markdown_text,
        flags=re.MULTILINE | re.DOTALL
    )
    
    # Convert markdown to HTML using markdown2 (without break-on-newline to avoid extra line breaks)
    html_content = markdown2.markdown(
        markdown_text,
        extras=['fenced-code-blocks', 'tables']
    )
    
    # Fix image URLs BEFORE sanitization
    # 1. Fix relative /content/ URLs
    html_content = re.sub(
        r'src="(/content/[^"]+)"',
        r'src="https://ordinals.com\1"',
        html_content
    )
    
    # 2. Fix bare inscription IDs (64-char hex + 'i' + number)
    html_content = re.sub(
        r'src="(?:[^"]*/)??([a-f0-9]{64}i\d+)"',
        r'src="https://ordinals.com/content/\1"',
        html_content
    )
    
    # 3. Add img-fluid class to all img tags that don't already have it
    html_content = re.sub(
        r'<img(?![^>]*class=)([^>]*)>',
        r'<img class="img-fluid"\1>',
        html_content
    )
    
    # Sanitize HTML to prevent XSS
    allowed_tags = [
        'p', 'br', 'strong', 'em', 'u', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        'ul', 'ol', 'li', 'blockquote', 'code', 'pre', 'a', 'img',
        'table', 'thead', 'tbody', 'tr', 'th', 'td',
        'figure', 'figcaption', 'small', 'hr', 'div', 'span'
    ]
    allowed_attrs = {
        'a': ['href', 'title', 'target'],
        'img': ['src', 'alt', 'title', 'class'],
        'code': ['class'],
        'figure': ['class'],
        'figcaption': ['class']
    }
    
    html_content = bleach.clean(
        html_content,
        tags=allowed_tags,
        attributes=allowed_attrs,
        strip=True
    )
    
    return html_content

def generate_user_menu():
    """Generate user menu HTML for navbar"""
    current_user = get_current_user()
    if current_user:
        user_role = current_user.get('role', 'user')
        is_admin = user_role in ['admin', 'editor'] or current_user['name'] in ['admin', 'Admin User']
        admin_link = '<li><a class="dropdown-item" href="/admin/">Admin Dashboard</a></li>' if is_admin else ''
        # Display name priority: displayName > oauthName > name > username
        display_name = (current_user.get('displayName') or
                       current_user.get('oauthName') or
                       current_user.get('name') or
                       current_user['username'])

        return f"""
        <div class="nav-item dropdown">
            <a class="nav-link dropdown-toggle" href="#" role="button" data-bs-toggle="dropdown" aria-expanded="false">
                {display_name}
            </a>
            <ul class="dropdown-menu">
                <li><a class="dropdown-item" href="/profile/">Profile</a></li>
                <li><a class="dropdown-item" href="/my-projects/">My Layers</a></li>
                <li><a class="dropdown-item" href="/submit/status/">My Submissions</a></li>
                {admin_link}
                <li><hr class="dropdown-divider"></li>
                <li><a class="dropdown-item" href="/logout/">Logout</a></li>
            </ul>
        </div>
        """
    else:
        return """
        <div class="nav-item">
            <a class="nav-link" href="#" onclick="event.preventDefault(); loginWithWeb3Auth(); return false;">Sign In</a>
        </div>
        """

def add_to_document_history(draft_name, action, user, details=""):
    """Add an entry to document history"""
    if draft_name not in DOCUMENT_HISTORY:
        DOCUMENT_HISTORY[draft_name] = []
    
    entry = {
        'action': action,
        'user': user,
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'details': details
    }
    DOCUMENT_HISTORY[draft_name].insert(0, entry)  # Add to beginning (most recent first)

def toggle_comment_like(draft_name, comment_id, user):
    """Toggle like on a comment"""
    like_key = f"{draft_name}:{comment_id}"
    if like_key not in COMMENT_LIKES:
        COMMENT_LIKES[like_key] = set()
    
    if user in COMMENT_LIKES[like_key]:
        COMMENT_LIKES[like_key].remove(user)
        return False  # Unliked
    else:
        COMMENT_LIKES[like_key].add(user)
        return True  # Liked

def get_comment_likes(draft_name, comment_id):
    """Get like count for a comment"""
    like_key = f"{draft_name}:{comment_id}"
    return len(COMMENT_LIKES.get(like_key, set()))

def is_comment_liked(draft_name, comment_id, user):
    """Check if user has liked a comment"""
    like_key = f"{draft_name}:{comment_id}"
    return user in COMMENT_LIKES.get(like_key, set())

def is_user_following_draft(draft_name, user):
    """Check if a user is following a specific draft"""
    if not user:
        return False
    return UserFollow.query.filter_by(user_id=user['id'], draft_name=draft_name).first() is not None

def get_user_follow(draft_name, user):
    """Get the UserFollow object for a user and draft"""
    if not user:
        return None
    return UserFollow.query.filter_by(user_id=user['id'], draft_name=draft_name).first()

def get_notification_controls(draft_name, user):
    """Generate HTML for notification level controls"""
    if not user:
        return ''

    follow = get_user_follow(draft_name, user)
    if not follow:
        return ''

    current_level = follow.notification_level
    options = []
    for level, description in UserFollow.NOTIFICATION_LEVELS.items():
        selected = 'selected' if level == current_level else ''
        options.append(f'<option value="{level}" {selected}>{description}</option>')

    return f'''
    <form method="post" action="/doc/draft/{draft_name}/update-notification/" class="mt-2">
        <label class="form-label small">Notification Level:</label>
        <select name="notification_level" class="form-select form-select-sm mb-1">
            {''.join(options)}
        </select>
        <button type="submit" class="btn btn-outline-secondary btn-sm w-100">Update Notifications</button>
    </form>
    '''

def should_notify_user(follow, event_type):
    """
    Determine if a user should be notified based on their notification level and event type.

    Event types:
    - 'comment': New comment added
    - 'revision': New revision uploaded
    - 'state_change': Document state changed
    - 'major_change': Major events (IESG actions, RFC publication)
    """
    level = follow.notification_level

    if level == 'none':
        return False
    elif level == 'comments':
        return event_type == 'comment'
    elif level == 'major':
        return event_type in ['major_change']
    elif level == 'significant':
        return event_type in ['revision', 'state_change', 'major_change']
    elif level == 'all':
        return True

    return False  # Default to no notification for unknown levels

def get_users_to_notify(draft_name, event_type):
    """
    Get all users who should be notified for a specific event on a document.
    Returns list of (user, follow) tuples.
    """
    follows = UserFollow.query.filter_by(draft_name=draft_name).all()
    users_to_notify = []

    for follow in follows:
        if should_notify_user(follow, event_type):
            user = User.query.get(follow.user_id)
            if user:
                users_to_notify.append((user, follow))

    return users_to_notify

def add_comment_reply(draft_name, parent_comment_id, reply_text, user):
    """Add a reply to a comment"""
    # Create reply in database
    reply = Comment(
        draft_name=draft_name,
        text=reply_text,
        author=user['name'],
        parent_id=int(parent_comment_id)
    )
    db.session.add(reply)
    db.session.commit()
    return reply

def build_comment_tree(draft_name):
    """Build a tree structure of comments with nested replies"""
    # Get all comments for this draft (including deleted ones, but mark them)
    all_comments = Comment.query.filter_by(draft_name=draft_name).order_by(Comment.timestamp).all()

    # Create a dictionary for quick lookup
    comment_dict = {}
    for comment in all_comments:
        comment_dict[comment.id] = {
            'id': str(comment.id),
            'author': comment.author,
            'date': comment.timestamp.strftime('%Y-%m-%d %H:%M'),
            'comment': comment.text if not comment.is_deleted else '[Deleted]',
            'avatar': ''.join([word[0].upper() for word in comment.author.split()[:2]]),
            'replies': [],
            'timestamp': comment.timestamp,
            'edited_at': comment.edited_at,
            'is_deleted': comment.is_deleted,
            'original_text': comment.original_text
        }

    # Build the tree
    top_level_comments = []
    for comment in all_comments:
        if comment.parent_id is None:
            # Top-level comment
            top_level_comments.append(comment_dict[comment.id])
        else:
            # Reply - add to parent's replies
            if comment.parent_id in comment_dict:
                comment_dict[comment.parent_id]['replies'].append(comment_dict[comment.id])

    return top_level_comments

def can_edit_delete_comment(comment, current_user):
    """Check if current user can edit/delete this comment"""
    if not current_user:
        return False
    if comment['author'] != current_user['name']:
        return False
    if comment.get('is_deleted', False):
        return False
    
    # Check time limit
    comment_time = comment.get('timestamp')
    if comment_time:
        time_diff = datetime.utcnow() - comment_time
        time_limit = timedelta(minutes=EDIT_DELETE_TIME_MINUTES)
        return time_diff <= time_limit
    return False

def render_comment_tree(comments, draft_name, level=0):
    """Recursively render comments and their nested replies"""
    if not comments:
        return ""
    
    indent_class = f"ms-{level * 4}" if level > 0 else ""
    html = f'<div class="{indent_class} mt-2">' if level > 0 else '<div class="mt-2">'

    for comment in comments:
        comment_id = comment.get('id', 'unknown')
        like_count = get_comment_likes(draft_name, comment_id)
        is_liked = is_comment_liked(draft_name, comment_id, get_current_user()['name']) if get_current_user() else False
        current_user = get_current_user()
        can_edit_delete = can_edit_delete_comment(comment, current_user)
        is_deleted = comment.get('is_deleted', False)
        edited_at = comment.get('edited_at')
        edited_text = f" (edited {edited_at.strftime('%Y-%m-%d %H:%M')})" if edited_at else ""

        # Like button styling
        like_btn_class = "btn-outline-danger" if is_liked else "btn-outline-secondary"
        like_icon = "❤️" if is_liked else "🤍"

        # Size styling based on nesting level
        avatar_size = max(30 - level * 5, 20)  # Decrease avatar size for nested replies
        font_size = max(14 - level * 2, 12)    # Decrease font size for nested replies
        card_class = "mb-2" if level > 0 else "mb-3"

        # Edit/Delete buttons HTML
        edit_delete_buttons = ""
        if can_edit_delete:
            edit_click = f"editComment('{comment_id}')"
            delete_click = f"deleteComment('{comment_id}')"
            edit_delete_buttons = f"""
                    <button class="btn btn-sm btn-outline-warning" onclick="{edit_click}" style="font-size: {font_size - 2}px;">
                        Edit
                    </button>
                    <button class="btn btn-sm btn-outline-danger" onclick="{delete_click}" style="font-size: {font_size - 2}px;">
                        Delete
                    </button>
            """

        # Build onclick handlers outside f-string
        like_click = f"toggleLike('{comment_id}')"
        reply_click = f"toggleReply('{comment_id}')"
        like_button = f'<button class="btn btn-sm {like_btn_class}" onclick="{like_click}" style="font-size: {font_size - 2}px;">{like_icon} {like_count}</button>' if not is_deleted else ''
        reply_button = f'<button class="btn btn-sm btn-outline-primary" onclick="{reply_click}" style="font-size: {font_size - 2}px;">Reply</button>' if not is_deleted else ''
        deleted_badge = '<small class="text-muted ms-2" style="font-style: italic;">[Deleted]</small>' if is_deleted else ''
        deleted_style = 'opacity: 0.5; font-style: italic;' if is_deleted else ''

        html += f"""
        <div class="card {card_class}" id="comment-{comment_id}">
            <div class="card-body py-2">
                <div class="d-flex align-items-center mb-1">
                    <div class="avatar bg-{"secondary" if level > 0 else "primary"} text-white rounded-circle me-2" style="width: {avatar_size}px; height: {avatar_size}px; display: flex; align-items: center; justify-content: center; font-weight: bold; font-size: {font_size - 2}px;">
                        {comment['avatar']}
                    </div>
                    <div>
                        <strong style="font-size: {font_size}px;">{comment['author']}</strong>
                        <small class="text-muted ms-2">{comment['date']}{edited_text}</small>
                        {deleted_badge}
                    </div>
                </div>
                <p class="mb-2" style="font-size: {font_size}px; {deleted_style}">{comment['comment']}</p>
                <div class="d-flex gap-2 align-items-center">
                    {like_button}
                    {reply_button}
                    {edit_delete_buttons}
                </div>

                <!-- Reply form (hidden by default) -->
                <div id="reply-form-{comment_id}" class="mt-3" style="display: none;">
                    <form method="POST" class="d-flex gap-2">
                        <input type="hidden" name="action" value="reply">
                        <input type="hidden" name="parent_comment_id" value="{comment_id}">
                        <input type="text" name="reply_text" class="form-control" placeholder="Write a reply..." required style="font-size: {font_size}px;">
                        <button type="submit" class="btn btn-primary btn-sm" style="font-size: {font_size - 2}px;">Reply</button>
                        <button type="button" class="btn btn-secondary btn-sm" onclick="{reply_click}" style="font-size: {font_size - 2}px;">Cancel</button>
                    </form>
                </div>

                <!-- Nested replies -->
                {render_comment_tree(comment.get('replies', []), draft_name, level + 1)}
            </div>
        </div>
        """

    html += '</div>'
    return html


# Load MLGH data from test files
def load_draft_data():
    """Load draft data from test files"""
    # Return empty list - test documents removed per user request
    drafts = []
    return drafts

def load_group_data():
    """Load group data from test files"""
    groups = []

    # Desirable Properties mapping for better names and descriptions
    dp_descriptions = {
        'dp1-federated-auth': {
            'title': 'Federated Authentication & Accountability',
            'desc': 'Developing standards for federated authentication systems that enable cross-platform identity verification while maintaining accountability and audit trails.'
        },
        'dp2-participant-agency': {
            'title': 'Participant Agency and Empowerment',
            'desc': 'Creating frameworks that empower participants with full control over their digital presence, decision-making authority, and ability to shape their environment.'
        },
        'dp3-adaptive-governance': {
            'title': 'Adaptive Governance Supporting an Exponentially Growing Community',
            'desc': 'Designing governance systems that can scale with exponential community growth while maintaining fairness, participation, and adaptability to emerging challenges.'
        },
        'dp4-data-sovereignty': {
            'title': 'Data Sovereignty and Privacy',
            'desc': 'Establishing protocols for complete data ownership, privacy by design, and user-controlled data portability across the Meta-Layer ecosystem.'
        },
        'dp5-decentralized-namespace': {
            'title': 'Decentralized Namespace',
            'desc': 'Developing decentralized naming systems that provide persistent, user-controlled identifiers and namespaces independent of centralized authorities.'
        },
        'dp6-commerce': {
            'title': 'Commerce',
            'desc': 'Creating secure, transparent commerce protocols that enable value exchange, micropayments, and economic interactions within the Meta-Layer.'
        },
        'dp7-simplicity-interoperability': {
            'title': 'Simplicity and Interoperability',
            'desc': 'Designing systems that reduce complexity while ensuring seamless interoperability between different platforms, tools, and communities.'
        },
        'dp8-collaborative-environment': {
            'title': 'Collaborative Environment and Meta-Communities',
            'desc': 'Building frameworks for meta-communities that span multiple platforms and enable fluid collaboration across organizational boundaries.'
        },
        'dp9-developer-incentives': {
            'title': 'Developer and Community Incentives',
            'desc': 'Creating incentive structures that reward developers and communities for contributing to the ecosystem while aligning with long-term sustainability.'
        },
        'dp10-education': {
            'title': 'Education',
            'desc': 'Developing educational frameworks and tools that help participants understand and effectively use the Meta-Layer capabilities.'
        },
        'dp21-multi-modal': {
            'title': 'Multi-modal',
            'desc': 'Enabling seamless interaction across multiple communication modalities including text, voice, video, AR/VR, and emerging interaction paradigms.'
        },
        'dp11-safe-ethical-ai': {
            'title': 'Safe and Ethical AI',
            'desc': 'Establishing ethical frameworks and safety protocols for AI systems operating within the Meta-Layer to ensure alignment with human values.'
        },
        'dp12-community-ai-governance': {
            'title': 'Community-Based AI Governance',
            'desc': 'Creating community-driven governance models for AI systems that ensure transparency, accountability, and collective oversight.'
        },
        'dp13-ai-containment': {
            'title': 'AI Containment',
            'desc': 'Developing containment strategies and technical measures to prevent AI systems from exceeding intended boundaries or causing unintended consequences.'
        },
        'dp14-trust-transparency': {
            'title': 'Trust and Transparency',
            'desc': 'Building trust through transparent decision-making, auditable processes, and verifiable system behaviors throughout the Meta-Layer.'
        },
        'dp15-security-provenance': {
            'title': 'Security and Provenance',
            'desc': 'Ensuring security through comprehensive provenance tracking, secure infrastructure, and verifiable data lineage across all interactions.'
        },
        'dp16-roadmap-milestones': {
            'title': 'Roadmap and Milestones',
            'desc': 'Developing structured roadmaps with clear milestones that guide the evolution of the Meta-Layer while maintaining community alignment.'
        },
        'dp17-financial-sustainability': {
            'title': 'Financial Sustainability',
            'desc': 'Creating financial models and incentive structures that ensure the long-term sustainability and equitable growth of the Meta-Layer ecosystem.'
        },
        'dp18-feedback-reputation': {
            'title': 'Feedback Loops and Reputation',
            'desc': 'Implementing feedback mechanisms and reputation systems that reward positive contributions and maintain community standards.'
        },
        'dp19-community-engagement': {
            'title': 'Amplifying Presence and Community Engagement',
            'desc': 'Developing systems that amplify community participation, enhance visibility of contributions, and strengthen community bonds.'
        },
        'dp20-community-ownership': {
            'title': 'Community Ownership',
            'desc': 'Ensuring community ownership through decentralized governance, shared decision-making, and equitable distribution of value and control.'
        }
    }

    try:
        with open('/home/ubuntu/datatracker/test/data/group-aliases', 'r') as f:
            for line in f:
                if line.startswith('#') or not line.strip():
                    continue
                # Extract group name from the line
                match = re.search(r'xfilter-([^:]+):', line)
                if match:
                    group_name = match.group(1)

                    # Use specific DP description if available, otherwise generate generic one
                    if group_name in dp_descriptions:
                        dp_info = dp_descriptions[group_name]
                        group_title = dp_info['title']
                        description = dp_info['desc']
                    else:
                        # Fallback for non-DP groups
                        group_title = group_name.replace('-', ' ').title()
                        description = f'The {group_title} Workgroup focuses on {group_title.lower()} standards and protocols for the Internet.'

                    groups.append({
                        'acronym': group_name,
                        'name': f'{group_title} Workgroup',
                        'type': 'Workgroup',
                        'state': 'Active',
                        'chairs': [f'Chair {i+1}' for i in range(1 + (hash(group_name) % 2))],
                        'description': description,
                        'members_require_approval': False  # default: join is instant
                    })
    except FileNotFoundError:
        print("Group aliases file not found")

    # Interface Governance Workgroup (ML-GOVERNANCE) - always include
    groups.append({
        'acronym': 'ml-governance',
        'name': 'Interface Governance Workgroup',
        'type': 'Workgroup',
        'state': 'Active',
        'chairs': [],
        'description': 'Developing governance practices and standards for the interface.',
        'about': 'Developing governance practices and standards for the interface that enable safe human-AI coexistence; connection with greater trust, consent, context; and even human-AI flourishing.',
        'members_require_approval': False
    })

    return groups

# Load the data
DRAFTS = load_draft_data()
GROUPS = load_group_data()

# HTML Templates
BASE_TEMPLATE = """
<!DOCTYPE html>
<html lang="en" data-theme="{theme}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <link rel="icon" type="image/png" href="/static/images/overweb_logo.png">
    <link rel="shortcut icon" type="image/png" href="/static/images/overweb_logo.png">
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
    
    {hypothesis_config}
    
    <style>
        :root {{
            /* Light theme (default) */
            --bg-color: #ffffff;
            --bg-secondary: #f7f9fa;
            --bg-tertiary: #e1e8ed;
            --text-primary: #14171a;
            --text-secondary: #657786;
            --text-muted: #aab8c2;
            --border-color: #e1e8ed;
            --border-hover: #ccd6dd;
            --accent-color: #1d9bf0;
            --accent-hover: #1a8cd8;
            --success-color: #00ba7c;
            --warning-color: #f7b529;
            --error-color: #f4212e;
            --navbar-bg: #ffffff;
            --navbar-text: #14171a;
            --navbar-border: #e1e8ed;
            --card-bg: #ffffff;
            --card-border: #e1e8ed;
            --input-bg: #ffffff;
            --input-border: #657786;
            --shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            --shadow-hover: 0 2px 8px rgba(0, 0, 0, 0.15);
        }}

        [data-theme="dark"] {{
            /* Dark theme */
            --bg-color: #000000;
            --bg-secondary: #16181c;
            --bg-tertiary: #1d1f23;
            --text-primary: #ffffff;
            --text-secondary: #8b98a5;
            --text-muted: #6c7b8a;
            --border-color: #2f3336;
            --border-hover: #3d4043;
            --accent-color: #1d9bf0;
            --accent-hover: #1a8cd8;
            --success-color: #00ba7c;
            --warning-color: #f7b529;
            --error-color: #f4212e;
            --navbar-bg: #16181c;
            --navbar-text: #ffffff;
            --navbar-border: #2f3336;
            --card-bg: #16181c;
            --card-border: #2f3336;
            --input-bg: #16181c;
            --input-border: #3d4043;
            --shadow: 0 1px 3px rgba(0, 0, 0, 0.3);
            --shadow-hover: 0 2px 8px rgba(0, 0, 0, 0.4);
        }}

        * {{
            box-sizing: border-box;
        }}

        body {{
            background-color: var(--bg-color);
            color: var(--text-primary);
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            line-height: 1.5;
            margin: 0;
            min-height: 100vh;
            transition: background-color 0.2s ease, color 0.2s ease;
        }}

        /* Modern navbar similar to X */
        .navbar {{
            background-color: var(--navbar-bg) !important;
            border-bottom: 1px solid var(--navbar-border);
            backdrop-filter: blur(10px);
            -webkit-backdrop-filter: blur(10px);
            box-shadow: var(--shadow);
            padding: 0;
            height: 53px;
            z-index: 1030 !important; /* Below Hypothesis sidebar/panels (~10000) so collapse & dropdown are visible */
            position: relative !important;
            overflow: visible !important;
        }}

        .navbar-brand {{
            color: var(--navbar-text) !important;
            font-weight: 700;
            font-size: 18px;
            padding: 16px 20px;
            margin: 0;
            display: flex;
            align-items: center;
            gap: 8px;
        }}

        .navbar-brand:hover {{
            color: var(--accent-color) !important;
        }}

        .navbar-brand img {{
            height: 24px;
            width: auto;
            object-fit: contain;
        }}

        /* White logo for dark mode */
        [data-theme="dark"] .navbar-brand img {{
            filter: brightness(0) invert(1);
        }}

        .navbar-nav {{
            align-items: center;
        }}

        .nav-link {{
            color: var(--text-secondary) !important;
            font-weight: 500;
            padding: 16px 20px;
            margin: 0;
            border-radius: 0;
            transition: all 0.2s ease;
        }}
        
        .waitlist-tab-flair {{
            background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
            border-radius: 8px 8px 0 0 !important;
            font-weight: 600 !important;
            position: relative;
        }}
        
        .waitlist-tab-flair::after {{
            content: '✨';
            margin-left: 6px;
            font-size: 0.9em;
        }}
        
        .waitlist-tab-flair.active {{
            background: linear-gradient(135deg, rgba(102, 126, 234, 0.2) 0%, rgba(118, 75, 162, 0.2) 100%);
        }}

        .nav-link:hover {{
            background-color: var(--bg-secondary);
            color: var(--accent-color) !important;
        }}

        .nav-link.active {{
            color: var(--accent-color) !important;
            border-bottom: 3px solid var(--accent-color);
            background-color: transparent;
        }}

        /* Theme toggle button */
        .theme-toggle {{
            background: none;
            border: none;
            color: var(--text-secondary);
            font-size: 18px;
            padding: 16px 20px;
            cursor: pointer;
            transition: color 0.2s ease;
        }}

        .theme-toggle:hover {{
            color: var(--accent-color);
        }}

        /* Cards with modern styling */
        .card {{
            background-color: var(--card-bg);
            border: 1px solid var(--card-border);
            border-radius: 16px;
            box-shadow: var(--shadow);
            transition: all 0.2s ease;
        }}

        .card:hover {{
            box-shadow: var(--shadow-hover);
            border-color: var(--border-hover);
        }}

        .card-header {{
            background-color: transparent;
            border-bottom: 1px solid var(--card-border);
            border-radius: 16px 16px 0 0 !important;
            padding: 16px 20px;
            font-weight: 700;
            color: var(--text-primary);
        }}

        .card-body {{
            padding: 20px;
        }}

        /* Buttons styled like X */
        .btn {{
            border-radius: 20px;
            font-weight: 700;
            padding: 8px 16px;
            transition: all 0.2s ease;
        }}

        .btn-primary {{
            background-color: var(--accent-color);
            border-color: var(--accent-color);
            color: white;
        }}

        .btn-primary:hover {{
            background-color: var(--accent-hover);
            border-color: var(--accent-hover);
            transform: translateY(-1px);
        }}

        .btn-outline-primary {{
            border-color: var(--text-secondary);
            color: var(--text-primary);
        }}

        .btn-outline-primary:hover {{
            background-color: var(--accent-color);
            border-color: var(--accent-color);
            color: white;
        }}

        .btn-outline-secondary {{
            border-color: var(--border-color);
            color: var(--text-secondary);
        }}

        .btn-outline-secondary:hover {{
            background-color: var(--bg-secondary);
            border-color: var(--border-hover);
            color: var(--text-primary);
        }}

        /* Form inputs */
        .form-control {{
            background-color: var(--input-bg) !important;
            border: 1px solid var(--input-border) !important;
            border-radius: 8px;
            color: var(--text-primary) !important;
            padding: 12px 16px;
            transition: all 0.2s ease;
        }}

        input.form-control, textarea.form-control, select.form-control {{
            color: var(--text-primary) !important;
            background-color: var(--input-bg) !important;
            border-color: var(--input-border) !important;
        }}

        [data-theme="dark"] input,
        [data-theme="dark"] textarea,
        [data-theme="dark"] select,
        [data-theme="dark"] input.form-control,
        [data-theme="dark"] textarea.form-control,
        [data-theme="dark"] select.form-control {{
            color: #ffffff !important;
            background-color: #16181c !important;
            border-color: #3d4043 !important;
        }}

        .form-control:focus {{
            border-color: var(--accent-color);
            box-shadow: 0 0 0 3px rgba(29, 155, 240, 0.1);
            background-color: var(--input-bg);
        }}

        .form-control::placeholder {{
            color: var(--text-muted);
        }}

        .form-select {{
            background-color: var(--input-bg) !important;
            border: 1px solid var(--input-border) !important;
            border-radius: 8px;
            color: var(--text-primary) !important;
            padding: 12px 16px;
            transition: all 0.2s ease;
        }}

        [data-theme="dark"] .form-select {{
            color: #ffffff !important;
            background-color: #16181c !important;
            border-color: #3d4043 !important;
        }}

        /* Modals */
        [data-theme="dark"] .modal-content {{
            background-color: var(--bg-secondary) !important;
            color: var(--text-primary) !important;
            border: 1px solid var(--border-color) !important;
        }}

        [data-theme="dark"] .modal-header {{
            border-bottom-color: var(--border-color) !important;
        }}

        [data-theme="dark"] .modal-footer {{
            border-top-color: var(--border-color) !important;
        }}

        /* App modals: same style for all (no default Bootstrap look) */
        .modal-content {{
            border-radius: 12px;
            border: 1px solid var(--border-color, #dee2e6);
            box-shadow: 0 0.5rem 2rem rgba(0,0,0,0.15);
        }}
        [data-theme="dark"] .modal-content {{
            box-shadow: 0 0.5rem 2rem rgba(0,0,0,0.4);
        }}
        .modal-header {{
            border-bottom: 1px solid var(--border-color, #dee2e6);
            padding: 1rem 1.25rem;
            border-radius: 12px 12px 0 0;
        }}
        .modal-body {{
            padding: 1.25rem;
        }}
        .modal-footer {{
            border-top: 1px solid var(--border-color, #dee2e6);
            padding: 1rem 1.25rem;
            border-radius: 0 0 12px 12px;
        }}
        .modal-title {{
            font-weight: 600;
        }}

        [data-theme="dark"] .btn-close {{
            filter: invert(1);
        }}

        [data-theme="dark"] .list-group-item {{
            background-color: var(--bg-secondary) !important;
            color: var(--text-primary) !important;
            border-color: var(--border-color) !important;
        }}

        /* Alerts */
        .alert {{
            border-radius: 12px;
            border: none;
            padding: 16px 20px;
        }}

        .alert-info {{
            background-color: rgba(29, 155, 240, 0.1);
            color: var(--accent-color);
        }}

        /* Badges */
        .badge {{
            border-radius: 12px;
            font-weight: 500;
            padding: 4px 8px;
        }}

        /* Tables */
        .table {{
            color: var(--text-primary);
            border-color: var(--border-color);
        }}

        .table thead th {{
            background-color: var(--bg-secondary);
            color: var(--text-primary);
            border-color: var(--border-color);
            font-weight: 600;
            padding: 12px;
        }}

        .table tbody td {{
            background-color: var(--card-bg);
            color: var(--text-primary);
            border-color: var(--border-color);
            padding: 12px;
        }}

        .table-hover tbody tr:hover td {{
            background-color: var(--bg-secondary);
            color: var(--text-primary);
        }}
        
        .table-hover tbody tr:hover td * {{
            color: var(--text-primary);
        }}

        .table-responsive {{
            border-radius: 8px;
            overflow: hidden;
        }}

        /* Pagination */
        .pagination .page-link {{
            background-color: var(--card-bg);
            color: var(--text-primary);
            border-color: var(--border-color);
        }}

        .pagination .page-link:hover {{
            background-color: var(--bg-secondary);
            color: var(--accent-color);
            border-color: var(--border-hover);
        }}

        .pagination .page-item.active .page-link {{
            background-color: var(--accent-color);
            border-color: var(--accent-color);
            color: white;
        }}

        .pagination .page-item.disabled .page-link {{
            background-color: var(--bg-secondary);
            color: var(--text-muted);
            border-color: var(--border-color);
        }}

        /* Dropdown menus in tables */
        .table .dropdown {{
            position: relative;
        }}

        .table .dropdown-menu {{
            background-color: var(--card-bg);
            border-color: var(--border-color);
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.3);
            min-width: 120px;
            z-index: 1050;
            margin-top: 4px;
        }}

        .table .dropdown-item {{
            color: var(--text-primary);
            padding: 8px 16px;
            cursor: pointer;
        }}

        .table .dropdown-item:hover {{
            background-color: var(--bg-secondary);
            color: var(--accent-color);
        }}

        .table .dropdown-toggle {{
            background-color: transparent;
            border-color: var(--accent-color);
            color: var(--accent-color);
        }}

        .table .dropdown-toggle:hover {{
            background-color: var(--accent-color);
            color: white;
            border-color: var(--accent-color);
        }}

        /* Ensure table doesn't clip dropdowns */
        .table-responsive {{
            overflow: visible !important;
        }}

        .card-body {{
            overflow: visible !important;
        }}

        /* Breadcrumbs */
        .breadcrumb {{
            background-color: transparent;
            padding: 0;
            margin-bottom: 20px;
        }}

        .breadcrumb-item a {{
            color: var(--text-secondary);
        }}

        .breadcrumb-item.active {{
            color: var(--text-primary);
            font-weight: 500;
        }}

        /* Flash messages */
        #flash-messages {{
            position: fixed;
            top: 70px;
            right: 20px;
            z-index: 1000;
            max-width: 400px;
        }}

        .flash-message {{
            margin-bottom: 10px;
            padding: 12px 16px;
            border-radius: 12px;
            font-weight: 500;
            box-shadow: var(--shadow);
        }}

        .flash-success {{
            background-color: rgba(0, 186, 124, 0.1);
            color: var(--success-color);
            border: 1px solid rgba(0, 186, 124, 0.2);
        }}

        .flash-error {{
            background-color: rgba(244, 33, 46, 0.1);
            color: var(--error-color);
            border: 1px solid rgba(244, 33, 46, 0.2);
        }}

        .flash-info {{
            background-color: rgba(247, 181, 41, 0.1);
            color: var(--warning-color);
            border: 1px solid rgba(247, 181, 41, 0.2);
        }}

        /* Avatar styling */
        .avatar {{
            border-radius: 50%;
            object-fit: cover;
        }}

        /* Wider content layout for better readability */
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            padding-left: 24px;
            padding-right: 24px;
        }}

        /* Responsive adjustments */
        @media (max-width: 768px) {{
            .navbar-brand {{
                font-size: 16px;
                padding: 16px 15px;
            }}

            .nav-link {{
                padding: 16px 12px;
                font-size: 14px;
            }}

            .theme-toggle {{
                padding: 16px 15px;
            }}

            .card {{
                border-radius: 12px;
            }}

            .card-header {{
                border-radius: 12px 12px 0 0 !important;
            }}

            .container {{
                padding-left: 15px;
                padding-right: 15px;
            }}
        }}

        @media (min-width: 1200px) {{
            .container {{
                padding-left: 40px;
                padding-right: 40px;
            }}
        }}

        /* Custom scrollbar */
        ::-webkit-scrollbar {{
            width: 8px;
        }}

        ::-webkit-scrollbar-track {{
            background: var(--bg-secondary);
        }}

        ::-webkit-scrollbar-thumb {{
            background: var(--border-color);
            border-radius: 4px;
        }}

        ::-webkit-scrollbar-thumb:hover {{
            background: var(--border-hover);
        }}

        /* Dropdown menu above navbar but below Hypothesis UI */
        .dropdown-menu {{
            z-index: 1050 !important;
            border-radius: 12px;
            border: 1px solid var(--border-color);
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.15);
            background-color: var(--card-bg);
            margin-top: 8px;
            overflow: visible !important;
            position: absolute !important;
            top: 100% !important;
            left: 0 !important;
            min-width: 200px;
        }}

        /* Ensure dropdown container doesn't clip */
        .dropdown {{
            position: relative !important;
            overflow: visible !important;
        }}

        /* Prevent any parent from clipping the dropdown */
        .navbar .dropdown {{
            overflow: visible !important;
        }}

        /* Dropdown above navbar */
        .navbar .dropdown-menu {{
            z-index: 1050 !important;
            position: absolute !important;
            top: 100% !important;
            left: 0 !important;
        }}

        .dropdown-item {{
            color: var(--text-primary);
            padding: 12px 16px;
            transition: background-color 0.2s ease;
        }}

        .dropdown-item:hover {{
            background-color: var(--bg-secondary);
            color: var(--accent-color);
        }}

        .dropdown-toggle {{
            border: none;
            background: none;
            color: var(--text-secondary);
            font-weight: 500;
            padding: 16px 12px;
            border-radius: 8px;
            transition: all 0.2s ease;
        }}

        .dropdown-toggle:hover {{
            background-color: var(--bg-secondary);
            color: var(--text-primary);
        }}

        .dropdown-toggle:focus {{
            box-shadow: 0 0 0 3px rgba(29, 155, 240, 0.1);
        }}
    </style>
</head>
<body>
    <nav class="navbar navbar-expand-lg">
        <div class="container">
            <a class="navbar-brand" href="/">
                <img src="/static/images/overweb_logo.png" alt="Overweb" />
                MLGH
            </a>
            <div class="navbar-nav">
                <a class="nav-link" href="/projects/">Layers</a>
                <a class="nav-link" href="/roles/">Roles</a>
                <a class="nav-link" href="/workgroups/">Workgroups</a>
                <a class="nav-link" href="/guilds/">Guilds</a>
                <a class="nav-link" href="/person/">People</a>
                <a class="nav-link" href="/waitlists/">Waitlists</a>
                <a class="nav-link" href="/role-images/">Imagery</a>
                <a class="nav-link" href="/doc/all/">Docs</a>
                <a class="nav-link" href="/submit/">Submit</a>
            </div>
            <div class="navbar-nav ms-auto">
                {user_menu}
                <button class="theme-toggle" id="theme-toggle" title="Toggle theme">
                    <i class="fas fa-moon"></i>
                </button>
            </div>
        </div>
    </nav>

    <div id="flash-messages"></div>
    {content}

    <div class="container-fluid mt-5 py-3" style="border-top: 1px solid var(--border-color); background-color: var(--bg-secondary);">
        <div class="text-center text-muted small">
            Build {build_number} | MLGH Datatracker
        </div>
    </div>

    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/js/bootstrap.bundle.min.js"></script>
    <script>
        // Theme switching functionality
        const themeToggle = document.getElementById('theme-toggle');
        const html = document.documentElement;
        const icon = themeToggle.querySelector('i');

        // Load saved theme - prefer user preference over localStorage
        const userTheme = html.getAttribute('data-theme') || 'dark';
        const savedTheme = userTheme !== 'light' && userTheme !== 'dark' && userTheme !== 'auto' ?
            (localStorage.getItem('theme') || 'dark') : userTheme;
        html.setAttribute('data-theme', savedTheme);
        updateThemeIcon(savedTheme);

        function updateThemeIcon(theme) {{
            if (theme === 'dark') {{
                icon.className = 'fas fa-sun';
                themeToggle.title = 'Switch to light mode';
            }} else {{
                icon.className = 'fas fa-moon';
                themeToggle.title = 'Switch to dark mode';
            }}
        }}

        themeToggle.addEventListener('click', () => {{
            const currentTheme = html.getAttribute('data-theme');
            const newTheme = currentTheme === 'dark' ? 'light' : 'dark';

            html.setAttribute('data-theme', newTheme);
            localStorage.setItem('theme', newTheme);
            updateThemeIcon(newTheme);
        }});

        // Flash message auto-hide
        setTimeout(() => {{
            const flashMessages = document.querySelectorAll('.flash-message');
            flashMessages.forEach(msg => {{
                msg.style.opacity = '0';
                setTimeout(() => msg.remove(), 300);
            }});
        }}, 5000);

        // Web3Auth Integration
        let web3auth = null;

        // Function to load script dynamically
        function loadScript(src) {{
            return new Promise((resolve, reject) => {{
                const script = document.createElement('script');
                script.src = src;
                script.onload = () => resolve();
                script.onerror = (e) => reject(e);
                document.head.appendChild(script);
            }});
        }}

        // Initialize Web3Auth after ensuring scripts are loaded
        async function initWeb3Auth() {{
            try {{
                await loadScript('https://cdn.jsdelivr.net/npm/web3@1.10.0/dist/web3.min.js');
                await loadScript('https://unpkg.com/@web3auth/modal@10.13.1/dist/modal.umd.min.js');

                await new Promise(resolve => {{
                    const checkWeb3Auth = () => {{
                        if (window.Modal && window.Modal.Web3Auth) {{
                            resolve();
                        }} else {{
                            setTimeout(checkWeb3Auth, 100);
                        }}
                    }};
                    checkWeb3Auth();
                }});

                const Web3AuthConstructor = window.Modal.Web3Auth;
                const web3AuthConfig = {{
                    clientId: "BKvRj4akAwrNHHk4UyYCC4zt9KWigdiuosCX5-idVNclsk9hPPQ4_b8grcl0JF4NhT26oLWb3O5K949SVv6lTGk",
                    web3AuthNetwork: 'sapphire_devnet',
                    chainConfig: {{
                        chainNamespace: 'eip155',
                        chainId: '0x1',
                        rpcTarget: 'https://rpc.ankr.com/eth',
                        displayName: 'Ethereum Mainnet',
                        blockExplorerUrl: 'https://etherscan.io',
                        ticker: 'ETH',
                        tickerName: 'Ethereum',
                    }},
                    uiConfig: {{
                        mode: 'dark',
                        theme: {{
                            primary: '#1d9bf0'
                        }},
                        loginMethodsOrder: ['google', 'twitter', 'email_passwordless', 'wallet'],
                        defaultLanguage: 'en',
                    }},
                }};

                web3auth = new Web3AuthConstructor(web3AuthConfig);
                await web3auth.init();
                console.log('Web3Auth initialized successfully');

                // Check if we should trigger login modal
                const urlParams = new URLSearchParams(window.location.search);
                if (urlParams.get('show_login') === '1') {{
                    // Remove the parameter from URL
                    window.history.replaceState({{}}, '', window.location.pathname);
                    // Show login modal
                    await loginWithWeb3Auth();
                }}
            }} catch (error) {{
                console.error('Web3Auth initialization failed:', error);
            }}
        }}

        async function loginWithWeb3Auth() {{
            if (!web3auth) {{
                alert("Web3Auth not initialized. Please refresh the page.");
                return;
            }}

            try {{
                // Logout first to ensure clean state
                try {{
                    await web3auth.logout();
                }} catch (e) {{
                    // Ignore if not logged in
                }}

                // Connect without specifying a provider - shows modal with all options
                const web3authProvider = await web3auth.connect();
                const userInfo = await web3auth.getUserInfo();
                
                console.log('User info received:', userInfo);
                
                // Get wallet address - with retry and error handling
                let evmAddress = '';
                try {{
                    if (web3authProvider) {{
                        const web3 = new Web3(web3authProvider);
                        // Wait a bit for provider to be ready
                        await new Promise(resolve => setTimeout(resolve, 500));
                        const accounts = await web3.eth.getAccounts();
                        if (accounts && accounts.length > 0) {{
                            evmAddress = accounts[0];
                        }}
                    }}
                }} catch (addrError) {{
                    console.warn('Could not get EVM address:', addrError);
                    // Not critical for social logins
                }}

                // Build the payload - handle different structures
                // For email_passwordless, verifierId might be different
                const finalVerifierId = userInfo.verifierId || userInfo.email || evmAddress || 'unknown';
                const finalTypeOfLogin = userInfo.typeOfLogin || 'unknown';
                
                const payload = {{
                    verifierId: finalVerifierId,
                    typeOfLogin: finalTypeOfLogin,
                    email: userInfo.email || '',
                    name: userInfo.name || userInfo.email?.split('@')[0] || '',
                    profileImage: userInfo.profileImage || '',
                    evmAddress: evmAddress || '',
                }};

                console.log('Sending payload:', payload);
                console.log('typeOfLogin:', finalTypeOfLogin);

                // Send to backend
                const response = await fetch('/api/auth/web3auth', {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify(payload)
                }});

                const result = await response.json();
                if (response.ok) {{
                    window.location.href = '/';
                }} else {{
                    console.error('Backend error:', result);
                    alert('Login failed: ' + (result.error || 'Unknown error'));
                }}
            }} catch (error) {{
                console.error('Login failed:', error);
                if (error.message && !error.message.includes('user closed')) {{
                    alert('Login failed: ' + error.message);
                }}
            }}
        }}

        // Initialize Web3Auth on page load
        if (document.readyState === 'loading') {{
            document.addEventListener('DOMContentLoaded', initWeb3Auth);
        }} else {{
            initWeb3Auth();
        }}

        // Make loginWithWeb3Auth available globally
        window.loginWithWeb3Auth = loginWithWeb3Auth;
        
        // Hypothesis Annotation Toggle
        function toggleAnnotations() {{
            const button = document.getElementById('toggle-annotations');
            const text = document.getElementById('annotations-text');
            const currentState = getCookie('annotations') || 'off';
            
            if (currentState === 'off') {{
                setCookie('annotations', 'on', 365);
                button.className = 'btn btn-success w-100 mb-2';
                text.textContent = 'Disable Annotations';
                // Reload page to load Hypothesis
                window.location.reload();
            }} else {{
                setCookie('annotations', 'off', 365);
                button.className = 'btn btn-outline-info w-100 mb-2';
                text.textContent = 'Enable Annotations';
                // Reload page to remove Hypothesis
                window.location.reload();
            }}
        }}
        
        // Cookie helper functions
        function setCookie(name, value, days) {{
            const expires = new Date();
            expires.setTime(expires.getTime() + (days * 24 * 60 * 60 * 1000));
            document.cookie = name + '=' + value + ';expires=' + expires.toUTCString() + ';path=/';
        }}
        
        function getCookie(name) {{
            const nameEQ = name + "=";
            const ca = document.cookie.split(';');
            for(let i = 0; i < ca.length; i++) {{
                let c = ca[i];
                while (c.charAt(0) == ' ') c = c.substring(1, c.length);
                if (c.indexOf(nameEQ) == 0) return c.substring(nameEQ.length, c.length);
            }}
            return null;
        }}
        
        // Update button state on page load
        document.addEventListener('DOMContentLoaded', function() {{
            const button = document.getElementById('toggle-annotations');
            const text = document.getElementById('annotations-text');
            if (button && text) {{
                const currentState = getCookie('annotations') || 'off';
                if (currentState === 'on') {{
                    button.className = 'btn btn-success w-100 mb-2';
                    text.textContent = 'Disable Annotations';
                }} else {{
                    button.className = 'btn btn-outline-info w-100 mb-2';
                    text.textContent = 'Enable Annotations';
                }}
            }}
            
            // Load annotation count
            const countElement = document.getElementById('annotation-count');
            if (countElement) {{
                const documentName = window.location.pathname.split('/').pop().replace('/', '');
                fetch(`/api/annotations/${{documentName}}/count`)
                    .then(response => response.json())
                    .then(data => {{
                        const count = data.count || 0;
                        countElement.innerHTML = `<i class="fas fa-comment-dots me-1"></i>${{count}} annotation${{count !== 1 ? 's' : ''}}`;
                    }})
                    .catch(error => {{
                        countElement.textContent = '';
                    }});
            }}
        }});
    </script>
</body>
</html>
"""

SUBMIT_TEMPLATE = """
<div class="container mt-4">
    <nav aria-label="breadcrumb">
        <ol class="breadcrumb">
            <li class="breadcrumb-item"><a href="/">Home</a></li>
            <li class="breadcrumb-item active">Submit Draft</li>
        </ol>
    </nav>
    
    <h1>Submit Internet-Draft</h1>
    <p class="lead">Submit a new Meta-Layer Draft to the MLGH datatracker</p>
    
    <div class="row">
        <div class="col-md-8">
            <div class="card">
                <div class="card-header">
                    <h5>Draft Submission Form</h5>
                </div>
                <div class="card-body">
                    <div id="flash-messages"></div>
                    
                    <!-- Tabs for Upload File vs From Ordinal -->
                    <ul class="nav nav-tabs mb-3" id="submissionTabs" role="tablist">
                        <li class="nav-item" role="presentation">
                            <button class="nav-link active" id="upload-tab" data-bs-toggle="tab" 
                                    data-bs-target="#upload" type="button" role="tab">
                                <i class="bi bi-upload"></i> Upload File
                            </button>
                        </li>
                        <li class="nav-item" role="presentation">
                            <button class="nav-link" id="ordinal-tab" data-bs-toggle="tab" 
                                    data-bs-target="#ordinal" type="button" role="tab">
                                <i class="bi bi-coin"></i> From Ordinal
                            </button>
                        </li>
                    </ul>
                    
                    <div class="tab-content" id="submissionTabContent">
                        <!-- Upload File Tab -->
                        <div class="tab-pane fade show active" id="upload" role="tabpanel">
                            <form method="POST" enctype="multipart/form-data" id="uploadForm">
                                <input type="hidden" name="sourceType" value="file">
                                
                                <div class="mb-3">
                                    <label for="title" class="form-label">Document Title *</label>
                                    <input type="text" class="form-control" id="title" name="title" required 
                                           placeholder="Enter the document title">
                                </div>
                                
                                <div class="mb-3">
                                    <label for="authors" class="form-label">Authors *</label>
                                    <input type="text" class="form-control" id="authors" name="authors" required 
                                           placeholder="Comma-separated list of authors (e.g., John Doe, Jane Smith)">
                                </div>
                                
                                <div class="mb-3">
                                    <label for="abstract" class="form-label">Abstract</label>
                                    <textarea class="form-control" id="abstract" name="abstract" rows="4" 
                                              placeholder="Brief description of the document"></textarea>
                                </div>
                                
                                {{LAYER_SELECTOR}}
                                
                                <div class="mb-3">
                                    <label for="group" class="form-label">Workgroup (Optional)</label>
                                    <select class="form-select" id="group" name="group">
                                        <option value="">Select a Workgroup</option>
                                        <option value="httpbis">HTTP</option>
                                        <option value="quic">QUIC</option>
                                        <option value="tls">TLS</option>
                                        <option value="dnsop">DNSOP</option>
                                        <option value="rtgwg">RTGWG</option>
                                    </select>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="file" class="form-label">Document File *</label>
                                    <input type="file" class="form-control" id="file" name="file" required 
                                           accept=".pdf,.txt,.xml,.doc,.docx">
                                    <div class="form-text">Supported formats: PDF, TXT, XML, DOC, DOCX (max 16MB)</div>
                                </div>
                                
                                <div class="mb-3">
                                    <div class="form-check">
                                        <input class="form-check-input" type="checkbox" id="terms" required>
                                        <label class="form-check-label" for="terms">
                                            I agree to the <a href="#" target="_blank">MLGH submission terms</a>
                                        </label>
                                    </div>
                                </div>
                                
                                <div class="d-grid gap-2 d-md-flex justify-content-md-end">
                                    <button type="submit" class="btn btn-primary">Submit Draft</button>
                                    <a href="/" class="btn btn-secondary">Cancel</a>
                                </div>
                            </form>
                        </div>
                        
                        <!-- From Ordinal Tab -->
                        <div class="tab-pane fade" id="ordinal" role="tabpanel">
                            <form method="POST" id="ordinalForm">
                                <input type="hidden" name="sourceType" value="ordinal">
                                <input type="hidden" name="ordinalContentUrl" id="ordinalContentUrl">
                                <input type="hidden" name="ordinalContentType" id="ordinalContentType">
                                <input type="hidden" name="inscriptionNumber" id="inscriptionNumber">
                                <input type="hidden" name="blockHeight" id="blockHeight">
                                <input type="hidden" name="inscriptionTimestamp" id="inscriptionTimestamp">
                                
                                <div class="mb-3">
                                    <label for="ordinalId" class="form-label">Inscription ID *</label>
                                    <div class="input-group">
                                        <input type="text" class="form-control" id="ordinalId" name="ordinalId" required 
                                               placeholder="Enter Bitcoin Ordinal inscription ID">
                                        <button class="btn btn-outline-secondary" type="button" id="previewBtn">
                                            <i class="bi bi-eye"></i> Preview
                                        </button>
                                    </div>
                                    <div class="form-text">Enter the inscription ID from ordinals.com</div>
                                </div>
                                
                                <!-- Preview Area -->
                                <div id="ordinalPreview" class="mb-3" style="display: none;">
                                    <div class="card">
                                        <div class="card-header">
                                            <h6 class="mb-0">Ordinal Preview</h6>
                                        </div>
                                        <div class="card-body">
                                            <div id="previewLoading" style="display: none;">
                                                <div class="text-center">
                                                    <div class="spinner-border text-primary" role="status">
                                                        <span class="visually-hidden">Loading...</span>
                                                    </div>
                                                    <p class="mt-2">Loading ordinal content...</p>
                                                </div>
                                            </div>
                                            <div id="previewError" class="alert alert-danger" style="display: none;"></div>
                                            <div id="previewContent"></div>
                                            <div id="previewMetadata" class="mt-3" style="display: none;">
                                                <hr>
                                                <h6>Metadata:</h6>
                                                <ul class="list-unstyled small">
                                                    <li><strong>Inscription ID:</strong> <span id="metaInscriptionId"></span></li>
                                                    <li><strong>Inscription Number:</strong> <span id="metaInscriptionNumber"></span></li>
                                                    <li><strong>Block Height:</strong> <span id="metaBlockHeight"></span></li>
                                                    <li><strong>Timestamp:</strong> <span id="metaTimestamp"></span></li>
                                                    <li><strong>Content Type:</strong> <span id="metaContentType"></span></li>
                                                    <li><strong>Content Size:</strong> <span id="metaContentSize"></span></li>
                                                </ul>
                                            </div>
                                        </div>
                                    </div>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="ordinalTitle" class="form-label">Document Title *</label>
                                    <input type="text" class="form-control" id="ordinalTitle" name="title" required 
                                           placeholder="Enter the document title">
                                </div>
                                
                                <div class="mb-3">
                                    <label for="ordinalAuthors" class="form-label">Authors *</label>
                                    <input type="text" class="form-control" id="ordinalAuthors" name="authors" required 
                                           placeholder="Comma-separated list of authors">
                                </div>
                                
                                <div class="mb-3">
                                    <label for="ordinalAbstract" class="form-label">Abstract</label>
                                    <textarea class="form-control" id="ordinalAbstract" name="abstract" rows="4" 
                                              placeholder="Brief description of the document"></textarea>
                                </div>
                                
                                {{LAYER_SELECTOR}}
                                
                                <div class="mb-3">
                                    <label for="ordinalGroup" class="form-label">Workgroup (Optional)</label>
                                    <select class="form-select" id="ordinalGroup" name="group">
                                        <option value="">Select a Workgroup</option>
                                        <option value="httpbis">HTTP</option>
                                        <option value="quic">QUIC</option>
                                        <option value="tls">TLS</option>
                                        <option value="dnsop">DNSOP</option>
                                        <option value="rtgwg">RTGWG</option>
                                    </select>
                                </div>
                                
                                <div class="mb-3">
                                    <div class="form-check">
                                        <input class="form-check-input" type="checkbox" id="ordinalTerms" required>
                                        <label class="form-check-label" for="ordinalTerms">
                                            I agree to the <a href="#" target="_blank">MLGH submission terms</a>
                                        </label>
                                    </div>
                                </div>
                                
                                <div class="d-grid gap-2 d-md-flex justify-content-md-end">
                                    <button type="submit" class="btn btn-primary" id="ordinalSubmitBtn" disabled>Submit Draft</button>
                                    <a href="/" class="btn btn-secondary">Cancel</a>
                                </div>
                            </form>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        
        <div class="col-md-4">
            <div class="card">
                <div class="card-header">
                    <h5>Submission Guidelines</h5>
                </div>
                <div class="card-body">
                    <h6>File Requirements:</h6>
                    <ul class="small">
                        <li>PDF format preferred</li>
                        <li>Maximum 16MB file size</li>
                        <li>Use standard MLGH formatting</li>
                    </ul>
                    
                    <h6>Ordinal Requirements:</h6>
                    <ul class="small">
                        <li>Content must be < 50KB</li>
                        <li>Supported: Images, Text, Markdown, HTML</li>
                        <li>Valid inscription ID required</li>
                    </ul>
                    
                    <h6>Content Requirements:</h6>
                    <ul class="small">
                        <li>Clear, descriptive title</li>
                        <li>Complete author information</li>
                        <li>Abstract describing the work</li>
                        <li>Proper MLGH document structure</li>
                    </ul>
                    
                    <h6>Review Process:</h6>
                    <ul class="small">
                        <li>Initial technical review</li>
                        <li>Workgroup consideration</li>
                        <li>IESG review (if applicable)</li>
                        <li>Publication decision</li>
                    </ul>
                </div>
            </div>
        </div>
    </div>
</div>

<script>
// Ordinal preview functionality
document.addEventListener('DOMContentLoaded', function() {
    console.log('🔨 BUILD """ + str(BUILD_NUMBER) + """ - Ordinals Module Loaded');
    
    const previewBtn = document.getElementById('previewBtn');
    const ordinalIdInput = document.getElementById('ordinalId');
    const ordinalPreview = document.getElementById('ordinalPreview');
    const previewLoading = document.getElementById('previewLoading');
    const previewError = document.getElementById('previewError');
    const previewContent = document.getElementById('previewContent');
    const previewMetadata = document.getElementById('previewMetadata');
    const ordinalSubmitBtn = document.getElementById('ordinalSubmitBtn');
    
    let previewData = null;
    
    previewBtn.addEventListener('click', async function() {
        const inscriptionId = ordinalIdInput.value.trim();
        
        if (!inscriptionId) {
            alert('Please enter an inscription ID');
            return;
        }
        
        // Show preview area and loading
        ordinalPreview.style.display = 'block';
        previewLoading.style.display = 'block';
        previewError.style.display = 'none';
        previewContent.innerHTML = '';
        previewMetadata.style.display = 'none';
        ordinalSubmitBtn.disabled = true;
        
        try {
            const response = await fetch('/api/ordinal/preview', {
                method: 'POST',
                headers: {
                    'Content-Type': 'application/json'
                },
                body: JSON.stringify({ inscriptionId })
            });
            
            const data = await response.json();
            
            previewLoading.style.display = 'none';
            
            if (!data.success) {
                previewError.textContent = data.error || 'Failed to load ordinal';
                previewError.style.display = 'block';
                return;
            }
            
            // Store preview data
            previewData = data;
            
            // Populate hidden fields
            document.getElementById('ordinalContentUrl').value = data.contentUrl;
            document.getElementById('ordinalContentType').value = data.contentType;
            document.getElementById('inscriptionNumber').value = data.inscriptionNumber || '';
            document.getElementById('blockHeight').value = data.blockHeight || '';
            document.getElementById('inscriptionTimestamp').value = data.timestamp || '';
            
            // Display content based on type
            displayOrdinalContent(data);
            
            // Display metadata
            displayMetadata(data);
            
            // Enable submit button
            ordinalSubmitBtn.disabled = false;
            
        } catch (error) {
            previewLoading.style.display = 'none';
            previewError.textContent = 'Error: ' + error.message;
            previewError.style.display = 'block';
        }
    });
    
    function displayOrdinalContent(data) {
        const contentType = data.contentType;
        const contentUrl = data.contentUrl;
        
        console.log('=== displayOrdinalContent DEBUG ===');
        console.log('contentType:', contentType);
        console.log('contentUrl:', contentUrl);
        console.log('startsWith image:', contentType.startsWith('image/'));
        console.log('includes text/plain:', contentType.includes('text/plain'));
        console.log('includes text/javascript:', contentType.includes('text/javascript'));
        console.log('includes application/json:', contentType.includes('application/json'));
        console.log('includes text/markdown:', contentType.includes('text/markdown'));
        console.log('includes text/html:', contentType.includes('text/html'));
        
        if (contentType.startsWith('image/')) {
            console.log('→ RENDERING AS IMAGE');
            // Display image
            previewContent.innerHTML = `<img src="${contentUrl}" class="img-fluid" alt="Ordinal content" style="max-height: 400px;">`;
        } else if (contentType.includes('text/plain') || contentType.includes('text/javascript') || contentType.includes('application/json')) {
            console.log('→ RENDERING AS TEXT/PLAIN (checking for markdown)');
            // Display plain text (handles charset parameters), but check if it's actually markdown
            fetch(contentUrl)
                .then(res => {
                    console.log('Fetch response status:', res.status);
                    return res.text();
                })
                .then(text => {
                    console.log('Text content length:', text.length);
                    console.log('First 100 chars:', text.substring(0, 100));
                    
                    // Check if text looks like markdown
                    const markdownPatterns = [
                        /^#{1,6}\s+.+$/m,              // Headers: # Header
                        /\[.+\]\(.+\)/,                // Links: [text](url)
                        /!\[.*\]\(.+\)/,               // Images: ![alt](url)
                        /^\s*[-*+]\s+.+$/m,            // Unordered lists
                        /^\s*\d+\.\s+.+$/m,            // Ordered lists
                        /```[\s\S]*?```/,              // Code blocks
                        /^\s*>\s+.+$/m,                // Blockquotes
                        /\*\*.+?\*\*/,                 // Bold
                        /__(.+?)__/,                   // Bold (alt)
                        /\*.+?\*/,                     // Italic
                        /_(.+?)_/                      // Italic (alt)
                    ];
                    
                    const looksLikeMarkdown = markdownPatterns.some(pattern => pattern.test(text));
                    
                    if (looksLikeMarkdown) {
                        console.log('→ DETECTED MARKDOWN in text/plain, converting...');
                        // Treat as markdown
                        fetch('/api/ordinal/convert-markdown', {
                            method: 'POST',
                            headers: { 'Content-Type': 'application/json' },
                            body: JSON.stringify({ markdown: text })
                        })
                        .then(res => {
                            console.log('✅ Markdown API response status:', res.status);
                            return res.json();
                        })
                        .then(result => {
                            console.log('✅ Markdown API result:', result);
                            if (result.success) {
                                console.log('✅ HTML length:', result.html.length);
                                console.log('📄 HTML first 500 chars:', result.html.substring(0, 500));
                                // Fix relative image URLs
                                let html = result.html;
                                const beforeFix = html;
                                html = html.replace(/src="\/content\//g, 'src="https://ordinals.com/content/');
                                html = html.replace(/src='\/content\//g, "src='https://ordinals.com/content/");
                                if (html !== beforeFix) {
                                    console.log('✅ FIXED relative image URLs in frontend');
                                    console.log('📄 Fixed HTML first 500 chars:', html.substring(0, 500));
                                } else {
                                    console.log('⚠️  No relative URLs found to fix in frontend');
                                }
                                previewContent.innerHTML = `<div class="border p-3" style="max-height: 400px; overflow-y: auto;">${html}</div>`;
                                console.log('✅ HTML injected into DOM');
                            } else {
                                console.error('❌ Markdown conversion failed:', result.error);
                                // Fallback to plain text
                                previewContent.innerHTML = `<pre class="border p-3" style="max-height: 400px; overflow-y: auto;">${escapeHtml(text)}</pre>`;
                            }
                        })
                        .catch(err => {
                            console.error('❌ Error calling markdown API:', err);
                            // Fallback to plain text
                            previewContent.innerHTML = `<pre class="border p-3" style="max-height: 400px; overflow-y: auto;">${escapeHtml(text)}</pre>`;
                        });
                    } else {
                        console.log('→ DISPLAYING AS PLAIN TEXT');
                        previewContent.innerHTML = `<pre class="border p-3" style="max-height: 400px; overflow-y: auto;">${escapeHtml(text)}</pre>`;
                    }
                })
                .catch(err => {
                    console.error('Error fetching text:', err);
                    previewContent.innerHTML = `<div class="alert alert-danger">Error loading text: ${err.message}</div>`;
                });
        } else if (contentType.includes('text/markdown')) {
            console.log('→ RENDERING AS MARKDOWN');
            // Display markdown (convert to HTML)
            fetch(contentUrl)
                .then(res => res.text())
                .then(markdown => {
                    return fetch('/api/ordinal/convert-markdown', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({ markdown })
                    });
                })
                .then(res => res.json())
                .then(result => {
                    if (result.success) {
                        // Fix relative image URLs to ordinals.com
                        let html = result.html;
                        html = html.replace(/src="\/content\//g, 'src="https://ordinals.com/content/');
                        html = html.replace(/src='\/content\//g, "src='https://ordinals.com/content/");
                        previewContent.innerHTML = `<div class="border p-3" style="max-height: 400px; overflow-y: auto;">${html}</div>`;
                    }
                });
        } else if (contentType.includes('text/html')) {
            console.log('→ RENDERING AS HTML');
            // Display HTML in sandboxed iframe
            previewContent.innerHTML = `<iframe src="${contentUrl}" sandbox="allow-same-origin" style="width: 100%; height: 400px; border: 1px solid var(--card-border);"></iframe>`;
        } else {
            console.log('→ UNSUPPORTED TYPE');
            previewContent.innerHTML = `<div class="alert alert-info">Content type: ${contentType}<br>Cannot preview this content type.</div>`;
        }
    }
    
    function displayMetadata(data) {
        document.getElementById('metaInscriptionId').textContent = data.inscriptionId;
        document.getElementById('metaInscriptionNumber').textContent = data.inscriptionNumber || 'N/A';
        document.getElementById('metaBlockHeight').textContent = data.blockHeight || 'N/A';
        document.getElementById('metaTimestamp').textContent = data.timestamp || 'N/A';
        document.getElementById('metaContentType').textContent = data.contentType;
        document.getElementById('metaContentSize').textContent = formatBytes(data.contentSize);
        previewMetadata.style.display = 'block';
    }
    
    function formatBytes(bytes) {
        if (bytes < 1024) return bytes + ' B';
        if (bytes < 1024 * 1024) return (bytes / 1024).toFixed(2) + ' KB';
        return (bytes / (1024 * 1024)).toFixed(2) + ' MB';
    }
    
    function escapeHtml(text) {
        const div = document.createElement('div');
        div.textContent = text;
        return div.innerHTML;
    }
});
</script>
"""

@app.before_request
def deployment_safety_check():
    """Block data modifications during deployment"""
    if DEPLOYMENT_MODE and request.method in ['POST', 'PUT', 'DELETE', 'PATCH']:
        # Allow deployment endpoints and static files
        if (request.path.startswith('/_deploy/') or
            request.path.startswith('/static/') or
            request.path in ['/login/', '/logout/']):
            return
        # Block all other data modifications
        print(f"🚨 BLOCKED {request.method} {request.path} - Deployment mode active")
        from flask import jsonify
        return jsonify({'error': 'Data modifications disabled during deployment'}), 403

# ================================================================
# HOST → LAYER MIDDLEWARE
# ================================================================

@app.after_request
def add_security_headers(response):
    """Add security headers including CSP for inline scripts"""
    if IS_DEVELOPMENT:
        # Permissive CSP for development (allows inline scripts)
        response.headers['Content-Security-Policy'] = "default-src 'self' 'unsafe-inline' 'unsafe-eval' data: https: http:; script-src 'self' 'unsafe-inline' 'unsafe-eval' https: http:; style-src 'self' 'unsafe-inline' https: http:;"
    return response

@app.before_request
def resolve_layer_from_host():
    """Resolve project/layer context from subdomain."""
    g.layer = None
    g.layer_slug = None

    host = request.host.split(':')[0].lower()

    if not host.endswith('.' + BASE_DOMAIN):
        return  # Not a subdomain of our domain

    subdomain = host[: -(len(BASE_DOMAIN) + 1)]  # Strip ".themetalayer.org"

    if not subdomain or '.' in subdomain:
        return  # Root domain or multi-level subdomain

    if subdomain in RESERVED_SUBDOMAINS:
        return  # Reserved, no layer context

    project = Project.query.filter_by(slug=subdomain).first()
    if project:
        g.layer = project
        g.layer_slug = subdomain

def calculate_pages_and_words(file_path, filename, max_size_mb=50, timeout_seconds=30):
    """
    Calculate pages and words from a file.
    Returns: (pages, words) tuple
    Defaults to (1, 0) if calculation fails
    
    Security features:
    - File size limit (default 50MB)
    - Processing timeout (default 30s)
    - Safe error handling
    """
    try:
        # Check file size (security: prevent memory exhaustion)
        file_size = os.path.getsize(file_path)
        max_size_bytes = max_size_mb * 1024 * 1024
        if file_size > max_size_bytes:
            print(f"[WARNING] File too large: {file_size} bytes (max {max_size_bytes})")
            return (1, 0)
        
        _, ext = os.path.splitext(filename.lower())
        words = 0
        pages = 1
        
        # Use signal for timeout (Unix-like systems only)
        import signal
        
        def timeout_handler(signum, frame):
            raise TimeoutError("File processing timeout")
        
        # Set timeout alarm (if supported)
        if hasattr(signal, 'SIGALRM'):
            signal.signal(signal.SIGALRM, timeout_handler)
            signal.alarm(timeout_seconds)
        
        try:
            if ext in ['.txt', '.xml']:
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
                words = len(content.split())
                pages = max(1, (words + 499) // 500)  # ~500 words per page
                
            elif ext == '.docx' and DOCX_SUPPORT:
                doc = docx.Document(file_path)
                content_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text)
                content = '\n\n'.join(content_parts)
                words = len(content.split())
                pages = max(1, (words + 499) // 500)
                
            elif ext == '.pdf' and PDF_SUPPORT:
                reader = PyPDF2.PdfReader(file_path)
                content_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(text)
                content = '\n\n'.join(content_parts)
                words = len(content.split())
                pages = len(reader.pages) if reader.pages else max(1, (words + 499) // 500)
        finally:
            # Cancel timeout alarm
            if hasattr(signal, 'SIGALRM'):
                signal.alarm(0)
        
        return (pages, words)
        
    except TimeoutError:
        print(f"[WARNING] File processing timeout for {filename}")
        return (1, 0)
    except Exception as e:
        print(f"[WARNING] Failed to calculate pages/words for {filename}: {e}")
        return (1, 0)  # Default fallback

@app.route('/submit/', methods=['GET', 'POST'])
@require_auth
def submit_draft():
    user_menu = generate_user_menu()
    current_theme = session.get('theme', get_current_user().get('theme', 'dark') if get_current_user() else 'dark')

    # Generate workgroup options dynamically
    group_options = '<option value="">Select a Workgroup</option>'
    for group in GROUPS:
        group_options += f'<option value="{group["acronym"]}">{group["name"]}</option>'

    # Replace the hardcoded options in the template (multiple occurrences for both tabs)
    submit_template = SUBMIT_TEMPLATE
    for _ in range(2):  # Replace in both upload and ordinal tabs
        submit_template = submit_template.replace(
            '''<option value="">Select a Workgroup</option>
                                        <option value="httpbis">HTTP</option>
                                        <option value="quic">QUIC</option>
                                        <option value="tls">TLS</option>
                                        <option value="dnsop">DNSOP</option>
                                        <option value="rtgwg">RTGWG</option>''',
            group_options,
            1  # Replace only one occurrence at a time
        )

    # Layer selector: required for project_id. Use g.layer from subdomain or dropdown.
    projects = Project.query.filter(Project.approval_status == 'approved').order_by(Project.name).all()
    if g.layer:
        layer_selector = f'''
                                <div class="mb-3">
                                    <label class="form-label">Layer</label>
                                    <p class="form-control-plaintext mb-0"><strong>{g.layer.name}</strong> <small class="text-muted">(from URL)</small></p>
                                    <input type="hidden" name="project_id" value="{g.layer.id}">
                                </div>'''
    elif projects:
        opts = '<option value="">Select a layer...</option>' + ''.join(
            f'<option value="{p.id}">{p.name}</option>' for p in projects
        )
        layer_selector = f'''
                                <div class="mb-3">
                                    <label for="project_id" class="form-label">Layer *</label>
                                    <select class="form-select" id="project_id" name="project_id" required>
                                        {opts}
                                    </select>
                                    <div class="form-text">Drafts are submitted to a specific layer.</div>
                                </div>'''
    else:
        layer_selector = '''
                                <div class="mb-3">
                                    <p class="text-warning mb-0">No approved layers available. Submit from a layer subdomain (e.g. overweb.themetalayer.org) or create a layer first.</p>
                                </div>'''
    submit_template = submit_template.replace('{{LAYER_SELECTOR}}', layer_selector)

    if request.method == 'POST':
        # Get common fields
        title = request.form.get('title', '').strip()
        authors = request.form.get('authors', '').strip()
        abstract = request.form.get('abstract', '').strip()
        group = request.form.get('group', '').strip()
        source_type = request.form.get('sourceType', 'file').strip()
        form_project_id = request.form.get('project_id', '').strip()
        project_id = form_project_id or (g.layer.id if g.layer else None)
        
        # Validate project_id when projects exist
        if not project_id and projects:
            flash('Please select a layer for this submission.', 'error')
            return BASE_TEMPLATE.format(title="Submit Internet-Draft - MLGH", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER, hypothesis_config="")
        
        # Process authors (comma-separated)
        authors_list = [a.strip() for a in authors.split(',') if a.strip()]
        
        # Generate submission ID
        import random
        import string
        submission_id = ''.join(random.choices(string.ascii_lowercase + string.digits, k=8))
        
        # Handle based on source type
        if source_type == 'ordinal':
            # Ordinal submission
            ordinal_id = request.form.get('ordinalId', '').strip()
            ordinal_content_url = request.form.get('ordinalContentUrl', '').strip()
            ordinal_content_type = request.form.get('ordinalContentType', '').strip()
            inscription_number = request.form.get('inscriptionNumber', '').strip()
            block_height = request.form.get('blockHeight', '').strip()
            inscription_timestamp = request.form.get('inscriptionTimestamp', '').strip()
            
            # Validation
            if not title or not authors or not ordinal_id:
                flash('Title, authors, and inscription ID are required', 'error')
                return BASE_TEMPLATE.format(title="Submit Internet-Draft - MLGH", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER, hypothesis_config="")
            
            if not ordinal_content_url:
                flash('Please preview the ordinal before submitting', 'error')
                return BASE_TEMPLATE.format(title="Submit Internet-Draft - MLGH", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER, hypothesis_config="")
            
            # Fetch ordinal content and calculate pages/words
            try:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                    'Accept': '*/*',
                    'Connection': 'keep-alive'
                    # Removed Accept-Encoding to avoid compression issues that cause wrong word counts
                }
                response = requests.get(ordinal_content_url, headers=headers, timeout=30)
                response.raise_for_status()
                content_text = response.text
                
                # Calculate pages and words from text
                word_count = len(content_text.split())
                chars_per_page = 3000
                page_count = max(1, (len(content_text) + chars_per_page - 1) // chars_per_page)
            except Exception as e:
                app.logger.error(f"Failed to fetch ordinal content for pages/words: {e}")
                # Use defaults if fetch fails
                page_count = 1
                word_count = 0
            
            # Create submission record with ordinal data
            current_user_info = get_current_user()
            app.logger.info(f"📝 CREATING SUBMISSION:")
            app.logger.info(f"   current_user_info: {current_user_info}")
            app.logger.info(f"   submitted_by will be: {current_user_info['name']}")
            
            # Get doc_type from form (default to 'draft')
            doc_type = request.form.get('doc_type', 'draft').strip() or 'draft'
            if doc_type not in ['draft', 'rfc']:
                doc_type = 'draft'
            
            submission = Submission(
                id=submission_id,
                title=title,
                authors=authors_list,
                abstract=abstract,
                group=group,
                project_id=project_id,
                submitted_by=current_user_info['name'],
                sourceType='ordinal',
                doc_type=doc_type,
                ordinalId=ordinal_id,
                ordinalContentUrl=ordinal_content_url,
                ordinalContentType=ordinal_content_type,
                inscriptionNumber=int(inscription_number) if inscription_number else None,
                blockHeight=int(block_height) if block_height else None,
                inscriptionTimestamp=datetime.strptime(inscription_timestamp.replace(' UTC', ''), '%Y-%m-%d %H:%M:%S') if inscription_timestamp else None,
                pages=page_count,
                words=word_count
            )
            
        else:
            # File upload submission
            file = request.files.get('file')
            
            # Validation
            if not title or not authors or not file:
                flash('Title, authors, and file are required', 'error')
                return BASE_TEMPLATE.format(title="Submit Internet-Draft - MLGH", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER, hypothesis_config="")
            
            # Security: Check file size (max 50MB)
            file.seek(0, os.SEEK_END)
            file_size = file.tell()
            file.seek(0)  # Reset to beginning
            max_size = 50 * 1024 * 1024  # 50MB
            if file_size > max_size:
                flash(f'File too large. Maximum size is 50MB. Your file is {file_size / (1024*1024):.1f}MB.', 'error')
                return BASE_TEMPLATE.format(title="Submit Internet-Draft - MLGH", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER, hypothesis_config="")
            
            # Save file
            filename = f"{submission_id}-{file.filename}"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            # Calculate pages and words
            pages, words = calculate_pages_and_words(file_path, filename)
            
            # Create submission record with file data
            submission = Submission(
                id=submission_id,
                title=title,
                authors=authors_list,
                abstract=abstract,
                group=group,
                project_id=project_id,
                filename=filename,
                file_path=file_path,
                submitted_by=get_current_user()['name'],
                sourceType='file',
                pages=pages,
                words=words
            )
        
        # ML numbers are assigned only when submissions are approved, not on submission
        # submission.ml_number will remain None until approval
        
        # Save to database
        db.session.add(submission)
        db.session.commit()

        # Log the action
        source_desc = f"from ordinal {submission.ordinalId}" if source_type == 'ordinal' else "via file upload"
        add_to_document_history(f"draft-{submission_id}", "submitted", get_current_user()['name'], f"New draft submitted {source_desc}: {title}")

        flash('Draft submitted successfully!', 'success')
        return redirect(f'/submit/status/{submission_id}/')

    return BASE_TEMPLATE.format(title="Submit Internet-Draft - MLGH", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/submit/revision/<draft_name>/', methods=['GET', 'POST'])
@require_auth
def submit_revision(draft_name):
    """Submit a new revision of an existing draft"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', get_current_user().get('theme', 'dark') if get_current_user() else 'dark')
    
    # Find the current draft
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)
    
    # If not found in DRAFTS, try to find as a submission
    submission = None
    if not draft:
        submission = Submission.query.filter_by(id=draft_name).first()
        if submission and submission.status == 'approved':
            draft = {
                'name': submission.id,
                'title': submission.title,
                'authors': ', '.join(submission.authors) if isinstance(submission.authors, list) else submission.authors,
                'abstract': submission.abstract or '',
                'group': submission.group or '',
                'rev': submission.revision_number or '00',
                'ml_number': submission.ml_number,
            }
        elif submission:
            flash('Cannot create revision of unapproved submission', 'error')
            return redirect(f'/submit/status/{submission.id}/')
    
    if not draft:
        flash('Draft not found', 'error')
        return redirect('/doc/all/')
    
    # Inherit project_id from parent draft (revisions belong to same layer)
    parent_sub = Submission.query.filter_by(id=draft_name).first()
    revision_project_id = parent_sub.project_id if parent_sub else (g.layer.id if g.layer else None)
    
    # Determine display ID (ML-Draft-XXX or internal ID)
    display_id = draft.get('ml_number', draft_name) or draft_name
    
    # Calculate new revision number
    current_rev = int(draft.get('rev', '00'))
    new_rev = f"{current_rev + 1:02d}"
    
    if request.method == 'POST':
        # Get form data
        title = request.form.get('title', '').strip()
        authors = request.form.get('authors', '').strip()
        abstract = request.form.get('abstract', '').strip()
        group = request.form.get('group', '').strip()
        what_changed = request.form.get('what_changed', '').strip()
        source_type = request.form.get('sourceType', 'file').strip()
        
        # Process authors
        authors_list = [a.strip() for a in authors.split(',') if a.strip()]
        
        # Generate submission ID
        import random
        import string
        submission_id = ''.join(random.choices(string.ascii_lowercase + string.digits, k=8))
        
        # Handle based on source type
        if source_type == 'ordinal':
            # Ordinal submission
            ordinal_id = request.form.get('ordinalId', '').strip()
            ordinal_content_url = request.form.get('ordinalContentUrl', '').strip()
            ordinal_content_type = request.form.get('ordinalContentType', '').strip()
            inscription_number = request.form.get('inscriptionNumber', '').strip()
            block_height = request.form.get('blockHeight', '').strip()
            inscription_timestamp = request.form.get('inscriptionTimestamp', '').strip()
            
            # Validation
            if not title or not authors or not ordinal_id:
                flash('Title, authors, and inscription ID are required', 'error')
                return redirect(f'/submit/revision/{draft_name}/')
            
            if not ordinal_content_url:
                flash('Please preview the ordinal before submitting', 'error')
                return redirect(f'/submit/revision/{draft_name}/')
            
            # Fetch ordinal content and calculate pages/words
            try:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                    'Accept': '*/*',
                    'Connection': 'keep-alive'
                }
                response = requests.get(ordinal_content_url, headers=headers, timeout=30)
                response.raise_for_status()
                content_text = response.text
                
                word_count = len(content_text.split())
                chars_per_page = 3000
                page_count = max(1, (len(content_text) + chars_per_page - 1) // chars_per_page)
            except Exception as e:
                app.logger.error(f"Failed to fetch ordinal content: {e}")
                page_count = 1
                word_count = 0
            
            # Create revision submission with ordinal data
            submission = Submission(
                id=submission_id,
                title=title,
                authors=authors_list,
                abstract=abstract,
                group=group,
                project_id=revision_project_id,
                submitted_by=get_current_user()['name'],
                sourceType='ordinal',
                doc_type='draft',
                ordinalId=ordinal_id,
                ordinalContentUrl=ordinal_content_url,
                ordinalContentType=ordinal_content_type,
                inscriptionNumber=int(inscription_number) if inscription_number else None,
                blockHeight=int(block_height) if block_height else None,
                inscriptionTimestamp=datetime.strptime(inscription_timestamp.replace(' UTC', ''), '%Y-%m-%d %H:%M:%S') if inscription_timestamp else None,
                pages=page_count,
                words=word_count,
                # Revision fields
                parent_draft_name=draft_name,
                revision_number=new_rev,
                what_changed=what_changed,
                is_revision=True
            )
        else:
            # File upload submission
            file = request.files.get('file')
            
            # Validation
            if not title or not authors or not file:
                flash('Title, authors, and file are required', 'error')
                return redirect(f'/submit/revision/{draft_name}/')
            
            # Security: Check file size (max 50MB)
            file.seek(0, os.SEEK_END)
            file_size = file.tell()
            file.seek(0)
            max_size = 50 * 1024 * 1024
            if file_size > max_size:
                flash(f'File too large. Maximum size is 50MB.', 'error')
                return redirect(f'/submit/revision/{draft_name}/')
            
            # Save file
            filename = f"{submission_id}-{file.filename}"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            # Calculate pages and words
            pages, words = calculate_pages_and_words(file_path, filename)
            
            # Create revision submission with file data
            submission = Submission(
                id=submission_id,
                title=title,
                authors=authors_list,
                abstract=abstract,
                group=group,
                project_id=revision_project_id,
                filename=filename,
                file_path=file_path,
                submitted_by=get_current_user()['name'],
                sourceType='file',
                pages=pages,
                words=words,
                # Revision fields
                parent_draft_name=draft_name,
                revision_number=new_rev,
                what_changed=what_changed,
                is_revision=True
            )
        
        # Save to database
        db.session.add(submission)
        db.session.commit()
        
        # Log the action
        source_desc = f"from ordinal {submission.ordinalId}" if source_type == 'ordinal' else "via file upload"
        change_desc = f" Changes: {what_changed[:100]}" if what_changed else ""
        add_to_document_history(
            draft_name,
            "revision_submitted",
            get_current_user()['name'],
            f"Revision {new_rev} submitted {source_desc}.{change_desc}"
        )
        
        flash(f'Revision {new_rev} submitted successfully!', 'success')
        return redirect(f'/submit/status/{submission_id}/')
    
    # GET: Show form with pre-populated data
    # Generate workgroup options
    group_options = '<option value="">Select a Workgroup</option>'
    for g in GROUPS:
        selected = 'selected' if g['acronym'] == draft.get('group', '') else ''
        group_options += f'<option value="{g["acronym"]}" {selected}>{g["name"]}</option>'
    
    revision_form = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/">Home</a></li>
                <li class="breadcrumb-item"><a href="/doc/draft/{draft_name}/">{display_id}</a></li>
                <li class="breadcrumb-item active">Submit Revision</li>
            </ol>
        </nav>
        
        <h1>Submit New Revision</h1>
        <p class="lead">Submit a new revision of {display_id}</p>
        
        <div class="alert alert-info">
            <i class="fas fa-info-circle me-2"></i>
            <strong>Current Revision:</strong> {draft.get('rev', '00')} → <strong>New Revision:</strong> {new_rev}
        </div>
        
        <form method="POST" enctype="multipart/form-data" id="revisionForm">
            <div class="mb-3">
                <label class="form-label">Draft Name</label>
                <input type="text" class="form-control" value="{display_id}" disabled>
                <input type="hidden" name="draft_name" value="{draft_name}">
                <small class="form-text text-muted">This field cannot be changed for revisions</small>
            </div>
            
            <div class="mb-3">
                <label class="form-label">Title *</label>
                <input type="text" class="form-control" name="title" value="{draft.get('title', '')}" required>
            </div>
            
            <div class="mb-3">
                <label class="form-label">Authors *</label>
                <input type="text" class="form-control" name="authors" value="{draft.get('authors', '')}" required>
                <small class="form-text text-muted">Comma-separated list</small>
            </div>
            
            <div class="mb-3">
                <label class="form-label">Abstract</label>
                <textarea class="form-control" name="abstract" rows="4">{draft.get('abstract', '')}</textarea>
            </div>
            
            <div class="mb-3">
                <label class="form-label">Workgroup</label>
                <select class="form-control" name="group">
                    {group_options}
                </select>
            </div>
            
            <div class="mb-3">
                <label class="form-label">What changed since the last revision?</label>
                <textarea class="form-control" name="what_changed" rows="3" 
                          placeholder="Example: Clarified workgroup role in determining rough consensus; added glossary; no change to core governance principles."></textarea>
                <small class="form-text text-muted">
                    Optional but recommended. Briefly describe substantive changes so reviewers and future readers 
                    can understand what evolved and why. Not required for minor or editorial edits.
                </small>
            </div>
            
            <ul class="nav nav-tabs" role="tablist">
                <li class="nav-item">
                    <a class="nav-link active" data-bs-toggle="tab" href="#upload" onclick="document.getElementById('sourceType').value='file'">Upload File</a>
                </li>
                <li class="nav-item">
                    <a class="nav-link" data-bs-toggle="tab" href="#ordinal" onclick="document.getElementById('sourceType').value='ordinal'">Bitcoin Ordinal</a>
                </li>
            </ul>
            
            <div class="tab-content mt-3">
                <div id="upload" class="tab-pane active">
                    <div class="mb-3">
                        <label class="form-label">Upload Document *</label>
                        <input type="file" class="form-control" name="file" accept=".txt,.pdf,.xml,.docx">
                        <small class="form-text text-muted">Supported formats: TXT, PDF, XML, DOCX</small>
                    </div>
                </div>
                
                <div id="ordinal" class="tab-pane">
                    <div class="mb-3">
                        <label class="form-label">Inscription ID *</label>
                        <input type="text" class="form-control" name="ordinalId" id="ordinalId" 
                               placeholder="e.g., 6fb976ab49dcec017f1e201e84395983204ae1a7c2abf7ced0a85d692e442799i0">
                        <small class="form-text text-muted">The unique inscription ID from Bitcoin</small>
                    </div>
                    
                    <div class="mb-3">
                        <button type="button" class="btn btn-secondary" onclick="previewOrdinal()">
                            <i class="fas fa-eye me-1"></i>Preview Ordinal
                        </button>
                    </div>
                    
                    <div id="ordinalPreview" class="mb-3" style="display: none;">
                        <div class="card">
                            <div class="card-header">
                                <h6>Ordinal Preview</h6>
                            </div>
                            <div class="card-body">
                                <div id="ordinalContent"></div>
                                <input type="hidden" name="ordinalContentUrl" id="ordinalContentUrl">
                                <input type="hidden" name="ordinalContentType" id="ordinalContentType">
                                <input type="hidden" name="inscriptionNumber" id="inscriptionNumber">
                                <input type="hidden" name="blockHeight" id="blockHeight">
                                <input type="hidden" name="inscriptionTimestamp" id="inscriptionTimestamp">
                            </div>
                        </div>
                    </div>
                </div>
            </div>
            
            <input type="hidden" name="sourceType" value="file" id="sourceType">
            
            <div class="mt-4">
                <button type="submit" class="btn btn-success btn-lg">
                    <i class="fas fa-upload me-2"></i>Submit Revision
                </button>
                <a href="/doc/draft/{draft_name}/" class="btn btn-secondary btn-lg ms-2">Cancel</a>
            </div>
        </form>
    </div>
    
    <script>
    async function previewOrdinal() {{
        const inscriptionId = document.getElementById('ordinalId').value.trim();
        if (!inscriptionId) {{
            alert('Please enter an inscription ID');
            return;
        }}
        
        // Show loading
        const preview = document.getElementById('ordinalPreview');
        const content = document.getElementById('ordinalContent');
        content.innerHTML = '<div class="text-center"><i class="fas fa-spinner fa-spin"></i> Loading ordinal...</div>';
        preview.style.display = 'block';
        
        try {{
            // Use our API endpoint to fetch ordinal metadata
            const response = await fetch('/api/ordinal/preview', {{
                method: 'POST',
                headers: {{
                    'Content-Type': 'application/json'
                }},
                body: JSON.stringify({{ inscriptionId }})
            }});
            
            const data = await response.json();
            
            if (!data.success) {{
                content.innerHTML = `<div class="alert alert-danger">Error: ${{data.error}}</div>`;
                return;
            }}
            
            // Fill in hidden form fields
            document.getElementById('ordinalContentUrl').value = data.contentUrl;
            document.getElementById('ordinalContentType').value = data.contentType;
            document.getElementById('inscriptionNumber').value = data.inscriptionNumber || '';
            document.getElementById('blockHeight').value = data.blockHeight || '';
            document.getElementById('inscriptionTimestamp').value = data.timestamp || '';
            
            // Fetch and display content
            const contentResponse = await fetch(data.contentUrl);
            const contentText = await contentResponse.text();
            
            // Check if it's markdown
            const isMarkdown = data.contentType.includes('markdown') || data.contentType.includes('text/plain');
            
            if (isMarkdown) {{
                // Convert markdown to HTML
                const convertResponse = await fetch('/api/ordinal/convert-markdown', {{
                    method: 'POST',
                    headers: {{
                        'Content-Type': 'application/json'
                    }},
                    body: JSON.stringify({{ markdown: contentText }})
                }});
                
                const convertData = await convertResponse.json();
                
                console.log('Markdown conversion response:', convertData);
                console.log('HTML length:', convertData.html ? convertData.html.length : 0);
                console.log('HTML preview:', convertData.html ? convertData.html.substring(0, 200) : 'none');
                
                if (convertData.success) {{
                    // Clear and set the preview content
                    content.innerHTML = '';
                    
                    // Create info alert
                    const infoDiv = document.createElement('div');
                    infoDiv.className = 'alert alert-info mb-3';
                    infoDiv.innerHTML = `<strong>Preview:</strong> Inscription #${{data.inscriptionNumber}} | Block: ${{data.blockHeight}} | Size: ${{(data.contentSize / 1024).toFixed(2)}} KB`;
                    content.appendChild(infoDiv);
                    
                    // Create content container
                    const contentDiv = document.createElement('div');
                    contentDiv.className = 'document-content';
                    contentDiv.style.cssText = 'font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; font-size: 1em; line-height: 1.6; max-height: 600px; overflow-y: auto; padding: 20px; border: 1px solid var(--border-color); border-radius: 8px; background: var(--input-bg); color: var(--text-primary);';
                    
                    // Set the HTML content
                    contentDiv.innerHTML = convertData.html;
                    content.appendChild(contentDiv);
                    
                    // Add CSS to ensure images display properly
                    const images = contentDiv.querySelectorAll('img');
                    images.forEach(img => {{
                        img.style.maxWidth = '100%';
                        img.style.height = 'auto';
                        img.style.display = 'block';
                        img.style.margin = '1em 0';
                    }});
                    
                    console.log('✅ Preview rendered successfully. HTML length:', convertData.html.length);
                    console.log('✅ Images found and styled:', images.length);
                    if (images.length > 0) {{
                        console.log('✅ First image src:', images[0].src);
                    }}
                }} else {{
                    console.error('Markdown conversion failed:', convertData.error);
                    content.innerHTML = `<div class="alert alert-danger">Conversion failed: ${{convertData.error}}</div>`;
                }}
            }} else {{
                content.innerHTML = `<pre style="max-height: 400px; overflow-y: auto; white-space: pre-wrap;">${{contentText.substring(0, 2000)}}</pre>`;
            }}
            
        }} catch (error) {{
            content.innerHTML = `<div class="alert alert-danger">Error loading ordinal: ${{error.message}}</div>`;
        }}
    }}
    </script>
    """
    
    return BASE_TEMPLATE.format(title=f"Submit Revision - {display_id}", theme=current_theme, user_menu=user_menu, content=revision_form, build_number=BUILD_NUMBER, hypothesis_config="")

SUBMISSION_STATUS_TEMPLATE = """
<div class="container mt-4">
    <nav aria-label="breadcrumb">
        <ol class="breadcrumb">
            <li class="breadcrumb-item"><a href="/">Home</a></li>
            <li class="breadcrumb-item"><a href="/submit/">Submit Draft</a></li>
            <li class="breadcrumb-item active">Submission Status</li>
        </ol>
    </nav>
    
    <h1>Submission Status</h1>
    <p class="lead">Track your Internet-Draft submission</p>

    <div id="flash-messages"></div>
    
    <div class="row">
        <div class="col-md-8">
            <div class="card">
                <div class="card-header">
                    <h5>
                        Submission Details
                        {% if is_revision %}
                        <span class="badge bg-success ms-2">Revision {{ revision_number }}</span>
                        {% endif %}
                    </h5>
                </div>
                <div class="card-body">
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Submission ID:</strong></div>
                        <div class="col-sm-9"><code>{{ submission.id }}</code></div>
                    </div>
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Status:</strong></div>
                        <div class="col-sm-9">
                            <span class="badge bg-primary">{{ submission_status_title }}</span>
                            {% if is_ordinal %}
                            <span class="badge bg-info ms-2"><i class="bi bi-coin"></i> Ordinal</span>
                            {% else %}
                            <span class="badge bg-secondary ms-2"><i class="bi bi-file-earmark"></i> File</span>
                            {% endif %}
                        </div>
                    </div>
                    {% if is_revision %}
                    <div class="alert alert-info mb-3">
                        <strong><i class="fas fa-code-branch me-2"></i>This is a revision</strong><br>
                        Revision <strong>{{ revision_number }}</strong> of
                        <a href="/doc/draft/{{ parent_draft_name }}/">{{ parent_draft_name }}</a>
                    </div>
                    {% endif %}
                    {% if what_changed %}
                    <div class="card mb-3">
                        <div class="card-header">
                            <strong>What changed (submitter's explanation)</strong>
                        </div>
                        <div class="card-body">
                            <p class="mb-0">{{ what_changed }}</p>
                        </div>
                    </div>
                    {% endif %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Title:</strong></div>
                        <div class="col-sm-9">{{ submission_title }}</div>
                    </div>
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Authors:</strong></div>
                        <div class="col-sm-9">{{ submission_authors_joined }}</div>
                    </div>
                    {% if ml_number %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>ML Number:</strong></div>
                        <div class="col-sm-9"><code>{{ ml_number }}</code></div>
                    </div>
                    {% endif %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Draft ID:</strong></div>
                        <div class="col-sm-9"><code>{{ submission_id }}</code></div>
                    </div>
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Submitted:</strong></div>
                        <div class="col-sm-9">{{ submission_submitted_at }}</div>
                    </div>
                    {% if is_file %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>File:</strong></div>
                        <div class="col-sm-9">
                            <code>{{ submission_filename }}</code>
                            <a href="/download/{{ submission_id }}" class="btn btn-sm btn-outline-primary ms-2">Download</a>
                        </div>
                    </div>
                    {% endif %}
                    {% if submission_abstract %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Abstract:</strong></div>
                        <div class="col-sm-9">{{ submission_abstract }}</div>
                    </div>
                    {% endif %}

                    {% if is_ordinal %}
                    <h6 class="mt-4">Ordinal Metadata</h6>
                    <div class="card mb-3" style="background-color: var(--bg-secondary);">
                        <div class="card-body">
                            <div class="row mb-2">
                                <div class="col-sm-4"><strong>Inscription ID:</strong></div>
                                <div class="col-sm-8">
                                    <a href="https://ordinals.com/inscription/{{ ordinal_id }}" target="_blank" class="text-decoration-none" style="color: var(--accent-color);">
                                        <code style="font-size: 0.85em;">{{ ordinal_id_short }}</code>
                                    </a>
                                </div>
                            </div>
                            {% if inscription_number %}
                            <div class="row mb-2">
                                <div class="col-sm-4"><strong>Inscription Number:</strong></div>
                                <div class="col-sm-8">{{ inscription_number }}</div>
                            </div>
                            {% endif %}
                            {% if block_height %}
                            <div class="row mb-2">
                                <div class="col-sm-4"><strong>Block Height:</strong></div>
                                <div class="col-sm-8">{{ block_height }}</div>
                            </div>
                            {% endif %}
                            {% if inscription_timestamp %}
                            <div class="row mb-2">
                                <div class="col-sm-4"><strong>Timestamp:</strong></div>
                                <div class="col-sm-8">{{ inscription_timestamp }}</div>
                            </div>
                            {% endif %}
                            <div class="row mb-2">
                                <div class="col-sm-4"><strong>Content Type:</strong></div>
                                <div class="col-sm-8"><code>{{ ordinal_content_type }}</code></div>
                            </div>
                            <div class="row">
                                <div class="col-sm-12">
                                    <a href="https://ordinals.com/inscription/{{ ordinal_id }}" target="_blank" class="btn btn-sm btn-outline-primary">
                                        <i class="bi bi-box-arrow-up-right"></i> View on Ordinals.com
                                    </a>
                                </div>
                            </div>
                        </div>
                    </div>
                    {% endif %}

                    <h6 class="mt-4">Content Preview</h6>
                    {% if content_preview_html %}
                    <div class="border rounded p-3" style="background-color: var(--input-bg); border-color: var(--input-border);">
                        {{ content_preview_html|safe }}
                    </div>
                    {% else %}
                    <div class="border rounded p-3" style="background-color: var(--input-bg); border-color: var(--input-border);">
                        <pre class="mb-0" style="font-size: 0.9em; max-height: 400px; overflow-y: auto; color: var(--text-primary);">{{ file_content }}</pre>
                    </div>
                    {% endif %}

                    {% if is_submitted and is_admin %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Actions:</strong></div>
                        <div class="col-sm-9">
                            <form method="POST" action="/submit/approve/{{ submission_id }}" style="display: inline;">
                                <button type="submit" class="btn btn-success btn-sm">Approve & Publish</button>
                            </form>
                            <form method="POST" action="/submit/reject/{{ submission_id }}" style="display: inline; margin-left: 10px;">
                                <button type="submit" class="btn btn-danger btn-sm">Reject</button>
                            </form>
                        </div>
                    </div>
                    {% endif %}
                </div>
            </div>
            
            <div class="card mt-3">
                <div class="card-header">
                    <h5>Review Timeline</h5>
                </div>
                <div class="card-body">
                    <div class="timeline">
                        <div class="timeline-item">
                            <div class="timeline-marker bg-success"></div>
                            <div class="timeline-content">
                                <h6>Submitted</h6>
                                <p class="text-muted small">{{ submission_submitted_at }}</p>
                            </div>
                        </div>
                        <div class="timeline-item">
                            {% if is_approved_or_rejected %}
                            <div class="timeline-marker bg-success"></div>
                            <div class="timeline-content">
                                <h6>Initial Review</h6>
                                <p class="text-muted small">
                                    {% if is_approved %}Completed{% else %}Rejected{% endif %}
                                    {% if submission_approved_at %}
                                    - {{ submission_approved_at }}
                                    {% elif submission_rejected_at %}
                                    - {{ submission_rejected_at }}
                                    {% endif %}
                                </p>
                            </div>
                            {% else %}
                            <div class="timeline-marker bg-secondary"></div>
                            <div class="timeline-content">
                                <h6>Initial Review</h6>
                                <p class="text-muted small">In Progress</p>
                            </div>
                            {% endif %}
                        </div>
                        {% if is_approved %}
                        <div class="timeline-item">
                            <div class="timeline-marker bg-primary"></div>
                            <div class="timeline-content">
                                <h6>Published</h6>
                                <p class="text-muted small">Available in document repository</p>
                            </div>
                        </div>
                        {% else %}
                        <div class="timeline-item">
                            <div class="timeline-marker bg-secondary"></div>
                            <div class="timeline-content">
                                <h6>Workgroup Review</h6>
                                <p class="text-muted small">Pending initial approval</p>
                            </div>
                        </div>
                        <div class="timeline-item">
                            <div class="timeline-marker bg-secondary"></div>
                            <div class="timeline-content">
                                <h6>MLSG Review</h6>
                                <p class="text-muted small">Pending workgroup review</p>
                            </div>
                        </div>
                        {% endif %}
                    </div>
                </div>
            </div>
        </div>
        
        <div class="col-md-4">
            <div class="card">
                <div class="card-header">
                    <h5>Actions</h5>
                </div>
                <div class="card-body">
                    <a href="/submit/" class="btn btn-primary w-100 mb-2">Submit Another Draft</a>
                    <a href="/doc/all/" class="btn btn-outline-secondary w-100 mb-2">View All Documents</a>
                    <a href="/" class="btn btn-outline-secondary w-100">Back to Home</a>
                </div>
            </div>
            
            <div class="card mt-3">
                <div class="card-header">
                    <h5>Need Help?</h5>
                </div>
                <div class="card-body">
                    <p class="small">If you have questions about your submission:</p>
                    <ul class="small">
                        <li>Check the <a href="#" target="_blank">submission guidelines</a></li>
                        <li>Contact the <a href="mailto:draft@metalayer.org">MLGH Secretariat</a></li>
                        <li>Join the <a href="#" target="_blank">MLGH discussion list</a></li>
                    </ul>
                </div>
            </div>
        </div>
    </div>
</div>

<style>
.timeline {
    position: relative;
    padding-left: 30px;
}

.timeline-item {
    position: relative;
    margin-bottom: 20px;
}

.timeline-marker {
    position: absolute;
    left: -25px;
    top: 5px;
    width: 12px;
    height: 12px;
    border-radius: 50%;
    border: 2px solid #fff;
    box-shadow: 0 0 0 2px #dee2e6;
}

.timeline-content h6 {
    margin-bottom: 5px;
    font-weight: 600;
}

.timeline-content p {
    margin-bottom: 0;
}
</style>
"""

@app.route('/submit/status/')
@require_auth
def submission_status():
    user_menu = generate_user_menu()
    current_theme = session.get('theme', get_current_user().get('theme', 'dark') if get_current_user() else 'dark')

    # Get user's submissions
    user_name = get_current_user()['name']
    submissions = Submission.query.filter_by(submitted_by=user_name).order_by(Submission.submitted_at.desc()).all()

    # Format submissions for template
    submissions_html = ""
    for submission in submissions:
        status_badge = {
            'submitted': 'badge bg-warning text-dark',
            'approved': 'badge bg-success',
            'rejected': 'badge bg-danger',
            'published': 'badge bg-info'
        }.get(submission.status, 'badge bg-secondary')
        
        # Get source type
        source_type = getattr(submission, 'sourceType', 'file')
        source_badge = '<span class="badge bg-info ms-2"><i class="bi bi-coin"></i> Ordinal</span>' if source_type == 'ordinal' else '<span class="badge bg-secondary ms-2"><i class="bi bi-file-earmark"></i> File</span>'
        
        # Get revision info
        is_revision = getattr(submission, 'is_revision', False)
        revision_number = getattr(submission, 'revision_number', '')
        parent_draft_name = getattr(submission, 'parent_draft_name', '')
        revision_badge = f'<span class="badge bg-success ms-2">Revision {revision_number}</span>' if is_revision and revision_number else ''
        
        # Get source info (inscription number or filename)
        if source_type == 'ordinal':
            inscription_number = getattr(submission, 'inscriptionNumber', None)
            ordinal_id = getattr(submission, 'ordinalId', None)
            if inscription_number:
                source_info = f'<p class="mb-2"><strong>Inscription:</strong> #{inscription_number}</p>'
            elif ordinal_id:
                shortened_id = shorten_inscription_id(ordinal_id, 8)
                source_info = f'<p class="mb-2"><strong>Inscription:</strong> <a href="https://ordinals.com/inscription/{ordinal_id}" target="_blank" class="text-decoration-none"><code>{shortened_id}</code></a></p>'
            else:
                source_info = ''
        else:
            filename = getattr(submission, 'filename', None)
            source_info = f'<p class="mb-2"><strong>File:</strong> {filename}</p>' if filename else '<p class="mb-2 text-muted"><em>No file</em></p>'

        submissions_html += f"""
        <div class="submission-item">
            <div class="card mb-3">
                <div class="card-header d-flex justify-content-between align-items-center">
                    <h6 class="mb-0">
                        <a href="/submit/status/{submission.id}/" class="text-decoration-none">
                            {submission.title}
                        </a>
                    </h6>
                    <div>
                        <span class="{status_badge}">{submission.status.title()}</span>
                        {source_badge}
                        {revision_badge}
                    </div>
                </div>
                <div class="card-body">
                    <div class="row">
                        <div class="col-md-8">
                            <p class="mb-2"><strong>Authors:</strong> {', '.join(submission.authors)}</p>
                            <p class="mb-2"><strong>Group:</strong> {submission.group or 'None'}</p>
                            <p class="mb-2"><strong>Submitted:</strong> {submission.submitted_at.strftime('%Y-%m-%d %H:%M')}</p>
                            {source_info}
                            {f'<p class="mb-2"><strong>Abstract:</strong> {submission.abstract[:100]}...</p>' if submission.abstract else ''}
                        </div>
                        <div class="col-md-4 text-end">
                            <a href="/submit/status/{submission.id}/" class="btn btn-sm btn-primary me-2">View Details</a>
                            <a href="/doc/draft/{submission.id}/" class="btn btn-sm btn-outline-primary">View Draft</a>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        """

    content = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/">Home</a></li>
                <li class="breadcrumb-item"><a href="/submit/">Submit Draft</a></li>
                <li class="breadcrumb-item active">My Submissions</li>
            </ol>
        </nav>

        <h1>My Submissions</h1>

        {f'<div class="alert alert-info">You have {len(submissions)} submission(s).</div>' if submissions else '<div class="alert alert-info">You have no submissions yet.</div>'}

        {submissions_html}

        <div class="mt-4">
            <a href="/submit/" class="btn btn-primary">Submit Another Draft</a>
            <a href="/" class="btn btn-secondary ms-2">Back to Home</a>
        </div>
    </div>
    """

    return BASE_TEMPLATE.format(title="My Submissions - MLGH", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/submit/status/<submission_id>/')
@require_auth
def submission_detail(submission_id):
    user_menu = generate_user_menu()
    current_theme = session.get('theme', get_current_user().get('theme', 'dark') if get_current_user() else 'dark')
    current_user = get_current_user()

    submission = Submission.query.filter_by(id=submission_id).first()
    if not submission:
        return "Submission not found", 404

    # Check if user owns this submission or is admin
    app.logger.info(f"🔐 ACCESS CHECK:")
    app.logger.info(f"   submission.submitted_by: {submission.submitted_by}")
    app.logger.info(f"   current_user['name']: {current_user['name']}")
    app.logger.info(f"   current_user.get('role'): {current_user.get('role')}")
    app.logger.info(f"   Match: {submission.submitted_by == current_user['name']}")
    
    if submission.submitted_by != current_user['name'] and current_user.get('role') not in ['admin', 'editor']:
        app.logger.warning(f"❌ ACCESS DENIED!")
        return "Access denied", 403
    
    app.logger.info(f"✅ ACCESS GRANTED")

    # Handle content preview based on source type
    file_content = "File preview not available"
    content_preview_html = ""
    source_type = getattr(submission, 'sourceType', 'file')
    
    print(f"=== BACKEND DEBUG: submission_detail ===")
    print(f"source_type: {source_type}")
    
    if source_type == 'ordinal':
        # Ordinal content - generate preview HTML
        ordinal_content_type = getattr(submission, 'ordinalContentType', '')
        ordinal_content_url = getattr(submission, 'ordinalContentUrl', '')
        
        print(f"ordinal_content_type: {ordinal_content_type}")
        print(f"ordinal_content_url: {ordinal_content_url}")
        
        if ordinal_content_type.startswith('image/'):
            print("→ Rendering as image")
            content_preview_html = f'<img src="{ordinal_content_url}" class="img-fluid" style="max-height: 400px;" alt="Ordinal content">'
        elif 'text/plain' in ordinal_content_type or 'text/javascript' in ordinal_content_type or 'application/json' in ordinal_content_type or 'application/javascript' in ordinal_content_type:
            print("→ Rendering as text/plain or text/javascript or application/json")
            # Fetch and display text-based content (handles charset parameters)
            try:
                import markdown2
                import bleach
                import re
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = requests.get(ordinal_content_url, headers=headers, timeout=10)
                text_content = response.text
                print(f"Text content length: {len(text_content)}")
                print(f"First 100 chars: {text_content[:100]}")
                
                # Check if content is markdown
                is_markdown = False
                if 'text/plain' in ordinal_content_type:
                    # Detect markdown patterns
                    markdown_patterns = [
                        r'^#{1,6}\s+.+$',  # Headers
                        r'\*\*.+\*\*',      # Bold
                        r'\*.+\*',          # Italic
                        r'^\s*[-*+]\s+',    # Lists
                        r'^\s*\d+\.\s+',    # Numbered lists
                        r'\[.+\]\(.+\)',    # Links
                        r'!\[.*\]\(.+\)'    # Images
                    ]
                    for pattern in markdown_patterns:
                        if re.search(pattern, text_content, re.MULTILINE):
                            is_markdown = True
                            print(f"→ DETECTED MARKDOWN (pattern: {pattern})")
                            break
                
                if is_markdown:
                    # Convert markdown to HTML
                    html_content = markdown2.markdown(text_content, extras=['fenced-code-blocks', 'tables', 'break-on-newline'])
                    
                    # Sanitize HTML
                    allowed_tags = ['p', 'br', 'strong', 'em', 'u', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                                  'ul', 'ol', 'li', 'a', 'img', 'code', 'pre', 'blockquote', 'table',
                                  'thead', 'tbody', 'tr', 'th', 'td', 'hr', 'div', 'span']
                    allowed_attrs = {'a': ['href', 'title', 'target'], 'img': ['src', 'alt', 'title', 'width', 'height']}
                    html_content = bleach.clean(html_content, tags=allowed_tags, attributes=allowed_attrs, strip=True)
                    
                    # Fix relative image URLs to point to ordinals.com
                    html_content = re.sub(
                        r'src="(/content/[^"]+)"',
                        r'src="https://ordinals.com\1"',
                        html_content
                    )
                    
                    content_preview_html = f'<div class="border p-3" style="max-height: 400px; overflow-y: auto;">{html_content}</div>'
                    file_content = ""  # Clear file_content since we're using content_preview_html
                else:
                    # Display as plain text
                    if len(text_content) > 2000:
                        file_content = text_content[:2000] + "..."
                    else:
                        file_content = text_content
                print(f"file_content set, length: {len(file_content) if file_content else 0}")
            except Exception as e:
                print(f"ERROR fetching text content: {e}")
                file_content = "Error loading ordinal text content"
        elif 'text/markdown' in ordinal_content_type:
            # Fetch, convert, and display markdown
            try:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = requests.get(ordinal_content_url, headers=headers, timeout=10)
                markdown_text = response.text
                if MARKDOWN_SUPPORT:
                    html_content = markdown2.markdown(
                        markdown_text,
                        extras=['fenced-code-blocks', 'tables', 'break-on-newline']
                    )
                    content_preview_html = f'<div class="border p-3" style="max-height: 400px; overflow-y: auto;">{html_content}</div>'
                else:
                    file_content = markdown_text[:2000] + ("..." if len(markdown_text) > 2000 else "")
            except:
                file_content = "Error loading ordinal markdown content"
        elif 'text/html' in ordinal_content_type:
            print("→ Rendering as HTML iframe")
            # Display HTML in iframe (handles charset parameters)
            content_preview_html = f'<iframe src="{ordinal_content_url}" sandbox="allow-same-origin" style="width: 100%; height: 400px; border: 1px solid var(--card-border);"></iframe>'
        else:
            print(f"→ UNSUPPORTED content type")
            file_content = f"Ordinal content type: {ordinal_content_type}\\nPreview not available for this content type."
    
    elif submission.file_path and os.path.exists(submission.file_path):
        # File upload - extract text for preview
        _, ext = os.path.splitext(submission.filename.lower())

        try:
            if ext in ['.txt', '.xml']:
                # Text-based files can be previewed directly
                with open(submission.file_path, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
                    # Limit preview to first 2000 characters
                    if len(content) > 2000:
                        file_content = content[:2000] + "..."
                    else:
                        file_content = content

            elif ext == '.docx':
                # Extract text from DOCX files
                from docx import Document
                doc = Document(submission.file_path)
                content = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content += paragraph.text + "\n"

                # Also check tables for content
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content += cell.text + "\n"

                if content.strip():
                    # Limit preview to first 2000 characters
                    if len(content) > 2000:
                        file_content = content[:2000] + "..."
                    else:
                        file_content = content
                else:
                    file_content = "DOCX file appears to be empty or contains no extractable text."

            elif ext == '.pdf':
                # For PDFs, use embedded viewer instead of text extraction
                file_size = os.path.getsize(submission.file_path)
                file_size_kb = file_size / 1024
                
                # Create an embedded PDF viewer
                content_preview_html = f'''
                <div class="pdf-viewer-container">
                    <div class="alert alert-info mb-3">
                        <i class="bi bi-file-pdf"></i> PDF Document ({file_size_kb:.1f} KB) - 
                        <a href="/download/{submission.id}" class="alert-link">Download PDF</a> for best viewing experience
                    </div>
                    <iframe src="/view/{submission.id}" 
                            type="application/pdf" 
                            style="width: 100%; height: 600px; border: 1px solid var(--card-border);"
                            title="PDF Preview">
                        <p>Your browser does not support PDF preview. 
                           <a href="/download/{submission.id}">Download the PDF</a> to view it.</p>
                    </iframe>
                </div>
                '''
                file_content = ""  # Clear file_content since we're using content_preview_html

            elif ext == '.doc':
                # Legacy DOC files - show file info
                file_size = os.path.getsize(submission.file_path)
                file_size_kb = file_size / 1024
                file_content = f"Legacy DOC file ({file_size_kb:.1f} KB)\\nText extraction not supported for legacy .doc format.\\nPlease convert to .docx for text preview."

            else:
                # Unknown file type
                file_size = os.path.getsize(submission.file_path)
                file_size_kb = file_size / 1024
                file_content = f"Unsupported file type: {ext} ({file_size_kb:.1f} KB)\\nPreview not available."

        except Exception as e:
            file_size = os.path.getsize(submission.file_path)
            file_size_kb = file_size / 1024
            file_content = f"Error extracting text from {ext[1:].upper()} file ({file_size_kb:.1f} KB): {str(e)}"

    # Use Flask's Jinja2 render_template_string properly - BEST PRACTICE
    # Prepare template variables
    template_vars = {
        'submission': submission,
        'current_user': current_user,
        'file_content': file_content,
        'content_preview_html': content_preview_html,
        'submission_id': submission.id,
        'submission_status': submission.status,
        'submission_status_title': submission.status.title(),
        'submission_title': submission.title or '',
        'submission_abstract': submission.abstract or '',
        'submission_authors': submission.authors,
        'submission_authors_joined': ', '.join(submission.authors) if submission.authors else '',
        'submission_draft_name': getattr(submission, 'draft_name', submission.id) or '',
        'submission_submitted_at': submission.submitted_at.strftime('%Y-%m-%d %H:%M:%S') if submission.submitted_at else '',
        'submission_filename': submission.filename or '',
        'submission_approved_at': submission.approved_at.strftime('%Y-%m-%d %H:%M:%S') if submission.approved_at else '',
        'submission_rejected_at': submission.rejected_at.strftime('%Y-%m-%d %H:%M:%S') if submission.rejected_at else '',
        'is_admin': current_user and (current_user.get('role') in ['admin', 'editor'] or current_user['name'] in ['admin', 'Admin User']),
        'is_approved': submission.status == 'approved',
        'is_rejected': submission.status == 'rejected',
        'is_approved_or_rejected': submission.status in ['approved', 'rejected'],
        'is_submitted': submission.status == 'submitted',
        # Ordinal-specific fields
        'source_type': source_type,
        'is_ordinal': source_type == 'ordinal',
        'is_file': source_type == 'file',
        'ordinal_id': getattr(submission, 'ordinalId', ''),
        'ordinal_id_short': shorten_inscription_id(getattr(submission, 'ordinalId', ''), 8),
        'ordinal_content_url': getattr(submission, 'ordinalContentUrl', ''),
        'ordinal_content_type': getattr(submission, 'ordinalContentType', ''),
        'inscription_number': getattr(submission, 'inscriptionNumber', None),
        'block_height': getattr(submission, 'blockHeight', None),
        'inscription_timestamp': getattr(submission, 'inscriptionTimestamp', None),
        'ml_number': submission.ml_number,
        # Revision fields
        'is_revision': getattr(submission, 'is_revision', False),
        'parent_draft_name': getattr(submission, 'parent_draft_name', ''),
        'revision_number': getattr(submission, 'revision_number', ''),
        'what_changed': getattr(submission, 'what_changed', '')
    }
    
    # Render the submission status template using Flask's Jinja2 engine
    # This properly handles all conditionals and preserves HTML structure
    rendered_content = render_template_string(SUBMISSION_STATUS_TEMPLATE, **template_vars)
    
    # Now use the rendered content in BASE_TEMPLATE (which uses Python .format())
    return BASE_TEMPLATE.format(title=f"Submission {submission.id} - MLGH", theme=current_theme, user_menu=user_menu, content=rendered_content, build_number=BUILD_NUMBER, hypothesis_config="")

LOGIN_TEMPLATE = """
<div class="container mt-4">
    <div class="row justify-content-center">
        <div class="col-md-6">
            <div class="card">
                <div class="card-header">
                    <h3 class="mb-0">Sign In</h3>
                </div>
                <div class="card-body">
                    <div id="flash-messages"></div>

                    <!-- Single Web3Auth Sign In Button -->
                    <div class="mb-4 text-center">
                        <p class="text-muted mb-3">Connect your account to continue</p>
                        <button type="button" class="btn btn-primary btn-lg" id="web3auth-signin-btn" onclick="loginWithWeb3Auth()">
                            <svg width="20" height="20" class="me-2" viewBox="0 0 24 24" fill="currentColor" style="vertical-align: middle;">
                                <path d="M12 2L2 7v10c0 5.55 3.84 10.74 9 12 5.16-1.26 9-6.45 9-12V7l-10-5zm0 18c-3.31 0-6-2.69-6-6s2.69-6 6-6 6 2.69 6 6-2.69 6-6 6z"/>
                            </svg>
                            Sign In with Web3Auth
                        </button>
                        <p class="text-muted mt-3 small">Sign in with Google, Twitter, Email, or connect your wallet</p>
                    </div>

                </div>
            </div>
        </div>
    </div>
</div>

<script>
// Web3Auth Integration
let web3auth = null;

// Function to load script dynamically
function loadScript(src) {
    return new Promise((resolve, reject) => {
        const script = document.createElement('script');
        script.src = src;
        script.onload = () => resolve();
        script.onerror = (e) => reject(e);
        document.head.appendChild(script);
    });
}

// Initialize Web3Auth after ensuring scripts are loaded
async function initWeb3Auth() {
    try {
        // Load Web3 first
        await loadScript('https://cdn.jsdelivr.net/npm/web3@1.10.0/dist/web3.min.js');

        // Load Web3Auth Modal
        await loadScript('https://unpkg.com/@web3auth/modal@10.13.1/dist/modal.umd.min.js');

        // Wait for Web3Auth to be available
        await new Promise(resolve => {
            const checkWeb3Auth = () => {
                if (window.Modal && window.Modal.Web3Auth) {
                    resolve();
                } else {
                    setTimeout(checkWeb3Auth, 100);
                }
            };
            checkWeb3Auth();
        });

        const Web3AuthConstructor = window.Modal.Web3Auth;

        const web3AuthConfig = {
            clientId: "BKvRj4akAwrNHHk4UyYCC4zt9KWigdiuosCX5-idVNclsk9hPPQ4_b8grcl0JF4NhT26oLWb3O5K949SVv6lTGk",
            web3AuthNetwork: 'sapphire_devnet',
            chainConfig: {
                chainNamespace: 'eip155',
                chainId: '0x1',
                rpcTarget: 'https://rpc.ankr.com/eth',
                displayName: 'Ethereum Mainnet',
                blockExplorerUrl: 'https://etherscan.io',
                ticker: 'ETH',
                tickerName: 'Ethereum',
            },
            // Disable modal for direct provider login
            modal: false,
            uiConfig: {
                theme: 'dark',
                loginMethodsOrder: ['google', 'twitter', 'email_passwordless', 'wallet'],
                defaultLanguage: 'en',
            },
            // Force account selection for Google OAuth
            loginConfig: {
                google: {
                    verifier: 'web3auth-google-sapphire-devnet',
                    typeOfLogin: 'google',
                    clientId: 'BKvRj4akAwrNHHk4UyYCC4zt9KWigdiuosCX5-idVNclsk9hPPQ4_b8grcl0JF4NhT26oLWb3O5K949SVv6lTGk',
                    // Force account selection and re-authentication
                    extraLoginOptions: {
                        prompt: 'login select_account',
                        access_type: 'offline'
                    },
                    // Also try query parameters
                    queryParameters: {
                        prompt: 'login select_account',
                        access_type: 'offline'
                    }
                }
            },
        };

        web3auth = new Web3AuthConstructor(web3AuthConfig);
        await web3auth.init();

        console.log('Web3Auth initialized successfully');

    } catch (error) {
        console.error('Web3Auth initialization failed:', error);
    }
}

async function loginWithWeb3Auth(buttonType) {
    if (!web3auth) {
        alert("Web3Auth not initialized. Please refresh the page.");
        return;
    }

    try {
        // Update button to show loading
        const btn = document.getElementById(buttonType + '-login-btn');
        if (btn) {
            btn.innerHTML = 'Connecting...';
        }

        // Connect to Web3Auth with specific login provider
        let loginProvider = null;
        if (buttonType === 'google') loginProvider = 'google';
        else if (buttonType === 'twitter') loginProvider = 'twitter';
        else if (buttonType === 'email') loginProvider = 'email_passwordless';
        else if (buttonType === 'wallet') loginProvider = 'wallet';

        // Connect directly to provider with forced auth
        console.log("Connecting directly to:", loginProvider);
        const web3authProvider = await web3auth.connect({
            loginProvider,
            extraLoginOptions: {
                prompt: 'login select_account',
                access_type: 'offline'
            }
        });

        if (web3authProvider) {
            // Get user info
            const userInfo = await web3auth.getUserInfo();

            // Get wallet address for wallet logins
            let walletAddress = null;
            try {
                const web3 = new Web3(web3authProvider);
                const accounts = await web3.eth.getAccounts();
                walletAddress = accounts[0];
            } catch (walletError) {
                console.log("No wallet address available:", walletError.message);
            }

            // Determine login type
            let loginType = 'unknown';
            if (userInfo.groupedAuthConnectionId) {
                if (userInfo.groupedAuthConnectionId.includes('google')) {
                    loginType = 'google';
                } else if (userInfo.groupedAuthConnectionId.includes('twitter')) {
                    loginType = 'twitter';
                } else if (userInfo.groupedAuthConnectionId.includes('email')) {
                    loginType = 'email';
                } else if (userInfo.groupedAuthConnectionId.includes('wallet')) {
                    loginType = 'wallet';
                }
            } else if (walletAddress) {
                loginType = 'wallet';
            }

            // Send user data to backend
            const requestData = {
                verifierId: userInfo.verifierId || userInfo.groupedAuthConnectionId || `user_${Date.now()}`,
                typeOfLogin: loginType,
                email: userInfo.email,
                name: userInfo.name,
                profileImage: userInfo.profileImage,
                oauthName: userInfo.name,
                evmAddress: walletAddress
            };

            const response = await fetch('/api/auth/web3auth', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify(requestData)
            });

            if (response.ok) {
                // Update button to show final success
                if (btn) {
                    btn.innerHTML = 'Success! Redirecting...';
                }
                // Redirect after a short delay
                setTimeout(() => {
                    window.location.href = '/';
                }, 1000);
            } else {
                const error = await response.json().catch(() => ({}));
                alert('Login failed: ' + (error.error || 'Unknown error'));

                // Reset button
                if (btn) {
                    if (buttonType === 'google') btn.innerHTML = '<svg width="18" height="18" class="me-2" viewBox="0 0 24 24"><path fill="currentColor" d="M22.56 12.25c0-.78-.07-1.53-.2-2.25H12v4.26h5.92c-.26 1.37-1.04 2.53-2.21 3.31v2.77h3.57c2.08-1.92 3.28-4.74 3.28-8.09z"/><path fill="currentColor" d="M12 23c2.97 0 5.46-.98 7.28-2.66l-3.57-2.77c-.98.66-2.23 1.06-3.71 1.06-2.86 0-5.29-1.93-6.16-4.53H2.18v2.84C3.99 20.53 7.7 23 12 23z"/><path fill="currentColor" d="M5.84 14.09c-.22-.66-.35-1.36-.35-2.09s.13-1.43.35-2.09V7.07H2.18C1.43 8.55 1 10.22 1 12s.43 3.45 1.18 4.93l2.85-2.22.81-.62z"/><path fill="currentColor" d="M12 5.38c1.62 0 3.06.56 4.21 1.64l3.15-3.15C17.45 2.09 14.97 1 12 1 7.7 1 3.99 3.47 2.18 7.07l3.66 2.84c.87-2.6 3.3-4.53 6.16-4.53z"/></svg>Continue with Google';
                    else if (buttonType === 'twitter') btn.innerHTML = '<svg width="18" height="18" class="me-2" viewBox="0 0 24 24" fill="currentColor"><path d="M23.953 4.57a10 10 0 01-2.825.775 4.958 4.958 0 002.163-2.723c-.951.555-2.005.959-3.127 1.184a4.92 4.92 0 00-8.384 4.482C7.69 8.095 4.067 6.13 1.64 3.162a4.822 4.822 0 00-.666 2.475c0 1.71.87 3.213 2.188 4.096a4.904 4.904 0 01-2.228-.616v.06a4.923 4.923 0 003.946 4.827 4.996 4.996 0 01-2.212.085 4.936 4.936 0 004.604 3.417 9.867 9.867 0 01-6.102 2.105c-.39 0-.779-.023-1.17-.067a13.995 13.995 0 007.557 2.209c9.053 0 13.998-7.496 13.998-13.985 0-.21 0-.42-.015-.63A9.935 9.935 0 0024 4.59z"/></svg>Continue with X (Twitter)';
                    else if (buttonType === 'email') btn.innerHTML = '<svg width="18" height="18" class="me-2" viewBox="0 0 24 24" fill="currentColor"><path d="M20 4H4c-1.1 0-1.99.9-1.99 2L2 18c0 1.1.9 2 2 2h16c1.1 0 2-.9 2-2V6c0-1.1-.9-2-2-2zm0 4l-8 5-8-5V6l8 5 8-5v2z"/></svg>Continue with Email';
                    else if (buttonType === 'wallet') btn.innerHTML = '<svg width="18" height="18" class="me-2" viewBox="0 0 24 24" fill="currentColor"><path d="M21 7.28V5c0-1.1-.9-2-2-2H5c-1.11 0-2 .9-2 2v14c0 1.1.89 2 2 2h14c1.1 0 2-.9 2-2v-2.28c.59-.35 1-.98 1-1.72V9c0-.74-.41-1.37-1-1.72zM20 9v6h-7V9h7zM5 7h14v10H5V7z"/><circle cx="16" cy="12" r="1.5"/></svg>Connect Wallet';
                }
            }
        }
    } catch (error) {
        console.error('Web3Auth login error:', error);
        alert('Login failed: ' + error.message);

        // Reset button
        const btn = document.getElementById(buttonType + '-login-btn');
        if (btn) {
            if (buttonType === 'google') btn.innerHTML = '<svg width="18" height="18" class="me-2" viewBox="0 0 24 24"><path fill="currentColor" d="M22.56 12.25c0-.78-.07-1.53-.2-2.25H12v4.26h5.92c-.26 1.37-1.04 2.53-2.21 3.31v2.77h3.57c2.08-1.92 3.28-4.74 3.28-8.09z"/><path fill="currentColor" d="M12 23c2.97 0 5.46-.98 7.28-2.66l-3.57-2.77c-.98.66-2.23 1.06-3.71 1.06-2.86 0-5.29-1.93-6.16-4.53H2.18v2.84C3.99 20.53 7.7 23 12 23z"/><path fill="currentColor" d="M5.84 14.09c-.22-.66-.35-1.36-.35-2.09s.13-1.43.35-2.09V7.07H2.18C1.43 8.55 1 10.22 1 12s.43 3.45 1.18 4.93l2.85-2.22.81-.62z"/><path fill="currentColor" d="M12 5.38c1.62 0 3.06.56 4.21 1.64l3.15-3.15C17.45 2.09 14.97 1 12 1 7.7 1 3.99 3.47 2.18 7.07l3.66 2.84c.87-2.6 3.3-4.53 6.16-4.53z"/></svg>Continue with Google';
            else if (buttonType === 'twitter') btn.innerHTML = '<svg width="18" height="18" class="me-2" viewBox="0 0 24 24" fill="currentColor"><path d="M23.953 4.57a10 10 0 01-2.825.775 4.958 4.958 0 002.163-2.723c-.951.555-2.005.959-3.127 1.184a4.92 4.92 0 00-8.384 4.482C7.69 8.095 4.067 6.13 1.64 3.162a4.822 4.822 0 00-.666 2.475c0 1.71.87 3.213 2.188 4.096a4.904 4.904 0 01-2.228-.616v.06a4.923 4.923 0 003.946 4.827 4.996 4.996 0 01-2.212.085 4.936 4.936 0 004.604 3.417 9.867 9.867 0 01-6.102 2.105c-.39 0-.779-.023-1.17-.067a13.995 13.995 0 007.557 2.209c9.053 0 13.998-7.496 13.998-13.985 0-.21 0-.42-.015-.63A9.935 9.935 0 0024 4.59z"/></svg>Continue with X (Twitter)';
            else if (buttonType === 'email') btn.innerHTML = '<svg width="18" height="18" class="me-2" viewBox="0 0 24 24" fill="currentColor"><path d="M20 4H4c-1.1 0-1.99.9-1.99 2L2 18c0 1.1.9 2 2 2h16c1.1 0 2-.9 2-2V6c0-1.1-.9-2-2-2zm0 4l-8 5-8-5V6l8 5 8-5v2z"/></svg>Continue with Email';
            else if (buttonType === 'wallet') btn.innerHTML = '<svg width="18" height="18" class="me-2" viewBox="0 0 24 24" fill="currentColor"><path d="M21 7.28V5c0-1.1-.9-2-2-2H5c-1.11 0-2 .9-2 2v14c0 1.1.89 2 2 2h14c1.1 0 2-.9 2-2v-2.28c.59-.35 1-.98 1-1.72V9c0-.74-.41-1.37-1-1.72zM20 9v6h-7V9h7zM5 7h14v10H5V7z"/><circle cx="16" cy="12" r="1.5"/></svg>Connect Wallet';
        }
    }
}

// Initialize Web3Auth when page loads
document.addEventListener('DOMContentLoaded', () => {
    initWeb3Auth();
});
</script>
"""

REGISTER_TEMPLATE = """
<div class="container mt-4">
    <div class="row justify-content-center">
        <div class="col-md-6">
            <div class="card">
                <div class="card-header">
                    <h3 class="mb-0">Create Account</h3>
                </div>
                <div class="card-body">
                    <div id="flash-messages"></div>
                    
                    <form method="POST">
                        <div class="mb-3">
                            <label for="username" class="form-label">Username</label>
                            <input type="text" class="form-control" id="username" name="username" required>
                        </div>
                        <div class="mb-3">
                            <label for="name" class="form-label">Full Name</label>
                            <input type="text" class="form-control" id="name" name="name" required>
                        </div>
                        <div class="mb-3">
                            <label for="email" class="form-label">Email</label>
                            <input type="email" class="form-control" id="email" name="email" required>
                        </div>
                        <div class="mb-3">
                            <label for="password" class="form-label">Password</label>
                            <input type="password" class="form-control" id="password" name="password" required minlength="6">
                        </div>
                        <div class="d-grid">
                            <button type="submit" class="btn btn-primary">Create Account</button>
                        </div>
                    </form>
                    
                    <hr>
                    <div class="text-center">
                        <p class="mb-0">Already have an account? <a href="/login/">Sign in</a></p>
                    </div>
                </div>
            </div>
        </div>
    </div>
</div>
"""

PROFILE_TEMPLATE = """
<div class="container mt-4">
    <div class="row">
        <div class="col-md-8">
            <div class="card">
                <div class="card-header">
                    <h3 class="mb-0">User Profile</h3>
                </div>
                <div class="card-body">
                    <div id="flash-messages"></div>
                    
                    <!-- Profile Information -->
                    <h5>Profile Information</h5>
                    <form method="POST">
                        <input type="hidden" name="action" value="update_profile">
                        <div class="mb-3">
                            <label for="name" class="form-label">Full Name</label>
                            <input type="text" class="form-control" id="name" name="name" value="{current_user_name}" required>
                        </div>
                        <div class="mb-3">
                            <label for="email" class="form-label">Email</label>
                            <input type="email" class="form-control" id="email" name="email" value="{current_user_email}" required>
                        </div>
                        <div class="mb-3">
                            <label class="form-label">Username</label>
                            <input type="text" class="form-control" value="{session_user}" readonly>
                        </div>
                        <button type="submit" class="btn btn-primary">Update Profile</button>
                    </form>
                    
                    <hr>
                    
                    <!-- Password Change -->
                    <h5>Change Password</h5>
                    <form method="POST">
                        <input type="hidden" name="action" value="update_password">
                        <div class="mb-3">
                            <label for="old_password" class="form-label">Current Password</label>
                            <input type="password" class="form-control" id="old_password" name="old_password" required>
                        </div>
                        <div class="mb-3">
                            <label for="new_password" class="form-label">New Password</label>
                            <input type="password" class="form-control" id="new_password" name="new_password" required minlength="6">
                        </div>
                        <button type="submit" class="btn btn-warning">Change Password</button>
                    </form>

                    <hr>

                    <!-- Theme Preferences -->
                    <h5>Theme Preferences</h5>
                    <form method="POST">
                        <input type="hidden" name="action" value="update_theme">
                        <div class="mb-3">
                            <label class="form-label">Preferred Theme</label>
                            <select class="form-select" name="theme" id="theme-select">
                                <option value="light" {light_selected}>Light Mode</option>
                                <option value="dark" {dark_selected}>Dark Mode</option>
                                <option value="auto" {auto_selected}>Auto (System)</option>
                            </select>
                            <div class="form-text">Choose your preferred theme. Auto will follow your system's preference.</div>
                        </div>
                        <button type="submit" class="btn btn-secondary">Save Theme Preference</button>
                    </form>
                </div>
            </div>
        </div>
        
        <div class="col-md-4">
            <div class="card">
                <div class="card-header">
                    <h5>Account Status</h5>
                </div>
                <div class="card-body">
                    <p><strong>Username:</strong> {session_user}</p>
                    <p><strong>Name:</strong> {current_user_name}</p>
                    <p><strong>Email:</strong> {current_user_email}</p>
                    <p><strong>Status:</strong> <span class="badge bg-success">Active</span></p>
                </div>
            </div>
        </div>
    </div>
</div>
"""

# Authentication routes
@app.route('/login/', methods=['GET'])
def login():
    """Redirect to home and trigger Web3Auth modal"""
    return redirect(url_for('home') + '?show_login=1')

@app.route('/logout/')
def logout():
    """User logout"""
    session.pop('user', None)
    flash('You have been logged out.', 'info')
    return redirect(url_for('home'))

@app.route('/register/', methods=['GET', 'POST'])
def register():
    """User registration"""
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip()
        
        # Check if username or email already exists
        existing_user = User.query.filter((User.username == username) | (User.email == email)).first()
        if existing_user:
            if existing_user.username == username:
                flash('Username already exists.', 'error')
            else:
                flash('Email already registered.', 'error')
        elif len(password) < 6:
            flash('Password must be at least 6 characters.', 'error')
        else:
            # Create new user in database
            new_user = User(
                username=username,
                password_hash=generate_password_hash(password),
                name=name,
                email=email,
                role='user',  # Default role
                theme='dark'  # Default theme
            )
            db.session.add(new_user)
            db.session.commit()

            session['user'] = username
            flash(f'Account created successfully! Welcome, {name}!', 'success')
            return redirect(url_for('home'))
    
    # Generate user menu for register page
    user_menu = """
    <div class="nav-item">
        <a class="nav-link" href="/login/">Sign In</a>
    </div>
    """
    return render_template_string(BASE_TEMPLATE.format(title="Register - MLGH", theme="light", user_menu=user_menu, content=REGISTER_TEMPLATE, build_number=BUILD_NUMBER, hypothesis_config=""))

# Ordinals API routes
@app.route('/api/ordinal/preview', methods=['POST'])
def preview_ordinal():
    """Preview ordinal content and fetch metadata"""
    try:
        data = request.get_json()
        inscription_id = data.get('inscriptionId', '').strip()
        
        if not inscription_id:
            return jsonify({'success': False, 'error': 'Inscription ID is required'}), 400
        
        # Validate inscription ID format (basic validation)
        if len(inscription_id) < 10 or not all(c.isalnum() or c in 'i-_' for c in inscription_id):
            return jsonify({'success': False, 'error': 'Invalid inscription ID format'}), 400
        
        # Build content URL
        content_url = f"https://ordinals.com/content/{inscription_id}"
        
        # Check size and content type with HEAD request
        try:
            # Add headers to avoid 403 errors from ordinals.com
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Accept': '*/*',
                'Connection': 'keep-alive'
                # Removed Accept-Encoding to avoid compression issues
            }
            
            head_response = requests.head(content_url, headers=headers, timeout=10, allow_redirects=True)
            
            if head_response.status_code == 404:
                return jsonify({'success': False, 'error': 'Inscription not found'}), 404
            
            if head_response.status_code == 403:
                return jsonify({
                    'success': False, 
                    'error': 'Access denied by ordinals.com. The inscription exists but cannot be accessed from the server. You can view it directly at: ' + content_url
                }), 403
            
            if head_response.status_code != 200:
                return jsonify({'success': False, 'error': f'Failed to fetch inscription (status: {head_response.status_code})'}), 400
            
            # Get content length
            content_length = int(head_response.headers.get('Content-Length', 0))
            max_size = 50 * 1024  # 50KB
            
            # If HEAD doesn't return Content-Length, try a GET request with streaming
            if content_length == 0:
                try:
                    get_response = requests.get(content_url, headers=headers, timeout=10, stream=True)
                    # Read just enough to check size
                    content_chunk = get_response.raw.read(max_size + 1)
                    content_length = len(content_chunk)
                    content_type = get_response.headers.get('Content-Type', 'unknown').lower()
                except:
                    # If we can't determine size, allow it but note it
                    content_length = 1  # Set to non-zero to proceed
            
            if content_length > max_size:
                size_kb = content_length / 1024
                return jsonify({
                    'success': False, 
                    'error': f'Content too large: {size_kb:.1f}KB (max 50KB)'
                }), 400
            
            # Get content type
            content_type = head_response.headers.get('Content-Type', 'unknown').lower()
            
            # Check if supported type
            supported_types = [
                'image/png', 'image/jpeg', 'image/jpg', 'image/gif', 
                'image/svg+xml', 'image/webp',
                'text/plain', 'text/markdown', 'text/html', 'text/javascript',
                'application/json', 'application/javascript'
            ]
            
            is_supported = any(st in content_type for st in supported_types)
            
            if not is_supported:
                return jsonify({
                    'success': False,
                    'error': f'Unsupported content type: {content_type}. Supported: images, text, markdown, HTML'
                }), 400
            
            # Fetch metadata from ordinals.com HTML (JSON API is disabled on public instance - returns 406)
            inscription_number = None
            block_height = None
            timestamp = None
            
            try:
                import re
                
                # Fetch HTML page
                page_url = f"https://ordinals.com/inscription/{inscription_id}"
                page_headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                app.logger.info(f"🔍 Scraping metadata from HTML: {page_url}")
                
                page_response = requests.get(page_url, headers=page_headers, timeout=10)
                app.logger.info(f"📡 HTML Response status: {page_response.status_code}")
                
                if page_response.status_code == 200:
                    html = page_response.text
                    
                    # Extract inscription number from title: <title>Inscription 117530382</title>
                    number_match = re.search(r'<title>Inscription (\d+)</title>', html)
                    if number_match:
                        inscription_number = int(number_match.group(1))
                        app.logger.info(f"   ✅ Inscription number: {inscription_number}")
                    
                    # Extract block height: <dt>height</dt><dd><a href=/block/933535>933535</a></dd>
                    height_match = re.search(r'<dt>height</dt>\s*<dd><a[^>]*>(\d+)</a></dd>', html)
                    if height_match:
                        block_height = int(height_match.group(1))
                        app.logger.info(f"   ✅ Block height: {block_height}")
                    
                    # Extract timestamp: <dt>timestamp</dt><dd><time>2026-01-23 15:15:43 UTC</time></dd>
                    time_match = re.search(r'<dt>timestamp</dt>\s*<dd><time[^>]*>([^<]+)</time></dd>', html)
                    if time_match:
                        timestamp = time_match.group(1).strip()
                        app.logger.info(f"   ✅ Timestamp: {timestamp}")
                    
                    app.logger.info(f"✅ Metadata scraped: number={inscription_number}, height={block_height}, time={timestamp}")
                else:
                    app.logger.warning(f"⚠️  HTML fetch failed ({page_response.status_code})")
            except Exception as e:
                app.logger.error(f"❌ Error scraping metadata: {e}")
                import traceback
                app.logger.error(traceback.format_exc())
                pass  # Metadata fetch failed, continue without it
            
            return jsonify({
                'success': True,
                'contentUrl': content_url,
                'contentType': content_type,
                'contentSize': content_length,
                'inscriptionId': inscription_id,
                'inscriptionNumber': inscription_number,
                'blockHeight': block_height,
                'timestamp': timestamp
            })
            
        except requests.Timeout:
            return jsonify({'success': False, 'error': 'Request timed out. Please try again.'}), 408
        except requests.ConnectionError:
            return jsonify({'success': False, 'error': 'Ordinals service unavailable. Please try again later.'}), 503
        except Exception as e:
            return jsonify({'success': False, 'error': f'Failed to fetch content: {str(e)}'}), 500
            
    except Exception as e:
        print(f"Ordinal preview error: {e}")
        return jsonify({'success': False, 'error': 'Internal server error'}), 500

@app.route('/api/ordinal/convert-markdown', methods=['POST'])
def convert_markdown():
    """Convert markdown to HTML with sanitization - uses shared processing function"""
    try:
        data = request.get_json()
        markdown_text = data.get('markdown', '')
        
        app.logger.info(f"📝 MARKDOWN CONVERSION REQUEST")
        app.logger.info(f"   Input length: {len(markdown_text)} chars")
        app.logger.info(f"   First 200 chars: {markdown_text[:200]}")
        
        if not markdown_text:
            return jsonify({'success': False, 'error': 'No markdown provided'}), 400
        
        # Use shared markdown processing function
        html_content = process_ordinal_markdown(markdown_text)
        
        app.logger.info(f"   ✅ Processed HTML length: {len(html_content)} chars")
        app.logger.info(f"   📄 HTML first 500 chars: {html_content[:500]}")
        
        return jsonify({'success': True, 'html': html_content})
        
    except Exception as e:
        print(f"Markdown conversion error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': 'Conversion failed'}), 500

# Web3Auth API routes
@app.route('/api/auth/web3auth', methods=['POST'])
def web3auth_login():
    """Web3Auth login endpoint"""
    # Rate limiting: 10 requests per 5 minutes per IP
    client_ip = request.remote_addr or request.environ.get('HTTP_X_FORWARDED_FOR', 'unknown')
    if not check_rate_limit(f"web3auth_{client_ip}", max_requests=10, window_seconds=300):
        return jsonify({'error': 'Rate limit exceeded. Try again later.'}), 429

    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        verifierId = data.get('verifierId')
        typeOfLogin = data.get('typeOfLogin')
        email = data.get('email')
        name = data.get('name')
        profileImage = data.get('profileImage')
        evmAddress = data.get('evmAddress')
        solanaAddress = data.get('solanaAddress')

        if not verifierId:
            return jsonify({'error': 'verifierId required'}), 400

        # Check if user exists in database
        user = User.query.filter_by(web3authVerifierId=verifierId).first()

        # If not found by verifierId, check by email
        if not user and email:
            user = User.query.filter_by(email=email).first()

        # If user exists, update their Web3Auth info and last login
        if user:
            # Update existing user with new Web3Auth data
            user.web3authVerifierId = verifierId
            user.typeOfLogin = typeOfLogin
            user.last_login = datetime.utcnow()
            if name:
                user.displayName = name
                user.displayNameSetAt = datetime.utcnow()
                user.oauthName = name
            if profileImage:
                user.profileImage = profileImage
            if evmAddress:
                user.evmAddress = evmAddress
            if solanaAddress:
                user.solanaAddress = solanaAddress
            db.session.commit()
        else:
            # Create new user
            # Generate handle (use email or wallet address)
            existing_handles = db.session.query(User.username).all()
            existing_handles = [handle[0] for handle in existing_handles]
            if typeOfLogin == 'wallet' and evmAddress:
                # For wallet login, generate handle from wallet address
                short_address = f"{evmAddress[:6]}...{evmAddress[-4:]}"
                handle = f"wallet_{short_address}"
                counter = 1
                while handle in existing_handles:
                    handle = f"wallet_{short_address}_{counter}"
                    counter += 1
            else:
                # For social login, generate from email
                base_handle = email.split('@')[0] if email else 'user'
                base_handle = re.sub(r'[^a-zA-Z0-9_]', '', base_handle)
                if len(base_handle) < 3:
                    base_handle = 'user'
                handle = base_handle
                counter = 1
                while handle in existing_handles:
                    handle = f"{base_handle}{counter}"
                    counter += 1

            # Create user
            user = User(
                web3authVerifierId=verifierId,
                typeOfLogin=typeOfLogin,
                displayName=name if name else None,
                displayNameSetAt=datetime.utcnow() if name else None,
                oauthName=name,
                email=email,
                profileImage=profileImage,
                evmAddress=evmAddress,
                solanaAddress=solanaAddress,
                username=handle,
                handle=handle,
                role='user',
                theme='dark',
                last_login=datetime.utcnow()
            )
            db.session.add(user)
            db.session.commit()

        # Create session
        session['user'] = user.username
        session['theme'] = user.theme

        # Return user data (excluding sensitive info)
        user_data = {
            'id': user.id,
            'username': user.username,
            'displayName': user.displayName,
            'oauthName': user.oauthName,
            'email': user.email,
            'profileImage': user.profileImage,
            'evmAddress': user.evmAddress,
            'typeOfLogin': user.typeOfLogin,
            'theme': user.theme
        }

        # Return sanitized user data (exclude sensitive fields)
        safe_user_data = {
            'id': user.id,
            'username': user.username,
            'displayName': user.displayName,
            'oauthName': user.oauthName,
            'email': user.email,
            'profileImage': user.profileImage,
            'evmAddress': user.evmAddress,
            'solanaAddress': user.solanaAddress,
            'typeOfLogin': user.typeOfLogin,
            'theme': user.theme
        }

        return jsonify({'success': True, 'user': safe_user_data})

    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Web3Auth login error: {e}")
        print(f"Traceback: {error_details}")
        print(f"Request data: {data if 'data' in locals() else 'N/A'}")
        db.session.rollback()
        return jsonify({'error': f'Authentication failed: {str(e)}'}), 500


@app.route('/api/user/me', methods=['GET'])
def get_user_profile():
    """Get current user profile"""
    username = session.get('user')
    if not username:
        return jsonify({'error': 'Unauthorized'}), 401

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({'error': 'User not found'}), 404

    user_data = {
        'id': user.id,
        'public_id': user.public_id,
        'username': user.username,
        'displayName': user.displayName,
        'displayNameSetAt': user.displayNameSetAt.isoformat() if user.displayNameSetAt else None,
        'oauthName': user.oauthName,
        'email': user.email,
        'profileImage': user.profileImage,
        'evmAddress': user.evmAddress,
        'solanaAddress': user.solanaAddress,
        'typeOfLogin': user.typeOfLogin,
        'theme': user.theme
    }

    # Return sanitized user data (exclude sensitive fields)
    safe_user_data = {
        'id': user.id,
        'public_id': user.public_id,
        'username': user.username,
        'displayName': user.displayName,
        'displayNameSetAt': user.displayNameSetAt.isoformat() if user.displayNameSetAt else None,
        'oauthName': user.oauthName,
        'email': user.email,
        'profileImage': user.profileImage,
        'evmAddress': user.evmAddress,
        'solanaAddress': user.solanaAddress,
        'typeOfLogin': user.typeOfLogin,
        'theme': user.theme
    }

    return jsonify({'user': safe_user_data})

@app.route('/api/user/display-name', methods=['PUT'])
def update_display_name():
    """Update user display name"""
    username = session.get('user')
    if not username:
        return jsonify({'error': 'Unauthorized'}), 401

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({'error': 'User not found'}), 404

    data = request.get_json()
    if not data or 'displayName' not in data:
        return jsonify({'error': 'Display name required'}), 400

    displayName = data['displayName'].strip()

    # Validation
    if not displayName:
        return jsonify({'error': 'Display name cannot be empty'}), 400

    if len(displayName) > 50:
        return jsonify({'error': 'Display name must be 50 characters or less'}), 400

    # Optional: Check for allowed characters
    import re
    if not re.match(r'^[a-zA-Z0-9\s\-_]+$', displayName):
        return jsonify({'error': 'Display name can only contain letters, numbers, spaces, hyphens, and underscores'}), 400

    # Update user
    user.displayName = displayName
    user.displayNameSetAt = datetime.utcnow()
    db.session.commit()

    return jsonify({'success': True, 'user': {
        'id': user.id,
        'displayName': user.displayName,
        'displayNameSetAt': user.displayNameSetAt.isoformat()
    }})

@app.route('/api/auth/logout', methods=['POST'])
def api_logout():
    """API logout endpoint"""
    session.pop('user', None)
    return jsonify({'success': True})

# ============================================================================
# Role Images API Endpoints
# ============================================================================

@app.route('/api/role-images/roles-with-stats/', methods=['GET'])
def api_role_images_roles_with_stats():
    """List roles with image count and vote count for role-images gallery page"""
    from sqlalchemy import func
    
    project_id = request.args.get('project_id')
    
    # Subquery: image_count and total_votes (upvotes+downvotes) per role_slug
    stats_subq = db.session.query(
        RoleImage.role_slug,
        func.count(RoleImage.id).label('image_count'),
        func.coalesce(func.sum(RoleImage.upvotes + RoleImage.downvotes), 0).label('vote_count')
    ).group_by(RoleImage.role_slug).subquery()
    
    query = db.session.query(
        Role,
        stats_subq.c.image_count,
        stats_subq.c.vote_count
    ).outerjoin(stats_subq, Role.role_slug == stats_subq.c.role_slug)
    
    if project_id:
        query = query.filter(Role.project_id == project_id)
    
    query = query.order_by(Role.project_id, Role.order, Role.title_guild)
    rows = query.all()
    
    # Build list with role dict + image_count, vote_count
    result = []
    for row in rows:
        role, image_count, vote_count = row
        d = role.to_dict()
        d['image_count'] = image_count or 0
        d['vote_count'] = int(vote_count or 0)
        d['project_name'] = role.project.name if role.project else None
        d['project_slug'] = role.project.slug if role.project else None
        result.append(d)
    
    return jsonify({'roles': result, 'count': len(result)})

@app.route('/api/roles/<role_slug>/images/', methods=['GET'])
def api_list_role_images(role_slug):
    """List role image proposals with vote counts"""
    # Get query parameters
    sort_by = request.args.get('sort', 'net_score')  # net_score, date, upvotes
    include_hidden = request.args.get('include_hidden', 'false').lower() == 'true'
    
    # Build query
    query = RoleImage.query.filter_by(role_slug=role_slug)
    
    # Filter hidden images unless user is admin
    current_user = get_current_user()
    if not (current_user and current_user.get('role') == 'admin') and not include_hidden:
        query = query.filter_by(is_hidden=False)
    
    # Apply sorting
    if sort_by == 'date':
        query = query.order_by(RoleImage.submitted_at.desc())
    elif sort_by == 'upvotes':
        query = query.order_by(RoleImage.upvotes.desc())
    else:  # net_score (default)
        query = query.order_by(RoleImage.net_score.desc(), RoleImage.submitted_at.desc())
    
    images = query.all()
    
    # Add current user's vote to each image
    result = []
    for img in images:
        img_dict = img.to_dict()
        if current_user:
            vote = RoleImageVote.query.filter_by(
                image_id=img.id,
                user_id=current_user['id']
            ).first()
            img_dict['user_vote'] = vote.value if vote else None
        else:
            img_dict['user_vote'] = None
        result.append(img_dict)
    
    return jsonify({'images': result, 'count': len(result)})

@app.route('/api/roles/<role_slug>/images/', methods=['POST'])
@require_auth
def api_submit_role_image(role_slug):
    """Submit a role image proposal"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    # Rate limiting: Check submissions in last 24 hours
    from datetime import timedelta
    recent_submissions = RoleImage.query.filter(
        RoleImage.submitted_by_id == current_user['id'],
        RoleImage.submitted_at >= datetime.utcnow() - timedelta(days=1)
    ).count()
    
    if recent_submissions >= 10:
        return jsonify({'error': 'Rate limit exceeded. Maximum 10 image proposals per day.'}), 429
    
    source_type = request.form.get('source_type')  # upload, url, ordinal
    
    if not source_type or source_type not in ['upload', 'url', 'ordinal']:
        return jsonify({'error': 'Invalid source_type. Must be upload, url, or ordinal.'}), 400
    
    # Generate image ID
    image_id = generate_role_image_id()
    
    # Create image record
    image = RoleImage(
        id=image_id,
        role_slug=role_slug,
        source_type=source_type,
        submitted_by_id=current_user['id']
    )
    
    # Handle different source types
    if source_type == 'upload':
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        image_url, err = upload_image_600x600(
            file, ROLE_IMAGE_UPLOAD_FOLDER, '/uploads/role_images', filename_prefix=image_id
        )
        if err:
            return jsonify({'error': err}), 400
        
        image.file_path = os.path.join(ROLE_IMAGE_UPLOAD_FOLDER, image_url.split('/')[-1])
        image.image_url = image_url
        
    elif source_type == 'url':
        image_url = request.form.get('image_url')
        if not image_url:
            return jsonify({'error': 'image_url required for url source type'}), 400
        image.image_url = image_url
        
    elif source_type == 'ordinal':
        inscription_id = request.form.get('inscription_id')
        if not inscription_id:
            return jsonify({'error': 'inscription_id required for ordinal source type'}), 400
        
        image.chain = request.form.get('chain', 'bitcoin')
        image.inscription_id = inscription_id
        image.content_type = request.form.get('content_type', 'image/png')
        image.image_url = f"https://ordinals.com/content/{inscription_id}"
    
    # Save to database
    db.session.add(image)
    db.session.commit()
    
    return jsonify({'success': True, 'image': image.to_dict()}), 201

@app.route('/api/role-images/<image_id>/', methods=['GET'])
def api_get_role_image(image_id):
    """Get role image details"""
    image = RoleImage.query.get_or_404(image_id)
    
    # Check if hidden
    current_user = get_current_user()
    if image.is_hidden and not (current_user and current_user.get('role') == 'admin'):
        return jsonify({'error': 'Image not found'}), 404
    
    img_dict = image.to_dict()
    
    # Add current user's vote
    if current_user:
        vote = RoleImageVote.query.filter_by(
            image_id=image.id,
            user_id=current_user['id']
        ).first()
        img_dict['user_vote'] = vote.value if vote else None
    else:
        img_dict['user_vote'] = None
    
    return jsonify(img_dict)

@app.route('/api/role-images/<image_id>/vote/', methods=['POST'])
@require_auth
def api_vote_role_image(image_id):
    """Vote on a role image (upvote/downvote)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    image = RoleImage.query.get_or_404(image_id)
    
    # Get vote value from request
    data = request.get_json()
    vote_value = data.get('value')  # 1 for upvote, -1 for downvote
    
    if vote_value not in [1, -1]:
        return jsonify({'error': 'Invalid vote value. Must be 1 (upvote) or -1 (downvote).'}), 400
    
    # Check for existing vote
    existing_vote = RoleImageVote.query.filter_by(
        image_id=image_id,
        user_id=current_user['id']
    ).first()
    
    if existing_vote:
        # Update existing vote
        existing_vote.value = vote_value
        existing_vote.updated_at = datetime.utcnow()
    else:
        # Create new vote
        vote = RoleImageVote(
            image_id=image_id,
            user_id=current_user['id'],
            value=vote_value
        )
        db.session.add(vote)
    
    db.session.commit()
    
    # Update vote counts
    update_image_vote_counts(image_id)
    
    # Return updated image
    image = RoleImage.query.get(image_id)
    return jsonify({'success': True, 'image': image.to_dict()})

@app.route('/api/role-images/<image_id>/vote/', methods=['DELETE'])
@require_auth
def api_remove_vote_role_image(image_id):
    """Remove vote from a role image"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    vote = RoleImageVote.query.filter_by(
        image_id=image_id,
        user_id=current_user['id']
    ).first()
    
    if not vote:
        return jsonify({'error': 'No vote found'}), 404
    
    db.session.delete(vote)
    db.session.commit()
    
    # Update vote counts
    update_image_vote_counts(image_id)
    
    # Return updated image
    image = RoleImage.query.get(image_id)
    return jsonify({'success': True, 'image': image.to_dict()})

@app.route('/api/role-images/<image_id>/promote/', methods=['POST'])
@require_auth
def api_promote_role_image(image_id):
    """Promote image to primary (Project Admin only)"""
    current_user = get_current_user()
    if not current_user or current_user.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    image = RoleImage.query.get_or_404(image_id)
    
    # Demote any existing primary image for this role
    RoleImage.query.filter_by(
        role_slug=image.role_slug,
        is_primary=True
    ).update({'is_primary': False})
    
    # Promote this image
    image.is_primary = True
    image.promoted_by_id = current_user['id']
    image.promoted_at = datetime.utcnow()
    
    db.session.commit()
    
    return jsonify({'success': True, 'image': image.to_dict()})

@app.route('/api/role-images/<image_id>/hide/', methods=['POST'])
@require_auth
def api_hide_role_image(image_id):
    """Hide image (Project Admin only)"""
    current_user = get_current_user()
    if not current_user or current_user.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    image = RoleImage.query.get_or_404(image_id)
    image.is_hidden = True
    db.session.commit()
    
    return jsonify({'success': True, 'image': image.to_dict()})

@app.route('/api/role-images/<image_id>/unhide/', methods=['POST'])
@require_auth
def api_unhide_role_image(image_id):
    """Unhide image (Project Admin only)"""
    current_user = get_current_user()
    if not current_user or current_user.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    image = RoleImage.query.get_or_404(image_id)
    image.is_hidden = False
    db.session.commit()
    
    return jsonify({'success': True, 'image': image.to_dict()})

@app.route('/api/role-images/<image_id>/', methods=['DELETE'])
@require_auth
def api_delete_role_image(image_id):
    """Remove image (Project Admin only)"""
    current_user = get_current_user()
    if not current_user or current_user.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    image = RoleImage.query.get_or_404(image_id)
    
    # Delete associated votes
    RoleImageVote.query.filter_by(image_id=image_id).delete()
    
    # Delete file if it exists
    if image.file_path and os.path.exists(image.file_path):
        try:
            os.remove(image.file_path)
        except Exception as e:
            app.logger.error(f"Failed to delete file {image.file_path}: {e}")
    
    db.session.delete(image)
    db.session.commit()
    
    return jsonify({'success': True})

@app.route('/api/role-images/<image_id>/note/', methods=['PATCH'])
@require_auth
def api_update_role_image_note(image_id):
    """Add/update admin note (Project Admin only)"""
    current_user = get_current_user()
    if not current_user or current_user.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    image = RoleImage.query.get_or_404(image_id)
    
    data = request.get_json()
    admin_note = data.get('admin_note', '')
    
    image.admin_note = admin_note
    db.session.commit()
    
    return jsonify({'success': True, 'image': image.to_dict()})

@app.route('/uploads/role_images/<filename>')
def serve_role_image(filename):
    """Serve uploaded role images"""
    return send_from_directory(ROLE_IMAGE_UPLOAD_FOLDER, filename)

@app.route('/api/upload/entity-image', methods=['POST'])
@require_auth
def api_upload_entity_image():
    """Upload an image for project/workgroup/guild/waitlist. Max 600×600, 5MB. Returns { image_url }."""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400
    entity_type = (request.form.get('entity_type') or 'entity').strip()[:20].replace('/', '_') or 'entity'
    prefix = entity_type.lower()
    image_url, err = upload_image_600x600(
        file, ENTITY_IMAGE_UPLOAD_FOLDER, '/uploads/entity_images', filename_prefix=prefix
    )
    if err:
        return jsonify({'error': err}), 400
    return jsonify({'success': True, 'image_url': image_url}), 201

@app.route('/uploads/entity_images/<filename>')
def serve_entity_image(filename):
    """Serve uploaded entity images (projects, workgroups, guilds, waitlists)"""
    return send_from_directory(ENTITY_IMAGE_UPLOAD_FOLDER, filename)

# ============================================================================
# Projects API Endpoints
# ============================================================================

@app.route('/api/projects/', methods=['GET'])
def api_list_projects():
    """List all projects with filtering"""
    # Get query parameters
    status = request.args.get('status')  # proposed, active, etc.
    approval_status = request.args.get('approval_status')  # pending, approved, rejected
    
    # Build query
    query = Project.query
    
    if status:
        query = query.filter_by(status=status)
    if approval_status:
        query = query.filter_by(approval_status=approval_status)
    
    # Order by last activity (most recent first)
    query = query.order_by(Project.last_activity.desc())
    
    projects = query.all()
    resp = jsonify({'projects': [p.to_dict() for p in projects], 'count': len(projects)})
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate'
    return resp

@app.route('/api/projects/', methods=['POST'])
@require_auth
def api_create_project():
    """Create a new project"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    data = request.get_json()
    name = data.get('name', '').strip()
    mission = data.get('mission') or data.get('mission_statement') or ''
    mission = mission.strip() if mission else None
    description = data.get('description', '').strip()
    
    if not name:
        return jsonify({'error': 'Layer name is required'}), 400
    
    # Check if project name already exists
    existing = Project.query.filter_by(name=name).first()
    if existing:
        return jsonify({'error': 'Layer name already exists'}), 400
    
    # Generate ID and slug
    project_id = generate_project_id()
    slug = create_slug(name)
    
    # Validate slug against reserved subdomains
    if slug in RESERVED_SUBDOMAINS:
        return jsonify({'error': f'The slug "{slug}" is reserved and cannot be used. Please choose a different project name.'}), 400
    
    # Ensure slug is unique
    counter = 1
    original_slug = slug
    while Project.query.filter_by(slug=slug).first():
        slug = f"{original_slug}-{counter}"
        counter += 1
        # Re-check reserved subdomains for modified slug
        if slug in RESERVED_SUBDOMAINS:
            counter += 1
            slug = f"{original_slug}-{counter}"
    
    # Create project
    project = Project(
        id=project_id,
        name=name,
        slug=slug,
        initiator_id=current_user['id'],
        mission=mission or None,
        description=description,
        status='proposed',
        approval_status='pending'
    )
    
    db.session.add(project)
    db.session.commit()
    
    return jsonify({'success': True, 'project': project.to_dict()}), 201

@app.route('/api/projects/<project_id>/', methods=['GET'])
def api_get_project(project_id):
    """Get project details"""
    project = Project.query.get_or_404(project_id)
    
    # Include workgroups count
    workgroups_count = Workgroup.query.filter_by(project_id=project_id).count()
    
    project_dict = project.to_dict()
    project_dict['workgroups_count'] = workgroups_count
    
    # Include membership for current user if authenticated
    current_user = get_current_user()
    if current_user:
        member = ProjectMember.query.filter_by(
            project_id=project_id, user_id=current_user['id'], status='active'
        ).first()
        project_dict['is_member'] = member is not None
        project_dict['member_role'] = member.role if member else None
    else:
        project_dict['is_member'] = False
        project_dict['member_role'] = None
    
    return jsonify(project_dict)

@app.route('/api/projects/<project_id>/', methods=['PATCH'])
@require_auth
def api_update_project(project_id):
    """Update project details"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    project = Project.query.get_or_404(project_id)
    
    # Check permissions (initiator, assigned admin, or site admin)
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Permission denied'}), 403
    
    data = request.get_json()
    
    # Update allowed fields
    if 'name' in data:
        name = (data['name'] or '').strip()
        if not name:
            return jsonify({'error': 'Layer name cannot be empty'}), 400
        if name != project.name:
            if Project.query.filter_by(name=name).first():
                return jsonify({'error': 'A project with this name already exists'}), 400
            project.name = name
            slug = create_slug(name)
            original_slug = slug
            counter = 1
            while Project.query.filter(Project.slug == slug, Project.id != project_id).first():
                slug = f'{original_slug}-{counter}'
                counter += 1
            project.slug = slug
    if 'mission' in data:
        project.mission = data['mission'] if data['mission'] else None
    if 'description' in data:
        project.description = data['description']
    if 'image_url' in data:
        project.image_url = data['image_url'].strip() if data['image_url'] else None
    if 'status' in data and data['status'] in ['proposed', 'active', 'stabilizing', 'maintaining', 'dormant', 'concluded', 'archived']:
        old_status = project.status
        project.status = data['status']
        
        # Record status change
        if old_status != project.status:
            status_change = StatusChange(
                id=generate_status_change_id(),
                entity_type='project',
                entity_id=project_id,
                field_name='status',
                from_value=old_status,
                to_value=project.status,
                note=data.get('status_reason'),
                changed_by_id=current_user['id']
            )
            db.session.add(status_change)
    
    if 'status_reason' in data:
        project.status_reason = data['status_reason']
    
    project.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({'success': True, 'project': project.to_dict()})

@app.route('/api/projects/<project_id>/approve/', methods=['POST'])
@require_auth
def api_approve_project(project_id):
    """Approve or reject a project (admin only)"""
    current_user = get_current_user()
    if not current_user or current_user.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    project = Project.query.get_or_404(project_id)
    
    data = request.get_json()
    action = data.get('action')  # 'approve' or 'reject'
    
    if action not in ['approve', 'reject']:
        return jsonify({'error': 'Invalid action. Must be approve or reject.'}), 400
    
    old_status = project.approval_status
    project.approval_status = 'approved' if action == 'approve' else 'rejected'
    project.approved_by_id = current_user['id']
    project.approved_at = datetime.utcnow()
    
    # Record status change
    status_change = StatusChange(
        id=generate_status_change_id(),
        entity_type='project',
        entity_id=project_id,
        field_name='approval_status',
        from_value=old_status,
        to_value=project.approval_status,
        note=data.get('note'),
        changed_by_id=current_user['id']
    )
    db.session.add(status_change)
    db.session.commit()
    
    return jsonify({'success': True, 'project': project.to_dict()})

# ============================================================================
# Project Admins API
# ============================================================================

@app.route('/api/projects/<project_id>/admins/', methods=['GET'])
def api_list_project_admins(project_id):
    """List project admins (owner + assigned). Only project admins can see this."""
    project = Project.query.get_or_404(project_id)
    current_user = get_current_user()
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project admins can view the admin list'}), 403
    
    owner = project.initiator
    owner_dict = {
        'user_id': owner.id,
        'username': owner.username,
        'display_name': owner.displayName or owner.username,
        'is_owner': True,
        'added_at': project.created_at.isoformat() if project.created_at else None
    }
    
    assigned = ProjectAdmin.query.filter_by(project_id=project_id).all()
    assigned_list = []
    for pa in assigned:
        u = pa.user
        assigned_list.append({
            'user_id': u.id,
            'username': u.username,
            'display_name': u.displayName or u.username,
            'is_owner': False,
            'added_at': pa.added_at.isoformat() if pa.added_at else None
        })
    
    return jsonify({
        'owner': owner_dict,
        'admins': assigned_list,
        'count': 1 + len(assigned_list)
    })


@app.route('/api/projects/<project_id>/admins/', methods=['POST'])
@require_auth
def api_add_project_admin(project_id):
    """Add a project admin. Only existing project admins can add."""
    project = Project.query.get_or_404(project_id)
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project admins can add admins'}), 403
    
    data = request.get_json()
    user_id = data.get('user_id')
    username = data.get('username')
    
    if user_id is not None:
        user_id = int(user_id)
    elif username:
        u = User.query.filter_by(username=username).first()
        if not u:
            return jsonify({'error': 'User not found'}), 404
        user_id = u.id
    else:
        return jsonify({'error': 'Provide user_id or username'}), 400
    
    if user_id == project.initiator_id:
        return jsonify({'error': 'Owner is already an admin'}), 400
    
    existing = ProjectAdmin.query.filter_by(project_id=project_id, user_id=user_id).first()
    if existing:
        return jsonify({'error': 'User is already a project admin'}), 400
    
    pa = ProjectAdmin(project_id=project_id, user_id=user_id)
    db.session.add(pa)
    db.session.commit()
    
    u = User.query.get(user_id)
    return jsonify({
        'success': True,
        'admin': {
            'user_id': u.id,
            'username': u.username,
            'display_name': u.displayName or u.username,
            'is_owner': False,
            'added_at': pa.added_at.isoformat() if pa.added_at else None
        }
    })


@app.route('/api/projects/<project_id>/admins/<int:user_id>/', methods=['DELETE'])
@require_auth
def api_remove_project_admin(project_id, user_id):
    """Remove a project admin. Owner cannot be removed."""
    project = Project.query.get_or_404(project_id)
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project admins can remove admins'}), 403
    
    if user_id == project.initiator_id:
        return jsonify({'error': 'Cannot remove the layer owner'}), 400
    
    pa = ProjectAdmin.query.filter_by(project_id=project_id, user_id=user_id).first()
    if not pa:
        return jsonify({'error': 'User is not an assigned admin'}), 404
    
    db.session.delete(pa)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/users/search/', methods=['GET'])
def api_search_users():
    """Search users by username or display name (for adding project admins)."""
    q = request.args.get('q', '').strip()
    if not q or len(q) < 2:
        return jsonify({'users': [], 'count': 0})
    
    users = User.query.filter(
        db.or_(
            User.username.ilike(f'%{q}%'),
            User.displayName.ilike(f'%{q}%'),
            User.name.ilike(f'%{q}%')
        )
    ).limit(20).all()
    
    return jsonify({
        'users': [{'id': u.id, 'username': u.username, 'display_name': u.displayName or u.username} for u in users],
        'count': len(users)
    })

@app.route('/api/projects/<project_id>/members/', methods=['GET'])
def api_list_project_members(project_id):
    """List project members"""
    project = Project.query.get_or_404(project_id)
    
    members = ProjectMember.query.filter_by(project_id=project_id, status='active').all()
    
    return jsonify({
        'members': [{
            'id': m.id,
            'user_id': m.user_id,
            'username': m.user.username,
            'display_name': m.user.displayName or m.user.username,
            'role': m.role,
            'joined_at': m.joined_at.isoformat() if m.joined_at else None,
            'referred_by': m.referred_by.displayName or m.referred_by.username if m.referred_by else None
        } for m in members]
    }), 200

@app.route('/api/projects/<project_id>/join/', methods=['POST'])
@require_auth
def api_join_project(project_id):
    """Join a project (with optional referral tracking)"""
    current_user_data = get_current_user()
    if not current_user_data:
        return jsonify({'error': 'Authentication required'}), 401
    
    user = User.query.get(current_user_data['id'])
    project = Project.query.get_or_404(project_id)
    
    # Check if already a member
    existing = ProjectMember.query.filter_by(project_id=project_id, user_id=user.id).first()
    if existing and existing.status == 'active':
        return jsonify({'error': 'Already a member of this project'}), 400
    
    data = request.get_json() or {}
    referral_code = data.get('referral_code')
    
    # Find referrer if referral code provided
    referred_by_id = None
    if referral_code:
        referrer = User.query.filter_by(referral_code=referral_code).first()
        if referrer and referrer.id != user.id:  # Can't refer yourself
            referred_by_id = referrer.id
    
    # Create or reactivate membership
    if existing:
        existing.status = 'active'
        existing.joined_at = datetime.utcnow()
        existing.left_at = None
        if referred_by_id and not existing.referred_by_id:  # Only set referrer if not already set
            existing.referred_by_id = referred_by_id
            existing.referral_code = referral_code
        member = existing
    else:
        member = ProjectMember(
            project_id=project_id,
            user_id=user.id,
            referred_by_id=referred_by_id,
            referral_code=referral_code,
            role='contributor'
        )
        db.session.add(member)
    
    db.session.commit()
    
    return jsonify({
        'message': 'Successfully joined project',
        'member': {
            'id': member.id,
            'project_id': member.project_id,
            'user_id': member.user_id,
            'role': member.role,
            'joined_at': member.joined_at.isoformat() if member.joined_at else None,
            'referred_by': member.referred_by.displayName or member.referred_by.username if member.referred_by else None
        }
    }), 201

@app.route('/api/projects/<project_id>/leave/', methods=['POST'])
@require_auth
def api_leave_project(project_id):
    """Leave a project"""
    current_user_data = get_current_user()
    if not current_user_data:
        return jsonify({'error': 'Authentication required'}), 401
    
    user = User.query.get(current_user_data['id'])
    project = Project.query.get_or_404(project_id)
    
    member = ProjectMember.query.filter_by(project_id=project_id, user_id=user.id, status='active').first()
    if not member:
        return jsonify({'error': 'Not a member of this project'}), 404
    
    member.status = 'left'
    member.left_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({'message': 'Successfully left project'}), 200

@app.route('/api/user/referral-code/', methods=['GET'])
@require_auth
def api_get_referral_code():
    """Get current user's referral code"""
    current_user_data = get_current_user()
    if not current_user_data:
        return jsonify({'error': 'Authentication required'}), 401
    
    user = User.query.get(current_user_data['id'])
    referral_code = get_or_create_referral_code(user)
    
    # Count referrals
    referral_count = ProjectMember.query.filter_by(referred_by_id=user.id).count()
    
    return jsonify({
        'referral_code': referral_code,
        'referral_count': referral_count,
        'referral_url': f"{request.host_url}?ref={referral_code}"
    }), 200


# ============================================================================
# Waitlist API Endpoints
# ============================================================================

@app.route('/api/projects/<project_id>/waitlists/', methods=['GET'])
def api_list_waitlists(project_id):
    """List waitlists for a project. Only active+visible ones for non-admins."""
    project = Project.query.get_or_404(project_id)
    current_user = get_current_user()
    is_admin = current_user and is_project_admin(project, current_user)
    
    query = Waitlist.query.filter_by(project_id=project_id, archived=False)
    
    if not is_admin:
        query = query.filter_by(active=True)
        if not current_user:
            query = query.filter_by(public=True)
        else:
            is_member = ProjectMember.query.filter_by(project_id=project_id, user_id=current_user['id'], status='active').first() is not None
            if not is_member:
                query = query.filter_by(public=True)
    
    waitlists = query.order_by(Waitlist.created_at.desc()).all()
    
    result = []
    for w in waitlists:
        d = w.to_dict()
        # Expose milestone objects as array for UI; to_dict has milestones as bool
        if d.get('milestones'):
            d['milestones'] = [{'id': m.id, 'title': m.title, 'description': m.description, 'threshold': m.threshold, 'action_type': m.action_type} for m in w.milestone_list.order_by(WaitlistMilestone.threshold).all()]
        else:
            d['milestones'] = []
        if current_user:
            entry = WaitlistEntry.query.filter_by(waitlist_id=w.id, user_id=current_user['id'], left_at=None).first()
            d['my_entry'] = {'position': entry.position, 'joined_at': entry.joined_at.isoformat()} if entry else None
            if w.referrals:
                user = User.query.get(current_user['id'])
                ref_code = get_or_create_referral_code(user)
                d['referral_url'] = f"{request.host_url}projects/{project.slug}/waitlist/{w.id}/?ref={ref_code}"
        else:
            d['my_entry'] = None
            d['referral_url'] = None
        result.append(d)
    
    return jsonify({'waitlists': result, 'count': len(result)})

@app.route('/api/projects/<project_id>/waitlists/', methods=['POST'])
@require_auth
def api_create_waitlist(project_id):
    """Create waitlist - project admin only"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    project = Project.query.get_or_404(project_id)
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project admins can create waitlists'}), 403
    
    data = request.get_json() or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'error': 'Name is required'}), 400
    
    from dateutil import parser as date_parser
    start_date = None
    if data.get('start_date'):
        try:
            start_date = date_parser.parse(data['start_date'])
        except Exception:
            return jsonify({'error': 'Invalid start_date'}), 400
    if not start_date:
        start_date = datetime.utcnow()
    
    closing_date = None
    if data.get('closing_date'):
        try:
            closing_date = date_parser.parse(data['closing_date'])
        except Exception:
            pass
    
    waitlist = Waitlist(
        project_id=project_id,
        name=name,
        description=data.get('description', ''),
        image_url=data.get('image_url'),
        public=data.get('public', True),
        referrals=data.get('referrals', False),
        active=data.get('active', True),
        start_date=start_date,
        closing_date=closing_date,
        max_number=data.get('max_number'),
        milestones=data.get('milestones', False),
        show_milestones=(data.get('show_milestones') or 'all')[:20]
    )
    db.session.add(waitlist)
    db.session.commit()
    
    return jsonify({'waitlist': waitlist.to_dict()}), 201

@app.route('/api/waitlists/<int:waitlist_id>/', methods=['GET'])
def api_get_waitlist(waitlist_id):
    """Get single waitlist with milestones and user's entry"""
    waitlist = Waitlist.query.get_or_404(waitlist_id)
    project = Project.query.get_or_404(waitlist.project_id)
    current_user = get_current_user()
    is_admin = current_user and is_project_admin(project, current_user)
    
    if not waitlist.active and not is_admin:
        return jsonify({'error': 'Waitlist not found'}), 404
    
    d = waitlist.to_dict()
    d['milestones'] = [{'id': m.id, 'title': m.title, 'description': m.description, 'threshold': m.threshold, 'action_type': m.action_type} for m in waitlist.milestone_list.order_by(WaitlistMilestone.threshold).all()]
    
    if current_user:
        entry = WaitlistEntry.query.filter_by(waitlist_id=waitlist_id, user_id=current_user['id'], left_at=None).first()
        d['my_entry'] = {'position': entry.position, 'joined_at': entry.joined_at.isoformat()} if entry else None
        if waitlist.referrals:
            user = User.query.get(current_user['id'])
            ref_code = get_or_create_referral_code(user)
            d['referral_url'] = f"{request.host_url}projects/{project.slug}/waitlist/{waitlist.id}/?ref={ref_code}"
    else:
        d['my_entry'] = None
        d['referral_url'] = None
    
    return jsonify(d)

@app.route('/api/waitlists/<int:waitlist_id>/', methods=['PATCH'])
@require_auth
def api_update_waitlist(waitlist_id):
    """Update waitlist - project admin only"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    waitlist = Waitlist.query.get_or_404(waitlist_id)
    project = Project.query.get_or_404(waitlist.project_id)
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project admins can edit waitlists'}), 403
    
    data = request.get_json() or {}
    if 'name' in data and data['name']:
        waitlist.name = data['name'].strip()
    if 'description' in data:
        waitlist.description = data['description']
    if 'public' in data:
        waitlist.public = bool(data['public'])
    if 'referrals' in data:
        waitlist.referrals = bool(data['referrals'])
    if 'active' in data:
        waitlist.active = bool(data['active'])
    if 'archived' in data:
        waitlist.archived = bool(data['archived'])
    if 'max_number' in data:
        waitlist.max_number = data['max_number'] if data['max_number'] is not None else None
    if 'milestones' in data:
        waitlist.milestones = bool(data['milestones'])
    if 'show_milestones' in data:
        waitlist.show_milestones = (data['show_milestones'] or 'all')[:20]
    if 'start_date' in data and data['start_date']:
        try:
            from dateutil import parser as date_parser
            waitlist.start_date = date_parser.parse(data['start_date'])
        except Exception:
            pass
    if 'closing_date' in data:
        if data['closing_date'] is None or data['closing_date'] == '':
            waitlist.closing_date = None
        else:
            try:
                from dateutil import parser as date_parser
                waitlist.closing_date = date_parser.parse(data['closing_date'])
            except Exception:
                pass
    
    db.session.commit()
    return jsonify({'waitlist': waitlist.to_dict()}), 200

@app.route('/api/waitlists/<int:waitlist_id>/entries/', methods=['GET'])
@require_auth
def api_list_waitlist_entries(waitlist_id):
    """List entries (names) - project admin only"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    waitlist = Waitlist.query.get_or_404(waitlist_id)
    project = Project.query.get_or_404(waitlist.project_id)
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project admins can view the entry list'}), 403
    
    entries = WaitlistEntry.query.filter_by(waitlist_id=waitlist_id, left_at=None).order_by(WaitlistEntry.position).all()
    return jsonify({
        'entries': [{'id': e.id, 'user_id': e.user_id, 'username': e.user.username, 'display_name': e.user.displayName or e.user.username, 'position': e.position, 'joined_at': e.joined_at.isoformat() if e.joined_at else None, 'referred_by': e.referred_by.displayName or e.referred_by.username if e.referred_by else None} for e in entries]
    }), 200

@app.route('/api/waitlists/<int:waitlist_id>/join/', methods=['POST'])
@require_auth
def api_join_waitlist(waitlist_id):
    """Join waitlist. If referred and not on project, add to project. Optional message."""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    waitlist = Waitlist.query.get_or_404(waitlist_id)
    project = Project.query.get_or_404(waitlist.project_id)
    
    now = datetime.utcnow()
    if now < waitlist.start_date:
        return jsonify({'error': 'Waitlist has not started yet'}), 400
    if waitlist.closing_date and now >= waitlist.closing_date:
        return jsonify({'error': 'Waitlist is closed'}), 400
    count = WaitlistEntry.query.filter_by(waitlist_id=waitlist_id, left_at=None).count()
    if waitlist.max_number is not None and count >= waitlist.max_number:
        return jsonify({'error': 'Waitlist is full'}), 400
    
    existing = WaitlistEntry.query.filter_by(waitlist_id=waitlist_id, user_id=current_user['id']).first()
    if existing:
        if existing.left_at:
            existing.left_at = None
            existing.position = count + 1
            existing.message = (request.get_json() or {}).get('message', existing.message)
            existing.referred_by_id = None
            existing.referral_code = None
            db.session.commit()
        else:
            return jsonify({'error': 'Already on waitlist'}), 400
    else:
        data = request.get_json() or {}
        message = data.get('message', '')
        referral_code = data.get('referral_code')
        source = data.get('source')  # e.g., 'embed:example.com', 'direct'
        source_url = data.get('source_url')  # Full URL where signup occurred
        referred_by_id = None
        if referral_code:
            referrer = User.query.filter_by(referral_code=referral_code).first()
            if referrer and referrer.id != current_user['id']:
                referred_by_id = referrer.id
                pm = ProjectMember.query.filter_by(project_id=project.id, user_id=current_user['id'], status='active').first()
                if not pm:
                    pm = ProjectMember(project_id=project.id, user_id=current_user['id'], referred_by_id=referred_by_id, referral_code=referral_code, role='contributor')
                    db.session.add(pm)
        
        entry = WaitlistEntry(waitlist_id=waitlist_id, user_id=current_user['id'], message=message, position=count + 1, referred_by_id=referred_by_id, referral_code=referral_code, source=source, source_url=source_url)
        db.session.add(entry)
        db.session.commit()
    
    entry = WaitlistEntry.query.filter_by(waitlist_id=waitlist_id, user_id=current_user['id'], left_at=None).first()
    return jsonify({'entry': {'position': entry.position, 'joined_at': entry.joined_at.isoformat()}, 'waitlist': waitlist.to_dict()}), 201

@app.route('/api/waitlists/<int:waitlist_id>/leave/', methods=['POST'])
@require_auth
def api_leave_waitlist(waitlist_id):
    """Leave waitlist"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    waitlist = Waitlist.query.get_or_404(waitlist_id)
    entry = WaitlistEntry.query.filter_by(waitlist_id=waitlist_id, user_id=current_user['id'], left_at=None).first()
    if not entry:
        return jsonify({'error': 'Not on waitlist'}), 404
    entry.left_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'message': 'Left waitlist'}), 200

@app.route('/api/waitlists/<int:waitlist_id>/milestones/', methods=['GET'])
def api_list_waitlist_milestones(waitlist_id):
    waitlist = Waitlist.query.get_or_404(waitlist_id)
    ms = waitlist.milestone_list.order_by(WaitlistMilestone.threshold).all()
    return jsonify({'milestones': [{'id': m.id, 'title': m.title, 'description': m.description, 'threshold': m.threshold, 'action_type': m.action_type} for m in ms]}), 200

@app.route('/api/waitlists/<int:waitlist_id>/milestones/', methods=['POST'])
@require_auth
def api_create_waitlist_milestone(waitlist_id):
    """Add milestone - project admin only"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    waitlist = Waitlist.query.get_or_404(waitlist_id)
    project = Project.query.get_or_404(waitlist.project_id)
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project admins can add milestones'}), 403
    
    data = request.get_json() or {}
    title = data.get('title', '').strip()
    if not title:
        return jsonify({'error': 'Title is required'}), 400
    threshold = data.get('threshold', 0)
    try:
        threshold = int(threshold)
    except (TypeError, ValueError):
        threshold = 0
    
    m = WaitlistMilestone(waitlist_id=waitlist_id, title=title, description=data.get('description'), threshold=threshold, action_type=data.get('action_type'), action_payload=data.get('action_payload'))
    db.session.add(m)
    db.session.commit()
    return jsonify({'milestone': {'id': m.id, 'title': m.title, 'description': m.description, 'threshold': m.threshold, 'action_type': m.action_type}}), 201

@app.route('/embed/waitlist/<int:waitlist_id>/')
def embed_waitlist_widget(waitlist_id):
    """Embeddable waitlist widget - returns HTML/JS that can be embedded in external pages"""
    waitlist = Waitlist.query.get_or_404(waitlist_id)
    project = Project.query.get_or_404(waitlist.project_id)
    
    # Check if waitlist is public or active
    if not waitlist.public and not waitlist.active:
        return "Waitlist not available", 404
    
    # Get current entry count
    entry_count = WaitlistEntry.query.filter_by(waitlist_id=waitlist_id, left_at=None).count()
    
    # Determine status
    now = datetime.utcnow()
    is_upcoming = now < waitlist.start_date
    is_closed = waitlist.archived or not waitlist.active or (waitlist.closing_date and now >= waitlist.closing_date)
    is_full = waitlist.max_number and entry_count >= waitlist.max_number
    
    # Get the base URL for API calls
    base_url = request.url_root.rstrip('/')
    
    widget_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
    <style>
        .waitlist-widget {{
            max-width: 600px;
            margin: 0 auto;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
        }}
        .waitlist-header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 2rem;
            border-radius: 12px 12px 0 0;
        }}
        .waitlist-body {{
            background: white;
            padding: 2rem;
            border: 1px solid #e1e8ed;
            border-top: none;
            border-radius: 0 0 12px 12px;
        }}
        .waitlist-stats {{
            display: flex;
            justify-content: space-around;
            margin: 1.5rem 0;
            padding: 1rem;
            background: #f7f9fa;
            border-radius: 8px;
        }}
        .stat {{
            text-align: center;
        }}
        .stat-value {{
            font-size: 2rem;
            font-weight: bold;
            color: #667eea;
        }}
        .stat-label {{
            font-size: 0.875rem;
            color: #657786;
            margin-top: 0.25rem;
        }}
        .join-button {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            padding: 1rem 2rem;
            font-size: 1.125rem;
            font-weight: 600;
            border-radius: 8px;
            cursor: pointer;
            width: 100%;
            transition: transform 0.2s;
        }}
        .join-button:hover {{
            transform: translateY(-2px);
        }}
        .join-button:disabled {{
            opacity: 0.6;
            cursor: not-allowed;
        }}
        .success-message {{
            background: #00ba7c;
            color: white;
            padding: 1rem;
            border-radius: 8px;
            margin-top: 1rem;
            text-align: center;
        }}
        .error-message {{
            background: #f4212e;
            color: white;
            padding: 1rem;
            border-radius: 8px;
            margin-top: 1rem;
            text-align: center;
        }}
    </style>
</head>
<body>
    <div class="waitlist-widget">
        <div class="waitlist-header">
            <h2 style="margin: 0;">{waitlist.name}</h2>
            <p style="margin: 0.5rem 0 0 0; opacity: 0.9;">{waitlist.description or ''}</p>
            <small style="opacity: 0.8;">Part of {project.name}</small>
        </div>
        <div class="waitlist-body">
            <div class="waitlist-stats">
                <div class="stat">
                    <div class="stat-value" id="entry-count">{entry_count}</div>
                    <div class="stat-label">Members</div>
                </div>
                {f'''<div class="stat">
                    <div class="stat-value">{waitlist.max_number - entry_count}</div>
                    <div class="stat-label">Spots Left</div>
                </div>''' if waitlist.max_number else ''}
            </div>
            
            <div id="join-section">
                {f'<p class="text-center text-muted">Opens {waitlist.start_date.strftime("%B %d, %Y")}</p>' if is_upcoming else ''}
                {f'<p class="text-center text-muted">This waitlist is closed</p>' if is_closed else ''}
                {f'<p class="text-center text-muted">Waitlist is full</p>' if is_full and not is_closed else ''}
                
                {'<button class="join-button" onclick="joinWaitlist()" id="join-btn" disabled>Join Waitlist</button>' if is_upcoming or is_closed or is_full else '<button class="join-button" onclick="joinWaitlist()" id="join-btn">Join Waitlist</button>'}
            </div>
            
            <div id="message-area"></div>
            
            <div style="text-align: center; margin-top: 1.5rem;">
                <small class="text-muted">
                    Powered by <a href="{base_url}" target="_blank" style="color: #667eea;">MLGH</a>
                </small>
            </div>
        </div>
    </div>
    
    <script src="https://cdn.jsdelivr.net/npm/web3auth@latest/dist/web3auth.umd.min.js"></script>
    <script>
        const WAITLIST_ID = {waitlist_id};
        const API_BASE = '{base_url}';
        const SOURCE_URL = window.location.href;
        const SOURCE_DOMAIN = window.location.hostname;
        
        let web3auth = null;
        let currentUser = null;
        
        // Initialize Web3Auth
        async function initWeb3Auth() {{
            try {{
                web3auth = new window.Web3auth.Web3Auth({{
                    clientId: "YOUR_WEB3AUTH_CLIENT_ID", // This should be configured
                    chainConfig: {{
                        chainNamespace: "eip155",
                        chainId: "0x1",
                    }},
                }});
                await web3auth.initModal();
            }} catch (error) {{
                console.error("Web3Auth init error:", error);
            }}
        }}
        
        async function joinWaitlist() {{
            const btn = document.getElementById('join-btn');
            const messageArea = document.getElementById('message-area');
            
            btn.disabled = true;
            btn.textContent = 'Connecting...';
            
            try {{
                // Check if user is authenticated
                const response = await fetch(API_BASE + '/api/user/', {{
                    credentials: 'include'
                }});
                
                if (!response.ok) {{
                    // Need to authenticate
                    messageArea.innerHTML = '<div class="error-message">Please sign in first. Redirecting to login...</div>';
                    setTimeout(() => {{
                        window.open(API_BASE + '/login/?redirect=' + encodeURIComponent(SOURCE_URL), '_blank');
                    }}, 1500);
                    btn.disabled = false;
                    btn.textContent = 'Join Waitlist';
                    return;
                }}
                
                currentUser = await response.json();
                
                // Join the waitlist
                const joinResponse = await fetch(API_BASE + '/api/waitlists/' + WAITLIST_ID + '/join/', {{
                    method: 'POST',
                    headers: {{
                        'Content-Type': 'application/json',
                    }},
                    credentials: 'include',
                    body: JSON.stringify({{
                        source: 'embed:' + SOURCE_DOMAIN,
                        source_url: SOURCE_URL
                    }})
                }});
                
                const data = await joinResponse.json();
                
                if (joinResponse.ok) {{
                    messageArea.innerHTML = '<div class="success-message"><i class="fas fa-check-circle"></i> You\'re on the list! Position: ' + data.entry.position + '</div>';
                    document.getElementById('entry-count').textContent = parseInt(document.getElementById('entry-count').textContent) + 1;
                    btn.style.display = 'none';
                }} else {{
                    messageArea.innerHTML = '<div class="error-message">' + (data.error || 'Failed to join waitlist') + '</div>';
                    btn.disabled = false;
                    btn.textContent = 'Join Waitlist';
                }}
            }} catch (error) {{
                console.error('Error:', error);
                messageArea.innerHTML = '<div class="error-message">An error occurred. Please try again.</div>';
                btn.disabled = false;
                btn.textContent = 'Join Waitlist';
            }}
        }}
        
        // Initialize on load
        // initWeb3Auth(); // Uncomment when Web3Auth is configured
    </script>
</body>
</html>
    """
    
    return widget_html, 200, {'Content-Type': 'text/html; charset=utf-8', 'X-Frame-Options': 'ALLOWALL'}


# ============================================================================
# Workgroups API Endpoints
# ============================================================================

@app.route('/api/projects/<project_id>/workgroups/', methods=['GET'])
def api_list_workgroups(project_id):
    """List workgroups for a project"""
    status = request.args.get('status')
    approval_status = request.args.get('approval_status')
    
    query = Workgroup.query.filter_by(project_id=project_id)
    
    if status:
        query = query.filter_by(status=status)
    if approval_status:
        query = query.filter_by(approval_status=approval_status)
    
    query = query.order_by(Workgroup.created_at.desc())
    workgroups = query.all()
    
    return jsonify({'workgroups': [wg.to_dict() for wg in workgroups], 'count': len(workgroups)})

@app.route('/api/projects/<project_id>/workgroups/', methods=['POST'])
@require_auth
def api_create_workgroup(project_id):
    """Create a new workgroup"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    # Verify project exists
    project = Project.query.get_or_404(project_id)
    
    data = request.get_json()
    name = data.get('name', '').strip()
    description = data.get('description', '').strip()
    
    if not name:
        return jsonify({'error': 'Workgroup name is required'}), 400
    
    # Generate slug (ID will be auto-generated by database)
    slug = create_slug(name)
    
    # Ensure slug is unique within project
    counter = 1
    original_slug = slug
    while Workgroup.query.filter_by(project_id=project_id, slug=slug).first():
        slug = f"{original_slug}-{counter}"
        counter += 1
    
    # Generate acronym from name (for consistency with legacy groups)
    acronym = create_slug(name)
    counter = 1
    original_acronym = acronym
    while Workgroup.query.filter_by(acronym=acronym).first():
        acronym = f"{original_acronym}-{counter}"
        counter += 1
    
    # Create workgroup (ID auto-increments)
    workgroup = Workgroup(
        acronym=acronym,
        name=name,
        slug=slug,
        project_id=project_id,
        coordinator_id=current_user['id'],  # Creator becomes coordinator
        description=description,
        status='active',
        approval_status='pending'
    )
    
    db.session.add(workgroup)
    db.session.commit()
    
    return jsonify({'success': True, 'workgroup': workgroup.to_dict()}), 201

@app.route('/api/workgroups/<workgroup_id>/', methods=['GET'])
def api_get_workgroup(workgroup_id):
    """Get workgroup details"""
    workgroup = Workgroup.query.get_or_404(workgroup_id)
    d = workgroup.to_dict()
    current_user = get_current_user()
    project = Project.query.get(workgroup.project_id)
    # Coordinator, project admin, or site admin can edit
    d['can_edit'] = bool(
        current_user
        and (
            workgroup.coordinator_id == current_user['id']
            or (project and is_project_admin(project, current_user))
            or current_user.get('role') == 'admin'
        )
    )
    return jsonify(d)

@app.route('/api/workgroups/<workgroup_id>/', methods=['PATCH'])
@require_auth
def api_update_workgroup(workgroup_id):
    """Update workgroup details (coordinator, project admin, or site admin)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    workgroup = Workgroup.query.get_or_404(workgroup_id)
    project = Project.query.get(workgroup.project_id)
    
    # Check permissions: coordinator, project admin, or site admin
    can_edit = (
        workgroup.coordinator_id == current_user['id']
        or (project and is_project_admin(project, current_user))
        or current_user.get('role') == 'admin'
    )
    if not can_edit:
        return jsonify({'error': 'Permission denied'}), 403
    
    data = request.get_json()
    
    if 'name' in data and data['name']:
        workgroup.name = data['name'].strip()
    if 'description' in data:
        workgroup.description = data['description']
    if 'image_url' in data:
        workgroup.image_url = data['image_url'].strip() if data['image_url'] else None
    if 'status' in data and data['status'] in ['active', 'inactive', 'completed', 'archived']:
        old_status = workgroup.status
        workgroup.status = data['status']
        
        if old_status != workgroup.status:
            status_change = StatusChange(
                id=generate_status_change_id(),
                entity_type='workgroup',
                entity_id=workgroup_id,
                field_name='status',
                from_value=old_status,
                to_value=workgroup.status,
                changed_by_id=current_user['id']
            )
            db.session.add(status_change)
    
    workgroup.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({'success': True, 'workgroup': workgroup.to_dict()})

@app.route('/api/workgroups/<workgroup_id>/approve/', methods=['POST'])
@require_auth
def api_approve_workgroup(workgroup_id):
    """Approve or reject a workgroup (project admin/initiator or site admin)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    workgroup = Workgroup.query.get_or_404(workgroup_id)
    
    # Check permissions: site admin, editor, or project initiator
    is_site_admin = current_user.get('role') in ['admin', 'editor']
    is_project_initiator = workgroup.project and is_project_admin(workgroup.project, current_user)
    
    if not (is_site_admin or is_project_initiator):
        return jsonify({'error': 'Only project admin or site admin can approve workgroups'}), 403
    
    data = request.get_json()
    action = data.get('action')
    
    if action not in ['approve', 'reject']:
        return jsonify({'error': 'Invalid action'}), 400
    
    old_status = workgroup.approval_status
    workgroup.approval_status = 'approved' if action == 'approve' else 'rejected'
    workgroup.approved_by_id = current_user['id']
    workgroup.approved_at = datetime.utcnow()
    
    status_change = StatusChange(
        id=generate_status_change_id(),
        entity_type='workgroup',
        entity_id=workgroup_id,
        field_name='approval_status',
        from_value=old_status,
        to_value=workgroup.approval_status,
        note=data.get('note'),
        changed_by_id=current_user['id']
    )
    db.session.add(status_change)
    db.session.commit()
    
    return jsonify({'success': True, 'workgroup': workgroup.to_dict()})

@app.route('/api/workgroups/<int:workgroup_id>/chairs/', methods=['GET'])
def api_list_workgroup_chairs(workgroup_id):
    """List chairs for a workgroup"""
    workgroup = Workgroup.query.get_or_404(workgroup_id)
    
    # Query chairs using acronym (legacy field)
    from sqlalchemy import text
    chairs_query = text("""
        SELECT id, group_acronym, chair_name, approved, set_at, user_id
        FROM working_group_chair
        WHERE group_acronym = :acronym
        ORDER BY set_at DESC
    """)
    
    result = db.session.execute(chairs_query, {'acronym': workgroup.acronym})
    chairs = []
    for row in result:
        chairs.append({
            'id': row[0],
            'group_acronym': row[1],
            'chair_name': row[2],
            'approved': bool(row[3]),
            'set_at': row[4],
            'user_id': row[5]
        })
    
    return jsonify({'chairs': chairs, 'count': len(chairs)})

@app.route('/api/workgroups/<int:workgroup_id>/members/', methods=['GET'])
def api_list_workgroup_members(workgroup_id):
    """List members for a workgroup"""
    workgroup = Workgroup.query.get_or_404(workgroup_id)
    
    # Query members using acronym (legacy field)
    from sqlalchemy import text
    members_query = text("""
        SELECT id, group_acronym, user_name, joined_at, user_id
        FROM working_group_member
        WHERE group_acronym = :acronym
        ORDER BY joined_at DESC
    """)
    
    result = db.session.execute(members_query, {'acronym': workgroup.acronym})
    members = []
    for row in result:
        members.append({
            'id': row[0],
            'group_acronym': row[1],
            'user_name': row[2],
            'joined_at': row[3],
            'user_id': row[4]
        })
    
    return jsonify({'members': members, 'count': len(members)})

@app.route('/api/workgroups/<int:workgroup_id>/join/', methods=['POST'])
@require_auth
def api_join_workgroup(workgroup_id):
    """Join a workgroup as a member"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    workgroup = Workgroup.query.get_or_404(workgroup_id)
    
    # Check if workgroup is approved
    if workgroup.approval_status != 'approved':
        return jsonify({'error': 'Workgroup must be approved before joining'}), 400
    
    # Check if user is already a member
    from sqlalchemy import text
    check_query = text("""
        SELECT id FROM working_group_member
        WHERE group_acronym = :acronym AND user_id = :user_id
    """)
    
    existing = db.session.execute(check_query, {
        'acronym': workgroup.acronym,
        'user_id': current_user['id']
    }).fetchone()
    
    if existing:
        return jsonify({'error': 'You are already a member of this workgroup'}), 400
    
    # Add member
    insert_query = text("""
        INSERT INTO working_group_member (group_acronym, user_id, user_name, joined_at)
        VALUES (:acronym, :user_id, :user_name, :joined_at)
    """)
    
    db.session.execute(insert_query, {
        'acronym': workgroup.acronym,
        'user_id': current_user['id'],
        'user_name': current_user.get('displayName') or current_user.get('username'),
        'joined_at': datetime.utcnow()
    })
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Successfully joined workgroup'})

@app.route('/api/workgroups/<int:workgroup_id>/nominate-chair/', methods=['POST'])
@require_auth
def api_nominate_chair(workgroup_id):
    """Nominate yourself as a chair/coordinator for a workgroup"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    data = request.get_json()
    statement = data.get('statement', '').strip()
    
    if not statement:
        return jsonify({'error': 'Statement is required'}), 400
    
    workgroup = Workgroup.query.get_or_404(workgroup_id)
    
    # Check if workgroup is approved
    if workgroup.approval_status != 'approved':
        return jsonify({'error': 'Workgroup must be approved before nominating chairs'}), 400
    
    # Check if user is already a chair
    from sqlalchemy import text
    check_query = text("""
        SELECT id FROM working_group_chair
        WHERE group_acronym = :acronym AND user_id = :user_id
    """)
    
    existing = db.session.execute(check_query, {
        'acronym': workgroup.acronym,
        'user_id': current_user['id']
    }).fetchone()
    
    if existing:
        return jsonify({'error': 'You are already nominated as a chair for this workgroup'}), 400
    
    # Add chair nomination (pending approval)
    insert_query = text("""
        INSERT INTO working_group_chair 
        (group_acronym, user_id, chair_name, approved, set_at, statement, nominated_by_user_id, is_self_nomination)
        VALUES (:acronym, :user_id, :chair_name, :approved, :set_at, :statement, :nominated_by, :is_self)
    """)
    
    db.session.execute(insert_query, {
        'acronym': workgroup.acronym,
        'user_id': current_user['id'],
        'chair_name': current_user.get('displayName') or current_user.get('username'),
        'approved': False,  # Requires approval
        'set_at': datetime.utcnow(),
        'statement': statement,
        'nominated_by': current_user['id'],  # Self-nomination
        'is_self': True
    })
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Chair nomination submitted for approval'})

# ============================================================================
# User Profile API Endpoints
# ============================================================================

@app.route('/api/user/<username>/', methods=['GET'])
def api_get_user_profile(username):
    """Get user profile data"""
    user = User.query.filter_by(username=username).first()
    if not user:
        user = User.query.filter_by(handle=username).first()
    
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    import json
    social_links = []
    if user.social_links:
        try:
            social_links = json.loads(user.social_links)
        except:
            social_links = []
    
    return jsonify({
        'id': user.id,
        'username': user.username,
        'displayName': user.displayName,
        'handle': user.handle,
        'profileImage': user.profileImage,
        'banner_image': user.banner_image,
        'headline': user.headline,
        'bio': user.bio,
        'social_links': social_links,
        'role': user.role,
        'created_at': user.created_at.isoformat() if user.created_at else None
    })

@app.route('/api/user/profile/', methods=['PUT'])
@require_auth
def api_update_user_profile():
    """Update current user's profile"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    user = User.query.get(current_user['id'])
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    data = request.get_json()
    
    # Update allowed fields
    if 'headline' in data:
        user.headline = data['headline'][:200]  # Max 200 chars
    
    if 'bio' in data:
        user.bio = data['bio']
    
    if 'social_links' in data:
        import json
        user.social_links = json.dumps(data['social_links'])
    
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Profile updated successfully'})

@app.route('/api/admin/chair-nominations/', methods=['GET'])
@require_role('admin')
def api_admin_get_chair_nominations():
    """Get all chair nominations with full details for admin dashboard"""
    from sqlalchemy import text
    
    query = text("""
        SELECT 
            wgc.id,
            wgc.chair_name,
            wgc.approved,
            wgc.set_at,
            wgc.statement,
            wgc.is_self_nomination,
            wgc.group_acronym,
            wg.name as workgroup_name,
            wg.slug as workgroup_slug,
            p.name as project_name,
            p.slug as project_slug,
            u.id as nominee_id,
            u.username as nominee_username,
            u.profileImage as nominee_profile_image,
            nominator.id as nominator_id,
            nominator.username as nominator_username,
            nominator.displayName as nominator_name
        FROM working_group_chair wgc
        LEFT JOIN working_group wg ON wgc.group_acronym = wg.acronym
        LEFT JOIN project p ON wg.project_id = p.id
        LEFT JOIN user u ON wgc.user_id = u.id
        LEFT JOIN user nominator ON wgc.nominated_by_user_id = nominator.id
        ORDER BY wgc.approved ASC, wgc.set_at DESC
    """)
    
    results = db.session.execute(query).fetchall()
    
    nominations = []
    for row in results:
        nominations.append({
            'id': row[0],
            'chair_name': row[1],
            'approved': bool(row[2]),
            'set_at': row[3],
            'statement': row[4],
            'is_self_nomination': bool(row[5]),
            'workgroup_acronym': row[6],
            'workgroup_name': row[7],
            'workgroup_slug': row[8],
            'project_name': row[9],
            'project_slug': row[10],
            'nominee_id': row[11],
            'nominee_username': row[12],
            'nominee_profile_image': row[13],
            'nominator_id': row[14],
            'nominator_username': row[15],
            'nominator_name': row[16]
        })
    
    return jsonify({'nominations': nominations, 'count': len(nominations)})

@app.route('/api/admin/chair-nominations/<int:nomination_id>/approve/', methods=['POST'])
@require_role('admin')
def api_admin_approve_chair_nomination(nomination_id):
    """Approve a chair nomination"""
    from sqlalchemy import text
    
    # Update the chair nomination to approved
    update_query = text("""
        UPDATE working_group_chair 
        SET approved = 1
        WHERE id = :id
    """)
    
    db.session.execute(update_query, {'id': nomination_id})
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Chair nomination approved'})

@app.route('/api/admin/chair-nominations/<int:nomination_id>/reject/', methods=['POST'])
@require_role('admin')
def api_admin_reject_chair_nomination(nomination_id):
    """Reject and delete a chair nomination"""
    from sqlalchemy import text
    
    # Delete the nomination
    delete_query = text("""
        DELETE FROM working_group_chair 
        WHERE id = :id
    """)
    
    db.session.execute(delete_query, {'id': nomination_id})
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Chair nomination rejected'})

@app.route('/uploads/profile_images/<filename>')
def serve_profile_image(filename):
    """Serve uploaded profile/banner images"""
    return send_from_directory(PROFILE_IMAGE_UPLOAD_FOLDER, filename)

@app.route('/api/user/upload-image', methods=['POST'])
@app.route('/api/user/upload-image/', methods=['POST'])
@require_auth
def api_upload_profile_image():
    """Upload profile or banner image. Max 600×600px, 5MB."""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    user = User.query.get(current_user['id'])
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    image_type = request.form.get('type', 'profile')  # 'profile' or 'banner'
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    prefix = f"{image_type}_{user.id}"
    image_url, err = upload_image_600x600(
        file, PROFILE_IMAGE_UPLOAD_FOLDER, '/uploads/profile_images', filename_prefix=prefix
    )
    if err:
        return jsonify({'error': err}), 400
    
    if image_type == 'profile':
        user.profileImage = image_url
    elif image_type == 'banner':
        user.banner_image = image_url
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'message': 'Image uploaded successfully',
        'url': image_url
    })

# ============================================================================
# Guilds API Endpoints
# ============================================================================

@app.route('/api/guilds/', methods=['GET'])
def api_list_guilds():
    """List all guilds"""
    status = request.args.get('status')
    
    query = Guild.query.filter_by(status=status) if status else Guild.query
    query = query.order_by(Guild.created_at.desc())
    guilds = query.all()
    
    return jsonify({'guilds': [g.to_dict() for g in guilds], 'count': len(guilds)})

@app.route('/api/guilds/', methods=['POST'])
@require_auth
def api_create_guild():
    """Create a new guild (instant registration)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    data = request.get_json()
    name = data.get('name', '').strip()
    description = data.get('description', '').strip()
    
    if not name:
        return jsonify({'error': 'Guild name is required'}), 400
    
    # Check if guild name already exists
    existing = Guild.query.filter_by(name=name).first()
    if existing:
        return jsonify({'error': 'Guild name already exists'}), 400
    
    # Generate ID and slug
    guild_id = generate_guild_id()
    slug = create_slug(name)
    
    counter = 1
    original_slug = slug
    while Guild.query.filter_by(slug=slug).first():
        slug = f"{original_slug}-{counter}"
        counter += 1
    
    # Create guild
    guild = Guild(
        id=guild_id,
        name=name,
        slug=slug,
        initiator_id=current_user['id'],
        description=description,
        status='active'
    )
    
    # Add initiator as admin member
    membership = GuildMembership(
        guild_id=guild_id,
        user_id=current_user['id'],
        role='initiator'
    )
    
    db.session.add(guild)
    db.session.add(membership)
    db.session.commit()
    
    return jsonify({'success': True, 'guild': guild.to_dict()}), 201

@app.route('/api/guilds/<guild_id>/', methods=['GET'])
def api_get_guild(guild_id):
    """Get guild details with members"""
    guild = Guild.query.get_or_404(guild_id)
    
    # Get members
    memberships = GuildMembership.query.filter_by(guild_id=guild_id).all()
    members = []
    for m in memberships:
        if m.user:
            members.append({
                'user_id': m.user_id,
                'username': m.user.username,
                'display_name': m.user.displayName or m.user.username,
                'role': m.role,
                'joined_at': m.joined_at.isoformat() if m.joined_at else None
            })
    
    guild_dict = guild.to_dict()
    guild_dict['members'] = members
    guild_dict['member_count'] = len(members)
    
    return jsonify(guild_dict)

@app.route('/api/guilds/<guild_id>/', methods=['PATCH'])
@require_auth
def api_update_guild(guild_id):
    """Update guild details (initiator/admin only)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    guild = Guild.query.get_or_404(guild_id)
    membership = GuildMembership.query.filter_by(guild_id=guild_id, user_id=current_user['id']).first()
    if not membership or membership.role not in ['initiator', 'admin']:
        return jsonify({'error': 'Only guild admins can edit'}), 403
    
    data = request.get_json()
    if 'name' in data and data['name']:
        name = data['name'].strip()
        if name != guild.name and Guild.query.filter_by(name=name).first():
            return jsonify({'error': 'A guild with this name already exists'}), 400
        guild.name = name
        guild.slug = create_slug(name)
    if 'description' in data:
        guild.description = data['description']
    if 'image_url' in data:
        guild.image_url = data['image_url'].strip() if data['image_url'] else None
    if 'status' in data and data['status'] in ['active', 'archived']:
        guild.status = data['status']
    
    guild.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({'success': True, 'guild': guild.to_dict()})

@app.route('/api/guilds/<guild_id>/invite/', methods=['POST'])
@require_auth
def api_invite_to_guild(guild_id):
    """Invite user to guild (admin/initiator only)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    guild = Guild.query.get_or_404(guild_id)
    
    # Check if user is admin or initiator
    membership = GuildMembership.query.filter_by(
        guild_id=guild_id,
        user_id=current_user['id']
    ).first()
    
    if not membership or membership.role not in ['initiator', 'admin']:
        return jsonify({'error': 'Only guild admins can invite members'}), 403
    
    data = request.get_json()
    email = data.get('email', '').strip().lower()
    
    if not email:
        return jsonify({'error': 'Email is required'}), 400
    
    # Check if user already a member
    invitee = User.query.filter_by(email=email).first()
    if invitee:
        existing_membership = GuildMembership.query.filter_by(
            guild_id=guild_id,
            user_id=invitee.id
        ).first()
        if existing_membership:
            return jsonify({'error': 'User is already a member'}), 400
    
    # Create invitation
    from datetime import timedelta
    invitation_id = generate_guild_invitation_id()
    token = generate_invitation_token()
    
    invitation = GuildInvitation(
        id=invitation_id,
        guild_id=guild_id,
        inviter_id=current_user['id'],
        invitee_email=email,
        invitee_id=invitee.id if invitee else None,
        token=token,
        expires_at=datetime.utcnow() + timedelta(days=7)
    )
    
    db.session.add(invitation)
    db.session.commit()
    
    # TODO: Send email with invitation link
    invitation_link = f"https://rfc.themetalayer.org/guilds/invite/{token}/"
    
    return jsonify({
        'success': True,
        'invitation_id': invitation_id,
        'invitation_link': invitation_link,
        'expires_at': invitation.expires_at.isoformat()
    }), 201

# ============================================================================
# Clusters API
# ============================================================================

@app.route('/api/projects/<project_id>/clusters/', methods=['GET'])
def api_list_clusters(project_id):
    """List clusters for a project. By default excludes archived (deleted) clusters. ?include_roles=1 adds roles per cluster."""
    project = Project.query.get_or_404(project_id)
    
    # Filter by status if provided; default to excluding archived so "deleted" clusters disappear
    status = request.args.get('status')
    include_roles = request.args.get('include_roles', '').lower() in ('1', 'true', 'yes')
    
    query = Cluster.query.filter_by(project_id=project_id)
    if status:
        query = query.filter_by(status=status)
    else:
        query = query.filter(Cluster.status != 'archived')
    
    clusters = query.order_by(Cluster.order, Cluster.name).all()
    
    result = []
    for c in clusters:
        d = c.to_dict()
        if include_roles:
            roles = Role.query.filter_by(cluster_id=c.id).order_by(Role.order, Role.title_guild).all()
            d['roles'] = [r.to_dict() for r in roles]
        result.append(d)
    
    return jsonify({'clusters': result, 'count': len(result)})

@app.route('/api/projects/<project_id>/clusters/', methods=['POST'])
@require_auth
def api_create_cluster(project_id):
    """Create a cluster in a project"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    project = Project.query.get_or_404(project_id)
    
    # Check permissions (project initiator or admin)
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project initiator or admins can create clusters'}), 403
    
    data = request.get_json()
    name = data.get('name', '').strip()
    description = data.get('description', '').strip()
    order = data.get('order', 0)
    
    if not name:
        return jsonify({'error': 'Cluster name is required'}), 400
    
    # Generate slug
    cluster_slug = create_slug(name)
    
    # Check for slug collision within project
    existing = Cluster.query.filter_by(project_id=project_id, cluster_slug=cluster_slug).first()
    if existing:
        # Add number suffix
        counter = 1
        while existing:
            cluster_slug = f"{create_slug(name)}-{counter}"
            existing = Cluster.query.filter_by(project_id=project_id, cluster_slug=cluster_slug).first()
            counter += 1
    
    cluster_id = generate_cluster_id()
    cluster = Cluster(
        id=cluster_id,
        project_id=project_id,
        cluster_slug=cluster_slug,
        name=name,
        description=description if description else None,
        order=order,
        created_by_id=current_user['id']
    )
    
    db.session.add(cluster)
    db.session.commit()
    
    return jsonify({'success': True, 'cluster': cluster.to_dict()}), 201

@app.route('/api/clusters/<cluster_id>/', methods=['GET'])
def api_get_cluster(cluster_id):
    """Get cluster details"""
    cluster = Cluster.query.get_or_404(cluster_id)
    
    cluster_dict = cluster.to_dict()
    
    # Add roles count
    roles_count = Role.query.filter_by(cluster_id=cluster_id).count()
    cluster_dict['roles_count'] = roles_count
    
    # Return under 'cluster' key for consistent frontend consumption
    return jsonify({'cluster': cluster_dict})

@app.route('/api/clusters/<cluster_id>/', methods=['PATCH'])
@require_auth
def api_update_cluster(cluster_id):
    """Update cluster (project initiator or admin only)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    cluster = Cluster.query.get_or_404(cluster_id)
    project = Project.query.get_or_404(cluster.project_id)
    
    # Check permissions
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project initiator or admins can update clusters'}), 403
    
    data = request.get_json()
    
    # Update fields
    if 'name' in data:
        name = data['name'].strip()
        if name:
            cluster.name = name
            # Update slug if name changed
            new_slug = create_slug(name)
            if new_slug != cluster.cluster_slug:
                # Check for collision
                existing = Cluster.query.filter_by(
                    project_id=cluster.project_id,
                    cluster_slug=new_slug
                ).filter(Cluster.id != cluster_id).first()
                if not existing:
                    cluster.cluster_slug = new_slug
    
    if 'description' in data:
        cluster.description = data['description'].strip() if data['description'] else None
    
    if 'order' in data:
        cluster.order = data['order']
    
    if 'status' in data and data['status'] in ['active', 'archived']:
        old_status = cluster.status
        cluster.status = data['status']
        
        # Record status change
        if old_status != cluster.status:
            status_change = StatusChange(
                id=generate_status_change_id(),
                entity_type='cluster',
                entity_id=cluster_id,
                field_name='status',
                from_value=old_status,
                to_value=cluster.status,
                changed_by_id=current_user['id']
            )
            db.session.add(status_change)
    
    db.session.commit()
    
    return jsonify({'success': True, 'cluster': cluster.to_dict()})

@app.route('/api/clusters/<cluster_id>/', methods=['DELETE'])
@require_auth
def api_delete_cluster(cluster_id):
    """Archive cluster (project initiator or admin only)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    cluster = Cluster.query.get_or_404(cluster_id)
    project = Project.query.get_or_404(cluster.project_id)
    
    # Check permissions
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project initiator or admins can archive clusters'}), 403
    
    old_status = cluster.status
    cluster.status = 'archived'
    
    # Record status change
    status_change = StatusChange(
        id=generate_status_change_id(),
        entity_type='cluster',
        entity_id=cluster_id,
        field_name='status',
        from_value=old_status,
        to_value='archived',
        changed_by_id=current_user['id']
    )
    db.session.add(status_change)
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Cluster archived'})

@app.route('/api/clusters/<cluster_id>/roles/', methods=['GET'])
def api_list_cluster_roles(cluster_id):
    """List roles in a cluster"""
    cluster = Cluster.query.get_or_404(cluster_id)
    
    # Filter by status if provided
    status = request.args.get('status')
    
    query = Role.query.filter_by(cluster_id=cluster_id)
    if status:
        query = query.filter_by(status=status)
    
    roles = query.order_by(Role.order, Role.title_guild).all()
    
    return jsonify({'roles': [r.to_dict() for r in roles], 'count': len(roles)})

# ============================================================================
# Roles API
# ============================================================================

@app.route('/api/projects/<project_id>/roles/', methods=['GET'])
def api_list_roles(project_id):
    """List roles for a project"""
    project = Project.query.get_or_404(project_id)
    
    # Filter by status if provided
    status = request.args.get('status')
    cluster_id = request.args.get('cluster_id')
    public_only = request.args.get('public_only', 'false').lower() == 'true'
    
    query = Role.query.filter_by(project_id=project_id)
    
    if status:
        query = query.filter_by(status=status)
    
    if cluster_id:
        query = query.filter_by(cluster_id=cluster_id)
    
    if public_only:
        query = query.filter_by(public_visible=True)
    
    roles = query.order_by(Role.order, Role.title_guild).all()
    
    return jsonify({'roles': [r.to_dict() for r in roles], 'count': len(roles)})

@app.route('/api/projects/<project_id>/roles/', methods=['POST'])
@require_auth
def api_create_role(project_id):
    """Create a role in a project"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    project = Project.query.get_or_404(project_id)
    
    data = request.get_json()
    title_guild = data.get('title_guild', '').strip()
    title_operational = data.get('title_operational', '').strip()
    description = data.get('description', '').strip()
    cluster_id = data.get('cluster_id')
    image_url = data.get('image_url', '').strip()
    order = data.get('order', 0)
    
    if not title_guild:
        return jsonify({'error': 'Guild title is required'}), 400
    
    if not description:
        return jsonify({'error': 'Description is required'}), 400
    
    # Validate cluster if provided
    if cluster_id:
        cluster = Cluster.query.filter_by(id=cluster_id, project_id=project_id).first()
        if not cluster:
            return jsonify({'error': 'Invalid cluster for this project'}), 400
    
    # Generate slug from guild title
    role_slug = create_slug(title_guild)
    
    # Check for slug collision within project
    existing = Role.query.filter_by(project_id=project_id, role_slug=role_slug).first()
    if existing:
        counter = 1
        while existing:
            role_slug = f"{create_slug(title_guild)}-{counter}"
            existing = Role.query.filter_by(project_id=project_id, role_slug=role_slug).first()
            counter += 1
    
    role_id = generate_role_id()
    role = Role(
        id=role_id,
        project_id=project_id,
        role_slug=role_slug,
        title_guild=title_guild,
        title_operational=title_operational if title_operational else None,
        description=description,
        image_url=image_url if image_url else None,
        cluster_id=cluster_id if cluster_id else None,
        order=order,
        status='draft',
        created_by_id=current_user['id']
    )
    
    # Set configuration options if provided
    if 'claim_requires_approval' in data:
        role.claim_requires_approval = data['claim_requires_approval']
    if 'badge_enabled' in data:
        role.badge_enabled = data['badge_enabled']
    if 'badge_requires_approval' in data:
        role.badge_requires_approval = data['badge_requires_approval']
    if 'public_visible' in data:
        role.public_visible = data['public_visible']
    
    db.session.add(role)
    db.session.commit()
    
    return jsonify({'success': True, 'role': role.to_dict()}), 201

@app.route('/api/projects/<project_id>/roles/import/', methods=['POST'])
@require_auth
def api_import_roles(project_id):
    """Import roles from JSON"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    project = Project.query.get_or_404(project_id)
    
    # Check permissions (project initiator or admin)
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project initiator or admins can import roles'}), 403
    
    data = request.get_json()
    roles_data = data.get('roles', [])
    
    if not roles_data or not isinstance(roles_data, list):
        return jsonify({'error': 'Invalid roles data'}), 400
    
    imported_roles = []
    errors = []
    
    for idx, role_data in enumerate(roles_data):
        try:
            title_guild = role_data.get('title_guild', '').strip()
            description = role_data.get('description', '').strip()
            
            if not title_guild or not description:
                errors.append(f"Role {idx}: Missing title_guild or description")
                continue
            
            # Generate slug
            role_slug = create_slug(title_guild)
            existing = Role.query.filter_by(project_id=project_id, role_slug=role_slug).first()
            if existing:
                counter = 1
                while existing:
                    role_slug = f"{create_slug(title_guild)}-{counter}"
                    existing = Role.query.filter_by(project_id=project_id, role_slug=role_slug).first()
                    counter += 1
            
            role_id = generate_role_id()
            role = Role(
                id=role_id,
                project_id=project_id,
                role_slug=role_slug,
                title_guild=title_guild,
                title_operational=role_data.get('title_operational'),
                description=description,
                image_url=role_data.get('image_url'),
                cluster_id=role_data.get('cluster_id'),
                order=role_data.get('order', 0),
                status='draft',
                created_by_id=current_user['id']
            )
            
            db.session.add(role)
            imported_roles.append(role)
            
        except Exception as e:
            errors.append(f"Role {idx}: {str(e)}")
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'imported_count': len(imported_roles),
        'roles': [r.to_dict() for r in imported_roles],
        'errors': errors
    }), 201

@app.route('/api/roles/<role_id>/', methods=['GET'])
def api_get_role(role_id):
    """Get role details"""
    role = Role.query.get_or_404(role_id)
    
    role_dict = role.to_dict()
    
    # Add claims count
    claims_count = Claim.query.filter_by(role_id=role_id, status='active').count()
    role_dict['active_claims_count'] = claims_count
    
    # Add cluster name for role detail display
    if role.cluster_id and role.cluster:
        role_dict['cluster_name'] = role.cluster.name
    else:
        role_dict['cluster_name'] = None
    
    # Add can_edit for project admin (role detail page Edit button)
    current_user = get_current_user()
    project = Project.query.get(role.project_id)
    role_dict['can_edit'] = bool(project and current_user and is_project_admin(project, current_user))
    
    return jsonify(role_dict)

@app.route('/api/roles/<role_id>/', methods=['PATCH'])
@require_auth
def api_update_role(role_id):
    """Update role (project initiator or admin only)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    role = Role.query.get_or_404(role_id)
    project = Project.query.get_or_404(role.project_id)
    
    # Check permissions
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project initiator or admins can update roles'}), 403
    
    data = request.get_json()
    
    # Update fields
    if 'title_guild' in data:
        title = data['title_guild'].strip()
        if title:
            role.title_guild = title
    
    if 'title_operational' in data:
        role.title_operational = data['title_operational'].strip() if data['title_operational'] else None
    
    if 'description' in data:
        desc = data['description'].strip()
        if desc:
            role.description = desc
    
    if 'image_url' in data:
        role.image_url = data['image_url'].strip() if data['image_url'] else None
    
    if 'cluster_id' in data:
        if data['cluster_id']:
            cluster = Cluster.query.filter_by(id=data['cluster_id'], project_id=role.project_id).first()
            if cluster:
                role.cluster_id = data['cluster_id']
        else:
            role.cluster_id = None
    
    if 'order' in data:
        role.order = data['order']
    
    if 'public_visible' in data:
        role.public_visible = data['public_visible']
    
    if 'claim_requires_approval' in data:
        role.claim_requires_approval = data['claim_requires_approval']
    
    if 'badge_enabled' in data:
        role.badge_enabled = data['badge_enabled']
    
    if 'badge_requires_approval' in data:
        role.badge_requires_approval = data['badge_requires_approval']
    
    db.session.commit()
    
    return jsonify({'success': True, 'role': role.to_dict()})

@app.route('/api/roles/<role_id>/approve/', methods=['POST'])
@require_auth
def api_approve_role(role_id):
    """Approve role (admin only)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    if not current_user.get('is_admin'):
        return jsonify({'error': 'Admin access required'}), 403
    
    role = Role.query.get_or_404(role_id)
    
    data = request.get_json()
    approve = data.get('approve', True)
    
    old_status = role.status
    
    if approve:
        role.status = 'approved'
        role.approved_by_id = current_user['id']
        role.approved_at = datetime.utcnow()
    else:
        role.status = 'draft'
        role.approved_by_id = None
        role.approved_at = None
    
    # Record status change
    status_change = StatusChange(
        id=generate_status_change_id(),
        entity_type='role',
        entity_id=role_id,
        field_name='status',
        from_value=old_status,
        to_value=role.status,
        note=data.get('note'),
        changed_by_id=current_user['id']
    )
    db.session.add(status_change)
    db.session.commit()
    
    return jsonify({'success': True, 'role': role.to_dict()})

@app.route('/api/roles/<role_id>/status/', methods=['POST'])
@require_auth
def api_change_role_status(role_id):
    """Change role status (project initiator or admin only)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    role = Role.query.get_or_404(role_id)
    project = Project.query.get_or_404(role.project_id)
    
    # Check permissions
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project initiator or admins can change role status'}), 403
    
    data = request.get_json()
    new_status = data.get('status')
    
    if new_status not in ['draft', 'approved', 'deprecated', 'archived']:
        return jsonify({'error': 'Invalid status'}), 400
    
    old_status = role.status
    role.status = new_status
    
    # Record status change
    status_change = StatusChange(
        id=generate_status_change_id(),
        entity_type='role',
        entity_id=role_id,
        field_name='status',
        from_value=old_status,
        to_value=new_status,
        note=data.get('note'),
        changed_by_id=current_user['id']
    )
    db.session.add(status_change)
    db.session.commit()
    
    return jsonify({'success': True, 'role': role.to_dict()})

@app.route('/api/roles/<role_id>/claims/', methods=['GET'])
def api_list_role_claims(role_id):
    """List claims for a role"""
    role = Role.query.get_or_404(role_id)
    
    # Filter by status if provided
    status = request.args.get('status')
    
    query = Claim.query.filter_by(role_id=role_id)
    if status:
        query = query.filter_by(status=status)
    
    claims = query.order_by(Claim.created_at.desc()).all()
    
    # Enrich with claimant name and username for profile link
    result = []
    for c in claims:
        d = c.to_dict()
        if c.claimant:
            d['claimant_name'] = c.claimant.displayName or c.claimant.username or c.claimant.name
            d['claimant_username'] = c.claimant.username
        else:
            d['claimant_name'] = None
            d['claimant_username'] = None
        result.append(d)
    
    return jsonify({'claims': result, 'count': len(result)})

# ============================================================================
# Claims API
# ============================================================================

@app.route('/api/projects/<project_id>/claims/', methods=['GET'])
def api_list_claims(project_id):
    """List claims for a project"""
    project = Project.query.get_or_404(project_id)
    
    # Filter by status if provided
    status = request.args.get('status')
    role_id = request.args.get('role_id')
    claimant_id = request.args.get('claimant_id')
    
    query = Claim.query.filter_by(project_id=project_id)
    
    if status:
        query = query.filter_by(status=status)
    
    if role_id:
        query = query.filter_by(role_id=role_id)
    
    if claimant_id:
        query = query.filter_by(claimant_id=int(claimant_id))
    
    claims = query.order_by(Claim.created_at.desc()).all()
    
    # Enrich with role name/slug and claimant name/username for project claims tab
    result = []
    for c in claims:
        d = c.to_dict()
        if c.role:
            d['role_name'] = c.role.title_operational or c.role.title_guild
            d['role_slug'] = c.role.role_slug
        else:
            d['role_name'] = None
            d['role_slug'] = None
        if c.claimant:
            d['claimant_name'] = c.claimant.displayName or c.claimant.username or getattr(c.claimant, 'name', None)
            d['claimant_username'] = c.claimant.username
        else:
            d['claimant_name'] = None
            d['claimant_username'] = None
        result.append(d)
    
    return jsonify({'claims': result, 'count': len(result)})

@app.route('/api/roles/<role_id>/claims/', methods=['POST'])
@require_auth
def api_create_claim(role_id):
    """Create a claim for a role"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    role = Role.query.get_or_404(role_id)
    
    # Check if role is approved
    if role.status != 'approved':
        return jsonify({'error': 'Can only claim approved roles'}), 400
    
    data = request.get_json()
    intent = data.get('intent', '').strip()
    evidence_links = data.get('evidence_links', [])
    # Accept term in months (default 3) or legacy days
    term_duration_months = data.get('term_duration_months')
    if term_duration_months is not None:
        months = int(term_duration_months) if term_duration_months else 3
        months = max(1, min(12, months))
        term_duration_days = {1: 30, 3: 90, 6: 182, 12: 365}.get(months, months * 30)
    else:
        term_duration_days = data.get('term_duration_days')
    
    # Check if user already has an active claim for this role
    existing_claim = Claim.query.filter_by(
        role_id=role_id,
        claimant_id=current_user['id'],
        status='active'
    ).first()
    
    if existing_claim:
        return jsonify({'error': 'You already have an active claim for this role'}), 400
    
    claim_id = generate_claim_id()
    
    # Determine if approval is required
    approval_required = role.claim_requires_approval
    initial_status = 'pending_approval' if approval_required else 'active'
    
    claim = Claim(
        id=claim_id,
        project_id=role.project_id,
        role_id=role_id,
        claimant_id=current_user['id'],
        intent=intent if intent else None,
        evidence_links=evidence_links,
        status=initial_status,
        approval_required=approval_required
    )
    
    # Set term if provided
    if term_duration_days:
        claim.term_start = datetime.utcnow().date()
        claim.term_duration_days = term_duration_days
        from datetime import timedelta
        claim.term_end = claim.term_start + timedelta(days=term_duration_days)
        claim.term_status = 'active'
    
    db.session.add(claim)
    db.session.commit()
    
    return jsonify({'success': True, 'claim': claim.to_dict()}), 201

@app.route('/api/claims/<claim_id>/', methods=['GET'])
def api_get_claim(claim_id):
    """Get claim details"""
    claim = Claim.query.get_or_404(claim_id)
    
    claim_dict = claim.to_dict()
    
    # Add role info
    role = Role.query.get(claim.role_id)
    if role:
        claim_dict['role'] = {
            'id': role.id,
            'title_guild': role.title_guild,
            'title_operational': role.title_operational
        }
    
    # Add claimant info
    claimant = User.query.get(claim.claimant_id)
    if claimant:
        claim_dict['claimant'] = {
            'id': claimant.id,
            'username': claimant.username,
            'name': claimant.name
        }
    
    return jsonify(claim_dict)

@app.route('/api/claims/<claim_id>/', methods=['PATCH'])
@require_auth
def api_update_claim(claim_id):
    """Update claim (claimant only)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    claim = Claim.query.get_or_404(claim_id)
    
    # Check permissions (only claimant can update)
    if claim.claimant_id != current_user['id']:
        return jsonify({'error': 'Only the claimant can update this claim'}), 403
    
    data = request.get_json()
    
    # Update fields
    if 'intent' in data:
        claim.intent = data['intent'].strip() if data['intent'] else None
    
    if 'evidence_links' in data:
        claim.evidence_links = data['evidence_links']
    
    db.session.commit()
    
    return jsonify({'success': True, 'claim': claim.to_dict()})

@app.route('/api/claims/<claim_id>/approve/', methods=['POST'])
@require_auth
def api_approve_claim(claim_id):
    """Approve claim (project initiator or admin only)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    claim = Claim.query.get_or_404(claim_id)
    project = Project.query.get_or_404(claim.project_id)
    
    # Check permissions
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project initiator or admins can approve claims'}), 403
    
    data = request.get_json()
    approve = data.get('approve', True)
    
    old_status = claim.status
    
    if approve:
        claim.status = 'active'
        claim.approved_by_id = current_user['id']
        claim.approved_at = datetime.utcnow()
    else:
        claim.status = 'revoked'
    
    # Record status change
    status_change = StatusChange(
        id=generate_status_change_id(),
        entity_type='claim',
        entity_id=claim_id,
        field_name='status',
        from_value=old_status,
        to_value=claim.status,
        note=data.get('note'),
        changed_by_id=current_user['id']
    )
    db.session.add(status_change)
    db.session.commit()
    
    return jsonify({'success': True, 'claim': claim.to_dict()})

@app.route('/api/claims/<claim_id>/status/', methods=['POST'])
@require_auth
def api_change_claim_status(claim_id):
    """Change claim status (claimant, project initiator, or admin)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    claim = Claim.query.get_or_404(claim_id)
    project = Project.query.get_or_404(claim.project_id)
    
    # Check permissions
    is_claimant = claim.claimant_id == current_user['id']
    is_padmin = is_project_admin(project, current_user)
    
    if not (is_claimant or is_padmin):
        return jsonify({'error': 'Permission denied'}), 403
    
    data = request.get_json()
    new_status = data.get('status')
    
    # Validate status
    valid_statuses = ['active', 'pending_approval', 'paused', 'expired', 'revoked']
    if new_status not in valid_statuses:
        return jsonify({'error': 'Invalid status'}), 400
    
    # Claimants can only pause/unpause their own claims
    if is_claimant and not is_padmin:
        if new_status not in ['paused', 'active']:
            return jsonify({'error': 'You can only pause or reactivate your claim'}), 403
    
    old_status = claim.status
    claim.status = new_status
    
    # Record status change
    status_change = StatusChange(
        id=generate_status_change_id(),
        entity_type='claim',
        entity_id=claim_id,
        field_name='status',
        from_value=old_status,
        to_value=new_status,
        note=data.get('note'),
        changed_by_id=current_user['id']
    )
    db.session.add(status_change)
    db.session.commit()
    
    return jsonify({'success': True, 'claim': claim.to_dict()})

# ============================================================================
# Badges API
# ============================================================================

@app.route('/api/projects/<project_id>/badges/', methods=['GET'])
def api_list_badges(project_id):
    """List badges for a project"""
    project = Project.query.get_or_404(project_id)
    
    # Filter by status if provided
    status = request.args.get('status')
    claim_id = request.args.get('claim_id')
    claimant_id = request.args.get('claimant_id')
    
    query = Badge.query.filter_by(project_id=project_id)
    
    if status:
        query = query.filter_by(status=status)
    
    if claim_id:
        query = query.filter_by(claim_id=claim_id)
    
    if claimant_id:
        query = query.filter_by(claimant_id=int(claimant_id))
    
    badges = query.order_by(Badge.created_at.desc()).all()
    
    return jsonify({'badges': [b.to_dict() for b in badges], 'count': len(badges)})

@app.route('/api/claims/<claim_id>/badges/', methods=['GET'])
def api_list_claim_badges(claim_id):
    """List badges for a claim"""
    claim = Claim.query.get_or_404(claim_id)
    
    badges = Badge.query.filter_by(claim_id=claim_id).order_by(Badge.created_at.desc()).all()
    
    return jsonify({'badges': [b.to_dict() for b in badges], 'count': len(badges)})

@app.route('/api/claims/<claim_id>/badges/', methods=['POST'])
@require_auth
def api_request_badge(claim_id):
    """Request a badge for a claim"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    claim = Claim.query.get_or_404(claim_id)
    role = Role.query.get_or_404(claim.role_id)
    
    # Check if badges are enabled for this role
    if not role.badge_enabled:
        return jsonify({'error': 'Badges are not enabled for this role'}), 400
    
    # Check if claim is active
    if claim.status != 'active':
        return jsonify({'error': 'Can only request badges for active claims'}), 400
    
    # Check if requester is claimant or project admin
    project = Project.query.get_or_404(claim.project_id)
    is_claimant = claim.claimant_id == current_user['id']
    is_padmin = is_project_admin(project, current_user)
    
    if not (is_claimant or is_padmin):
        return jsonify({'error': 'Only the claimant or project admins can request badges'}), 403
    
    data = request.get_json()
    badge_type = data.get('badge_type', 'role_badge')
    evidence_links = data.get('evidence_links', [])
    custody_mode = data.get('custody_mode', 'user_wallet')
    btc_taproot_address = data.get('btc_taproot_address', '').strip()
    
    # Validate badge type
    if badge_type not in ['role_badge', 'founding_wave_badge', 'term_renewal_marker']:
        return jsonify({'error': 'Invalid badge type'}), 400
    
    # Validate custody mode
    if custody_mode not in ['user_wallet', 'overweb_treasury']:
        return jsonify({'error': 'Invalid custody mode'}), 400
    
    # If user_wallet, require BTC address
    if custody_mode == 'user_wallet' and not btc_taproot_address:
        return jsonify({'error': 'BTC Taproot address is required for user wallet custody'}), 400
    
    badge_id = generate_badge_id()
    
    # Determine initial status based on role configuration
    initial_status = 'requested' if role.badge_requires_approval else 'approved'
    
    badge = Badge(
        id=badge_id,
        project_id=claim.project_id,
        claim_id=claim_id,
        role_id=claim.role_id,
        claimant_id=claim.claimant_id,
        requested_by_id=current_user['id'],
        badge_type=badge_type,
        status=initial_status,
        evidence_links=evidence_links,
        custody_mode=custody_mode,
        btc_taproot_address=btc_taproot_address if btc_taproot_address else None
    )
    
    # If auto-approved, set approval fields
    if initial_status == 'approved':
        badge.approved_by_id = current_user['id']
        badge.approved_at = datetime.utcnow()
    
    db.session.add(badge)
    db.session.commit()
    
    return jsonify({'success': True, 'badge': badge.to_dict()}), 201

@app.route('/api/badges/<badge_id>/', methods=['GET'])
def api_get_badge(badge_id):
    """Get badge details"""
    badge = Badge.query.get_or_404(badge_id)
    
    badge_dict = badge.to_dict()
    
    # Add role info
    role = Role.query.get(badge.role_id)
    if role:
        badge_dict['role'] = {
            'id': role.id,
            'title_guild': role.title_guild,
            'title_operational': role.title_operational
        }
    
    # Add claimant info
    claimant = User.query.get(badge.claimant_id)
    if claimant:
        badge_dict['claimant'] = {
            'id': claimant.id,
            'username': claimant.username,
            'name': claimant.name
        }
    
    return jsonify(badge_dict)

@app.route('/api/badges/<badge_id>/approve/', methods=['POST'])
@require_auth
def api_approve_badge(badge_id):
    """Approve badge (project initiator or admin only)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    badge = Badge.query.get_or_404(badge_id)
    project = Project.query.get_or_404(badge.project_id)
    
    # Check permissions
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project initiator or admins can approve badges'}), 403
    
    data = request.get_json()
    approve = data.get('approve', True)
    approval_note = data.get('approval_note', '').strip()
    
    old_status = badge.status
    
    if approve:
        badge.status = 'approved'
        badge.approved_by_id = current_user['id']
        badge.approved_at = datetime.utcnow()
        badge.approval_note = approval_note if approval_note else None
    else:
        badge.status = 'denied'
        badge.approval_note = approval_note if approval_note else None
    
    # Record status change
    status_change = StatusChange(
        id=generate_status_change_id(),
        entity_type='badge',
        entity_id=badge_id,
        field_name='status',
        from_value=old_status,
        to_value=badge.status,
        note=approval_note,
        changed_by_id=current_user['id']
    )
    db.session.add(status_change)
    db.session.commit()
    
    return jsonify({'success': True, 'badge': badge.to_dict()})

@app.route('/api/badges/<badge_id>/issue/', methods=['POST'])
@require_auth
def api_issue_badge(badge_id):
    """Issue badge with inscription details (admin only)"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    if not current_user.get('is_admin'):
        return jsonify({'error': 'Admin access required'}), 403
    
    badge = Badge.query.get_or_404(badge_id)
    
    # Check if badge is approved
    if badge.status != 'approved':
        return jsonify({'error': 'Badge must be approved before issuance'}), 400
    
    data = request.get_json()
    inscription_id = data.get('inscription_id', '').strip()
    tx_ref = data.get('tx_ref', '').strip()
    chain = data.get('chain', 'bitcoin').strip()
    
    if not inscription_id:
        return jsonify({'error': 'Inscription ID is required'}), 400
    
    old_status = badge.status
    badge.status = 'issued'
    badge.issuance_kind = 'ordinal'
    badge.inscription_id = inscription_id
    badge.tx_ref = tx_ref if tx_ref else None
    badge.chain = chain
    
    # Record status change
    status_change = StatusChange(
        id=generate_status_change_id(),
        entity_type='badge',
        entity_id=badge_id,
        field_name='status',
        from_value=old_status,
        to_value='issued',
        note=f"Issued with inscription {inscription_id}",
        changed_by_id=current_user['id']
    )
    db.session.add(status_change)
    db.session.commit()
    
    return jsonify({'success': True, 'badge': badge.to_dict()})

# ================================================================
# VOTING ROUTES
# ================================================================

@app.route('/api/projects/<project_id>/submissions/', methods=['GET'])
def api_list_project_submissions(project_id):
    """List approved drafts (not RFCs) for a project - eligible for voting"""
    project = Project.query.get_or_404(project_id)
    
    # Simple criteria: approved drafts for this project (not RFCs)
    submissions = Submission.query.filter(
        Submission.project_id == project_id,
        Submission.status == 'approved',
        Submission.doc_type == 'draft'
    ).order_by(Submission.submitted_at.desc()).all()
    
    return jsonify({
        'submissions': [{
            'id': s.id,
            'public_id': s.public_id,
            'title': s.title,
            'draft_name': s.draft_name,
            'ml_number': s.ml_number,
            'group': s.group,
            'status': s.status,
            'submitted_at': s.submitted_at.isoformat() if s.submitted_at else None
        } for s in submissions]
    })

@app.route('/api/projects/<project_id>/votes/', methods=['POST'])
@require_auth
def api_create_vote(project_id):
    """Create a new vote for a project"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    project = Project.query.get_or_404(project_id)
    
    # Check if user is project admin or site admin
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project admins can create votes'}), 403
    
    data = request.get_json()
    title = data.get('title', '').strip()
    description = data.get('description', '').strip()
    submission_id = data.get('submission_id', '').strip()
    start_at_str = data.get('start_at', '').strip()
    end_at_str = data.get('end_at', '').strip()
    quorum_count = data.get('quorum_count')
    win_threshold = data.get('win_threshold', 0.5)
    
    # Validation
    if not title:
        return jsonify({'error': 'Vote title is required'}), 400
    if not submission_id:
        return jsonify({'error': 'Submission ID is required'}), 400
    if not start_at_str or not end_at_str:
        return jsonify({'error': 'Start and end times are required'}), 400
    if quorum_count is None or quorum_count < 1:
        return jsonify({'error': 'Quorum count must be at least 1'}), 400
    if not (0.0 <= win_threshold <= 1.0):
        return jsonify({'error': 'Win threshold must be between 0.0 and 1.0'}), 400
    
    # Check submission exists
    submission = Submission.query.get(submission_id)
    if not submission:
        return jsonify({'error': 'Submission not found'}), 404
    
    # Parse dates
    try:
        start_at = datetime.fromisoformat(start_at_str.replace('Z', '+00:00'))
        end_at = datetime.fromisoformat(end_at_str.replace('Z', '+00:00'))
    except ValueError:
        return jsonify({'error': 'Invalid date format. Use ISO 8601 format.'}), 400
    
    # Validate dates (strip timezone info to compare naive datetimes)
    start_at = start_at.replace(tzinfo=None)
    end_at = end_at.replace(tzinfo=None)
    now = datetime.utcnow()
    if start_at <= now:
        return jsonify({'error': 'Start time must be in the future'}), 400
    if end_at <= start_at:
        return jsonify({'error': 'End time must be after start time'}), 400
    
    # Create vote
    vote = Vote(
        project_id=project_id,
        submission_id=submission_id,
        created_by_id=current_user['id'],
        title=title,
        description=description or None,
        start_at=start_at,
        end_at=end_at,
        quorum_count=quorum_count,
        win_threshold=win_threshold,
        status='scheduled'
    )
    
    db.session.add(vote)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'vote': {
            'id': vote.id,
            'public_id': vote.public_id,
            'title': vote.title,
            'description': vote.description,
            'status': vote.status,
            'start_at': vote.start_at.isoformat(),
            'end_at': vote.end_at.isoformat(),
            'quorum_count': vote.quorum_count,
            'win_threshold': vote.win_threshold
        }
    }), 201

@app.route('/api/projects/<project_id>/votes/', methods=['GET'])
def api_list_votes(project_id):
    """List votes for a project"""
    try:
        project = Project.query.get_or_404(project_id)
        
        status_filter = request.args.get('status')
        query = Vote.query.filter_by(project_id=project_id)
        
        if status_filter:
            query = query.filter_by(status=status_filter)
        
        votes = query.order_by(Vote.created_at.desc()).all()
        
        return jsonify({
            'votes': [{
                'id': v.id,
                'public_id': v.public_id,
                'title': v.title,
                'description': v.description,
                'status': v.status,
                'result': v.result,
                'start_at': v.start_at.isoformat() if v.start_at else None,
                'end_at': v.end_at.isoformat() if v.end_at else None,
                'created_at': v.created_at.isoformat() if v.created_at else None
            } for v in votes]
        })
    except Exception as e:
        app.logger.error(f"Error in api_list_votes for project {project_id}: {e}")
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/api/votes/<vote_id>/', methods=['GET'])
def api_get_vote(vote_id):
    """Get vote details"""
    # Support both integer ID and public_id UUID
    if len(vote_id) == 36 and '-' in vote_id:
        vote = Vote.query.filter_by(public_id=vote_id).first_or_404()
    else:
        vote = Vote.query.get_or_404(vote_id)
    
    # Count ballots
    ballot_count = Ballot.query.filter_by(vote_id=vote.id).count()
    eligible_count = VoteEligibilitySnapshot.query.filter_by(vote_id=vote.id, is_eligible=True).count()
    
    # Parse result summary if available
    result_summary = None
    if vote.result_summary:
        try:
            result_summary = json.loads(vote.result_summary)
        except:
            pass
    
    return jsonify({
        'id': vote.id,
        'public_id': vote.public_id,
        'project_id': vote.project_id,
        'submission_id': vote.submission_id,
        'title': vote.title,
        'description': vote.description,
        'status': vote.status,
        'result': vote.result,
        'result_summary': result_summary,
        'start_at': vote.start_at.isoformat() if vote.start_at else None,
        'end_at': vote.end_at.isoformat() if vote.end_at else None,
        'quorum_count': vote.quorum_count,
        'win_threshold': vote.win_threshold,
        'ballot_count': ballot_count,
        'eligible_count': eligible_count,
        'created_at': vote.created_at.isoformat() if vote.created_at else None,
        'closed_at': vote.closed_at.isoformat() if vote.closed_at else None
    })

@app.route('/api/votes/<vote_id>/ballot/', methods=['POST'])
@require_auth
def api_cast_ballot(vote_id):
    """Cast or update a ballot"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    # Support both integer ID and public_id UUID
    if len(vote_id) == 36 and '-' in vote_id:
        vote = Vote.query.filter_by(public_id=vote_id).first_or_404()
    else:
        vote = Vote.query.get_or_404(vote_id)
    
    # Check vote is active
    if vote.status != 'active':
        return jsonify({'error': 'Vote is not active'}), 400
    
    # Check eligibility
    eligibility = VoteEligibilitySnapshot.query.filter_by(
        vote_id=vote.id,
        person_id=current_user['id']
    ).first()
    
    if not eligibility or not eligibility.is_eligible:
        return jsonify({'error': 'You are not eligible to vote in this election'}), 403
    
    data = request.get_json()
    choice = data.get('choice', '').strip().lower()
    
    if choice not in ['yes', 'no', 'abstain']:
        return jsonify({'error': 'Choice must be yes, no, or abstain'}), 400
    
    # Check for existing ballot
    existing_ballot = Ballot.query.filter_by(
        vote_id=vote.id,
        person_id=current_user['id']
    ).first()
    
    if existing_ballot:
        # Update existing ballot
        existing_ballot.choice = choice
        existing_ballot.cast_at = datetime.utcnow()
    else:
        # Create new ballot
        ballot = Ballot(
            vote_id=vote.id,
            person_id=current_user['id'],
            choice=choice
        )
        db.session.add(ballot)
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'choice': choice,
        'cast_at': datetime.utcnow().isoformat()
    })

@app.route('/api/votes/<vote_id>/cancel/', methods=['POST'])
@require_auth
def api_cancel_vote(vote_id):
    """Cancel a vote"""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'error': 'Authentication required'}), 401
    
    # Support both integer ID and public_id UUID
    if len(vote_id) == 36 and '-' in vote_id:
        vote = Vote.query.filter_by(public_id=vote_id).first_or_404()
    else:
        vote = Vote.query.get_or_404(vote_id)
    
    project = Project.query.get_or_404(vote.project_id)
    
    # Check if user is project admin or site admin
    if not is_project_admin(project, current_user):
        return jsonify({'error': 'Only project admins can cancel votes'}), 403
    
    # Check vote can be canceled
    if vote.status not in ['scheduled', 'active']:
        return jsonify({'error': 'Only scheduled or active votes can be canceled'}), 400
    
    vote.status = 'canceled'
    vote.result = 'canceled'
    vote.closed_at = datetime.utcnow()
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'vote': {
            'id': vote.id,
            'public_id': vote.public_id,
            'status': vote.status,
            'result': vote.result
        }
    })

@app.route('/my-projects/')
@require_auth
def my_projects():
    """Redirect to current user's profile with My Projects tab active"""
    current_user = get_current_user()
    if not current_user:
        return redirect('/')
    user = User.query.get(current_user['id'])
    if not user:
        return redirect('/')
    return redirect(f'/profile/{user.username}/#my-projects')

@app.route('/profile/', methods=['GET', 'POST'])
@require_auth
def profile():
    """User profile management - redirects to new profile page"""
    # Redirect GET requests to the new profile edit page
    if request.method == 'GET':
        return redirect(url_for('profile_edit'))
    
    current_user = get_current_user()
    
    if request.method == 'POST':
        action = request.form.get('action')
        user = User.query.filter_by(username=session['user']).first()

        if action == 'update_password':
            old_password = request.form.get('old_password', '').strip()
            new_password = request.form.get('new_password', '').strip()
            
            if check_password_hash(user.password_hash, old_password):
                if len(new_password) >= 6:
                    user.password_hash = generate_password_hash(new_password)
                    db.session.commit()
                    flash('Password updated successfully!', 'success')
                else:
                    flash('New password must be at least 6 characters.', 'error')
            else:
                flash('Current password is incorrect.', 'error')
        
        elif action == 'update_profile':
            name = request.form.get('name', '').strip()
            email = request.form.get('email', '').strip()
            
            # Check if email is already taken by another user
            existing_email = User.query.filter(User.email == email, User.username != session['user']).first()
            if existing_email:
                flash('Email already registered to another account.', 'error')
            else:
                if name:
                    user.name = name
                if email:
                    user.email = email
                db.session.commit()
                flash('Profile updated successfully!', 'success')

        elif action == 'update_theme':
            theme = request.form.get('theme', 'dark').strip()
            if theme in ['light', 'dark', 'auto']:
                user.theme = theme
                db.session.commit()
                session['theme'] = theme  # Update session immediately
                flash('Theme preference updated successfully!', 'success')
            else:
                flash('Invalid theme selection.', 'error')
    
    # Generate user menu
    user_menu = generate_user_menu()
    
    current_theme = current_user.get('theme', 'dark')
    light_selected = 'selected' if current_theme == 'light' else ''
    dark_selected = 'selected' if current_theme == 'dark' else ''
    auto_selected = 'selected' if current_theme == 'auto' else ''
    
    profile_content = PROFILE_TEMPLATE.format(
        current_user_name=current_user['name'],
        current_user_email=current_user['email'],
        current_user_theme=current_theme,
        light_selected=light_selected,
        dark_selected=dark_selected,
        auto_selected=auto_selected,
        session_user=session['user']
    )
    return render_template_string(BASE_TEMPLATE.format(title="Profile - MLGH", theme=current_theme, user_menu=user_menu, content=profile_content, build_number=BUILD_NUMBER, hypothesis_config=""))

@app.route('/admin/')
@require_role('admin')
def admin_dashboard():
    user_menu = generate_user_menu()

    # Enhanced admin statistics
    total_users = User.query.count()
    total_groups = len(GROUPS)
    total_submissions = Submission.query.count()
    approved_drafts = Submission.query.filter(Submission.status.in_(['approved', 'published'])).count()
    pending_chairs = WorkingGroupChair.query.filter_by(approved=False).count()
    
    # New statistics for Projects/Workgroups/Guilds
    total_projects = Project.query.count()
    pending_projects = Project.query.filter_by(approval_status='pending').count()
    total_workgroups = Workgroup.query.count()
    pending_workgroups = Workgroup.query.filter_by(approval_status='pending').count()
    total_guilds = Guild.query.count()
    total_roles = Role.query.count()
    pending_roles = Role.query.filter_by(status='draft').count()  # Roles use 'draft' status before approval
    total_claims = Claim.query.count()
    pending_claims = Claim.query.filter_by(status='pending_approval').count()  # Claims use 'pending_approval'
    total_badges = Badge.query.count()
    pending_badges = Badge.query.filter_by(status='requested').count()  # Badges use 'requested' status

    # Recent activity and alerts
    pending_submissions = Submission.query.filter_by(status='submitted').count()
    recent_submissions = Submission.query.order_by(Submission.submitted_at.desc()).limit(5).all()
    recent_users = User.query.order_by(User.created_at.desc()).limit(5).all()

    # Get most active drafts (by comment count or views if we had them)
    # For now, just show recent submissions as proxy
    active_drafts = Submission.query.order_by(Submission.submitted_at.desc()).limit(10).all()

    # Most active users (by login frequency - simplified)
    active_users = User.query.order_by(User.last_login.desc()).limit(10).all()

    # Build recent activity feed
    activity_html = ""
    for submission in recent_submissions[:3]:  # Show last 3 submissions
        activity_html += f"""
        <div class="activity-item mb-2">
            <small class="text-muted">
                <i class="fas fa-file-alt me-1"></i>
                New submission: <strong>{submission.title[:50]}...</strong>
                by {submission.submitted_by}
                <span class="float-end">{submission.submitted_at.strftime('%m/%d %H:%M')}</span>
            </small>
        </div>
        """

    for user in recent_users[:2]:  # Show last 2 new users
        activity_html += f"""
        <div class="activity-item mb-2">
            <small class="text-muted">
                <i class="fas fa-user-plus me-1"></i>
                New user: <strong>{user.name}</strong> ({user.email})
                <span class="float-end">{user.created_at.strftime('%m/%d %H:%M')}</span>
            </small>
        </div>
        """

    # Build alerts section
    alerts_html = ""
    if pending_submissions > 0:
        alerts_html += f"""
        <div class="alert alert-warning alert-dismissible fade show" role="alert">
            <i class="fas fa-exclamation-triangle me-2"></i>
            <strong>{pending_submissions}</strong> draft submission(s) pending review
            <a href="/admin/submissions/" class="alert-link">Review now</a>
            <button type="button" class="btn-close" data-bs-dismiss="alert"></button>
        </div>
        """

    if pending_chairs > 0:
        alerts_html += f"""
        <div class="alert alert-info alert-dismissible fade show" role="alert">
            <i class="fas fa-users me-2"></i>
            <strong>{pending_chairs}</strong> chair nomination(s) pending approval
            <a href="/admin/chair-nominations/" class="alert-link">Review nominations</a>
            <button type="button" class="btn-close" data-bs-dismiss="alert"></button>
        </div>
        """
    
    if pending_projects > 0:
        alerts_html += f"""
        <div class="alert alert-primary alert-dismissible fade show" role="alert">
            <i class="fas fa-project-diagram me-2"></i>
            <strong>{pending_projects}</strong> layer(s) pending approval
            <a href="/admin/projects/" class="alert-link">Review now</a>
            <button type="button" class="btn-close" data-bs-dismiss="alert"></button>
        </div>
        """
    
    if pending_workgroups > 0:
        alerts_html += f"""
        <div class="alert alert-warning alert-dismissible fade show" role="alert">
            <i class="fas fa-users me-2"></i>
            <strong>{pending_workgroups}</strong> workgroup(s) pending approval
            <a href="/admin/workgroups/" class="alert-link">Review now</a>
            <button type="button" class="btn-close" data-bs-dismiss="alert"></button>
        </div>
        """
    
    if pending_roles > 0:
        alerts_html += f"""
        <div class="alert alert-secondary alert-dismissible fade show" role="alert">
            <i class="fas fa-user-tag me-2"></i>
            <strong>{pending_roles}</strong> role(s) pending approval
            <a href="/admin/roles/" class="alert-link">Review now</a>
            <button type="button" class="btn-close" data-bs-dismiss="alert"></button>
        </div>
        """
    
    if pending_badges > 0:
        alerts_html += f"""
        <div class="alert alert-success alert-dismissible fade show" role="alert">
            <i class="fas fa-award me-2"></i>
            <strong>{pending_badges}</strong> badge(s) pending issuance
            <a href="/admin/badges/" class="alert-link">Issue now</a>
            <button type="button" class="btn-close" data-bs-dismiss="alert"></button>
        </div>
        """

    content = f"""
    <div class="container mt-4">
        <!-- Alerts Section -->
        <div id="admin-alerts" class="mb-4">
            {alerts_html}
        </div>

        <div class="row">
            <div class="col-12">
                <div class="d-flex justify-content-between align-items-center mb-4">
                    <h1>Admin Dashboard</h1>
                    <div>
                        <a href="/admin/users/" class="btn btn-outline-primary me-2"><i class="fas fa-users me-1"></i>Users</a>
                        <a href="/admin/submissions/" class="btn btn-outline-success me-2"><i class="fas fa-file-alt me-1"></i>Submissions</a>
                        <a href="/admin/projects/" class="btn btn-outline-info me-2"><i class="fas fa-project-diagram me-1"></i>Layers</a>
                        <a href="/admin/workgroups/" class="btn btn-outline-warning me-2"><i class="fas fa-users me-1"></i>Workgroups</a>
                        <a href="/admin/roles/" class="btn btn-outline-secondary me-2"><i class="fas fa-user-tag me-1"></i>Roles</a>
                        <a href="/admin/badges/" class="btn btn-outline-primary"><i class="fas fa-award me-1"></i>Badges</a>
                    </div>
                </div>

                <!-- Statistics Cards -->
                <div class="row mb-4">
                    <div class="col-md-3">
                        <div class="card h-100">
                            <div class="card-body text-center">
                                <h4 class="text-primary mb-1">{total_users}</h4>
                                <p class="mb-0 small">Total Users</p>
                            </div>
                        </div>
                    </div>
                    <div class="col-md-3">
                        <div class="card h-100">
                            <div class="card-body text-center">
                                <h4 class="text-info mb-1">{total_projects}</h4>
                                <p class="mb-0 small">Layers</p>
                                {f'<small class="text-warning">({pending_projects} pending)</small>' if pending_projects > 0 else ''}
                            </div>
                        </div>
                    </div>
                    <div class="col-md-3">
                        <div class="card h-100">
                            <div class="card-body text-center">
                                <h4 class="text-success mb-1">{total_workgroups}</h4>
                                <p class="mb-0 small">Workgroups</p>
                                {f'<small class="text-warning">({pending_workgroups} pending)</small>' if pending_workgroups > 0 else ''}
                            </div>
                        </div>
                    </div>
                    <div class="col-md-3">
                        <div class="card h-100">
                            <div class="card-body text-center">
                                <h4 class="text-secondary mb-1">{total_roles}</h4>
                                <p class="mb-0 small">Roles</p>
                                {f'<small class="text-warning">({pending_roles} pending)</small>' if pending_roles > 0 else ''}
                            </div>
                        </div>
                    </div>
                </div>
                
                <div class="row mb-4">
                    <div class="col-md-3">
                        <div class="card h-100">
                            <div class="card-body text-center">
                                <h4 class="text-warning mb-1">{total_submissions}</h4>
                                <p class="mb-0 small">Submissions</p>
                                {f'<small class="text-danger">({pending_submissions} pending)</small>' if pending_submissions > 0 else ''}
                            </div>
                        </div>
                    </div>
                    <div class="col-md-3">
                        <div class="card h-100">
                            <div class="card-body text-center">
                                <h4 class="text-primary mb-1">{total_badges}</h4>
                                <p class="mb-0 small">Badges</p>
                                {f'<small class="text-warning">({pending_badges} pending)</small>' if pending_badges > 0 else ''}
                            </div>
                        </div>
                    </div>
                    <div class="col-md-3">
                        <div class="card h-100">
                            <div class="card-body text-center">
                                <h4 class="text-success mb-1">{total_claims}</h4>
                                <p class="mb-0 small">Claims</p>
                                {f'<small class="text-warning">({pending_claims} pending)</small>' if pending_claims > 0 else ''}
                            </div>
                        </div>
                    </div>
                    <div class="col-md-3">
                        <div class="card h-100">
                            <div class="card-body text-center">
                                <h4 class="text-info mb-1">{total_guilds}</h4>
                                <p class="mb-0 small">Guilds</p>
                            </div>
                        </div>
                    </div>
                </div>

                <div class="row">
                    <!-- Recent Activity -->
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header d-flex justify-content-between align-items-center">
                                <h5 class="mb-0">Recent Activity</h5>
                                <span class="badge bg-primary">Live</span>
                            </div>
                            <div class="card-body">
                                {activity_html}
                                <hr>
                                <a href="/admin/activity/" class="btn btn-sm btn-outline-primary">View All Activity</a>
                            </div>
                        </div>
                    </div>

                    <!-- Quick Actions -->
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <h5>Quick Actions</h5>
                            </div>
                            <div class="card-body">
                                <div class="d-grid gap-2">
                                    <a href="/admin/submissions/" class="btn btn-success">
                                        <i class="fas fa-check-circle me-2"></i>Review Submissions ({pending_submissions} pending)
                                    </a>
                                    <a href="/admin/users/" class="btn btn-primary">
                                        <i class="fas fa-users me-2"></i>Manage Users ({total_users} total)
                                    </a>
                                    <a href="/group/" class="btn btn-info">
                                        <i class="fas fa-users-cog me-2"></i>Manage Workgroups ({pending_chairs} pending coordinators)
                                    </a>
                                    <a href="/admin/analytics/" class="btn btn-secondary">
                                        <i class="fas fa-chart-bar me-2"></i>View Analytics
                                    </a>
                                </div>
                            </div>
                        </div>
                    </div>
                </div>

                <!-- Content Management Section -->
                <div class="row mt-4">
                    <div class="col-12">
                        <h3 class="mb-3">Content Management</h3>
                    </div>
                </div>

                <div class="row">
                    <!-- Most Active Drafts -->
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <h5>Recent Draft Submissions</h5>
                            </div>
                            <div class="card-body">
                                <div class="list-group list-group-flush">
                                    {"".join([f'''
                                    <a href="/doc/draft/{draft.id}/" class="list-group-item list-group-item-action d-flex justify-content-between align-items-center">
                                        <div>
                                            <strong>{draft.title[:40]}...</strong>
                                            <br><small class="text-muted">by {draft.submitted_by} • {draft.submitted_at.strftime('%m/%d')}</small>
                                        </div>
                                        <span class="badge bg-{'warning' if draft.status == 'submitted' else 'success'}">{draft.status}</span>
                                    </a>
                                    ''' for draft in active_drafts[:5]])}
                                </div>
                            </div>
                        </div>
                    </div>

                    <!-- Active Users -->
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <h5>Recent User Activity</h5>
                            </div>
                            <div class="card-body">
                                <div class="list-group list-group-flush">
                                    {"".join([f'''
                                    <div class="list-group-item d-flex justify-content-between align-items-center">
                                        <div>
                                            <strong>{user.name}</strong>
                                            <br><small class="text-muted">{user.email} • {user.role}</small>
                                        </div>
                                        <small class="text-muted">
                                            {user.last_login.strftime('%m/%d %H:%M') if user.last_login else 'Never logged in'}
                                        </small>
                                    </div>
                                    ''' for user in active_users[:5]])}
                                </div>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </div>
    </div>
    """

    return BASE_TEMPLATE.format(
        title="Admin Dashboard - MLGH",
        theme=get_current_user().get('theme', 'dark'),
        content=content,
        user_menu=user_menu, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/admin/users/')
@require_role('admin')
def admin_users():
    user_menu = generate_user_menu()
    current_theme = get_current_user().get('theme', 'dark')

    # Get all users with pagination
    page = request.args.get('page', 1, type=int)
    per_page = 20
    search = request.args.get('search', '').strip()
    role_filter = request.args.get('role', '')

    query = User.query

    if search:
        query = query.filter(
            db.or_(
                User.username.contains(search),
                User.name.contains(search),
                User.email.contains(search)
            )
        )

    if role_filter:
        query = query.filter_by(role=role_filter)

    users = query.order_by(User.created_at.desc()).paginate(page=page, per_page=per_page, error_out=False)
    total_users = query.count()

    # Build user rows
    user_rows = ""
    for user in users.items:
        role_badge = {
            'admin': 'badge bg-danger',
            'editor': 'badge bg-warning',
            'user': 'badge bg-secondary'
        }.get(user.role, 'badge bg-secondary')

        last_login = user.last_login.strftime('%Y-%m-%d %H:%M') if user.last_login else 'Never'

        user_rows += f"""
        <tr>
            <td>
                <strong>{user.name}</strong><br>
                <small class="text-muted">@{user.username}</small>
            </td>
            <td>{user.email}</td>
            <td><span class="{role_badge}">{user.role.title()}</span></td>
            <td>{user.theme.title()}</td>
            <td>{user.created_at.strftime('%Y-%m-%d')}</td>
            <td>{last_login}</td>
            <td>
                <div class="btn-group btn-group-sm" role="group">
                    <div class="dropdown">
                        <button class="btn btn-outline-primary btn-sm dropdown-toggle" type="button" id="roleDropdown{user.username}" data-bs-toggle="dropdown" aria-expanded="false" data-bs-offset="0,4">
                            <i class="fas fa-user-edit"></i>
                        </button>
                        <ul class="dropdown-menu dropdown-menu-end" aria-labelledby="roleDropdown{user.username}">
                            <li><a class="dropdown-item" href="#" onclick="changeRole('{user.username}', 'user'); return false;">User</a></li>
                            <li><a class="dropdown-item" href="#" onclick="changeRole('{user.username}', 'editor'); return false;">Editor</a></li>
                            <li><a class="dropdown-item" href="#" onclick="changeRole('{user.username}', 'admin'); return false;">Admin</a></li>
                            <li><hr class="dropdown-divider"></li>
                            <li><a class="dropdown-item" href="/admin/users/{user.id}/add-coordinator"><i class="fas fa-user-tie me-1"></i>Add as coordinator</a></li>
                        </ul>
                    </div>
                    <button class="btn btn-outline-danger btn-sm ms-1" onclick="deleteUser('{user.username}')">
                        <i class="fas fa-trash"></i>
                    </button>
                </div>
            </td>
        </tr>
        """

    # Role filter options
    role_options = f"""
    <option value="">All Roles</option>
    <option value="admin" {'selected' if role_filter == 'admin' else ''}>Admin</option>
    <option value="editor" {'selected' if role_filter == 'editor' else ''}>Editor</option>
    <option value="user" {'selected' if role_filter == 'user' else ''}>User</option>
    """

    content = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/admin/">Admin Dashboard</a></li>
                <li class="breadcrumb-item active">User Management</li>
            </ol>
        </nav>

        <div class="d-flex justify-content-between align-items-center mb-4">
            <h1>User Management</h1>
            <div>
                <span class="badge bg-info me-2">Total: {total_users} users</span>
            </div>
        </div>

        <!-- Filters and Search -->
        <div class="card mb-4">
            <div class="card-body">
                <form method="GET" class="row g-3">
                    <div class="col-md-6">
                        <label for="search" class="form-label">Search Users</label>
                        <input type="text" class="form-control" id="search" name="search"
                               value="{search}" placeholder="Name, username, or email">
                    </div>
                    <div class="col-md-4">
                        <label for="role" class="form-label">Filter by Role</label>
                        <select class="form-select" id="role" name="role">
                            {role_options}
                        </select>
                    </div>
                    <div class="col-md-2 d-flex align-items-end">
                        <button type="submit" class="btn btn-primary me-2">
                            <i class="fas fa-search me-1"></i>Filter
                        </button>
                        <a href="/admin/users/" class="btn btn-outline-secondary">
                            <i class="fas fa-times"></i>
                        </a>
                    </div>
                </form>
            </div>
        </div>

        <!-- Users Table -->
        <div class="card">
            <div class="card-header">
                <h5 class="mb-0">Users ({users.total} total)</h5>
            </div>
            <div class="card-body p-0">
                <div class="table-responsive">
                    <table class="table table-hover mb-0">
                        <thead>
                            <tr>
                                <th>Name</th>
                                <th>Email</th>
                                <th>Role</th>
                                <th>Theme</th>
                                <th>Joined</th>
                                <th>Last Login</th>
                                <th>Actions</th>
                            </tr>
                        </thead>
                        <tbody>
                            {user_rows}
                        </tbody>
                    </table>
                </div>
            </div>
        </div>

        <!-- Pagination -->
        {f'''
        <nav aria-label="User pagination" class="mt-4">
            <ul class="pagination justify-content-center">
                {f'<li class="page-item {"disabled" if not users.has_prev else ""}"><a class="page-link" href="?page={users.prev_num}&search={search}&role={role_filter}">Previous</a></li>' if users.has_prev else ''}
                {''.join([f'<li class="page-item {"active" if i == users.page else ""}"><a class="page-link" href="?page={i}&search={search}&role={role_filter}">{i}</a></li>' for i in users.iter_pages()])}
                {f'<li class="page-item {"disabled" if not users.has_next else ""}"><a class="page-link" href="?page={users.next_num}&search={search}&role={role_filter}">Next</a></li>' if users.has_next else ''}
            </ul>
        </nav>
        ''' if users.pages > 1 else ''}
        </div>

    <script>
        function changeRole(username, newRole) {{
            console.log('Changing role for', username, 'to', newRole);
            
            fetch('/admin/users/' + username + '/role', {{
                method: 'POST',
                headers: {{
                    'Content-Type': 'application/json',
                }},
                body: JSON.stringify({{ role: newRole }})
            }})
            .then(response => {{
                console.log('Response status:', response.status);
                return response.json();
            }})
            .then(data => {{
                console.log('Response data:', data);
                if (data.success) {{
                    location.reload();
                }} else {{
                    alert('Error: ' + (data.message || 'Unknown error'));
                }}
            }})
            .catch(error => {{
                console.error('Error:', error);
                alert('Error updating role: ' + error.message);
            }});
        }}

        function deleteUser(username) {{
            if (confirm("Are you sure you want to delete user " + username + "? This action cannot be undone.")) {{
                console.log('Deleting user:', username);
                fetch('/admin/users/' + username + '/delete', {{
                    method: 'POST',
                    headers: {{
                        'Content-Type': 'application/json',
                    }}
                }})
                .then(response => {{
                    console.log('Delete response status:', response.status);
                    return response.json();
                }})
                .then(data => {{
                    console.log('Delete response data:', data);
                    if (data.success) {{
                        // Just reload without alert
                        location.reload();
                    }} else {{
                        alert('Error: ' + (data.message || 'Unknown error'));
                    }}
                }})
                .catch(error => {{
                    console.error('Error:', error);
                    alert('Error deleting user: ' + error.message);
                }});
            }}
        }}
    </script>
    """

    return BASE_TEMPLATE.format(
        title="User Management - MLGH",
        theme=current_theme,
        user_menu=user_menu,
        content=content, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/admin/users/<username>/role', methods=['POST'])
@require_role('admin')
def change_user_role(username):
    data = request.get_json()
    new_role = data.get('role', '')

    if new_role not in ['user', 'editor', 'admin']:
        return jsonify({'success': False, 'message': 'Invalid role'}), 400

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({'success': False, 'message': 'User not found'}), 404

    # Prevent admin from demoting themselves
    current_admin = get_current_user()
    if user.id == current_admin['id'] and new_role != 'admin':
        return jsonify({'success': False, 'message': 'Cannot change your own admin role'}), 400

    user.role = new_role
    db.session.commit()

    # Log the action
    add_to_document_history(f"user-{user.id}", "role_changed", current_admin['name'],
                           f"Changed {user.name}'s role to {new_role}")

    return jsonify({'success': True, 'message': f'Role changed to {new_role}'})

@app.route('/admin/users/<username>/delete', methods=['POST'])
@require_role('admin')
def delete_user(username):
    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({'success': False, 'message': 'User not found'}), 404

    # Prevent admin from deleting themselves
    current_admin = get_current_user()
    if user.id == current_admin['id']:
        return jsonify({'success': False, 'message': 'Cannot delete your own account'}), 400

    # Log before deletion
    add_to_document_history(f"user-{user.id}", "user_deleted", current_admin['name'],
                           f"Deleted user {user.name} ({user.email})")

    db.session.delete(user)
    db.session.commit()

    return jsonify({'success': True, 'message': 'User deleted successfully'})

@app.route('/admin/submissions/')
@require_role('admin')
def admin_submissions():
    user_menu = generate_user_menu()
    current_theme = get_current_user().get('theme', 'dark')

    # Get submissions with filters
    status_filter = request.args.get('status', 'submitted')
    page = request.args.get('page', 1, type=int)
    per_page = 10

    query = Submission.query

    if status_filter and status_filter != 'all':
        query = query.filter_by(status=status_filter)

    submissions = query.order_by(Submission.submitted_at.desc()).paginate(
        page=page, per_page=per_page, error_out=False)

    # Build submission cards
    submission_cards = ""
    for submission in submissions.items:
        status_badge = {
            'submitted': 'badge bg-warning text-dark',
            'approved': 'badge bg-success',
            'rejected': 'badge bg-danger',
            'published': 'badge bg-info'
        }.get(submission.status, 'badge bg-secondary')

        # Get revision info
        is_revision = getattr(submission, 'is_revision', False)
        revision_number = getattr(submission, 'revision_number', '')
        parent_draft_name = getattr(submission, 'parent_draft_name', '')
        revision_badge = f'<span class="badge bg-success ms-2">Revision {revision_number}</span>' if is_revision and revision_number else ''

        # Get source info (file or ordinal)
        source_type = getattr(submission, 'sourceType', 'file')
        if source_type == 'ordinal':
            inscription_number = getattr(submission, 'inscriptionNumber', None)
            ordinal_id = getattr(submission, 'ordinalId', None)
            block_height = getattr(submission, 'blockHeight', None)
            if inscription_number:
                source_info = f'<span class="badge bg-info"><i class="bi bi-coin"></i> Ordinal</span> Inscription #{inscription_number}'
                if block_height:
                    source_info += f' (Block {block_height})'
            elif ordinal_id:
                source_info = f'<span class="badge bg-info"><i class="bi bi-coin"></i> Ordinal</span> {ordinal_id[:16]}...'
            else:
                source_info = '<span class="badge bg-info"><i class="bi bi-coin"></i> Ordinal</span>'
        else:
            file_size = "N/A"
            if submission.file_path and os.path.exists(submission.file_path):
                file_size = f"{os.path.getsize(submission.file_path) / 1024:.1f} KB"
            source_info = f'<span class="badge bg-secondary"><i class="bi bi-file-earmark"></i> File</span> {submission.filename} ({file_size})'

        action_buttons = ""
        if submission.status == 'submitted':
            action_buttons = f"""
            <button class="btn btn-success btn-sm me-2" onclick="approveSubmission('{submission.id}')">
                <i class="fas fa-check me-1"></i>Approve
            </button>
            <button class="btn btn-danger btn-sm me-2" onclick="rejectSubmission('{submission.id}')">
                <i class="fas fa-times me-1"></i>Reject
            </button>
            <button class="btn btn-info btn-sm" onclick="publishAsRFC('{submission.id}')">
                <i class="fas fa-star me-1"></i>Publish as RFC
            </button>
            """
        elif submission.status == 'approved':
            action_buttons = f"""
            <button class="btn btn-info btn-sm me-2" onclick="publishAsRFC('{submission.id}')">
                <i class="fas fa-star me-1"></i>Publish as RFC
            </button>
            <button class="btn btn-warning btn-sm" onclick="unapproveSubmission('{submission.id}')">
                <i class="fas fa-undo me-1"></i>Unapprove
            </button>
            """

        # Add revision context if this is a revision
        revision_context = ""
        if is_revision and parent_draft_name:
            revision_context = f'<p class="mb-2"><strong>Revision of:</strong> <a href="/doc/draft/{parent_draft_name}/" class="text-decoration-none">{parent_draft_name}</a></p>'

        submission_cards += f"""
        <div class="card mb-3">
            <div class="card-header d-flex justify-content-between align-items-center">
                <h6 class="mb-0">
                    <a href="/doc/draft/{submission.id}/" class="text-decoration-none">
                        {submission.title}
                    </a>
                </h6>
                <div>
                    <span class="{status_badge}">{submission.status.title()}</span>
                    {revision_badge}
                </div>
            </div>
            <div class="card-body">
                <div class="row">
                    <div class="col-md-8">
                        <p class="mb-2"><strong>Authors:</strong> {', '.join(submission.authors)}</p>
                        <p class="mb-2"><strong>Group:</strong> {submission.group or 'None'}</p>
                        <p class="mb-2"><strong>Submitted:</strong> {submission.submitted_at.strftime('%Y-%m-%d %H:%M')} by {submission.submitted_by}</p>
                        <p class="mb-2"><strong>Source:</strong> {source_info}</p>
                        {revision_context}
                        {f'<p class="mb-2"><strong>Abstract:</strong> {submission.abstract[:200]}...</p>' if submission.abstract else ''}
                    </div>
                    <div class="col-md-4">
                        <div class="d-grid gap-2">
                            <a href="/doc/draft/{submission.id}/" class="btn btn-outline-primary btn-sm">
                                <i class="fas fa-eye me-1"></i>View Draft
                            </a>
                            {action_buttons}
                        </div>
                    </div>
                </div>
            </div>
        </div>
        """

    # Status filter options
    status_options = f"""
    <option value="all" {'selected' if status_filter == 'all' else ''}>All Submissions</option>
    <option value="submitted" {'selected' if status_filter == 'submitted' else ''}>Pending Review</option>
    <option value="approved" {'selected' if status_filter == 'approved' else ''}>Approved</option>
    <option value="rejected" {'selected' if status_filter == 'rejected' else ''}>Rejected</option>
    <option value="published" {'selected' if status_filter == 'published' else ''}>Published</option>
    """

    content = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/admin/">Admin Dashboard</a></li>
                <li class="breadcrumb-item active">Submission Management</li>
            </ol>
        </nav>

        <div class="d-flex justify-content-between align-items-center mb-4">
            <h1>Submission Management</h1>
            <div>
                <select class="form-select form-select-sm" onchange="changeStatusFilter(this.value)">
                    {status_options}
                </select>
            </div>
        </div>

        <!-- Statistics -->
        <div class="row mb-4">
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-warning">{Submission.query.filter_by(status='submitted').count()}</h4>
                        <p class="mb-0 small">Pending Review</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-success">{Submission.query.filter_by(status='approved').count()}</h4>
                        <p class="mb-0 small">Approved</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-danger">{Submission.query.filter_by(status='rejected').count()}</h4>
                        <p class="mb-0 small">Rejected</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-info">{Submission.query.filter_by(status='published').count()}</h4>
                        <p class="mb-0 small">Published as RFC</p>
                    </div>
                </div>
            </div>
        </div>

        <!-- Submissions -->
        <div id="submissions-container">
            {submission_cards}
        </div>

        <!-- Pagination -->
        {f'''
        <nav aria-label="Submission pagination" class="mt-4">
            <ul class="pagination justify-content-center">
                {f'<li class="page-item {"disabled" if not submissions.has_prev else ""}"><a class="page-link" href="?page={submissions.prev_num}&status={status_filter}">Previous</a></li>' if submissions.has_prev else ''}
                {''.join([f'<li class="page-item {"active" if i == submissions.page else ""}"><a class="page-link" href="?page={i}&status={status_filter}">{i}</a></li>' for i in submissions.iter_pages()])}
                {f'<li class="page-item {"disabled" if not submissions.has_next else ""}"><a class="page-link" href="?page={submissions.next_num}&status={status_filter}">Next</a></li>' if submissions.has_next else ''}
            </ul>
        </nav>
        ''' if submissions.pages > 1 else ''}
    </div>

    <script>
        function changeStatusFilter(status) {{
            window.location.href = '?status=' + status;
        }}

        function approveSubmission(submissionId) {{
            if (confirm('Approve this draft submission? It will be marked as approved and ready for publication.')) {{
                updateSubmissionStatus(submissionId, 'approved');
            }}
        }}

        function rejectSubmission(submissionId) {{
            const reason = prompt('Reason for rejection (optional):');
            updateSubmissionStatus(submissionId, 'rejected', reason);
        }}

        function unapproveSubmission(submissionId) {{
            if (confirm('Remove approval for this submission?')) {{
                updateSubmissionStatus(submissionId, 'submitted');
            }}
        }}

        function publishAsRFC(submissionId) {{
            const rfcNumber = prompt('Enter RFC number:');
            if (rfcNumber && confirm('Publish as RFC ' + rfcNumber + '?')) {{
                updateSubmissionStatus(submissionId, 'published', null, rfcNumber);
            }}
        }}

        function updateSubmissionStatus(submissionId, status, reason = null, rfcNumber = null) {{
            const data = {{ status: status }};
            if (reason) data.reason = reason;
            if (rfcNumber) data.rfc_number = rfcNumber;

                fetch('/admin/submissions/' + submissionId + '/status', {{
                method: 'POST',
                headers: {{
                    'Content-Type': 'application/json',
                }},
                body: JSON.stringify(data)
            }})
            .then(response => response.json())
            .then(data => {{
                if (data.success) {{
                    location.reload();
                }} else {{
                    alert('Error: ' + data.message);
                }}
            }})
            .catch(error => {{
                console.error('Error:', error);
                alert('Error updating submission status');
            }});
        }}
    </script>
    """

    return BASE_TEMPLATE.format(
        title="Submission Management - MLGH",
        theme=current_theme,
        user_menu=user_menu,
        content=content, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/admin/submissions/<submission_id>/status', methods=['POST'])
@require_role('admin')
def update_submission_status(submission_id):
    data = request.get_json()
    new_status = data.get('status', '')
    reason = data.get('reason', '')
    rfc_number = data.get('rfc_number', '')

    if new_status not in ['submitted', 'approved', 'rejected', 'published']:
        return jsonify({'success': False, 'message': 'Invalid status'}), 400

    submission = Submission.query.filter_by(id=submission_id).first()
    if not submission:
        return jsonify({'success': False, 'message': 'Submission not found'}), 404

    old_status = submission.status
    submission.status = new_status

    # Check for duplicate revision numbers BEFORE approving (for revisions that already have ML number)
    if new_status == 'approved':
        is_revision = getattr(submission, 'is_revision', False)
        revision_num = getattr(submission, 'revision_number', '')
        ml_num = submission.ml_number
        
        if is_revision and revision_num and ml_num:
            # Check if this revision number already exists
            existing_revision = Submission.query.filter(
                Submission.ml_number == ml_num,
                Submission.revision_number == revision_num,
                Submission.status == 'approved',
                Submission.id != submission_id
            ).first()
            
            if existing_revision:
                # Find next available revision number
                all_revisions = Submission.query.filter(
                    Submission.ml_number == ml_num,
                    Submission.status == 'approved',
                    Submission.revision_number.isnot(None)
                ).all()
                
                existing_nums = []
                for rev in all_revisions:
                    try:
                        existing_nums.append(int(rev.revision_number))
                    except (ValueError, TypeError):
                        pass
                
                next_num = 1
                while next_num in existing_nums:
                    next_num += 1
                
                old_revision = submission.revision_number
                submission.revision_number = f"{next_num:02d}"
                app.logger.warning(f"⚠️ Duplicate revision {old_revision} detected for {ml_num}, auto-assigned {submission.revision_number}")

    # Assign ML number when approving
    if new_status == 'approved' and not submission.ml_number:
        # Check if this is a revision
        is_revision = getattr(submission, 'is_revision', False)
        parent_draft_name = getattr(submission, 'parent_draft_name', '')
        
        if is_revision and parent_draft_name:
            # This is a revision - find the parent draft and use its ML number
            parent_submission = Submission.query.filter_by(id=parent_draft_name).first()
            if parent_submission and parent_submission.ml_number:
                # AUTO-ASSIGN next available revision number if duplicate detected
                revision_num = getattr(submission, 'revision_number', '')
                if revision_num:
                    existing_revision = Submission.query.filter(
                        Submission.ml_number == parent_submission.ml_number,
                        Submission.revision_number == revision_num,
                        Submission.status == 'approved',
                        Submission.id != submission_id  # Exclude current submission
                    ).first()
                    
                    if existing_revision:
                        # Find the next available revision number
                        all_revisions = Submission.query.filter(
                            Submission.ml_number == parent_submission.ml_number,
                            Submission.status == 'approved',
                            Submission.revision_number.isnot(None)
                        ).all()
                        
                        # Get all existing revision numbers as integers
                        existing_nums = []
                        for rev in all_revisions:
                            try:
                                existing_nums.append(int(rev.revision_number))
                            except (ValueError, TypeError):
                                pass
                        
                        # Find next available number
                        next_num = 1
                        while next_num in existing_nums:
                            next_num += 1
                        
                        # Update submission with new revision number
                        old_revision = submission.revision_number
                        submission.revision_number = f"{next_num:02d}"
                        
                        app.logger.warning(f"⚠️ Duplicate revision {old_revision} detected for {parent_submission.ml_number}, auto-assigned {submission.revision_number}")
                
                # Use the parent's ML number for the revision
                submission.ml_number = parent_submission.ml_number
                submission.approved_at = datetime.utcnow()
                app.logger.info(f"✅ Revision {submission_id} inherits ML number {parent_submission.ml_number} from parent {parent_draft_name} via admin status update")
            else:
                app.logger.warning(f"⚠️ Parent draft {parent_draft_name} not found or has no ML number, assigning new ML number")
                try:
                    doc_type = getattr(submission, 'doc_type', 'draft') or 'draft'
                    ml_number = get_next_ml_number(doc_type)
                    submission.ml_number = ml_number
                    submission.approved_at = datetime.utcnow()
                    app.logger.info(f"✅ Assigned new ML number {ml_number} to revision {submission_id} via admin status update")
                except Exception as e:
                    app.logger.error(f"❌ Failed to assign ML number to revision {submission_id} via admin status update: {e}")
        else:
            # This is a new draft - assign a new ML number
            try:
                doc_type = getattr(submission, 'doc_type', 'draft') or 'draft'
                ml_number = get_next_ml_number(doc_type)
                submission.ml_number = ml_number
                submission.approved_at = datetime.utcnow()
                app.logger.info(f"✅ Assigned ML number {ml_number} to submission {submission_id} via admin status update")
            except Exception as e:
                app.logger.error(f"❌ Failed to assign ML number to submission {submission_id} via admin status update: {e}")
            # Continue with status change even if ML assignment fails

    if new_status == 'rejected' and reason:
        submission.rejected_at = datetime.utcnow()

    if new_status == 'published':
        # Store RFC number directly in submission table
        if rfc_number:
            try:
                submission.rfc_number = int(rfc_number)
            except (ValueError, TypeError):
                return jsonify({'success': False, 'message': 'Invalid RFC number'}), 400
        # Update ML number to RFC format if it was a draft
        if submission.ml_number and submission.ml_number.startswith('ML-Draft-'):
            # Convert ML-Draft-001 to ML-RFC-001
            draft_num = submission.ml_number.split('-')[-1]
            submission.ml_number = f"ML-RFC-{draft_num}"
        submission.doc_type = 'rfc'

    db.session.commit()

    # Log the action
    admin_user = get_current_user()
    action_details = f"Changed status from {old_status} to {new_status}"
    if reason:
        action_details += f" - Reason: {reason}"
    if rfc_number:
        action_details += f" - Published as RFC {rfc_number}"

    add_to_document_history(f"submission-{submission.id}", "status_changed",
                           admin_user['name'], action_details)

    return jsonify({'success': True, 'message': f'Status updated to {new_status}'})

def get_next_ml_number(doc_type='draft'):
    """Get the next ML number (ML-Draft-001 or ML-RFC-001)
    
    Args:
        doc_type: 'draft' or 'rfc'
    
    Returns:
        str: ML-Draft-001, ML-RFC-001, etc.
    """
    # Find the highest existing ML number for this document type
    prefix = f"ML-{doc_type.capitalize()}-"
    max_ml = db.session.query(db.func.max(Submission.ml_number)).filter(
        Submission.ml_number.like(f"{prefix}%")
    ).scalar()
    
    if max_ml:
        # Extract number from ML-Draft-XXX or ML-RFC-XXX format
        try:
            current_num = int(max_ml.split('-')[-1])
            next_num = current_num + 1
        except (ValueError, IndexError):
            next_num = 1
    else:
        next_num = 1
    
    # Use 3 digits for 1-999, 4 digits for 1000+
    if next_num < 1000:
        return f"{prefix}{next_num:03d}"
    else:
        return f"{prefix}{next_num:04d}"

@app.route('/submit/approve/<submission_id>', methods=['POST'])
@require_role('admin')
def approve_submission(submission_id):
    submission = Submission.query.filter_by(id=submission_id).first()
    if not submission:
        flash('Submission not found', 'error')
        return redirect('/admin/submissions/')

    # Check for duplicate revision numbers FIRST (for revisions that already have ML number)
    is_revision = getattr(submission, 'is_revision', False)
    revision_num = getattr(submission, 'revision_number', '')
    ml_num = submission.ml_number
    
    if is_revision and revision_num and ml_num:
        # Check if this revision number already exists
        existing_revision = Submission.query.filter(
            Submission.ml_number == ml_num,
            Submission.revision_number == revision_num,
            Submission.status == 'approved',
            Submission.id != submission_id
        ).first()
        
        if existing_revision:
            # Find next available revision number
            all_revisions = Submission.query.filter(
                Submission.ml_number == ml_num,
                Submission.status == 'approved',
                Submission.revision_number.isnot(None)
            ).all()
            
            existing_nums = []
            for rev in all_revisions:
                try:
                    existing_nums.append(int(rev.revision_number))
                except (ValueError, TypeError):
                    pass
            
            next_num = 1
            while next_num in existing_nums:
                next_num += 1
            
            old_revision = submission.revision_number
            submission.revision_number = f"{next_num:02d}"
            flash(f'⚠️ Revision {old_revision} already exists. Auto-assigned revision {submission.revision_number} instead.', 'warning')
            app.logger.warning(f"⚠️ Duplicate revision {old_revision} detected for {ml_num}, auto-assigned {submission.revision_number}")

    # Check if this is a revision
    is_revision = getattr(submission, 'is_revision', False)
    parent_draft_name = getattr(submission, 'parent_draft_name', '')
    
    if is_revision and parent_draft_name:
        # This is a revision - find the parent draft and use its ML number
        parent_submission = Submission.query.filter_by(id=parent_draft_name).first()
        if parent_submission and parent_submission.ml_number:
            # AUTO-ASSIGN next available revision number if duplicate detected
            revision_num = getattr(submission, 'revision_number', '')
            if revision_num:
                existing_revision = Submission.query.filter(
                    Submission.ml_number == parent_submission.ml_number,
                    Submission.revision_number == revision_num,
                    Submission.status == 'approved',
                    Submission.id != submission_id  # Exclude current submission
                ).first()
                
                if existing_revision:
                    # Find the next available revision number
                    all_revisions = Submission.query.filter(
                        Submission.ml_number == parent_submission.ml_number,
                        Submission.status == 'approved',
                        Submission.revision_number.isnot(None)
                    ).all()
                    
                    # Get all existing revision numbers as integers
                    existing_nums = []
                    for rev in all_revisions:
                        try:
                            existing_nums.append(int(rev.revision_number))
                        except (ValueError, TypeError):
                            pass
                    
                    # Find next available number
                    next_num = 1
                    while next_num in existing_nums:
                        next_num += 1
                    
                    # Update submission with new revision number
                    old_revision = submission.revision_number
                    submission.revision_number = f"{next_num:02d}"
                    
                    flash(f'⚠️ Revision {old_revision} already exists. Auto-assigned revision {submission.revision_number} instead.', 'warning')
                    app.logger.warning(f"⚠️ Duplicate revision {old_revision} detected for {parent_submission.ml_number}, auto-assigned {submission.revision_number}")
            
            # Use the parent's ML number for the revision
            # Use the parent's ML number for the revision
            submission.ml_number = parent_submission.ml_number
            app.logger.info(f"✅ Revision {submission_id} inherits ML number {parent_submission.ml_number} from parent {parent_draft_name}")
        else:
            app.logger.warning(f"⚠️ Parent draft {parent_draft_name} not found or has no ML number")
            # Assign a new ML number as fallback
            if not submission.ml_number:
                try:
                    doc_type = getattr(submission, 'doc_type', 'draft') or 'draft'
                    ml_number = get_next_ml_number(doc_type)
                    submission.ml_number = ml_number
                    app.logger.info(f"✅ Assigned new ML number {ml_number} to revision {submission_id}")
                except Exception as e:
                    app.logger.error(f"❌ Failed to assign ML number to revision {submission_id}: {e}")
    else:
        # This is a new draft - assign ML number if not already assigned
        if not submission.ml_number:
            try:
                doc_type = getattr(submission, 'doc_type', 'draft') or 'draft'
                ml_number = get_next_ml_number(doc_type)
                submission.ml_number = ml_number
                app.logger.info(f"✅ Assigned ML number {ml_number} to submission {submission_id}")
            except Exception as e:
                app.logger.error(f"❌ Failed to assign ML number to submission {submission_id}: {e}")
                # Continue with approval even if ML assignment fails
    
    submission.status = 'approved'
    submission.approved_at = datetime.utcnow()
    
    try:
        db.session.commit()
        app.logger.info(f"✅ Successfully approved submission {submission_id}")
    except Exception as e:
        app.logger.error(f"❌ Failed to commit approval for submission {submission_id}: {e}")
        db.session.rollback()
        flash(f'Failed to approve submission: {str(e)}', 'error')
        return redirect('/admin/submissions/')

    # Log the action
    admin_user = get_current_user()
    action_desc = f"Approved revision {submission.revision_number} of {parent_draft_name}" if is_revision else f"Approved submission: {submission.title}"
    add_to_document_history(f"submission-{submission.id}", "approved", admin_user['name'], action_desc)

    flash_msg = f'Revision {submission.id} approved successfully! ML number: {submission.ml_number}' if is_revision else f'Submission {submission.id} approved successfully! Assigned ML number: {submission.ml_number}'
    flash(flash_msg, 'success')
    return redirect(f'/submit/status/{submission_id}/')

@app.route('/submit/reject/<submission_id>', methods=['POST'])
@require_role('admin')
def reject_submission(submission_id):
    submission = Submission.query.filter_by(id=submission_id).first()
    if not submission:
        flash('Submission not found', 'error')
        return redirect('/admin/submissions/')

    submission.status = 'rejected'
    submission.rejected_at = datetime.utcnow()
    db.session.commit()

    # Log the action
    admin_user = get_current_user()
    add_to_document_history(f"submission-{submission.id}", "rejected", admin_user['name'],
                           f"Rejected submission: {submission.title}")

    flash(f'Submission {submission.id} rejected!', 'warning')
    return redirect(f'/submit/status/{submission_id}/')

@app.route('/view/<submission_id>')
@require_auth
def view_submission(submission_id):
    """View a submission file inline (for PDFs and other viewable files)"""
    submission = Submission.query.filter_by(id=submission_id).first()
    if not submission:
        return "Submission not found", 404

    # Check if user owns this submission or is admin
    current_user = get_current_user()
    if submission.submitted_by != current_user['name'] and current_user.get('role') not in ['admin', 'editor']:
        return "Access denied", 403

    if not submission.file_path or not os.path.exists(submission.file_path):
        return "File not found", 404

    # Serve file inline (not as attachment) for viewing in browser
    return send_file(submission.file_path, as_attachment=False, download_name=submission.filename)

@app.route('/download/<submission_id>')
@require_auth
def download_submission(submission_id):
    """Download a submission file"""
    submission = Submission.query.filter_by(id=submission_id).first()
    if not submission:
        return "Submission not found", 404

    # Check if user owns this submission or is admin
    current_user = get_current_user()
    if submission.submitted_by != current_user['name'] and current_user.get('role') not in ['admin', 'editor']:
        return "Access denied", 403

    if not submission.file_path or not os.path.exists(submission.file_path):
        return "File not found", 404

    return send_file(submission.file_path, as_attachment=True, download_name=submission.filename)

@app.route('/admin/analytics/')
@require_role('admin')
def admin_analytics():
    user_menu = generate_user_menu()
    current_theme = get_current_user().get('theme', 'dark')

    # Most active drafts (recent submissions)
    active_drafts = Submission.query.order_by(Submission.submitted_at.desc()).limit(20).all()

    # Most active users (by recent logins and submissions)
    active_users = User.query.order_by(User.last_login.desc()).limit(20).all()

    # User role distribution
    role_stats = db.session.query(User.role, db.func.count(User.id)).group_by(User.role).all()
    role_data = {role: count for role, count in role_stats}

    # Submission status distribution
    status_stats = db.session.query(Submission.status, db.func.count(Submission.id)).group_by(Submission.status).all()
    status_data = {status: count for status, count in status_stats}

    # Recent activity (last 30 days)
    thirty_days_ago = datetime.utcnow() - timedelta(days=30)
    recent_users = User.query.filter(User.created_at >= thirty_days_ago).count()
    recent_submissions = Submission.query.filter(Submission.submitted_at >= thirty_days_ago).count()

    # Build active drafts table
    draft_rows = ""
    for i, draft in enumerate(active_drafts, 1):
        draft_rows += f"""
        <tr>
            <td>{i}</td>
            <td>
                <a href="/doc/draft/{draft.id}/" class="text-decoration-none">
                    {draft.title[:60]}{'...' if len(draft.title) > 60 else ''}
                </a>
            </td>
            <td>{', '.join(draft.authors[:2])}{'...' if len(draft.authors) > 2 else ''}</td>
            <td>{draft.group or 'None'}</td>
            <td>{draft.submitted_at.strftime('%Y-%m-%d')}</td>
            <td><span class="badge bg-{ 'warning' if draft.status == 'submitted' else 'success' if draft.status == 'approved' else 'danger' if draft.status == 'rejected' else 'info'}">{draft.status}</span></td>
        </tr>
        """

    # Build active users table
    user_rows = ""
    for i, user in enumerate(active_users, 1):
        user_rows += f"""
        <tr>
            <td>{i}</td>
            <td>
                <strong>{user.name}</strong><br>
                <small class="text-muted">@{user.username}</small>
            </td>
            <td>{user.email}</td>
            <td><span class="badge bg-{ 'danger' if user.role == 'admin' else 'warning' if user.role == 'editor' else 'secondary'}">{user.role.title()}</span></td>
            <td>{user.created_at.strftime('%Y-%m-%d')}</td>
            <td>{user.last_login.strftime('%Y-%m-%d %H:%M') if user.last_login else 'Never'}</td>
        </tr>
        """

    content = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/admin/">Admin Dashboard</a></li>
                <li class="breadcrumb-item active">Analytics</li>
            </ol>
        </nav>

        <h1 class="mb-4">Analytics Dashboard</h1>

        <!-- Overview Stats -->
        <div class="row mb-4">
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-info">{recent_users}</h4>
                        <p class="mb-0 small">New Users (30 days)</p>
        </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-success">{recent_submissions}</h4>
                        <p class="mb-0 small">New Submissions (30 days)</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-primary">{len(active_drafts)}</h4>
                        <p class="mb-0 small">Total Submissions</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-warning">{len(active_users)}</h4>
                        <p class="mb-0 small">Registered Users</p>
                    </div>
                </div>
            </div>
        </div>

        <div class="row">
            <!-- Most Active Drafts -->
            <div class="col-md-6">
                <div class="card">
                    <div class="card-header">
                        <h5>Most Active Drafts</h5>
                        <small class="text-muted">Recent submissions and activity</small>
                    </div>
                    <div class="card-body p-0">
                        <div class="table-responsive">
                            <table class="table table-hover mb-0">
                                <thead class="table-light">
                                    <tr>
                                        <th>#</th>
                                        <th>Title</th>
                                        <th>Authors</th>
                                        <th>Group</th>
                                        <th>Date</th>
                                        <th>Status</th>
                                    </tr>
                                </thead>
                                <tbody>
                                    {draft_rows}
                                </tbody>
                            </table>
                        </div>
                    </div>
                </div>
            </div>

            <!-- Most Active Users -->
            <div class="col-md-6">
                <div class="card">
                    <div class="card-header">
                        <h5>Most Active Users</h5>
                        <small class="text-muted">Users by recent login activity</small>
                    </div>
                    <div class="card-body p-0">
                        <div class="table-responsive">
                            <table class="table table-hover mb-0">
                                <thead class="table-light">
                                    <tr>
                                        <th>#</th>
                                        <th>Name</th>
                                        <th>Email</th>
                                        <th>Role</th>
                                        <th>Joined</th>
                                        <th>Last Login</th>
                                    </tr>
                                </thead>
                                <tbody>
                                    {user_rows}
                                </tbody>
                            </table>
                        </div>
                    </div>
                </div>
            </div>
        </div>

        <!-- Distribution Charts (Text-based) -->
        <div class="row mt-4">
            <div class="col-md-6">
                <div class="card">
                    <div class="card-header">
                        <h5>User Role Distribution</h5>
                    </div>
                    <div class="card-body">
                        <div class="mb-3">
                            <strong>Admin:</strong> {role_data.get('admin', 0)} users
                            <div class="progress mb-2">
                                <div class="progress-bar bg-danger" style="width: {(role_data.get('admin', 0) / max(1, sum(role_data.values()))) * 100:.1f}%"></div>
                            </div>
                        </div>
                        <div class="mb-3">
                            <strong>Editor:</strong> {role_data.get('editor', 0)} users
                            <div class="progress mb-2">
                                <div class="progress-bar bg-warning" style="width: {(role_data.get('editor', 0) / max(1, sum(role_data.values()))) * 100:.1f}%"></div>
                            </div>
                        </div>
                        <div class="mb-3">
                            <strong>User:</strong> {role_data.get('user', 0)} users
                            <div class="progress mb-2">
                                <div class="progress-bar bg-secondary" style="width: {(role_data.get('user', 0) / max(1, sum(role_data.values()))) * 100:.1f}%"></div>
                            </div>
                        </div>
                    </div>
                </div>
            </div>

            <div class="col-md-6">
                <div class="card">
                    <div class="card-header">
                        <h5>Submission Status Distribution</h5>
                    </div>
                    <div class="card-body">
                        <div class="mb-3">
                            <strong>Submitted:</strong> {status_data.get('submitted', 0)}
                            <div class="progress mb-2">
                                <div class="progress-bar bg-warning" style="width: {(status_data.get('submitted', 0) / max(1, sum(status_data.values()))) * 100:.1f}%"></div>
                            </div>
                        </div>
                        <div class="mb-3">
                            <strong>Approved:</strong> {status_data.get('approved', 0)}
                            <div class="progress mb-2">
                                <div class="progress-bar bg-success" style="width: {(status_data.get('approved', 0) / max(1, sum(status_data.values()))) * 100:.1f}%"></div>
                            </div>
                        </div>
                        <div class="mb-3">
                            <strong>Published:</strong> {status_data.get('published', 0)}
                            <div class="progress mb-2">
                                <div class="progress-bar bg-info" style="width: {(status_data.get('published', 0) / max(1, sum(status_data.values()))) * 100:.1f}%"></div>
                            </div>
                        </div>
                        <div class="mb-3">
                            <strong>Rejected:</strong> {status_data.get('rejected', 0)}
                            <div class="progress mb-2">
                                <div class="progress-bar bg-danger" style="width: {(status_data.get('rejected', 0) / max(1, sum(status_data.values()))) * 100:.1f}%"></div>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        </div>
        """
    
    return BASE_TEMPLATE.format(
        title="Analytics - MLGH",
        theme=current_theme,
        user_menu=user_menu,
        content=content, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/admin/chairs/')
@require_auth
def admin_chairs():
    current_user = get_current_user()
    current_theme = session.get('theme', 'dark')

    # Generate user menu
    user_menu = generate_user_menu()

    # Get statistics from database (same source as workgroup page "Add Coord")
    all_chairs = WorkingGroupChair.query.all()
    total_chairs = len(all_chairs)
    approved_chairs = sum(1 for c in all_chairs if c.approved)
    pending_chairs = total_chairs - approved_chairs

    # Build chair list from database (Approve only for Pending; Active coordinators only get Delete)
    chair_list = ""
    for chair in all_chairs:
        status_badge = 'success' if chair.approved else 'warning'
        status_text = 'Active' if chair.approved else 'Pending'
        set_at_str = chair.set_at.strftime('%Y-%m-%d') if chair.set_at else 'N/A'
        if chair.approved:
            actions = f'<a href="/admin/chairs/{chair.id}/delete" class="btn btn-sm btn-outline-danger" onclick="return confirm(\'Remove this coordinator?\')">Delete</a>'
        else:
            actions = f'<a href="/admin/chairs/{chair.id}/approve" class="btn btn-sm btn-outline-success" onclick="return confirm(\'Approve this coordinator?\')">Approve</a> <a href="/admin/chairs/{chair.id}/delete" class="btn btn-sm btn-outline-danger" onclick="return confirm(\'Delete this coordinator?\')">Delete</a>'
        chair_list += f"""
        <tr>
            <td>{chair.chair_name}</td>
            <td>N/A</td>
            <td><code>{chair.group_acronym}</code></td>
            <td><span class="badge bg-{status_badge}">{status_text}</span></td>
            <td>{set_at_str}</td>
            <td>{actions}</td>
        </tr>
        """

    # Build pending coordinator request rows (before content f-string so variable is defined)
    pending_coord_requests = CoordinatorRequest.query.filter_by(status='pending').order_by(CoordinatorRequest.requested_at.desc()).all()
    coord_request_rows = ""
    for req in pending_coord_requests:
        req_at = req.requested_at.strftime('%Y-%m-%d %H:%M') if req.requested_at else ''
        coord_request_rows += f"""
        <tr>
            <td>{req.display_name or req.username}</td>
            <td><code>{req.username}</code></td>
            <td><code>{req.group_acronym}</code></td>
            <td>{req_at}</td>
            <td>
                <a href="/admin/coordinator_requests/{req.id}/approve" class="btn btn-sm btn-success">Approve</a>
                <a href="/admin/coordinator_requests/{req.id}/reject" class="btn btn-sm btn-outline-danger" onclick="return confirm('Reject this request?')">Reject</a>
            </td>
        </tr>
        """
    if not coord_request_rows:
        coord_request_rows = '<tr><td colspan="5" class="text-center text-muted py-3">No pending coordinator requests.</td></tr>'

    content = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/admin/">Admin Dashboard</a></li>
                <li class="breadcrumb-item active">Coordinator Management</li>
            </ol>
        </nav>

        <div class="mb-4">
            <h1 class="mb-1">Coordinator Management</h1>
            <p class="text-muted mb-0">Manage workgroup coordinators. Add coordinators from <a href="/admin/users/">User Management</a> or <a href="/person/">People</a> (admin actions).</p>
        </div>

        <!-- Statistics Cards -->
        <div class="row mb-4">
            <div class="col-md-4">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-primary">{total_chairs}</h4>
                        <small class="text-muted">Total Coordinators</small>
                    </div>
                </div>
            </div>
            <div class="col-md-4">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-success">{approved_chairs}</h4>
                        <small class="text-muted">Active Coordinators</small>
                    </div>
                </div>
            </div>
            <div class="col-md-4">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-warning">{pending_chairs}</h4>
                        <small class="text-muted">Pending Approval</small>
                    </div>
                </div>
            </div>
        </div>

        <!-- Pending coordinator requests (user-requested role; we have their id) -->
        <div class="card mb-4">
            <div class="card-header">
                <h5 class="mb-0">Pending coordinator requests</h5>
                <small class="text-muted">Users requested coordinator role; approve to grant.</small>
            </div>
            <div class="card-body p-0">
                <div class="table-responsive">
                    <table class="table table-hover mb-0">
                        <thead class="table-light">
                            <tr>
                                <th>Display name</th>
                                <th>Username (id)</th>
                                <th>Workgroup</th>
                                <th>Requested</th>
                                <th>Actions</th>
                            </tr>
                        </thead>
                        <tbody>
                            {coord_request_rows}
                        </tbody>
                    </table>
                </div>
            </div>
        </div>

        <!-- Coordinators Table -->
        <div class="card">
            <div class="card-header">
                <h5 class="mb-0">Workgroup Coordinators</h5>
            </div>
            <div class="card-body p-0">
                <div class="table-responsive">
                    <table class="table table-hover mb-0">
                        <thead class="table-light">
                            <tr>
                                <th>Name</th>
                                <th>Email</th>
                                <th>Group</th>
                                <th>Status</th>
                                <th>Added</th>
                                <th>Actions</th>
                            </tr>
                        </thead>
                        <tbody>
                            {chair_list if chair_list else '<tr><td colspan="6" class="text-center text-muted py-4">No coordinators yet. Add from User Management or People (admin).</td></tr>'}
                        </tbody>
                    </table>
                </div>
            </div>
        </div>

        <div class="mt-4">
            <a href="/admin/member_requests/" class="btn btn-outline-primary">View member requests</a>
        </div>
    </div>
    """

    return BASE_TEMPLATE.format(
        title="Coordinator Management - MLGH",
        theme=current_theme,
        user_menu=user_menu,
        content=content, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/admin/users/<int:user_id>/add-coordinator', methods=['GET', 'POST'])
@require_role('admin')
def add_coordinator_for_user(user_id):
    """Add an existing user as coordinator for a workgroup (by user id, workgroup from list)."""
    user_menu = generate_user_menu()
    current_theme = get_current_user().get('theme', 'dark')
    target = User.query.get(user_id)
    if not target:
        flash('User not found', 'error')
        return redirect('/admin/users/')
    display_name = target.name or target.displayName or target.oauthName or target.username

    if request.method == 'POST':
        group_acronym = request.form.get('group_acronym', '').strip()
        if not group_acronym:
            flash('Please select a workgroup', 'error')
        else:
            # Check valid group
            if not any(g['acronym'] == group_acronym for g in GROUPS):
                flash('Invalid workgroup', 'error')
            else:
                existing = WorkingGroupChair.query.filter_by(
                    group_acronym=group_acronym,
                    user_id=user_id
                ).first()
                if existing:
                    flash(f'{display_name} is already a coordinator for {group_acronym}', 'error')
                else:
                    chair = WorkingGroupChair(
                        group_acronym=group_acronym,
                        chair_name=display_name,
                        user_id=user_id,
                        approved=True
                    )
                    db.session.add(chair)
                    db.session.commit()
                    flash(f'Added {display_name} as coordinator for {group_acronym}', 'success')
                    return redirect('/admin/chairs/')

    group_options = ''.join(
        f'<option value="{g["acronym"]}">{g["acronym"]} – {g.get("name", g["acronym"])}</option>'
        for g in GROUPS
    )
    content = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/admin/">Admin Dashboard</a></li>
                <li class="breadcrumb-item"><a href="/admin/users/">User Management</a></li>
                <li class="breadcrumb-item"><a href="/admin/chairs/">Coordinator Management</a></li>
                <li class="breadcrumb-item active">Add as coordinator</li>
            </ol>
        </nav>
        <div class="row justify-content-center">
            <div class="col-md-6">
                <div class="card">
                    <div class="card-header">
                        <h5 class="mb-0">Add as coordinator</h5>
                    </div>
                    <div class="card-body">
                        <p class="mb-3">User: <strong>{display_name}</strong> (@{target.username})</p>
                        <form method="POST">
                            <div class="mb-3">
                                <label for="group_acronym" class="form-label">Workgroup *</label>
                                <select class="form-select" id="group_acronym" name="group_acronym" required>
                                    <option value="">Select workgroup</option>
                                    {group_options}
                                </select>
                            </div>
                            <div class="d-flex gap-2">
                                <button type="submit" class="btn btn-primary">Add as coordinator</button>
                                <a href="/admin/users/" class="btn btn-secondary">Cancel</a>
                            </div>
                        </form>
                    </div>
                </div>
            </div>
        </div>
    </div>
    """
    return render_page("Add as coordinator - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/admin/coordinator_requests/<int:req_id>/approve')
@require_role('admin')
def approve_coordinator_request(req_id):
    """Approve a user's coordinator request; creates WorkingGroupChair with their user id."""
    req = CoordinatorRequest.query.get(req_id)
    if not req or req.status != 'pending':
        flash('Request not found or already handled', 'error')
        return redirect('/admin/chairs/')
    admin_user = get_current_user()
    # Create coordinator record (linked to user id)
    chair = WorkingGroupChair(
        group_acronym=req.group_acronym,
        chair_name=req.display_name or req.username,
        user_id=req.user_id,
        approved=True
    )
    db.session.add(chair)
    req.status = 'approved'
    req.reviewed_at = datetime.utcnow()
    req.reviewed_by = admin_user.get('name') or admin_user.get('username')
    db.session.commit()
    flash(f'Coordinator request approved: {req.display_name or req.username} for {req.group_acronym}', 'success')
    return redirect('/admin/chairs/')

@app.route('/admin/coordinator_requests/<int:req_id>/reject')
@require_role('admin')
def reject_coordinator_request(req_id):
    req = CoordinatorRequest.query.get(req_id)
    if not req or req.status != 'pending':
        flash('Request not found or already handled', 'error')
        return redirect('/admin/chairs/')
    admin_user = get_current_user()
    req.status = 'rejected'
    req.reviewed_at = datetime.utcnow()
    req.reviewed_by = admin_user.get('name') or admin_user.get('username')
    db.session.commit()
    flash(f'Coordinator request rejected: {req.display_name or req.username}', 'warning')
    return redirect('/admin/chairs/')

@app.route('/admin/chairs/<int:chair_id>/approve')
@require_auth
def approve_chair(chair_id):
    chair = WorkingGroupChair.query.get(chair_id)
    if chair:
        chair.approved = True
        db.session.commit()
        flash('Coordinator approved successfully', 'success')
    else:
        flash('Coordinator not found', 'error')
    return redirect('/admin/chairs/')

@app.route('/admin/chairs/<int:chair_id>/delete')
@require_auth
def delete_chair(chair_id):
    chair = WorkingGroupChair.query.get(chair_id)
    if chair:
        db.session.delete(chair)
        db.session.commit()
        flash('Coordinator deleted successfully', 'success')
    else:
        flash('Coordinator not found', 'error')
    return redirect('/admin/chairs/')

# ============================================================================
# Admin Dashboards for Projects/Workgroups/Guilds/Roles/Badges
# ============================================================================

@app.route('/admin/projects/')
@require_role('admin')
def admin_projects():
    """Admin dashboard for managing projects"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    # Use same DB counts as admin dashboard so badge matches "Review now" alert
    pending_projects = Project.query.filter_by(approval_status='pending').order_by(Project.last_activity.desc()).all()
    approved_projects = Project.query.filter_by(approval_status='approved').order_by(Project.last_activity.desc()).all()
    rejected_projects = Project.query.filter_by(approval_status='rejected').order_by(Project.last_activity.desc()).all()
    pending_count = len(pending_projects)
    approved_count = len(approved_projects)
    rejected_count = len(rejected_projects)
    
    def _escape(s):
        if not s:
            return ''
        return str(s).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br>')
    
    def _project_row_html(p, show_actions=False):
        mission = ('<p class="mb-2">' + _escape(p.mission) + '</p>') if p.mission else ''
        wg_count = Workgroup.query.filter_by(project_id=p.id).count()
        created = p.created_at.strftime('%x') if p.created_at else ''
        safe_id = str(p.id).replace('&', '&amp;').replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')
        actions = ''
        if show_actions:
            actions = '''
                <div class="btn-group-vertical ms-3">
                    <button type="button" class="btn btn-sm btn-success btn-approve-project" data-project-id="''' + safe_id + '''">
                        <i class="fas fa-check me-1"></i>Approve
                    </button>
                    <button type="button" class="btn btn-sm btn-danger btn-reject-project" data-project-id="''' + safe_id + '''">
                        <i class="fas fa-times me-1"></i>Reject
                    </button>
                </div>
            '''
        return '''
            <div class="list-group-item">
                <div class="d-flex justify-content-between align-items-start">
                    <div class="flex-grow-1">
                        <h5><a href="/projects/''' + (p.slug or '') + '''/" target="_blank">''' + _escape(p.name) + '''</a></h5>
                        ''' + mission + '''
                        <p class="mb-2">''' + _escape(p.description or 'No description') + '''</p>
                        <small class="text-muted">
                            Created: ''' + created + ''' | Status: ''' + _escape(p.status) + ''' | Workgroups: ''' + str(wg_count) + '''
                        </small>
                    </div>
                    ''' + actions + '''
                </div>
            </div>
        '''
    
    def _list_html(projects, show_actions=False):
        if not projects:
            return '<div class="alert alert-info">No projects in this category</div>'
        parts = ['<div class="list-group">']
        for p in projects:
            parts.append(_project_row_html(p, show_actions))
        parts.append('</div>')
        return ''.join(parts)
    
    content = """
    <div class="container mt-4" id="manage-projects-container" data-server-pending=""" + str(pending_count) + """ data-server-approved=""" + str(approved_count) + """ data-server-rejected=""" + str(rejected_count) + """>
        <h1 class="mb-4">Manage Layers</h1>
        <div id="project-load-error" class="alert alert-danger d-none" role="alert"></div>
        
        <ul class="nav nav-tabs mb-4" id="projectTabs" role="tablist">
            <li class="nav-item" role="presentation">
                <button class="nav-link active" id="pending-tab" data-bs-toggle="tab" data-bs-target="#pending" type="button">
                    Pending Approval <span class="badge bg-warning ms-2" id="pending-count">""" + str(pending_count) + """</span>
                </button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="approved-tab" data-bs-toggle="tab" data-bs-target="#approved" type="button">
                    Approved <span class="badge bg-success ms-2" id="approved-count">""" + str(approved_count) + """</span>
                </button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="rejected-tab" data-bs-toggle="tab" data-bs-target="#rejected" type="button">
                    Rejected <span class="badge bg-danger ms-2" id="rejected-count">""" + str(rejected_count) + """</span>
                </button>
            </li>
        </ul>
        
        <div class="tab-content" id="projectTabContent">
            <div class="tab-pane fade show active" id="pending">
                <div id="pending-projects">""" + _list_html(pending_projects, show_actions=True) + """</div>
            </div>
            <div class="tab-pane fade" id="approved">
                <div id="approved-projects">""" + _list_html(approved_projects, show_actions=False) + """</div>
            </div>
            <div class="tab-pane fade" id="rejected">
                <div id="rejected-projects">""" + _list_html(rejected_projects, show_actions=False) + """</div>
            </div>
        </div>
    </div>
    
    <script>
    async function loadProjects() {
        const errEl = document.getElementById('project-load-error');
        const container = document.getElementById('manage-projects-container');
        const serverPending = container ? parseInt(container.getAttribute('data-server-pending') || '0', 10) : 0;
        try {
            const response = await fetch('/api/projects/?_t=' + Date.now(), { credentials: 'include', cache: 'no-store' });
            if (!response.ok) {
                throw new Error('API returned ' + response.status);
            }
            const data = await response.json();
            const projects = Array.isArray(data.projects) ? data.projects : [];
            function approvalStatus(p) { return (p && p.approval_status != null) ? String(p.approval_status).toLowerCase() : ''; }
            const pending = projects.filter(p => approvalStatus(p) === 'pending');
            const approved = projects.filter(p => approvalStatus(p) === 'approved');
            const rejected = projects.filter(p => approvalStatus(p) === 'rejected');
            
            document.getElementById('pending-count').textContent = pending.length;
            document.getElementById('approved-count').textContent = approved.length;
            document.getElementById('rejected-count').textContent = rejected.length;
            
            if (serverPending > 0 && pending.length === 0) {
                document.getElementById('pending-projects').innerHTML =
                    '<div class="alert alert-warning">The list from the server did not load. Showing server-rendered list below. <a href="#" onclick="loadProjects(); return false;">Refresh the list</a>.</div>' +
                    document.getElementById('pending-projects').innerHTML;
            } else {
                displayProjects('pending-projects', pending, true);
            }
            displayProjects('approved-projects', approved, false);
            displayProjects('rejected-projects', rejected, false);
            if (errEl) { errEl.classList.add('d-none'); errEl.textContent = ''; }
        } catch (error) {
            console.error('Error loading projects:', error);
            const msg = 'Failed to load projects: ' + (error.message || 'Please refresh or check console.');
            const pendingEl = document.getElementById('pending-projects');
            if (pendingEl) {
                pendingEl.innerHTML = '<div class="alert alert-danger">' + msg + (serverPending > 0 ? ' The server reports ' + serverPending + ' pending project(s).' : '') + ' <a href="#" onclick="loadProjects(); return false;">Try again</a> or reload the page.</div>';
            }
            if (errEl) {
                errEl.textContent = msg;
                errEl.classList.remove('d-none');
            }
        }
    }
    
    function displayProjects(containerId, projects, showActions) {
        const container = document.getElementById(containerId);
        
        if (projects.length === 0) {
            container.innerHTML = '<div class="alert alert-info">No projects in this category</div>';
            return;
        }
        
        function escapeHtml(text) {
            if (!text) return '';
            return String(text).replace(/&/g, '&amp;').replace(/</g, '&lt;').replace(/>/g, '&gt;').replace(/\\n/g, '<br>');
        }
        
        let html = '<div class="list-group">';
        projects.forEach(project => {
            const missionHtml = project.mission ? '<p class="mb-2">' + escapeHtml(project.mission) + '</p>' : '';
            const descHtml = '<p class="mb-2">' + escapeHtml(project.description || 'No description') + '</p>';
            const actionsHtml = showActions ? 
                '<div class="btn-group-vertical ms-3">' +
                    '<button type="button" class="btn btn-sm btn-success btn-approve-project" data-project-id="' + project.id + '">' +
                        '<i class="fas fa-check me-1"></i>Approve' +
                    '</button>' +
                    '<button type="button" class="btn btn-sm btn-danger btn-reject-project" data-project-id="' + project.id + '">' +
                        '<i class="fas fa-times me-1"></i>Reject' +
                    '</button>' +
                '</div>' : '';
            
            html += '<div class="list-group-item">' +
                '<div class="d-flex justify-content-between align-items-start">' +
                    '<div class="flex-grow-1">' +
                        '<h5><a href="/projects/' + project.slug + '/" target="_blank">' + escapeHtml(project.name) + '</a></h5>' +
                        missionHtml +
                        descHtml +
                        '<small class="text-muted">' +
                            'Created: ' + new Date(project.created_at).toLocaleDateString() + ' | ' +
                            'Status: ' + project.status + ' | ' +
                            'Workgroups: ' + (project.workgroups_count || 0) +
                        '</small>' +
                    '</div>' +
                    actionsHtml +
                '</div>' +
            '</div>';
        });
        html += '</div>';
        
        container.innerHTML = html;
    }
    
    async function approveProject(projectId) {
        if (!confirm('Approve this project?')) return;
        
        try {
            const response = await fetch(`/api/projects/${projectId}/approve/`, {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                credentials: 'include',
                body: JSON.stringify({action: 'approve'})
            });
            
            if (response.ok) {
                alert('Layer approved successfully');
                window.location.reload();
            } else {
                const data = await response.json();
                alert('Error: ' + (data.error || 'Failed to approve'));
            }
        } catch (error) {
            console.error('Error:', error);
            alert('Error approving project');
        }
    }
    
    async function rejectProject(projectId) {
        const note = prompt('Reason for rejection (optional):');
        if (note === null) return;
        
        try {
            const response = await fetch(`/api/projects/${projectId}/approve/`, {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                credentials: 'include',
                body: JSON.stringify({action: 'reject', note: note})
            });
            
            if (response.ok) {
                alert('Layer rejected');
                window.location.reload();
            } else {
                const data = await response.json();
                alert('Error: ' + (data.error || 'Failed to reject'));
            }
        } catch (error) {
            console.error('Error:', error);
            alert('Error rejecting project');
        }
    }
    
    // Projects are server-rendered; no API load on page init
    // Event delegation for approve/reject buttons
    const manageContainer = document.getElementById('manage-projects-container');
    if (manageContainer) {
        manageContainer.addEventListener('click', function(e) {
            const approveBtn = e.target.closest('.btn-approve-project');
            const rejectBtn = e.target.closest('.btn-reject-project');
            if (approveBtn) {
                e.preventDefault();
                approveProject(approveBtn.getAttribute('data-project-id'));
            }
            if (rejectBtn) {
                e.preventDefault();
                rejectProject(rejectBtn.getAttribute('data-project-id'));
            }
        });
    }
    </script>
    """
    
    return render_page("Admin: Manage Layers - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/admin/workgroups/')
@require_role('admin')
def admin_workgroups():
    """Admin dashboard for managing workgroups"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    
    content = """
    <div class="container mt-4">
        <h1 class="mb-4">Manage Workgroups</h1>
        
        <ul class="nav nav-tabs mb-4" id="workgroupTabs" role="tablist">
            <li class="nav-item" role="presentation">
                <button class="nav-link active" id="pending-tab" data-bs-toggle="tab" data-bs-target="#pending" type="button">
                    Pending Approval <span class="badge bg-warning ms-2" id="pending-count">0</span>
                </button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="approved-tab" data-bs-toggle="tab" data-bs-target="#approved" type="button">
                    Approved <span class="badge bg-success ms-2" id="approved-count">0</span>
                </button>
            </li>
        </ul>
        
        <div class="tab-content" id="workgroupTabContent">
            <div class="tab-pane fade show active" id="pending">
                <div id="pending-workgroups"></div>
            </div>
            <div class="tab-pane fade" id="approved">
                <div id="approved-workgroups"></div>
            </div>
        </div>
    </div>
    
    <script>
    async function loadWorkgroups() {
        try {
            // Load all projects first
            const projectsResp = await fetch('/api/projects/');
            const projectsData = await projectsResp.json();
            
            let allWorkgroups = [];
            
            // Load workgroups from all projects
            for (const project of projectsData.projects) {
                const wgResp = await fetch(`/api/projects/${project.id}/workgroups/`);
                const wgData = await wgResp.json();
                
                // Add project info to each workgroup
                wgData.workgroups.forEach(wg => {
                    wg.project_name = project.name;
                    wg.project_slug = project.slug;
                    allWorkgroups.push(wg);
                });
            }
            
            const pending = allWorkgroups.filter(wg => wg.approval_status === 'pending');
            const approved = allWorkgroups.filter(wg => wg.approval_status === 'approved');
            
            document.getElementById('pending-count').textContent = pending.length;
            document.getElementById('approved-count').textContent = approved.length;
            
            displayWorkgroups('pending-workgroups', pending, true);
            displayWorkgroups('approved-workgroups', approved, false);
        } catch (error) {
            console.error('Error loading workgroups:', error);
        }
    }
    
    function displayWorkgroups(containerId, workgroups, showActions) {
        const container = document.getElementById(containerId);
        
        if (workgroups.length === 0) {
            container.innerHTML = '<div class="alert alert-info">No workgroups in this category</div>';
            return;
        }
        
        let html = '<div class="list-group">';
        workgroups.forEach(wg => {
            html += `
                <div class="list-group-item">
                    <div class="d-flex justify-content-between align-items-start">
                        <div class="flex-grow-1">
                            <h5><a href="/workgroups/${wg.slug}/" target="_blank">${wg.name}</a></h5>
                            <p class="mb-2">${wg.description || 'No description'}</p>
                            <small class="text-muted">
                                Layer: <a href="/projects/${wg.project_slug}/" target="_blank">${wg.project_name}</a> | 
                                Created: ${new Date(wg.created_at).toLocaleDateString()} | 
                                Status: ${wg.status}
                            </small>
                        </div>
                        ${showActions ? `
                            <div class="btn-group-vertical ms-3">
                                <button class="btn btn-sm btn-success" onclick="approveWorkgroup('${wg.id}')">
                                    <i class="fas fa-check me-1"></i>Approve
                                </button>
                                <button class="btn btn-sm btn-danger" onclick="rejectWorkgroup('${wg.id}')">
                                    <i class="fas fa-times me-1"></i>Reject
                                </button>
                            </div>
                        ` : ''}
                    </div>
                </div>
            `;
        });
        html += '</div>';
        
        container.innerHTML = html;
    }
    
    async function approveWorkgroup(workgroupId) {
        if (!confirm('Approve this workgroup?')) return;
        
        try {
            const response = await fetch(`/api/workgroups/${workgroupId}/approve/`, {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({action: 'approve'})
            });
            
            if (response.ok) {
                alert('Workgroup approved successfully');
                loadWorkgroups();
            } else {
                const data = await response.json();
                alert('Error: ' + (data.error || 'Failed to approve'));
            }
        } catch (error) {
            console.error('Error:', error);
            alert('Error approving workgroup');
        }
    }
    
    async function rejectWorkgroup(workgroupId) {
        const note = prompt('Reason for rejection (optional):');
        if (note === null) return;
        
        try {
            const response = await fetch(`/api/workgroups/${workgroupId}/approve/`, {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({action: 'reject', note: note})
            });
            
            if (response.ok) {
                alert('Workgroup rejected');
                loadWorkgroups();
            } else {
                const data = await response.json();
                alert('Error: ' + (data.error || 'Failed to reject'));
            }
        } catch (error) {
            console.error('Error:', error);
            alert('Error rejecting workgroup');
        }
    }
    
    // Load workgroups on page load
    loadWorkgroups();
    </script>
    """
    
    return render_page("Admin: Manage Workgroups - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/admin/chair-nominations/')
@require_role('admin')
def admin_chair_nominations():
    """Admin dashboard for managing chair nominations"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    
    content = """
    <div class="container mt-4">
        <h1 class="mb-4">Chair/Coordinator Nominations</h1>
        <p class="lead">Review and approve workgroup chair nominations</p>
        
        <ul class="nav nav-tabs mb-4" id="chairTabs" role="tablist">
            <li class="nav-item" role="presentation">
                <button class="nav-link active" id="pending-tab" data-bs-toggle="tab" data-bs-target="#pending" type="button">
                    Pending <span class="badge bg-warning ms-2" id="pending-count">0</span>
                </button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="approved-tab" data-bs-toggle="tab" data-bs-target="#approved" type="button">
                    Approved <span class="badge bg-success ms-2" id="approved-count">0</span>
                </button>
            </li>
        </ul>
        
        <div class="tab-content" id="chairTabContent">
            <div class="tab-pane fade show active" id="pending">
                <div id="pending-nominations"></div>
            </div>
            <div class="tab-pane fade" id="approved">
                <div id="approved-nominations"></div>
            </div>
        </div>
    </div>
    
    <script>
    async function loadNominations() {
        try {
            const response = await fetch('/api/admin/chair-nominations/');
            const data = await response.json();
            
            // Count nominations by status
            const pendingNoms = data.nominations.filter(n => !n.approved);
            const approvedNoms = data.nominations.filter(n => n.approved);
            
            document.getElementById('pending-count').textContent = pendingNoms.length;
            document.getElementById('approved-count').textContent = approvedNoms.length;
            
            // Render pending nominations
            let pendingHtml = '';
            if (pendingNoms.length > 0) {
                pendingHtml = '<div class="row">';
                pendingNoms.forEach(nom => {
                    pendingHtml += renderNominationCard(nom, 'pending');
                });
                pendingHtml += '</div>';
            } else {
                pendingHtml = '<div class="alert alert-info">No pending chair nominations</div>';
            }
            document.getElementById('pending-nominations').innerHTML = pendingHtml;
            
            // Render approved nominations
            let approvedHtml = '';
            if (approvedNoms.length > 0) {
                approvedHtml = '<div class="row">';
                approvedNoms.forEach(nom => {
                    approvedHtml += renderNominationCard(nom, 'approved');
                });
                approvedHtml += '</div>';
            } else {
                approvedHtml = '<div class="alert alert-info">No approved chair nominations</div>';
            }
            document.getElementById('approved-nominations').innerHTML = approvedHtml;
            
        } catch (error) {
            console.error('Error loading nominations:', error);
            document.getElementById('pending-nominations').innerHTML = '<div class="alert alert-danger">Error loading nominations</div>';
        }
    }
    
    function renderNominationCard(nom, status) {
        const selfNomBadge = nom.is_self_nomination ? '<span class="badge bg-info ms-2">Self-Nomination</span>' : '';
        const statusBadge = nom.approved ? '<span class="badge bg-success">Approved</span>' : '<span class="badge bg-warning">Pending</span>';
        
        return `
            <div class="col-md-6 mb-4">
                <div class="card">
                    <div class="card-header">
                        <div class="d-flex justify-content-between align-items-center">
                            <h5 class="mb-0">${nom.workgroup_name}</h5>
                            ${statusBadge}
                        </div>
                    </div>
                    <div class="card-body">
                        <div class="d-flex align-items-center mb-3">
                            <img 
                                src="${nom.nominee_profile_image || '/static/images/default-avatar.png'}" 
                                class="rounded-circle me-3" 
                                style="width: 60px; height: 60px; object-fit: cover;"
                                onerror="this.src='/static/images/default-avatar.png'"
                            >
                            <div>
                                <h6 class="mb-0">
                                    <a href="/profile/${nom.nominee_username}/" target="_blank">
                                        ${nom.chair_name}
                                    </a>
                                    ${selfNomBadge}
                                </h6>
                                <small class="text-muted">Nominated ${new Date(nom.set_at).toLocaleDateString()}</small>
                            </div>
                        </div>
                        
                        ${nom.nominator_name && !nom.is_self_nomination ? `
                            <p class="mb-2"><small><strong>Nominated by:</strong> 
                                <a href="/profile/${nom.nominator_username}/" target="_blank">${nom.nominator_name}</a>
                            </small></p>
                        ` : ''}
                        
                        <div class="mb-3">
                            <strong>Statement:</strong>
                            <p class="mt-1">${nom.statement || 'No statement provided'}</p>
                        </div>
                        
                        <div class="mb-2">
                            <strong>Workgroup:</strong> 
                            <a href="/workgroups/${nom.workgroup_slug}/" target="_blank">${nom.workgroup_name}</a>
                        </div>
                        
                        <div class="mb-3">
                            <strong>Layer:</strong> 
                            <a href="/projects/${nom.project_slug}/" target="_blank">${nom.project_name}</a>
                        </div>
                        
                        ${status === 'pending' ? `
                            <div class="d-flex gap-2">
                                <button class="btn btn-success flex-fill" onclick="approveNomination(${nom.id})">
                                    <i class="fas fa-check me-2"></i>Approve
                                </button>
                                <button class="btn btn-danger flex-fill" onclick="rejectNomination(${nom.id})">
                                    <i class="fas fa-times me-2"></i>Reject
                                </button>
                            </div>
                        ` : ''}
                    </div>
                </div>
            </div>
        `;
    }
    
    async function approveNomination(nominationId) {
        if (!confirm('Approve this chair nomination?')) return;
        
        try {
            const response = await fetch(`/api/admin/chair-nominations/${nominationId}/approve/`, {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' }
            });
            
            const data = await response.json();
            if (response.ok) {
                alert('Nomination approved!');
                loadNominations();
            } else {
                alert(data.error || 'Failed to approve nomination');
            }
        } catch (error) {
            console.error('Error:', error);
            alert('Failed to approve nomination');
        }
    }
    
    async function rejectNomination(nominationId) {
        const reason = prompt('Reason for rejection (optional):');
        if (reason === null) return; // User cancelled
        
        try {
            const response = await fetch(`/api/admin/chair-nominations/${nominationId}/reject/`, {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({ reason: reason })
            });
            
            const data = await response.json();
            if (response.ok) {
                alert('Nomination rejected');
                loadNominations();
            } else {
                alert(data.error || 'Failed to reject nomination');
            }
        } catch (error) {
            console.error('Error:', error);
            alert('Failed to reject nomination');
        }
    }
    
    // Load nominations on page load
    loadNominations();
    </script>
    """
    
    return render_page("Admin: Chair Nominations - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/admin/roles/')
@require_role('admin')
def admin_roles():
    """Admin dashboard for managing roles"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    
    content = """
    <div class="container mt-4">
        <h1 class="mb-4">Manage Roles</h1>
        
        <ul class="nav nav-tabs mb-4" id="roleTabs" role="tablist">
            <li class="nav-item" role="presentation">
                <button class="nav-link active" id="draft-tab" data-bs-toggle="tab" data-bs-target="#draft" type="button">
                    Draft <span class="badge bg-secondary ms-2" id="draft-count">0</span>
                </button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="approved-tab" data-bs-toggle="tab" data-bs-target="#approved" type="button">
                    Approved <span class="badge bg-success ms-2" id="approved-count">0</span>
                </button>
            </li>
        </ul>
        
        <div class="tab-content" id="roleTabContent">
            <div class="tab-pane fade show active" id="draft">
                <div id="draft-roles"></div>
            </div>
            <div class="tab-pane fade" id="approved">
                <div id="approved-roles"></div>
            </div>
        </div>
    </div>
    
    <script>
    async function loadRoles() {
        try {
            // Load all projects first
            const projectsResp = await fetch('/api/projects/');
            const projectsData = await projectsResp.json();
            
            let allRoles = [];
            
            // Load roles from all projects
            for (const project of projectsData.projects) {
                const rolesResp = await fetch(`/api/projects/${project.id}/roles/`);
                const rolesData = await rolesResp.json();
                
                // Add project info to each role
                rolesData.roles.forEach(role => {
                    role.project_name = project.name;
                    role.project_slug = project.slug;
                    allRoles.push(role);
                });
            }
            
            const draft = allRoles.filter(r => r.status === 'draft');
            const approved = allRoles.filter(r => r.status === 'approved');
            
            document.getElementById('draft-count').textContent = draft.length;
            document.getElementById('approved-count').textContent = approved.length;
            
            displayRoles('draft-roles', draft, true);
            displayRoles('approved-roles', approved, false);
        } catch (error) {
            console.error('Error loading roles:', error);
        }
    }
    
    function displayRoles(containerId, roles, showActions) {
        const container = document.getElementById(containerId);
        
        if (roles.length === 0) {
            container.innerHTML = '<div class="alert alert-info">No roles in this category</div>';
            return;
        }
        
        let html = '<div class="list-group">';
        roles.forEach(role => {
            html += `
                <div class="list-group-item">
                    <div class="d-flex justify-content-between align-items-start">
                        <div class="flex-grow-1">
                            <h5>${role.title_guild}</h5>
                            ${role.title_operational ? `<h6 class="text-muted">${role.title_operational}</h6>` : ''}
                            <p class="mb-2">${role.description.substring(0, 200)}...</p>
                            <small class="text-muted">
                                Layer: <a href="/projects/${role.project_slug}/" target="_blank">${role.project_name}</a> | 
                                Created: ${new Date(role.created_at).toLocaleDateString()} | 
                                Public: ${role.public_visible ? 'Yes' : 'No'}
                            </small>
                        </div>
                        ${showActions ? `
                            <div class="btn-group-vertical ms-3">
                                <button class="btn btn-sm btn-success" onclick="approveRole('${role.id}')">
                                    <i class="fas fa-check me-1"></i>Approve
                                </button>
                            </div>
                        ` : ''}
                    </div>
                </div>
            `;
        });
        html += '</div>';
        
        container.innerHTML = html;
    }
    
    async function approveRole(roleId) {
        if (!confirm('Approve this role?')) return;
        
        try {
            const response = await fetch(`/api/roles/${roleId}/approve/`, {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({action: 'approve'})
            });
            
            if (response.ok) {
                alert('Role approved successfully');
                loadRoles();
            } else {
                const data = await response.json();
                alert('Error: ' + (data.error || 'Failed to approve'));
            }
        } catch (error) {
            console.error('Error:', error);
            alert('Error approving role');
        }
    }
    
    // Load roles on page load
    loadRoles();
    </script>
    """
    
    return render_page("Admin: Manage Roles - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/admin/badges/')
@require_role('admin')
def admin_badges():
    """Admin dashboard for managing and issuing badges"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    
    content = """
    <div class="container mt-4">
        <h1 class="mb-4">Manage Badges</h1>
        
        <ul class="nav nav-tabs mb-4" id="badgeTabs" role="tablist">
            <li class="nav-item" role="presentation">
                <button class="nav-link active" id="requested-tab" data-bs-toggle="tab" data-bs-target="#requested" type="button">
                    Requested <span class="badge bg-warning ms-2" id="requested-count">0</span>
                </button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="approved-tab" data-bs-toggle="tab" data-bs-target="#approved" type="button">
                    Approved <span class="badge bg-success ms-2" id="approved-count">0</span>
                </button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="issued-tab" data-bs-toggle="tab" data-bs-target="#issued" type="button">
                    Issued <span class="badge bg-primary ms-2" id="issued-count">0</span>
                </button>
            </li>
        </ul>
        
        <div class="tab-content" id="badgeTabContent">
            <div class="tab-pane fade show active" id="requested">
                <div id="requested-badges"></div>
            </div>
            <div class="tab-pane fade" id="approved">
                <div id="approved-badges"></div>
            </div>
            <div class="tab-pane fade" id="issued">
                <div id="issued-badges"></div>
            </div>
        </div>
    </div>
    
    <script>
    async function loadBadges() {
        try {
            // Load all projects first
            const projectsResp = await fetch('/api/projects/');
            const projectsData = await projectsResp.json();
            
            let allBadges = [];
            
            // Load badges from all projects
            for (const project of projectsData.projects) {
                const badgesResp = await fetch(`/api/projects/${project.id}/badges/`);
                const badgesData = await badgesResp.json();
                
                // Add project info to each badge
                badgesData.badges.forEach(badge => {
                    badge.project_name = project.name;
                    badge.project_slug = project.slug;
                    allBadges.push(badge);
                });
            }
            
            const requested = allBadges.filter(b => b.status === 'requested');
            const approved = allBadges.filter(b => b.status === 'approved');
            const issued = allBadges.filter(b => b.status === 'issued');
            
            document.getElementById('requested-count').textContent = requested.length;
            document.getElementById('approved-count').textContent = approved.length;
            document.getElementById('issued-count').textContent = issued.length;
            
            displayBadges('requested-badges', requested, 'approve');
            displayBadges('approved-badges', approved, 'issue');
            displayBadges('issued-badges', issued, 'none');
        } catch (error) {
            console.error('Error loading badges:', error);
        }
    }
    
    function displayBadges(containerId, badges, actionType) {
        const container = document.getElementById(containerId);
        
        if (badges.length === 0) {
            container.innerHTML = '<div class="alert alert-info">No badges in this category</div>';
            return;
        }
        
        let html = '<div class="list-group">';
        badges.forEach(badge => {
            html += `
                <div class="list-group-item">
                    <div class="d-flex justify-content-between align-items-start">
                        <div class="flex-grow-1">
                            <h5>Badge: ${badge.badge_type}</h5>
                            <p class="mb-2">
                                <strong>Claim ID:</strong> ${badge.claim_id}<br>
                                <strong>Claimant ID:</strong> ${badge.claimant_id}<br>
                                <strong>Custody:</strong> ${badge.custody_mode}<br>
                                ${badge.btc_taproot_address ? `<strong>BTC Address:</strong> ${badge.btc_taproot_address}<br>` : ''}
                                ${badge.inscription_id ? `<strong>Inscription:</strong> ${badge.inscription_id}<br>` : ''}
                            </p>
                            <small class="text-muted">
                                Layer: <a href="/projects/${badge.project_slug}/" target="_blank">${badge.project_name}</a> | 
                                Created: ${new Date(badge.created_at).toLocaleDateString()}
                            </small>
                        </div>
                        ${actionType === 'approve' ? `
                            <div class="btn-group-vertical ms-3">
                                <button class="btn btn-sm btn-success" onclick="approveBadge('${badge.id}')">
                                    <i class="fas fa-check me-1"></i>Approve
                                </button>
                                <button class="btn btn-sm btn-danger" onclick="denyBadge('${badge.id}')">
                                    <i class="fas fa-times me-1"></i>Deny
                                </button>
                            </div>
                        ` : actionType === 'issue' ? `
                            <div class="btn-group-vertical ms-3">
                                <button class="btn btn-sm btn-primary" onclick="issueBadge('${badge.id}')">
                                    <i class="fas fa-certificate me-1"></i>Issue
                                </button>
                            </div>
                        ` : ''}
                    </div>
                </div>
            `;
        });
        html += '</div>';
        
        container.innerHTML = html;
    }
    
    async function approveBadge(badgeId) {
        const note = prompt('Approval note (optional):');
        if (note === null) return;
        
        try {
            const response = await fetch(`/api/badges/${badgeId}/approve/`, {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({action: 'approve', approval_note: note})
            });
            
            if (response.ok) {
                alert('Badge approved successfully');
                loadBadges();
            } else {
                const data = await response.json();
                alert('Error: ' + (data.error || 'Failed to approve'));
            }
        } catch (error) {
            console.error('Error:', error);
            alert('Error approving badge');
        }
    }
    
    async function denyBadge(badgeId) {
        const note = prompt('Reason for denial:');
        if (!note) return;
        
        try {
            const response = await fetch(`/api/badges/${badgeId}/approve/`, {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({approve: false, approval_note: note})
            });
            
            if (response.ok) {
                alert('Badge denied');
                loadBadges();
            } else {
                const data = await response.json();
                alert('Error: ' + (data.error || 'Failed to deny'));
            }
        } catch (error) {
            console.error('Error:', error);
            alert('Error denying badge');
        }
    }
    
    async function issueBadge(badgeId) {
        const inscriptionId = prompt('Enter inscription ID:');
        if (!inscriptionId) return;
        
        const txRef = prompt('Enter transaction reference (optional):');
        
        try {
            const response = await fetch(`/api/badges/${badgeId}/issue/`, {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({
                    inscription_id: inscriptionId,
                    tx_ref: txRef || null,
                    chain: 'bitcoin'
                })
            });
            
            if (response.ok) {
                alert('Badge issued successfully');
                loadBadges();
            } else {
                const data = await response.json();
                alert('Error: ' + (data.error || 'Failed to issue'));
            }
        } catch (error) {
            console.error('Error:', error);
            alert('Error issuing badge');
        }
    }
    
    // Load badges on page load
    loadBadges();
    </script>
    """
    
    return render_page("Admin: Manage Badges - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/admin/member_requests/')
@require_role('admin')
def admin_member_requests():
    """List pending workgroup member requests (when workgroup has members_require_approval=True). Default: no approval."""
    current_user = get_current_user()
    current_theme = session.get('theme', 'dark')
    user_menu = generate_user_menu()

    pending = WorkgroupMemberRequest.query.filter_by(status='pending').order_by(WorkgroupMemberRequest.requested_at.desc()).all()
    rows = ""
    for req in pending:
        req_at = req.requested_at.strftime('%Y-%m-%d %H:%M') if req.requested_at else ''
        rows += f"""
        <tr>
            <td>{req.user_name}</td>
            <td><code>{req.group_acronym}</code></td>
            <td>{req_at}</td>
            <td>
                <a href="/admin/member_requests/{req.id}/approve" class="btn btn-sm btn-success">Approve</a>
                <a href="/admin/member_requests/{req.id}/reject" class="btn btn-sm btn-outline-danger" onclick="return confirm('Reject this request?')">Reject</a>
            </td>
        </tr>
        """
    if not rows:
        rows = '<tr><td colspan="4" class="text-center text-muted py-4">No pending member requests. (Default is no approval; join is instant.)</td></tr>'

    content = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/admin/">Admin Dashboard</a></li>
                <li class="breadcrumb-item"><a href="/admin/chairs/">Coordinator Management</a></li>
                <li class="breadcrumb-item active">Member requests</li>
            </ol>
        </nav>
        <h1 class="mb-2">Member requests</h1>
        <p class="text-muted">When a workgroup has approval required, join requests appear here. Default: no approval (instant join).</p>
        <div class="card">
            <div class="card-body p-0">
                <div class="table-responsive">
                    <table class="table table-hover mb-0">
                        <thead class="table-light">
                            <tr><th>User</th><th>Workgroup</th><th>Requested</th><th>Actions</th></tr>
                        </thead>
                        <tbody>{rows}</tbody>
                    </table>
                </div>
            </div>
        </div>
        <div class="mt-3">
            <a href="/admin/chairs/" class="btn btn-secondary">Back to Coordinator Management</a>
        </div>
    </div>
    """
    return BASE_TEMPLATE.format(title="Member requests - MLGH", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/admin/member_requests/<int:req_id>/approve')
@require_role('admin')
def approve_member_request(req_id):
    req = WorkgroupMemberRequest.query.get(req_id)
    if not req or req.status != 'pending':
        flash('Request not found or already handled', 'error')
        return redirect('/admin/member_requests/')
    # Membership is by user_id only; avoid duplicate
    if req.user_id:
        existing = WorkingGroupMember.query.filter_by(group_acronym=req.group_acronym, user_id=req.user_id).first()
        if not existing:
            membership = WorkingGroupMember(
                group_acronym=req.group_acronym,
                user_id=req.user_id,
                user_name=req.user_name or ''
            )
            db.session.add(membership)
    req.status = 'approved'
    req.reviewed_at = datetime.utcnow()
    req.reviewed_by = get_current_user().get('name') or get_current_user().get('username')
    db.session.commit()
    flash(f'Member approved: {req.user_name} for {req.group_acronym}', 'success')
    return redirect('/admin/member_requests/')

@app.route('/admin/member_requests/<int:req_id>/reject')
@require_role('admin')
def reject_member_request(req_id):
    req = WorkgroupMemberRequest.query.get(req_id)
    if not req or req.status != 'pending':
        flash('Request not found or already handled', 'error')
        return redirect('/admin/member_requests/')
    req.status = 'rejected'
    req.reviewed_at = datetime.utcnow()
    req.reviewed_by = get_current_user().get('name') or get_current_user().get('username')
    db.session.commit()
    flash(f'Member request rejected: {req.user_name}', 'warning')
    return redirect('/admin/member_requests/')

# Routes
@app.route('/')
def home():
    # Generate user menu
    current_user = get_current_user()
    current_theme = current_user.get('theme', 'dark') if current_user else 'dark'  # Default to dark
    user_menu = generate_user_menu()
    
    # Count documents: DRAFTS + approved/published submissions
    doc_count = len(DRAFTS) + Submission.query.filter(Submission.status.in_(['approved', 'published'])).count()
    
    return BASE_TEMPLATE.format(title="MLGH", theme=current_theme, user_menu=user_menu, content=f"""
    
    <div class="container mt-4">
        <div class="row">
            <div class="col-md-8">
                <p class="lead">Welcome to the Governance Hub for the Meta-Layer!</p>

                <div class="row">
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <h5><i class="fas fa-project-diagram me-2"></i>Layers</h5>
                            </div>
                            <div class="card-body">
                                <p>Browse and discover MLTF layers and their workgroups.</p>
                                <a href="/projects/" class="btn btn-primary">View Layers</a>
                            </div>
                        </div>
                    </div>
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <h5><i class="fas fa-users me-2"></i>Workgroups</h5>
                            </div>
                            <div class="card-body">
                                <p>Browse workgroups across all projects and their activities.</p>
                                <a href="/workgroups/" class="btn btn-primary">View Workgroups</a>
                            </div>
                        </div>
                    </div>
                </div>
                
                <div class="row mt-4">
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <h5><i class="fas fa-shield-alt me-2"></i>Guilds</h5>
                            </div>
                            <div class="card-body">
                                <p>Cross-project collaboration groups and communities.</p>
                                <a href="/guilds/" class="btn btn-primary">View Guilds</a>
                            </div>
                        </div>
                    </div>
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <h5><i class="fas fa-user-tag me-2"></i>Roles</h5>
                            </div>
                            <div class="card-body">
                                <p>Explore and claim roles across all projects.</p>
                                <a href="/roles/" class="btn btn-primary">Browse Roles</a>
                            </div>
                        </div>
                    </div>
                </div>
                
                <div class="row mt-4">
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <h5><i class="fas fa-user-friends me-2"></i>People</h5>
                            </div>
                            <div class="card-body">
                                <p>Directory of Meta-Layer participants and contributors.</p>
                                <a href="/person/" class="btn btn-primary">View People</a>
                            </div>
                        </div>
                    </div>
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <h5><i class="fas fa-file-alt me-2"></i>Documents</h5>
                            </div>
                            <div class="card-body">
                                <p>View the latest Meta-Layer documents including drafts and RFCs.</p>
                                <a href="/doc/all/" class="btn btn-primary">View All Documents</a>
                            </div>
                        </div>
                    </div>
                </div>
                
                <div class="row mt-4">
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <h5><i class="fas fa-images me-2"></i>Role Images</h5>
                            </div>
                            <div class="card-body">
                                <p>Browse and vote on visual representations for roles across all projects.</p>
                                <a href="/role-images/" class="btn btn-primary">View Gallery</a>
                            </div>
                        </div>
                    </div>
                    <div class="col-md-6">
                        <div class="card">
                            <div class="card-header">
                                <h5><i class="fas fa-list-ol me-2"></i>Waitlists</h5>
                            </div>
                            <div class="card-body">
                                <p>Join waitlists for upcoming projects, features, and opportunities.</p>
                                <a href="/waitlists/" class="btn btn-primary">View Waitlists</a>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
            <div class="col-md-4">
                <div class="card">
                    <div class="card-header">
                        <h5>Quick Stats</h5>
                    </div>
                    <div class="card-body">
                        <p><strong>Documents:</strong> {doc_count}</p>
                        <p><strong>Workgroups:</strong> {len(GROUPS)}</p>
                        <p><strong>Last Updated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
                    </div>
                </div>
            </div>
        </div>
    </div>
    """, build_number=BUILD_NUMBER, hypothesis_config="")

# ================================================================
# UUID-BASED CANONICAL ROUTES
# ================================================================

@app.route('/p/<public_id>')
def person_by_public_id(public_id):
    """Resolve person by public_id UUID"""
    user = User.query.filter_by(public_id=public_id).first_or_404()
    return redirect(url_for('user_profile', username=user.username or user.handle))

@app.route('/layer/<public_id>')
def layer_by_public_id(public_id):
    """Resolve project/layer by public_id UUID"""
    project = Project.query.filter_by(public_id=public_id).first_or_404()
    return redirect(url_for('project_detail', project_slug=project.slug))

@app.route('/draft/<public_id>')
def draft_by_public_id(public_id):
    """Resolve draft/submission by public_id UUID"""
    submission = Submission.query.filter_by(public_id=public_id).first_or_404()
    return redirect(url_for('draft_detail', draft_name=submission.draft_name))

@app.route('/vote/<public_id>')
def vote_by_public_id_redirect(public_id):
    """Resolve vote by public_id UUID - redirect to detail page"""
    vote = Vote.query.filter_by(public_id=public_id).first_or_404()
    # Redirect to the full vote detail page
    return redirect(url_for('vote_detail', vote_public_id=vote.public_id))

@app.route('/votes/<vote_public_id>/')
def vote_detail(vote_public_id):
    """Vote detail page"""
    vote = Vote.query.filter_by(public_id=vote_public_id).first_or_404()
    project = Project.query.get_or_404(vote.project_id)
    submission = Submission.query.get_or_404(vote.submission_id)
    
    current_user = get_current_user()
    current_theme = current_user.get('theme', 'dark') if current_user else 'dark'
    user_menu = generate_user_menu()
    
    # Check if current user is eligible
    is_eligible = False
    has_voted = False
    user_ballot = None
    
    if current_user:
        eligibility = VoteEligibilitySnapshot.query.filter_by(
            vote_id=vote.id,
            person_id=current_user['id']
        ).first()
        is_eligible = eligibility and eligibility.is_eligible
        
        user_ballot = Ballot.query.filter_by(
            vote_id=vote.id,
            person_id=current_user['id']
        ).first()
        has_voted = user_ballot is not None
    
    # Get ballot counts
    ballot_count = Ballot.query.filter_by(vote_id=vote.id).count()
    eligible_count = VoteEligibilitySnapshot.query.filter_by(vote_id=vote.id, is_eligible=True).count()
    
    # Parse result summary
    result_summary = None
    if vote.result_summary:
        try:
            result_summary = json.loads(vote.result_summary)
        except:
            pass
    
    # Status badge colors
    status_colors = {
        'scheduled': 'secondary',
        'active': 'primary',
        'closed': 'success',
        'canceled': 'danger'
    }
    status_color = status_colors.get(vote.status, 'secondary')
    
    # Result badge colors
    result_colors = {
        'passed': 'success',
        'failed': 'danger',
        'no_quorum': 'warning',
        'canceled': 'secondary'
    }
    result_color = result_colors.get(vote.result, 'secondary') if vote.result else 'secondary'
    
    # Build result HTML if available
    result_html = ''
    if vote.result:
        result_text = vote.result.upper() if vote.result else 'PENDING'
        result_html = f'<p><strong>Result:</strong> <span class="badge bg-{result_color}">{result_text}</span></p>'
    
    # Build ballot form if eligible and active
    ballot_form_html = ''
    if vote.status == 'active' and is_eligible:
        voted_msg = ''
        if has_voted:
            voted_msg = f'<p class="text-success">✓ You have already voted: <strong>{user_ballot.choice.upper()}</strong></p>'
        ballot_form_html = f'''<div class="card mb-3">
            <div class="card-header"><h5>Cast Your Ballot</h5></div>
            <div class="card-body">
                <p>You are eligible to vote in this election.</p>
                {voted_msg}
                <div class="btn-group" role="group">
                    <button class="btn btn-success" onclick="castBallot('yes')">Vote YES</button>
                    <button class="btn btn-danger" onclick="castBallot('no')">Vote NO</button>
                    <button class="btn btn-secondary" onclick="castBallot('abstain')">Abstain</button>
                </div>
                <div id="ballot-status" class="mt-2"></div>
            </div>
        </div>'''
    
    # Build results HTML
    results_html = ''
    if vote.status == 'closed' or result_summary:
        if result_summary:
            quorum_text = 'Yes' if result_summary['quorum_met'] else 'No'
            results_content = f'''<p><strong>Yes:</strong> {result_summary['yes']}</p>
                <p><strong>No:</strong> {result_summary['no']}</p>
                <p><strong>Abstain:</strong> {result_summary['abstain']}</p>
                <p><strong>Total Votes Cast:</strong> {result_summary['votes_cast']} / {result_summary['eligible']} eligible</p>
                <p><strong>Quorum Met:</strong> {quorum_text}</p>
                <p><strong>Yes Ratio:</strong> {int(result_summary['yes_ratio'] * 100)}%</p>'''
        else:
            results_content = f'<p>Votes cast: {ballot_count} / {eligible_count} eligible</p>'
        results_html = f'''<div class="card mb-3">
            <div class="card-header"><h5>Results</h5></div>
            <div class="card-body">{results_content}</div>
        </div>'''
    
    # Build user status HTML
    user_status_html = ''
    if current_user:
        status_text = '✓ Eligible' if is_eligible else '✗ Not Eligible'
        user_status_html = f'<p><strong>Your Status:</strong> {status_text}</p>'
    
    content = f'''
    <div class="container mt-4">
        <div class="row">
            <div class="col-md-8">
                <h1>{vote.title}</h1>
                <p class="lead">{vote.description or ''}</p>
                
                <div class="card mb-3">
                    <div class="card-header">
                        <h5>Vote Details</h5>
                    </div>
                    <div class="card-body">
                        <p><strong>Status:</strong> <span class="badge bg-{status_color}">{vote.status.upper()}</span></p>
                        {result_html}
                        <p><strong>Layer:</strong> <a href="/projects/{project.slug}/">{project.name}</a></p>
                        <p><strong>Draft:</strong> <a href="/doc/draft/{submission.draft_name}/">{submission.title}</a></p>
                        <p><strong>Start:</strong> {vote.start_at.strftime('%Y-%m-%d %H:%M UTC')}</p>
                        <p><strong>End:</strong> {vote.end_at.strftime('%Y-%m-%d %H:%M UTC')}</p>
                        <p><strong>Quorum Required:</strong> {vote.quorum_count} votes</p>
                        <p><strong>Win Threshold:</strong> {int(vote.win_threshold * 100)}%</p>
                    </div>
                </div>
                
                {ballot_form_html}
                {results_html}
            </div>
            
            <div class="col-md-4">
                <div class="card">
                    <div class="card-header">
                        <h5>Participation</h5>
                    </div>
                    <div class="card-body">
                        <p><strong>Eligible Voters:</strong> {eligible_count}</p>
                        <p><strong>Ballots Cast:</strong> {ballot_count}</p>
                        {user_status_html}
                    </div>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    function castBallot(choice) {{
        fetch('/api/votes/{vote.public_id}/ballot/', {{
            method: 'POST',
            headers: {{'Content-Type': 'application/json'}},
            body: JSON.stringify({{choice: choice}})
        }})
        .then(res => res.json())
        .then(data => {{
            if (data.success) {{
                document.getElementById('ballot-status').innerHTML = '<div class="alert alert-success">Ballot cast successfully: ' + choice.toUpperCase() + '</div>';
                setTimeout(() => location.reload(), 1500);
            }} else {{
                document.getElementById('ballot-status').innerHTML = '<div class="alert alert-danger">Error: ' + data.error + '</div>';
            }}
        }})
        .catch(err => {{
            document.getElementById('ballot-status').innerHTML = '<div class="alert alert-danger">Error casting ballot</div>';
        }});
    }}
    </script>
    '''
    
    return BASE_TEMPLATE.format(title=f"Vote: {vote.title}", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/doc/active/')
def active_documents():
    """Show active documents (alias for all documents)"""
    return all_documents()

@app.route('/doc/all/')
def all_documents():
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    
    # Get all documents: published drafts + approved/published submissions
    all_docs = []
    
    # Add published drafts from DRAFTS list
    all_docs.extend(DRAFTS)
    
    # Add approved/published submissions from database
    # Exclude revisions - we only want to show the original draft or the latest approved revision
    approved_submissions = Submission.query.filter(
        Submission.status.in_(['approved', 'published']),
        Submission.is_revision == False  # Only show non-revision submissions
    ).all()
    
    # For each approved submission, check if there's a newer approved revision
    for submission in approved_submissions:
        # Check if there's an approved revision for this draft
        latest_revision = Submission.query.filter(
            Submission.parent_draft_name == submission.id,
            Submission.is_revision == True,
            Submission.status.in_(['approved', 'published'])
        ).order_by(Submission.revision_number.desc()).first()
        
        # Use the latest approved revision if it exists, otherwise use the original
        display_submission = latest_revision if latest_revision else submission
        
        # Use stored pages and words values (calculated on submission)
        pages = display_submission.pages if display_submission.pages else 1
        words = display_submission.words if display_submission.words else 0
        
        # Get revision info for display
        is_revision = getattr(display_submission, 'is_revision', False)
        revision_number = getattr(display_submission, 'revision_number', '')
        
        all_docs.append({
            'name': display_submission.id,
            'title': display_submission.title,
            'authors': display_submission.authors if isinstance(display_submission.authors, list) else [display_submission.authors] if display_submission.authors else [],
            'group': display_submission.group or 'N/A',
            'status': display_submission.status,
            'rev': revision_number if is_revision else '00',
            'pages': pages,
            'words': words,
            'date': display_submission.submitted_at.strftime('%Y-%m-%d') if display_submission.submitted_at else '',
            'abstract': display_submission.abstract or '',
            'ml_number': display_submission.ml_number,
            'is_revision': is_revision,
            'revision_number': revision_number
        })
    
    docs_html = ""
    for draft in all_docs:
        display_id = draft.get('ml_number') or draft['name']
        is_revision = draft.get('is_revision', False)
        revision_number = draft.get('revision_number', '')
        revision_badge = f'<span class="badge bg-success ms-2">Revision {revision_number}</span>' if is_revision and revision_number else ''
        
        docs_html += f"""
        <div class="col-md-6 document-card">
            <div class="card">
                <div class="card-body">
                    <h5 class="card-title document-title">
                        <a href="/doc/draft/{draft['name']}/">{display_id}</a>
                        {revision_badge}
                    </h5>
                    <p class="card-text">{draft['title']}</p>
                    <div class="document-meta">
                        <span class="badge bg-secondary status-badge">{draft['status']}</span>
                        <span class="ms-2">Rev: {draft['rev']}</span>
                        <span class="ms-2">{draft['pages']} pages</span>
                        <span class="ms-2">{draft['words']} words</span>
                    </div>
                    <div class="mt-2">
                        <small class="text-muted">
                            Authors: {', '.join(draft['authors']) if draft['authors'] else 'N/A'}<br>
                            Group: {draft['group']}<br>
                            Date: {draft['date']}
                        </small>
                    </div>
                    <div class="mt-2">
                        <a href="/doc/draft/{draft['name']}/comments/" class="btn btn-sm btn-outline-primary">Comments</a>
                        <a href="/doc/draft/{draft['name']}/history/" class="btn btn-sm btn-outline-secondary">History</a>
                        <a href="/doc/draft/{draft['name']}/revisions/" class="btn btn-sm btn-outline-info">Revisions</a>
                    </div>
                </div>
            </div>
        </div>
        """
    
    content = f"""
    <div class="container mt-4">
        <h1>All Documents</h1>
        <p>Showing {len(all_docs)} documents</p>
        
        <div class="row">
            {docs_html}
        </div>
    </div>
    """

    return BASE_TEMPLATE.format(title="All Documents - MLGH", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/doc/draft/<path:draft_name>.txt')
def draft_text(draft_name):
    """Serve draft content as plain text"""
    # First try to find in DRAFTS (published documents)
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)
    
    # If not found in DRAFTS, try to find as a submission ID
    submission = None
    if not draft:
        submission = Submission.query.filter_by(id=draft_name).first()
        if submission:
            draft = {
                'name': submission.id,
                'title': submission.title,
                'authors': submission.authors,
                'abstract': submission.abstract or 'Abstract not available for this draft.',
                'status': submission.status,
                'group': submission.group,
                'date': submission.submitted_at.strftime('%Y-%m-%d') if submission.submitted_at else '',
            }
    
    if not draft:
        return "Document not found", 404
    
    # Load document content
    document_content = "Document content not available."
    
    # Try to get content from submission file first
    if submission and submission.file_path and os.path.exists(submission.file_path):
        _, ext = os.path.splitext(submission.filename.lower())
        try:
            if ext in ['.txt', '.xml']:
                with open(submission.file_path, 'r', encoding='utf-8', errors='replace') as f:
                    document_content = f.read()
            elif ext == '.docx':
                from docx import Document
                doc = Document(submission.file_path)
                content_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content_parts.append(cell.text)
                document_content = '\n\n'.join(content_parts)
            elif ext == '.pdf':
                from PyPDF2 import PdfReader
                reader = PdfReader(submission.file_path)
                content_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(text)
                document_content = '\n\n'.join(content_parts)
                # Clean up PDF text
                import re
                document_content = re.sub(r'\n+', '\n', document_content)
                document_content = re.sub(r' +', ' ', document_content)
            else:
                document_content = f"Document content cannot be displayed for {ext.upper()} files. Please download to view."
        except Exception as e:
            document_content = f"Error loading document content: {str(e)}"
    
    # If no submission content, try to get from DRAFTS data
    elif draft and 'name' in draft:
        # Generate text content from draft data
        document_content = f"""INTERNET-DRAFT                                               {', '.join(draft.get('authors', []))}
Intended status: Informational                            Meta-Layer Initiative
Expires: {draft.get('date', 'TBD')}                                      {draft.get('date', 'TBD')}


{draft.get('title', 'Document Title')}


Abstract

{draft.get('abstract', 'Abstract not available.')}


1. Introduction

This document describes {draft.get('title', 'the subject matter')}.

The content of this draft is currently being developed and will be available
in the full document once published.

2. Status of This Memo

This Internet-Draft is submitted in full conformance with the provisions
of BCP 78 and BCP 79.

Meta-Layer Drafts are working documents of the Meta-Layer Task Force
(MLGH). These documents represent proposals and specifications for the
Meta-Layer ecosystem. The list of current Meta-Layer Drafts is available
in the MLGH datatracker.

Internet-Drafts are draft documents valid for a maximum of six months and
may be updated, replaced, or obsoleted by other documents at any time. It is
inappropriate to use Internet-Drafts as reference material or to cite them
other than as "work in progress."

This Internet-Draft will expire on {draft.get('date', 'TBD')}.


3. References

[MLGH] MLGH Datatracker, https://rfc.themetalayer.org/

Authors' Addresses

{chr(10).join([f'{author} <email@example.com>' for author in draft.get('authors', [])])}

Meta-Layer Initiative
"""
    
    # Return as plain text
    from flask import Response
    return Response(document_content, mimetype='text/plain; charset=utf-8')

@app.route('/doc/draft/<draft_name>/')
def draft_detail(draft_name):
    # First try to find in DRAFTS (published documents)
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)

    # If not found in DRAFTS, try to find as a submission ID or ML number
    submission = None
    if not draft:
        # Try to find by submission ID first
        submission = Submission.query.filter_by(id=draft_name).first()
        
        # If not found by ID, try to find by ML number
        if not submission:
            submission = Submission.query.filter_by(ml_number=draft_name).first()
        if submission:
            # Calculate pages and words for ordinals
            source_type = getattr(submission, 'sourceType', 'file')
            pages_count = 1
            words_count = 0
            ordinal_content_url = getattr(submission, 'ordinalContentUrl', None)
            ordinal_content_type = getattr(submission, 'ordinalContentType', '')
            
            if source_type == 'ordinal':
                # Fetch ordinal content to calculate words and pages
                
                if ordinal_content_url and ('text/' in ordinal_content_type or 'application/json' in ordinal_content_type):
                    try:
                        import requests
                        headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                        }
                        response = requests.get(ordinal_content_url, headers=headers, timeout=10)
                        if response.status_code == 200:
                            text_content = response.text
                            words_count = len(text_content.split())
                            # Estimate pages (assuming ~500 words per page)
                            pages_count = max(1, (words_count + 499) // 500)
                    except Exception as e:
                        app.logger.warning(f"Failed to fetch ordinal content for word/page count: {e}")
                        # Keep defaults
                        pass
            
            # Create a draft-like object from the submission
            draft = {
                'name': submission.id,
                'title': submission.title,
                'authors': submission.authors,
                'abstract': submission.abstract or 'Abstract not available for this draft.',
                'status': submission.status,
                'group': submission.group,
                'date': submission.submitted_at.strftime('%Y-%m-%d') if submission.submitted_at else '',
                'rev': '00',  # Default revision for submissions
                'pages': pages_count,
                'words': words_count,
                'stream': 'mltf',  # Default stream
                'ml_number': submission.ml_number,
                # Ordinal metadata
                'sourceType': source_type,
                'ordinalId': getattr(submission, 'ordinalId', None),
                'inscriptionNumber': getattr(submission, 'inscriptionNumber', None),
                'blockHeight': getattr(submission, 'blockHeight', None),
                'inscriptionTimestamp': getattr(submission, 'inscriptionTimestamp', None),
                'ordinalContentType': ordinal_content_type,
                # Revision metadata
                'is_revision': getattr(submission, 'is_revision', False),
                'revision_number': getattr(submission, 'revision_number', ''),
                'parent_draft_name': getattr(submission, 'parent_draft_name', '')
            }

    if not draft:
        return "Document not found", 404
    
    # Load full document content
    document_content = "Document content not available."
    calculated_pages = draft.get('pages', 1)
    calculated_words = draft.get('words', 0)

    # Try to get content from ordinal first
    if submission and draft.get('sourceType') == 'ordinal':
        ordinal_content_url = getattr(submission, 'ordinalContentUrl', None)
        ordinal_content_type = getattr(submission, 'ordinalContentType', '')
        
        if ordinal_content_url and ('text/' in ordinal_content_type or 'application/json' in ordinal_content_type):
            try:
                import requests
                import markdown2
                import bleach
                import re
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = requests.get(ordinal_content_url, headers=headers, timeout=10)
                if response.status_code == 200:
                    raw_content = response.text
                    # Calculate words and pages from ordinal text
                    words = len(raw_content.split())
                    calculated_pages = max(1, (words + 499) // 500)
                    calculated_words = words
                    # Update draft with calculated values
                    draft['pages'] = calculated_pages
                    draft['words'] = calculated_words
                    
                    # Check if content is markdown
                    is_markdown = False
                    if 'text/plain' in ordinal_content_type or 'text/markdown' in ordinal_content_type:
                        # Detect markdown patterns
                        markdown_patterns = [
                            r'^#{1,6}\s+.+$',  # Headers
                            r'\*\*.+\*\*',      # Bold
                            r'\*.+\*',          # Italic
                            r'^\s*[-*+]\s+',    # Lists
                            r'^\s*\d+\.\s+',    # Numbered lists
                            r'\[.+\]\(.+\)',    # Links
                            r'!\[.*\]\(.+\)'    # Images
                        ]
                        for pattern in markdown_patterns:
                            if re.search(pattern, raw_content, re.MULTILINE):
                                is_markdown = True
                                break
                    
                    if is_markdown:
                        # Use shared markdown processing function for consistency
                        document_content = process_ordinal_markdown(raw_content)
                    else:
                        # Display as plain text
                        document_content = raw_content
            except Exception as e:
                app.logger.warning(f"Failed to fetch ordinal content for display: {e}")
                document_content = f"Error loading ordinal content: {str(e)}"
        elif ordinal_content_url and ordinal_content_type.startswith('image/'):
            document_content = f'<img src="{ordinal_content_url}" class="img-fluid" style="max-width: 100%;" alt="Ordinal image content">'
        else:
            document_content = f"Ordinal content type: {ordinal_content_type}\nPreview not available for this content type."
    
    # Try to get content from submission file
    elif submission and submission.file_path and os.path.exists(submission.file_path):
        _, ext = os.path.splitext(submission.filename.lower())
        try:
            if ext in ['.txt', '.xml']:
                with open(submission.file_path, 'r', encoding='utf-8', errors='replace') as f:
                    document_content = f.read()
                # Calculate words and pages from text
                words = len(document_content.split())
                # Estimate pages (assuming ~500 words per page)
                calculated_pages = max(1, (words + 499) // 500)
                calculated_words = words
            elif ext == '.docx':
                from docx import Document
                doc = Document(submission.file_path)
                content_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content_parts.append(cell.text)
                document_content = '\n\n'.join(content_parts)
                # Calculate words and pages
                words = len(document_content.split())
                calculated_pages = max(1, (words + 499) // 500)
                calculated_words = words
            elif ext == '.pdf':
                # For PDFs, extract metadata but display with embedded viewer
                from PyPDF2 import PdfReader
                reader = PdfReader(submission.file_path)
                # Get page count and estimate words
                calculated_pages = len(reader.pages) if reader.pages else 1
                # Estimate words (roughly 250-300 words per page for PDFs)
                calculated_words = calculated_pages * 275
                
                # Create embedded PDF viewer for display
                file_size = os.path.getsize(submission.file_path)
                file_size_kb = file_size / 1024
                document_content = f'''
<div class="pdf-viewer-container">
    <div class="alert alert-info mb-3">
        <i class="bi bi-file-pdf"></i> PDF Document ({calculated_pages} pages, ~{calculated_words} words, {file_size_kb:.1f} KB)
    </div>
    <iframe src="/view/{draft_name}" 
            type="application/pdf" 
            style="width: 100%; height: 800px; border: 1px solid var(--card-border); border-radius: 4px;"
            title="PDF Document Viewer">
        <p>Your browser does not support PDF preview. 
           <a href="/download/{draft_name}">Download the PDF</a> to view it.</p>
    </iframe>
</div>
'''
            else:
                document_content = f"Document content cannot be displayed for {ext.upper()} files. Please download to view."
        except Exception as e:
            document_content = f"Error loading document content: {str(e)}"
        # Update draft with calculated values
        if submission:
            draft['pages'] = calculated_pages
            draft['words'] = calculated_words

    # If no submission content, try to get from DRAFTS data
    elif draft and 'name' in draft:
        # For demo purposes, generate some sample content based on the draft
        document_content = f"""INTERNET-DRAFT                                               {', '.join(draft.get('authors', []))}
Intended status: Informational                            Meta-Layer Initiative
Expires: {draft.get('date', 'TBD')}                                      {draft.get('date', 'TBD')}


{draft.get('title', 'Document Title')}


Abstract

{draft.get('abstract', 'Abstract not available.')}


1. Introduction

This document describes {draft.get('title', 'the subject matter')}.

The content of this draft is currently being developed and will be available
in the full document once published.

2. Status of This Memo

This Internet-Draft is submitted in full conformance with the provisions
of BCP 78 and BCP 79.

Meta-Layer Drafts are working documents of the Meta-Layer Task Force
(MLGH). These documents represent proposals and specifications for the
Meta-Layer ecosystem. The list of current Meta-Layer Drafts is available
in the MLGH datatracker.

Internet-Drafts are draft documents valid for a maximum of six months and
may be updated, replaced, or obsoleted by other documents at any time. It is
inappropriate to use Internet-Drafts as reference material or to cite them
other than as "work in progress."

This Internet-Draft will expire on {draft.get('date', 'TBD')}.


3. References

[MLGH] MLGH Datatracker, https://rfc.themetalayer.org/

Authors' Addresses

{chr(10).join([f'{author} <email@example.com>' for author in draft.get('authors', [])])}

Meta-Layer Initiative
"""

    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    # Show ML number only if approved, otherwise show submission ID
    if draft.get('status') == 'approved' and draft.get('ml_number'):
        display_id = draft.get('ml_number')
    else:
        display_id = draft['name']
    # Check if this is a revision
    is_revision = draft.get('is_revision', False)
    revision_number = draft.get('revision_number', '')
    revision_badge = f'<span class="badge bg-success ms-2">Revision {revision_number}</span>' if is_revision and revision_number else ''
    
    # Determine content styling based on source type
    if draft.get('sourceType') == 'ordinal':
        content_style = "font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; font-size: 1em; line-height: 1.6;"
    else:
        content_style = "font-family: 'Courier New', monospace; font-size: 0.9em; line-height: 1.4; white-space: pre-wrap;"
    
    content = f"""
    <div class="container mt-4">
        <h1>{display_id} {revision_badge}</h1>
        <p class="lead">{draft['title']}</p>

        <div class="row">
            <div class="col-md-8">
                <div class="card">
                    <div class="card-header">
                        <h5>Document Information</h5>
                    </div>
                    <div class="card-body">
                        <table class="table" style="color: var(--text-primary) !important;">
                            <tr><td style="color: var(--text-secondary) !important;"><strong>ID:</strong></td><td style="color: var(--text-primary) !important;">{display_id}</td></tr>
                            <tr><td style="color: var(--text-secondary) !important;"><strong>Title:</strong></td><td style="color: var(--text-primary) !important;">{draft['title']}</td></tr>
                            <tr><td style="color: var(--text-secondary) !important;"><strong>Status:</strong></td><td style="color: var(--text-primary) !important;"><span class="badge bg-secondary">{draft['status']}</span></td></tr>
                            <tr><td style="color: var(--text-secondary) !important;"><strong>Authors:</strong></td><td style="color: var(--text-primary) !important;">{', '.join(draft['authors'])}</td></tr>
                            <tr><td style="color: var(--text-secondary) !important;"><strong>Group:</strong></td><td style="color: var(--text-primary) !important;">{draft['group'] or 'N/A'}</td></tr>
                            <tr><td style="color: var(--text-secondary) !important;"><strong>Date:</strong></td><td style="color: var(--text-primary) !important;">{draft['date']}</td></tr>
                            {f'<tr><td colspan="2" style="padding-top: 15px;"><hr style="border-color: var(--border-color);"></td></tr><tr><td style="color: var(--text-secondary) !important;"><strong>Source:</strong></td><td style="color: var(--text-primary) !important;"><span class="badge bg-info"><i class="bi bi-coin"></i> Bitcoin Ordinal</span></td></tr>' if draft.get('sourceType') == 'ordinal' else f'<tr><td style="color: var(--text-secondary) !important;"><strong>Revision:</strong></td><td style="color: var(--text-primary) !important;">{draft["rev"]}</td></tr><tr><td style="color: var(--text-secondary) !important;"><strong>Pages:</strong></td><td style="color: var(--text-primary) !important;">{draft["pages"]}</td></tr><tr><td style="color: var(--text-secondary) !important;"><strong>Words:</strong></td><td style="color: var(--text-primary) !important;">{draft["words"]}</td></tr>'}
                            {f'<tr><td style="color: var(--text-secondary) !important;"><strong>Inscription #:</strong></td><td style="color: var(--text-primary) !important;">{draft["inscriptionNumber"]}</td></tr>' if draft.get('sourceType') == 'ordinal' and draft.get('inscriptionNumber') else ''}
                            {f'<tr><td style="color: var(--text-secondary) !important;"><strong>Block Height:</strong></td><td style="color: var(--text-primary) !important;">{draft["blockHeight"]}</td></tr>' if draft.get('sourceType') == 'ordinal' and draft.get('blockHeight') else ''}
                            {f'<tr><td style="color: var(--text-secondary) !important;"><strong>Timestamp:</strong></td><td style="color: var(--text-primary) !important;">{draft["inscriptionTimestamp"].strftime("%Y-%m-%d %H:%M UTC") if draft.get("inscriptionTimestamp") else "N/A"}</td></tr>' if draft.get('sourceType') == 'ordinal' else ''}
                            {f'<tr><td style="color: var(--text-secondary) !important;"><strong>Content Type:</strong></td><td style="color: var(--text-primary) !important;">{draft["ordinalContentType"]}</td></tr>' if draft.get('sourceType') == 'ordinal' and draft.get('ordinalContentType') else ''}
                            {f'<tr><td style="color: var(--text-secondary) !important;"><strong>Inscription ID:</strong></td><td style="color: var(--text-primary) !important;"><a href="https://ordinals.com/inscription/{draft["ordinalId"]}" target="_blank" class="text-decoration-none" style="color: var(--accent-color) !important;"><code style="font-family: monospace; font-size: 0.85em;">{shorten_inscription_id(draft["ordinalId"], 8)}</code></a></td></tr>' if draft.get('sourceType') == 'ordinal' and draft.get('ordinalId') else ''}
                        </table>
                    </div>
                </div>
                
                <div class="card mt-3">
                    <div class="card-header">
                        <h5>Abstract</h5>
                    </div>
                    <div class="card-body">
                        <p>{draft.get('abstract', 'Abstract not available for this draft.')}</p>
                    </div>
                </div>

                <!-- Full Document Content -->
                <div class="card mt-3">
                    <div class="card-header d-flex justify-content-between align-items-center">
                        <h5 class="mb-0">Document Content</h5>
                        <div>
                            {'' if draft.get('sourceType') == 'ordinal' else f'''
                            <a href="/download/{draft['name']}" class="btn btn-sm btn-outline-primary" target="_blank">
                                <i class="fas fa-download me-1"></i>Download
                            </a>
                            <a href="/doc/draft/{draft['name']}.txt" class="btn btn-sm btn-outline-secondary" target="_blank">
                                <i class="fas fa-external-link-alt me-1"></i>View TXT
                            </a>
                            '''}
                        </div>
                    </div>
                    <div class="card-body">
                        <div class="document-content" style="{content_style} background-color: var(--input-bg) !important; color: var(--text-primary) !important; padding: 20px; border-radius: 8px; max-height: 800px; overflow-y: auto; border: 1px solid var(--input-border);">
{document_content}
                        </div>
                    </div>
                </div>
            </div>
            
            <div class="col-md-4">
                <div class="card">
                    <div class="card-header">
                        <h5>Actions</h5>
                    </div>
                    <div class="card-body">
                        {f'<a href="/doc/draft/{draft["name"]}/comments/" class="btn btn-primary w-100 mb-2">View Comments ({Comment.query.filter_by(draft_name=draft_name).count()})</a>' if draft.get('status') == 'approved' else ''}
                        {f'<div class="small text-muted mb-2" id="annotation-count">Loading annotation count...</div>' if HYPOTHESIS_ENABLED else ''}
                        <a href="/doc/draft/{draft['name']}/history/" class="btn btn-secondary w-100 mb-2">View History</a>
                        <a href="/doc/draft/{draft['name']}/revisions/" class="btn btn-info w-100 mb-2">View Revisions</a>
                        
                        {f'''<div class="border-top pt-2 mt-2">
                            <h6 class="text-muted mb-2">Annotations</h6>
                            <button id="toggle-annotations" class="btn btn-outline-info w-100 mb-2" onclick="toggleAnnotations()">
                                <i class="fas fa-comment-dots me-1"></i>
                                <span id="annotations-text">Enable Annotations</span>
                            </button>
                            {'<div class="alert alert-info small mt-2" role="alert"><i class="fas fa-user-plus me-1"></i><strong>First time?</strong> <a href="https://hypothes.is/signup" target="_blank" class="alert-link">Create free Hypothesis account</a> (30 seconds) to annotate and highlight text.</div>' if not current_user or not current_user.get('hypothesis_account') else ''}
                            <small class="text-muted d-block">
                                Powered by <a href="https://hypothes.is" target="_blank" class="text-decoration-none">Hypothesis</a>. 
                                Public annotations visible to everyone.
                            </small>
                        </div>''' if HYPOTHESIS_ENABLED else ''}
                        {f'<a href="/submit/revision/{draft["name"]}/" class="btn btn-success w-100 mb-2"><i class="fas fa-plus me-1"></i>Submit New Revision</a>' if current_user and draft.get('status') == 'approved' else ''}
                        {'' if draft.get('sourceType') == 'ordinal' else f'<a href="/download/{draft["name"]}" class="btn btn-outline-primary w-100 mb-2">Download Document</a>'}
                        {'<form method="post" action="/doc/draft/' + draft['name'] + '/follow/" style="display: inline;" class="mb-2"><select name="notification_level" class="form-select form-select-sm mb-1"><option value="all">All changes & comments</option><option value="significant">Significant changes only</option><option value="major">Major changes only</option><option value="comments">Comments only</option><option value="none">No notifications</option></select><button type="submit" class="btn btn-success w-100"><i class="fas fa-bell me-1"></i>Follow Document</button></form>' if current_user and draft.get('status') == 'approved' and not is_user_following_draft(draft_name, current_user) else ''}
                        {'<form method="post" action="/doc/draft/' + draft['name'] + '/unfollow/" style="display: inline;" class="mb-2"><button type="submit" class="btn btn-warning w-100"><i class="fas fa-bell-slash me-1"></i>Unfollow Document</button></form>' if current_user and draft.get('status') == 'approved' and is_user_following_draft(draft_name, current_user) else ''}
                        {get_notification_controls(draft_name, current_user) if current_user and draft.get('status') == 'approved' and is_user_following_draft(draft_name, current_user) else ''}
                        {'' if not current_user else ''}
                    </div>
                </div>
                
                <div class="card mt-3" id="votes-card" style="display: none;">
                    <div class="card-header">
                        <h5>Votes</h5>
                    </div>
                    <div class="card-body" id="votes-container">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
                
                {f'''<div class="card mt-3">
                    <div class="card-header">
                        <h5>Quick Comment</h5>
                    </div>
                    <div class="card-body">
                        <form method="POST" action="/doc/draft/{draft['name']}/comments/">
                            <div class="mb-3">
                                <textarea class="form-control" name="comment" rows="3" placeholder="Add a quick comment..." required></textarea>
                    </div>
                            <button type="submit" class="btn btn-success btn-sm w-100">Post Comment</button>
                        </form>
        </div>
    </div>''' if draft.get('status') == 'approved' else ''}
    
                <div class="card mt-3">
                    <div class="card-header">
                        <h5>Related Documents</h5>
                    </div>
                <div class="card-body">
                        <p>Related documents would appear here in the real datatracker.</p>
                    </div>
                    </div>
                </div>
            </div>
        </div>
        
        <script>
        const submissionId = '{draft["name"] if submission else draft_name}';
        
        async function loadDraftVotes() {{
            try {{
                // Find votes that reference this submission
                const allProjects = await fetch('/api/projects/').then(r => r.json());
                let votes = [];
                
                for (const proj of allProjects.projects || []) {{
                    const res = await fetch(`/api/projects/${{proj.id}}/votes/`);
                    const data = await res.json();
                    const matchingVotes = (data.votes || []).filter(v => v.submission_id === submissionId);
                    votes.push(...matchingVotes);
                }}
                
                if (votes.length === 0) {{
                    return; // Keep card hidden
                }}
                
                document.getElementById('votes-card').style.display = 'block';
                const container = document.getElementById('votes-container');
                
                let html = '';
                for (const v of votes) {{
                    const statusBadge = v.status === 'active' ? '<span class="badge bg-success">Active</span>' : v.status === 'closed' ? '<span class="badge bg-secondary">Closed</span>' : '<span class="badge bg-info">Scheduled</span>';
                    const resultBadge = v.result ? '<span class="badge bg-' + (v.result === 'passed' ? 'success' : v.result === 'failed' ? 'danger' : 'warning') + ' ms-1">' + v.result + '</span>' : '';
                    html += '<div class="mb-3">';
                    html += '<h6><a href="/votes/' + v.public_id + '/">' + v.title + '</a> ' + statusBadge + resultBadge + '</h6>';
                    html += '<p class="small mb-1">' + (v.description || '') + '</p>';
                    html += '<p class="small text-muted mb-0">Ends: ' + new Date(v.end_at).toLocaleString() + '</p>';
                    html += '</div>';
                }}
                
                container.innerHTML = html;
            }} catch (e) {{
                console.error('Error loading votes:', e);
            }}
        }}
        
        loadDraftVotes();
        </script>
        """
    
    # Add document_content to the template
    content = content.replace('{document_content}', document_content)

    # Use ML number for title if approved and available, otherwise use draft name (submission ID)
    if draft.get('status') == 'approved' and draft.get('ml_number'):
        title_id = draft.get('ml_number')
    else:
        title_id = draft['name']
    
    # Generate Hypothesis configuration for this document
    hypothesis_config = generate_hypothesis_config(document_name=draft['name'], document_type='draft')
    
    return BASE_TEMPLATE.format(
        title=f"{title_id} - MLGH", 
        theme=current_theme, 
        user_menu=user_menu, 
        content=content, 
        build_number=BUILD_NUMBER,
        hypothesis_config=hypothesis_config
    )

@app.route('/doc/draft/<draft_name>/comments/', methods=['GET', 'POST'])
@require_auth
def draft_comments(draft_name):
    # First try to find in DRAFTS (published documents)
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)
    
    # If not found in DRAFTS, try to find as a submission ID
    submission = None
    if not draft:
        submission = Submission.query.filter_by(id=draft_name).first()
        if submission:
            draft = {
                'name': submission.id,
                'title': submission.title,
                'authors': submission.authors,
                'status': submission.status,
                'group': submission.group,
                'date': submission.submitted_at.strftime('%Y-%m-%d') if submission.submitted_at else '',
                'ml_number': submission.ml_number
            }
    
    if not draft:
        return "Document not found", 404
    
    # Get display ID (ML number if available, otherwise draft name)
    display_id = draft.get('ml_number') or draft_name

    user_menu = generate_user_menu()
    current_theme = session.get('theme', get_current_user().get('theme', 'dark') if get_current_user() else 'dark')
    current_user = get_current_user()

    # Handle new comment submission
    if request.method == 'POST':
        action = request.form.get('action', 'comment')
        
        if action == 'comment':
            comment_text = request.form.get('comment', '').strip()
            if comment_text:
                # Create new comment in database
                new_comment = Comment(
                    draft_name=draft_name,
                    text=comment_text,
                    author=current_user['name']
                )
                db.session.add(new_comment)
                db.session.commit()
                
                # Add to document history
                add_to_document_history(draft_name, 'Comment added', current_user['name'], f'Added comment: {comment_text[:50]}...')
                
                flash('Comment added successfully!', 'success')
                return redirect(f'/doc/draft/{draft_name}/comments/')
            else:
                flash('Please enter a comment.', 'error')
        
        elif action == 'like':
            comment_id = request.form.get('comment_id')
            if comment_id:
                liked = toggle_comment_like(draft_name, comment_id, current_user['name'])
                action_text = 'liked' if liked else 'unliked'
                flash(f'Comment {action_text}!', 'success')
                return redirect(f'/doc/draft/{draft_name}/comments/')
            else:
                flash('Invalid comment ID.', 'error')
        
        elif action == 'reply':
            parent_comment_id = request.form.get('parent_comment_id')
            reply_text = request.form.get('reply_text', '').strip()
            if reply_text and parent_comment_id:
                add_comment_reply(draft_name, parent_comment_id, reply_text, current_user)
                flash('Reply added successfully!', 'success')
                return redirect(f'/doc/draft/{draft_name}/comments/')
            else:
                flash('Please enter a reply.', 'error')
    
        elif action == 'edit':
            comment_id = request.form.get('comment_id')
            new_text = request.form.get('new_text', '').strip()
            if comment_id and new_text:
                comment = Comment.query.filter_by(id=int(comment_id)).first()
                if comment and comment.author == current_user['name']:
                    # Check time limit
                    time_diff = datetime.utcnow() - comment.timestamp
                    time_limit = timedelta(minutes=EDIT_DELETE_TIME_MINUTES)
                    if time_diff <= time_limit and not comment.is_deleted:
                        # Store original text if first edit
                        if not comment.original_text:
                            comment.original_text = comment.text
                        comment.text = new_text
                        comment.edited_at = datetime.utcnow()
                        db.session.commit()
                        flash('Comment updated successfully!', 'success')
                    else:
                        flash('Edit time limit has expired.', 'error')
                else:
                    flash('You can only edit your own comments.', 'error')
                return redirect(f'/doc/draft/{draft_name}/comments/')
            else:
                flash('Invalid comment or empty text.', 'error')
        
        elif action == 'delete':
            comment_id = request.form.get('comment_id')
            if comment_id:
                comment = Comment.query.filter_by(id=int(comment_id)).first()
                if comment and comment.author == current_user['name']:
                    # Check time limit
                    time_diff = datetime.utcnow() - comment.timestamp
                    time_limit = timedelta(minutes=EDIT_DELETE_TIME_MINUTES)
                    if time_diff <= time_limit and not comment.is_deleted:
                        comment.is_deleted = True
                        comment.text = '[Deleted]'
                        db.session.commit()
                        flash('Comment deleted successfully!', 'success')
                    else:
                        flash('Delete time limit has expired.', 'error')
                else:
                    flash('You can only delete your own comments.', 'error')
                return redirect(f'/doc/draft/{draft_name}/comments/')
            else:
                flash('Invalid comment ID.', 'error')

    # Get comments for this draft and build comment tree
    all_comments = build_comment_tree(draft_name)

    # Render the comment tree with nested replies
    comments_html = render_comment_tree(all_comments, draft_name)

    content = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/">Home</a></li>
                <li class="breadcrumb-item"><a href="/doc/all/">Documents</a></li>
                <li class="breadcrumb-item"><a href="/doc/draft/{draft_name}/">{display_id}</a></li>
                <li class="breadcrumb-item active">Comments</li>
            </ol>
        </nav>
        
        <h1>Comments for {display_id}</h1>
        <p class="lead">{draft['title']}</p>

        <div class="mb-4">
            <a href="/doc/draft/{draft_name}/" class="btn btn-secondary me-2">
                <i class="fas fa-arrow-left me-1"></i>Back to Draft
            </a>
            <a href="/doc/draft/{draft_name}/history/" class="btn btn-outline-secondary me-2">History</a>
            <a href="/doc/draft/{draft_name}/revisions/" class="btn btn-outline-secondary">Revisions</a>
        </div>
        
        <div class="row">
            <div class="col-md-8">
                <h3>Comments ({len(all_comments)})</h3>
                <div id="flash-messages"></div>
                {comments_html}
                
                <div class="card mt-4">
                    <div class="card-header">
                        <h5>Add a Comment</h5>
                    </div>
                    <div class="card-body">
                        <form method="POST">
                            <div class="mb-3">
                                <label for="comment" class="form-label">Your Comment</label>
                                <textarea class="form-control" id="comment" name="comment" rows="4" placeholder="Enter your comment here..." required></textarea>
                            </div>
                            <button type="submit" class="btn btn-primary">Submit Comment</button>
                        </form>
                    </div>
                </div>
            </div>
            
            <div class="col-md-4">
                <div class="card">
                    <div class="card-header">
                        <h5>Document Info</h5>
                    </div>
                    <div class="card-body">
                        <p><strong>Title:</strong> {draft['title']}</p>
                        <p><strong>Authors:</strong> {', '.join(draft['authors'])}</p>
                        <p><strong>Status:</strong> <span class="badge bg-secondary">{draft['status']}</span></p>
                        <p><strong>Last Updated:</strong> {draft['date']}</p>
                    </div>
                </div>
            </div>
        </div>
    </div>
    
    <script>
        function toggleLike(commentId) {{
            // Create a form to submit the like action
            const form = document.createElement('form');
            form.method = 'POST';
            form.style.display = 'none';

            const actionInput = document.createElement('input');
            actionInput.type = 'hidden';
            actionInput.name = 'action';
            actionInput.value = 'like';

            const commentIdInput = document.createElement('input');
            commentIdInput.type = 'hidden';
            commentIdInput.name = 'comment_id';
            commentIdInput.value = commentId;

            form.appendChild(actionInput);
            form.appendChild(commentIdInput);
            document.body.appendChild(form);
            form.submit();
        }}
        
        function toggleReply(commentId) {{
            // Find the reply form for this comment
            const replyForm = document.getElementById('reply-form-' + commentId);
            if (replyForm) {{
                // Toggle visibility
                if (replyForm.style.display === 'none' || replyForm.style.display === '') {{
                replyForm.style.display = 'block';
            }} else {{
                replyForm.style.display = 'none';
            }}
            }}
        }}

        function editComment(commentId) {{
            const commentCard = document.getElementById('comment-' + commentId);
            if (!commentCard) return;
            
            // Find the comment text element
            const commentText = commentCard.querySelector('p.mb-2');
            if (!commentText) return;
            
            const currentText = commentText.textContent.trim();
            
            // Create edit form
            const editForm = document.createElement('form');
            editForm.method = 'POST';
            editForm.style.marginTop = '10px';
            
            const textarea = document.createElement('textarea');
            textarea.className = 'form-control';
            textarea.name = 'new_text';
            textarea.rows = 3;
            textarea.value = currentText;
            textarea.required = true;
            
            const actionInput = document.createElement('input');
            actionInput.type = 'hidden';
            actionInput.name = 'action';
            actionInput.value = 'edit';
            
            const commentIdInput = document.createElement('input');
            commentIdInput.type = 'hidden';
            commentIdInput.name = 'comment_id';
            commentIdInput.value = commentId;
            
            const buttonDiv = document.createElement('div');
            buttonDiv.className = 'd-flex gap-2 mt-2';
            
            const saveBtn = document.createElement('button');
            saveBtn.type = 'submit';
            saveBtn.className = 'btn btn-sm btn-primary';
            saveBtn.textContent = 'Save';
            
            const cancelBtn = document.createElement('button');
            cancelBtn.type = 'button';
            cancelBtn.className = 'btn btn-sm btn-secondary';
            cancelBtn.textContent = 'Cancel';
            cancelBtn.onclick = function() {{
                commentText.style.display = 'block';
                editForm.remove();
            }};
            
            buttonDiv.appendChild(saveBtn);
            buttonDiv.appendChild(cancelBtn);
            
            editForm.appendChild(actionInput);
            editForm.appendChild(commentIdInput);
            editForm.appendChild(textarea);
            editForm.appendChild(buttonDiv);
            
            // Hide original text and show edit form
            commentText.style.display = 'none';
            commentText.parentNode.insertBefore(editForm, commentText.nextSibling);
        }}

        function deleteComment(commentId) {{
            if (!confirm('Are you sure you want to delete this comment? This action cannot be undone.')) {{
                return;
            }}
            
            // Create a form to submit the delete action
            const form = document.createElement('form');
            form.method = 'POST';
            form.style.display = 'none';
            
            const actionInput = document.createElement('input');
            actionInput.type = 'hidden';
            actionInput.name = 'action';
            actionInput.value = 'delete';
            
            const commentIdInput = document.createElement('input');
            commentIdInput.type = 'hidden';
            commentIdInput.name = 'comment_id';
            commentIdInput.value = commentId;
            
            form.appendChild(actionInput);
            form.appendChild(commentIdInput);
            document.body.appendChild(form);
            form.submit();
        }}
    </script>
"""

    return BASE_TEMPLATE.format(title=f"Comments - {draft_name}", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/doc/draft/<draft_name>/history/')
def draft_history(draft_name):
    # First try to find in DRAFTS (published documents)
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)
    
    # If not found in DRAFTS, try to find as a submission ID
    submission = None
    if not draft:
        submission = Submission.query.filter_by(id=draft_name).first()
        if submission:
            draft = {
                'name': submission.id,
                'title': submission.title,
                'authors': submission.authors,
                'status': submission.status,
                'group': submission.group,
                'date': submission.submitted_at.strftime('%Y-%m-%d') if submission.submitted_at else '',
                'ml_number': submission.ml_number,
            }
    
    if not draft:
        return "Document not found", 404
    
    # Determine display ID (ML-Draft-XXX or internal ID)
    display_id = draft.get('ml_number', draft_name) or draft_name
    
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')

    # Get history for this draft
    history = DocumentHistory.query.filter_by(draft_name=draft_name).order_by(DocumentHistory.timestamp.desc()).all()
    
    history_html = ""
    if history:
        for entry in history:
            history_html += f"""
            <div class="card mb-3">
                        <div class="card-body">
                    <div class="d-flex justify-content-between align-items-start mb-2">
                        <span class="badge bg-primary">{entry.action}</span>
                        <small class="text-muted">{entry.timestamp.strftime('%Y-%m-%d %H:%M')}</small>
                            </div>
                    <p class="mb-1"><strong>User:</strong> {entry.user}</p>
                    <p class="mb-0">{entry.details}</p>
                        </div>
                    </div>
            """
    else:
        history_html = """
        <div class="alert alert-info">
            <i class="fas fa-info-circle me-2"></i>
            No history available for this draft.
        </div>
        """
    
    content = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/">Home</a></li>
                <li class="breadcrumb-item"><a href="/doc/all/">Documents</a></li>
                <li class="breadcrumb-item"><a href="/doc/draft/{draft_name}/">{display_id}</a></li>
                <li class="breadcrumb-item active">History</li>
            </ol>
        </nav>
        
        <h1>History for {display_id}</h1>
        <p class="lead">{draft['title']}</p>
        
        <div class="mb-4">
            <a href="/doc/draft/{draft_name}/" class="btn btn-secondary me-2">
                <i class="fas fa-arrow-left me-1"></i>Back to Draft
            </a>
            <a href="/doc/draft/{draft_name}/comments/" class="btn btn-outline-secondary me-2">Comments</a>
            <a href="/doc/draft/{draft_name}/revisions/" class="btn btn-outline-secondary">Revisions</a>
        </div>

                {history_html}
            </div>
    """

    return BASE_TEMPLATE.format(title=f"History - {display_id}", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/doc/draft/<draft_name>/follow/', methods=['POST'])
def follow_draft(draft_name):
    current_user = get_current_user()
    if not current_user:
        flash('You must be logged in to follow documents.', 'error')
        return redirect(url_for('draft_detail', draft_name=draft_name))

    # Check if already following
    existing_follow = UserFollow.query.filter_by(user_id=current_user['id'], draft_name=draft_name).first()
    if existing_follow:
        flash('You are already following this document.', 'info')
    else:
        # Get notification level from form, default to 'all'
        notification_level = request.form.get('notification_level', 'all')
        follow = UserFollow(
            user_id=current_user['id'],
            draft_name=draft_name,
            notification_level=notification_level
        )
        db.session.add(follow)
        db.session.commit()
        level_desc = UserFollow.NOTIFICATION_LEVELS.get(notification_level, 'All changes and comments')
        flash(f'You are now following this document with {level_desc.lower()} notifications.', 'success')

    return redirect(url_for('draft_detail', draft_name=draft_name))

@app.route('/doc/draft/<draft_name>/unfollow/', methods=['POST'])
def unfollow_draft(draft_name):
    current_user = get_current_user()
    if not current_user:
        flash('You must be logged in to unfollow documents.', 'error')
        return redirect(url_for('draft_detail', draft_name=draft_name))

    follow = UserFollow.query.filter_by(user_id=current_user['id'], draft_name=draft_name).first()
    if follow:
        db.session.delete(follow)
        db.session.commit()
        flash('You have stopped following this document.', 'success')
    else:
        flash('You were not following this document.', 'info')

    return redirect(url_for('draft_detail', draft_name=draft_name))

@app.route('/doc/draft/<draft_name>/update-notification/', methods=['POST'])
def update_notification_level(draft_name):
    current_user = get_current_user()
    if not current_user:
        flash('You must be logged in to update notification settings.', 'error')
        return redirect(url_for('draft_detail', draft_name=draft_name))

    follow = UserFollow.query.filter_by(user_id=current_user['id'], draft_name=draft_name).first()
    if not follow:
        flash('You are not following this document.', 'error')
        return redirect(url_for('draft_detail', draft_name=draft_name))

    notification_level = request.form.get('notification_level', 'all')
    if notification_level in UserFollow.NOTIFICATION_LEVELS:
        follow.notification_level = notification_level
        db.session.commit()
        level_desc = UserFollow.NOTIFICATION_LEVELS[notification_level]
        flash(f'Notification level updated to: {level_desc}', 'success')
    else:
        flash('Invalid notification level.', 'error')

    return redirect(url_for('draft_detail', draft_name=draft_name))

@app.route('/doc/draft/<draft_name>/revisions/')
def draft_revisions(draft_name):
    current_user = get_current_user()
    # First try to find in DRAFTS (published documents)
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)
    
    # If not found in DRAFTS, try to find as a submission ID
    submission = None
    original_submission_id = None
    if not draft:
        submission = Submission.query.filter_by(id=draft_name).first()
        if submission:
            # Determine the original submission ID
            # If this submission is a revision, use its parent_draft_name
            # Otherwise, this IS the original
            if getattr(submission, 'is_revision', False) and getattr(submission, 'parent_draft_name', ''):
                original_submission_id = submission.parent_draft_name
            else:
                original_submission_id = submission.id
            
            # For the revisions page, just show the requested draft as-is
            # Don't try to find the "latest" - show what was requested
            draft = {
                'name': submission.id,
                'title': submission.title,
                'authors': submission.authors,
                'status': submission.status,
                'group': submission.group,
                'date': submission.submitted_at.strftime('%Y-%m-%d') if submission.submitted_at else '',
                'rev': getattr(submission, 'revision_number', '00') or '00',
                'pages': submission.pages or 1,
                'words': submission.words or 0,
                'is_revision': getattr(submission, 'is_revision', False),
                'parent_draft_name': getattr(submission, 'parent_draft_name', ''),
                'original_submission_id': original_submission_id,  # Always the true original
                'ml_number': submission.ml_number,
            }
    
    if not draft:
        return "Document not found", 404
    
    # Determine display ID (ML-Draft-XXX or internal ID)
    display_id = draft.get('ml_number', draft_name) or draft_name

    # Calculate pages and words from document content if it's a submission
    calculated_pages = draft.get('pages', 1)
    calculated_words = draft.get('words', 0)
    
    if submission and submission.file_path and os.path.exists(submission.file_path):
        _, ext = os.path.splitext(submission.filename.lower())
        try:
            if ext in ['.txt', '.xml']:
                with open(submission.file_path, 'r', encoding='utf-8', errors='replace') as f:
                    document_content = f.read()
                # Calculate words and pages from text
                words = len(document_content.split())
                # Estimate pages (assuming ~500 words per page)
                calculated_pages = max(1, (words + 499) // 500)
                calculated_words = words
            elif ext == '.docx':
                from docx import Document
                doc = Document(submission.file_path)
                content_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content_parts.append(cell.text)
                document_content = '\n\n'.join(content_parts)
                # Calculate words and pages
                words = len(document_content.split())
                calculated_pages = max(1, (words + 499) // 500)
                calculated_words = words
            elif ext == '.pdf':
                from PyPDF2 import PdfReader
                reader = PdfReader(submission.file_path)
                content_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(text)
                document_content = '\n\n'.join(content_parts)
                # Clean up PDF text
                import re
                document_content = re.sub(r'\n+', '\n', document_content)
                document_content = re.sub(r' +', ' ', document_content)
                # Calculate words and pages
                words = len(document_content.split())
                calculated_pages = len(reader.pages) if reader.pages else max(1, (words + 499) // 500)
                calculated_words = words
        except Exception as e:
            # If calculation fails, keep default values
            pass
        
        # Update draft with calculated values
        draft['pages'] = calculated_pages
        draft['words'] = calculated_words

    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')

    # Get the original submission ID to find all revisions
    original_id = draft.get('original_submission_id', draft['name'])
    
    # Get the original submission for display
    original_submission = Submission.query.filter_by(id=original_id).first()
    
    # Find all approved/published revisions for this draft
    all_revisions = Submission.query.filter(
        Submission.parent_draft_name == original_id,
        Submission.is_revision == True,
        Submission.status.in_(['approved', 'published'])
    ).order_by(Submission.revision_number.desc()).all()
    
    # Build revision list HTML - include ALL revisions (current and historical)
    revisions_list_html = ""
    for rev in all_revisions:
        status_badge_class = {
            'submitted': 'bg-warning text-dark',
            'approved': 'bg-success',
            'rejected': 'bg-danger',
            'published': 'bg-info'
        }.get(rev.status, 'bg-secondary')
        
        what_changed = getattr(rev, 'what_changed', '')
        what_changed_html = f'<p class="mb-2"><strong>What changed:</strong> {what_changed}</p>' if what_changed else ''
        
        # Check if this is the current revision
        is_current = (rev.id == draft['name'])
        current_badge = '<span class="badge bg-primary ms-2">Current</span>' if is_current else ''
        
        revisions_list_html += f"""
        <div class="card mb-3">
            <div class="card-header d-flex justify-content-between align-items-center">
                <h6 class="mb-0">
                    <a href="/doc/draft/{rev.id}/" class="text-decoration-none">Revision {rev.revision_number}</a>
                    {current_badge}
                </h6>
                <span class="badge {status_badge_class}">{rev.status.title()}</span>
            </div>
            <div class="card-body">
                <p class="mb-2"><strong>Published:</strong> {rev.approved_at.strftime('%Y-%m-%d') if rev.approved_at and rev.status == 'approved' else (rev.submitted_at.strftime('%Y-%m-%d') if rev.submitted_at else 'N/A')}</p>
                <p class="mb-2"><strong>Pages:</strong> {rev.pages or 1} | <strong>Words:</strong> {rev.words or 0}</p>
                {what_changed_html}
            </div>
        </div>
        """
    
    # Show revision history
    revisions_html = f"""
                <h4>Revision History</h4>
                {revisions_list_html if revisions_list_html else '<p class="text-muted">No revisions yet.</p>'}
                
                <div class="card mt-3">
                    <div class="card-header d-flex justify-content-between align-items-center">
                        <h6 class="mb-0">
                            <a href="/doc/draft/{original_id}/" class="text-decoration-none">Original Version (Rev 00)</a>
                        </h6>
                        <span class="badge bg-success">Approved</span>
                    </div>
                    <div class="card-body">
                        <p class="mb-2"><strong>Published:</strong> {original_submission.approved_at.strftime('%Y-%m-%d') if original_submission and original_submission.approved_at else (original_submission.submitted_at.strftime('%Y-%m-%d') if original_submission and original_submission.submitted_at else draft['date'])}</p>
                        <p class="mb-0"><strong>Pages:</strong> {original_submission.pages if original_submission else 1} | <strong>Words:</strong> {original_submission.words if original_submission else 0}</p>
                    </div>
                </div>

    <div class="alert alert-info mt-3">
        <i class="fas fa-info-circle me-2"></i>
        Detailed revision history and diff viewing would be implemented in a full datatracker system.
            </div>
    """

    content = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/">Home</a></li>
                <li class="breadcrumb-item"><a href="/doc/all/">Documents</a></li>
                <li class="breadcrumb-item"><a href="/doc/draft/{draft_name}/">{display_id}</a></li>
                <li class="breadcrumb-item active">Revisions</li>
            </ol>
        </nav>

        <h1>Revisions for {display_id}</h1>
        <p class="lead">{draft['title']}</p>

        <div class="mb-4">
            <a href="/doc/draft/{draft_name}/" class="btn btn-secondary me-2">
                <i class="fas fa-arrow-left me-1"></i>Back to Draft
            </a>
            {f'<a href="/submit/revision/{draft_name}/" class="btn btn-success me-2"><i class="fas fa-plus me-1"></i>Submit New Revision</a>' if current_user and draft.get('status') == 'approved' else ''}
            <a href="/doc/draft/{draft_name}/comments/" class="btn btn-outline-secondary me-2">Comments</a>
            <a href="/doc/draft/{draft_name}/history/" class="btn btn-outline-secondary">History</a>
        </div>

        {revisions_html}
    </div>
    """

    return BASE_TEMPLATE.format(title=f"Revisions - {display_id}", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/group/')
def groups():
    user_menu = generate_user_menu()
    current_theme = get_current_user().get('theme', 'dark') if get_current_user() else 'light'
    groups_html = ""
    for group in GROUPS:
        # Get chair information from database
        all_chairs = WorkingGroupChair.query.filter_by(group_acronym=group['acronym']).all()
        if all_chairs:
            chair_names = []
            for chair in all_chairs:
                chair_name = chair.chair_name
                if not chair.approved:
                    chair_name += " (Pending)"
                chair_names.append(chair_name)
            chair_display = ", ".join(chair_names)
        else:
            chair_display = "TBD"

        groups_html += f"""
        <div class="col-md-6">
            <div class="card mb-3">
                <div class="card-body">
                    <h5 class="card-title">
                        <a href="/group/{group['acronym']}/">{group['acronym']}</a>
                    </h5>
                    <p class="card-text">{group['name']}</p>
                    <div class="document-meta">
                        <span class="badge bg-primary">{group['type']}</span>
                        <span class="badge bg-success ms-2">{group['state']}</span>
                    </div>
                    <div class="mt-2">
                        <small class="text-muted">
                            Coordinator: {chair_display}<br>
                            {group['description']}
                        </small>
                    </div>
                </div>
            </div>
        </div>
        """

    # Get theme from session or user preference
    current_theme = session.get('theme', 'dark')

    content = f"""
    <div class="container mt-4">
        <div class="row">
            <div class="col-12">
                <h1 class="mb-4">Workgroups</h1>
                <p class="lead mb-4">Browse the Meta-Layer Desirable Properties workgroups.</p>

                <div class="row">
                    {groups_html}
                </div>
            </div>
        </div>
    </div>
    """

    return BASE_TEMPLATE.format(
        title="Workgroups - MLGH",
        theme=current_theme,
        content=content,
        user_menu=user_menu, build_number=BUILD_NUMBER, hypothesis_config="")
@app.route('/group/<acronym>/')
def group_detail(acronym):
    """Display individual workgroup details"""
    # Find the group - handle both full acronyms and short forms (DP1, DP2, etc.)
    group = None
    full_acronym = acronym  # Default to the URL parameter

    for g in GROUPS:
        if g['acronym'] == acronym:
            group = g
            full_acronym = g['acronym']
            break
        # Also check for short form (dp1 -> dp1-federated-auth, DP1 -> dp1-federated-auth)
        if acronym.lower().startswith('dp') and g['acronym'].startswith(acronym.lower() + '-'):
            group = g
            full_acronym = g['acronym']
            break

    if not group:
        return f"Workgroup '{acronym}' not found. Available: {[g['acronym'] for g in GROUPS]}", 404

    user_menu = generate_user_menu()
    current_user = get_current_user()

    # Membership is by user_id only
    is_member = False
    pending_member_request = False
    if current_user and current_user.get('id'):
        membership = WorkingGroupMember.query.filter_by(
            group_acronym=full_acronym,
            user_id=current_user['id']
        ).first()
        is_member = membership is not None
        if not is_member:
            req = WorkgroupMemberRequest.query.filter_by(
                group_acronym=full_acronym,
                user_id=current_user['id'],
                status='pending'
            ).first()
            pending_member_request = req is not None

    # Get chair information using the full acronym (coordinator by user_id only)
    all_chairs = WorkingGroupChair.query.filter_by(group_acronym=full_acronym).all()
    is_coordinator = False
    has_pending_coord_request = False
    if current_user and current_user.get('id') and all_chairs:
        for chair in all_chairs:
            if chair.user_id == current_user['id']:
                is_coordinator = True
                break
    if current_user and current_user.get('id') and not is_coordinator:
        cr = CoordinatorRequest.query.filter_by(
            group_acronym=full_acronym,
            user_id=current_user['id'],
            status='pending'
        ).first()
        has_pending_coord_request = cr is not None
    if all_chairs:
        approved_chairs = [chair.chair_name for chair in all_chairs if chair.approved]
        pending_chairs = [chair.chair_name for chair in all_chairs if not chair.approved]

        if approved_chairs:
            chair_name = ", ".join(approved_chairs)
            if pending_chairs:
                chair_name += f" (Pending: {', '.join(pending_chairs)})"
        else:
            chair_name = f"Pending: {', '.join(pending_chairs)}"
        chair_approved = len(approved_chairs) > 0
    else:
        chair_name = "TBD"
        chair_approved = False

    join_button = ""
    if current_user and pending_member_request:
        join_button = '<span class="badge bg-warning">Membership request pending</span>'
    elif current_user and not is_member:
        join_button = f'<button class="btn btn-primary" onclick="joinGroup(\'{full_acronym}\')">Join Workgroup</button>'
    elif current_user and is_member:
        join_button = f'<span class="badge bg-success">Member</span> <button class="btn btn-outline-danger btn-sm ms-2" onclick="leaveGroup(\'{full_acronym}\')">Leave</button>'

    # Coordinator status: show tag if coordinator, else pending or request button
    coord_request_ui = ""
    if current_user and current_user.get('id'):
        if is_coordinator:
            coord_request_ui = '<span class="badge bg-primary ms-2">Coordinator</span>'
        elif has_pending_coord_request:
            coord_request_ui = '<span class="badge bg-warning ms-2">Coordinator request pending</span>'
        else:
            coord_request_ui = f'<button class="btn btn-outline-secondary btn-sm ms-2" onclick="requestCoordinator(\'{full_acronym}\')">Request coordinator role</button>'

    # Coordinator management is done via request/approve in Admin only; no inline box on workgroup page
    chair_management = ""

    # Get theme from session or user preference
    current_theme = session.get('theme', current_user.get('theme', 'dark') if current_user else 'dark')

    content = f"""
    <div class="container mt-4">
        <div class="row">
            <div class="col-12">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/">Home</a></li>
                        <li class="breadcrumb-item"><a href="/group/">Workgroups</a></li>
                        <li class="breadcrumb-item active">{group['name']}</li>
            </ol>
        </nav>
        
                <div class="card mb-4">
                    <div class="card-body">
                        <div class="d-flex justify-content-between align-items-start">
                            <div>
                                <h1 class="card-title mb-2">{group['name']}</h1>
                                <p class="text-muted mb-3">{group['acronym'].upper()}</p>
                                <p class="card-text">{group['description']}</p>
                            </div>
                            <div class="text-end">
                                {join_button}
                                {coord_request_ui}
                            </div>
                        </div>
                    </div>
                </div>
        
        <div class="row">
            <div class="col-md-8">
                        <div class="card mb-4">
                            <div class="card-header">
                                <h5 class="mb-0">About</h5>
                </div>
                            <div class="card-body">
                                <p>{group.get('about', group['description'])}</p>
            </div>
                        </div>
                    </div>
            <div class="col-md-4">
                        <div class="card mb-4">
                    <div class="card-header">
                                <h5 class="mb-0">Leadership</h5>
                    </div>
                    <div class="card-body">
                                <p><strong>Coordinator:</strong> {chair_name}</p>
                                {'<span class="badge bg-warning">Pending Approval</span>' if not chair_approved and chair_name != "TBD" else ''}
                            </div>
                        </div>

                        {chair_management}
                    </div>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    function joinGroup(acronym) {{
        fetch(`/group/${{acronym}}/join`, {{
            method: 'POST',
            headers: {{
                'Content-Type': 'application/json',
            }}
        }})
        .then(response => response.json())
        .then(data => {{
            if (data.success) {{
                location.reload();
            }} else {{
                alert('Error joining group: ' + data.message);
            }}
        }})
        .catch(error => {{
            console.error('Error:', error);
            alert('Error joining group');
        }});
    }}

    function leaveGroup(acronym) {{
        fetch(`/group/${{acronym}}/leave`, {{
            method: 'POST',
            headers: {{
                'Content-Type': 'application/json',
            }}
        }})
        .then(response => response.json())
        .then(data => {{
            if (data.success) {{
                location.reload();
            }} else {{
                alert('Error leaving group: ' + data.message);
            }}
        }})
        .catch(error => {{
            console.error('Error:', error);
            alert('Error leaving group');
        }});
    }}

    function requestCoordinator(acronym) {{
        fetch(`/group/${{acronym}}/request_coordinator`, {{
            method: 'POST',
            headers: {{ 'Content-Type': 'application/json' }}
        }})
        .then(response => response.json())
        .then(data => {{
            if (data.success) {{ location.reload(); }}
            else {{ alert(data.message || 'Request failed'); }}
        }})
        .catch(() => alert('Request failed'));
    }}

    function addChair(acronym) {{
        const input = document.getElementById(`new-chair-input-${{acronym}}`);
        const chairName = input.value.trim();
        if (!chairName) {{
            alert('Please enter a chair name');
            return;
        }}

        fetch(`/group/${{acronym}}/add_chair`, {{
            method: 'POST',
            headers: {{
                'Content-Type': 'application/json',
            }},
            body: JSON.stringify({{ chair_name: chairName }})
        }})
        .then(response => response.json())
        .then(data => {{
            if (data.success) {{
                location.reload();
            }} else {{
                alert('Error adding chair: ' + data.message);
            }}
        }})
        .catch(error => {{
            console.error('Error:', error);
            alert('Error adding chair');
        }});
    }}

    function updateChairs(acronym) {{
        const select = document.getElementById(`chair-select-${{acronym}}`);
        const selectedOptions = Array.from(select.selectedOptions);
        const chairIds = selectedOptions.map(option => parseInt(option.value));

        fetch(`/group/${{acronym}}/update_chairs`, {{
            method: 'POST',
            headers: {{
                'Content-Type': 'application/json',
            }},
            body: JSON.stringify({{ chair_ids: chairIds }})
        }})
        .then(response => response.json())
        .then(data => {{
            if (data.success) {{
                location.reload();
            }} else {{
                alert('Error updating chairs: ' + data.message);
            }}
        }})
        .catch(error => {{
            console.error('Error:', error);
            alert('Error updating chairs');
        }});
    }}

    function removeChair(acronym) {{
        const select = document.getElementById(`chair-select-${{acronym}}`);
        const selectedOptions = Array.from(select.selectedOptions);

        if (selectedOptions.length === 0) {{
            alert('Please select chairs to remove');
            return;
        }}

        if (confirm('Are you sure you want to remove ' + selectedOptions.length + ' chair(s)?')) {{
            const chairIds = selectedOptions.map(option => parseInt(option.value));

            fetch(`/group/${{acronym}}/remove_chairs`, {{
                method: 'POST',
                headers: {{
                    'Content-Type': 'application/json',
                }},
                body: JSON.stringify({{ chair_ids: chairIds }})
            }})
            .then(response => response.json())
            .then(data => {{
                if (data.success) {{
                    location.reload();
                }} else {{
                    alert('Error removing chairs: ' + data.message);
                }}
            }})
            .catch(error => {{
                console.error('Error:', error);
                alert('Error removing chairs');
            }});
        }}
    }}
    </script>
    """

    return BASE_TEMPLATE.format(
        title=f"{group['name']} - MLGH",
        theme=current_theme,
        content=content,
        user_menu=user_menu, build_number=BUILD_NUMBER, hypothesis_config="")

@app.route('/group/<acronym>/join', methods=['POST'])
@require_auth
def join_group(acronym):
    current_user = get_current_user()
    if not current_user:
        return jsonify({'success': False, 'message': 'Not authenticated'}), 401

    # Resolve full acronym
    full_acronym = acronym
    group = None
    for g in GROUPS:
        if g['acronym'] == acronym or (acronym.lower().startswith('dp') and g['acronym'].startswith(acronym.lower() + '-')):
            full_acronym = g['acronym']
            group = g
            break
    if not group:
        group = next((g for g in GROUPS if g['acronym'] == acronym), None)
    if not group:
        full_acronym = acronym

    # Membership is by user_id only
    user_id = current_user.get('id')
    if not user_id:
        return jsonify({'success': False, 'message': 'You must be logged in to join'}), 400
    existing = WorkingGroupMember.query.filter_by(
        group_acronym=full_acronym,
        user_id=user_id
    ).first()
    if existing:
        return jsonify({'success': False, 'message': 'Already a member'}), 400

    # If workgroup requires approval, create a request instead of immediate membership (default: no approval)
    require_approval = group.get('members_require_approval', False) if group else False
    if require_approval:
        pending = WorkgroupMemberRequest.query.filter_by(
            group_acronym=full_acronym,
            user_id=user_id,
            status='pending'
        ).first()
        if pending:
            return jsonify({'success': False, 'message': 'Membership request already pending'}), 400
        req = WorkgroupMemberRequest(
            group_acronym=full_acronym,
            user_id=user_id,
            user_name=current_user['name'],
            status='pending'
        )
        db.session.add(req)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Membership requested; pending approval'})

    # Default: instant join (store user_id for stable membership)
    membership = WorkingGroupMember(
        group_acronym=full_acronym,
        user_id=user_id,
        user_name=current_user.get('name') or current_user.get('username', '')
    )
    db.session.add(membership)
    db.session.commit()
    return jsonify({'success': True, 'message': 'Joined successfully'})

@app.route('/group/<acronym>/leave', methods=['POST'])
@require_auth
def leave_group(acronym):
    current_user = get_current_user()
    if not current_user:
        return jsonify({'success': False, 'message': 'Not authenticated'}), 401

    # Resolve full acronym (same as join) so URL variant matches DB
    full_acronym = acronym
    for g in GROUPS:
        if g['acronym'] == acronym or (acronym.lower().startswith('dp') and g['acronym'].startswith(acronym.lower() + '-')):
            full_acronym = g['acronym']
            break

    # Membership is by user_id only
    user_id = current_user.get('id')
    if not user_id:
        return jsonify({'success': False, 'message': 'You must be logged in to leave'}), 400
    
    # Debug: log what we're looking for
    print(f"DEBUG leave_group: acronym={acronym}, user_id={user_id}, full_acronym={full_acronym}")
    
    # Check what's actually in the database
    all_memberships = WorkingGroupMember.query.filter_by(user_id=user_id).all()
    print(f"DEBUG: All memberships for user {user_id}: {[(m.group_acronym, m.user_name) for m in all_memberships]}")
    
    membership = WorkingGroupMember.query.filter_by(
        group_acronym=full_acronym,
        user_id=user_id
    ).first()
    print(f"DEBUG leave_group: membership found={membership is not None}")

    if not membership:
        error_msg = 'Not a member (user_id=' + str(user_id) + ', group=' + str(full_acronym) + ', original_acronym=' + str(acronym) + ')'
        return jsonify({'success': False, 'message': error_msg}), 400

    db.session.delete(membership)
    db.session.commit()

    return jsonify({'success': True, 'message': 'Left successfully'})

@app.route('/group/<acronym>/request_coordinator', methods=['POST'])
@require_auth
def request_coordinator(acronym):
    """User requests coordinator role; creates pending request (we have their user id)."""
    current_user = get_current_user()
    if not current_user:
        return jsonify({'success': False, 'message': 'Not authenticated'}), 401

    # Resolve full acronym
    full_acronym = acronym
    for g in GROUPS:
        if g['acronym'] == acronym or (acronym.lower().startswith('dp') and g['acronym'].startswith(acronym.lower() + '-')):
            full_acronym = g['acronym']
            break

    # Already a coordinator? (by user_id only)
    existing_chair = WorkingGroupChair.query.filter_by(
        group_acronym=full_acronym,
        user_id=current_user['id']
    ).first()
    if existing_chair:
        return jsonify({'success': False, 'message': 'You are already a coordinator'}), 400

    # Already have a pending request? (by user_id only)
    existing = CoordinatorRequest.query.filter_by(
        group_acronym=full_acronym,
        user_id=current_user['id'],
        status='pending'
    ).first()
    if existing:
        return jsonify({'success': False, 'message': 'You already have a pending request'}), 400

    req = CoordinatorRequest(
        group_acronym=full_acronym,
        user_id=current_user['id'],
        username=current_user.get('username', ''),
        display_name=current_user.get('name') or current_user.get('displayName') or current_user.get('username', ''),
        status='pending'
    )
    db.session.add(req)
    db.session.commit()
    return jsonify({'success': True, 'message': 'Coordinator role requested; pending approval'})

@app.route('/group/<acronym>/add_chair', methods=['POST'])
@require_role('admin')
def add_group_chair(acronym):
    data = request.get_json()
    chair_name = data.get('chair_name', '').strip()
    if not chair_name:
        return jsonify({'success': False, 'message': 'Coordinator name required'}), 400

    # Check if chair already exists
    existing = WorkingGroupChair.query.filter_by(group_acronym=acronym, chair_name=chair_name).first()
    if existing:
        return jsonify({'success': False, 'message': 'Coordinator already exists'}), 400

    # Add new chair (unapproved)
    chair = WorkingGroupChair(
        group_acronym=acronym,
        chair_name=chair_name,
        approved=False
    )
    db.session.add(chair)
    db.session.commit()

    return jsonify({'success': True, 'message': 'Coordinator added successfully'})

@app.route('/group/<acronym>/update_chairs', methods=['POST'])
@require_role('admin')
def update_group_chairs(acronym):
    data = request.get_json()
    chair_ids = data.get('chair_ids', [])

    # Mark all chairs as unapproved first
    WorkingGroupChair.query.filter_by(group_acronym=acronym).update({'approved': False})

    # Approve selected chairs
    if chair_ids:
        WorkingGroupChair.query.filter(
            WorkingGroupChair.group_acronym == acronym,
            WorkingGroupChair.id.in_(chair_ids)
        ).update({'approved': True})

    db.session.commit()

    return jsonify({'success': True, 'message': 'Coordinators updated successfully'})

@app.route('/group/<acronym>/remove_chairs', methods=['POST'])
@require_role('admin')
def remove_group_chairs(acronym):
    data = request.get_json()
    chair_ids = data.get('chair_ids', [])

    if not chair_ids:
        return jsonify({'success': False, 'message': 'No chairs selected'}), 400

    # Remove selected chairs
    WorkingGroupChair.query.filter(
        WorkingGroupChair.group_acronym == acronym,
        WorkingGroupChair.id.in_(chair_ids)
    ).delete()

    db.session.commit()

    return jsonify({'success': True, 'message': 'Coordinators removed successfully'})

@app.route('/person/')
def people():
    """People directory: list users; admins get Add as coordinator and other actions."""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    is_admin = current_user and current_user.get('role') == 'admin'
    is_editor_or_admin = current_user and current_user.get('role') in ('admin', 'editor')

    users = User.query.order_by(User.username).all()
    group_options = ''.join(
        f'<option value="{g["acronym"]}">{g["acronym"]}</option>' for g in GROUPS
    )
    rows = []
    for u in users:
        display = u.name or u.displayName or u.oauthName or u.username
        # Coordinator: workgroups where they are a coordinator
        coord_groups = WorkingGroupChair.query.filter_by(user_id=u.id).all()
        coord_acronyms = ' '.join(c.group_acronym for c in coord_groups)
        coord_badges = ' '.join(
            f'<span class="badge bg-secondary me-1">{c.group_acronym}</span>'
            for c in coord_groups
        ) if coord_groups else '<span class="text-muted">—</span>'
        # Member: workgroups they are a member of
        member_groups = WorkingGroupMember.query.filter_by(user_id=u.id).all()
        member_acronyms = ' '.join(m.group_acronym for m in member_groups)
        member_badges = ' '.join(
            f'<span class="badge bg-info me-1">{m.group_acronym}</span>'
            for m in member_groups
        ) if member_groups else '<span class="text-muted">—</span>'
        # Combined for filter: member + coordinator acronyms
        all_groups = (member_acronyms + ' ' + coord_acronyms).strip() or ''
        # Role (only shown to editor/admin)
        role_badge = f'<span class="badge bg-{"danger" if u.role == "admin" else "warning" if u.role == "editor" else "secondary"}">{u.role or "user"}</span>'
        # Last active (last_login)
        if u.last_login:
            last_active = u.last_login.strftime('%Y-%m-%d')
        else:
            last_active = '<span class="text-muted">Never</span>'
        # Submissions count (match by submitted_by string to user's names)
        name_variants = [x for x in (u.name, u.displayName, u.oauthName, u.username) if x]
        submissions_count = Submission.query.filter(Submission.submitted_by.in_(name_variants)).count() if name_variants else 0
        # Documents followed count
        follows_count = UserFollow.query.filter_by(user_id=u.id).count()
        # Comments count (site document comments, not Hypothesis)
        comments_count = Comment.query.filter(Comment.author.in_(name_variants)).count() if name_variants else 0
        if is_admin:
            actions_td = f'<td><a href="/admin/users/{u.id}/add-coordinator" class="btn btn-outline-primary btn-sm">Add as coordinator</a></td>'
        else:
            actions_td = ''
        # data-search and data-groups for client-side filter
        search_text = f"{display} {u.username}".lower()
        role_td = f'<td>{role_badge}</td>' if is_editor_or_admin else ''
        rows.append(f"""
        <tr data-search="{search_text}" data-groups="{all_groups}">
            <td><strong>{display}</strong><br><small class="text-muted">@{u.username}</small></td>
            {role_td}
            <td>{member_badges}</td>
            <td>{coord_badges}</td>
            <td>{last_active}</td>
            <td>{submissions_count}</td>
            <td>{follows_count}</td>
            <td>{comments_count}</td>
            {actions_td}
        </tr>
        """)

    num_cols = 7 + (1 if is_editor_or_admin else 0) + (1 if is_admin else 0)  # Name, [Role], Member, Coordinator, Last active, Submissions, Documents followed, Comments, [Actions]
    table_rows = ''.join(rows) if rows else f'<tr><td colspan="{num_cols}" class="text-center text-muted py-4">No users yet.</td></tr>'
    role_th = '<th>Role</th>' if is_editor_or_admin else ''
    actions_th = '<th>Actions</th>' if is_admin else ''
    content = f"""
    <div class="container mt-4">
        <nav aria-label="breadcrumb">
            <ol class="breadcrumb">
                <li class="breadcrumb-item"><a href="/">Home</a></li>
                <li class="breadcrumb-item active">People</li>
            </ol>
        </nav>
        <h1 class="mb-2">People</h1>
        <p class="text-muted mb-4">Directory of MLGH participants. Member and coordinator workgroups and activity at a glance.</p>
        <div class="card">
            <div class="card-body">
                <div class="row g-2 mb-3">
                    <div class="col-md-6">
                        <label class="form-label small text-muted mb-0">Search</label>
                        <input type="text" id="people-search" class="form-control" placeholder="Type to search by name or username..." autocomplete="off">
                    </div>
                    <div class="col-md-4">
                        <label class="form-label small text-muted mb-0">Workgroup</label>
                        <select id="people-workgroup" class="form-select">
                            <option value="">All workgroups</option>
                            {group_options}
                        </select>
                    </div>
                </div>
            </div>
            <div class="card-body p-0 pt-0">
                <div class="table-responsive">
                    <table class="table table-hover mb-0" id="people-table">
                        <thead class="table-light">
                            <tr>
                                <th>Name</th>
                                {role_th}
                                <th>Member</th>
                                <th>Coordinator</th>
                                <th>Last active</th>
                                <th>Submissions</th>
                                <th>Documents followed</th>
                                <th>Comments</th>
                                {actions_th}
                            </tr>
                        </thead>
                        <tbody>{table_rows}</tbody>
                    </table>
                </div>
            </div>
        </div>
    </div>
    <script>
    (function() {{
        var searchEl = document.getElementById('people-search');
        var workgroupEl = document.getElementById('people-workgroup');
        var rows = document.querySelectorAll('#people-table tbody tr[data-search]');
        function filterPeople() {{
            var q = (searchEl && searchEl.value) ? searchEl.value.toLowerCase().trim() : '';
            var group = (workgroupEl && workgroupEl.value) ? workgroupEl.value.trim() : '';
            rows.forEach(function(tr) {{
                var show = true;
                if (q && tr.getAttribute('data-search').indexOf(q) === -1) show = false;
                if (group) {{
                    var groups = (tr.getAttribute('data-groups') || '').split(/\\s+/).filter(Boolean);
                    if (groups.indexOf(group) === -1) show = false;
                }}
                tr.style.display = show ? '' : 'none';
            }});
        }}
        if (searchEl) searchEl.addEventListener('input', filterPeople);
        if (searchEl) searchEl.addEventListener('keyup', filterPeople);
        if (workgroupEl) workgroupEl.addEventListener('change', filterPeople);
    }})();
    </script>
    """
    return render_page("People - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/meeting/')
def meetings():
    """Meetings - coming soon"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')

    content = """
    <div class="container mt-4">
        <div class="row justify-content-center">
            <div class="col-md-8">
                <div class="text-center">
                    <i class="fas fa-calendar fa-4x text-muted mb-4"></i>
                    <h1 class="mb-3">Meetings</h1>
                    <p class="lead text-muted mb-4">Coming Soon</p>
                    <p class="mb-4">Information about upcoming MLGH meetings and sessions will be available here. Stay tuned for announcements about our first events.</p>
                    <a href="/" class="btn btn-primary">Return to Home</a>
                </div>
            </div>
        </div>
    </div>
    """

    return render_page("Meetings - MLGH", content, theme=session.get('theme', 'dark'), user_menu=user_menu)

# ============================================================================
# Role Images Pages
# ============================================================================

@app.route('/roles/<role_slug>/images/')
def role_images_gallery(role_slug):
    """Gallery of role image proposals with voting"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    is_admin = current_user and current_user.get('role') == 'admin'
    
    # Load role for display name and link back to role
    role = Role.query.filter_by(role_slug=role_slug).first()
    role_title = role.title_guild if role else role_slug
    
    content = f"""
    <div class="container mt-4">
        <div class="row mb-4">
            <div class="col-md-12">
                <h1>Role Images: <a href="/roles/{role_slug}/" class="text-decoration-none">{role_title}</a></h1>
                <p class="lead">Community-proposed images for this role</p>
            </div>
        </div>
        
        <div class="row mb-4">
            <div class="col-md-6">
                <label for="sort-select" class="form-label">Sort by:</label>
                <select id="sort-select" class="form-select" onchange="loadImages()">
                    <option value="net_score">Net Score (Votes)</option>
                    <option value="upvotes">Most Upvotes</option>
                    <option value="date">Most Recent</option>
                </select>
            </div>
            <div class="col-md-6 text-end">
                {'<button class="btn btn-primary" data-bs-toggle="modal" data-bs-target="#submitImageModal"><i class="fas fa-plus me-2"></i>Submit Image</button>' if current_user else '<a href="/login/" class="btn btn-primary"><i class="fas fa-sign-in-alt me-2"></i>Login to Submit</a>'}
            </div>
        </div>
        
        <div id="images-container" class="row">
            <div class="col-12 text-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        </div>
    </div>
    
    <!-- Submit Image Modal -->
    <div class="modal fade" id="submitImageModal" tabindex="-1">
        <div class="modal-dialog">
            <div class="modal-content">
                <div class="modal-header">
                    <h5 class="modal-title">Submit Role Image</h5>
                    <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                </div>
                <div class="modal-body">
                    <form id="submitImageForm">
                        <div class="mb-3">
                            <label class="form-label">Source Type</label>
                            <select class="form-select" id="sourceType" onchange="toggleSourceFields()">
                                <option value="upload">Upload Image File</option>
                                <option value="url">Image URL</option>
                                <option value="ordinal">Bitcoin Ordinal</option>
                            </select>
                        </div>
                        
                        <div id="uploadField" class="mb-3">
                            <label for="imageFile" class="form-label">Image File</label>
                            <input type="file" class="form-control" id="imageFile" accept="image/*">
                            <small class="text-muted">Max 600×600 px, 5MB. Formats: PNG, JPG, GIF, WebP, SVG</small>
                        </div>
                        
                        <div id="urlField" class="mb-3" style="display:none;">
                            <label for="imageUrl" class="form-label">Image URL</label>
                            <input type="url" class="form-control" id="imageUrl" placeholder="https://example.com/image.png">
                        </div>
                        
                        <div id="ordinalFields" style="display:none;">
                            <div class="mb-3">
                                <label for="inscriptionId" class="form-label">Inscription ID</label>
                                <input type="text" class="form-control" id="inscriptionId" placeholder="a455e1c4...e9aa72i0">
                            </div>
                            <div class="mb-3">
                                <label for="contentType" class="form-label">Content Type</label>
                                <input type="text" class="form-control" id="contentType" value="image/png">
                            </div>
                        </div>
                    </form>
                </div>
                <div class="modal-footer">
                    <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                    <button type="button" class="btn btn-primary" onclick="submitImage()">Submit</button>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    const roleSlug = '{role_slug}';
    const isAdmin = {'true' if is_admin else 'false'};
    
    function toggleSourceFields() {{
        const sourceType = document.getElementById('sourceType').value;
        document.getElementById('uploadField').style.display = sourceType === 'upload' ? 'block' : 'none';
        document.getElementById('urlField').style.display = sourceType === 'url' ? 'block' : 'none';
        document.getElementById('ordinalFields').style.display = sourceType === 'ordinal' ? 'block' : 'none';
    }}
    
    async function loadImages() {{
        const sortBy = document.getElementById('sort-select').value;
        const container = document.getElementById('images-container');
        
        try {{
            const response = await fetch(`/api/roles/${{roleSlug}}/images/?sort=${{sortBy}}`);
            const data = await response.json();
            
            if (data.images.length === 0) {{
                container.innerHTML = '<div class="col-12 text-center py-5"><p class="text-muted">No images yet. Be the first to submit one!</p></div>';
                return;
            }}
            
            container.innerHTML = data.images.map(img => {{
                // Determine image source
                let imgSrc = img.image_url;
                if (img.source_type === 'upload' && img.file_path) {{
                    imgSrc = `/uploads/role_images/${{img.file_path.split('/').pop()}}`;
                }} else if (img.source_type === 'ordinal') {{
                    imgSrc = `https://ordinals.com/content/${{img.inscription_id}}`;
                }}
                
                return `
                <div class="col-md-4 mb-4">
                    <div class="card h-100">
                        <a href="/roles/${{roleSlug}}/images/${{img.id}}/">
                            <img src="${{imgSrc}}" class="card-img-top" alt="Role image" style="height: 250px; object-fit: cover;" onerror="this.src='data:image/svg+xml,%3Csvg xmlns=\\'http://www.w3.org/2000/svg\\' width=\\'200\\' height=\\'200\\'%3E%3Crect fill=\\'%23ddd\\' width=\\'200\\' height=\\'200\\'/%3E%3Ctext fill=\\'%23999\\' x=\\'50%25\\' y=\\'50%25\\' dominant-baseline=\\'middle\\' text-anchor=\\'middle\\' font-family=\\'sans-serif\\' font-size=\\'16\\'%3EImage not available%3C/text%3E%3C/svg%3E';">
                        </a>
                        <div class="card-body">
                            ${{img.is_primary ? '<span class="badge bg-success mb-2">Primary Image</span>' : ''}}
                            ${{img.is_hidden && isAdmin ? '<span class="badge bg-warning mb-2">Hidden</span>' : ''}}
                            <p class="small text-muted mb-2">
                                By ${{img.submitted_by_name}}<br>
                                ${{new Date(img.submitted_at).toLocaleDateString()}}
                            </p>
                            <div class="d-flex justify-content-between align-items-center">
                                <div class="btn-group">
                                    <button class="btn btn-sm ${{img.user_vote === 1 ? 'btn-success' : 'btn-outline-success'}}" onclick="vote('${{img.id}}', 1, event)">
                                        <i class="fas fa-thumbs-up"></i> ${{img.upvotes}}
                                    </button>
                                    <button class="btn btn-sm ${{img.user_vote === -1 ? 'btn-danger' : 'btn-outline-danger'}}" onclick="vote('${{img.id}}', -1, event)">
                                        <i class="fas fa-thumbs-down"></i> ${{img.downvotes}}
                                    </button>
                                </div>
                                <span class="badge bg-primary">Score: ${{img.net_score}}</span>
                            </div>
                        </div>
                    </div>
                </div>
                `;
            }}).join('');
        }} catch (error) {{
            console.error('Error loading images:', error);
            container.innerHTML = '<div class="col-12 text-center py-5"><p class="text-danger">Error loading images</p></div>';
        }}
    }}
    
    async function vote(imageId, value, event) {{
        event.preventDefault();
        event.stopPropagation();
        
        try {{
            const response = await fetch(`/api/role-images/${{imageId}}/vote/`, {{
                method: 'POST',
                headers: {{'Content-Type': 'application/json'}},
                body: JSON.stringify({{value}})
            }});
            
            if (response.ok) {{
                loadImages();
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to vote');
            }}
        }} catch (error) {{
            console.error('Error voting:', error);
            alert('Error voting on image');
        }}
    }}
    
    async function submitImage() {{
        const sourceType = document.getElementById('sourceType').value;
        const formData = new FormData();
        formData.append('source_type', sourceType);
        
        if (sourceType === 'upload') {{
            const file = document.getElementById('imageFile').files[0];
            if (!file) {{
                alert('Please select a file');
                return;
            }}
            formData.append('file', file);
        }} else if (sourceType === 'url') {{
            const url = document.getElementById('imageUrl').value;
            if (!url) {{
                alert('Please enter an image URL');
                return;
            }}
            formData.append('image_url', url);
        }} else if (sourceType === 'ordinal') {{
            const inscriptionId = document.getElementById('inscriptionId').value;
            if (!inscriptionId) {{
                alert('Please enter an inscription ID');
                return;
            }}
            formData.append('inscription_id', inscriptionId);
            formData.append('content_type', document.getElementById('contentType').value);
            formData.append('chain', 'bitcoin');
        }}
        
        try {{
            const response = await fetch(`/api/roles/${{roleSlug}}/images/`, {{
                method: 'POST',
                body: formData
            }});
            
            if (response.ok) {{
                bootstrap.Modal.getInstance(document.getElementById('submitImageModal')).hide();
                document.getElementById('submitImageForm').reset();
                loadImages();
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to submit image');
            }}
        }} catch (error) {{
            console.error('Error submitting image:', error);
            alert('Error submitting image');
        }}
    }}
    
    // Load images on page load
    loadImages();
    </script>
    """
    
    return render_page(f"Role Images: {role_slug} - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/roles/<role_slug>/images/<image_id>/')
def role_image_detail(role_slug, image_id):
    """Detailed view of a single role image proposal"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    is_admin = current_user and current_user.get('role') == 'admin'
    
    # Fetch image
    image = RoleImage.query.get_or_404(image_id)
    
    # Check if hidden
    if image.is_hidden and not is_admin:
        flash('Image not found', 'error')
        return redirect(url_for('role_images_gallery', role_slug=role_slug))
    
    # Get user's vote
    user_vote = None
    if current_user:
        vote = RoleImageVote.query.filter_by(
            image_id=image_id,
            user_id=current_user['id']
        ).first()
        user_vote = vote.value if vote else None
    
    content = f"""
    <div class="container mt-4">
        <div class="row mb-3">
            <div class="col-md-12">
                <a href="/roles/{role_slug}/images/" class="btn btn-outline-secondary btn-sm">
                    <i class="fas fa-arrow-left me-1"></i> Back to Gallery
                </a>
            </div>
        </div>
        
        <div class="row">
            <div class="col-md-8">
                <div class="card">
                    <div class="card-body">
                        <h2 class="card-title mb-3">
                            Role Image for {role_slug}
                            {f'<span class="badge bg-success ms-2">Primary</span>' if image.is_primary else ''}
                            {f'<span class="badge bg-warning ms-2">Hidden</span>' if image.is_hidden and is_admin else ''}
                        </h2>
                        
                        <div class="text-center mb-4">
                            {'<iframe src="' + image.image_url + '" style="width: 100%; height: 500px; border: none;"></iframe>' if image.source_type == 'ordinal' and image.content_type and 'html' in image.content_type.lower() else '<img src="' + image.image_url + '" class="img-fluid" alt="Role image" style="max-height: 500px;">'}
                        </div>
                        
                        <div class="d-flex justify-content-between align-items-center mb-4">
                            <div class="btn-group" role="group">
                                <button class="btn {'btn-success' if user_vote == 1 else 'btn-outline-success'}" onclick="vote(1)" {'disabled' if not current_user else ''}>
                                    <i class="fas fa-thumbs-up"></i> Upvote ({image.upvotes})
                                </button>
                                <button class="btn {'btn-danger' if user_vote == -1 else 'btn-outline-danger'}" onclick="vote(-1)" {'disabled' if not current_user else ''}>
                                    <i class="fas fa-thumbs-down"></i> Downvote ({image.downvotes})
                                </button>
                                {f'<button class="btn btn-outline-secondary" onclick="removeVote()">Remove Vote</button>' if user_vote and current_user else ''}
                            </div>
                            <div>
                                <h4 class="mb-0">
                                    <span class="badge bg-primary">Net Score: {image.net_score}</span>
                                </h4>
                            </div>
                        </div>
                        
                        {'<div class="alert alert-info"><i class="fas fa-info-circle me-2"></i>Please <a href="/login/">login</a> to vote on images.</div>' if not current_user else ''}
                    </div>
                </div>
            </div>
            
            <div class="col-md-4">
                <div class="card mb-3">
                    <div class="card-header">
                        <h5 class="mb-0">Image Details</h5>
                    </div>
                    <div class="card-body">
                        <dl class="row mb-0">
                            <dt class="col-sm-5">Submitted by:</dt>
                            <dd class="col-sm-7">{image.submitted_by.displayName or image.submitted_by.username if image.submitted_by else 'Unknown'}</dd>
                            
                            <dt class="col-sm-5">Submitted:</dt>
                            <dd class="col-sm-7">{image.submitted_at.strftime('%Y-%m-%d %H:%M') if image.submitted_at else 'Unknown'}</dd>
                            
                            <dt class="col-sm-5">Source:</dt>
                            <dd class="col-sm-7"><span class="badge bg-info">{image.source_type}</span></dd>
                            
                            {f'<dt class="col-sm-5">Chain:</dt><dd class="col-sm-7">{image.chain}</dd>' if image.chain else ''}
                            
                            {f'<dt class="col-sm-5">Inscription ID:</dt><dd class="col-sm-7"><a href="https://ordinals.com/inscription/{image.inscription_id}" target="_blank" class="text-break small">{image.inscription_id[:20]}...</a></dd>' if image.inscription_id else ''}
                            
                            {f'<dt class="col-sm-5">Content Type:</dt><dd class="col-sm-7">{image.content_type}</dd>' if image.content_type else ''}
                            
                            {f'<dt class="col-sm-5">Promoted by:</dt><dd class="col-sm-7">{image.promoted_by.displayName or image.promoted_by.username if image.promoted_by else "N/A"}</dd>' if image.is_primary else ''}
                            
                            {f'<dt class="col-sm-5">Promoted at:</dt><dd class="col-sm-7">{image.promoted_at.strftime("%Y-%m-%d %H:%M") if image.promoted_at else "N/A"}</dd>' if image.is_primary else ''}
                        </dl>
                    </div>
                </div>
                
                {f'''<div class="card mb-3">
                    <div class="card-header bg-danger text-white">
                        <h5 class="mb-0">Admin Actions</h5>
                    </div>
                    <div class="card-body">
                        {f'<button class="btn btn-success w-100 mb-2" onclick="promoteImage()"><i class="fas fa-star me-2"></i>Promote to Primary</button>' if not image.is_primary else '<button class="btn btn-warning w-100 mb-2" onclick="demoteImage()"><i class="fas fa-star-half-alt me-2"></i>Demote from Primary</button>'}
                        
                        {f'<button class="btn btn-warning w-100 mb-2" onclick="hideImage()"><i class="fas fa-eye-slash me-2"></i>Hide Image</button>' if not image.is_hidden else '<button class="btn btn-info w-100 mb-2" onclick="unhideImage()"><i class="fas fa-eye me-2"></i>Unhide Image</button>'}
                        
                        <button class="btn btn-danger w-100 mb-3" onclick="deleteImage()">
                            <i class="fas fa-trash me-2"></i>Delete Image
                        </button>
                        
                        <hr>
                        
                        <div class="mb-2">
                            <label for="adminNote" class="form-label">Admin Note:</label>
                            <textarea class="form-control" id="adminNote" rows="3">{image.admin_note or ''}</textarea>
                        </div>
                        <button class="btn btn-primary w-100" onclick="saveNote()">
                            <i class="fas fa-save me-2"></i>Save Note
                        </button>
                    </div>
                </div>''' if is_admin else ''}
            </div>
        </div>
    </div>
    
    <script>
    const imageId = '{image_id}';
    const roleSlug = '{role_slug}';
    
    async function vote(value) {{
        try {{
            const response = await fetch(`/api/role-images/${{imageId}}/vote/`, {{
                method: 'POST',
                headers: {{'Content-Type': 'application/json'}},
                body: JSON.stringify({{value}})
            }});
            
            if (response.ok) {{
                location.reload();
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to vote');
            }}
        }} catch (error) {{
            console.error('Error voting:', error);
            alert('Error voting on image');
        }}
    }}
    
    async function removeVote() {{
        try {{
            const response = await fetch(`/api/role-images/${{imageId}}/vote/`, {{
                method: 'DELETE'
            }});
            
            if (response.ok) {{
                location.reload();
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to remove vote');
            }}
        }} catch (error) {{
            console.error('Error removing vote:', error);
            alert('Error removing vote');
        }}
    }}
    
    async function promoteImage() {{
        if (!confirm('Promote this image to primary role image?')) return;
        
        try {{
            const response = await fetch(`/api/role-images/${{imageId}}/promote/`, {{
                method: 'POST'
            }});
            
            if (response.ok) {{
                location.reload();
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to promote image');
            }}
        }} catch (error) {{
            console.error('Error promoting image:', error);
            alert('Error promoting image');
        }}
    }}
    
    async function demoteImage() {{
        if (!confirm('Demote this image from primary?')) return;
        
        try {{
            const response = await fetch(`/api/role-images/${{imageId}}/promote/`, {{
                method: 'POST',
                headers: {{'Content-Type': 'application/json'}},
                body: JSON.stringify({{demote: true}})
            }});
            
            if (response.ok) {{
                location.reload();
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to demote image');
            }}
        }} catch (error) {{
            console.error('Error demoting image:', error);
            alert('Error demoting image');
        }}
    }}
    
    async function hideImage() {{
        if (!confirm('Hide this image from public view?')) return;
        
        try {{
            const response = await fetch(`/api/role-images/${{imageId}}/hide/`, {{
                method: 'POST'
            }});
            
            if (response.ok) {{
                location.reload();
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to hide image');
            }}
        }} catch (error) {{
            console.error('Error hiding image:', error);
            alert('Error hiding image');
        }}
    }}
    
    async function unhideImage() {{
        if (!confirm('Unhide this image?')) return;
        
        try {{
            const response = await fetch(`/api/role-images/${{imageId}}/unhide/`, {{
                method: 'POST'
            }});
            
            if (response.ok) {{
                location.reload();
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to unhide image');
            }}
        }} catch (error) {{
            console.error('Error unhiding image:', error);
            alert('Error unhiding image');
        }}
    }}
    
    async function deleteImage() {{
        if (!confirm('Permanently delete this image? This cannot be undone.')) return;
        
        try {{
            const response = await fetch(`/api/role-images/${{imageId}}/`, {{
                method: 'DELETE'
            }});
            
            if (response.ok) {{
                window.location.href = `/roles/${{roleSlug}}/images/`;
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to delete image');
            }}
        }} catch (error) {{
            console.error('Error deleting image:', error);
            alert('Error deleting image');
        }}
    }}
    
    async function saveNote() {{
        const note = document.getElementById('adminNote').value;
        
        try {{
            const response = await fetch(`/api/role-images/${{imageId}}/note/`, {{
                method: 'PATCH',
                headers: {{'Content-Type': 'application/json'}},
                body: JSON.stringify({{admin_note: note}})
            }});
            
            if (response.ok) {{
                alert('Note saved successfully');
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to save note');
            }}
        }} catch (error) {{
            console.error('Error saving note:', error);
            alert('Error saving note');
        }}
    }}
    </script>
    """
    
    return render_page(f"Image Detail: {role_slug} - MLGH", content, theme=current_theme, user_menu=user_menu)

# ============================================================================
# Projects, Workgroups, and Guilds UI Pages
# ============================================================================

@app.route('/projects/')
def projects_directory():
    """Projects directory page"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = f"""
    <div class="container mt-4">
        <div class="row mb-4">
            <div class="col-md-8">
                <h1>Layers Directory</h1>
                <p class="lead">Browse and discover MLTF layers</p>
            </div>
            <div class="col-md-4 text-end">
                {'<a href="/projects/create/" class="btn btn-primary"><i class="fas fa-plus me-2"></i>Create Layer</a>' if current_user else '<a href="/login/" class="btn btn-primary"><i class="fas fa-sign-in-alt me-2"></i>Login to Create</a>'}
            </div>
        </div>
        
        <div class="row mb-4">
            <div class="col-md-4">
                <label for="status-filter" class="form-label">Status:</label>
                <select id="status-filter" class="form-select" onchange="loadProjects()">
                    <option value="">All Statuses</option>
                    <option value="proposed">Proposed</option>
                    <option value="active">Active</option>
                    <option value="stabilizing">Stabilizing</option>
                    <option value="maintaining">Maintaining</option>
                    <option value="dormant">Dormant</option>
                    <option value="concluded">Concluded</option>
                    <option value="archived">Archived</option>
                </select>
            </div>
            <div class="col-md-4">
                <label for="approval-filter" class="form-label">Approval:</label>
                <select id="approval-filter" class="form-select" onchange="loadProjects()">
                    <option value="">All</option>
                    <option value="pending">Pending</option>
                    <option value="approved">Approved</option>
                    <option value="rejected">Rejected</option>
                </select>
            </div>
            <div class="col-md-4">
                <label for="search-input" class="form-label">Search:</label>
                <input type="text" id="search-input" class="form-control" placeholder="Search layers..." onkeyup="filterProjects()">
            </div>
        </div>
        
        <div id="projects-container" class="row">
            <div class="col-12 text-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    let allProjects = [];
    
    async function loadProjects() {{
        const statusFilter = document.getElementById('status-filter').value;
        const approvalFilter = document.getElementById('approval-filter').value;
        
        let url = '/api/projects/';
        const params = new URLSearchParams();
        if (statusFilter) params.append('status', statusFilter);
        if (approvalFilter) params.append('approval_status', approvalFilter);
        if (params.toString()) url += '?' + params.toString();
        
        try {{
            const response = await fetch(url);
            const data = await response.json();
            allProjects = data.projects;
            displayProjects(allProjects);
        }} catch (error) {{
            console.error('Error loading projects:', error);
            document.getElementById('projects-container').innerHTML = '<div class="col-12"><div class="alert alert-danger">Error loading projects</div></div>';
        }}
    }}
    
    function filterProjects() {{
        const searchTerm = document.getElementById('search-input').value.toLowerCase();
        const filtered = allProjects.filter(p => 
            p.name.toLowerCase().includes(searchTerm) ||
            (p.description && p.description.toLowerCase().includes(searchTerm))
        );
        displayProjects(filtered);
    }}
    
    function displayProjects(projects) {{
        const container = document.getElementById('projects-container');
        
        if (projects.length === 0) {{
            container.innerHTML = '<div class="col-12"><div class="alert alert-info">No projects found</div></div>';
            return;
        }}
        
        let html = '';
        projects.forEach(project => {{
            const statusBadge = (project.approval_status === 'approved' && project.status === 'proposed') ? '' : getStatusBadge(project.status);
            const approvalBadge = getApprovalBadge(project.approval_status);
            
            const projectImgHtml = project.image_url ? `<div class="card-img-top overflow-hidden" style="height: 140px; background: var(--bg-secondary, #f8f9fa);"><img src="${{project.image_url}}" alt="${{project.name}}" class="w-100 h-100 object-fit-cover"></div>` : '';
            html += `
                <div class="col-md-6 col-lg-4 mb-4">
                    <div class="card h-100">
                        ${{projectImgHtml}}
                        <div class="card-body">
                            <h5 class="card-title">
                                <a href="/projects/${{project.slug}}/">${{project.name}}</a>
                            </h5>
                            <div class="mb-2">
                                ${{statusBadge}}
                                ${{approvalBadge}}
                            </div>
                            ${{project.mission ? '<p class="card-text fw-medium">' + (project.mission || '').replace(/&/g, '&amp;').replace(/</g, '&lt;').replace(/>/g, '&gt;').replace(/\\n/g, '<br>') + '</p>' : ''}}
                            <p class="card-text text-muted">${{(project.description || 'No description').replace(/&/g, '&amp;').replace(/</g, '&lt;').replace(/>/g, '&gt;').replace(/\\n/g, '<br>')}}</p>
                            <div class="mt-3">
                                <small class="text-muted">
                                    <i class="fas fa-users me-1"></i> ${{project.workgroups_count || 0}} workgroups
                                </small>
                            </div>
                        </div>
                        <div class="card-footer">
                            <small class="text-muted">Created ${{new Date(project.created_at).toLocaleDateString()}}</small>
                        </div>
                    </div>
                </div>
            `;
        }});
        
        container.innerHTML = html;
    }}
    
    function getStatusBadge(status) {{
        const badges = {{
            'proposed': '<span class="badge bg-info">Proposed</span>',
            'active': '<span class="badge bg-success">Active</span>',
            'stabilizing': '<span class="badge bg-primary">Stabilizing</span>',
            'maintaining': '<span class="badge bg-secondary">Maintaining</span>',
            'dormant': '<span class="badge bg-warning">Dormant</span>',
            'concluded': '<span class="badge bg-dark">Concluded</span>',
            'archived': '<span class="badge bg-secondary">Archived</span>'
        }};
        return badges[status] || '';
    }}
    
    function getApprovalBadge(approval) {{
        const badges = {{
            'pending': '<span class="badge bg-warning">Pending Approval</span>',
            'approved': '<span class="badge bg-success">Approved</span>',
            'rejected': '<span class="badge bg-danger">Rejected</span>'
        }};
        return badges[approval] || '';
    }}
    
    // Load projects on page load
    loadProjects();
    </script>
    """
    
    return render_page("Layers Directory - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/workgroups/')
def workgroups_directory():
    """Workgroups directory page"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = f"""
    <div class="container mt-4">
        <div class="row mb-4">
            <div class="col-md-8">
                <h1>Workgroups Directory</h1>
                <p class="lead">Browse workgroups across all projects</p>
            </div>
            <div class="col-md-4 text-end">
                <a href="/projects/" class="btn btn-secondary mb-2 w-100"><i class="fas fa-arrow-left me-2"></i>Back to Layers</a>
                {'<button class="btn btn-primary w-100" onclick="showCreateWorkgroupModal()"><i class="fas fa-plus me-2"></i>Create Workgroup</button>' if current_user else ''}
            </div>
        </div>
        
        <div class="row mb-4">
            <div class="col-md-4">
                <label for="project-filter" class="form-label">Layer:</label>
                <select id="project-filter" class="form-select" onchange="loadWorkgroups()">
                    <option value="">All Layers</option>
                </select>
            </div>
            <div class="col-md-4">
                <label for="status-filter" class="form-label">Status:</label>
                <select id="status-filter" class="form-select" onchange="loadWorkgroups()">
                    <option value="">All Statuses</option>
                    <option value="active">Active</option>
                    <option value="inactive">Inactive</option>
                    <option value="completed">Completed</option>
                    <option value="archived">Archived</option>
                </select>
            </div>
            <div class="col-md-4">
                <label for="search-input" class="form-label">Search:</label>
                <input type="text" id="search-input" class="form-control" placeholder="Search workgroups..." onkeyup="filterWorkgroups()">
            </div>
        </div>
        
        <div id="workgroups-container" class="row">
            <div class="col-12 text-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    let allWorkgroups = [];
    let allProjects = [];
    
    async function loadProjects() {{
        try {{
            const response = await fetch('/api/projects/?approval_status=approved');
            const data = await response.json();
            allProjects = data.projects;
            
            const select = document.getElementById('project-filter');
            allProjects.forEach(project => {{
                const option = document.createElement('option');
                option.value = project.id;
                option.textContent = project.name;
                select.appendChild(option);
            }});
        }} catch (error) {{
            console.error('Error loading projects:', error);
        }}
    }}
    
    async function loadWorkgroups() {{
        const projectFilter = document.getElementById('project-filter').value;
        const statusFilter = document.getElementById('status-filter').value;
        
        try {{
            allWorkgroups = [];
            
            if (projectFilter) {{
                // Load workgroups for specific project
                let url = `/api/projects/${{projectFilter}}/workgroups/`;
                if (statusFilter) url += `?status=${{statusFilter}}`;
                
                const response = await fetch(url);
                const data = await response.json();
                allWorkgroups = data.workgroups;
            }} else {{
                // Load workgroups from all projects
                for (const project of allProjects) {{
                    let url = `/api/projects/${{project.id}}/workgroups/`;
                    if (statusFilter) url += `?status=${{statusFilter}}`;
                    
                    const response = await fetch(url);
                    const data = await response.json();
                    allWorkgroups = allWorkgroups.concat(data.workgroups);
                }}
            }}
            
            displayWorkgroups(allWorkgroups);
        }} catch (error) {{
            console.error('Error loading workgroups:', error);
            document.getElementById('workgroups-container').innerHTML = '<div class="col-12"><div class="alert alert-danger">Error loading workgroups</div></div>';
        }}
    }}
    
    function filterWorkgroups() {{
        const searchTerm = document.getElementById('search-input').value.toLowerCase();
        const filtered = allWorkgroups.filter(wg => 
            wg.name.toLowerCase().includes(searchTerm) ||
            (wg.description && wg.description.toLowerCase().includes(searchTerm))
        );
        displayWorkgroups(filtered);
    }}
    
    function displayWorkgroups(workgroups) {{
        const container = document.getElementById('workgroups-container');
        
        if (workgroups.length === 0) {{
            container.innerHTML = '<div class="col-12"><div class="alert alert-info">No workgroups found</div></div>';
            return;
        }}
        
        let html = '';
        workgroups.forEach(wg => {{
            const statusBadge = getStatusBadge(wg.status);
            const approvalBadge = getApprovalBadge(wg.approval_status);
            const project = allProjects.find(p => p.id === wg.project_id);
            
            const wgImgHtml = wg.image_url ? `<div class="card-img-top overflow-hidden" style="height: 140px; background: var(--bg-secondary, #f8f9fa);"><img src="${{wg.image_url}}" alt="${{wg.name}}" class="w-100 h-100 object-fit-cover"></div>` : '';
            html += `
                <div class="col-md-6 col-lg-4 mb-4">
                    <div class="card h-100">
                        ${{wgImgHtml}}
                        <div class="card-body">
                            <h5 class="card-title">
                                <a href="/workgroups/${{wg.slug}}/">${{wg.name}}</a>
                            </h5>
                            <div class="mb-2">
                                ${{statusBadge}}
                                ${{approvalBadge}}
                            </div>
                            <p class="card-text text-muted">${{wg.description || 'No description'}}</p>
                            ${{project ? `<div class="mt-2"><small class="text-muted"><i class="fas fa-project-diagram me-1"></i> ${{project.name}}</small></div>` : ''}}
                        </div>
                        <div class="card-footer">
                            <small class="text-muted">Created ${{new Date(wg.created_at).toLocaleDateString()}}</small>
                        </div>
                    </div>
                </div>
            `;
        }});
        
        container.innerHTML = html;
    }}
    
    function getStatusBadge(status) {{
        const badges = {{
            'active': '<span class="badge bg-success">Active</span>',
            'inactive': '<span class="badge bg-warning">Inactive</span>',
            'completed': '<span class="badge bg-primary">Completed</span>',
            'archived': '<span class="badge bg-secondary">Archived</span>'
        }};
        return badges[status] || '';
    }}
    
    function getApprovalBadge(approval) {{
        const badges = {{
            'pending': '<span class="badge bg-warning">Pending Approval</span>',
            'approved': '<span class="badge bg-success">Approved</span>',
            'rejected': '<span class="badge bg-danger">Rejected</span>'
        }};
        return badges[approval] || '';
    }}
    
    // Load data on page load
    function showCreateWorkgroupModal() {{
        const modalHtml = `
            <div class="modal fade" id="createWorkgroupModal" tabindex="-1">
                <div class="modal-dialog modal-lg">
                    <div class="modal-content">
                        <div class="modal-header">
                            <h5 class="modal-title">Create New Workgroup</h5>
                            <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                        </div>
                        <div class="modal-body">
                            <div id="wg-alert-container"></div>
                            
                            <form id="createWorkgroupForm">
                                <div class="mb-3">
                                    <label for="wg-project" class="form-label">Layer *</label>
                                    <select class="form-select" id="wg-project" required>
                                        <option value="">Select a project...</option>
                                    </select>
                                    <div class="form-text">Select the project this workgroup belongs to</div>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="wg-name" class="form-label">Workgroup Name *</label>
                                    <input type="text" class="form-control" id="wg-name" required>
                                    <div class="form-text">A clear, descriptive name for the workgroup</div>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="wg-description" class="form-label">Description *</label>
                                    <textarea class="form-control" id="wg-description" rows="4" required></textarea>
                                    <div class="form-text">Describe the workgroup's purpose and goals</div>
                                </div>
                                
                                <div class="alert alert-info">
                                    <i class="fas fa-info-circle me-2"></i>
                                    <strong>Note:</strong> New workgroups require approval from the layer admin before becoming active.
                                </div>
                            </form>
                        </div>
                        <div class="modal-footer">
                            <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                            <button type="button" class="btn btn-primary" id="submitWorkgroupBtn">
                                <i class="fas fa-plus me-2"></i>Create Workgroup
                            </button>
                        </div>
                    </div>
                </div>
            </div>
        `;
        
        if (!document.getElementById('createWorkgroupModal')) {{
            document.body.insertAdjacentHTML('beforeend', modalHtml);
        }}
        
        // Populate project dropdown
        const select = document.getElementById('wg-project');
        select.innerHTML = '<option value="">Select a project...</option>';
        allProjects.forEach(project => {{
            const option = document.createElement('option');
            option.value = project.id;
            option.textContent = project.name;
            select.appendChild(option);
        }});
        
        const modal = new bootstrap.Modal(document.getElementById('createWorkgroupModal'));
        modal.show();
        
        document.getElementById('submitWorkgroupBtn').onclick = async () => {{
            const projectId = document.getElementById('wg-project').value;
            const name = document.getElementById('wg-name').value.trim();
            const description = document.getElementById('wg-description').value.trim();
            
            if (!projectId) {{
                document.getElementById('wg-alert-container').innerHTML = `
                    <div class="alert alert-danger">
                        <i class="fas fa-exclamation-circle me-2"></i>
                        Please select a project
                    </div>
                `;
                return;
            }}
            
            if (!name || !description) {{
                document.getElementById('wg-alert-container').innerHTML = `
                    <div class="alert alert-danger">
                        <i class="fas fa-exclamation-circle me-2"></i>
                        Name and description are required
                    </div>
                `;
                return;
            }}
            
            const submitBtn = document.getElementById('submitWorkgroupBtn');
            submitBtn.disabled = true;
            submitBtn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Creating...';
            
            try {{
                const response = await fetch(`/api/projects/${{projectId}}/workgroups/`, {{
                    method: 'POST',
                    headers: {{'Content-Type': 'application/json'}},
                    body: JSON.stringify({{ name, description }})
                }});
                
                const data = await response.json();
                
                if (response.ok) {{
                    modal.hide();
                    loadWorkgroups();
                    alert('Workgroup created successfully! It will be visible once approved by the layer admin.');
                }} else {{
                    throw new Error(data.error || 'Failed to create workgroup');
                }}
            }} catch (error) {{
                document.getElementById('wg-alert-container').innerHTML = `
                    <div class="alert alert-danger">
                        <i class="fas fa-exclamation-circle me-2"></i>
                        ${{error.message}}
                    </div>
                `;
                submitBtn.disabled = false;
                submitBtn.innerHTML = '<i class="fas fa-plus me-2"></i>Create Workgroup';
            }}
        }};
    }}
    
    loadProjects().then(() => loadWorkgroups());
    </script>
    """
    
    return render_page("Workgroups Directory - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/waitlists/')
def waitlists_directory():
    """Waitlists directory page"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = f"""
    <div class="container mt-4">
        <div class="row mb-4">
            <div class="col-md-8">
                <h1>Waitlists Directory</h1>
                <p class="lead">Join waitlists for upcoming projects, features, and opportunities</p>
            </div>
        </div>
        
        <div class="row mb-4">
            <div class="col-md-4">
                <label for="project-filter" class="form-label">Layer:</label>
                <select id="project-filter" class="form-select" onchange="loadWaitlists()">
                    <option value="">All Layers</option>
                </select>
            </div>
            <div class="col-md-4">
                <label for="status-filter" class="form-label">Status:</label>
                <select id="status-filter" class="form-select" onchange="loadWaitlists()">
                    <option value="">All</option>
                    <option value="active">Active</option>
                    <option value="upcoming">Upcoming</option>
                    <option value="closed">Closed</option>
                </select>
            </div>
            <div class="col-md-4">
                <label for="search-input" class="form-label">Search:</label>
                <input type="text" id="search-input" class="form-control" placeholder="Search waitlists..." onkeyup="filterWaitlists()">
            </div>
        </div>
        
        <div id="waitlists-container" class="row">
            <div class="col-12 text-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    let allWaitlists = [];
    let allProjects = [];
    
    async function loadProjects() {{
        try {{
            const response = await fetch('/api/projects/?approval_status=approved');
            const data = await response.json();
            allProjects = data.projects;
            
            const select = document.getElementById('project-filter');
            allProjects.forEach(project => {{
                const option = document.createElement('option');
                option.value = project.id;
                option.textContent = project.name;
                select.appendChild(option);
            }});
        }} catch (error) {{
            console.error('Error loading projects:', error);
        }}
    }}
    
    async function loadWaitlists() {{
        const projectFilter = document.getElementById('project-filter').value;
        const statusFilter = document.getElementById('status-filter').value;
        
        try {{
            allWaitlists = [];
            
            if (projectFilter) {{
                // Load waitlists for specific project
                const response = await fetch(`/api/projects/${{projectFilter}}/waitlists/`);
                const data = await response.json();
                allWaitlists = data.waitlists || [];
            }} else {{
                // Load waitlists from all projects
                for (const project of allProjects) {{
                    const response = await fetch(`/api/projects/${{project.id}}/waitlists/`);
                    const data = await response.json();
                    if (data.waitlists) {{
                        allWaitlists = allWaitlists.concat(data.waitlists);
                    }}
                }}
            }}
            
            // Filter by status
            if (statusFilter) {{
                const now = new Date();
                allWaitlists = allWaitlists.filter(wl => {{
                    const startDate = new Date(wl.start_date);
                    const closingDate = wl.closing_date ? new Date(wl.closing_date) : null;
                    const isFull = wl.max_number && wl.entry_count >= wl.max_number;
                    
                    if (statusFilter === 'active') {{
                        return wl.active && !wl.archived && now >= startDate && (!closingDate || now <= closingDate) && !isFull;
                    }} else if (statusFilter === 'upcoming') {{
                        return wl.active && !wl.archived && now < startDate;
                    }} else if (statusFilter === 'closed') {{
                        return wl.archived || !wl.active || (closingDate && now > closingDate) || isFull;
                    }}
                    return true;
                }});
            }}
            
            displayWaitlists(allWaitlists);
        }} catch (error) {{
            console.error('Error loading waitlists:', error);
            document.getElementById('waitlists-container').innerHTML = '<div class="col-12"><div class="alert alert-danger">Error loading waitlists</div></div>';
        }}
    }}
    
    function filterWaitlists() {{
        const searchTerm = document.getElementById('search-input').value.toLowerCase();
        const filtered = allWaitlists.filter(wl => 
            wl.name.toLowerCase().includes(searchTerm) ||
            (wl.description && wl.description.toLowerCase().includes(searchTerm))
        );
        displayWaitlists(filtered);
    }}
    
    function displayWaitlists(waitlists) {{
        const container = document.getElementById('waitlists-container');
        
        if (waitlists.length === 0) {{
            container.innerHTML = '<div class="col-12"><div class="alert alert-info">No waitlists found</div></div>';
            return;
        }}
        
        let html = '';
        waitlists.forEach(wl => {{
            const project = allProjects.find(p => p.id === wl.project_id);
            const now = new Date();
            const startDate = new Date(wl.start_date);
            const closingDate = wl.closing_date ? new Date(wl.closing_date) : null;
            const isFull = wl.max_number && wl.count >= wl.max_number;
            
            let statusBadge = '';
            let statusText = '';
            
            if (!wl.active || wl.archived) {{
                statusBadge = '<span class="badge bg-secondary">Closed</span>';
                statusText = 'This waitlist is closed';
            }} else if (now < startDate) {{
                statusBadge = '<span class="badge bg-info">Upcoming</span>';
                statusText = `Opens ${{startDate.toLocaleDateString()}}`;
            }} else if (isFull) {{
                statusBadge = '<span class="badge bg-warning">Full</span>';
                statusText = `${{wl.count}} / ${{wl.max_number}} spots filled`;
            }} else if (closingDate && now > closingDate) {{
                statusBadge = '<span class="badge bg-secondary">Closed</span>';
                statusText = `Closed ${{closingDate.toLocaleDateString()}}`;
            }} else {{
                statusBadge = '<span class="badge bg-success">Active</span>';
                if (wl.max_number) {{
                    statusText = `${{wl.count}} / ${{wl.max_number}} spots filled`;
                }} else {{
                    statusText = `${{wl.count}} member${{wl.count !== 1 ? 's' : ''}}`;
                }}
            }}
            
            const imgHtml = wl.image_url ? `<div class="card-img-top overflow-hidden" style="height: 140px; background: var(--bg-secondary, #f8f9fa);"><img src="${{wl.image_url}}" alt="${{wl.name}}" class="w-100 h-100 object-fit-cover"></div>` : '';
            html += `
                <div class="col-md-6 col-lg-4 mb-4">
                    <div class="card h-100">
                        ${{imgHtml}}
                        <div class="card-body">
                            <h5 class="card-title">
                                <a href="/projects/${{project ? project.slug : wl.project_id}}/waitlist/${{wl.id}}/">${{wl.name}}</a>
                            </h5>
                            <div class="mb-2">
                                ${{statusBadge}}
                                ${{wl.referrals ? '<span class="badge bg-primary ms-1"><i class="fas fa-users"></i> Referrals</span>' : ''}}
                                ${{wl.milestones ? '<span class="badge bg-info ms-1"><i class="fas fa-flag"></i> Milestones</span>' : ''}}
                            </div>
                            <p class="card-text text-muted small mb-2">
                                <i class="fas fa-project-diagram me-1"></i>
                                <a href="/projects/${{project ? project.slug : wl.project_id}}/">${{project ? project.name : 'Unknown Layer'}}</a>
                            </p>
                            <p class="card-text">${{wl.description || 'No description'}}</p>
                            <div class="mt-3">
                                <small class="text-muted">${{statusText}}</small>
                            </div>
                        </div>
                        <div class="card-footer">
                            <small class="text-muted">Created ${{new Date(wl.created_at).toLocaleDateString()}}</small>
                        </div>
                    </div>
                </div>
            `;
        }});
        
        container.innerHTML = html;
    }}
    
    // Load data on page load
    loadProjects().then(() => loadWaitlists());
    </script>
    """
    
    return render_page("Waitlists Directory - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/guilds/')
def guilds_directory():
    """Guilds directory page"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = f"""
    <div class="container mt-4">
        <div class="row mb-4">
            <div class="col-md-8">
                <h1>Guilds Directory</h1>
                <p class="lead">Cross-project collaboration groups</p>
            </div>
            <div class="col-md-4 text-end">
                {'<a href="/guilds/create/" class="btn btn-primary"><i class="fas fa-plus me-2"></i>Create Guild</a>' if current_user else '<a href="/login/" class="btn btn-primary"><i class="fas fa-sign-in-alt me-2"></i>Login to Create</a>'}
            </div>
        </div>
        
        <div class="row mb-4">
            <div class="col-md-6">
                <label for="status-filter" class="form-label">Status:</label>
                <select id="status-filter" class="form-select" onchange="loadGuilds()">
                    <option value="">All Statuses</option>
                    <option value="active">Active</option>
                    <option value="archived">Archived</option>
                </select>
            </div>
            <div class="col-md-6">
                <label for="search-input" class="form-label">Search:</label>
                <input type="text" id="search-input" class="form-control" placeholder="Search guilds..." onkeyup="filterGuilds()">
            </div>
        </div>
        
        <div id="guilds-container" class="row">
            <div class="col-12 text-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    let allGuilds = [];
    
    async function loadGuilds() {{
        const statusFilter = document.getElementById('status-filter').value;
        
        let url = '/api/guilds/';
        if (statusFilter) url += `?status=${{statusFilter}}`;
        
        try {{
            const response = await fetch(url);
            const data = await response.json();
            allGuilds = data.guilds;
            displayGuilds(allGuilds);
        }} catch (error) {{
            console.error('Error loading guilds:', error);
            document.getElementById('guilds-container').innerHTML = '<div class="col-12"><div class="alert alert-danger">Error loading guilds</div></div>';
        }}
    }}
    
    function filterGuilds() {{
        const searchTerm = document.getElementById('search-input').value.toLowerCase();
        const filtered = allGuilds.filter(g => 
            g.name.toLowerCase().includes(searchTerm) ||
            (g.description && g.description.toLowerCase().includes(searchTerm))
        );
        displayGuilds(filtered);
    }}
    
    function displayGuilds(guilds) {{
        const container = document.getElementById('guilds-container');
        
        if (guilds.length === 0) {{
            container.innerHTML = '<div class="col-12"><div class="alert alert-info">No guilds found</div></div>';
            return;
        }}
        
        let html = '';
        guilds.forEach(guild => {{
            const statusBadge = guild.status === 'active' 
                ? '<span class="badge bg-success">Active</span>' 
                : '<span class="badge bg-secondary">Archived</span>';
            
            const guildImgHtml = guild.image_url ? `<div class="card-img-top overflow-hidden" style="height: 140px; background: var(--bg-secondary, #f8f9fa);"><img src="${{"guild.image_url"}}" alt="${{"guild.name"}}" class="w-100 h-100 object-fit-cover"></div>` : '';
            html += `
                <div class="col-md-6 col-lg-4 mb-4">
                    <div class="card h-100">
                        ${{"guildImgHtml"}}
                        <div class="card-body">
                            <h5 class="card-title">
                                <a href="/guilds/${{guild.slug}}/">${{guild.name}}</a>
                            </h5>
                            <div class="mb-2">
                                ${{statusBadge}}
                            </div>
                            <p class="card-text text-muted">${{guild.description || 'No description'}}</p>
                            <div class="mt-3">
                                <small class="text-muted">
                                    <i class="fas fa-users me-1"></i> ${{guild.members_count || 0}} members
                                </small>
                            </div>
                        </div>
                        <div class="card-footer">
                            <small class="text-muted">Created ${{new Date(guild.created_at).toLocaleDateString()}}</small>
                        </div>
                    </div>
                </div>
            `;
        }});
        
        container.innerHTML = html;
    }}
    
    // Load guilds on page load
    loadGuilds();
    </script>
    """
    
    return render_page("Guilds Directory - MLGH", content, theme=current_theme, user_menu=user_menu)

def _render_project_detail(project_slug, waitlist_id=None):
    """Shared logic for project detail page. waitlist_id when from /projects/<slug>/waitlist/<id>/"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    project_obj = Project.query.filter_by(slug=project_slug).first()
    show_admin_tab = bool(project_obj and current_user and is_project_admin(project_obj, current_user))
    initial_waitlist_id = int(waitlist_id) if waitlist_id else None
    
    admin_tab_html = ''
    admin_tab_pane_html = ''
    admin_tab_listener = ''
    if show_admin_tab:
        admin_tab_html = '''
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="admin-tab" data-bs-toggle="tab" data-bs-target="#admin" type="button">Admin</button>
            </li>
        '''
        admin_tab_pane_html = '''
            <div class="tab-pane fade" id="admin">
                <div id="admin-content"></div>
            </div>
        '''
        admin_tab_listener = "document.getElementById('admin-tab').addEventListener('shown.bs.tab', loadAdmins);"
    
    content = f"""
    <div class="container mt-4">
        <div id="project-header" class="mb-4">
            <div class="d-flex justify-content-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        </div>
        
        <ul class="nav nav-tabs mb-4" id="projectTabs" role="tablist">
            <li class="nav-item" role="presentation">
                <button class="nav-link active" id="overview-tab" data-bs-toggle="tab" data-bs-target="#overview" type="button">Overview</button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="workgroups-tab" data-bs-toggle="tab" data-bs-target="#workgroups" type="button">Workgroups</button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="clusters-tab" data-bs-toggle="tab" data-bs-target="#clusters" type="button">Clusters</button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="roles-tab" data-bs-toggle="tab" data-bs-target="#roles" type="button">Roles</button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="claims-tab" data-bs-toggle="tab" data-bs-target="#claims" type="button">Claims</button>
            </li>
            <li class="nav-item" role="presentation">
                <button class="nav-link" id="votes-tab" data-bs-toggle="tab" data-bs-target="#votes" type="button">Votes</button>
            </li>
            {admin_tab_html}
            <li id="waitlist-tabs-marker" class="nav-item d-none"></li>
        </ul>
        
        <div class="tab-content" id="projectTabContent">
            <div class="tab-pane fade show active" id="overview">
                <div id="overview-content"></div>
            </div>
            <div class="tab-pane fade" id="workgroups">
                <div id="workgroups-content"></div>
            </div>
            <div class="tab-pane fade" id="clusters">
                <div id="clusters-content"></div>
            </div>
            <div class="tab-pane fade" id="roles">
                <div id="roles-content"></div>
            </div>
            <div class="tab-pane fade" id="claims">
                <div id="claims-content"></div>
            </div>
            <div class="tab-pane fade" id="votes">
                <div id="votes-content"></div>
            </div>
            {admin_tab_pane_html}
            <div id="waitlist-panes-marker" class="d-none"></div>
        </div>
    </div>
    
    <div class="modal fade" id="joinProjectModal" tabindex="-1">
        <div class="modal-dialog">
            <div class="modal-content">
                <div class="modal-header">
                    <h5 class="modal-title"><i class="fas fa-plus me-2"></i>Join Layer</h5>
                    <button type="button" class="btn-close" data-bs-dismiss="modal" aria-label="Close"></button>
                </div>
                <div class="modal-body">
                    <p class="text-muted mb-3">You will join this project as a contributor. Optionally add a referral code if you were invited via a link.</p>
                    <div class="mb-0">
                        <label for="join-project-referral" class="form-label">Referral code (optional)</label>
                        <input type="text" class="form-control" id="join-project-referral" placeholder="Leave blank if none">
                    </div>
                </div>
                <div class="modal-footer">
                    <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                    <button type="button" class="btn btn-primary" id="join-project-confirm-btn" onclick="submitJoinProjectModal()"><i class="fas fa-check me-2"></i>Join</button>
                </div>
            </div>
        </div>
    </div>
    
    <div class="modal fade" id="joinWaitlistModal" tabindex="-1">
        <div class="modal-dialog">
            <div class="modal-content">
                <div class="modal-header">
                    <h5 class="modal-title" id="join-waitlist-modal-title"><i class="fas fa-list-alt me-2"></i>Join Waitlist</h5>
                    <button type="button" class="btn-close" data-bs-dismiss="modal" aria-label="Close"></button>
                </div>
                <div class="modal-body">
                    <p class="text-muted mb-3" id="join-waitlist-modal-desc">Add an optional message for the waitlist owner.</p>
                    <div class="mb-0">
                        <label for="join-waitlist-message" class="form-label">Message (optional)</label>
                        <textarea class="form-control" id="join-waitlist-message" rows="3" placeholder="Leave blank to skip"></textarea>
                    </div>
                </div>
                <div class="modal-footer">
                    <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                    <button type="button" class="btn btn-primary" id="join-waitlist-confirm-btn" onclick="submitJoinWaitlistModal()"><i class="fas fa-check me-2"></i>Join</button>
                </div>
            </div>
        </div>
    </div>
    
    <div class="modal fade" id="addAdminModal" tabindex="-1">
        <div class="modal-dialog">
            <div class="modal-content">
                <div class="modal-header">
                    <h5 class="modal-title">Add layer admin</h5>
                    <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                </div>
                <div class="modal-body">
                    <label class="form-label">Search by username or name</label>
                    <input type="text" class="form-control" id="add-admin-username" placeholder="Type to search..." oninput="searchUsersForAdmin()">
                    <div id="add-admin-results" class="mt-3"></div>
                </div>
            </div>
        </div>
    </div>
    
    <div class="modal fade" id="embedCodeModal" tabindex="-1">
        <div class="modal-dialog modal-lg">
            <div class="modal-content">
                <div class="modal-header">
                    <h5 class="modal-title">Embed Waitlist Widget</h5>
                    <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                </div>
                <div class="modal-body">
                    <p class="lead" id="embed-waitlist-name"></p>
                    <p class="text-muted">Copy and paste this code into your website to embed the waitlist widget. Signups will be tracked with the source URL.</p>
                    
                    <div class="mb-4">
                        <label class="form-label"><strong>Embed Code (iframe)</strong></label>
                        <div class="input-group">
                            <textarea class="form-control font-monospace" id="embed-code-iframe" rows="3" readonly></textarea>
                            <button class="btn btn-outline-primary" onclick="copyEmbedCode('iframe')"><i class="fas fa-copy"></i> Copy</button>
                        </div>
                        <small class="text-muted">Recommended: Simple iframe embed with automatic sizing</small>
                    </div>
                    
                    <div class="mb-4">
                        <label class="form-label"><strong>Direct Widget URL</strong></label>
                        <div class="input-group">
                            <input type="text" class="form-control font-monospace" id="embed-url" readonly>
                            <button class="btn btn-outline-primary" onclick="copyEmbedCode('url')"><i class="fas fa-copy"></i> Copy</button>
                        </div>
                        <small class="text-muted">Use this URL to embed in an iframe or link directly</small>
                    </div>
                    
                    <div class="alert alert-info">
                        <i class="fas fa-info-circle me-2"></i>
                        <strong>Source Tracking:</strong> All signups from the embedded widget will be tracked with the source domain and URL where the signup occurred.
                    </div>
                </div>
            </div>
        </div>
    </div>
    
    <div class="modal fade" id="createVoteModal" tabindex="-1">
        <div class="modal-dialog modal-lg">
            <div class="modal-content">
                <div class="modal-header">
                    <h5 class="modal-title"><i class="fas fa-vote-yea me-2"></i>Create Vote</h5>
                    <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                </div>
                <div class="modal-body">
                    <div id="create-vote-alert" class="alert d-none" role="alert"></div>
                    <form id="createVoteForm">
                        <div class="mb-3">
                            <label for="vote-title" class="form-label">Title *</label>
                            <input type="text" class="form-control" id="vote-title" required placeholder="e.g. Approve ML-DRAFT-001">
                        </div>
                        <div class="mb-3">
                            <label for="vote-description" class="form-label">Description</label>
                            <textarea class="form-control" id="vote-description" rows="2" placeholder="What is being decided"></textarea>
                        </div>
                        <div class="mb-3">
                            <label for="vote-submission-id" class="form-label">Draft to vote on *</label>
                            <select class="form-select" id="vote-submission-id" required>
                                <option value="">Loading drafts...</option>
                            </select>
                            <div class="form-text">Select an approved draft from this layer's workgroups</div>
                        </div>
                        <div class="row">
                            <div class="col-md-6 mb-3">
                                <label for="vote-start" class="form-label">Start (<span id="timezone-start">your local time</span>) *</label>
                                <input type="datetime-local" class="form-control" id="vote-start" required>
                            </div>
                            <div class="col-md-6 mb-3">
                                <label for="vote-end" class="form-label">End (<span id="timezone-end">your local time</span>) *</label>
                                <input type="datetime-local" class="form-control" id="vote-end" required>
                            </div>
                        </div>
                        <div class="row">
                            <div class="col-md-6 mb-3">
                                <label for="vote-quorum" class="form-label">Quorum (min votes) *</label>
                                <input type="number" class="form-control" id="vote-quorum" required min="1" value="1">
                            </div>
                            <div class="col-md-6 mb-3">
                                <label for="vote-threshold" class="form-label">Win threshold (0–1) *</label>
                                <input type="number" class="form-control" id="vote-threshold" required min="0" max="1" step="0.01" value="0.5" placeholder="0.5 = majority">
                            </div>
                        </div>
                    </form>
                </div>
                <div class="modal-footer">
                    <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                    <button type="button" class="btn btn-primary" id="create-vote-submit-btn" onclick="submitCreateVote()"><i class="fas fa-check me-2"></i>Create Vote</button>
                </div>
            </div>
        </div>
    </div>
    
    <div class="modal fade" id="editProjectModal" tabindex="-1">
        <div class="modal-dialog modal-lg">
            <div class="modal-content">
                <div class="modal-header">
                    <h5 class="modal-title">Edit Layer</h5>
                    <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                </div>
                <div class="modal-body" style="max-height: 70vh; overflow-y: auto;">
                    <div id="edit-project-alert" class="alert d-none" role="alert"></div>
                    <form id="editProjectForm">
                        <div class="card bg-secondary bg-opacity-10 mb-4">
                            <div class="card-body py-3">
                                <h6 class="card-title mb-2"><i class="fas fa-user-shield me-2"></i>Layer Admins</h6>
                                <p class="text-muted small mb-2">Admins can manage workgroups, roles, claims, and other admins. The owner cannot be removed.</p>
                                <div id="edit-modal-admins-list" class="list-group mb-3"></div>
                                <div class="mb-0">
                                    <label class="form-label small">Add admin</label>
                                    <div class="input-group input-group-sm">
                                        <input type="text" class="form-control" id="edit-modal-add-admin-q" placeholder="Search by username (min 2 chars)..." oninput="searchUsersForEditModalAdmin()">
                                        <button type="button" class="btn btn-outline-primary" onclick="searchUsersForEditModalAdmin()"><i class="fas fa-search"></i></button>
                                    </div>
                                    <div id="edit-modal-add-admin-results" class="mt-2"></div>
                                </div>
                            </div>
                        </div>
                        <hr>
                        <div class="mb-3">
                            <label for="edit-project-name" class="form-label">Layer Name *</label>
                            <input type="text" class="form-control" id="edit-project-name" required maxlength="255">
                        </div>
                        <div class="mb-3">
                            <label for="edit-project-mission" class="form-label">Mission</label>
                            <textarea class="form-control" id="edit-project-mission" rows="3"></textarea>
                            <div class="form-text">Core purpose and values (line breaks preserved)</div>
                        </div>
                        <div class="mb-3">
                            <label for="edit-project-description" class="form-label">Description</label>
                            <textarea class="form-control" id="edit-project-description" rows="4"></textarea>
                            <div class="form-text">Line breaks are preserved when displayed</div>
                        </div>
                        <div class="mb-3">
                            <label for="edit-project-image-url" class="form-label">Image (optional)</label>
                            <input type="url" class="form-control mb-2" id="edit-project-image-url" placeholder="https://example.com/image.png or upload below">
                            <div class="input-group">
                                <input type="file" class="form-control" id="edit-project-image-file" accept="image/*">
                                <button class="btn btn-outline-primary" type="button" onclick="uploadProjectImage()">
                                    <i class="fas fa-upload"></i> Upload
                                </button>
                            </div>
                            <div class="form-text">Layer logo or banner image. Max 600×600px, 5MB. Upload or paste URL above.</div>
                            <div id="edit-project-image-upload-status" class="mt-1"></div>
                        </div>
                        <div class="mb-3">
                            <label for="edit-project-status" class="form-label">Status</label>
                            <select class="form-select" id="edit-project-status">
                                <option value="proposed">Proposed</option>
                                <option value="active">Active</option>
                                <option value="stabilizing">Stabilizing</option>
                                <option value="maintaining">Maintaining</option>
                                <option value="dormant">Dormant</option>
                                <option value="concluded">Concluded</option>
                                <option value="archived">Archived</option>
                            </select>
                        </div>
                        <div class="mb-3">
                            <label for="edit-project-status-reason" class="form-label">Status Reason (optional)</label>
                            <input type="text" class="form-control" id="edit-project-status-reason" placeholder="e.g. reason for status change">
                        </div>
                    </form>
                </div>
                <div class="modal-footer">
                    <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                    <button type="button" class="btn btn-primary" id="edit-project-save-btn" onclick="saveProjectEdit()"><i class="fas fa-save me-2"></i>Save</button>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    let project = null;
    const projectSlug = '{project_slug}';
    const initialWaitlistId = {json.dumps(initial_waitlist_id)};
    const isAuthenticated = {'true' if current_user else 'false'};
    const isAdmin = {('true' if current_user and current_user.get('is_admin') else 'false')};
    const isProjectAdmin = {'true' if show_admin_tab else 'false'};
    
    const referralRef = {json.dumps(request.args.get('ref') or '')};
    
    function escapeHtml(text) {{
        if (!text) return '';
        return String(text).split('&').join('&amp;').split('<').join('&lt;').split('>').join('&gt;').split('\\n').join('<br>');
    }}
    
    function escapeHtmlBasic(text) {{
        if (!text) return '';
        return String(text).split('&').join('&amp;').split('<').join('&lt;').split('>').join('&gt;');
    }}
    
    async function loadProject() {{
        try {{
            const response = await fetch('/api/projects/');
            const data = await response.json();
            project = data.projects.find(p => p.slug === projectSlug);
            
            if (!project) {{
                document.getElementById('project-header').innerHTML = '<div class="alert alert-danger">Layer not found</div>';
                return;
            }}
            const detailResp = await fetch('/api/projects/' + project.id + '/');
            const detail = await detailResp.json();
            project.is_member = detail.is_member === true;
            project.member_role = detail.member_role || null;
            
            displayProjectHeader();
            loadOverview();
            const wlResp = await fetch(`/api/projects/${{project.id}}/waitlists/`);
            const wlData = await wlResp.json().catch(() => ({{ waitlists: [], count: 0 }}));
            const enabledWaitlists = (wlData.waitlists || []).filter(w => w.active !== false);
            buildWaitlistTabs(enabledWaitlists);
            if (initialWaitlistId) {{
                const wl = enabledWaitlists.find(w => w.id === initialWaitlistId);
                if (wl) {{
                    document.getElementById('waitlist-tab-' + wl.id)?.click();
                }} else {{
                    showWaitlistInactiveMessage(initialWaitlistId);
                }}
            }}
        }} catch (error) {{
            console.error('Error loading project:', error);
            document.getElementById('project-header').innerHTML = '<div class="alert alert-danger">Error loading project</div>';
        }}
    }}
    
    function showJoinProjectModal() {{
        if (!isAuthenticated) {{ alert('Please sign in to join this project'); return; }}
        const refInput = document.getElementById('join-project-referral');
        if (refInput) refInput.value = referralRef || '';
        const modal = new bootstrap.Modal(document.getElementById('joinProjectModal'));
        modal.show();
    }}
    
    async function submitJoinProjectModal() {{
        const refInput = document.getElementById('join-project-referral');
        const ref = refInput && refInput.value ? refInput.value.trim() : (referralRef || '');
        const body = ref ? {{ referral_code: ref }} : {{}};
        try {{
            const res = await fetch('/api/projects/' + project.id + '/join/', {{
                method: 'POST',
                headers: {{ 'Content-Type': 'application/json' }},
                body: JSON.stringify(body)
            }});
            const data = await res.json();
            if (res.ok) {{
                project.is_member = true;
                project.member_role = data.member && data.member.role ? data.member.role : 'contributor';
                displayProjectHeader();
                bootstrap.Modal.getInstance(document.getElementById('joinProjectModal')).hide();
            }} else {{ alert(data.error || 'Failed to join'); }}
        }} catch (e) {{ console.error(e); alert('Failed to join project'); }}
    }}
    
    async function leaveProject() {{
        if (!confirm('Leave this project?')) return;
        try {{
            const res = await fetch('/api/projects/' + project.id + '/leave/', {{ method: 'POST' }});
            if (res.ok) {{
                project.is_member = false;
                project.member_role = null;
                displayProjectHeader();
            }} else {{ const d = await res.json(); alert(d.error || 'Failed to leave'); }}
        }} catch (e) {{ alert('Failed to leave project'); }}
    }}
    
    function displayProjectHeader() {{
        const statusBadge = (project.approval_status === 'approved' && project.status === 'proposed') ? '' : getStatusBadge(project.status);
        const approvalBadge = getApprovalBadge(project.approval_status);
        const isJoined = project.is_member || isProjectAdmin;
        let actionsHtml = '';
        if (project.mission) {{
            actionsHtml += '<div class="mb-3"><strong>Mission</strong><p class="mb-0 small">' + escapeHtml(project.mission || '') + '</p></div>';
        }}
        if (isAuthenticated) {{
            if (isJoined) {{
                actionsHtml += '<div class="mb-3"><span class="badge bg-success">Joined</span></div>';
            }} else {{
                actionsHtml += '<div class="mb-3"><button class="btn btn-primary btn-sm w-100" onclick="showJoinProjectModal()"><i class="fas fa-plus me-2"></i>Join Layer</button></div>';
            }}
        }}
        if (isProjectAdmin) {{
            actionsHtml += '<div class="mb-3"><button class="btn btn-outline-primary btn-sm w-100" onclick="createWaitlist()"><i class="fas fa-plus me-2"></i>Create Waitlist</button></div>';
            actionsHtml += '<div class="mb-3"><button class="btn btn-outline-primary btn-sm w-100" onclick="showCreateVoteModal()"><i class="fas fa-vote-yea me-2"></i>Create Vote</button></div>';
        }}
        actionsHtml += '<div class="mb-2"><a href="/projects/" class="btn btn-outline-secondary btn-sm w-100"><i class="fas fa-arrow-left me-2"></i>Back to Layers</a></div>';
        if (isProjectAdmin) {{
            actionsHtml += '<button class="btn btn-secondary btn-sm w-100" onclick="editProject()"><i class="fas fa-edit me-2"></i>Edit</button>';
        }}
        const imageHtml = project.image_url ? '<div class="card mb-3"><div class="card-body p-2 text-center"><img src="' + project.image_url + '" alt="' + escapeHtmlBasic(project.name) + '" class="img-fluid rounded" style="max-height: 200px; max-width: 100%;"></div></div>' : '';
        document.getElementById('project-header').innerHTML =
            '<div class="row">' +
                '<div class="col-md-8">' +
                    '<h1>' + escapeHtml(project.name) + '</h1>' +
                    '<div class="mb-3">' + statusBadge + approvalBadge + '</div>' +
                    '<p class="lead">' + escapeHtml(project.description || 'No description') + '</p>' +
                '</div>' +
                '<div class="col-md-4">' +
                    imageHtml +
                    '<div class="card">' +
                        '<div class="card-header py-2"><strong>Actions</strong></div>' +
                        '<div class="card-body py-3">' + actionsHtml + '</div>' +
                    '</div>' +
                '</div>' +
            '</div>';
    }}
    
    function loadOverview() {{
        document.getElementById('overview-content').innerHTML = `
            <div class="row">
                <div class="col-md-6">
                    <div class="card mb-4">
                        <div class="card-header"><h5>Layer Information</h5></div>
                        <div class="card-body">
                            <p><strong>Status:</strong> ${{project.status}}</p>
                            <p><strong>Approval:</strong> ${{project.approval_status}}</p>
                            <p><strong>Created:</strong> ${{new Date(project.created_at).toLocaleDateString()}}</p>
                            <p><strong>Last Activity:</strong> ${{project.last_activity_at ? new Date(project.last_activity_at).toLocaleDateString() : 'Never'}}</p>
                        </div>
                    </div>
                </div>
                <div class="col-md-6">
                    <div class="card mb-4">
                        <div class="card-header"><h5>Quick Stats</h5></div>
                        <div class="card-body">
                            <p><strong>Workgroups:</strong> ${{project.workgroups_count || 0}}</p>
                            <p><strong>Roles:</strong> <span id="roles-count">Loading...</span></p>
                            <p><strong>Active Claims:</strong> <span id="claims-count">Loading...</span></p>
                        </div>
                    </div>
                </div>
            </div>
        `;
        
        loadRolesCounts();
    }}
    
    async function loadRolesCounts() {{
        try {{
            const rolesResp = await fetch(`/api/projects/${{project.id}}/roles/`);
            const rolesData = await rolesResp.json();
            document.getElementById('roles-count').textContent = rolesData.count;
            
            const claimsResp = await fetch(`/api/projects/${{project.id}}/claims/?status=active`);
            const claimsData = await claimsResp.json();
            document.getElementById('claims-count').textContent = claimsData.count;
        }} catch (error) {{
            console.error('Error loading counts:', error);
        }}
    }}
    
    async function loadVotes() {{
        const container = document.getElementById('votes-content');
        if (!container || !project) {{
            console.log('loadVotes: container or project missing', {{container: !!container, project: !!project}});
            return;
        }}
        container.innerHTML = '<div class="py-4 text-center"><div class="spinner-border text-primary"></div></div>';
        try {{
            console.log('loadVotes: fetching from /api/projects/' + project.id + '/votes/');
            const res = await fetch(`/api/projects/${{project.id}}/votes/`);
            console.log('loadVotes: response status', res.status, res.ok);
            
            const data = await res.json();
            console.log('loadVotes: received data', data);
            
            if (!res.ok) {{
                console.error('loadVotes: API error', res.status, data.error || 'Unknown error');
                container.innerHTML = '<div class="alert alert-danger">Error loading votes: ' + (data.error || 'HTTP ' + res.status) + '</div>';
                return;
            }}
            const votes = data.votes || [];
            if (votes.length === 0) {{
                container.innerHTML = '<div class="alert alert-info">No votes yet. Layer admins can create a vote using the Create Vote button above.</div>';
                return;
            }}
            let html = '<div class="list-group">';
            votes.forEach(v => {{
                const statusBadge = v.status === 'active' ? '<span class="badge bg-success">Active</span>' : v.status === 'closed' ? '<span class="badge bg-secondary">Closed</span>' : v.status === 'scheduled' ? '<span class="badge bg-info">Scheduled</span>' : '<span class="badge bg-warning">' + (v.status || '') + '</span>';
                const resultBadge = v.result ? '<span class="badge bg-' + (v.result === 'passed' ? 'success' : v.result === 'failed' ? 'danger' : v.result === 'no_quorum' ? 'warning' : 'secondary') + ' ms-1">' + v.result + '</span>' : '';
                html += '<a href="/votes/' + v.public_id + '/" class="list-group-item list-group-item-action">';
                html += '<div class="d-flex w-100 justify-content-between"><h6 class="mb-1">' + escapeHtmlBasic(v.title) + '</h6>' + statusBadge + resultBadge + '</div>';
                html += '<p class="mb-1 small text-muted">' + escapeHtmlBasic(v.description || '') + '</p>';
                html += '<small>Start: ' + new Date(v.start_at).toLocaleString() + ' &middot; End: ' + new Date(v.end_at).toLocaleString() + '</small>';
                html += '</a>';
            }});
            html += '</div>';
            container.innerHTML = html;
        }} catch (error) {{
            console.error('Error loading votes:', error);
            container.innerHTML = '<div class="alert alert-danger">Error loading votes: ' + error.message + '</div>';
        }}
    }}
    
    async function showCreateVoteModal() {{
        document.getElementById('create-vote-alert').classList.add('d-none');
        document.getElementById('createVoteForm').reset();
        
        // Set timezone labels (handle both modal variants)
        const tzAbbr = new Date().toLocaleTimeString('en-us', {{timeZoneName:'short'}}).split(' ').pop();
        const startLabel = document.getElementById('timezone-start') || document.getElementById('timezone-start-at');
        const endLabel = document.getElementById('timezone-end') || document.getElementById('timezone-end-at');
        if (startLabel) startLabel.textContent = tzAbbr || 'your local time';
        if (endLabel) endLabel.textContent = tzAbbr || 'your local time';
        
        // Set default times: next hour + 7 days
        const now = new Date();
        const nextHour = new Date(now.getFullYear(), now.getMonth(), now.getDate(), now.getHours() + 1, 0, 0);
        const sevenDaysLater = new Date(nextHour.getTime() + 7 * 24 * 60 * 60 * 1000);
        
        const formatDatetimeLocal = (d) => {{
            const pad = (n) => String(n).padStart(2, '0');
            return `${{d.getFullYear()}}-${{pad(d.getMonth()+1)}}-${{pad(d.getDate())}}T${{pad(d.getHours())}}:${{pad(d.getMinutes())}}`;
        }};
        
        // Set default times (handle both modal variants)
        const startInput = document.getElementById('vote-start') || document.getElementById('vote-start-at');
        const endInput = document.getElementById('vote-end') || document.getElementById('vote-end-at');
        if (startInput) startInput.value = formatDatetimeLocal(nextHour);
        if (endInput) endInput.value = formatDatetimeLocal(sevenDaysLater);
        document.getElementById('vote-quorum').value = '1';
        document.getElementById('vote-threshold').value = '0.5';
        
        // Load submissions for this project
        const submissionSelect = document.getElementById('vote-submission-id');
        submissionSelect.innerHTML = '<option value="">Loading...</option>';
        try {{
            const res = await fetch(`/api/projects/${{project.id}}/submissions/`);
            const data = await res.json();
            const submissions = data.submissions || [];
            if (submissions.length === 0) {{
                submissionSelect.innerHTML = '<option value="">No approved drafts available</option>';
            }} else {{
                submissionSelect.innerHTML = '<option value="">Select a draft...</option>';
                submissions.forEach(s => {{
                    const label = s.ml_number ? `${{s.ml_number}} - ${{s.title}}` : `${{s.title}} (${{s.id}})`;
                    submissionSelect.innerHTML += `<option value="${{s.id}}">${{label}}</option>`;
                }});
            }}
        }} catch (e) {{
            submissionSelect.innerHTML = '<option value="">Error loading drafts</option>';
        }}
        
        const modal = new bootstrap.Modal(document.getElementById('createVoteModal'));
        modal.show();
    }}
    
    async function submitCreateVote() {{
        const title = document.getElementById('vote-title').value.trim();
        const description = document.getElementById('vote-description').value.trim();
        const submission_id = document.getElementById('vote-submission-id').value.trim();
        const startVal = document.getElementById('vote-start').value;
        const endVal = document.getElementById('vote-end').value;
        const quorum = parseInt(document.getElementById('vote-quorum').value, 10);
        const threshold = parseFloat(document.getElementById('vote-threshold').value);
        
        const alertEl = document.getElementById('create-vote-alert');
        alertEl.classList.add('d-none');
        if (!title || !submission_id || !startVal || !endVal) {{
            alertEl.textContent = 'Title, Submission ID, Start, and End are required';
            alertEl.className = 'alert alert-danger';
            alertEl.classList.remove('d-none');
            return;
        }}
        const startAt = new Date(startVal).toISOString();
        const endAt = new Date(endVal).toISOString();
        
        const btn = document.getElementById('create-vote-submit-btn');
        btn.disabled = true;
        btn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Creating...';
        
        try {{
            const res = await fetch(`/api/projects/${{project.id}}/votes/`, {{
                method: 'POST',
                headers: {{ 'Content-Type': 'application/json' }},
                body: JSON.stringify({{
                    title: title,
                    description: description,
                    submission_id: submission_id,
                    start_at: startAt,
                    end_at: endAt,
                    quorum_count: quorum,
                    win_threshold: threshold
                }})
            }});
            // Read body once as text, then parse as JSON
            const rawText = await res.text();
            let data = {{}};
            try {{ data = JSON.parse(rawText); }} catch {{}}
            
            if (res.ok) {{
                bootstrap.Modal.getInstance(document.getElementById('createVoteModal')).hide();
                document.getElementById('votes-tab').click();
                loadVotes();
            }} else {{
                const msg = data.error || rawText.slice(0, 300) || 'HTTP ' + res.status;
                alertEl.textContent = msg;
                alertEl.className = 'alert alert-danger';
                alertEl.classList.remove('d-none');
            }}
        }} catch (e) {{
            console.error('Create vote fetch error:', e);
            alertEl.textContent = 'Network error: ' + e.message;
            alertEl.className = 'alert alert-danger';
            alertEl.classList.remove('d-none');
        }}
        btn.disabled = false;
        btn.innerHTML = '<i class="fas fa-check me-2"></i>Create Vote';
    }}
    
    // Tab event listeners
    document.getElementById('workgroups-tab').addEventListener('shown.bs.tab', loadWorkgroups);
    document.getElementById('clusters-tab').addEventListener('shown.bs.tab', loadClusters);
    document.getElementById('roles-tab').addEventListener('shown.bs.tab', loadRoles);
    document.getElementById('claims-tab').addEventListener('shown.bs.tab', loadClaims);
    document.getElementById('votes-tab').addEventListener('shown.bs.tab', loadVotes);
    {admin_tab_listener}
    
    function buildWaitlistTabs(waitlists) {{
        const marker = document.getElementById('waitlist-tabs-marker');
        const paneMarker = document.getElementById('waitlist-panes-marker');
        if (!marker || !paneMarker) return;
        while (marker.previousElementSibling && marker.previousElementSibling.id && marker.previousElementSibling.id.startsWith('waitlist-tab-li-')) {{
            marker.previousElementSibling.remove();
        }}
        document.querySelectorAll('[id^="waitlist-pane-"]').forEach(el => el.remove());
        if (waitlists.length === 0) return;
        waitlists.forEach((w, idx) => {{
            const li = document.createElement('li');
            li.className = 'nav-item';
            li.id = 'waitlist-tab-li-' + w.id;
            li.innerHTML = '<button class="nav-link" id="waitlist-tab-' + w.id + '" data-bs-toggle="tab" data-bs-target="#waitlist-pane-' + w.id + '" type="button" data-waitlist-id="' + w.id + '">' + escapeHtmlBasic(w.name || '') + '</button>';
            marker.parentNode.insertBefore(li, marker);
            const pane = document.createElement('div');
            pane.className = 'tab-pane fade' + (idx === 0 ? '' : '');
            pane.id = 'waitlist-pane-' + w.id;
            pane.dataset.waitlistId = w.id;
            pane.innerHTML = '<div class="py-4 text-center"><div class="spinner-border text-primary"></div></div>';
            paneMarker.parentNode.insertBefore(pane, paneMarker);
            li.querySelector('button').addEventListener('shown.bs.tab', () => loadWaitlistPane(w.id));
        }});
    }}
    
    function showWaitlistInactiveMessage(waitlistId) {{
        const header = document.getElementById('project-header');
        const alert = document.createElement('div');
        alert.className = 'alert alert-warning alert-dismissible fade show';
        alert.innerHTML = '<i class="fas fa-exclamation-triangle me-2"></i>This waitlist is no longer active.<button type="button" class="btn-close" data-bs-dismiss="alert"></button>';
        header.insertAdjacentElement('afterend', alert);
    }}
    
    async function loadWaitlistPane(waitlistId) {{
        const pane = document.getElementById('waitlist-pane-' + waitlistId);
        if (!pane || !project) return;
        pane.innerHTML = '<div class="py-4 text-center"><div class="spinner-border text-primary"></div></div>';
        try {{
            const res = await fetch(`/api/projects/${{project.id}}/waitlists/`);
            const data = await res.json();
            const w = (data.waitlists || []).find(x => x.id === waitlistId);
            if (!w) {{
                pane.innerHTML = '<div class="alert alert-warning">Waitlist not found or no longer active.</div>';
                return;
            }}
            const started = w.started !== false;
            const closed = w.closed === true;
            const full = w.full === true;
            const canJoin = isAuthenticated && started && !closed && !full && !w.my_entry;
            const countStr = w.max_number != null ? `${{w.count}} of ${{w.max_number}}` : `${{w.count}}`;
            const closingStr = w.closing_date ? new Date(w.closing_date).toLocaleDateString() : '-';
            const link = w.referral_url || (window.location.origin + '/projects/' + projectSlug + '/waitlist/' + w.id + '/');
            const wName = escapeHtmlBasic(w.name || '');
            const wDesc = escapeHtmlBasic(w.description || 'No description');
            let milestonesHtml = '';
            if (w.milestones && w.milestones.length) {{
                const items = w.milestones.map(function(m) {{
                    return '<li><strong>' + escapeHtmlBasic(m.title || '') + '</strong> at ' + m.threshold + ' - ' + escapeHtmlBasic(m.description || '') + '</li>';
                }});
                milestonesHtml = '<div class="mt-3"><h6>Milestones</h6><ul class="list-unstyled small">' + items.join('') + '</ul></div>';
            }}
            let actionHtml = '';
            if (w.my_entry) {{
                actionHtml = '<span class="badge bg-success">Joined</span><span class="text-muted">Position #' + w.my_entry.position + '</span>';
            }} else if (canJoin) {{
                const wlNameForJs = (w.name || 'Waitlist').split("\\\\").join("\\\\\\\\").split("'").join("\\\\'");
                actionHtml = '<button class="btn btn-primary btn-sm" onclick="showJoinWaitlistModal(' + w.id + ', \\'' + wlNameForJs + '\\')">Join</button>';
            }} else if (!started) {{
                actionHtml = '<span class="badge bg-secondary">Not started</span>';
            }} else if (full) {{
                actionHtml = '<span class="badge bg-secondary">Full</span>';
            }} else if (closed) {{
                actionHtml = '<span class="badge bg-secondary">Closed</span>';
            }} else if (!isAuthenticated) {{
                actionHtml = '<a href="/login/" class="btn btn-primary btn-sm">Sign in to join</a>';
            }}
            const leaveBtn = w.my_entry ? '<button class="btn btn-outline-danger btn-sm" onclick="leaveWaitlist(' + w.id + ')">Leave</button>' : '';
            const linkHtml = w.referrals ? 'Your referral link: <a href="' + link + '" target="_blank">' + link + '</a>' : 'Link: <a href="' + link + '">' + link + '</a>';
            const dateStarted = w.start_date ? new Date(w.start_date).toLocaleDateString() : '-';
            const visibility = w.public ? 'Public' : 'Private';
            
            const wImg = (w.image_url) ? '<div class="mb-3"><img src="' + w.image_url + '" alt="' + wName + '" class="img-fluid rounded" style="max-height: 180px;"></div>' : '';
            const html = '<div class="card mb-4"><div class="card-body">' +
                '<nav aria-label="breadcrumb"><ol class="breadcrumb mb-2"><li class="breadcrumb-item"><a href="/projects/">Layers</a></li><li class="breadcrumb-item"><a href="/projects/' + projectSlug + '/">' + escapeHtmlBasic(project.name) + '</a></li><li class="breadcrumb-item active">' + wName + ' Waitlist</li></ol></nav>' +
                wImg +
                '<h5 class="card-title">' + wName + '</h5>' +
                '<p class="text-muted">' + wDesc + '</p>' +
                '<p class="small mb-2">Date started: ' + dateStarted + ' · ' + visibility + '</p>' +
                '<p class="small">' + linkHtml + '</p>' +
                '<div class="d-flex flex-wrap align-items-center gap-3 mt-3">' +
                actionHtml +
                '<span class="text-muted">' + countStr + ' on waitlist</span>' +
                '<span class="text-muted">Closing: ' + closingStr + '</span>' +
                leaveBtn +
                '</div>' + milestonesHtml + '</div></div>';
            pane.innerHTML = html;
        }} catch (e) {{
            console.error(e);
            pane.innerHTML = '<div class="alert alert-danger">Error loading waitlist</div>';
        }}
    }}
    
    let pendingWaitlistId = null;
    
    function showJoinWaitlistModal(waitlistId, waitlistName) {{
        if (!isAuthenticated) {{ alert('Please sign in to join this waitlist'); return; }}
        pendingWaitlistId = waitlistId;
        const titleEl = document.getElementById('join-waitlist-modal-title');
        if (titleEl) titleEl.innerHTML = '<i class="fas fa-list-alt me-2"></i>Join: ' + escapeHtmlBasic(waitlistName || 'Waitlist');
        const msgEl = document.getElementById('join-waitlist-message');
        if (msgEl) msgEl.value = '';
        const modal = new bootstrap.Modal(document.getElementById('joinWaitlistModal'));
        modal.show();
    }}
    
    async function submitJoinWaitlistModal() {{
        if (!pendingWaitlistId) return;
        const msgEl = document.getElementById('join-waitlist-message');
        const msg = msgEl ? msgEl.value : '';
        try {{
            const body = {{ message: msg || '' }};
            if (referralRef) body.referral_code = referralRef;
            const res = await fetch('/api/waitlists/' + pendingWaitlistId + '/join/', {{ method: 'POST', headers: {{ 'Content-Type': 'application/json' }}, body: JSON.stringify(body) }});
            const data = await res.json();
            if (res.ok) {{
                loadWaitlistPane(pendingWaitlistId);
                bootstrap.Modal.getInstance(document.getElementById('joinWaitlistModal')).hide();
            }} else {{ alert(data.error || 'Failed to join'); }}
        }} catch (e) {{ alert('Failed to join'); }}
        pendingWaitlistId = null;
    }}
    
    async function leaveWaitlist(waitlistId) {{
        if (!confirm('Leave this waitlist?')) return;
        try {{
            const res = await fetch(`/api/waitlists/${{waitlistId}}/leave/`, {{ method: 'POST' }});
            if (res.ok) loadWaitlistPane(waitlistId); else {{ const d = await res.json(); alert(d.error || 'Failed to leave'); }}
        }} catch (e) {{ alert('Failed to leave'); }}
    }}
    
    async function loadAdmins() {{
        const container = document.getElementById('admin-content');
        if (!container) return;
        container.innerHTML = '<div class="text-center py-5"><div class="spinner-border text-primary"></div></div>';
        
        try {{
            const response = await fetch(`/api/projects/${{project.id}}/admins/`);
            if (response.status === 403) {{
                container.innerHTML = '<div class="alert alert-warning">You do not have permission to view layer admins.</div>';
                return;
            }}
            const data = await response.json();
            
            let html = `
                <div class="d-flex justify-content-between align-items-center mb-3">
                    <h4>Layer admins</h4>
                    <button class="btn btn-primary btn-sm" onclick="showAddAdminModal()"><i class="fas fa-plus me-2"></i>Add admin</button>
                </div>
                <p class="text-muted">Admins can manage workgroups, roles, claims, and other admins. The owner cannot be removed.</p>
                <div class="list-group">
                    <div class="list-group-item d-flex justify-content-between align-items-center">
                        <div>
                            <a href="/profile/${{data.owner.username}}/" class="fw-bold text-decoration-none">${{data.owner.display_name}}</a>
                            <span class="badge bg-primary ms-2">Owner</span>
                        </div>
                        <span class="text-muted">—</span>
                    </div>
            `;
            (data.admins || []).forEach(a => {{
                html += `
                    <div class="list-group-item d-flex justify-content-between align-items-center">
                        <a href="/profile/${{a.username}}/" class="text-decoration-none">${{a.display_name}}</a>
                        <button class="btn btn-outline-danger btn-sm" onclick="removeAdmin(${{a.user_id}}, this)">Remove</button>
                    </div>
                `;
            }});
            html += '</div>';
            
            // Add pending workgroups section for approval
            html += '<hr class="my-4"><h4 class="mb-3">Pending workgroups</h4>';
            const wgResponse = await fetch(`/api/projects/${{project.id}}/workgroups/?approval_status=pending`);
            const wgData = await wgResponse.json();
            if (wgData.workgroups && wgData.workgroups.length > 0) {{
                html += '<div class="list-group">';
                wgData.workgroups.forEach(wg => {{
                    html += `
                        <div class="list-group-item">
                            <div class="d-flex justify-content-between align-items-start">
                                <div>
                                    <h6 class="mb-1">${{wg.name}}</h6>
                                    <p class="mb-1 text-muted small">${{wg.description || 'No description'}}</p>
                                </div>
                                <div class="btn-group btn-group-sm">
                                    <button class="btn btn-success" onclick="approveWorkgroup('${{wg.id}}')"><i class="fas fa-check me-1"></i>Approve</button>
                                    <button class="btn btn-danger" onclick="rejectWorkgroup('${{wg.id}}')"><i class="fas fa-times me-1"></i>Reject</button>
                                </div>
                            </div>
                        </div>
                    `;
                }});
                html += '</div>';
            }} else {{
                html += '<p class="text-muted">No pending workgroups</p>';
            }}
            
            // Add waitlist management section
            html += '<hr class="my-4"><div class="d-flex justify-content-between align-items-center mb-3"><h4>Waitlists</h4><button class="btn btn-primary btn-sm" onclick="createWaitlist()"><i class="fas fa-plus me-2"></i>Create Waitlist</button></div>';
            const wlResponse = await fetch(`/api/projects/${{project.id}}/waitlists/`);
            const wlData = await wlResponse.json();
            if (wlData.waitlists && wlData.waitlists.length > 0) {{
                html += '<div class="list-group">';
                wlData.waitlists.forEach(wl => {{
                    const statusBadge = wl.active ? '<span class="badge bg-success">Active</span>' : '<span class="badge bg-secondary">Inactive</span>';
                    const wlNameEsc = escapeHtmlBasic(wl.name || '');
                    const wlDescEsc = escapeHtmlBasic(wl.description || 'No description');
                    const wlNameAttr = (wl.name || '').split("\\\\").join("\\\\\\\\").split("'").join("\\\\'");
                    html += '<div class="list-group-item"><div class="d-flex justify-content-between align-items-start">' +
                        '<div class="flex-grow-1">' +
                        '<h6 class="mb-1">' + wlNameEsc + ' ' + statusBadge + '</h6>' +
                        '<p class="mb-1 text-muted small">' + wlDescEsc + '</p>' +
                        '<p class="mb-0 small text-muted">Members: ' + wl.count + (wl.max_number ? ' / ' + wl.max_number : '') + '</p>' +
                        '</div><div class="btn-group btn-group-sm">' +
                        '<button class="btn btn-outline-primary" onclick="showEmbedCode(' + wl.id + ', \\'' + wlNameAttr + '\\')"><i class="fas fa-code"></i></button>' +
                        '<a href="/projects/' + projectSlug + '/waitlist/' + wl.id + '/" class="btn btn-outline-secondary" target="_blank"><i class="fas fa-external-link-alt"></i></a>' +
                        '</div></div></div>';
                }});
                html += '</div>';
            }} else {{
                html += '<p class="text-muted">No waitlists yet. Create one to start collecting signups.</p>';
            }}
            
            container.innerHTML = html;
        }} catch (error) {{
            console.error('Error loading admins:', error);
            container.innerHTML = '<div class="alert alert-danger">Error loading admins</div>';
        }}
    }}
    
    async function approveWorkgroup(wgId) {{
        if (!confirm('Approve this workgroup?')) return;
        try {{
            const response = await fetch(`/api/workgroups/${{wgId}}/approve/`, {{
                method: 'POST',
                headers: {{'Content-Type': 'application/json'}},
                body: JSON.stringify({{action: 'approve'}})
            }});
            if (response.ok) {{
                loadAdmins();
                loadWorkgroups();
                alert('Workgroup approved!');
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to approve workgroup');
            }}
        }} catch (e) {{
            alert('Failed to approve workgroup');
        }}
    }}
    
    async function rejectWorkgroup(wgId) {{
        if (!confirm('Reject this workgroup?')) return;
        try {{
            const response = await fetch(`/api/workgroups/${{wgId}}/approve/`, {{
                method: 'POST',
                headers: {{'Content-Type': 'application/json'}},
                body: JSON.stringify({{action: 'reject'}})
            }});
            if (response.ok) {{
                loadAdmins();
                loadWorkgroups();
                alert('Workgroup rejected');
            }} else {{
                const data = await response.json();
                alert(data.error || 'Failed to reject workgroup');
            }}
        }} catch (e) {{
            alert('Failed to reject workgroup');
        }}
    }}
    
    function showAddAdminModal() {{
        const modal = new bootstrap.Modal(document.getElementById('addAdminModal'));
        document.getElementById('add-admin-username').value = '';
        document.getElementById('add-admin-results').innerHTML = '';
        modal.show();
    }}
    
    async function searchUsersForAdmin() {{
        const q = document.getElementById('add-admin-username').value.trim();
        const resultsEl = document.getElementById('add-admin-results');
        if (q.length < 2) {{ resultsEl.innerHTML = ''; return; }}
        const response = await fetch('/api/users/search/?q=' + encodeURIComponent(q));
        const data = await response.json();
        if (data.users.length === 0) {{
            resultsEl.innerHTML = '<p class="text-muted small">No users found</p>';
            return;
        }}
        resultsEl.innerHTML = data.users.map(u => `
            <div class="d-flex justify-content-between align-items-center border-bottom py-2">
                <span>${{u.display_name}} <small class="text-muted">@${{u.username}}</small></span>
                <button class="btn btn-sm btn-primary" onclick="addAdmin(${{u.id}})">Add</button>
            </div>
        `).join('');
    }}
    
    async function addAdmin(userId) {{
        try {{
            const response = await fetch(`/api/projects/${{project.id}}/admins/`, {{
                method: 'POST',
                headers: {{ 'Content-Type': 'application/json' }},
                body: JSON.stringify({{ user_id: userId }})
            }});
            const data = await response.json();
            if (response.ok) {{
                bootstrap.Modal.getInstance(document.getElementById('addAdminModal')).hide();
                loadAdmins();
            }} else {{
                alert(data.error || 'Failed to add admin');
            }}
        }} catch (e) {{
            alert('Failed to add admin');
        }}
    }}
    
    async function removeAdmin(userId, btn) {{
        const displayName = (btn && btn.closest('.list-group-item')) ? btn.closest('.list-group-item').querySelector('a').textContent : 'this user';
        if (!confirm('Remove "' + displayName + '" as layer admin?')) return;
        try {{
            const response = await fetch(`/api/projects/${{project.id}}/admins/${{userId}}/`, {{ method: 'DELETE' }});
            const data = await response.json();
            if (response.ok) {{
                loadAdmins();
            }} else {{
                alert(data.error || 'Failed to remove admin');
            }}
        }} catch (e) {{
            alert('Failed to remove admin');
        }}
    }}
    
    async function loadWorkgroups() {{
        document.getElementById('workgroups-content').innerHTML = '<div class="text-center py-5"><div class="spinner-border text-primary"></div></div>';
        
        try {{
            const response = await fetch(`/api/projects/${{project.id}}/workgroups/`);
            const data = await response.json();
            
            let html = `
                <div class="d-flex justify-content-between mb-3">
                    <h4>Workgroups (${{data.count}})</h4>
                    ${{isAuthenticated ? '<button class="btn btn-primary btn-sm" onclick="createWorkgroup()"><i class="fas fa-plus me-2"></i>Create Workgroup</button>' : ''}}
                </div>
            `;
            
            if (data.workgroups.length === 0) {{
                html += '<div class="alert alert-info">No workgroups yet</div>';
            }} else {{
                html += '<div class="row">';
                data.workgroups.forEach(wg => {{
                    const approvalBadge = wg.approval_status === 'pending' ? '<span class="badge bg-warning ms-2">Pending Approval</span>' : (wg.approval_status === 'rejected' ? '<span class="badge bg-danger ms-2">Rejected</span>' : '');
                    const wgImg = wg.image_url ? `<div class="card-img-top overflow-hidden" style="height: 120px; background: var(--bg-secondary, #f8f9fa);"><img src="${{wg.image_url}}" alt="${{wg.name}}" class="w-100 h-100 object-fit-cover"></div>` : '';
                    html += `
                        <div class="col-md-6 mb-3">
                            <div class="card">
                                ${{wgImg}}
                                <div class="card-body">
                                    <h5 class="card-title"><a href="/workgroups/${{wg.slug}}/">${{wg.name}}</a></h5>
                                    <p class="card-text text-muted">${{wg.description || 'No description'}}</p>
                                    <span class="badge bg-${{wg.status === 'active' ? 'success' : 'secondary'}}">${{wg.status}}</span>
                                    ${{approvalBadge}}
                                </div>
                            </div>
                        </div>
                    `;
                }});
                html += '</div>';
            }}
            
            document.getElementById('workgroups-content').innerHTML = html;
        }} catch (error) {{
            console.error('Error loading workgroups:', error);
            document.getElementById('workgroups-content').innerHTML = '<div class="alert alert-danger">Error loading workgroups</div>';
        }}
    }}
    
    async function loadClusters() {{
        document.getElementById('clusters-content').innerHTML = '<div class="text-center py-5"><div class="spinner-border text-primary"></div></div>';
        
        try {{
            const response = await fetch(`/api/projects/${{project.id}}/clusters/?include_roles=1`);
            const data = await response.json();
            
            let html = `
                <div class="d-flex justify-content-between mb-3">
                    <h4>Role Clusters (${{data.count}})</h4>
                    ${{isProjectAdmin ? '<button class="btn btn-primary btn-sm" onclick="createCluster()"><i class="fas fa-plus me-2"></i>Create Cluster</button>' : ''}}
                </div>
                <p class="text-muted">Clusters group related roles together for better organization.</p>
            `;
            
            const clusters = Array.isArray(data.clusters) ? data.clusters : [];
            if (clusters.length === 0) {{
                html += '<div class="alert alert-info">No clusters yet. Create one to organize your roles!</div>';
            }} else {{
                html += '<div class="row">';
                clusters.forEach(cluster => {{
                    if (!cluster) return;
                    const cName = (cluster.name != null && cluster.name !== '') ? cluster.name : 'Unnamed';
                    const cNameEsc = (cName + '').replace(/'/g, "\\\\'");
                    const roles = cluster.roles || [];
                    const rolesHtml = roles.length
                        ? '<ul class="list-unstyled mb-0 mt-2 small">' + roles.map(r => '<li><a href="/roles/' + (r.role_slug || r.slug || '') + '/">' + (r.title_guild || r.title_operational || 'Role') + '</a></li>').join('') + '</ul>'
                        : '<p class="text-muted small mb-0 mt-2">No roles in this cluster</p>';
                    html += `
                        <div class="col-md-6 mb-3">
                            <div class="card">
                                <div class="card-body">
                                    <div class="d-flex justify-content-between align-items-start">
                                        <div>
                                            <h5 class="card-title">${{cName}}</h5>
                                            <p class="card-text text-muted">${{(cluster.description || 'No description')}}</p>
                                            <small class="text-muted">Order: ${{cluster.order != null ? cluster.order : '—'}}</small>
                                            <div class="mt-2"><strong>Roles:</strong> ${{rolesHtml}}</div>
                                        </div>
                                        ${{isProjectAdmin ? `
                                            <div class="btn-group btn-group-sm">
                                                <button class="btn btn-outline-secondary" onclick="editCluster('${{cluster.id || ''}}')">
                                                    <i class="fas fa-edit"></i>
                                                </button>
                                                <button class="btn btn-outline-danger" onclick="deleteCluster('${{cluster.id || ''}}', '${{cNameEsc}}')">
                                                    <i class="fas fa-trash"></i>
                                                </button>
                                            </div>
                                        ` : ''}}
                                    </div>
                                </div>
                            </div>
                        </div>
                    `;
                }});
                html += '</div>';
            }}
            
            document.getElementById('clusters-content').innerHTML = html;
        }} catch (error) {{
            console.error('Error loading clusters:', error);
            document.getElementById('clusters-content').innerHTML = '<div class="alert alert-danger">Error loading clusters</div>';
        }}
    }}
    
    async function loadRoles() {{
        document.getElementById('roles-content').innerHTML = '<div class="text-center py-5"><div class="spinner-border text-primary"></div></div>';
        
        try {{
            const response = await fetch(`/api/projects/${{project.id}}/roles/`);
            const data = await response.json();
            
            let html = `
                <div class="d-flex justify-content-between mb-3">
                    <h4>Roles (${{data.count}})</h4>
                    ${{isProjectAdmin ? '<button class="btn btn-primary btn-sm" onclick="createRole()"><i class="fas fa-plus me-2"></i>Create Role</button>' : ''}}
                </div>
            `;
            
            if (data.roles.length === 0) {{
                html += '<div class="alert alert-info">No roles yet</div>';
            }} else {{
                html += '<div class="row">';
                data.roles.forEach(role => {{
                    html += `
                        <div class="col-md-6 mb-3">
                            <div class="card">
                                <div class="card-body">
                                    <h5 class="card-title">
                                        <a href="/roles/${{role.role_slug}}/">${{role.title_guild}}</a>
                                    </h5>
                                    ${{role.title_operational ? `<h6 class="card-subtitle mb-2 text-muted">${{role.title_operational}}</h6>` : ''}}
                                    <p class="card-text">${{role.description.substring(0, 150)}}...</p>
                                    <span class="badge bg-${{role.status === 'approved' ? 'success' : 'warning'}}">${{role.status}}</span>
                                    ${{role.public_visible ? '<span class="badge bg-info ms-2">Public</span>' : ''}}
                                </div>
                            </div>
                        </div>
                    `;
                }});
                html += '</div>';
            }}
            
            document.getElementById('roles-content').innerHTML = html;
        }} catch (error) {{
            console.error('Error loading roles:', error);
            document.getElementById('roles-content').innerHTML = '<div class="alert alert-danger">Error loading roles</div>';
        }}
    }}
    
    async function loadClaims() {{
        document.getElementById('claims-content').innerHTML = '<div class="text-center py-5"><div class="spinner-border text-primary"></div></div>';
        
        try {{
            const response = await fetch(`/api/projects/${{project.id}}/claims/`);
            const data = await response.json();
            
            let html = `<h4 class="mb-3">Claims (${{data.count}})</h4>`;
            
            if (!data.claims || data.claims.length === 0) {{
                html += '<div class="alert alert-info">No claims yet</div>';
            }} else {{
                html += '<div class="table-responsive"><table class="table table-hover"><thead><tr><th>Claim</th><th>Status</th><th>Role</th><th>User</th></tr></thead><tbody>';
                data.claims.forEach((claim, idx) => {{
                    const claimName = claim.role_name ? (claim.intent ? (claim.intent.substring(0, 40) + (claim.intent.length > 40 ? '…' : '')) : ('Claim: ' + claim.role_name)) : ('Claim #' + (claim.id || '').toString().slice(-6));
                    const roleName = claim.role_name || ('Role ' + (claim.role_id || '').toString().slice(-6));
                    const roleLink = claim.role_slug ? `/roles/${{claim.role_slug}}/` : '#';
                    const userName = claim.claimant_name || ('User #' + (claim.claimant_id || ''));
                    const userLink = claim.claimant_username ? '/profile/' + claim.claimant_username + '/' : '#';
                    const statusClass = claim.status === 'active' ? 'success' : (claim.status === 'pending_approval' ? 'warning' : 'secondary');
                    html += `
                        <tr class="project-claim-row" data-claim-index="${{idx}}" tabindex="0" title="Hover for claim details">
                            <td>${{claimName}}</td>
                            <td><span class="badge bg-${{statusClass}}">${{claim.status || '—'}}</span></td>
                            <td>${{claim.role_slug ? '<a href="' + roleLink + '">' + roleName + '</a>' : roleName}}</td>
                            <td>${{claim.claimant_username ? '<a href="' + userLink + '">' + userName + '</a>' : userName}}</td>
                        </tr>
                    `;
                }});
                html += '</tbody></table></div>';
            }}
            
            document.getElementById('claims-content').innerHTML = html;
            
            // Attach claim popover (same content as role detail page) on hover
            if (data.claims && data.claims.length > 0) {{
                function getClaimPopoverContent(c) {{
                    const intent = c.intent ? '<p class="mb-2"><strong>Intent:</strong><br><span style="white-space: pre-wrap; word-wrap: break-word;">' + (c.intent || '').replace(/</g, '&lt;').replace(/>/g, '&gt;') + '</span></p>' : '';
                    const links = (c.evidence_links || []).filter(u => u && u.trim());
                    const evidenceHtml = links.length ? links.map(u => '<a href="' + u + '" target="_blank" rel="noopener">' + u + '</a>').join('<br>') : '<span class="text-muted">No evidence yet</span>';
                    const termStr = c.term_duration_days ? (c.term_duration_days + ' days' + (c.term_end ? ', until ' + new Date(c.term_end).toLocaleDateString() : '')) : 'Indefinite';
                    return '<div class="text-start" style="min-width: 280px; max-width: 480px; white-space: normal; word-wrap: break-word;">' + intent +
                        '<p class="mb-2"><strong>Supporting work:</strong><br>' + evidenceHtml + '</p>' +
                        '<p class="mb-2"><strong>Term:</strong> ' + termStr + '</p>' +
                        '<p class="mb-0 small text-muted">Claimed: ' + new Date(c.created_at).toLocaleDateString() + '</p></div>';
                }}
                document.querySelectorAll('.project-claim-row').forEach(el => {{
                    const idx = parseInt(el.getAttribute('data-claim-index'), 10);
                    const claim = data.claims[idx];
                    if (claim) {{
                        new bootstrap.Popover(el, {{ content: getClaimPopoverContent(claim), html: true, trigger: 'hover focus', placement: 'auto', container: 'body' }});
                    }}
                }});
            }}
        }} catch (error) {{
            console.error('Error loading claims:', error);
            document.getElementById('claims-content').innerHTML = '<div class="alert alert-danger">Error loading claims</div>';
        }}
    }}
    
    async function loadWaitlists() {{
        document.getElementById('waitlist-content').innerHTML = '<div class="text-center py-5"><div class="spinner-border text-primary"></div></div>';
        try {{
            const response = await fetch(`/api/projects/${{project.id}}/waitlists/`);
            const data = await response.json();
            
            let html = `<div class="d-flex justify-content-between align-items-center mb-4">
                <h4>Waitlists (${{data.count}})</h4>
                ${{isProjectAdmin ? '<button class="btn btn-primary btn-sm" onclick="createWaitlist()"><i class="fas fa-plus me-2"></i>Create Waitlist</button>' : ''}}
            </div>`;
            
            if (!data.waitlists || data.waitlists.length === 0) {{
                html += '<div class="alert alert-info">No waitlists yet</div>';
            }} else {{
                data.waitlists.forEach(wl => {{
                    const started = wl.started;
                    const closed = wl.closed;
                    const full = wl.full;
                    const canJoin = isAuthenticated && started && !closed && !full && !wl.my_entry;
                    const statusBadge = full ? '<span class="badge bg-danger">Full</span>' : (closed ? '<span class="badge bg-secondary">Closed</span>' : (started ? '<span class="badge bg-success">Open</span>' : '<span class="badge bg-warning">Not Started</span>'));
                    const myEntry = wl.my_entry;
                    
                    html += `
                    <div class="card mb-4 waitlist-card" id="waitlist-${{wl.id}}" style="border-left: 4px solid var(--accent-color);">
                        <div class="card-body">
                            <div class="row">
                                <div class="col-md-8">
                                    <nav aria-label="breadcrumb">
                                        <ol class="breadcrumb">
                                            <li class="breadcrumb-item"><a href="/projects/${{project.slug}}/">Layer</a></li>
                                            <li class="breadcrumb-item active">${{wl.name}} Waitlist</li>
                                        </ol>
                                    </nav>
                                    <h3 class="mb-3">${{wl.name}}</h3>
                                    <p class="lead">${{wl.description || 'No description'}}</p>
                                    <p class="text-muted mb-2"><strong>Started:</strong> ${{new Date(wl.start_date).toLocaleDateString()}}</p>
                                    <p class="text-muted mb-2"><strong>Visibility:</strong> ${{wl.public ? 'Public' : 'Private (layer members or link)'}}</p>
                                    ${{wl.referral_url && myEntry ? '<p class="mb-2"><strong>Your referral link:</strong> <code class="user-select-all">' + wl.referral_url + '</code> <button class="btn btn-sm btn-outline-primary" onclick="copyText(\\''+wl.referral_url+'\\')"><i class="fas fa-copy"></i></button></p>' : ''}}
                                    ${{!myEntry && wl.referrals ? '<p class="text-muted mb-2"><em>Join to get your referral link</em></p>' : ''}}
                                    
                                    ${{wl.milestones && wl.milestones.length > 0 ? '<hr><h5 class="mt-3">Milestones</h5><ul class="list-unstyled">' + wl.milestones.map(m => '<li class="mb-2"><strong>' + m.title + '</strong> (at ' + m.threshold + ' members)' + (m.description ? '<br><small class="text-muted">' + m.description + '</small>' : '') + '</li>').join('') + '</ul>' : ''}}
                                </div>
                                <div class="col-md-4">
                                    <div class="card">
                                        <div class="card-body">
                                            <h5 class="card-title">Actions</h5>
                                            ${{myEntry ? '<div class="mb-3"><span class="badge bg-success fs-6">Joined</span> <span class="text-muted">#' + myEntry.position + '</span><br><button class="btn btn-outline-danger btn-sm mt-2" onclick="leaveWaitlist(' + wl.id + ')">Leave</button></div>' : (canJoin ? '<button class="btn btn-primary w-100 mb-3" onclick="joinWaitlist(' + wl.id + ')">Join Waitlist</button>' : '<p class="text-muted">' + (started ? (closed ? 'Closed' : (full ? 'Full' : 'Login to join')) : 'Not started') + '</p>')}}
                                            <p class="mb-2"><strong>On waitlist:</strong> ${{wl.count}}${{wl.max_number ? ' of ' + wl.max_number : ''}}</p>
                                            ${{wl.closing_date ? '<p class="mb-2"><strong>Closes:</strong> ' + new Date(wl.closing_date).toLocaleDateString() + '</p>' : ''}}
                                            ${{statusBadge}}
                                            ${{isProjectAdmin ? '<hr><button class="btn btn-outline-primary btn-sm w-100 mt-2" onclick="showEmbedCode(' + wl.id + ', this.dataset.wlName)" data-wl-name="' + wl.name + '"><i class="fas fa-code me-2"></i>Get Embed Code</button>' : ''}}
                                        </div>
                                    </div>
                                </div>
                            </div>
                        </div>
                    </div>
                    `;
                }});
            }}
            
            document.getElementById('waitlist-content').innerHTML = html;
        }} catch (error) {{
            console.error('Error loading waitlists:', error);
            document.getElementById('waitlist-content').innerHTML = '<div class="alert alert-danger">Error loading waitlists</div>';
        }}
    }}
    
    async function joinWaitlist(wlId) {{
        if (!isAuthenticated) {{ alert('Please sign in to join'); return; }}
        const msg = prompt('Optional message:');
        if (msg === null) return;
        try {{
            const body = {{ message: msg }};
            if (referralRef) body.referral_code = referralRef;
            const res = await fetch(`/api/waitlists/${{wlId}}/join/`, {{
                method: 'POST',
                headers: {{ 'Content-Type': 'application/json' }},
                body: JSON.stringify(body)
            }});
            const d = await res.json();
            if (res.ok) {{
                alert('Joined! Position: #' + d.entry.position);
                loadWaitlists();
            }} else {{ alert(d.error || 'Failed to join'); }}
        }} catch (e) {{ alert('Failed to join waitlist'); }}
    }}
    
    async function leaveWaitlist(wlId) {{
        if (!confirm('Leave this waitlist?')) return;
        try {{
            const res = await fetch(`/api/waitlists/${{wlId}}/leave/`, {{ method: 'POST' }});
            if (res.ok) {{ loadWaitlists(); }} else {{ alert('Failed to leave'); }}
        }} catch (e) {{ alert('Failed to leave waitlist'); }}
    }}
    
    function copyText(text) {{
        navigator.clipboard.writeText(text).then(() => {{
            const btn = event.target.closest('button');
            if (btn) {{ const o = btn.innerHTML; btn.innerHTML = '<i class="fas fa-check"></i>'; setTimeout(() => btn.innerHTML = o, 1500); }}
        }});
    }}
    
    function showEmbedCode(waitlistId, waitlistName) {{
        const baseUrl = window.location.origin;
        const embedUrl = `${{baseUrl}}/embed/waitlist/${{waitlistId}}/`;
        const iframeCode = `<iframe src="${{embedUrl}}" width="100%" height="600" frameborder="0" style="border: none; border-radius: 12px;"></iframe>`;
        
        document.getElementById('embed-waitlist-name').textContent = waitlistName;
        document.getElementById('embed-url').value = embedUrl;
        document.getElementById('embed-code-iframe').value = iframeCode;
        
        const modal = new bootstrap.Modal(document.getElementById('embedCodeModal'));
        modal.show();
    }}
    
    function copyEmbedCode(type) {{
        const elementId = type === 'iframe' ? 'embed-code-iframe' : 'embed-url';
        const element = document.getElementById(elementId);
        element.select();
        navigator.clipboard.writeText(element.value).then(() => {{
            const btn = event.target.closest('button');
            if (btn) {{
                const original = btn.innerHTML;
                btn.innerHTML = '<i class="fas fa-check"></i> Copied!';
                setTimeout(() => btn.innerHTML = original, 2000);
            }}
        }});
    }}
    
    async function uploadWaitlistImage() {{
        const fileInput = document.getElementById('wl-image-file');
        const statusEl = document.getElementById('wl-image-upload-status');
        const urlInput = document.getElementById('wl-image-url');
        
        if (!fileInput.files || !fileInput.files[0]) {{
            statusEl.innerHTML = '<small class="text-danger">Please select a file first</small>';
            return;
        }}
        
        const formData = new FormData();
        formData.append('file', fileInput.files[0]);
        formData.append('entity_type', 'waitlist');
        
        statusEl.innerHTML = '<small class="text-info"><i class="fas fa-spinner fa-spin"></i> Uploading...</small>';
        
        try {{
            const response = await fetch('/api/upload/entity-image', {{
                method: 'POST',
                credentials: 'include',
                body: formData
            }});
            
            const data = await response.json();
            
            if (response.ok && data.image_url) {{
                urlInput.value = data.image_url;
                statusEl.innerHTML = '<small class="text-success"><i class="fas fa-check"></i> Uploaded successfully</small>';
                fileInput.value = '';
            }} else {{
                statusEl.innerHTML = `<small class="text-danger">${{data.error || 'Upload failed'}}</small>`;
            }}
        }} catch (error) {{
            console.error('Upload error:', error);
            statusEl.innerHTML = '<small class="text-danger">Upload failed. Please try again.</small>';
        }}
    }}
    
    function createWaitlist() {{
        const modalHtml = `
            <div class="modal fade" id="createWaitlistModal" tabindex="-1">
                <div class="modal-dialog modal-lg">
                    <div class="modal-content">
                        <div class="modal-header">
                            <h5 class="modal-title">Create Waitlist</h5>
                            <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                        </div>
                        <div class="modal-body">
                            <div id="wl-alert-container"></div>
                            <form id="createWaitlistForm">
                                <div class="mb-3">
                                    <label for="wl-name" class="form-label">Waitlist Name *</label>
                                    <input type="text" class="form-control" id="wl-name" required>
                                </div>
                                <div class="mb-3">
                                    <label for="wl-description" class="form-label">Description</label>
                                    <textarea class="form-control" id="wl-description" rows="3"></textarea>
                                </div>
                                <div class="row">
                                    <div class="col-md-6 mb-3">
                                        <label for="wl-start-date" class="form-label">Start Date *</label>
                                        <input type="date" class="form-control" id="wl-start-date" required>
                                    </div>
                                    <div class="col-md-6 mb-3">
                                        <label for="wl-closing-date" class="form-label">Closing Date (optional)</label>
                                        <input type="date" class="form-control" id="wl-closing-date">
                                    </div>
                                </div>
                                <div class="mb-3">
                                    <label for="wl-max-number" class="form-label">Max Number of Entries (optional)</label>
                                    <input type="number" class="form-control" id="wl-max-number" min="1">
                                </div>
                                <div class="form-check mb-3">
                                    <input class="form-check-input" type="checkbox" id="wl-public" checked>
                                    <label class="form-check-label" for="wl-public">Public (visible to all)</label>
                                </div>
                                <div class="form-check mb-3">
                                    <input class="form-check-input" type="checkbox" id="wl-referrals" checked>
                                    <label class="form-check-label" for="wl-referrals">Enable Referrals</label>
                                </div>
                                <div class="form-check mb-3">
                                    <input class="form-check-input" type="checkbox" id="wl-active" checked>
                                    <label class="form-check-label" for="wl-active">Active</label>
                                </div>
                                <div class="mb-3">
                                    <label for="wl-image-url" class="form-label">Image (optional)</label>
                                    <input type="url" class="form-control mb-2" id="wl-image-url" placeholder="https://example.com/image.png or upload below">
                                    <div class="input-group">
                                        <input type="file" class="form-control" id="wl-image-file" accept="image/*">
                                        <button class="btn btn-outline-primary" type="button" onclick="uploadWaitlistImage()">
                                            <i class="fas fa-upload"></i> Upload
                                        </button>
                                    </div>
                                    <div class="form-text">Waitlist banner or icon. Max 600×600px, 5MB. Upload or paste URL above.</div>
                                    <div id="wl-image-upload-status" class="mt-1"></div>
                                </div>
                            </form>
                        </div>
                        <div class="modal-footer">
                            <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                            <button type="button" class="btn btn-primary" id="submitWaitlistBtn">
                                <i class="fas fa-plus me-2"></i>Create Waitlist
                            </button>
                        </div>
                    </div>
                </div>
            </div>
        `;
        if (!document.getElementById('createWaitlistModal')) {{
            document.body.insertAdjacentHTML('beforeend', modalHtml);
        }}
        document.getElementById('wl-alert-container').innerHTML = '';
        document.getElementById('wl-name').value = '';
        document.getElementById('wl-description').value = '';
        document.getElementById('wl-image-url').value = '';
        document.getElementById('wl-start-date').value = new Date().toISOString().split('T')[0];
        document.getElementById('wl-closing-date').value = '';
        document.getElementById('wl-max-number').value = '';
        document.getElementById('wl-public').checked = true;
        document.getElementById('wl-referrals').checked = true;
        document.getElementById('wl-active').checked = true;
        if (document.getElementById('wl-image-url')) document.getElementById('wl-image-url').value = '';
        if (document.getElementById('wl-image-file')) document.getElementById('wl-image-file').value = '';
        if (document.getElementById('wl-image-upload-status')) document.getElementById('wl-image-upload-status').innerHTML = '';
        const modal = new bootstrap.Modal(document.getElementById('createWaitlistModal'));
        modal.show();
        document.getElementById('submitWaitlistBtn').onclick = async () => {{
            const name = document.getElementById('wl-name').value.trim();
            const description = document.getElementById('wl-description').value.trim();
            const image_url = document.getElementById('wl-image-url').value.trim();
            const startDate = document.getElementById('wl-start-date').value;
            const closingDate = document.getElementById('wl-closing-date').value;
            const maxNumber = document.getElementById('wl-max-number').value;
            const isPublic = document.getElementById('wl-public').checked;
            const referrals = document.getElementById('wl-referrals').checked;
            const active = document.getElementById('wl-active').checked;
            if (!name || !startDate) {{
                document.getElementById('wl-alert-container').innerHTML = '<div class="alert alert-danger">Name and start date are required.</div>';
                return;
            }}
            const btn = document.getElementById('submitWaitlistBtn');
            btn.disabled = true;
            btn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Creating...';
            try {{
                const body = {{ name, description, image_url: image_url || null, start_date: startDate, public: isPublic, referrals, active }};
                if (closingDate) body.closing_date = closingDate;
                if (maxNumber) body.max_number = parseInt(maxNumber, 10);
                const res = await fetch(`/api/projects/${{project.id}}/waitlists/`, {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify(body)
                }});
                const d = await res.json();
                if (res.ok) {{
                    modal.hide();
                    loadWaitlists();
                }} else {{
                    document.getElementById('wl-alert-container').innerHTML = '<div class="alert alert-danger">' + (d.error || 'Failed to create waitlist') + '</div>';
                    btn.disabled = false;
                    btn.innerHTML = '<i class="fas fa-plus me-2"></i>Create Waitlist';
                }}
            }} catch (e) {{
                document.getElementById('wl-alert-container').innerHTML = '<div class="alert alert-danger">Failed to create waitlist</div>';
                btn.disabled = false;
                btn.innerHTML = '<i class="fas fa-plus me-2"></i>Create Waitlist';
            }}
        }};
    }}
    
    function getStatusBadge(status) {{
        const badges = {{
            'proposed': '<span class="badge bg-info">Proposed</span>',
            'active': '<span class="badge bg-success">Active</span>',
            'stabilizing': '<span class="badge bg-primary">Stabilizing</span>',
            'maintaining': '<span class="badge bg-secondary">Maintaining</span>',
            'dormant': '<span class="badge bg-warning">Dormant</span>',
            'concluded': '<span class="badge bg-dark">Concluded</span>',
            'archived': '<span class="badge bg-secondary">Archived</span>'
        }};
        return badges[status] || '';
    }}
    
    function getApprovalBadge(approval) {{
        const badges = {{
            'pending': '<span class="badge bg-warning">Pending Approval</span>',
            'approved': '<span class="badge bg-success">Approved</span>',
            'rejected': '<span class="badge bg-danger">Rejected</span>'
        }};
        return badges[approval] || '';
    }}
    
    async function uploadProjectImage() {{
        const fileInput = document.getElementById('edit-project-image-file');
        const statusEl = document.getElementById('edit-project-image-upload-status');
        const urlInput = document.getElementById('edit-project-image-url');
        
        if (!fileInput.files || !fileInput.files[0]) {{
            statusEl.innerHTML = '<small class="text-danger">Please select a file first</small>';
            return;
        }}
        
        const formData = new FormData();
        formData.append('file', fileInput.files[0]);
        formData.append('entity_type', 'project');
        
        statusEl.innerHTML = '<small class="text-info"><i class="fas fa-spinner fa-spin"></i> Uploading...</small>';
        
        try {{
            const response = await fetch('/api/upload/entity-image', {{
                method: 'POST',
                credentials: 'include',
                body: formData
            }});
            
            const data = await response.json();
            
            if (response.ok && data.image_url) {{
                urlInput.value = data.image_url;
                statusEl.innerHTML = '<small class="text-success"><i class="fas fa-check"></i> Uploaded successfully</small>';
                fileInput.value = '';
            }} else {{
                statusEl.innerHTML = `<small class="text-danger">${{data.error || 'Upload failed'}}</small>`;
            }}
        }} catch (error) {{
            console.error('Upload error:', error);
            statusEl.innerHTML = '<small class="text-danger">Upload failed. Please try again.</small>';
        }}
    }}
    
    function editProject() {{
        if (!project) return;
        document.getElementById('edit-project-name').value = project.name || '';
        document.getElementById('edit-project-mission').value = project.mission || '';
        document.getElementById('edit-project-description').value = project.description || '';
        document.getElementById('edit-project-image-url').value = project.image_url || '';
        document.getElementById('edit-project-image-file').value = '';
        document.getElementById('edit-project-image-upload-status').innerHTML = '';
        document.getElementById('edit-project-status').value = project.status || 'proposed';
        document.getElementById('edit-project-status-reason').value = project.status_reason || '';
        document.getElementById('edit-modal-add-admin-q').value = '';
        document.getElementById('edit-modal-add-admin-results').innerHTML = '';
        const alertEl = document.getElementById('edit-project-alert');
        alertEl.classList.add('d-none');
        alertEl.textContent = '';
        loadEditModalAdmins();
        const modal = new bootstrap.Modal(document.getElementById('editProjectModal'));
        modal.show();
    }}
    
    async function loadEditModalAdmins() {{
        const container = document.getElementById('edit-modal-admins-list');
        if (!container || !project) return;
        container.innerHTML = '<div class="list-group-item text-muted small">Loading...</div>';
        try {{
            const response = await fetch('/api/projects/' + project.id + '/admins/', {{ credentials: 'include' }});
            if (response.status === 403) {{
                container.innerHTML = '<div class="list-group-item text-warning small">You need layer admin access to view or manage admins.</div>';
                return;
            }}
            if (!response.ok) {{
                container.innerHTML = '<div class="list-group-item text-danger small">Failed to load admins (status ' + response.status + '). Try refreshing.</div>';
                return;
            }}
            const data = await response.json();
            let html = `
                <div class="list-group-item d-flex justify-content-between align-items-center py-2">
                    <div>
                        <a href="/profile/${{data.owner.username}}/" class="fw-bold text-decoration-none small">${{data.owner.display_name}}</a>
                        <span class="badge bg-primary ms-2">Owner</span>
                    </div>
                    <span class="text-muted small">-</span>
                </div>
            `;
            (data.admins || []).forEach(a => {{
                html += `
                    <div class="list-group-item d-flex justify-content-between align-items-center py-2" id="edit-modal-admin-${{a.user_id}}">
                        <a href="/profile/${{a.username}}/" class="text-decoration-none small">${{a.display_name}}</a>
                        <button type="button" class="btn btn-outline-danger btn-sm" onclick="removeAdminFromEditModal(${{a.user_id}})">Remove</button>
                    </div>
                `;
            }});
            container.innerHTML = html || '<div class="list-group-item text-muted small">No assigned admins</div>';
        }} catch (e) {{
            console.error('loadEditModalAdmins error:', e);
            container.innerHTML = '<div class="list-group-item text-danger small">Failed to load admins. Check console for details.</div>';
        }}
    }}
    
    function searchUsersForEditModalAdmin() {{
        const q = document.getElementById('edit-modal-add-admin-q').value.trim();
        const resultsEl = document.getElementById('edit-modal-add-admin-results');
        if (q.length < 2) {{ resultsEl.innerHTML = ''; return; }}
        fetch('/api/users/search/?q=' + encodeURIComponent(q))
            .then(r => r.json())
            .then(data => {{
                if (!data.users || data.users.length === 0) {{
                    resultsEl.innerHTML = '<p class="text-muted small mb-0">No users found</p>';
                    return;
                }}
                resultsEl.innerHTML = data.users.map(u => `
                    <div class="d-flex justify-content-between align-items-center border-bottom py-1 small">
                        <span>${{u.display_name}} <small class="text-muted">@${{u.username}}</small></span>
                        <button type="button" class="btn btn-sm btn-primary" onclick="addAdminFromEditModal(${{u.id}})">Add</button>
                    </div>
                `).join('');
            }})
            .catch(() => {{ resultsEl.innerHTML = '<p class="text-danger small mb-0">Search failed</p>'; }});
    }}
    
    async function addAdminFromEditModal(userId) {{
        try {{
            const response = await fetch('/api/projects/' + project.id + '/admins/', {{
                method: 'POST',
                headers: {{ 'Content-Type': 'application/json' }},
                credentials: 'include',
                body: JSON.stringify({{ user_id: userId }})
            }});
            const data = await response.json();
            if (response.ok) {{
                document.getElementById('edit-modal-add-admin-q').value = '';
                document.getElementById('edit-modal-add-admin-results').innerHTML = '';
                loadEditModalAdmins();
                if (typeof loadAdmins === 'function') loadAdmins();
            }} else {{
                alert(data.error || 'Failed to add admin');
            }}
        }} catch (e) {{
            alert('Failed to add admin');
        }}
    }}
    
    async function removeAdminFromEditModal(userId) {{
        if (!confirm('Remove this user as layer admin?')) return;
        try {{
            const response = await fetch('/api/projects/' + project.id + '/admins/' + userId + '/', {{
                method: 'DELETE',
                credentials: 'include'
            }});
            const data = await response.json();
            if (response.ok) {{
                loadEditModalAdmins();
                if (typeof loadAdmins === 'function') loadAdmins();
            }} else {{
                alert(data.error || 'Failed to remove admin');
            }}
        }} catch (e) {{
            alert('Failed to remove admin');
        }}
    }}
    
    async function saveProjectEdit() {{
        if (!project) return;
        const name = document.getElementById('edit-project-name').value.trim();
        const mission = document.getElementById('edit-project-mission').value;
        const description = document.getElementById('edit-project-description').value;
        const image_url = document.getElementById('edit-project-image-url').value.trim();
        const status = document.getElementById('edit-project-status').value;
        const status_reason = document.getElementById('edit-project-status-reason').value.trim();
        const alertEl = document.getElementById('edit-project-alert');
        const saveBtn = document.getElementById('edit-project-save-btn');
        
        if (!name) {{
            alertEl.textContent = 'Layer name is required.';
            alertEl.className = 'alert alert-danger';
            alertEl.classList.remove('d-none');
            return;
        }}
        
        saveBtn.disabled = true;
        alertEl.classList.add('d-none');
        try {{
            const res = await fetch('/api/projects/' + project.id + '/', {{
                method: 'PATCH',
                headers: {{ 'Content-Type': 'application/json' }},
                credentials: 'include',
                body: JSON.stringify({{ name: name, mission: mission || null, description: description, image_url: image_url || null, status: status, status_reason: status_reason || null }})
            }});
            let data;
            try {{ data = await res.json(); }} catch (_) {{
                alertEl.textContent = res.status === 401 ? 'Please sign in to edit projects.' : 'Server error. Please try again.';
                alertEl.className = 'alert alert-danger';
                alertEl.classList.remove('d-none');
                saveBtn.disabled = false;
                return;
            }}
            if (res.ok) {{
                bootstrap.Modal.getInstance(document.getElementById('editProjectModal')).hide();
                project = data.project;
                if (project.slug !== projectSlug) {{
                    window.location.href = '/projects/' + project.slug + '/';
                    return;
                }}
                displayProjectHeader();
                loadOverview();
            }} else {{
                alertEl.textContent = data.error || 'Failed to update project';
                alertEl.className = 'alert alert-danger';
                alertEl.classList.remove('d-none');
            }}
        }} catch (e) {{
            alertEl.textContent = 'Network error. Please try again.';
            alertEl.className = 'alert alert-danger';
            alertEl.classList.remove('d-none');
        }}
        saveBtn.disabled = false;
    }}
    
    function createWorkgroup() {{
        const modalHtml = `
            <div class="modal fade" id="projectCreateWorkgroupModal" tabindex="-1">
                <div class="modal-dialog">
                    <div class="modal-content">
                        <div class="modal-header">
                            <h5 class="modal-title">Create Workgroup</h5>
                            <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                        </div>
                        <div class="modal-body">
                            <div id="project-wg-alert-container"></div>
                            <form id="projectCreateWorkgroupForm">
                                <div class="mb-3">
                                    <label for="project-wg-name" class="form-label">Workgroup Name *</label>
                                    <input type="text" class="form-control" id="project-wg-name" required>
                                </div>
                                <div class="mb-3">
                                    <label for="project-wg-description" class="form-label">Description *</label>
                                    <textarea class="form-control" id="project-wg-description" rows="3" required></textarea>
                                </div>
                                <p class="text-muted small">New workgroups require approval from the layer admin before becoming active.</p>
                            </form>
                        </div>
                        <div class="modal-footer">
                            <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                            <button type="button" class="btn btn-primary" id="projectSubmitWorkgroupBtn">
                                <i class="fas fa-plus me-2"></i>Create Workgroup
                            </button>
                        </div>
                    </div>
                </div>
            </div>
        `;
        if (!document.getElementById('projectCreateWorkgroupModal')) {{
            document.body.insertAdjacentHTML('beforeend', modalHtml);
        }}
        document.getElementById('project-wg-alert-container').innerHTML = '';
        document.getElementById('project-wg-name').value = '';
        document.getElementById('project-wg-description').value = '';
        const modal = new bootstrap.Modal(document.getElementById('projectCreateWorkgroupModal'));
        modal.show();
        document.getElementById('projectSubmitWorkgroupBtn').onclick = async () => {{
            const name = document.getElementById('project-wg-name').value.trim();
            const description = document.getElementById('project-wg-description').value.trim();
            if (!name || !description) {{
                document.getElementById('project-wg-alert-container').innerHTML = '<div class="alert alert-danger">Name and description are required.</div>';
                return;
            }}
            const btn = document.getElementById('projectSubmitWorkgroupBtn');
            btn.disabled = true;
            btn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Creating...';
            try {{
                const response = await fetch(`/api/projects/${{project.id}}/workgroups/`, {{
                    method: 'POST',
                    headers: {{'Content-Type': 'application/json'}},
                    body: JSON.stringify({{ name, description }})
                }});
                const data = await response.json();
                if (response.ok) {{
                    modal.hide();
                    loadWorkgroups();
                    alert('Workgroup created! It will be visible once approved by the layer admin.');
                }} else {{
                    throw new Error(data.error || 'Failed to create workgroup');
                }}
            }} catch (err) {{
                document.getElementById('project-wg-alert-container').innerHTML = '<div class="alert alert-danger">' + (err.message || 'Failed to create workgroup') + '</div>';
            }}
            btn.disabled = false;
            btn.innerHTML = '<i class="fas fa-plus me-2"></i>Create Workgroup';
        }};
    }}
    
    function createRole() {{
        // Create modal HTML
        const modalHtml = `
            <div class="modal fade" id="createRoleModal" tabindex="-1">
                <div class="modal-dialog modal-lg">
                    <div class="modal-content">
                        <div class="modal-header">
                            <h5 class="modal-title">Create New Role</h5>
                            <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                        </div>
                        <div class="modal-body">
                            <div id="role-alert-container"></div>
                            
                            <form id="createRoleForm">
                                <div class="mb-3">
                                    <label for="role-title-guild" class="form-label">Guild Title *</label>
                                    <input type="text" class="form-control" id="role-title-guild" required>
                                    <div class="form-text">The formal/ceremonial title (e.g., "Keeper of the Keys")</div>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="role-title-operational" class="form-label">Operational Title</label>
                                    <input type="text" class="form-control" id="role-title-operational">
                                    <div class="form-text">Optional: The practical title (e.g., "Security Lead")</div>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="role-description" class="form-label">Description *</label>
                                    <textarea class="form-control" id="role-description" rows="4" required></textarea>
                                    <div class="form-text">Describe the role's responsibilities and purpose</div>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="role-cluster" class="form-label">Cluster</label>
                                    <select class="form-select" id="role-cluster">
                                        <option value="">No cluster</option>
                                    </select>
                                    <div class="form-text">Optional: Group this role with others</div>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="role-image-url" class="form-label">Image URL</label>
                                    <input type="url" class="form-control" id="role-image-url">
                                    <div class="form-text">Optional: URL to role image/icon</div>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="role-order" class="form-label">Display Order</label>
                                    <input type="number" class="form-control" id="role-order" value="0">
                                    <div class="form-text">Lower numbers appear first</div>
                                </div>
                                
                                <div class="mb-3">
                                    <div class="form-check">
                                        <input class="form-check-input" type="checkbox" id="role-public-visible" checked>
                                        <label class="form-check-label" for="role-public-visible">
                                            Public Visible
                                        </label>
                                    </div>
                                    <div class="form-check">
                                        <input class="form-check-input" type="checkbox" id="role-claim-approval">
                                        <label class="form-check-label" for="role-claim-approval">
                                            Claims Require Approval
                                        </label>
                                    </div>
                                    <div class="form-check">
                                        <input class="form-check-input" type="checkbox" id="role-badge-enabled" checked>
                                        <label class="form-check-label" for="role-badge-enabled">
                                            Badge Enabled
                                        </label>
                                    </div>
                                    <div class="form-check">
                                        <input class="form-check-input" type="checkbox" id="role-badge-approval">
                                        <label class="form-check-label" for="role-badge-approval">
                                            Badges Require Approval
                                        </label>
                                    </div>
                                </div>
                            </form>
                        </div>
                        <div class="modal-footer">
                            <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                            <button type="button" class="btn btn-primary" id="submitRoleBtn">
                                <i class="fas fa-plus me-2"></i>Create Role
                            </button>
                        </div>
                    </div>
                </div>
            </div>
        `;
        
        // Add modal to page if not exists
        if (!document.getElementById('createRoleModal')) {{
            document.body.insertAdjacentHTML('beforeend', modalHtml);
        }}
        
        // Load clusters for dropdown
        fetch(`/api/projects/${{project.id}}/clusters/`)
            .then(r => r.json())
            .then(data => {{
                const select = document.getElementById('role-cluster');
                (data.clusters || []).forEach(cluster => {{
                    if (!cluster) return;
                    const option = document.createElement('option');
                    option.value = cluster.id || '';
                    option.textContent = (cluster.name != null && cluster.name !== '') ? cluster.name : 'Unnamed';
                    select.appendChild(option);
                }});
            }});
        
        // Show modal
        const modal = new bootstrap.Modal(document.getElementById('createRoleModal'));
        modal.show();
        
        // Handle form submission
        document.getElementById('submitRoleBtn').onclick = async () => {{
            const titleGuild = document.getElementById('role-title-guild').value.trim();
            const titleOperational = document.getElementById('role-title-operational').value.trim();
            const description = document.getElementById('role-description').value.trim();
            
            if (!titleGuild || !description) {{
                document.getElementById('role-alert-container').innerHTML = `
                    <div class="alert alert-danger">
                        <i class="fas fa-exclamation-circle me-2"></i>
                        Guild title and description are required
                    </div>
                `;
                return;
            }}
            
            const submitBtn = document.getElementById('submitRoleBtn');
            submitBtn.disabled = true;
            submitBtn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Creating...';
            
            const formData = {{
                title_guild: titleGuild,
                title_operational: titleOperational || null,
                description: description,
                cluster_id: document.getElementById('role-cluster').value || null,
                image_url: document.getElementById('role-image-url').value || null,
                order: parseInt(document.getElementById('role-order').value) || 0,
                public_visible: document.getElementById('role-public-visible').checked,
                claim_requires_approval: document.getElementById('role-claim-approval').checked,
                badge_enabled: document.getElementById('role-badge-enabled').checked,
                badge_requires_approval: document.getElementById('role-badge-approval').checked
            }};
            
            try {{
                const response = await fetch(`/api/projects/${{project.id}}/roles/`, {{
                    method: 'POST',
                    headers: {{'Content-Type': 'application/json'}},
                    body: JSON.stringify(formData)
                }});
                
                const data = await response.json();
                
                if (response.ok) {{
                    modal.hide();
                    loadRoles(); // Reload roles list
                    alert('Role created successfully!');
                }} else {{
                    throw new Error(data.error || 'Failed to create role');
                }}
            }} catch (error) {{
                document.getElementById('role-alert-container').innerHTML = `
                    <div class="alert alert-danger">
                        <i class="fas fa-exclamation-circle me-2"></i>
                        ${{error.message}}
                    </div>
                `;
                submitBtn.disabled = false;
                submitBtn.innerHTML = '<i class="fas fa-plus me-2"></i>Create Role';
            }}
        }};
    }}
    
    function createCluster() {{
        const modalHtml = `
            <div class="modal fade" id="createClusterModal" tabindex="-1">
                <div class="modal-dialog">
                    <div class="modal-content">
                        <div class="modal-header">
                            <h5 class="modal-title">Create New Cluster</h5>
                            <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                        </div>
                        <div class="modal-body">
                            <div id="cluster-alert-container"></div>
                            
                            <form id="createClusterForm">
                                <div class="mb-3">
                                    <label for="cluster-name" class="form-label">Cluster Name *</label>
                                    <input type="text" class="form-control" id="cluster-name" required>
                                    <div class="form-text">A descriptive name for this role group</div>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="cluster-description" class="form-label">Description</label>
                                    <textarea class="form-control" id="cluster-description" rows="3"></textarea>
                                    <div class="form-text">Optional: Describe the purpose of this cluster</div>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="cluster-order" class="form-label">Display Order</label>
                                    <input type="number" class="form-control" id="cluster-order" value="0">
                                    <div class="form-text">Lower numbers appear first</div>
                                </div>
                            </form>
                        </div>
                        <div class="modal-footer">
                            <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                            <button type="button" class="btn btn-primary" id="submitClusterBtn">
                                <i class="fas fa-plus me-2"></i>Create Cluster
                            </button>
                        </div>
                    </div>
                </div>
            </div>
        `;
        
        if (!document.getElementById('createClusterModal')) {{
            document.body.insertAdjacentHTML('beforeend', modalHtml);
        }}
        
        // Reset form and button so each open is fresh (and not stuck from a previous submit)
        document.getElementById('cluster-name').value = '';
        document.getElementById('cluster-description').value = '';
        document.getElementById('cluster-order').value = '0';
        document.getElementById('cluster-alert-container').innerHTML = '';
        const submitClusterBtn = document.getElementById('submitClusterBtn');
        submitClusterBtn.disabled = false;
        submitClusterBtn.innerHTML = '<i class="fas fa-plus me-2"></i>Create Cluster';
        
        const modal = new bootstrap.Modal(document.getElementById('createClusterModal'));
        modal.show();
        
        document.getElementById('submitClusterBtn').onclick = async () => {{
            const name = document.getElementById('cluster-name').value.trim();
            const description = document.getElementById('cluster-description').value.trim();
            const order = parseInt(document.getElementById('cluster-order').value) || 0;
            
            if (!name) {{
                document.getElementById('cluster-alert-container').innerHTML = `
                    <div class="alert alert-danger">
                        <i class="fas fa-exclamation-circle me-2"></i>
                        Cluster name is required
                    </div>
                `;
                return;
            }}
            
            const submitBtn = document.getElementById('submitClusterBtn');
            submitBtn.disabled = true;
            submitBtn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Creating...';
            
            try {{
                const response = await fetch(`/api/projects/${{project.id}}/clusters/`, {{
                    method: 'POST',
                    headers: {{'Content-Type': 'application/json'}},
                    body: JSON.stringify({{ name, description, order }})
                }});
                
                const data = await response.json();
                
                if (response.ok) {{
                    document.getElementById('cluster-name').value = '';
                    document.getElementById('cluster-description').value = '';
                    document.getElementById('cluster-order').value = '0';
                    document.getElementById('cluster-alert-container').innerHTML = '';
                    submitBtn.disabled = false;
                    submitBtn.innerHTML = '<i class="fas fa-plus me-2"></i>Create Cluster';
                    modal.hide();
                    loadClusters();
                    alert('Cluster created successfully!');
                }} else {{
                    throw new Error(data.error || 'Failed to create cluster');
                }}
            }} catch (error) {{
                document.getElementById('cluster-alert-container').innerHTML = `
                    <div class="alert alert-danger">
                        <i class="fas fa-exclamation-circle me-2"></i>
                        ${{error.message}}
                    </div>
                `;
                submitBtn.disabled = false;
                submitBtn.innerHTML = '<i class="fas fa-plus me-2"></i>Create Cluster';
            }}
        }};
    }}
    
    async function editCluster(clusterId) {{
        try {{
            const response = await fetch(`/api/clusters/${{clusterId}}/`);
            const data = await response.json();
            const cluster = data.cluster || data;
            if (!cluster) {{
                throw new Error('Cluster not found');
            }}
            const cName = (cluster.name != null && cluster.name !== '') ? String(cluster.name) : '';
            const cDesc = (cluster.description != null) ? String(cluster.description) : '';
            const cOrder = (cluster.order != null && cluster.order !== '') ? cluster.order : 0;
            
            const modalHtml = `
                <div class="modal fade" id="editClusterModal" tabindex="-1">
                    <div class="modal-dialog">
                        <div class="modal-content">
                            <div class="modal-header">
                                <h5 class="modal-title">Edit Cluster</h5>
                                <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                            </div>
                            <div class="modal-body">
                                <div id="edit-cluster-alert-container"></div>
                                
                                <form id="editClusterForm">
                                    <div class="mb-3">
                                        <label for="edit-cluster-name" class="form-label">Cluster Name *</label>
                                        <input type="text" class="form-control" id="edit-cluster-name" value="${{cName.replace(/"/g, '&quot;')}}" required>
                                    </div>
                                    
                                    <div class="mb-3">
                                        <label for="edit-cluster-description" class="form-label">Description</label>
                                        <textarea class="form-control" id="edit-cluster-description" rows="3">${{cDesc.replace(/</g, '&lt;').replace(/>/g, '&gt;')}}</textarea>
                                    </div>
                                    
                                    <div class="mb-3">
                                        <label for="edit-cluster-order" class="form-label">Display Order</label>
                                        <input type="number" class="form-control" id="edit-cluster-order" value="${{cOrder}}">
                                    </div>
                                </form>
                            </div>
                            <div class="modal-footer">
                                <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                                <button type="button" class="btn btn-primary" id="updateClusterBtn">
                                    <i class="fas fa-save me-2"></i>Save Changes
                                </button>
                            </div>
                        </div>
                    </div>
                </div>
            `;
            
            if (document.getElementById('editClusterModal')) {{
                document.getElementById('editClusterModal').remove();
            }}
            document.body.insertAdjacentHTML('beforeend', modalHtml);
            
            const modal = new bootstrap.Modal(document.getElementById('editClusterModal'));
            modal.show();
            
            document.getElementById('updateClusterBtn').onclick = async () => {{
                const name = document.getElementById('edit-cluster-name').value.trim();
                const description = document.getElementById('edit-cluster-description').value.trim();
                const order = parseInt(document.getElementById('edit-cluster-order').value) || 0;
                
                if (!name) {{
                    document.getElementById('edit-cluster-alert-container').innerHTML = `
                        <div class="alert alert-danger">
                            <i class="fas fa-exclamation-circle me-2"></i>
                            Cluster name is required
                        </div>
                    `;
                    return;
                }}
                
                const submitBtn = document.getElementById('updateClusterBtn');
                submitBtn.disabled = true;
                submitBtn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Saving...';
                
                try {{
                    const response = await fetch(`/api/clusters/${{clusterId}}/`, {{
                        method: 'PATCH',
                        headers: {{'Content-Type': 'application/json'}},
                        body: JSON.stringify({{ name, description, order }})
                    }});
                    
                    const data = await response.json();
                    
                    if (response.ok) {{
                        modal.hide();
                        loadClusters();
                        alert('Cluster updated successfully!');
                    }} else {{
                        throw new Error(data.error || 'Failed to update cluster');
                    }}
                }} catch (error) {{
                    document.getElementById('edit-cluster-alert-container').innerHTML = `
                        <div class="alert alert-danger">
                            <i class="fas fa-exclamation-circle me-2"></i>
                            ${{error.message}}
                        </div>
                    `;
                    submitBtn.disabled = false;
                    submitBtn.innerHTML = '<i class="fas fa-save me-2"></i>Save Changes';
                }}
            }};
        }} catch (error) {{
            alert('Error loading cluster: ' + error.message);
        }}
    }}
    
    async function deleteCluster(clusterId, clusterName) {{
        if (!confirm(`Are you sure you want to delete the cluster "${{clusterName}}"? This will unassign all roles from this cluster.`)) {{
            return;
        }}
        
        try {{
            const response = await fetch(`/api/clusters/${{clusterId}}/`, {{
                method: 'DELETE'
            }});
            
            const data = await response.json();
            
            if (response.ok) {{
                loadClusters();
                alert('Cluster deleted successfully!');
            }} else {{
                throw new Error(data.error || 'Failed to delete cluster');
            }}
        }} catch (error) {{
            alert('Error deleting cluster: ' + error.message);
        }}
    }}
    
    // Load project on page load
    loadProject();
    </script>
    """
    
    return render_page(f"Layer: {project_slug} - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/projects/<project_slug>/')
def project_detail(project_slug):
    """Project detail page"""
    return _render_project_detail(project_slug)

@app.route('/projects/<project_slug>/waitlist/<int:waitlist_id>/')
def project_detail_waitlist(project_slug, waitlist_id):
    """Project detail with specific waitlist tab (for referral links)"""
    return _render_project_detail(project_slug, waitlist_id=waitlist_id)

@app.route('/projects/create/')
@require_auth
def create_project_page():
    """Create project form page"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = """
    <div class="container mt-4">
        <div class="row">
            <div class="col-md-8 offset-md-2">
                <h1 class="mb-4">Create New Layer</h1>
                
                <div id="alert-container"></div>
                
                <form id="createProjectForm">
                    <div class="mb-3">
                        <label for="name" class="form-label">Layer Name *</label>
                        <input type="text" class="form-control" id="name" required>
                        <div class="form-text">A clear, descriptive name for your layer</div>
                    </div>
                    
                    <div class="mb-3">
                        <label for="mission" class="form-label">Mission</label>
                        <textarea class="form-control" id="mission" rows="3" style="white-space: pre-wrap;"></textarea>
                        <div class="form-text">Optional: The layer's core purpose and values (line breaks preserved)</div>
                    </div>
                    
                    <div class="mb-3">
                        <label for="description" class="form-label">Description *</label>
                        <textarea class="form-control" id="description" rows="4" required style="white-space: pre-wrap;"></textarea>
                        <div class="form-text">Explain what this layer is about and its goals (line breaks preserved)</div>
                    </div>
                    
                    <div class="mb-3">
                        <label for="repo_url" class="form-label">Repository URL</label>
                        <input type="url" class="form-control" id="repo_url" placeholder="https://github.com/...">
                    </div>
                    
                    <div class="mb-3">
                        <label for="website_url" class="form-label">Website URL</label>
                        <input type="url" class="form-control" id="website_url" placeholder="https://...">
                    </div>
                    
                    <div class="alert alert-info">
                        <i class="fas fa-info-circle me-2"></i>
                        <strong>Note:</strong> New layers start with "proposed" status and require admin approval before becoming active.
                        You will be the layer owner; you can add more admins after creation via <strong>Edit</strong> on the layer page.
                    </div>
                    
                    <div class="d-flex gap-2">
                        <button type="submit" class="btn btn-primary" id="submitBtn">
                            <i class="fas fa-plus me-2"></i>Create Layer
                        </button>
                        <a href="/projects/" class="btn btn-secondary">Cancel</a>
                    </div>
                </form>
            </div>
        </div>
    </div>
    
    <script>
    document.getElementById('createProjectForm').addEventListener('submit', async (e) => {
        e.preventDefault();
        
        const submitBtn = document.getElementById('submitBtn');
        submitBtn.disabled = true;
        submitBtn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Creating...';
        
        const formData = {
            name: document.getElementById('name').value,
            mission: document.getElementById('mission').value,
            description: document.getElementById('description').value,
            repo_url: document.getElementById('repo_url').value,
            website_url: document.getElementById('website_url').value
        };
        
        try {
            const response = await fetch('/api/projects/', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify(formData)
            });
            
            const data = await response.json();
            
            if (response.ok) {
                document.getElementById('alert-container').innerHTML = `
                    <div class="alert alert-success">
                        <i class="fas fa-check-circle me-2"></i>
                        Project created successfully! Redirecting...
                    </div>
                `;
                setTimeout(() => {
                    window.location.href = `/projects/${data.project.slug}/`;
                }, 1500);
            } else {
                throw new Error(data.error || 'Failed to create project');
            }
        } catch (error) {
            document.getElementById('alert-container').innerHTML = `
                <div class="alert alert-danger">
                    <i class="fas fa-exclamation-circle me-2"></i>
                    ${error.message}
                </div>
            `;
            submitBtn.disabled = false;
            submitBtn.innerHTML = '<i class="fas fa-plus me-2"></i>Create Layer';
        }
    });
    </script>
    """
    
    return render_page("Create Layer - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/guilds/<guild_slug>/')
def guild_detail(guild_slug):
    """Guild detail page with members"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = f"""
    <div class="container mt-4">
        <div id="guild-header" class="mb-4">
            <div class="d-flex justify-content-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        </div>
        
        <div class="row">
            <div class="col-md-8">
                <div class="card mb-4">
                    <div class="card-header"><h5>About</h5></div>
                    <div class="card-body" id="guild-about">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
                
                <div class="card">
                    <div class="card-header"><h5>Members</h5></div>
                    <div class="card-body" id="guild-members">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
            </div>
            
            <div class="col-md-4">
                <div class="card mb-4">
                    <div class="card-header"><h5>Quick Actions</h5></div>
                    <div class="card-body" id="guild-actions">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
                
                <div class="card">
                    <div class="card-header"><h5>Statistics</h5></div>
                    <div class="card-body" id="guild-stats">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    let guild = null;
    const guildSlug = '{guild_slug}';
    const isAuthenticated = {'true' if current_user else 'false'};
    const currentUserId = {current_user['id'] if current_user else 'null'};
    
    async function loadGuild() {{
        try {{
            // Find guild by slug
            const response = await fetch('/api/guilds/');
            const data = await response.json();
            guild = data.guilds.find(g => g.slug === guildSlug);
            
            if (!guild) {{
                document.getElementById('guild-header').innerHTML = '<div class="alert alert-danger">Guild not found</div>';
                return;
            }}
            
            // Load full guild details with members
            const detailResponse = await fetch(`/api/guilds/${{guild.id}}/`);
            guild = await detailResponse.json();
            
            displayGuildHeader();
            displayGuildAbout();
            displayGuildMembers();
            displayGuildActions();
            displayGuildStats();
        }} catch (error) {{
            console.error('Error loading guild:', error);
            document.getElementById('guild-header').innerHTML = '<div class="alert alert-danger">Error loading guild</div>';
        }}
    }}
    
    function displayGuildHeader() {{
        const statusBadge = guild.status === 'active' 
            ? '<span class="badge bg-success">Active</span>' 
            : '<span class="badge bg-secondary">Archived</span>';
        
        const isInitiator = isAuthenticated && guild.initiator_id === currentUserId;
        
        document.getElementById('guild-header').innerHTML =
            '<div class="row">' +
                '<div class="col-md-8">' +
                    '<h1>' + (guild.name || '') + '</h1>' +
                    '<div class="mb-3">' + statusBadge + '</div>' +
                '</div>' +
                '<div class="col-md-4 text-end">' +
                    '${{isInitiator ? \'<button class="btn btn-secondary me-2" onclick="editGuild()"><i class="fas fa-edit me-2"></i>Edit</button>\' : \'\'}}' +
                    '<a href="/guilds/" class="btn btn-outline-secondary"><i class="fas fa-arrow-left me-2"></i>Back</a>' +
                '</div>' +
            '</div>';
    }}
    
    function displayGuildAbout() {{
        document.getElementById('guild-about').innerHTML = `
            <p>${{guild.description || 'No description provided'}}</p>
            <hr>
            <p><strong>Created:</strong> ${{new Date(guild.created_at).toLocaleDateString()}}</p>
            <p><strong>Last Updated:</strong> ${{guild.updated_at ? new Date(guild.updated_at).toLocaleDateString() : 'Never'}}</p>
        `;
    }}
    
    function displayGuildMembers() {{
        if (!guild.members || guild.members.length === 0) {{
            document.getElementById('guild-members').innerHTML = '<p class="text-muted">No members yet</p>';
            return;
        }}
        
        let html = '<div class="list-group">';
        guild.members.forEach(member => {{
            const roleClass = member.role === 'initiator' ? 'primary' : member.role === 'admin' ? 'success' : 'secondary';
            html += `
                <div class="list-group-item d-flex justify-content-between align-items-center">
                    <div>
                        <strong>${{member.username}}</strong>
                        ${{member.name ? `<br><small class="text-muted">${{member.name}}</small>` : ''}}
                    </div>
                    <span class="badge bg-${{roleClass}}">${{member.role}}</span>
                </div>
            `;
        }});
        html += '</div>';
        
        document.getElementById('guild-members').innerHTML = html;
    }}
    
    function displayGuildActions() {{
        const userMembership = guild.members ? guild.members.find(m => m.user_id === currentUserId) : null;
        const isAdmin = userMembership && (userMembership.role === 'initiator' || userMembership.role === 'admin');
        
        let html = '';
        
        // Add image at the top if available
        if (guild.image_url) {{
            html += '<div class="mb-3 text-center"><img src="' + guild.image_url + '" alt="' + (guild.name || '') + '" class="img-fluid rounded" style="max-height: 180px;"></div>';
        }}
        
        if (!isAuthenticated) {{
            html += '<a href="/login/" class="btn btn-primary w-100 mb-2"><i class="fas fa-sign-in-alt me-2"></i>Login to Join</a>';
        }} else if (!userMembership) {{
            html += '<p class="text-muted">Request an invitation from a guild admin to join</p>';
        }} else {{
            if (isAdmin) {{
                html += '<button class="btn btn-primary w-100 mb-2" onclick="inviteMember()"><i class="fas fa-user-plus me-2"></i>Invite Member</button>';
                html += '<button class="btn btn-secondary w-100 mb-2" onclick="manageGuild()"><i class="fas fa-cog me-2"></i>Manage Guild</button>';
            }}
            html += `<p class="text-muted mt-2">Your role: <strong>${{userMembership.role}}</strong></p>`;
        }}
        
        document.getElementById('guild-actions').innerHTML = html;
    }}
    
    function displayGuildStats() {{
        const memberCount = guild.members ? guild.members.length : 0;
        const adminCount = guild.members ? guild.members.filter(m => m.role === 'admin' || m.role === 'initiator').length : 0;
        
        document.getElementById('guild-stats').innerHTML = `
            <p><strong>Total Members:</strong> ${{memberCount}}</p>
            <p><strong>Admins:</strong> ${{adminCount}}</p>
            <p><strong>Status:</strong> ${{guild.status}}</p>
        `;
    }}
    
    function inviteMember() {{
        const email = prompt('Enter email address to invite:');
        if (!email) return;
        
        fetch(`/api/guilds/${{guild.id}}/invite/`, {{
            method: 'POST',
            headers: {{'Content-Type': 'application/json'}},
            body: JSON.stringify({{email: email}})
        }})
        .then(response => response.json())
        .then(data => {{
            if (data.success) {{
                alert(`Invitation sent! Link: ${{data.invitation_link}}`);
            }} else {{
                alert('Error: ' + (data.error || 'Failed to send invitation'));
            }}
        }})
        .catch(error => {{
            console.error('Error:', error);
            alert('Error sending invitation');
        }});
    }}
    
    function manageGuild() {{
        alert('Guild management functionality coming soon');
    }}
    
    // Load guild on page load
    async function uploadGuildImage() {{
        const fileInput = document.getElementById('edit-guild-image-file');
        const statusEl = document.getElementById('edit-guild-image-upload-status');
        const urlInput = document.getElementById('edit-guild-image-url');
        
        if (!fileInput.files || !fileInput.files[0]) {{
            statusEl.innerHTML = '<small class="text-danger">Please select a file first</small>';
            return;
        }}
        
        const formData = new FormData();
        formData.append('file', fileInput.files[0]);
        formData.append('entity_type', 'guild');
        
        statusEl.innerHTML = '<small class="text-info"><i class="fas fa-spinner fa-spin"></i> Uploading...</small>';
        
        try {{
            const response = await fetch('/api/upload/entity-image', {{
                method: 'POST',
                credentials: 'include',
                body: formData
            }});
            
            const data = await response.json();
            
            if (response.ok && data.image_url) {{
                urlInput.value = data.image_url;
                statusEl.innerHTML = '<small class="text-success"><i class="fas fa-check"></i> Uploaded successfully</small>';
                fileInput.value = '';
            }} else {{
                statusEl.innerHTML = `<small class="text-danger">${{data.error || 'Upload failed'}}</small>`;
            }}
        }} catch (error) {{
            console.error('Upload error:', error);
            statusEl.innerHTML = '<small class="text-danger">Upload failed. Please try again.</small>';
        }}
    }}
    
    function editGuild() {{
        const modalHtml = `
            <div class="modal fade" id="editGuildModal" tabindex="-1">
                <div class="modal-dialog modal-lg">
                    <div class="modal-content">
                        <div class="modal-header">
                            <h5 class="modal-title">Edit Guild</h5>
                            <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                        </div>
                        <div class="modal-body">
                            <div id="edit-guild-alert-container"></div>
                            
                            <form id="editGuildForm">
                                <div class="mb-3">
                                    <label for="edit-guild-name" class="form-label">Guild Name *</label>
                                    <input type="text" class="form-control" id="edit-guild-name" value="${{guild.name}}" required>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="edit-guild-description" class="form-label">Description *</label>
                                    <textarea class="form-control" id="edit-guild-description" rows="4" required>${{guild.description}}</textarea>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="edit-guild-image-url" class="form-label">Image (optional)</label>
                                    <input type="url" class="form-control mb-2" id="edit-guild-image-url" value="${{guild.image_url || ''}}" placeholder="https://example.com/image.png or upload below">
                                    <div class="input-group">
                                        <input type="file" class="form-control" id="edit-guild-image-file" accept="image/*">
                                        <button class="btn btn-outline-primary" type="button" onclick="uploadGuildImage()">
                                            <i class="fas fa-upload"></i> Upload
                                        </button>
                                    </div>
                                    <div class="form-text">Guild logo or banner. Max 600×600px, 5MB. Upload or paste URL above.</div>
                                    <div id="edit-guild-image-upload-status" class="mt-1"></div>
                                </div>
                                
                                <div class="mb-3">
                                    <label for="edit-guild-status" class="form-label">Status</label>
                                    <select class="form-select" id="edit-guild-status">
                                        <option value="active" ${{guild.status === 'active' ? 'selected' : ''}}>Active</option>
                                        <option value="archived" ${{guild.status === 'archived' ? 'selected' : ''}}>Archived</option>
                                    </select>
                                </div>
                            </form>
                        </div>
                        <div class="modal-footer">
                            <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                            <button type="button" class="btn btn-primary" id="updateGuildBtn">
                                <i class="fas fa-save me-2"></i>Save Changes
                            </button>
                        </div>
                    </div>
                </div>
            </div>
        `;
        
        if (document.getElementById('editGuildModal')) {{
            document.getElementById('editGuildModal').remove();
        }}
        document.body.insertAdjacentHTML('beforeend', modalHtml);
        
        const modal = new bootstrap.Modal(document.getElementById('editGuildModal'));
        modal.show();
        
        document.getElementById('updateGuildBtn').onclick = async () => {{
            const name = document.getElementById('edit-guild-name').value.trim();
            const description = document.getElementById('edit-guild-description').value.trim();
            const image_url = document.getElementById('edit-guild-image-url') ? document.getElementById('edit-guild-image-url').value.trim() : '';
            const status = document.getElementById('edit-guild-status').value;
            
            if (!name || !description) {{
                document.getElementById('edit-guild-alert-container').innerHTML = `
                    <div class="alert alert-danger">
                        <i class="fas fa-exclamation-circle me-2"></i>
                        Name and description are required
                    </div>
                `;
                return;
            }}
            
            const submitBtn = document.getElementById('updateGuildBtn');
            submitBtn.disabled = true;
            submitBtn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Saving...';
            
            try {{
                const response = await fetch(`/api/guilds/${{guild.id}}/`, {{
                    method: 'PATCH',
                    headers: {{'Content-Type': 'application/json'}},
                    body: JSON.stringify({{ name, description, image_url: image_url || null, status }})
                }});
                
                const data = await response.json();
                
                if (response.ok) {{
                    modal.hide();
                    loadGuild(); // Reload guild
                    alert('Guild updated successfully!');
                }} else {{
                    throw new Error(data.error || 'Failed to update guild');
                }}
            }} catch (error) {{
                document.getElementById('edit-guild-alert-container').innerHTML = `
                    <div class="alert alert-danger">
                        <i class="fas fa-exclamation-circle me-2"></i>
                        ${{error.message}}
                    </div>
                `;
                submitBtn.disabled = false;
                submitBtn.innerHTML = '<i class="fas fa-save me-2"></i>Save Changes';
            }}
        }};
    }}
    
    loadGuild();
    </script>
    """
    
    return render_page(f"Guild: {guild_slug} - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/guilds/create/')
@require_auth
def create_guild_page():
    """Create guild form page"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = """
    <div class="container mt-4">
        <div class="row">
            <div class="col-md-8 offset-md-2">
                <h1 class="mb-4">Create New Guild</h1>
                
                <div id="alert-container"></div>
                
                <form id="createGuildForm">
                    <div class="mb-3">
                        <label for="name" class="form-label">Guild Name *</label>
                        <input type="text" class="form-control" id="name" required>
                        <div class="form-text">A clear, descriptive name for your guild</div>
                    </div>
                    
                    <div class="mb-3">
                        <label for="description" class="form-label">Description *</label>
                        <textarea class="form-control" id="description" rows="4" required></textarea>
                        <div class="form-text">Explain what this guild is about and its purpose</div>
                    </div>
                    
                    <div class="alert alert-info">
                        <i class="fas fa-info-circle me-2"></i>
                        <strong>Note:</strong> Guilds are instantly created with no approval required. You will automatically become the guild initiator and admin.
                    </div>
                    
                    <div class="d-flex gap-2">
                        <button type="submit" class="btn btn-primary" id="submitBtn">
                            <i class="fas fa-plus me-2"></i>Create Guild
                        </button>
                        <a href="/guilds/" class="btn btn-secondary">Cancel</a>
                    </div>
                </form>
            </div>
        </div>
    </div>
    
    <script>
    document.getElementById('createGuildForm').addEventListener('submit', async (e) => {
        e.preventDefault();
        
        const submitBtn = document.getElementById('submitBtn');
        submitBtn.disabled = true;
        submitBtn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Creating...';
        
        const formData = {
            name: document.getElementById('name').value,
            description: document.getElementById('description').value
        };
        
        try {
            const response = await fetch('/api/guilds/', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify(formData)
            });
            
            const data = await response.json();
            
            if (response.ok) {
                document.getElementById('alert-container').innerHTML = `
                    <div class="alert alert-success">
                        <i class="fas fa-check-circle me-2"></i>
                        Guild created successfully! Redirecting...
                    </div>
                `;
                setTimeout(() => {
                    window.location.href = `/guilds/${data.guild.slug}/`;
                }, 1500);
            } else {
                throw new Error(data.error || 'Failed to create guild');
            }
        } catch (error) {
            document.getElementById('alert-container').innerHTML = `
                <div class="alert alert-danger">
                    <i class="fas fa-exclamation-circle me-2"></i>
                    ${error.message}
                </div>
            `;
            submitBtn.disabled = false;
            submitBtn.innerHTML = '<i class="fas fa-plus me-2"></i>Create Guild';
        }
    });
    </script>
    """
    
    return render_page("Create Guild - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/workgroups/<workgroup_slug>/')
def workgroup_detail(workgroup_slug):
    """Workgroup detail page"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = f"""
    <div class="container mt-4">
        <div id="workgroup-header" class="mb-4">
            <div class="d-flex justify-content-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        </div>
        
        <div class="row">
            <div class="col-md-8">
                <div class="card mb-4">
                    <div class="card-header"><h5>About</h5></div>
                    <div class="card-body" id="workgroup-about">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
                
                <div class="card mb-4">
                    <div class="card-header">
                        <div class="d-flex justify-content-between align-items-center">
                            <h5 class="mb-0">Chairs / Coordinators</h5>
                            {'<button class="btn btn-sm btn-success" onclick="nominateForChair()" id="nominate-btn" style="display:none;"><i class="fas fa-star me-1"></i>Nominate for Chair</button>' if current_user else ''}
                        </div>
                    </div>
                    <div class="card-body" id="workgroup-chairs">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
                
                <div class="card">
                    <div class="card-header">
                        <div class="d-flex justify-content-between align-items-center">
                            <h5 class="mb-0">Members</h5>
                            {'<button class="btn btn-sm btn-primary" onclick="joinWorkgroup()" id="join-btn" style="display:none;"><i class="fas fa-user-plus me-1"></i>Join</button>' if current_user else ''}
                        </div>
                    </div>
                    <div class="card-body" id="workgroup-members">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
            </div>
            
            <div class="col-md-4">
                <div class="card mb-4">
                    <div class="card-header"><h5>Details</h5></div>
                    <div class="card-body" id="workgroup-details">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
                
                <div class="card">
                    <div class="card-header"><h5>Charter & Goals</h5></div>
                    <div class="card-body" id="workgroup-charter">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
            </div>
        </div>
    </div>
    
    <!-- Chair Nomination Modal -->
    <div class="modal fade" id="nominateChairModal" tabindex="-1">
        <div class="modal-dialog">
            <div class="modal-content">
                <div class="modal-header">
                    <h5 class="modal-title">Nominate for Chair/Coordinator</h5>
                    <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                </div>
                <div class="modal-body">
                    <form id="nominateChairForm">
                        <div class="mb-3">
                            <label class="form-label">Workgroup</label>
                            <p class="form-control-plaintext" id="modal-workgroup-name"></p>
                        </div>
                        <div class="mb-3">
                            <label for="nomination-statement" class="form-label">Statement <span class="text-danger">*</span></label>
                            <textarea 
                                class="form-control" 
                                id="nomination-statement" 
                                rows="4" 
                                required
                                placeholder="Explain why you would be a good chair/coordinator for this workgroup..."
                            ></textarea>
                            <div class="form-text">Share your relevant experience, vision, and commitment to leading this workgroup.</div>
                        </div>
                        <div class="alert alert-info">
                            <i class="fas fa-info-circle me-2"></i>
                            Your nomination will be reviewed by layer administrators and workgroup coordinators.
                        </div>
                    </form>
                </div>
                <div class="modal-footer">
                    <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                    <button type="button" class="btn btn-success" onclick="submitChairNomination()">
                        <i class="fas fa-paper-plane me-2"></i>Submit Nomination
                    </button>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    let workgroup = null;
    let project = null;
    const workgroupSlug = '{workgroup_slug}';
    const isAuthenticated = {'true' if current_user else 'false'};
    
    async function loadWorkgroup() {{
        try {{
            // Load all projects first to find the workgroup
            const projectsResp = await fetch('/api/projects/');
            const projectsData = await projectsResp.json();
            
            // Search for workgroup across all projects
            for (const proj of projectsData.projects) {{
                const wgResp = await fetch(`/api/projects/${{proj.id}}/workgroups/`);
                const wgData = await wgResp.json();
                const found = wgData.workgroups.find(wg => wg.slug === workgroupSlug);
                
                if (found) {{
                    workgroup = found;
                    project = proj;
                    break;
                }}
            }}
            
            if (!workgroup) {{
                document.getElementById('workgroup-header').innerHTML = '<div class="alert alert-danger">Workgroup not found</div>';
                return;
            }}
            
            // Load full workgroup details
            const detailResp = await fetch(`/api/workgroups/${{workgroup.id}}/`);
            workgroup = await detailResp.json();
            
            displayWorkgroupHeader();
            displayWorkgroupAbout();
            displayWorkgroupCharter();
            displayWorkgroupDetails();
            loadChairs();
            loadMembers();
        }} catch (error) {{
            console.error('Error loading workgroup:', error);
            document.getElementById('workgroup-header').innerHTML = '<div class="alert alert-danger">Error loading workgroup</div>';
        }}
    }}
    
    function displayWorkgroupHeader() {{
        const statusBadge = getStatusBadge(workgroup.status);
        const approvalBadge = getApprovalBadge(workgroup.approval_status);
        
        // Use project data if available, otherwise use workgroup's project_name
        const projectSlug = project ? project.slug : '';
        const projectName = project ? project.name : (workgroup.project_name || 'Layer');
        
        document.getElementById('workgroup-header').innerHTML = `
            <div class="row">
                <div class="col-md-8">
                    <nav aria-label="breadcrumb">
                        <ol class="breadcrumb">
                            <li class="breadcrumb-item"><a href="/projects/">Layers</a></li>
                            ${{projectSlug ? `<li class="breadcrumb-item"><a href="/projects/${{projectSlug}}/">${{projectName}}</a></li>` : `<li class="breadcrumb-item">${{projectName}}</li>`}}
                            <li class="breadcrumb-item active">${{workgroup.name}}</li>
                        </ol>
                    </nav>
                    <h1>${{workgroup.name}}</h1>
                    <div class="mb-3">
                        ${{statusBadge}}
                        ${{approvalBadge}}
                    </div>
                </div>
                <div class="col-md-4">
                    ${{workgroup.image_url ? `<div class="card mb-3"><div class="card-body p-2 text-center"><img src="${{workgroup.image_url}}" alt="${{workgroup.name}}" class="img-fluid rounded" style="max-height: 200px;"></div></div>` : ''}}
                    <div class="text-end">
                        ${{workgroup.can_edit ? '<button type="button" class="btn btn-outline-secondary me-2" onclick="editWorkgroup()"><i class="fas fa-edit me-2"></i>Edit Workgroup</button>' : ''}}
                        ${{projectSlug ? `<a href="/projects/${{projectSlug}}/" class="btn btn-outline-secondary"><i class="fas fa-arrow-left me-2"></i>Back to Layer</a>` : '<a href="/workgroups/" class="btn btn-outline-secondary"><i class="fas fa-arrow-left me-2"></i>Back to Workgroups</a>'}}
                    </div>
                </div>
            </div>
        `;
    }}
    
    function displayWorkgroupAbout() {{
        document.getElementById('workgroup-about').innerHTML = `
            <p>${{workgroup.description || 'No description provided'}}</p>
        `;
    }}
    
    function displayWorkgroupCharter() {{
        let html = '';
        if (workgroup.charter) {{
            html = `<p>${{workgroup.charter}}</p>`;
        }} else {{
            html = '<p class="text-muted">No charter defined yet</p>';
        }}
        
        if (workgroup.goals) {{
            html += '<h6 class="mt-3">Goals</h6>';
            html += `<p>${{workgroup.goals}}</p>`;
        }}
        
        document.getElementById('workgroup-charter').innerHTML = html;
    }}
    
    function displayWorkgroupDetails() {{
        const projectSlug = project ? project.slug : '';
        const projectName = project ? project.name : (workgroup.project_name || 'Unknown Project');
        
        document.getElementById('workgroup-details').innerHTML = `
            <p><strong>Layer:</strong> ${{projectSlug ? `<a href="/projects/${{projectSlug}}/">${{projectName}}</a>` : projectName}}</p>
            <p><strong>Status:</strong> ${{workgroup.status}}</p>
            <p><strong>Approval:</strong> ${{workgroup.approval_status}}</p>
            <p><strong>Created:</strong> ${{new Date(workgroup.created_at).toLocaleDateString()}}</p>
            ${{workgroup.coordinator_name ? `<p><strong>Coordinator:</strong> ${{workgroup.coordinator_name}}</p>` : ''}}
        `;
    }}
    
    function getStatusBadge(status) {{
        const badges = {{
            'active': '<span class="badge bg-success">Active</span>',
            'inactive': '<span class="badge bg-warning">Inactive</span>',
            'completed': '<span class="badge bg-primary">Completed</span>',
            'archived': '<span class="badge bg-secondary">Archived</span>'
        }};
        return badges[status] || '';
    }}
    
    function getApprovalBadge(approval) {{
        const badges = {{
            'pending': '<span class="badge bg-warning">Pending Approval</span>',
            'approved': '<span class="badge bg-success">Approved</span>',
            'rejected': '<span class="badge bg-danger">Rejected</span>'
        }};
        return badges[approval] || '';
    }}
    
    async function loadChairs() {{
        try {{
            // Load chairs from working_group_chair table using acronym
            const response = await fetch(`/api/workgroups/${{workgroup.id}}/chairs/`);
            const data = await response.json();
            
            // Check if current user is already a chair
            const currentUserId = {current_user['id'] if current_user else 'null'};
            let isCurrentUserChair = false;
            if (isAuthenticated && data.chairs) {{
                isCurrentUserChair = data.chairs.some(c => c.user_id === currentUserId);
            }}
            
            // Show/hide nominate button (only if user is a member and not already a chair)
            const nominateBtn = document.getElementById('nominate-btn');
            if (nominateBtn) {{
                // We'll check membership status from the members list
                if (isAuthenticated && !isCurrentUserChair && workgroup.approval_status === 'approved') {{
                    nominateBtn.style.display = 'block';
                }} else {{
                    nominateBtn.style.display = 'none';
                }}
            }}
            
            let html = '';
            if (data.chairs && data.chairs.length > 0) {{
                html = '<div class="list-group">';
                data.chairs.forEach(chair => {{
                    html += `
                        <div class="list-group-item">
                            <div class="d-flex justify-content-between align-items-center">
                                <div>
                                    <strong>${{chair.chair_name}}</strong>
                                    ${{chair.approved ? '<span class="badge bg-success ms-2">Approved</span>' : '<span class="badge bg-warning ms-2">Pending</span>'}}
                                </div>
                            </div>
                        </div>
                    `;
                }});
                html += '</div>';
            }} else {{
                html = '<p class="text-muted">No chairs assigned yet</p>';
            }}
            
            document.getElementById('workgroup-chairs').innerHTML = html;
        }} catch (error) {{
            console.error('Error loading chairs:', error);
            document.getElementById('workgroup-chairs').innerHTML = '<p class="text-muted">No chairs assigned yet</p>';
        }}
    }}
    
    async function loadMembers() {{
        try {{
            // Load members from working_group_member table using acronym
            const response = await fetch(`/api/workgroups/${{workgroup.id}}/members/`);
            const data = await response.json();
            
            // Check if current user is already a member
            const currentUserId = {current_user['id'] if current_user else 'null'};
            let isCurrentUserMember = false;
            if (isAuthenticated && data.members) {{
                isCurrentUserMember = data.members.some(m => m.user_id === currentUserId);
            }}
            
            // Show/hide join button
            const joinBtn = document.getElementById('join-btn');
            if (joinBtn) {{
                if (isAuthenticated && !isCurrentUserMember && workgroup.approval_status === 'approved') {{
                    joinBtn.style.display = 'block';
                }} else {{
                    joinBtn.style.display = 'none';
                }}
            }}
            
            let html = '';
            if (data.members && data.members.length > 0) {{
                html = `<p class="text-muted mb-2">${{data.members.length}} member(s)</p>`;
                html += '<div class="list-group">';
                data.members.forEach(member => {{
                    html += `
                        <div class="list-group-item">
                            <strong>${{member.user_name}}</strong>
                            <small class="text-muted d-block">Joined: ${{new Date(member.joined_at).toLocaleDateString()}}</small>
                        </div>
                    `;
                }});
                html += '</div>';
            }} else {{
                html = '<p class="text-muted">No members yet</p>';
            }}
            
            document.getElementById('workgroup-members').innerHTML = html;
        }} catch (error) {{
            console.error('Error loading members:', error);
            document.getElementById('workgroup-members').innerHTML = '<p class="text-muted">No members yet</p>';
        }}
    }}
    
    async function joinWorkgroup() {{
        if (!isAuthenticated) {{
            alert('Please sign in to join this workgroup');
            return;
        }}
        
        if (!confirm('Join this workgroup?')) {{
            return;
        }}
        
        try {{
            const response = await fetch(`/api/workgroups/${{workgroup.id}}/join/`, {{
                method: 'POST',
                headers: {{ 'Content-Type': 'application/json' }}
            }});
            
            const data = await response.json();
            
            if (response.ok) {{
                alert('Successfully joined workgroup!');
                loadMembers(); // Reload members list
            }} else {{
                alert(data.error || 'Failed to join workgroup');
            }}
        }} catch (error) {{
            console.error('Error joining workgroup:', error);
            alert('Failed to join workgroup');
        }}
    }}
    
    function nominateForChair() {{
        if (!isAuthenticated) {{
            alert('Please sign in to nominate yourself for chair');
            return;
        }}
        
        // Populate modal
        document.getElementById('modal-workgroup-name').textContent = workgroup.name;
        document.getElementById('nomination-statement').value = '';
        
        // Show modal
        const modal = new bootstrap.Modal(document.getElementById('nominateChairModal'));
        modal.show();
    }}
    
    async function submitChairNomination() {{
        const statement = document.getElementById('nomination-statement').value.trim();
        
        if (!statement) {{
            alert('Please provide a statement for your nomination');
            return;
        }}
        
        try {{
            const response = await fetch(`/api/workgroups/${{workgroup.id}}/nominate-chair/`, {{
                method: 'POST',
                headers: {{ 'Content-Type': 'application/json' }},
                body: JSON.stringify({{ statement: statement }})
            }});
            
            const data = await response.json();
            
            if (response.ok) {{
                // Close modal
                const modal = bootstrap.Modal.getInstance(document.getElementById('nominateChairModal'));
                modal.hide();
                
                alert('Chair nomination submitted! It will require approval.');
                loadChairs(); // Reload chairs list
            }} else {{
                alert(data.error || 'Failed to nominate for chair');
            }}
        }} catch (error) {{
            console.error('Error nominating for chair:', error);
            alert('Failed to nominate for chair');
        }}
    }}
    
    async function uploadWorkgroupImage() {{
        const fileInput = document.getElementById('edit-wg-image-file');
        const statusEl = document.getElementById('edit-wg-image-upload-status');
        const urlInput = document.getElementById('edit-wg-image-url');
        
        if (!fileInput.files || !fileInput.files[0]) {{
            statusEl.innerHTML = '<small class="text-danger">Please select a file first</small>';
            return;
        }}
        
        const formData = new FormData();
        formData.append('file', fileInput.files[0]);
        formData.append('entity_type', 'workgroup');
        
        statusEl.innerHTML = '<small class="text-info"><i class="fas fa-spinner fa-spin"></i> Uploading...</small>';
        
        try {{
            const response = await fetch('/api/upload/entity-image', {{
                method: 'POST',
                credentials: 'include',
                body: formData
            }});
            
            const data = await response.json();
            
            if (response.ok && data.image_url) {{
                urlInput.value = data.image_url;
                statusEl.innerHTML = '<small class="text-success"><i class="fas fa-check"></i> Uploaded successfully</small>';
                fileInput.value = '';
            }} else {{
                statusEl.innerHTML = `<small class="text-danger">${{data.error || 'Upload failed'}}</small>`;
            }}
        }} catch (error) {{
            console.error('Upload error:', error);
            statusEl.innerHTML = '<small class="text-danger">Upload failed. Please try again.</small>';
        }}
    }}
    
    function editWorkgroup() {{
        const modalHtml = `
            <div class="modal fade" id="editWorkgroupModal" tabindex="-1">
                <div class="modal-dialog">
                    <div class="modal-content">
                        <div class="modal-header">
                            <h5 class="modal-title">Edit Workgroup</h5>
                            <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                        </div>
                        <div class="modal-body">
                            <div id="edit-workgroup-alert-container"></div>
                            <form id="editWorkgroupForm">
                                <div class="mb-3">
                                    <label for="edit-wg-name" class="form-label">Name *</label>
                                    <input type="text" class="form-control" id="edit-wg-name" required>
                                </div>
                                <div class="mb-3">
                                    <label for="edit-wg-description" class="form-label">Description</label>
                                    <textarea class="form-control" id="edit-wg-description" rows="3"></textarea>
                                </div>
                                <div class="mb-3">
                                    <label for="edit-wg-image-url" class="form-label">Image (optional)</label>
                                    <input type="url" class="form-control mb-2" id="edit-wg-image-url" placeholder="https://example.com/image.png or upload below">
                                    <div class="input-group">
                                        <input type="file" class="form-control" id="edit-wg-image-file" accept="image/*">
                                        <button class="btn btn-outline-primary" type="button" onclick="uploadWorkgroupImage()">
                                            <i class="fas fa-upload"></i> Upload
                                        </button>
                                    </div>
                                    <div class="form-text">Workgroup logo or banner. Max 600×600px, 5MB. Upload or paste URL above.</div>
                                    <div id="edit-wg-image-upload-status" class="mt-1"></div>
                                </div>
                                <div class="mb-3">
                                    <label for="edit-wg-status" class="form-label">Status</label>
                                    <select class="form-select" id="edit-wg-status">
                                        <option value="active">Active</option>
                                        <option value="inactive">Inactive</option>
                                        <option value="completed">Completed</option>
                                        <option value="archived">Archived</option>
                                    </select>
                                </div>
                            </form>
                        </div>
                        <div class="modal-footer">
                            <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                            <button type="button" class="btn btn-primary" id="editWorkgroupSubmitBtn">
                                <i class="fas fa-save me-2"></i>Save Changes
                            </button>
                        </div>
                    </div>
                </div>
            </div>
        `;
        if (!document.getElementById('editWorkgroupModal')) {{
            document.body.insertAdjacentHTML('beforeend', modalHtml);
        }}
        document.getElementById('edit-wg-name').value = workgroup.name || '';
        document.getElementById('edit-wg-description').value = workgroup.description || '';
        const wgImgEl = document.getElementById('edit-wg-image-url');
        if (wgImgEl) wgImgEl.value = workgroup.image_url || '';
        const wgImgFileEl = document.getElementById('edit-wg-image-file');
        if (wgImgFileEl) wgImgFileEl.value = '';
        const wgImgStatusEl = document.getElementById('edit-wg-image-upload-status');
        if (wgImgStatusEl) wgImgStatusEl.innerHTML = '';
        document.getElementById('edit-wg-status').value = workgroup.status || 'active';
        document.getElementById('edit-workgroup-alert-container').innerHTML = '';
        const modal = new bootstrap.Modal(document.getElementById('editWorkgroupModal'));
        modal.show();
        document.getElementById('editWorkgroupSubmitBtn').onclick = async () => {{
            const name = document.getElementById('edit-wg-name').value.trim();
            const description = document.getElementById('edit-wg-description').value.trim();
            const image_url = document.getElementById('edit-wg-image-url') ? document.getElementById('edit-wg-image-url').value.trim() : '';
            const status = document.getElementById('edit-wg-status').value;
            if (!name) {{
                document.getElementById('edit-workgroup-alert-container').innerHTML = '<div class="alert alert-danger">Name is required.</div>';
                return;
            }}
            const btn = document.getElementById('editWorkgroupSubmitBtn');
            btn.disabled = true;
            btn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Saving...';
            try {{
                const response = await fetch(`/api/workgroups/${{workgroup.id}}/`, {{
                    method: 'PATCH',
                    headers: {{'Content-Type': 'application/json'}},
                    body: JSON.stringify({{ name, description, image_url: image_url || null, status }})
                }});
                if (!response.ok) {{
                    const data = await response.json();
                    throw new Error(data.error || 'Failed to update workgroup');
                }}
                const data = await response.json();
                workgroup.name = data.workgroup.name;
                workgroup.description = data.workgroup.description;
                workgroup.image_url = data.workgroup.image_url || null;
                workgroup.status = data.workgroup.status;
                modal.hide();
                displayWorkgroupHeader();
                displayWorkgroupAbout();
                displayWorkgroupDetails();
            }} catch (err) {{
                document.getElementById('edit-workgroup-alert-container').innerHTML = '<div class="alert alert-danger">' + (err.message || 'Failed to update workgroup') + '</div>';
            }}
            btn.disabled = false;
            btn.innerHTML = '<i class="fas fa-save me-2"></i>Save Changes';
        }};
    }}
    
    // Load workgroup on page load
    loadWorkgroup();
    </script>
    """
    
    return render_page(f"Workgroup: {workgroup_slug} - MLGH", content, theme=current_theme, user_menu=user_menu)

# ============================================================================
# User Profile Routes
# ============================================================================

@app.route('/profile/<username>/')
def user_profile(username):
    """User profile page"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    # Fetch user from database
    profile_user = User.query.filter_by(username=username).first()
    if not profile_user:
        # Try by handle
        profile_user = User.query.filter_by(handle=username).first()
    
    if not profile_user:
        return render_page("User Not Found - MLGH", f"""
            <div class="container mt-5">
                <div class="alert alert-danger">
                    <h4>User Not Found</h4>
                    <p>The user "{username}" does not exist.</p>
                    <a href="/" class="btn btn-primary">Back to Home</a>
                </div>
            </div>
        """, theme=current_theme, user_menu=user_menu)
    
    # Check if viewing own profile
    is_own_profile = current_user and current_user['id'] == profile_user.id
    
    # Parse social links (stored as JSON string)
    import json
    social_links = []
    if profile_user.social_links:
        try:
            social_links = json.loads(profile_user.social_links)
        except:
            social_links = []
    
    # Get user stats
    from sqlalchemy import text
    
    # Count projects initiated
    projects_count = db.session.execute(text("""
        SELECT COUNT(*) FROM project WHERE initiator_id = :user_id
    """), {'user_id': profile_user.id}).scalar() or 0
    
    # Count workgroups coordinated
    workgroups_count = db.session.execute(text("""
        SELECT COUNT(*) FROM working_group WHERE coordinator_id = :user_id
    """), {'user_id': profile_user.id}).scalar() or 0
    
    # Count workgroup memberships
    memberships_count = db.session.execute(text("""
        SELECT COUNT(*) FROM working_group_member WHERE user_id = :user_id
    """), {'user_id': profile_user.id}).scalar() or 0
    
    # Count chair positions
    chair_count = db.session.execute(text("""
        SELECT COUNT(*) FROM working_group_chair WHERE user_id = :user_id AND approved = 1
    """), {'user_id': profile_user.id}).scalar() or 0
    
    # Count submissions (submission table uses submitted_by string, not submitted_by_id)
    name_variants = [x for x in (profile_user.name, profile_user.displayName, profile_user.oauthName, profile_user.username) if x]
    submissions_count = Submission.query.filter(Submission.submitted_by.in_(name_variants)).count() if name_variants else 0
    
    # Count comments
    comments_count = Comment.query.filter(Comment.author.in_(name_variants)).count() if name_variants else 0
    
    # Get recent activity (simplified for now)
    recent_projects = db.session.execute(text("""
        SELECT name, slug, created_at FROM project 
        WHERE initiator_id = :user_id 
        ORDER BY created_at DESC LIMIT 5
    """), {'user_id': profile_user.id}).fetchall()
    
    # Recent submissions (submission has submitted_by string and submitted_at, not created_at)
    if name_variants:
        recent_submissions_q = Submission.query.filter(
            Submission.submitted_by.in_(name_variants)
        ).order_by(Submission.submitted_at.desc()).limit(5).all()
        recent_submissions = [(s.draft_name or f"Draft {s.id}", s.submitted_at, s.id, s.status == 'approved') for s in recent_submissions_q]
        all_submissions_q = Submission.query.filter(
            Submission.submitted_by.in_(name_variants)
        ).order_by(Submission.submitted_at.desc()).all()
    else:
        recent_submissions = []
        all_submissions_q = []
    
    # Get coordinated workgroups
    coordinated_workgroups = Workgroup.query.filter_by(coordinator_id=profile_user.id).order_by(Workgroup.created_at.desc()).all()
    
    # Get memberships
    memberships_q = db.session.execute(text("""
        SELECT wg.name, wg.slug, wgm.joined_at 
        FROM working_group_member wgm
        JOIN working_group wg ON wgm.group_acronym = wg.acronym
        WHERE wgm.user_id = :user_id
        ORDER BY wgm.joined_at DESC
    """), {'user_id': profile_user.id}).fetchall()
    
    # Get chair positions
    chairs_q = db.session.execute(text("""
        SELECT wg.name, wg.slug, wgc.set_at, wgc.approved
        FROM working_group_chair wgc
        JOIN working_group wg ON wgc.group_acronym = wg.acronym
        WHERE wgc.user_id = :user_id
        ORDER BY wgc.set_at DESC
    """), {'user_id': profile_user.id}).fetchall()
    
    # Get project memberships (projects user has joined)
    project_memberships = ProjectMember.query.filter_by(user_id=profile_user.id, status='active').order_by(ProjectMember.joined_at.desc()).all()
    
    # Get referral stats if viewing own profile
    referral_code = None
    referral_count = 0
    if current_user and current_user['id'] == profile_user.id:
        referral_code = get_or_create_referral_code(profile_user)
        referral_count = ProjectMember.query.filter_by(referred_by_id=profile_user.id).count()
    
    # Get project memberships
    project_memberships = ProjectMember.query.filter_by(user_id=profile_user.id, status='active').order_by(ProjectMember.joined_at.desc()).all()
    
    # Get referral stats if viewing own profile
    referral_code = None
    referral_count = 0
    if is_own_profile:
        referral_code = get_or_create_referral_code(profile_user)
        referral_count = ProjectMember.query.filter_by(referred_by_id=profile_user.id).count()
    
    content = f"""
    <style>
        .profile-banner {{
            width: 100%;
            height: 300px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            background-size: cover;
            background-position: center;
            position: relative;
        }}
        
        .profile-header {{
            position: relative;
            margin-top: -80px;
            padding: 0 2rem;
        }}
        
        .profile-avatar {{
            width: 160px;
            height: 160px;
            border-radius: 50%;
            border: 6px solid var(--bg-primary);
            background: var(--bg-secondary);
            object-fit: cover;
            display: block;
            image-rendering: pixelated;
            image-rendering: -moz-crisp-edges;
            image-rendering: crisp-edges;
        }}
        
        .profile-stats {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(120px, 1fr));
            gap: 1rem;
            margin-top: 1.5rem;
        }}
        
        .stat-card {{
            background: var(--bg-secondary);
            padding: 1rem;
            border-radius: 12px;
            text-align: center;
            border: 1px solid var(--border-color);
        }}
        
        .stat-value {{
            font-size: 2rem;
            font-weight: 700;
            color: var(--text-primary);
            display: block;
        }}
        
        .stat-label {{
            font-size: 0.875rem;
            color: var(--text-secondary);
            text-transform: uppercase;
            letter-spacing: 0.05em;
        }}
        
        .social-links a {{
            display: inline-block;
            margin-right: 1rem;
            color: var(--text-secondary);
            font-size: 1.5rem;
            transition: color 0.2s;
        }}
        
        .social-links a:hover {{
            color: var(--text-primary);
        }}
    </style>
    
    <!-- Banner -->
    <div class="profile-banner" style="{'background-image: url(' + profile_user.banner_image + ');' if profile_user.banner_image else ''}">
    </div>
    
    <!-- Profile Header -->
    <div class="container">
        <div class="profile-header">
            <div class="row">
                <div class="col-md-8">
                    <img 
                        src="{profile_user.profileImage or '/static/images/default-avatar.png'}" 
                        alt="{profile_user.displayName or profile_user.username}" 
                        class="profile-avatar"
                        onerror="this.src='/static/images/default-avatar.png'"
                    >
                    <h1 class="mt-3">{profile_user.displayName or profile_user.username}</h1>
                    {f'<p class="text-muted">@{profile_user.handle}</p>' if profile_user.handle else ''}
                    {f'<p class="lead mt-2">{profile_user.headline}</p>' if profile_user.headline else ''}
                    
                    {f'''<div class="social-links mt-3">
                        {''.join([f'<a href="{link.get("url")}" target="_blank" title="{link.get("platform")}"><i class="fab fa-{link.get("icon", "link")}"></i></a>' for link in social_links])}
                    </div>''' if social_links else ''}
                </div>
                <div class="col-md-4 text-end mt-5">
                    {'<a href="/profile/edit/" class="btn btn-primary"><i class="fas fa-edit me-2"></i>Edit Profile</a>' if is_own_profile else ''}
                </div>
            </div>
            
            <!-- Stats -->
            <div class="profile-stats mt-4">
                <div class="stat-card">
                    <span class="stat-value">{projects_count}</span>
                    <span class="stat-label">Initiated</span>
                </div>
                <div class="stat-card">
                    <span class="stat-value">{len(project_memberships)}</span>
                    <span class="stat-label">Layers</span>
                </div>
                <div class="stat-card">
                    <span class="stat-value">{workgroups_count}</span>
                    <span class="stat-label">Coordinating</span>
                </div>
                <div class="stat-card">
                    <span class="stat-value">{memberships_count}</span>
                    <span class="stat-label">Workgroups</span>
                </div>
                <div class="stat-card">
                    <span class="stat-value">{chair_count}</span>
                    <span class="stat-label">Chair</span>
                </div>
                <div class="stat-card">
                    <span class="stat-value">{submissions_count}</span>
                    <span class="stat-label">Submissions</span>
                </div>
                <div class="stat-card">
                    <span class="stat-value">{comments_count}</span>
                    <span class="stat-label">Comments</span>
                </div>
                {f'''<div class="stat-card">
                    <span class="stat-value">{referral_count}</span>
                    <span class="stat-label">Referrals</span>
                </div>''' if is_own_profile else ''}
            </div>
            
            {f'''<!-- Referral Code -->
            <div class="alert alert-info mt-3">
                <strong><i class="fas fa-share-alt me-2"></i>Your Referral Code:</strong> 
                <code id="referral-code">{referral_code}</code>
                <button class="btn btn-sm btn-outline-primary ms-2" onclick="copyReferralLink(this)">
                    <i class="fas fa-copy me-1"></i>Copy Link
                </button>
                <small class="d-block mt-2">Share this link to get credit when people join projects!</small>
            </div>''' if is_own_profile and referral_code else ''}
        </div>
        
        <!-- Content Tabs -->
        <div class="row mt-5">
            <div class="col-12">
                <ul class="nav nav-tabs" role="tablist">
                    <li class="nav-item">
                        <button class="nav-link active" data-bs-toggle="tab" data-bs-target="#about-tab">About</button>
                    </li>
                    <li class="nav-item">
                        <button class="nav-link" data-bs-toggle="tab" data-bs-target="#activity-tab">Activity</button>
                    </li>
                    <li class="nav-item">
                        <button class="nav-link" data-bs-toggle="tab" data-bs-target="#projects-tab">Initiated</button>
                    </li>
                    <li class="nav-item">
                        <button class="nav-link" data-bs-toggle="tab" data-bs-target="#my-projects-tab" id="my-projects">My Projects</button>
                    </li>
                    <li class="nav-item">
                        <button class="nav-link" data-bs-toggle="tab" data-bs-target="#coordinating-tab">Coordinating</button>
                    </li>
                    <li class="nav-item">
                        <button class="nav-link" data-bs-toggle="tab" data-bs-target="#memberships-tab">Memberships</button>
                    </li>
                    <li class="nav-item">
                        <button class="nav-link" data-bs-toggle="tab" data-bs-target="#chair-tab">Chair</button>
                    </li>
                    <li class="nav-item">
                        <button class="nav-link" data-bs-toggle="tab" data-bs-target="#submissions-tab">Submissions</button>
                    </li>
                </ul>
                
                <div class="tab-content mt-4">
                    <!-- About Tab -->
                    <div class="tab-pane fade show active" id="about-tab">
                        <div class="card">
                            <div class="card-body">
                                <h5 class="card-title">Bio</h5>
                                {f'<p>{profile_user.bio}</p>' if profile_user.bio else '<p class="text-muted">No bio provided yet.</p>'}
                                
                                <h5 class="card-title mt-4">Details</h5>
                                <p><strong>Member since:</strong> {profile_user.created_at.strftime('%B %Y') if profile_user.created_at else 'Unknown'}</p>
                                {f'<p><strong>Email:</strong> {profile_user.email}</p>' if profile_user.email else ''}
                                {f'<p><strong>Role:</strong> <span class="badge bg-primary">{profile_user.role}</span></p>' if profile_user.role else ''}
                            </div>
                        </div>
                    </div>
                    
                    <!-- Activity Tab -->
                    <div class="tab-pane fade" id="activity-tab">
                        <div class="card">
                            <div class="card-body">
                                <h5 class="card-title">Recent Activity</h5>
                                <div class="list-group list-group-flush">
                                    {_format_activity_items(recent_projects, recent_submissions)}
                                </div>
                            </div>
                        </div>
                    </div>
                    
                    <!-- Projects Tab -->
                    <div class="tab-pane fade" id="projects-tab">
                        <div class="card">
                            <div class="card-body">
                                <h5 class="card-title">Initiated Projects</h5>
                                <div class="list-group list-group-flush">
                                    {''.join([f'<a href="/projects/{p[1]}/" class="list-group-item list-group-item-action"><strong>{p[0]}</strong><br><small class="text-muted">Created {p[2]}</small></a>' for p in recent_projects]) if recent_projects else '<p class="text-muted">No projects yet.</p>'}
                                </div>
                            </div>
                        </div>
                    </div>
                    
                    <!-- My Projects Tab -->
                    <div class="tab-pane fade" id="my-projects-tab">
                        <div class="card">
                            <div class="card-body">
                                <h5 class="card-title">Layer Memberships</h5>
                                <div class="list-group list-group-flush">
                                    {''.join([f'''<a href="/projects/{pm.project.slug}/" class="list-group-item list-group-item-action">
                                        <div class="d-flex justify-content-between align-items-center">
                                            <div>
                                                <strong>{pm.project.name}</strong>
                                                <br><small class="text-muted">Role: {pm.role or "Member"} • Joined {pm.joined_at.strftime("%b %Y") if pm.joined_at else "Unknown"}</small>
                                                {f'<br><small class="text-success"><i class="fas fa-user-plus me-1"></i>Referred by {pm.referred_by.displayName or pm.referred_by.username}</small>' if pm.referred_by else ''}
                                            </div>
                                        </div>
                                    </a>''' for pm in project_memberships]) if project_memberships else '<p class="text-muted">Not a member of any projects yet.</p>'}
                                </div>
                            </div>
                        </div>
                    </div>
                    
                    <!-- Coordinating Tab -->
                    <div class="tab-pane fade" id="coordinating-tab">
                        <div class="card">
                            <div class="card-body">
                                <h5 class="card-title">Coordinating Workgroups</h5>
                                <div class="list-group list-group-flush">
                                    {''.join([f'<a href="/workgroups/{wg.slug}/" class="list-group-item list-group-item-action"><strong>{wg.name}</strong><br><small class="text-muted">Status: {wg.status}</small></a>' for wg in coordinated_workgroups]) if coordinated_workgroups else '<p class="text-muted">Not coordinating any workgroups yet.</p>'}
                                </div>
                            </div>
                        </div>
                    </div>
                    
                    <!-- Memberships Tab -->
                    <div class="tab-pane fade" id="memberships-tab">
                        <div class="card">
                            <div class="card-body">
                                <h5 class="card-title">Workgroup Memberships</h5>
                                <div class="list-group list-group-flush">
                                    {''.join([f'<a href="/workgroups/{m[1]}/" class="list-group-item list-group-item-action"><strong>{m[0]}</strong><br><small class="text-muted">Joined {m[2]}</small></a>' for m in memberships_q]) if memberships_q else '<p class="text-muted">Not a member of any workgroups yet.</p>'}
                                </div>
                            </div>
                        </div>
                    </div>
                    
                    <!-- Chair Tab -->
                    <div class="tab-pane fade" id="chair-tab">
                        <div class="card">
                            <div class="card-body">
                                <h5 class="card-title">Chair Positions</h5>
                                <div class="list-group list-group-flush">
                                    {''.join([f'<a href="/workgroups/{c[1]}/" class="list-group-item list-group-item-action"><strong>{c[0]}</strong><br><small class="text-muted">{"Approved" if c[3] else "Pending approval"} - Set {c[2]}</small></a>' for c in chairs_q]) if chairs_q else '<p class="text-muted">No chair positions yet.</p>'}
                                </div>
                            </div>
                        </div>
                    </div>
                    
                    <!-- Submissions Tab -->
                    <div class="tab-pane fade" id="submissions-tab">
                        <div class="card">
                            <div class="card-body">
                                <h5 class="card-title">All Submissions</h5>
                                <div class="list-group list-group-flush">
                                    {''.join([f'<a href="/submit/status/{s.id}/" class="list-group-item list-group-item-action"><strong>{s.draft_name or s.id}</strong><br><small class="text-muted">{s.status.title()} - Submitted {s.submitted_at.strftime("%Y-%m-%d") if s.submitted_at else "Unknown"}</small></a>' for s in all_submissions_q]) if all_submissions_q else '<p class="text-muted">No submissions yet.</p>'}
                                </div>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    (function() {{
        var hash = window.location.hash;
        if (hash === '#my-projects') {{
            var tab = document.querySelector('[data-bs-target="#my-projects-tab"]');
            if (tab) bootstrap.Tab.getOrCreateInstance(tab).show();
        }}
    }})();
    function copyReferralLink(btn) {{
        var el = document.getElementById('referral-code');
        if (!el) return;
        var url = window.location.origin + '/?ref=' + el.textContent;
        navigator.clipboard.writeText(url).then(function() {{
            if (btn) {{
                var orig = btn.innerHTML;
                btn.innerHTML = '<i class="fas fa-check me-1"></i>Copied!';
                setTimeout(function() {{ btn.innerHTML = orig; }}, 2000);
            }}
        }}).catch(function() {{ alert('Failed to copy'); }});
    }}
    </script>
    """
    
    return render_page(f"{profile_user.displayName or profile_user.username} - MLGH", content, theme=current_theme, user_menu=user_menu)

def _format_activity_items(projects, submissions):
    """Helper function to format activity items"""
    from datetime import datetime
    items = []
    
    for project in projects:
        # project[2] is created_at from raw SQL (might be string or datetime)
        date = project[2]
        if isinstance(date, str):
            try:
                date = datetime.fromisoformat(date.replace('Z', '+00:00'))
            except:
                date = datetime.utcnow()
        items.append((date, 'project', f'Created project <strong>{project[0]}</strong>', f'/projects/{project[1]}/'))
    
    for submission in submissions:
        # submission is a tuple (draft_name, submitted_at, id, approved) from recent_submissions
        draft_name = submission[0]
        date = submission[1] if submission[1] else datetime.utcnow()
        submission_id = submission[2]
        approved = submission[3]
        status = "Approved" if approved else "Draft"
        items.append((date, 'submission', f'Submitted <strong>{draft_name}</strong> <span class="badge bg-{"success" if approved else "secondary"}">{status}</span>', f'/submit/status/{submission_id}/'))
    
    # Sort by date (all datetime objects now)
    items.sort(key=lambda x: x[0], reverse=True)
    
    html = ''
    for date, type, text, link in items[:10]:  # Show latest 10
        icon = 'fa-folder' if type == 'project' else 'fa-file-alt'
        html += f'''
        <a href="{link}" class="list-group-item list-group-item-action">
            <div class="d-flex justify-content-between align-items-center">
                <div>
                    <i class="fas {icon} me-2"></i>
                    {text}
                </div>
                <small class="text-muted">{date.strftime('%b %d, %Y') if date else ''}</small>
            </div>
        </a>
        '''
    
    return html if html else '<p class="text-muted">No recent activity.</p>'

@app.route('/profile/edit/')
@require_auth
def profile_edit():
    """Profile edit page"""
    current_user_data = get_current_user()
    if not current_user_data:
        return redirect('/login')
    
    user = User.query.get(current_user_data['id'])
    if not user:
        return redirect('/')
    
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    
    # Parse social links
    import json
    social_links = []
    if user.social_links:
        try:
            social_links = json.loads(user.social_links)
        except:
            social_links = []
    
    # Social link platforms
    platforms = [
        {'name': 'Twitter', 'icon': 'twitter', 'placeholder': 'https://twitter.com/username'},
        {'name': 'GitHub', 'icon': 'github', 'placeholder': 'https://github.com/username'},
        {'name': 'LinkedIn', 'icon': 'linkedin', 'placeholder': 'https://linkedin.com/in/username'},
        {'name': 'Website', 'icon': 'globe', 'placeholder': 'https://yourwebsite.com'},
    ]
    
    content = f"""
    <div class="container mt-4">
        <div class="row">
            <div class="col-md-8 mx-auto">
                <h1>Edit Profile</h1>
                <p class="text-muted">Update your profile information</p>
                
                <div class="card mb-4">
                    <div class="card-header">
                        <h5 class="mb-0">Profile Images</h5>
                    </div>
                    <div class="card-body">
                        <div class="row">
                            <div class="col-md-6">
                                <label class="form-label">Profile Picture</label>
                                <div class="text-center mb-3">
                                    <img 
                                        id="profile-image-preview" 
                                        src="{user.profileImage or '/static/images/default-avatar.png'}" 
                                        class="img-thumbnail rounded-circle" 
                                        style="width: 150px; height: 150px; object-fit: cover;"
                                        onerror="this.src='/static/images/default-avatar.png'"
                                    >
                                </div>
                                <input 
                                    type="file" 
                                    class="form-control" 
                                    id="profile-image-file"
                                    accept="image/*"
                                    onchange="previewImage(this, 'profile-image-preview')"
                                >
                                <div class="form-text">Max 600×600px, 5MB. PNG, JPG, GIF, WebP, SVG</div>
                                <button class="btn btn-primary btn-sm mt-2 w-100" onclick="uploadProfileImage()">
                                    <i class="fas fa-upload me-2"></i>Upload Profile Picture
                                </button>
                            </div>
                            <div class="col-md-6">
                                <label class="form-label">Banner Image</label>
                                <div class="mb-3" style="height: 150px; overflow: hidden; border-radius: 8px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);">
                                    <img 
                                        id="banner-image-preview" 
                                        src="{user.banner_image or ''}" 
                                        class="w-100" 
                                        style="height: 150px; object-fit: cover; display: {'block' if user.banner_image else 'none'};"
                                    >
                                </div>
                                <input 
                                    type="file" 
                                    class="form-control" 
                                    id="banner-image-file"
                                    accept="image/*"
                                    onchange="previewBannerImage(this)"
                                >
                                <div class="form-text">Max 600×600px, 5MB. PNG, JPG, GIF, WebP, SVG</div>
                                <button class="btn btn-primary btn-sm mt-2 w-100" onclick="uploadBannerImage()">
                                    <i class="fas fa-upload me-2"></i>Upload Banner
                                </button>
                            </div>
                        </div>
                    </div>
                </div>
                
                <div class="card mb-4">
                    <div class="card-header">
                        <h5 class="mb-0">Basic Information</h5>
                    </div>
                    <div class="card-body">
                        <form id="profile-form">
                            <div class="mb-3">
                                <label for="headline" class="form-label">Headline</label>
                                <input 
                                    type="text" 
                                    class="form-control" 
                                    id="headline"
                                    maxlength="200"
                                    placeholder="Your professional headline..."
                                    value="{user.headline or ''}"
                                >
                                <div class="form-text">A short description of what you do (max 200 characters)</div>
                            </div>
                            
                            <div class="mb-3">
                                <label for="bio" class="form-label">Bio</label>
                                <textarea 
                                    class="form-control" 
                                    id="bio"
                                    rows="4"
                                    placeholder="Tell us about yourself..."
                                >{user.bio or ''}</textarea>
                                <div class="form-text">A longer description of your background and interests</div>
                            </div>
                        </form>
                    </div>
                </div>
                
                <div class="card mb-4">
                    <div class="card-header">
                        <h5 class="mb-0">Social Links</h5>
                    </div>
                    <div class="card-body">
                        <div id="social-links-container">
                            {_render_social_link_inputs(platforms, social_links)}
                        </div>
                    </div>
                </div>
                
                <div class="d-flex justify-content-between">
                    <a href="/profile/{user.username}/" class="btn btn-secondary">
                        <i class="fas fa-times me-2"></i>Cancel
                    </a>
                    <button class="btn btn-primary" onclick="saveProfile()">
                        <i class="fas fa-save me-2"></i>Save Changes
                    </button>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    function previewImage(input, previewId) {{
        if (input.files && input.files[0]) {{
            const reader = new FileReader();
            reader.onload = function(e) {{
                document.getElementById(previewId).src = e.target.result;
            }};
            reader.readAsDataURL(input.files[0]);
        }}
    }}
    
    function previewBannerImage(input) {{
        if (input.files && input.files[0]) {{
            const reader = new FileReader();
            const preview = document.getElementById('banner-image-preview');
            reader.onload = function(e) {{
                preview.src = e.target.result;
                preview.style.display = 'block';
            }};
            reader.readAsDataURL(input.files[0]);
        }}
    }}
    
    async function uploadProfileImage() {{
        const fileInput = document.getElementById('profile-image-file');
        if (!fileInput.files || !fileInput.files[0]) {{
            alert('Please select an image first');
            return;
        }}
        
        const formData = new FormData();
        formData.append('file', fileInput.files[0]);
        formData.append('type', 'profile');
        
        try {{
            const response = await fetch('/api/user/upload-image', {{
                method: 'POST',
                credentials: 'include',
                body: formData
            }});
            
            const data = await response.json();
            if (response.ok) {{
                alert('Profile picture uploaded successfully!');
                location.reload();
            }} else {{
                alert(data.error || 'Failed to upload image');
            }}
        }} catch (error) {{
            console.error('Error uploading image:', error);
            alert('Failed to upload image');
        }}
    }}
    
    async function uploadBannerImage() {{
        const fileInput = document.getElementById('banner-image-file');
        if (!fileInput.files || !fileInput.files[0]) {{
            alert('Please select an image first');
            return;
        }}
        
        const formData = new FormData();
        formData.append('file', fileInput.files[0]);
        formData.append('type', 'banner');
        
        try {{
            const response = await fetch('/api/user/upload-image', {{
                method: 'POST',
                credentials: 'include',
                body: formData
            }});
            
            const data = await response.json();
            if (response.ok) {{
                alert('Banner image uploaded successfully!');
                location.reload();
            }} else {{
                alert(data.error || 'Failed to upload image');
            }}
        }} catch (error) {{
            console.error('Error uploading image:', error);
            alert('Failed to upload image');
        }}
    }}
    
    async function saveProfile() {{
        const headline = document.getElementById('headline').value;
        const bio = document.getElementById('bio').value;
        
        // Collect social links
        const socialLinks = [];
        document.querySelectorAll('[data-social-platform]').forEach(input => {{
            const url = input.value.trim();
            if (url) {{
                socialLinks.push({{
                    platform: input.dataset.socialPlatform,
                    icon: input.dataset.socialIcon,
                    url: url
                }});
            }}
        }});
        
        try {{
            const response = await fetch('/api/user/profile/', {{
                method: 'PUT',
                headers: {{ 'Content-Type': 'application/json' }},
                body: JSON.stringify({{
                    headline: headline,
                    bio: bio,
                    social_links: socialLinks
                }})
            }});
            
            const data = await response.json();
            if (response.ok) {{
                alert('Profile updated successfully!');
                window.location.href = '/profile/{user.username}/';
            }} else {{
                alert(data.error || 'Failed to update profile');
            }}
        }} catch (error) {{
            console.error('Error updating profile:', error);
            alert('Failed to update profile');
        }}
    }}
    </script>
    """
    
    return render_page("Edit Profile - MLGH", content, theme=current_theme, user_menu=user_menu)

def _render_social_link_inputs(platforms, existing_links):
    """Helper to render social link input fields"""
    html = ''
    
    for platform in platforms:
        # Find existing link for this platform
        existing = next((link for link in existing_links if link.get('platform') == platform['name']), None)
        value = existing['url'] if existing else ''
        
        html += f'''
        <div class="mb-3">
            <label class="form-label">
                <i class="fab fa-{platform['icon']} me-2"></i>{platform['name']}
            </label>
            <input 
                type="url" 
                class="form-control" 
                data-social-platform="{platform['name']}"
                data-social-icon="{platform['icon']}"
                placeholder="{platform['placeholder']}"
                value="{value}"
            >
        </div>
        '''
    
    return html

@app.route('/roles/')
def roles_directory():
    """Roles directory page - browse all roles across projects"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = f"""
    <div class="container mt-4">
        <div class="row mb-4">
            <div class="col-md-8">
                <h1>Roles Directory</h1>
                <p class="lead">Browse and claim roles across all projects</p>
            </div>
        </div>
        
        <div class="row mb-4">
            <div class="col-md-4">
                <label for="project-filter" class="form-label">Layer:</label>
                <select id="project-filter" class="form-select" onchange="loadRoles()">
                    <option value="">All Layers</option>
                </select>
            </div>
            <div class="col-md-4">
                <label for="status-filter" class="form-label">Status:</label>
                <select id="status-filter" class="form-select" onchange="loadRoles()">
                    <option value="">All Statuses</option>
                    <option value="approved">Approved</option>
                    <option value="draft">Draft</option>
                    <option value="deprecated">Deprecated</option>
                </select>
            </div>
            <div class="col-md-4">
                <label for="search-input" class="form-label">Search:</label>
                <input type="text" id="search-input" class="form-control" placeholder="Search roles..." onkeyup="filterRoles()">
            </div>
        </div>
        
        <div id="roles-container" class="row">
            <div class="col-12 text-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    let allRoles = [];
    let allProjects = [];
    
    async function loadProjects() {{
        try {{
            const response = await fetch('/api/projects/?approval_status=approved');
            const data = await response.json();
            allProjects = data.projects;
            
            const select = document.getElementById('project-filter');
            allProjects.forEach(project => {{
                const option = document.createElement('option');
                option.value = project.id;
                option.textContent = project.name;
                select.appendChild(option);
            }});
        }} catch (error) {{
            console.error('Error loading projects:', error);
        }}
    }}
    
    async function loadRoles() {{
        const projectFilter = document.getElementById('project-filter').value;
        const statusFilter = document.getElementById('status-filter').value;
        
        try {{
            allRoles = [];
            
            if (projectFilter) {{
                // Load roles for specific project
                let url = `/api/projects/${{projectFilter}}/roles/`;
                if (statusFilter) url += `?status=${{statusFilter}}`;
                
                const response = await fetch(url);
                const data = await response.json();
                allRoles = data.roles;
            }} else {{
                // Load roles from all projects
                for (const project of allProjects) {{
                    let url = `/api/projects/${{project.id}}/roles/`;
                    if (statusFilter) url += `?status=${{statusFilter}}`;
                    
                    const response = await fetch(url);
                    const data = await response.json();
                    allRoles = allRoles.concat(data.roles.map(r => ({{...r, project_name: project.name, project_slug: project.slug}})));
                }}
            }}
            
            displayRoles(allRoles);
        }} catch (error) {{
            console.error('Error loading roles:', error);
            document.getElementById('roles-container').innerHTML = '<div class="col-12"><div class="alert alert-danger">Error loading roles</div></div>';
        }}
    }}
    
    function filterRoles() {{
        const searchTerm = document.getElementById('search-input').value.toLowerCase();
        const filtered = allRoles.filter(r => 
            r.title_guild.toLowerCase().includes(searchTerm) ||
            (r.description && r.description.toLowerCase().includes(searchTerm))
        );
        displayRoles(filtered);
    }}
    
    function displayRoles(roles) {{
        const container = document.getElementById('roles-container');
        
        if (roles.length === 0) {{
            container.innerHTML = '<div class="col-12"><div class="alert alert-info">No roles found</div></div>';
            return;
        }}
        
        let html = '';
        roles.forEach(role => {{
            const statusBadge = role.status === 'approved' 
                ? '<span class="badge bg-success">Approved</span>' 
                : role.status === 'draft'
                ? '<span class="badge bg-warning">Draft</span>'
                : '<span class="badge bg-secondary">Deprecated</span>';
            
            const claimBadge = role.claim_requires_approval 
                ? '<span class="badge bg-info"><i class="fas fa-check-circle me-1"></i>Approval Required</span>'
                : '<span class="badge bg-success"><i class="fas fa-bolt me-1"></i>Instant Claim</span>';
            
            html += `
                <div class="col-md-6 col-lg-4 mb-4">
                    <div class="card h-100">
                        <div class="card-body">
                            <h5 class="card-title">
                                <a href="/roles/${{role.slug}}/">${{role.title_guild}}</a>
                            </h5>
                            <div class="mb-2">
                                ${{statusBadge}}
                                ${{claimBadge}}
                            </div>
                            <p class="card-text text-muted small">${{role.description.substring(0, 100)}}...</p>
                            <div class="mt-3">
                                <small class="text-muted">
                                    <i class="fas fa-project-diagram me-1"></i> ${{role.project_name || 'Unknown Project'}}
                                </small>
                            </div>
                        </div>
                        <div class="card-footer">
                            <small class="text-muted">
                                <i class="fas fa-hand-paper me-1"></i> ${{role.claims_count || 0}} claims
                            </small>
                        </div>
                    </div>
                </div>
            `;
        }});
        
        container.innerHTML = html;
    }}
    
    // Load data on page load
    loadProjects().then(() => loadRoles());
    </script>
    """
    
    return render_page("Roles Directory - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/role-images/')
def role_images_directory():
    """Global role images directory - browse all roles with images"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = f"""
    <div class="container mt-4">
        <div class="row mb-4">
            <div class="col-md-8">
                <h1><i class="fas fa-images me-2"></i>Role Images Gallery</h1>
                <p class="lead">Browse and vote on role images across all projects</p>
            </div>
        </div>
        
        <div class="row mb-4">
            <div class="col-md-6">
                <label for="project-filter" class="form-label">Layer:</label>
                <select id="project-filter" class="form-select" onchange="loadRoleImages()">
                    <option value="">All Layers</option>
                </select>
            </div>
            <div class="col-md-6">
                <label for="search-input" class="form-label">Search Roles:</label>
                <input type="text" id="search-input" class="form-control" placeholder="Search roles..." onkeyup="filterRoles()">
            </div>
        </div>
        
        <div id="roles-container" class="row">
            <div class="col-12 text-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    let allRoles = [];
    let allProjects = [];
    
    async function loadProjects() {{
        try {{
            const response = await fetch('/api/projects/');
            const data = await response.json();
            allProjects = data.projects;
            
            const select = document.getElementById('project-filter');
            allProjects.forEach(project => {{
                const option = document.createElement('option');
                option.value = project.id;
                option.textContent = project.name;
                select.appendChild(option);
            }});
        }} catch (error) {{
            console.error('Error loading projects:', error);
        }}
    }}
    
    async function loadRoleImages() {{
        const projectFilter = document.getElementById('project-filter').value;
        
        document.getElementById('roles-container').innerHTML = `
            <div class="col-12 text-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        `;
        
        try {{
            let url = '/api/role-images/roles-with-stats/';
            if (projectFilter) url += '?project_id=' + encodeURIComponent(projectFilter);
            
            const response = await fetch(url);
            const data = await response.json();
            allRoles = (data.roles || []).filter(r => (r.image_count || 0) > 0);
            
            displayRoles(allRoles);
        }} catch (error) {{
            console.error('Error loading roles:', error);
            document.getElementById('roles-container').innerHTML = '<div class="col-12"><div class="alert alert-danger">Error loading roles</div></div>';
        }}
    }}
    
    function displayRoles(roles) {{
        const container = document.getElementById('roles-container');
        
        if (roles.length === 0) {{
            container.innerHTML = '<div class="col-12"><div class="alert alert-info">No roles found</div></div>';
            return;
        }}
        
        let html = '';
        roles.forEach(role => {{
            const projectName = role.project_name || 'Unknown Project';
            const roleSlug = role.role_slug || role.slug || '';
            const imageCount = role.image_count != null ? role.image_count : 0;
            const voteCount = role.vote_count != null ? role.vote_count : 0;
            
            html += `
                <div class="col-md-4 mb-4">
                    <div class="card h-100">
                        <div class="card-body">
                            <h5 class="card-title">
                                <a href="/roles/${{roleSlug}}/images/" class="text-decoration-none">${{role.title_guild}}</a>
                            </h5>
                            ${{role.title_operational ? `<h6 class="card-subtitle mb-2 text-muted">${{role.title_operational}}</h6>` : ''}}
                            <p class="card-text text-muted small">
                                <i class="fas fa-project-diagram me-1"></i>${{projectName}}
                            </p>
                            <p class="card-text">${{role.description ? role.description.substring(0, 100) + '...' : ''}}</p>
                            <div class="d-flex justify-content-between align-items-center flex-wrap gap-2">
                                <a href="/roles/${{roleSlug}}/images/" class="btn btn-primary btn-sm">
                                    <i class="fas fa-images me-1"></i>View Images
                                </a>
                                <span class="text-muted small">${{imageCount}} image${{imageCount !== 1 ? 's' : ''}} · ${{voteCount}} vote${{voteCount !== 1 ? 's' : ''}}</span>
                            </div>
                        </div>
                    </div>
                </div>
            `;
        }});
        
        container.innerHTML = html;
    }}
    
    function filterRoles() {{
        const searchTerm = document.getElementById('search-input').value.toLowerCase();
        const filtered = allRoles.filter(role => 
            role.title_guild.toLowerCase().includes(searchTerm) ||
            (role.title_operational && role.title_operational.toLowerCase().includes(searchTerm)) ||
            role.description.toLowerCase().includes(searchTerm)
        );
        displayRoles(filtered);
    }}
    
    // Load data on page load
    loadProjects().then(() => loadRoleImages());
    </script>
    """
    
    return render_page("Role Images Gallery - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/roles/<role_slug>/')
def role_detail(role_slug):
    """Role detail page with claims"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = f"""
    <div class="container mt-4">
        <div id="role-header" class="mb-4">
            <div class="d-flex justify-content-center py-5">
                <div class="spinner-border text-primary" role="status">
                    <span class="visually-hidden">Loading...</span>
                </div>
            </div>
        </div>
        
        <div class="row">
            <div class="col-md-8">
                <div class="card mb-4">
                    <div class="card-header"><h5>Description</h5></div>
                    <div class="card-body" id="role-description">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
                
                <div class="card">
                    <div class="card-header">
                        <div class="d-flex justify-content-between align-items-center">
                            <h5 class="mb-0">Active Claims</h5>
                            <span id="role-claim-btn-placeholder"></span>
                        </div>
                    </div>
                    <div class="card-body" id="role-claims">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
            </div>
            
            <div class="col-md-4">
                <div class="card mb-4">
                    <div class="card-header"><h5>Role Details</h5></div>
                    <div class="card-body" id="role-details">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
                
                <div class="card">
                    <div class="card-header"><h5>Configuration</h5></div>
                    <div class="card-body" id="role-config">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    let role = null;
    let project = null;
    const roleSlug = '{role_slug}';
    const isAuthenticated = {'true' if current_user else 'false'};
    const currentUserId = {current_user['id'] if current_user else 'null'};
    
    async function loadRole() {{
        try {{
            // Load all projects to find the role
            const projectsResp = await fetch('/api/projects/');
            const projectsData = await projectsResp.json();
            
            // Search for role across all projects
            for (const proj of projectsData.projects) {{
                const rolesResp = await fetch(`/api/projects/${{proj.id}}/roles/`);
                const rolesData = await rolesResp.json();
                const found = rolesData.roles.find(r => r.slug === roleSlug);
                
                if (found) {{
                    role = found;
                    project = proj;
                    break;
                }}
            }}
            
            if (!role) {{
                document.getElementById('role-header').innerHTML = '<div class="alert alert-danger">Role not found</div>';
                return;
            }}
            
            // Load full role details
            const detailResp = await fetch(`/api/roles/${{role.id}}/`);
            role = await detailResp.json();
            
            displayRoleHeader();
            displayRoleDescription();
            displayRoleDetails();
            displayRoleConfig();
            loadClaims();
        }} catch (error) {{
            console.error('Error loading role:', error);
            document.getElementById('role-header').innerHTML = '<div class="alert alert-danger">Error loading role</div>';
        }}
    }}
    
    function displayRoleHeader() {{
        const statusBadge = getStatusBadge(role.status);
        const editBtn = (role.can_edit) ? '<button type="button" class="btn btn-outline-secondary mb-2 me-2" onclick="editRole()"><i class="fas fa-edit me-2"></i>Edit Role</button>' : '';
        
        document.getElementById('role-header').innerHTML = `
            <div class="row">
                <div class="col-md-8">
                    <nav aria-label="breadcrumb">
                        <ol class="breadcrumb">
                            <li class="breadcrumb-item"><a href="/projects/">Layers</a></li>
                            <li class="breadcrumb-item"><a href="/projects/${{project.slug}}/">${{project.name}}</a></li>
                            <li class="breadcrumb-item active">${{role.title_guild}}</li>
                        </ol>
                    </nav>
                    <h1>${{role.title_guild}}</h1>
                    ${{role.title_operational ? `<h5 class="text-muted">${{role.title_operational}}</h5>` : ''}}
                    <div class="mb-3">
                        ${{statusBadge}}
                        ${{role.public_visible ? '<span class="badge bg-info ms-2">Public</span>' : ''}}
                    </div>
                </div>
                <div class="col-md-4 text-end">
                    ${{editBtn}}
                    <a href="/roles/${{roleSlug}}/images/" class="btn btn-outline-primary mb-2"><i class="fas fa-images me-2"></i>View Images</a>
                </div>
            </div>
        `;
    }}
    
    function displayRoleDescription() {{
        let html = `<p>${{role.description}}</p>`;
        document.getElementById('role-description').innerHTML = html;
    }}
    
    function displayRoleDetails() {{
        const clusterLine = role.cluster_name
            ? `<p><strong>Cluster:</strong> <a href="/projects/${{project.slug}}/#clusters">${{role.cluster_name}}</a></p>`
            : (role.cluster_id ? '<p><strong>Cluster:</strong> <span class="text-muted">—</span></p>' : '');
        
        const imageHtml = role.image_url ? `<div class="mb-3 text-center"><img src="${{role.image_url}}" alt="${{role.title_guild}}" class="img-fluid rounded" style="max-height: 200px;"></div>` : '';
        
        document.getElementById('role-details').innerHTML = `
            ${{imageHtml}}
            <p><strong>Layer:</strong> <a href="/projects/${{project.slug}}/">${{project.name}}</a></p>
            ${{clusterLine}}
            <p><strong>Status:</strong> ${{role.status}}</p>
            <p><strong>Visibility:</strong> ${{role.public_visible ? 'Public' : 'Private'}}</p>
            <p><strong>Active Claims:</strong> ${{role.active_claims_count || 0}}</p>
            <p><strong>Created:</strong> ${{new Date(role.created_at).toLocaleDateString()}}</p>
        `;
    }}
    
    function displayRoleConfig() {{
        document.getElementById('role-config').innerHTML = `
            <p><strong>Claim Approval:</strong> ${{role.claim_requires_approval ? 'Required' : 'Not Required'}}</p>
            <p><strong>Badges:</strong> ${{role.badge_enabled ? 'Enabled' : 'Disabled'}}</p>
            ${{role.badge_enabled ? `<p><strong>Badge Approval:</strong> ${{role.badge_requires_approval ? 'Required' : 'Not Required'}}</p>` : ''}}
        `;
    }}
    
    function getClaimPopoverContent(claim) {{
        const intent = claim.intent ? '<p class="mb-2"><strong>Intent:</strong><br><span style="white-space: pre-wrap; word-wrap: break-word;">' + (claim.intent || '').replace(/</g, '&lt;').replace(/>/g, '&gt;') + '</span></p>' : '';
        const links = (claim.evidence_links || []).filter(u => u && u.trim());
        const evidenceHtml = links.length ? links.map(u => '<a href="' + u + '" target="_blank" rel="noopener">' + u + '</a>').join('<br>') : '<span class="text-muted">No evidence yet</span>';
        const termStr = claim.term_duration_days 
            ? (claim.term_duration_days + ' days' + (claim.term_end ? ', until ' + new Date(claim.term_end).toLocaleDateString() : '')) 
            : 'Indefinite';
        return '<div class="text-start" style="min-width: 280px; max-width: 480px; white-space: normal; word-wrap: break-word;">' + intent +
            '<p class="mb-2"><strong>Supporting work:</strong><br>' + evidenceHtml + '</p>' +
            '<p class="mb-2"><strong>Term:</strong> ' + termStr + '</p>' +
            '<p class="mb-0 small text-muted">Claimed: ' + new Date(claim.created_at).toLocaleDateString() + '</p></div>';
    }}
    
    async function loadClaims() {{
        const container = document.getElementById('role-claims');
        const btnPlaceholder = document.getElementById('role-claim-btn-placeholder');
        
        if (!role.public_visible) {{
            if (btnPlaceholder) btnPlaceholder.innerHTML = '';
            container.innerHTML = '<p class="text-muted">Claims are only visible for public roles.</p>';
            return;
        }}
        
        try {{
            const response = await fetch(`/api/roles/${{role.id}}/claims/`);
            const data = await response.json();
            const claimsData = data.claims || [];
            const activeClaims = claimsData.filter(c => c.status === 'active' || c.status === 'pending_approval');
            const hasClaimed = isAuthenticated && claimsData.some(c => Number(c.claimant_id) === Number(currentUserId));
            
            if (btnPlaceholder) {{
                if (hasClaimed) {{
                    btnPlaceholder.innerHTML = '';
                }} else if (isAuthenticated) {{
                    btnPlaceholder.innerHTML = '<button class="btn btn-sm btn-primary" onclick="claimRole()"><i class="fas fa-hand-paper me-2"></i>Claim This Role</button>';
                }} else {{
                    btnPlaceholder.innerHTML = '<a href="/login/" class="btn btn-sm btn-primary">Login to Claim</a>';
                }}
            }}
            
            if (activeClaims.length === 0) {{
                container.innerHTML = '<p class="text-muted">No active claims yet</p>';
                return;
            }}
            
            const claimsDataDisplay = activeClaims;
            let html = '<div class="list-group">';
            claimsDataDisplay.forEach((claim, idx) => {{
                const claimantName = claim.claimant_name || ('User #' + claim.claimant_id);
                const claimantUsername = claim.claimant_username || '';
                const profileLink = claimantUsername ? '/profile/' + claimantUsername + '/' : '#';
                const nameDisplay = profileLink !== '#' 
                    ? '<a href="' + profileLink + '" class="text-decoration-none">' + claimantName + '</a>' 
                    : claimantName;
                
                html += `
                    <div class="list-group-item claim-list-item" data-claim-index="${{idx}}">
                        <div class="d-flex justify-content-between align-items-center">
                            <h6 class="mb-0">${{nameDisplay}}</h6>
                            <span class="badge bg-success">Active</span>
                        </div>
                        <small class="text-muted">Claimed: ${{new Date(claim.created_at).toLocaleDateString()}}</small>
                    </div>
                `;
            }});
            html += '</div>';
            
            container.innerHTML = html;
            
            container.querySelectorAll('.claim-list-item').forEach(el => {{
                const idx = parseInt(el.getAttribute('data-claim-index'), 10);
                const claim = claimsDataDisplay[idx];
                new bootstrap.Popover(el, {{
                    content: getClaimPopoverContent(claim),
                    html: true,
                    trigger: 'hover focus',
                    placement: 'auto',
                    container: 'body'
                }});
            }});
        }} catch (error) {{
            console.error('Error loading claims:', error);
            container.innerHTML = '<div class="alert alert-danger">Error loading claims</div>';
        }}
    }}
    
    function getStatusBadge(status) {{
        const badges = {{
            'draft': '<span class="badge bg-secondary">Draft</span>',
            'approved': '<span class="badge bg-success">Approved</span>',
            'deprecated': '<span class="badge bg-warning">Deprecated</span>',
            'archived': '<span class="badge bg-dark">Archived</span>'
        }};
        return badges[status] || '';
    }}
    
    function claimRole() {{
        if (role.status !== 'approved') {{
            alert('This role must be approved before it can be claimed');
            return;
        }}
        
        window.location.href = `/roles/${{roleSlug}}/claim/`;
    }}
    
    function editRole() {{
        const modalHtml = `
            <div class="modal fade" id="editRoleModal" tabindex="-1">
                <div class="modal-dialog modal-lg">
                    <div class="modal-content">
                        <div class="modal-header">
                            <h5 class="modal-title">Edit Role</h5>
                            <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                        </div>
                        <div class="modal-body">
                            <div id="edit-role-alert-container"></div>
                            <form id="editRoleForm">
                                <div class="mb-3">
                                    <label for="edit-role-title-guild" class="form-label">Guild Title *</label>
                                    <input type="text" class="form-control" id="edit-role-title-guild" required>
                                </div>
                                <div class="mb-3">
                                    <label for="edit-role-title-operational" class="form-label">Operational Title</label>
                                    <input type="text" class="form-control" id="edit-role-title-operational">
                                </div>
                                <div class="mb-3">
                                    <label for="edit-role-description" class="form-label">About / Description *</label>
                                    <textarea class="form-control" id="edit-role-description" rows="5" required></textarea>
                                </div>
                                <div class="mb-3">
                                    <label for="edit-role-cluster" class="form-label">Cluster</label>
                                    <select class="form-select" id="edit-role-cluster">
                                        <option value="">No cluster</option>
                                    </select>
                                </div>
                            </form>
                        </div>
                        <div class="modal-footer">
                            <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                            <button type="button" class="btn btn-primary" id="editRoleSubmitBtn">
                                <i class="fas fa-save me-2"></i>Save Changes
                            </button>
                        </div>
                    </div>
                </div>
            </div>
        `;
        if (!document.getElementById('editRoleModal')) {{
            document.body.insertAdjacentHTML('beforeend', modalHtml);
        }}
        document.getElementById('edit-role-title-guild').value = role.title_guild || '';
        document.getElementById('edit-role-title-operational').value = role.title_operational || '';
        document.getElementById('edit-role-description').value = role.description || '';
        document.getElementById('edit-role-alert-container').innerHTML = '';
        const clusterSelect = document.getElementById('edit-role-cluster');
        clusterSelect.innerHTML = '<option value="">No cluster</option>';
        fetch(`/api/projects/${{project.id}}/clusters/`).then(r => r.json()).then(d => {{
            (d.clusters || []).forEach(c => {{
                const opt = document.createElement('option');
                opt.value = c.id || '';
                opt.textContent = (c.name != null && c.name !== '') ? c.name : 'Unnamed';
                clusterSelect.appendChild(opt);
            }});
            clusterSelect.value = role.cluster_id || '';
        }});
        const modal = new bootstrap.Modal(document.getElementById('editRoleModal'));
        modal.show();
        document.getElementById('editRoleSubmitBtn').onclick = async () => {{
            const titleGuild = document.getElementById('edit-role-title-guild').value.trim();
            const titleOperational = document.getElementById('edit-role-title-operational').value.trim();
            const description = document.getElementById('edit-role-description').value.trim();
            const clusterId = document.getElementById('edit-role-cluster').value || null;
            if (!titleGuild || !description) {{
                document.getElementById('edit-role-alert-container').innerHTML = '<div class="alert alert-danger">Guild title and description are required.</div>';
                return;
            }}
            const btn = document.getElementById('editRoleSubmitBtn');
            btn.disabled = true;
            btn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Saving...';
            try {{
                const response = await fetch(`/api/roles/${{role.id}}/`, {{
                    method: 'PATCH',
                    headers: {{'Content-Type': 'application/json'}},
                    body: JSON.stringify({{ title_guild: titleGuild, title_operational: titleOperational || null, description, cluster_id: clusterId }})
                }});
                if (!response.ok) {{
                    const data = await response.json();
                    throw new Error(data.error || 'Failed to update role');
                }}
                role.title_guild = titleGuild;
                role.title_operational = titleOperational || null;
                role.description = description;
                role.cluster_id = clusterId;
                role.cluster_name = clusterId ? clusterSelect.options[clusterSelect.selectedIndex].text : null;
                modal.hide();
                displayRoleHeader();
                displayRoleDescription();
                displayRoleDetails();
            }} catch (err) {{
                document.getElementById('edit-role-alert-container').innerHTML = '<div class="alert alert-danger">' + (err.message || 'Failed to update role') + '</div>';
            }}
            btn.disabled = false;
            btn.innerHTML = '<i class="fas fa-save me-2"></i>Save Changes';
        }};
    }}
    
    // Load role on page load
    loadRole();
    </script>
    """
    
    return render_page(f"Role: {role_slug} - MLGH", content, theme=current_theme, user_menu=user_menu)

@app.route('/roles/<role_slug>/claim/')
@require_auth
def claim_role_page(role_slug):
    """Claim role form page"""
    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()
    
    content = f"""
    <div class="container mt-4">
        <div class="row">
            <div class="col-md-8 offset-md-2">
                <h1 class="mb-4" id="page-title">Claim Role</h1>
                
                <div id="alert-container"></div>
                
                <div id="role-info" class="card mb-4">
                    <div class="card-body text-center">
                        <div class="spinner-border text-primary"></div>
                    </div>
                </div>
                
                <form id="claimRoleForm" style="display: none;">
                    <div class="mb-3">
                        <label for="intent" class="form-label">Intent Statement</label>
                        <textarea class="form-control" id="intent" rows="4" placeholder="Describe your intent in claiming this role and how you plan to contribute..."></textarea>
                        <div class="form-text">Optional: Explain your motivation and plans</div>
                    </div>
                    
                    <div class="mb-3">
                        <label for="evidence_links" class="form-label">Supporting work</label>
                        <textarea class="form-control" id="evidence_links" rows="3" placeholder="https://example.com/my-work
https://github.com/username/project"></textarea>
                        <div class="form-text">Optional: Links to relevant work or contributions (one per line)</div>
                    </div>
                    
                    <div class="mb-3">
                        <label for="term_duration_months" class="form-label">Term duration (months)</label>
                        <select class="form-select" id="term_duration_months">
                            <option value="1">1 month</option>
                            <option value="3" selected>3 months</option>
                            <option value="6">6 months</option>
                            <option value="12">12 months</option>
                        </select>
                        <div class="form-text">Time limit for this claim</div>
                    </div>
                    
                    <div id="approval-notice" class="alert alert-warning" style="display: none;">
                        <i class="fas fa-exclamation-triangle me-2"></i>
                        <strong>Note:</strong> This role requires approval. Your claim will be pending until reviewed by a layer admin.
                    </div>
                    
                    <div class="d-flex gap-2">
                        <button type="submit" class="btn btn-primary" id="submitBtn">
                            <i class="fas fa-hand-paper me-2"></i>Submit Claim
                        </button>
                        <a href="/roles/{role_slug}/" class="btn btn-secondary">Cancel</a>
                    </div>
                </form>
            </div>
        </div>
    </div>
    
    <script>
    let role = null;
    let project = null;
    const roleSlug = '{role_slug}';
    
    async function loadRole() {{
        try {{
            // Load all projects to find the role
            const projectsResp = await fetch('/api/projects/');
            const projectsData = await projectsResp.json();
            
            // Search for role across all projects
            for (const proj of projectsData.projects) {{
                const rolesResp = await fetch(`/api/projects/${{proj.id}}/roles/`);
                const rolesData = await rolesResp.json();
                const found = rolesData.roles.find(r => r.slug === roleSlug);
                
                if (found) {{
                    role = found;
                    project = proj;
                    break;
                }}
            }}
            
            if (!role) {{
                document.getElementById('alert-container').innerHTML = '<div class="alert alert-danger">Role not found</div>';
                return;
            }}
            
            // Load full role details
            const detailResp = await fetch(`/api/roles/${{role.id}}/`);
            role = await detailResp.json();
            
            displayRoleInfo();
            
            if (role.claim_requires_approval) {{
                document.getElementById('approval-notice').style.display = 'block';
            }}
            
            document.getElementById('claimRoleForm').style.display = 'block';
        }} catch (error) {{
            console.error('Error loading role:', error);
            document.getElementById('alert-container').innerHTML = '<div class="alert alert-danger">Error loading role</div>';
        }}
    }}
    
    function displayRoleInfo() {{
        document.getElementById('page-title').textContent = `Claim Role: ${{role.title_guild}}`;
        
        document.getElementById('role-info').innerHTML = `
            <div class="card-body">
                <h5>${{role.title_guild}}</h5>
                ${{role.title_operational ? `<h6 class="text-muted">${{role.title_operational}}</h6>` : ''}}
                <p class="mt-3">${{role.description}}</p>
                <p class="mb-0"><strong>Layer:</strong> <a href="/projects/${{project.slug}}/">${{project.name}}</a></p>
            </div>
        `;
    }}
    
    document.getElementById('claimRoleForm').addEventListener('submit', async (e) => {{
        e.preventDefault();
        
        const submitBtn = document.getElementById('submitBtn');
        submitBtn.disabled = true;
        submitBtn.innerHTML = '<span class="spinner-border spinner-border-sm me-2"></span>Submitting...';
        
        // Parse evidence links
        const evidenceText = document.getElementById('evidence_links').value.trim();
        const evidenceLinks = evidenceText ? evidenceText.split('\\n').filter(l => l.trim()) : [];
        
        const termEl = document.getElementById('term_duration_months');
        const termVal = (termEl && termEl.value !== undefined && termEl.value !== '') ? termEl.value : '3';
        const termMonths = parseInt(termVal, 10) || 3;
        
        const formData = {{
            intent: document.getElementById('intent').value.trim() || null,
            evidence_links: evidenceLinks,
            term_duration_months: termMonths
        }};
        
        try {{
            const response = await fetch(`/api/roles/${{role.id}}/claims/`, {{
                method: 'POST',
                headers: {{'Content-Type': 'application/json'}},
                body: JSON.stringify(formData)
            }});
            
            const data = await response.json();
            
            if (response.ok) {{
                const statusMsg = role.claim_requires_approval 
                    ? 'Your claim has been submitted and is pending approval.'
                    : 'Your claim has been submitted successfully!';
                    
                document.getElementById('alert-container').innerHTML = `
                    <div class="alert alert-success">
                        <i class="fas fa-check-circle me-2"></i>
                        ${{statusMsg}} Redirecting...
                    </div>
                `;
                setTimeout(() => {{
                    window.location.href = `/roles/${{roleSlug}}/`;
                }}, 2000);
            }} else {{
                throw new Error(data.error || 'Failed to submit claim');
            }}
        }} catch (error) {{
            document.getElementById('alert-container').innerHTML = `
                <div class="alert alert-danger">
                    <i class="fas fa-exclamation-circle me-2"></i>
                    ${{error.message}}
                </div>
            `;
            submitBtn.disabled = false;
            submitBtn.innerHTML = '<i class="fas fa-hand-paper me-2"></i>Submit Claim';
        }}
    }});
    
    // Load role on page load
    loadRole();
    </script>
    """
    
    return render_page(f"Claim Role: {role_slug} - MLGH", content, theme=current_theme, user_menu=user_menu)

# Deployment API endpoint (development only)
@app.route('/_deploy/reload', methods=['POST'])
def reload_app():
    """Reload the application - development only"""
    if not IS_DEVELOPMENT:
        return jsonify({'error': 'Not available in production'}), 403
    
    # Clear Python cache
    import shutil
    cache_dirs = []
    for root, dirs, files in os.walk(os.path.dirname(os.path.abspath(__file__))):
        if '__pycache__' in dirs:
            cache_dirs.append(os.path.join(root, '__pycache__'))
        for file in files:
            if file.endswith('.pyc'):
                try:
                    os.remove(os.path.join(root, file))
                except:
                    pass
    
    for cache_dir in cache_dirs:
        try:
            shutil.rmtree(cache_dir)
        except:
            pass
    
    # Touch the file to trigger reload if using file watcher
    # For systemd, we'll return a signal to restart
    return jsonify({
        'status': 'success',
        'message': 'Cache cleared. Service restart required.',
        'restart_command': 'systemctl --user restart datatracker-dev.service'
    })

@app.route('/_deploy/status', methods=['GET'])
def deployment_status():
    """Check deployment status - comprehensive status endpoint"""
    import subprocess
    from datetime import datetime
    
    # Get git info
    git_branch = 'unknown'
    git_commit = 'unknown'
    git_commit_short = 'unknown'
    try:
        result = subprocess.run(['git', 'branch', '--show-current'], 
                               capture_output=True, text=True, timeout=2, cwd=os.path.dirname(os.path.abspath(__file__)))
        if result.returncode == 0:
            git_branch = result.stdout.strip() or 'unknown'
    except:
        pass
    
    try:
        result = subprocess.run(['git', 'rev-parse', 'HEAD'], 
                               capture_output=True, text=True, timeout=2, cwd=os.path.dirname(os.path.abspath(__file__)))
        if result.returncode == 0:
            git_commit = result.stdout.strip() or 'unknown'
            git_commit_short = git_commit[:8] if len(git_commit) > 8 else git_commit
    except:
        pass
    
    # Check service status
    service_name = f'datatracker{"-dev" if IS_DEVELOPMENT else ""}.service'
    service_active = None
    try:
        result = subprocess.run(['systemctl', '--user', 'is-active', service_name],
                               capture_output=True, text=True, timeout=2)
        service_active = result.returncode == 0
    except:
        pass
    
    # Check database
    db_exists = os.path.exists(DB_PATH)
    db_size = 0
    if db_exists:
        try:
            db_size = os.path.getsize(DB_PATH)
        except:
            pass
    
    status = {
        'environment': ENV,
        'port': PORT,
        'database': {
            'path': DB_PATH,
            'exists': db_exists,
            'size_bytes': db_size,
            'size_mb': round(db_size / 1024 / 1024, 2) if db_size > 0 else 0
        },
        'git': {
            'branch': git_branch,
            'commit': git_commit,
            'commit_short': git_commit_short
        },
        'service': {
            'name': service_name,
            'active': service_active
        },
        'deployed_at': datetime.now().isoformat(),
        'version': '2026-01-17-v2'
    }

    return jsonify(status)

# Draft pages: any file in drafts/ is served at /test/<filename>
_DRAFTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'drafts')

def _safe_draft_path(filename):
    """Resolve filename under _DRAFTS_DIR; return None if path escapes drafts."""
    if not filename or ".." in filename or filename.startswith("/"):
        return None
    path = os.path.normpath(os.path.join(_DRAFTS_DIR, filename))
    if not path.startswith(_DRAFTS_DIR):
        return None
    return path

@app.route('/test/', methods=['GET'])
@app.route('/test/<path:filename>', methods=['GET'])
def draft_page(filename=None):
    """Serve drafts/<filename> as text/html. /test/ serves digitalartifacts.htm."""
    name = filename or "digitalartifacts.htm"
    path = _safe_draft_path(name)
    if not path or not os.path.isfile(path):
        return jsonify({'error': 'Draft page not found'}), 404
    return send_file(path, mimetype='text/html; charset=utf-8')

@app.route('/_deploy/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    import subprocess
    
    # Check database connection
    db_healthy = False
    try:
        db.session.execute(db.text('SELECT 1'))
        db_healthy = True
    except:
        pass
    
    # Check service status
    service_name = f'datatracker{"-dev" if IS_DEVELOPMENT else ""}.service'
    service_healthy = None
    try:
        result = subprocess.run(['systemctl', '--user', 'is-active', service_name],
                               capture_output=True, text=True, timeout=2)
        service_healthy = result.returncode == 0
    except:
        pass
    
    overall_healthy = db_healthy and (service_healthy is True)
    
    return jsonify({
        'status': 'healthy' if overall_healthy else 'unhealthy',
        'database': 'connected' if db_healthy else 'disconnected',
        'service': 'active' if service_healthy else 'inactive',
        'timestamp': datetime.now().isoformat()
    }), 200 if overall_healthy else 503

@app.route('/api/annotations/<document_name>/count')
def annotation_count(document_name):
    """Get annotation count for a document"""
    annotations = get_document_annotations(document_name, 'draft')
    return jsonify({
        'count': len(annotations),
        'document': document_name
    })

@app.route('/_deploy/test', methods=['GET'])
def deployment_test():
    """Show a visible test page"""
    return f"""
    <html>
    <head><title>Deployment Test</title></head>
    <body style="font-family: Arial, sans-serif; padding: 20px;">
        <h1 style="color: red;">🚨 DEPLOYMENT TEST PAGE 🚨</h1>
        <div style="background-color: #ffcccc; border: 3px solid red; padding: 20px; margin: 20px 0; border-radius: 10px;">
            <h2 style="color: red;">If you can see this page, the deployment worked!</h2>
            <p><strong>Environment:</strong> {ENV}</p>
            <p><strong>Port:</strong> {PORT}</p>
            <p><strong>Database:</strong> {DB_PATH}</p>
            <p><strong>Time:</strong> {__import__('datetime').datetime.now()}</p>
        </div>
        <p><a href="/">← Back to main site</a></p>
    </body>
    </html>
    """

# ================================================================
# CLI COMMANDS
# ================================================================

import click

@app.cli.command('manage-votes')
@click.argument('action')
def manage_votes_cli(action):
    """Manage vote lifecycle. Usage: flask manage-votes tick"""
    if action == 'tick':
        now = datetime.utcnow()
        
        # Activate scheduled votes whose start_at has passed
        scheduled = Vote.query.filter(
            Vote.status == 'scheduled',
            Vote.start_at <= now
        ).all()
        
        activated_count = 0
        for vote in scheduled:
            success, msg = activate_vote(vote)
            if success:
                activated_count += 1
            print(f"  Activate vote {vote.id}: {msg}")
        
        # Close active votes whose end_at has passed
        active = Vote.query.filter(
            Vote.status == 'active',
            Vote.end_at <= now
        ).all()
        
        closed_count = 0
        for vote in active:
            success, msg = close_vote(vote)
            if success:
                closed_count += 1
            print(f"  Close vote {vote.id}: {msg}")
        
        print(f"✅ Tick complete: {activated_count} activated, {closed_count} closed")
    else:
        print(f"Unknown action: {action}. Use 'tick'.")

# ================================================================
# CLI COMMANDS
# ================================================================

import click

@app.cli.command('manage-votes')
@click.argument('action')
def manage_votes_cli(action):
    """Manage vote lifecycle. Usage: flask manage-votes tick"""
    if action == 'tick':
        now = datetime.utcnow()
        
        # Activate scheduled votes whose start_at has passed
        scheduled = Vote.query.filter(
            Vote.status == 'scheduled',
            Vote.start_at <= now
        ).all()
        
        activated_count = 0
        for vote in scheduled:
            success, msg = activate_vote(vote)
            if success:
                activated_count += 1
            print(f"  Activate vote {vote.id}: {msg}")
        
        # Close active votes whose end_at has passed
        active = Vote.query.filter(
            Vote.status == 'active',
            Vote.end_at <= now
        ).all()
        
        closed_count = 0
        for vote in active:
            success, msg = close_vote(vote)
            if success:
                closed_count += 1
            print(f"  Close vote {vote.id}: {msg}")
        
        print(f"Tick complete: {activated_count} activated, {closed_count} closed")
    else:
        print(f"Unknown action: {action}. Use 'tick'.")

if __name__ == '__main__':
    # Initialize deployment safety checks
    init_deployment_safety()
    # Initialize database on startup
    init_db()
    print(f"🚀 Starting MLGH Datatracker - BUILD {BUILD_NUMBER}")
    print(f"Environment: {ENV} mode on port {PORT}")
    print(f"Database: {DB_PATH}")
    # Disable reloader when running under systemd (detected by systemd environment)
    # The reloader can cause hanging in systemd services
    use_reloader = DEBUG and not os.environ.get('INVOCATION_ID')  # systemd sets INVOCATION_ID
    app.run(host='0.0.0.0', port=PORT, debug=DEBUG, use_reloader=use_reloader)

