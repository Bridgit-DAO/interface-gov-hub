"""Resolve draft context and load document body for record + reader views."""
from __future__ import annotations

import os
from html import escape
from typing import Any, Dict, Optional, Tuple

from services.documents import DRAFTS, submission_file_pages_words
from services.submissions import get_readable_submission_by_ref, get_submission_by_ref
from services.book_html import load_book_html_for_reader, looks_like_html_document
from services.submission_preview_md import markdown_to_safe_preview_html, text_looks_like_markdown
from services.ordinals import (
    process_ordinal_markdown,
    looks_like_html_inscription,
    format_ordinal_html_iframe_preview,
)


def _submission_uses_display_ordinal(submission) -> bool:
    if not submission:
        return False
    src = (getattr(submission, 'displayBodySource', None) or 'file').strip().lower()
    return src == 'ordinal' and bool(getattr(submission, 'displayOrdinalContentUrl', None))


def _load_ordinal_body(url: str, content_type: str, draft: dict) -> Tuple[str, bool, int, int]:
    from flask import current_app

    octype = content_type or ''
    document_content = ''
    render_html = False
    pages = draft.get('pages', 1)
    words = draft.get('words', 0)

    if url and ('text/' in octype or 'application/json' in octype):
        try:
            import requests
            from services.url_safety import validate_ordinals_fetch_url

            safe_url = validate_ordinals_fetch_url(url)
            headers = {
                'User-Agent': (
                    'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                    '(KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                )
            }
            response = requests.get(safe_url, headers=headers, timeout=10)
            if response.status_code == 200:
                raw_content = response.text
                words = len(raw_content.split())
                pages = max(1, (words + 499) // 500)
                draft['pages'] = pages
                draft['words'] = words

                is_markdown = False
                for pattern in (
                    r'^#{1,6}\s+.+$',
                    r'\*\*.+\*\*',
                    r'\*.+\*',
                    r'^\s*[-*+]\s+',
                    r'^\s*\d+\.\s+',
                    r'\[.+\]\(.+\)',
                    r'!\[.*\]\(.+\)',
                ):
                    import re

                    if re.search(pattern, raw_content, re.MULTILINE):
                        is_markdown = True
                        break

                if looks_like_html_inscription(raw_content, content_type):
                    document_content = format_ordinal_html_iframe_preview(url)
                    render_html = True
                elif is_markdown:
                    document_content = process_ordinal_markdown(raw_content)
                    render_html = True
                else:
                    document_content = raw_content
        except Exception as exc:
            current_app.logger.warning('Failed to fetch ordinal content: %s', exc)
            document_content = f'Error loading ordinal content: {exc}'
    elif url and octype.startswith('image/'):
        document_content = (
            f'<img src="{escape(url, quote=True)}" class="img-fluid" '
            f'style="max-width: 100%;" alt="Ordinal image content">'
        )
        render_html = True
    else:
        document_content = (
            f'Ordinal content type: {escape(octype)}\n'
            'Preview not available for this content type.'
        )

    return document_content, render_html, pages, words


def build_draft_context(
    draft_name: str,
    *,
    prefer_latest_revision: bool = False,
) -> Tuple[Optional[Dict[str, Any]], Optional[Any]]:
    """
    Resolve draft dict + submission for a URL ref (draft name, id, or ml_number).
    Returns (None, None) when not found.

    With ``prefer_latest_revision`` an ML number resolves to the latest approved
    revision rather than the Rev 00 parent row, so reading views serve the body
    the catalog advertises.
    """
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)
    submission = None
    lookup = get_readable_submission_by_ref if prefer_latest_revision else get_submission_by_ref

    if not draft:
        submission = lookup(draft_name)
        if submission:
            source_type = getattr(submission, 'sourceType', 'file')
            pages_count, words_count = submission_file_pages_words(submission)
            dbs = getattr(submission, 'displayBodySource', None) or 'file'
            displaying_linked = (
                dbs.strip().lower() == 'ordinal'
                and bool(getattr(submission, 'displayOrdinalContentUrl', None))
            )
            draft = {
                'name': submission.draft_name or submission.id,
                'title': submission.title,
                'authors': submission.authors,
                'abstract': submission.abstract or 'Abstract not available for this draft.',
                'status': submission.status,
                'group': submission.group,
                'date': submission.submitted_at.strftime('%Y-%m-%d') if submission.submitted_at else '',
                'rev': getattr(submission, 'revision_number', '') or '00',
                'pages': pages_count,
                'words': words_count,
                'stream': 'mltf',
                'ml_number': submission.ml_number,
                'sourceType': source_type,
                'ordinalId': getattr(submission, 'ordinalId', None),
                'inscriptionNumber': getattr(submission, 'inscriptionNumber', None),
                'blockHeight': getattr(submission, 'blockHeight', None),
                'inscriptionTimestamp': getattr(submission, 'inscriptionTimestamp', None),
                'ordinalContentType': getattr(submission, 'ordinalContentType', ''),
                'is_revision': getattr(submission, 'is_revision', False),
                'revision_number': getattr(submission, 'revision_number', ''),
                'parent_draft_name': getattr(submission, 'parent_draft_name', ''),
                'displayBodySource': dbs,
                'displayOrdinalId': getattr(submission, 'displayOrdinalId', None),
                'displayingLinkedOrdinal': displaying_linked,
                'document_category': getattr(submission, 'document_category', None) or 'document',
            }
            try:
                from services.layer_tags import tags_for_subject
                from models.layer_tag import SUBJECT_SUBMISSION
                draft['tags'] = tags_for_subject(SUBJECT_SUBMISSION, submission.id)
            except Exception:
                draft['tags'] = []

    if not draft:
        return None, None

    if not submission:
        submission = lookup(draft.get('name')) or lookup(draft_name)
        if submission:
            dbs = getattr(submission, 'displayBodySource', None) or 'file'
            draft['displayBodySource'] = dbs
            draft['displayOrdinalId'] = getattr(submission, 'displayOrdinalId', None)
            draft['displayingLinkedOrdinal'] = (
                dbs.strip().lower() == 'ordinal'
                and bool(getattr(submission, 'displayOrdinalContentUrl', None))
            )
            draft['document_category'] = getattr(submission, 'document_category', None) or draft.get(
                'document_category', 'document'
            )
            try:
                from services.layer_tags import tags_for_subject
                from models.layer_tag import SUBJECT_SUBMISSION
                draft['tags'] = tags_for_subject(SUBJECT_SUBMISSION, submission.id)
            except Exception:
                draft['tags'] = draft.get('tags') or []

    return draft, submission


def draft_display_id(draft: dict) -> str:
    if draft.get('status') == 'approved' and draft.get('ml_number'):
        return str(draft['ml_number'])
    return str(draft.get('name') or '')


def load_draft_document_body(
    draft: dict,
    submission,
    draft_name: str,
    *,
    pdf_iframe_height: str = '800px',
) -> Tuple[str, bool, int, int]:
    """
    Load body HTML/text for record or reader view.
    Returns (content, render_as_html, pages, words).
    """
    document_content = 'Document content not available.'
    calculated_pages = draft.get('pages', 1)
    calculated_words = draft.get('words', 0)
    render_document_as_html = False

    if submission and _submission_uses_display_ordinal(submission):
        return _load_ordinal_body(
            submission.displayOrdinalContentUrl,
            submission.displayOrdinalContentType or '',
            draft,
        )

    if submission and draft.get('sourceType') == 'ordinal':
        return _load_ordinal_body(
            getattr(submission, 'ordinalContentUrl', None),
            getattr(submission, 'ordinalContentType', '') or '',
            draft,
        )

    if submission and submission.file_path and os.path.exists(submission.file_path):
        _, ext = os.path.splitext((submission.filename or '').lower())
        try:
            if ext in ['.txt', '.xml', '.md', '.markdown', '.htm', '.html']:
                with open(submission.file_path, 'r', encoding='utf-8', errors='replace') as f:
                    raw_text = f.read()
                calculated_words = len(raw_text.split())
                calculated_pages = max(1, (calculated_words + 499) // 500)
                if ext in ('.htm', '.html') or looks_like_html_document(raw_text):
                    document_content = load_book_html_for_reader(raw_text)
                    render_document_as_html = True
                elif ext in ('.md', '.markdown'):
                    md_html = markdown_to_safe_preview_html(raw_text)
                    if md_html:
                        document_content = md_html
                        render_document_as_html = True
                    else:
                        document_content = raw_text
                elif ext == '.txt':
                    if text_looks_like_markdown(raw_text):
                        md_html = markdown_to_safe_preview_html(raw_text)
                        if md_html:
                            document_content = md_html
                            render_document_as_html = True
                        else:
                            document_content = raw_text
                    else:
                        document_content = raw_text
                else:
                    document_content = raw_text
            elif ext == '.docx':
                from services.draft_reader import docx_body_to_safe_html

                try:
                    html_body, docx_words = docx_body_to_safe_html(submission.file_path)
                except Exception as docx_exc:
                    document_content = f'Error loading .docx content: {docx_exc}'
                else:
                    document_content = html_body
                    calculated_words = docx_words
                    calculated_pages = max(1, (calculated_words + 499) // 500)
                    render_document_as_html = True
            elif ext == '.pdf':
                from PyPDF2 import PdfReader

                reader = PdfReader(submission.file_path)
                calculated_pages = len(reader.pages) if reader.pages else 1
                calculated_words = calculated_pages * 275
                file_size_kb = os.path.getsize(submission.file_path) / 1024
                render_document_as_html = True
                document_content = f'''
<div class="pdf-viewer-container">
    <div class="alert alert-info mb-3">
        <i class="bi bi-file-pdf"></i> PDF Document ({calculated_pages} pages, ~{calculated_words} words, {file_size_kb:.1f} KB)
    </div>
    <iframe src="/view/{escape(draft_name, quote=True)}"
            type="application/pdf"
            style="width: 100%; height: {pdf_iframe_height}; border: 1px solid var(--card-border); border-radius: 4px;"
            title="PDF Document Viewer">
        <p>Your browser does not support PDF preview.
           <a href="/download/{escape(draft_name, quote=True)}">Download the PDF</a> to view it.</p>
    </iframe>
</div>
'''
            else:
                document_content = (
                    f'Document content cannot be displayed for {ext.upper()} files. Please download to view.'
                )
        except Exception as exc:
            document_content = f'Error loading document content: {exc}'

        draft['pages'] = calculated_pages
        draft['words'] = calculated_words
        return document_content, render_document_as_html, calculated_pages, calculated_words

    if draft and draft.get('name'):
        authors = draft.get('authors') or []
        document_content = f"""INTERNET-DRAFT                                               {', '.join(authors)}
Intended status: Informational                            Meta-Layer Initiative
Expires: {draft.get('date', 'TBD')}                                      {draft.get('date', 'TBD')}


{draft.get('title', 'Document Title')}


Abstract

{draft.get('abstract', 'Abstract not available.')}


1. Introduction

This document describes {draft.get('title', 'the subject matter')}.

The content of this draft is currently being developed and will be available
in the full document once published.
"""
        return document_content, False, calculated_pages, calculated_words

    return document_content, render_document_as_html, calculated_pages, calculated_words


def docx_body_to_safe_html(docx_path: str) -> Tuple[str, int]:
    """
    Convert a .docx file's body to a sanitized HTML fragment that preserves
    document order (paragraphs and tables interleaved, in source order) and
    renders tables as real <table> markup.

    Returns (html, word_count). HTML is built by hand-escaping every text node
    with html.escape, so no user-controlled content can be interpreted as HTML.

    Heading-style paragraphs are rendered as <h1>..<h6>; everything else is
    <p>...</p>. Empty paragraphs are dropped. The first row of a table is
    rendered as <th> cells (header convention).
    """
    from html import escape as _escape

    from docx import Document as _DocxDocument
    from docx.oxml.ns import qn as _qn
    from docx.text.paragraph import Paragraph as _Paragraph
    from docx.table import Table as _Table

    doc = _DocxDocument(docx_path)
    body = doc.element.body

    # Build a fast lookup of which tables are header-styled (first row = <th>).
    # Heuristic: a table is treated as a header table when its first row's
    # cells are bold OR its style name contains "Header"/"Table Grid".
    def _is_header_table(tbl) -> bool:
        try:
            first = tbl.rows[0]
            cells = [c for c in first.cells]
            if not cells:
                return False
            # All-bold check (any cell not bold → not header)
            for c in cells:
                for para in c.paragraphs:
                    for run in para.runs:
                        if not run.bold:
                            return False
            return True
        except Exception:
            return False

    def _cell_shading(tc) -> str:
        """Return a hex color string if the cell has explicit shading, else ''."""
        try:
            tcPr = tc.find(_qn('w:tcPr'))
            if tcPr is None:
                return ''
            shd = tcPr.find(_qn('w:shd'))
            if shd is None:
                return ''
            fill = shd.get(_qn('w:fill')) or ''
            fill = fill.strip().lower()
            # 'auto' / 'FFFFFF' treated as no shading to avoid white-on-white
            if not fill or fill in ('auto', 'ffffff'):
                return ''
            return fill
        except Exception:
            return ''

    def _cell_style_attrs(tc) -> str:
        """Inline style attributes for a cell (background-color, color).

        Light fills (perceived luminance > 0.6) keep dark text; dark fills keep
        light text. Auto/white fills are skipped so we don't override the page.
        """
        fill = _cell_shading(tc)
        if not fill:
            return ''
        try:
            r = int(fill[0:2], 16)
            g = int(fill[2:4], 16)
            b = int(fill[4:6], 16)
        except ValueError:
            return ''
        luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255.0
        text_color = '#0a1628' if luminance > 0.6 else '#eef2ff'
        # Use 90% alpha tint for the dark theme so the row blends with the
        # navy card background instead of glowing against it.
        style = f'background-color:#{fill};color:{text_color};'
        return style

    parts: list = []
    word_count = 0

    for child in body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            para = _Paragraph(child, doc)
            text = para.text or ''
            stripped = text.strip()
            if not stripped:
                continue
            word_count += len(stripped.split())
            style_name = (para.style.name or '').lower() if para.style else ''
            safe = _escape(stripped, quote=False)
            if style_name.startswith('heading'):
                # "Heading 1" → h1
                level = ''.join(ch for ch in style_name if ch.isdigit()) or '2'
                try:
                    n = max(1, min(6, int(level)))
                except ValueError:
                    n = 2
                parts.append(f'<h{n}>{safe}</h{n}>')
            elif style_name in ('title', 'subtitle'):
                tag_name = 'h1' if style_name == 'title' else 'h2'
                parts.append(f'<{tag_name}>{safe}</{tag_name}>')
            else:
                parts.append(f'<p>{safe}</p>')
        elif tag == 'tbl':
            tbl = _Table(child, doc)
            try:
                rows = list(tbl.rows)
            except Exception:
                continue
            if not rows:
                continue
            header = _is_header_table(tbl)
            # Normalize each row to the same column count so the table is rectangular.
            col_count = 0
            for row in rows:
                col_count = max(col_count, len(row.cells))
            if col_count == 0:
                continue
            # Drop `table-striped` so Bootstrap doesn't recolor odd rows with
            # its light-mode body color (`#212529`), which is unreadable on the
            # dark card background. `table-bordered` keeps visible cell edges.
            tbl_html = ['<table class="table table-bordered docx-table">']
            if header:
                tbl_html.append('<thead>')
            else:
                tbl_html.append('<tbody>')
            for ri, row in enumerate(rows):
                cells = list(row.cells)
                # python-docx repeats the same cell object for horizontally-merged
                # grid spans. Dedupe by identity to avoid printing the same content
                # multiple times in a single row.
                seen = set()
                unique_cells = []
                for c in cells:
                    key = id(c._tc)
                    if key in seen:
                        continue
                    seen.add(key)
                    unique_cells.append(c)
                row_tag = 'th' if (header and ri == 0) else 'td'
                tbl_html.append('<tr>')
                for ci in range(col_count):
                    cell = unique_cells[ci] if ci < len(unique_cells) else None
                    if cell is None:
                        tbl_html.append(f'<{row_tag}></{row_tag}>')
                        continue
                    style_attrs = _cell_style_attrs(cell._tc)
                    style_part = f' style="{style_attrs}"' if style_attrs else ''
                    # Render each paragraph in the cell as a separate line.
                    inner = []
                    for p in cell.paragraphs:
                        t = (p.text or '').strip()
                        if not t:
                            continue
                        word_count += len(t.split())
                        inner.append(_escape(t, quote=False))
                    cell_html = '<br>'.join(inner) if inner else ''
                    tbl_html.append(f'<{row_tag}{style_part}>{cell_html}</{row_tag}>')
                tbl_html.append('</tr>')
            if header:
                tbl_html.append('</thead><tbody>')
                tbl_html.append('</tbody>')
            else:
                tbl_html.append('</tbody>')
            tbl_html.append('</table>')
            parts.append(''.join(tbl_html))
        # sectPr / sdt / other nodes → ignored (page breaks, content controls)

    return ''.join(parts), word_count
