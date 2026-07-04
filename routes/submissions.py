"""Submissions API and actions: list, approve, reject, view, download, admin/submissions."""
import html as html_mod
import os
import random
import re
import string
from datetime import datetime
from typing import Optional

import requests
from flask import Blueprint, jsonify, request, redirect, url_for, flash, send_file, current_app, session, render_template_string, g

from extensions import db
from models import Submission, Layer
from services.identity import get_current_user, require_auth, require_role
from services.submissions import get_submission_by_ref, add_to_document_history, get_next_ml_number, can_edit_submission_metadata
from services.documents import load_draft_data, revision_notes_to_safe_html
from services.directory_ui import gh_page_header, gh_breadcrumb, gh_living_module
from services.ordinals import (
    shorten_inscription_id,
    looks_like_html_inscription,
    format_ordinal_html_iframe_preview,
)
from services.utils import coerce_storage_bool
from services.submission_uploads import save_submission_upload
from services.url_safety import validate_ordinals_fetch_url

bp = Blueprint('submissions', __name__, url_prefix='')


def _submission_is_revision(submission) -> bool:
    """SQLite may store is_revision as TEXT '0'/'1'; bool('0') is True in Python."""
    return coerce_storage_bool(getattr(submission, 'is_revision', False), default=False)


def _strip_immortalize_from_submit_template(template: str) -> str:
    """Remove Immortalize tab nav and pane when product rollout has immortalize off."""
    template = re.sub(
        r'<!-- GH_IMMORTALIZE_NAV -->.*?<!-- /GH_IMMORTALIZE_NAV -->',
        '',
        template,
        flags=re.DOTALL,
    )
    template = re.sub(
        r'<!-- GH_IMMORTALIZE_PANE -->.*?<!-- /GH_IMMORTALIZE_PANE -->',
        '',
        template,
        flags=re.DOTALL,
    )
    return template


def _layer_prefix_for_submission(submission) -> str:
    """Return the 2-letter draft prefix for a submission's primary layer.

    Honours an explicit per-draft override on ``submission.prefix_code`` —
    the submit form sets this when a layer has more than one prefix and the
    author chose a non-default code. Falls back to the layer's default
    ``LayerPrefix`` row, then to 'ML' as a last resort.
    """
    if submission is None:
        return 'ML'
    override = (getattr(submission, 'prefix_code', None) or '').strip().upper()
    if override and _is_valid_prefix_code(override):
        return override
    layer_id = (
        getattr(submission, 'primary_layer_id', None)
        or getattr(submission, 'layer_id', None)
    )
    if not layer_id:
        return 'ML'
    try:
        from models import LayerPrefix  # untracked WIP model
        row = LayerPrefix.query.filter_by(layer_id=layer_id, is_default=True).first()
        if row is None:
            return 'ML'
        prefix = (getattr(row, 'prefix', None) or '').strip().upper()
        return prefix or 'ML'
    except Exception:
        return 'ML'


def _is_valid_prefix_code(value: object) -> bool:
    """Two uppercase ASCII letters, mirroring services.layer_prefixes format."""
    import re as _re
    return bool(_re.match(r'^[A-Z]{2}$', (str(value) if value is not None else '').strip().upper()))


# ---------------------------------------------------------------------------
# Submission form routes (submit_draft, submit_revision, submission_status, submission_detail)
# ---------------------------------------------------------------------------

SUBMISSION_STATUS_TEMPLATE = """
<div class="gh-page container mt-4">
    <nav aria-label="breadcrumb" class="gh-detail-breadcrumb mb-3">
        <ol class="breadcrumb">
            <li class="breadcrumb-item"><a href="{{ home_url }}">Home</a></li>
            <li class="breadcrumb-item"><a href="{{ submit_url }}">Submit Draft</a></li>
            <li class="breadcrumb-item active">Submission Status</li>
        </ol>
    </nav>

    <header class="gh-page-header">
        <div class="gh-page-header-main">
            <div class="gh-page-header-icon"><i class="fas fa-clipboard-check"></i></div>
            <div>
                <h1 class="gh-page-title">{{ status_page_heading }}</h1>
                <p class="gh-page-lead">Track your Internet-Draft submission</p>
            </div>
        </div>
    </header>

    <div id="flash-messages"></div>

    <div class="row">
        <div class="col-md-8">
            <div class="living-module mb-4">
                <div class="living-module-header">
                    <div class="living-module-icon"><i class="fas fa-file-alt"></i></div>
                    <h5 class="living-module-title">
                        Submission Details
                        {% if is_revision and revision_number %}
                        <span class="badge bg-success ms-2">Revision {{ revision_number }}</span>
                        {% endif %}
                    </h5>
                </div>
                <div class="living-module-body">
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Submission ID:</strong></div>
                        <div class="col-sm-9"><code>{{ submission.id }}</code></div>
                    </div>
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Status:</strong></div>
                        <div class="col-sm-9">
                            <span class="badge bg-primary">{{ submission_status_title }}</span>
                            {% if is_ordinal %}
                            <span class="badge bg-info ms-2"><i class="bi bi-coin"></i> Ordinal</span>
                            {% else %}
                            <span class="badge bg-secondary ms-2"><i class="bi bi-file-earmark"></i> File</span>
                            {% endif %}
                        </div>
                    </div>
                    {% if is_revision and revision_number and parent_draft_name %}
                    <div class="alert alert-info mb-3">
                        <strong><i class="fas fa-code-branch me-2"></i>This is a revision</strong><br>
                        Revision <strong>{{ revision_number }}</strong> of
                        <a href="{{ parent_draft_url }}">{{ parent_draft_name }}</a>
                    </div>
                    {% endif %}
                    {% if what_changed_html %}
                    <div class="card mb-3">
                        <div class="card-header">
                            <strong>What changed (submitter's explanation)</strong>
                        </div>
                        <div class="card-body revision-notes">
                            {{ what_changed_html | safe }}
                        </div>
                    </div>
                    {% endif %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Title:</strong></div>
                        <div class="col-sm-9">{{ submission_title }}</div>
                    </div>
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Authors:</strong></div>
                        <div class="col-sm-9">{{ submission_authors_joined }}</div>
                    </div>
                    {% if ml_number %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>ML Number:</strong></div>
                        <div class="col-sm-9"><code>{{ ml_number }}</code></div>
                    </div>
                    {% endif %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Draft ID:</strong></div>
                        <div class="col-sm-9"><code>{{ submission_id }}</code></div>
                    </div>
                    {% if artifact_id and artifact_layer_slug %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Artifact:</strong></div>
                        <div class="col-sm-9"><a href="/layers/{{ artifact_layer_slug }}/artifacts/{{ artifact_id }}/" class="btn btn-outline-secondary btn-sm">View Artifact</a></div>
                    </div>
                    {% endif %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Submitted:</strong></div>
                        <div class="col-sm-9">{{ submission_submitted_at }}</div>
                    </div>
                    {% if is_file %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>File:</strong></div>
                        <div class="col-sm-9">
                            <code>{{ submission_filename }}</code>
                            <a href="{{ download_url }}" class="btn btn-sm btn-outline-primary ms-2">Download</a>
                        </div>
                    </div>
                    {% endif %}
                    {% if submission_abstract %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Abstract:</strong></div>
                        <div class="col-sm-9">{{ submission_abstract }}</div>
                    </div>
                    {% endif %}

                    {% if is_ordinal %}
                    <h6 class="mt-4">Ordinal Metadata</h6>
                    <div class="card mb-3" style="background-color: var(--bg-secondary);">
                        <div class="card-body">
                            <div class="row mb-2">
                                <div class="col-sm-4"><strong>Inscription ID:</strong></div>
                                <div class="col-sm-8">
                                    <a href="https://ordinals.com/inscription/{{ ordinal_id }}" target="_blank" class="text-decoration-none" style="color: var(--accent-color);">
                                        <code style="font-size: 0.85em;">{{ ordinal_id_short }}</code>
                                    </a>
                                </div>
                            </div>
                            {% if inscription_number %}
                            <div class="row mb-2">
                                <div class="col-sm-4"><strong>Inscription Number:</strong></div>
                                <div class="col-sm-8">{{ inscription_number }}</div>
                            </div>
                            {% endif %}
                            {% if block_height %}
                            <div class="row mb-2">
                                <div class="col-sm-4"><strong>Block Height:</strong></div>
                                <div class="col-sm-8">{{ block_height }}</div>
                            </div>
                            {% endif %}
                            {% if inscription_timestamp %}
                            <div class="row mb-2">
                                <div class="col-sm-4"><strong>Timestamp:</strong></div>
                                <div class="col-sm-8">{{ inscription_timestamp }}</div>
                            </div>
                            {% endif %}
                            <div class="row mb-2">
                                <div class="col-sm-4"><strong>Content Type:</strong></div>
                                <div class="col-sm-8"><code>{{ ordinal_content_type }}</code></div>
                            </div>
                            <div class="row">
                                <div class="col-sm-12">
                                    <a href="https://ordinals.com/inscription/{{ ordinal_id }}" target="_blank" class="btn btn-sm btn-outline-primary">
                                        <i class="bi bi-box-arrow-up-right"></i> View on Ordinals.com
                                    </a>
                                </div>
                            </div>
                        </div>
                    </div>
                    {% endif %}

                    <h6 class="mt-4">Content Preview</h6>
                    {% if content_preview_html %}
                    <div class="border rounded p-3" style="background-color: var(--input-bg); border-color: var(--input-border);">
                        {{ content_preview_html|safe }}
                    </div>
                    {% else %}
                    <div class="border rounded p-3" style="background-color: var(--input-bg); border-color: var(--input-border);">
                        <pre class="mb-0" style="font-size: 0.9em; max-height: 400px; overflow-y: auto; color: var(--text-primary);">{{ file_content }}</pre>
                    </div>
                    {% endif %}

                    {% if is_submitted and is_admin %}
                    <div class="row mb-3">
                        <div class="col-sm-3"><strong>Actions:</strong></div>
                        <div class="col-sm-9">
                            <form method="POST" action="{{ approve_url }}" style="display: inline;">
                                <button type="submit" class="btn btn-success btn-sm">Approve & Publish</button>
                            </form>
                            <form method="POST" action="{{ reject_url }}" style="display: inline; margin-left: 10px;">
                                <button type="submit" class="btn btn-danger btn-sm">Reject</button>
                            </form>
                        </div>
                    </div>
                    {% endif %}
                </div>
            </div>

            <div class="card mt-3">
                <div class="card-header">
                    <h5>Review Timeline</h5>
                </div>
                <div class="card-body">
                    <div class="timeline">
                        <div class="timeline-item">
                            <div class="timeline-marker bg-success"></div>
                            <div class="timeline-content">
                                <h6>Submitted</h6>
                                <p class="text-muted small">{{ submission_submitted_at }}</p>
                            </div>
                        </div>
                        <div class="timeline-item">
                            {% if is_approved_or_rejected %}
                            <div class="timeline-marker bg-success"></div>
                            <div class="timeline-content">
                                <h6>Initial Review</h6>
                                <p class="text-muted small">
                                    {% if is_approved %}Completed{% else %}Rejected{% endif %}
                                    {% if submission_approved_at %}
                                    - {{ submission_approved_at }}
                                    {% elif submission_rejected_at %}
                                    - {{ submission_rejected_at }}
                                    {% endif %}
                                </p>
                            </div>
                            {% else %}
                            <div class="timeline-marker bg-secondary"></div>
                            <div class="timeline-content">
                                <h6>Initial Review</h6>
                                <p class="text-muted small">In Progress</p>
                            </div>
                            {% endif %}
                        </div>
                        {% if is_approved %}
                        <div class="timeline-item">
                            <div class="timeline-marker bg-primary"></div>
                            <div class="timeline-content">
                                <h6>Published</h6>
                                <p class="text-muted small">Available in document repository</p>
                            </div>
                        </div>
                        {% else %}
                        <div class="timeline-item">
                            <div class="timeline-marker bg-secondary"></div>
                            <div class="timeline-content">
                                <h6>Workgroup Review</h6>
                                <p class="text-muted small">Pending initial approval</p>
                            </div>
                        </div>
                        <div class="timeline-item">
                            <div class="timeline-marker bg-secondary"></div>
                            <div class="timeline-content">
                                <h6>MLSG Review</h6>
                                <p class="text-muted small">Pending workgroup review</p>
                            </div>
                        </div>
                        {% endif %}
                    </div>
                </div>
            </div>
        </div>

        <div class="col-md-4">
            <div class="card">
                <div class="card-header">
                    <h5>Actions</h5>
                </div>
                <div class="card-body">
                    <a href="{{ submit_url }}" class="btn btn-primary w-100 mb-2">Submit Another Draft</a>
                    <a href="{{ doc_all_url }}" class="btn btn-outline-secondary w-100 mb-2">View All Documents</a>
                    <a href="{{ home_url }}" class="btn btn-outline-secondary w-100">Back to Home</a>
                </div>
            </div>

            <div class="card mt-3">
                <div class="card-header">
                    <h5>Need Help?</h5>
                </div>
                <div class="card-body">
                    <p class="small">If you have questions about your submission:</p>
                    <ul class="small">
                        <li>Check the <a href="#" target="_blank">submission guidelines</a></li>
                        <li>Contact the <a href="mailto:info@metalayer.org">GovHub</a></li>
                        <li>Join the <a href="#" target="_blank">MLGH discussion list</a></li>
                    </ul>
                </div>
            </div>
        </div>
    </div>
</div>

<style>
.timeline {
    position: relative;
    padding-left: 30px;
}

.timeline-item {
    position: relative;
    margin-bottom: 20px;
}

.timeline-marker {
    position: absolute;
    left: -25px;
    top: 5px;
    width: 12px;
    height: 12px;
    border-radius: 50%;
    border: 2px solid #fff;
    box-shadow: 0 0 0 2px #dee2e6;
}

.timeline-content h6 {
    margin-bottom: 5px;
    font-weight: 600;
}

.timeline-content p {
    margin-bottom: 0;
}
</style>
"""


def _submit_document_meta_fields_html(
    *,
    selected_category: str = '',
    tags_value: str = '',
    compact: bool = False,
) -> str:
    from services.document_categories import document_category_options_html
    import html as html_mod

    tags_esc = html_mod.escape(tags_value or '', quote=True)
    if compact:
        return f'''
            <div class="mb-3">
                <label class="form-label">Document type</label>
                <select class="form-select" name="document_category">
                    {document_category_options_html(selected_category)}
                </select>
            </div>
            <div class="mb-3">
                <label class="form-label">Tags <span class="text-muted">(optional)</span></label>
                <input type="text" class="form-control" name="document_tags" value="{tags_esc}"
                       placeholder="governance, climate-policy (comma-separated)">
                <small class="form-text text-muted">Up to 10 layer tags (shared with artifacts).</small>
            </div>'''
    return f'''
                                <div class="mb-3">
                                    <label for="document_category" class="form-label">Document type</label>
                                    <select class="form-select" id="document_category" name="document_category">
                                        {document_category_options_html(selected_category)}
                                    </select>
                                </div>
                                <div class="mb-3">
                                    <label for="document_tags" class="form-label">Tags <span class="text-muted">(optional)</span></label>
                                    <input type="text" class="form-control" id="document_tags" name="document_tags"
                                           value="{tags_esc}"
                                           placeholder="governance, climate-policy (comma-separated)">
                                    <div class="form-text">Up to 10 layer tags. Same vocabulary as artifacts on this layer.</div>
                                </div>'''


def _apply_submission_document_meta(submission, form, user_id: Optional[str]) -> None:
    from flask import current_app
    from services.document_categories import normalize_document_category
    from services.layer_tags import (
        document_tags_enabled,
        parse_tag_slugs,
        set_submission_tags,
        sync_submission_tags_to_artifact,
    )

    submission.document_category = normalize_document_category(
        form.get('document_category') if hasattr(form, 'get') else None
    )
    if document_tags_enabled(current_app.config) and submission.layer_id:
        raw = form.get('document_tags', '') if hasattr(form, 'get') else ''
        slugs = parse_tag_slugs(raw)
        set_submission_tags(submission, slugs, user_id)
        sync_submission_tags_to_artifact(submission)


def _build_prefix_selector_inner_html(layer_id, layer_prefixes):
    """Inner markup for the per-draft prefix selector, per branch.

    The wrapper ``<div id="submit-prefix-selector-wrap">`` is the contract
    used by ``_GhRefreshSubmitPrefixSelector`` (and the workgroup_links.js
    change handler) and is emitted by ``_build_prefix_selector_html`` —
    helpers here return only the *contents* of that wrapper.
    """
    if not layer_id:
        return (
            '<label class="form-label" data-gh-i18n="prefix.label">Prefix</label>'
            '<div id="submit-prefix-selector-body" class="small text-muted">'
            'Select a layer to see available prefixes.'
            '</div>'
        )

    if not layer_prefixes:
        # Zero-prefixes fallback: render the system default ``ML`` inline and
        # tell the user it's the auto-fallback (not an admin-misconfiguration
        # error). The hidden ``#upload-prefix-code`` / ``#ordinal-prefix-code``
        # fields are populated server-side via the parent wrapper so the form
        # still submits a valid prefix without an extra round-trip.
        return (
            '<label class="form-label" data-gh-i18n="prefix.label">Prefix</label>'
            '<div class="d-flex align-items-center gap-2">'
            '<span class="font-monospace fs-5 fw-bold">ML</span>'
            '<span class="text-muted small">No prefixes configured for this layer — '
            "drafts will use the system default <code>ML</code>.</span>"
            '</div>'
            '<div class="form-text">Layer admins can add additional prefixes from '
            "the layer's Admin → Prefixes card.</div>"
        )

    if len(layer_prefixes) == 1:
        p = layer_prefixes[0]
        code = html_mod.escape(p.prefix or "")
        return (
            '<label class="form-label" data-gh-i18n="prefix.label">Prefix</label>'
            '<div class="d-flex align-items-center gap-2">'
            f'<span class="font-monospace fs-5 fw-bold">{code}</span>'
            '<span class="badge bg-success">Default</span>'
            '</div>'
            '<div class="form-text">This is the only prefix for this layer; the new '
            'draft will use it automatically.</div>'
        )

    # >1 prefix → dropdown options
    options = []
    for p in layer_prefixes:
        label = p.prefix
        if p.is_default:
            label += ' (default)'
        sel = ' selected' if p.is_default else ''
        options.append(
            f'<option value="{html_mod.escape(p.prefix or "")}"{sel}>'
            f'{html_mod.escape(label)}</option>'
        )
    return (
        '<label for="submit-prefix-select" class="form-label" '
        'data-gh-i18n="prefix.label">Prefix</label>'
        '<select class="form-select" id="submit-prefix-select" '
        'style="max-width: 14rem;">'
        f'{"".join(options)}'
        '</select>'
        '<div class="form-text">This layer has more than one prefix — pick which '
        "code to use for this draft's identifier.</div>"
    )


def _gh_prefix_selector_refresh_script():
    """Inline script that defines ``window._GhRefreshSubmitPrefixSelector``.

    The wrapper div is always present (emitted by
    ``_build_prefix_selector_html``); on layer change the change handler
    in ``services/workgroup_links.py`` calls this function which fetches
    the live prefixes and rebuilds the wrapper content for all three
    outcomes (zero / one / many), then keeps the hidden
    ``#upload-prefix-code`` and ``#ordinal-prefix-code`` fields in sync.
    """
    return (
        '<script>(function(){\n'
        'function _syncPrefixFields(val) {\n'
        '  var u=document.getElementById("upload-prefix-code");\n'
        '  if (u) u.value = val || "";\n'
        '  var o=document.getElementById("ordinal-prefix-code");\n'
        '  if (o) o.value = val || "";\n'
        '}\n'
        'function _clearPrefixFields() { _syncPrefixFields(""); }\n'
        'window._GhRefreshSubmitPrefixSelector = function(layerId) {\n'
        '  var wrap = document.getElementById("submit-prefix-selector-wrap");\n'
        '  if (!wrap) return;\n'
        '  if (!layerId) {\n'
        '    wrap.style.display = "none";\n'
        '    wrap.setAttribute("data-prefix-state", "no-layer");\n'
        '    wrap.innerHTML = \'<label class="form-label" data-gh-i18n="prefix.label">Prefix</label>\'\n'
        '      + \'<div id="submit-prefix-selector-body" class="small text-muted">Select a layer to see available prefixes.</div>\';\n'
        '    _clearPrefixFields();\n'
        '    return;\n'
        '  }\n'
        '  fetch("/api/layers/" + encodeURIComponent(layerId) + "/prefixes/", {credentials:"same-origin"})\n'
        '    .then(function(r){ return r.json(); })\n'
        '    .then(function(data){\n'
        '      var items = (data && data.prefixes) || [];\n'
        '      if (!items.length) {\n'
        '        wrap.setAttribute("data-prefix-state", "default");\n'
        '        wrap.innerHTML = \'<label class="form-label">Prefix</label>\'\n'
        '          + \'<div class="d-flex align-items-center gap-2">\'\n'
        '          + \'<span class="font-monospace fs-5 fw-bold">ML</span>\'\n'
        '          + \'<span class="text-muted small">No prefixes configured for this layer &mdash; drafts will use the system default <code>ML</code>.</span>\'\n'
        '          + \'</div>\';\n'
        '        wrap.style.display = "";\n'
        '        _syncPrefixFields("ML");\n'
        '        return;\n'
        '      }\n'
        '      if (items.length === 1) {\n'
        '        wrap.setAttribute("data-prefix-state", "single");\n'
        '        wrap.innerHTML = \'<label class="form-label">Prefix</label>\'\n'
        '          + \'<div class="d-flex align-items-center gap-2">\'\n'
        '          + \'<span class="font-monospace fs-5 fw-bold">\' + (items[0].prefix || "") + \'</span>\'\n'
        '          + \'<span class="badge bg-success">Default</span></div>\'\n'
        '          + \'<div class="form-text">This is the only prefix for this layer; the new draft will use it automatically.</div>\';\n'
        '        wrap.style.display = "";\n'
        '        _syncPrefixFields(items[0].prefix || "");\n'
        '        return;\n'
        '      }\n'
        '      var html = \'<label for="submit-prefix-select" class="form-label">Prefix</label>\'\n'
        '        + \'<select class="form-select" id="submit-prefix-select" style="max-width: 14rem;">\';\n'
        '      var def = items.find(function(x){ return x.is_default; }) || items[0];\n'
        '      items.forEach(function(p){\n'
        '        var lbl = p.prefix + (p.is_default ? " (default)" : "");\n'
        '        var sel = p.is_default ? " selected" : "";\n'
        '        html += \'<option value="\' + p.prefix + \'"\' + sel + \'>\' + lbl + \'</option>\';\n'
        '      });\n'
        '      html += \'</select><div class="form-text">This layer has more than one prefix — pick which code to use for this draft\\u2019s identifier.</div>\';\n'
        '      wrap.setAttribute("data-prefix-state", "multi");\n'
        '      wrap.innerHTML = html;\n'
        '      wrap.style.display = "";\n'
        '      var newSel = document.getElementById("submit-prefix-select");\n'
        '      _syncPrefixFields(def.prefix || "");\n'
        '      if (newSel) newSel.addEventListener("change", function(){ _syncPrefixFields(newSel.value); });\n'
        '    })\n'
        '    .catch(function(e){ console.warn("prefix fetch failed", e); });\n'
        '};\n'
        # If the wrapper was rendered with a concrete layer already chosen
        # (single or multi), wire the dropdown-or-hidden-field sync to that
        # initial value so the form has the prefix code even before any
        # user interaction.
        'function _initStaticSelectors() {\n'
        '  var wrap = document.getElementById("submit-prefix-selector-wrap");\n'
        '  if (!wrap) return;\n'
        '  var state = wrap.getAttribute("data-prefix-state");\n'
        '  var sel = document.getElementById("submit-prefix-select");\n'
        '  if (sel) {\n'
        '    _syncPrefixFields(sel.value);\n'
        '    sel.addEventListener("change", function(){ _syncPrefixFields(sel.value); });\n'
        '    return;\n'
        '  }\n'
        '  if (state === "single" || state === "default") {\n'
        '    // "single" → layer has exactly one configured prefix.\n'
        '    // "default" → layer has zero configured prefixes; server stamped\n'
        '    // the system code ``ML`` into data-prefix-default as the\n'
        '    // fallback. Either way, mirror it into the hidden form fields.\n'
        '    var def = wrap.getAttribute("data-prefix-default") || "";\n'
        '    _syncPrefixFields(def);\n'
        '  }\n'
        '}\n'
        'if (document.readyState==="loading") {\n'
        ' document.addEventListener("DOMContentLoaded", _initStaticSelectors);\n'
        '} else { _initStaticSelectors(); }\n'
        '})();</script>'
    )


def _build_prefix_selector_html(layer_id, layer_prefixes):
    """Per-draft prefix selector for the submit form.

    0 prefixes: warning (admin hasn't set one yet).
    1 prefix:   read-only badge (drafts always use this code).
    >1 prefix:  dropdown — author picks the code for this draft.
    The selection is mirrored into the hidden #upload-prefix-code and
    #ordinal-prefix-code fields on change so the server can read it.

    The wrapper div and the ``_GhRefreshSubmitPrefixSelector`` script are
    ALWAYS emitted, regardless of branch, so the layer-change handler in
    ``services/workgroup_links.py`` can always re-render the selector.
    """
    if not layer_id:
        state = 'no-layer'
        display_style = ' style="display:none;"'
    elif not layer_prefixes:
        # Zero-configured-prefixes fallback: render the system default ``ML``
        # inline. The hidden ``#upload-prefix-code`` / ``#ordinal-prefix-code``
        # fields default to ``ML`` via ``data-prefix-default`` so the static
        # init script below syncs them without an extra round-trip.
        state = 'default'
        display_style = ''
    elif len(layer_prefixes) == 1:
        state = 'single'
        display_style = ''
    else:
        state = 'multi'
        display_style = ''

    if state == 'single':
        default_attr = f' data-prefix-default="{html_mod.escape(layer_prefixes[0].prefix or "")}"'
    elif state == 'default':
        # Server-side fallback for layers with zero configured prefixes:
        # the hidden #upload-prefix-code / #ordinal-prefix-code fields
        # default to the system code ``ML`` (the static init script picks
        # this up via getAttribute("data-prefix-default")).
        default_attr = ' data-prefix-default="ML"'
    else:
        default_attr = ''

    inner_html = _build_prefix_selector_inner_html(layer_id, layer_prefixes)
    refresh_script = _gh_prefix_selector_refresh_script()
    return (
        f'<div class="mb-3" id="submit-prefix-selector-wrap" '
        f'data-prefix-state="{state}"{display_style}{default_attr}>'
        f'{inner_html}'
        f'{refresh_script}'
        '</div>'
    )


def _build_submit_form_template(
    *,
    effective_layer,
    layers,
    selected_group: str = '',
    build_number: int,
    empty_state_message: str = 'No approved layers available. Submit from a layer subdomain (e.g. overweb.themetalayer.org) or create a layer first.',
) -> str:
    """Build submit page HTML with DB-backed workgroup options for the layer."""
    from templates.html_templates import SUBMIT_TEMPLATE
    from services.workgroup_links import (
        submit_workgroup_layer_script,
        workgroup_select_options_html,
    )
    from services.product_rollout import is_feature_enabled
    from services.page_heroes import render_page_hero_html
    from services.layer_prefixes import list_prefixes

    layer_id = effective_layer.id if effective_layer else None
    group_options = workgroup_select_options_html(layer_id, selected_group)
    workgroup_script = submit_workgroup_layer_script(fixed_layer_id=layer_id)

    # Build the per-draft prefix selector. 0 prefixes → warning; 1 → read-only;
    # >1 → dropdown. The server re-validates on submit, so this is just UX.
    layer_prefixes = list_prefixes(layer_id) if layer_id else []
    prefix_selector_html = _build_prefix_selector_html(layer_id, layer_prefixes)

    submit_template = SUBMIT_TEMPLATE.replace('{{WORKGROUP_OPTIONS}}', group_options)
    submit_template = submit_template.replace('{{WORKGROUP_LAYER_SCRIPT}}', workgroup_script)
    submit_template = submit_template.replace('{{PREFIX_SELECTOR}}', prefix_selector_html)

    if effective_layer:
        layer_selector_shared = f'''
                    <div class="mb-3">
                        <label class="form-label">Layer *</label>
                        <p class="form-control-plaintext mb-0"><strong>{effective_layer.name}</strong> <small class="text-muted">(from layer view)</small></p>
                    </div>'''
        layer_hidden_field = (
            f'<input type="hidden" name="layer_id" class="submit-layer-id-field" '
            f'value="{effective_layer.id}">'
        )
    elif layers:
        opts = '<option value="">Select a layer...</option>' + ''.join(
            f'<option value="{p.id}">{p.name}</option>' for p in layers
        )
        layer_selector_shared = f'''
                    <div class="mb-3">
                        <label for="layer_id" class="form-label">Layer *</label>
                        <select class="form-select" id="layer_id" required>
                            {opts}
                        </select>
                        <div class="form-text">Drafts are submitted to a specific layer. Workgroups update when you change layer.</div>
                    </div>'''
        layer_hidden_field = (
            '<input type="hidden" name="layer_id" class="submit-layer-id-field" value="">'
        )
    else:
        layer_selector_shared = f'''
                    <div class="mb-3">
                        <p class="text-warning mb-0">{html_mod.escape(empty_state_message)}</p>
                    </div>'''
        layer_hidden_field = ''

    submit_template = submit_template.replace('{{LAYER_SELECTOR_SHARED}}', layer_selector_shared)
    submit_template = submit_template.replace('{{LAYER_HIDDEN_FIELD}}', layer_hidden_field)
    submit_template = submit_template.replace('{{DOCUMENT_META_FIELDS}}', _submit_document_meta_fields_html())
    stripe_pk = os.environ.get('STRIPE_PUBLISHABLE_KEY', '')
    submit_template = submit_template.replace('{{STRIPE_PK}}', stripe_pk)
    submit_template = submit_template.replace(
        '{{PAGE_HERO}}',
        render_page_hero_html('submit_draft'),
    )
    offer_tier = effective_layer and getattr(effective_layer, 'offer_tier_pricing', False)
    submit_template = submit_template.replace('{{OFFER_TIER_PRICING}}', 'true' if offer_tier else 'false')
    submit_template = submit_template.replace('{build_number}', str(build_number))

    if not is_feature_enabled('immortalize', layer=effective_layer):
        submit_template = _strip_immortalize_from_submit_template(submit_template)
    return submit_template


def _render_submit_revision_form(
    *,
    draft,
    draft_name,
    display_id,
    new_rev,
    revision_layer_id,
    selected_group: str = '',
    user_menu,
    current_theme,
    build_number,
):
    """Render revision submit form with DB-backed workgroup options."""
    from services.rendering import _format_base_template
    from services.workgroup_links import workgroup_select_options_html

    group_options = workgroup_select_options_html(
        revision_layer_id,
        selected_group or draft.get('group', ''),
    )
    tag_slugs = draft.get('tag_slugs') or []
    if isinstance(tag_slugs, list) and tag_slugs and isinstance(tag_slugs[0], dict):
        tags_value = ', '.join(t.get('slug') or t.get('label') or '' for t in tag_slugs)
    else:
        tags_value = draft.get('document_tags') or ''
    doc_meta_html = _submit_document_meta_fields_html(
        selected_category=draft.get('document_category') or '',
        tags_value=tags_value,
        compact=True,
    )
    draft_detail_url = url_for('documents.draft_detail', draft_name=draft_name)
    revision_form = f"""
    <div class="gh-page container mt-4">
        {gh_breadcrumb([('Home', url_for('pages.home')), (display_id, draft_detail_url), ('Submit Revision', None)])}
        {gh_page_header('Submit New Revision', f'Submit a new revision of {display_id} (rev {draft.get("rev", "00")} → {new_rev})', 'fa-code-branch')}

        {gh_living_module('Revision form', f'''
        <form method="POST" enctype="multipart/form-data" id="revisionForm">
            <div class="mb-3">
                <label class="form-label">Draft Name</label>
                <input type="text" class="form-control" value="{display_id}" disabled>
                <input type="hidden" name="draft_name" value="{draft_name}">
                <small class="form-text text-muted">This field cannot be changed for revisions</small>
            </div>

            <div class="mb-3">
                <label class="form-label">Title *</label>
                <input type="text" class="form-control" name="title" value="{draft.get('title', '')}" required>
            </div>

            <div class="mb-3">
                <label class="form-label">Authors *</label>
                <input type="text" class="form-control" name="authors" value="{draft.get('authors', '')}" required>
                <small class="form-text text-muted">Comma-separated list</small>
            </div>

            <div class="mb-3">
                <label class="form-label">Abstract</label>
                <textarea class="form-control" name="abstract" rows="4">{draft.get('abstract', '')}</textarea>
            </div>

            <div class="mb-3">
                <label class="form-label">Workgroup</label>
                <select class="form-control" name="group">
                    {group_options}
                </select>
            </div>
            {doc_meta_html}

            <div class="mb-3">
                <label class="form-label">What changed since the last revision?</label>
                <textarea class="form-control" name="what_changed" rows="3"
                          placeholder="Example: Clarified workgroup role in determining rough consensus; added glossary; no change to core governance principles."></textarea>
                <small class="form-text text-muted">
                    Optional but recommended. Briefly describe substantive changes so reviewers and future readers
                    can understand what evolved and why. Not required for minor or editorial edits.
                </small>
            </div>

            <ul class="nav nav-tabs" role="tablist">
                <li class="nav-item">
                    <a class="nav-link active" data-bs-toggle="tab" href="#upload" onclick="document.getElementById('sourceType').value='file'">Upload File</a>
                </li>
                <li class="nav-item">
                    <a class="nav-link" data-bs-toggle="tab" href="#ordinal" onclick="document.getElementById('sourceType').value='ordinal'">Bitcoin Ordinal</a>
                </li>
            </ul>

            <div class="tab-content mt-3">
                <div id="upload" class="tab-pane active">
                    <div class="mb-3">
                        <label class="form-label">Upload Document *</label>
                        <input type="file" class="form-control" name="file" accept=".txt,.pdf,.xml,.docx">
                        <small class="form-text text-muted">Supported formats: TXT, PDF, XML, DOCX</small>
                    </div>
                </div>

                <div id="ordinal" class="tab-pane">
                    <div class="mb-3">
                        <label class="form-label">Inscription ID *</label>
                        <input type="text" class="form-control" name="ordinalId" id="ordinalId"
                               placeholder="e.g., 6fb976ab49dcec017f1e201e84395983204ae1a7c2abf7ced0a85d692e442799i0">
                        <small class="form-text text-muted">The unique inscription ID from Bitcoin</small>
                    </div>

                    <div class="mb-3">
                        <button type="button" class="btn btn-secondary" onclick="previewOrdinal()">
                            <i class="fas fa-eye me-1"></i>Preview Ordinal
                        </button>
                    </div>

                    <div id="ordinalPreview" class="mb-3" style="display: none;">
                        <div class="card">
                            <div class="card-header">
                                <h6>Ordinal Preview</h6>
                            </div>
                            <div class="card-body">
                                <div id="ordinalContent"></div>
                                <input type="hidden" name="ordinalContentUrl" id="ordinalContentUrl">
                                <input type="hidden" name="ordinalContentType" id="ordinalContentType">
                                <input type="hidden" name="inscriptionNumber" id="inscriptionNumber">
                                <input type="hidden" name="blockHeight" id="blockHeight">
                                <input type="hidden" name="inscriptionTimestamp" id="inscriptionTimestamp">
                            </div>
                        </div>
                    </div>
                </div>
            </div>

            <input type="hidden" name="sourceType" value="file" id="sourceType">

            <div class="mt-4">
                <button type="submit" class="btn btn-success btn-lg">
                    <i class="fas fa-upload me-2"></i>Submit Revision
                </button>
                <a href="{draft_detail_url}" class="btn btn-secondary btn-lg ms-2">Cancel</a>
            </div>
        </form>
        ''', 'fa-edit')}
    </div>

    <script>
    async function previewOrdinal() {{
        const inscriptionId = document.getElementById('ordinalId').value.trim();
        if (!inscriptionId) {{
            await GhDialog.alert({{ title: 'Notice', message: 'Please enter an inscription ID', variant: 'info' }});
            return;
        }}

        const preview = document.getElementById('ordinalPreview');
        const content = document.getElementById('ordinalContent');
        content.innerHTML = '<div class="text-center"><i class="fas fa-spinner fa-spin"></i> Loading ordinal...</div>';
        preview.style.display = 'block';

        try {{
            const response = await fetch('/api/ordinal/preview', {{
                method: 'POST',
                headers: {{
                    'Content-Type': 'application/json'
                }},
                body: JSON.stringify({{ inscriptionId }})
            }});

            const data = await response.json();

            if (!data.success) {{
                content.innerHTML = `<div class="alert alert-danger">Error: ${{data.error}}</div>`;
                return;
            }}

            document.getElementById('ordinalContentUrl').value = data.contentUrl;
            document.getElementById('ordinalContentType').value = data.contentType;
            document.getElementById('inscriptionNumber').value = data.inscriptionNumber || '';
            document.getElementById('blockHeight').value = data.blockHeight || '';
            document.getElementById('inscriptionTimestamp').value = data.timestamp || '';

            const contentResponse = await fetch(data.contentUrl);
            const contentText = await contentResponse.text();

            const isMarkdown = data.contentType.includes('markdown') || data.contentType.includes('text/plain');

            if (isMarkdown) {{
                const convertResponse = await fetch('/api/ordinal/convert-markdown', {{
                    method: 'POST',
                    headers: {{
                        'Content-Type': 'application/json'
                    }},
                    body: JSON.stringify({{ markdown: contentText }})
                }});

                const convertData = await convertResponse.json();

                if (convertData.success) {{
                    content.innerHTML = '';
                    const infoDiv = document.createElement('div');
                    infoDiv.className = 'alert alert-info mb-3';
                    infoDiv.innerHTML = `<strong>Preview:</strong> Inscription #${{data.inscriptionNumber}} | Block: ${{data.blockHeight}} | Size: ${{(data.contentSize / 1024).toFixed(2)}} KB`;
                    content.appendChild(infoDiv);
                    const previewDiv = document.createElement('div');
                    previewDiv.innerHTML = convertData.html;
                    content.appendChild(previewDiv);
                    previewDiv.querySelectorAll('img').forEach(img => {{
                        img.style.maxWidth = '100%';
                        img.style.display = 'block';
                        img.style.margin = '1em 0';
                    }});
                }} else {{
                    content.innerHTML = `<div class="alert alert-danger">Conversion failed: ${{convertData.error}}</div>`;
                }}
            }} else {{
                content.innerHTML = `<pre style="max-height: 400px; overflow-y: auto; white-space: pre-wrap;">${{contentText.substring(0, 2000)}}</pre>`;
            }}

        }} catch (error) {{
            content.innerHTML = `<div class="alert alert-danger">Error loading ordinal: ${{error.message}}</div>`;
        }}
    }}
    </script>
    """

    return _format_base_template(
        title=f"Submit Revision - {display_id}",
        theme=current_theme,
        user_menu=user_menu,
        content=revision_form,
        build_number=build_number,
    )


@bp.route('/submit/', methods=['GET', 'POST'])
@require_auth
def submit_draft():
    from services.rendering import _format_base_template, generate_user_menu
    from services.identity import get_current_user
    from config import BUILD_NUMBER
    from services.documents import calculate_pages_and_words
    from services.workgroup_links import workgroup_belongs_to_layer

    user_menu = generate_user_menu()
    current_theme = session.get('theme', get_current_user().get('theme', 'dark') if get_current_user() else 'dark')

    current_user = get_current_user() or {}
    user_id = current_user.get('id')
    user_role = (current_user.get('role') or '').strip().lower()

    # Layer dropdown rules:
    #   * Admins see the full ``visible_layers_for_user`` (pending + active,
    #     everything approved) so they can review any submission target.
    #   * Non-admins only see layers where they are a ``LayerAdmin`` (or the
    #     layer's initiator/owner), filtered to ``approval_status='approved'``
    #     and deduped by name. Empty result renders a friendly message.
    from services.layer_prefixes import visible_layers_for_user, _layer_admin_layer_ids

    if user_role == 'admin':
        layers = visible_layers_for_user(user_id)
    else:
        admin_layer_ids = _layer_admin_layer_ids(user_id)
        if admin_layer_ids:
            layers = (
                Layer.query
                .filter(
                    Layer.id.in_(admin_layer_ids),
                    Layer.approval_status == 'approved',
                )
                .group_by(Layer.name)
                .order_by(Layer.name)
                .all()
            )
        else:
            layers = []
    layer_from_param = None
    if request.args.get('layer'):
        layer_from_param = Layer.query.filter_by(slug=request.args.get('layer').strip()).first()
    elif request.args.get('layer_id'):
        layer_from_param = Layer.query.get(request.args.get('layer_id').strip())
    effective_layer = g.get('layer') or layer_from_param
    # Non-admins can only target a layer they're a member of. If a layer
    # came in via query param / subdomain but the user isn't a member,
    # treat as no effective layer (so the dropdown shows).
    if effective_layer and user_role != 'admin' and effective_layer.id not in set(_layer_admin_layer_ids(user_id)):
        effective_layer = None
    if not effective_layer and request.method == 'POST' and request.form.get('layer_id'):
        effective_layer = Layer.query.get(request.form.get('layer_id').strip())

    selected_group = request.form.get('group', '').strip() if request.method == 'POST' else ''
    if not layers and user_role != 'admin':
        empty_state_message = (
            "You're not a member of any layer yet, so the layer dropdown is empty. "
            "Ask a layer admin to add you as a layer admin, or create a new layer from "
            "the Layers page to start submitting drafts."
        )
    elif not layers:
        empty_state_message = (
            'No approved layers available. Submit from a layer subdomain '
            '(e.g. overweb.themetalayer.org) or create a layer first.'
        )
    else:
        empty_state_message = ''
    submit_template = _build_submit_form_template(
        effective_layer=effective_layer,
        layers=layers,
        selected_group=selected_group,
        build_number=BUILD_NUMBER,
        empty_state_message=empty_state_message,
    )

    if request.method == 'POST':
        # Get common fields
        title = request.form.get('title', '').strip()
        authors = request.form.get('authors', '').strip()
        abstract = request.form.get('abstract', '').strip()
        group = request.form.get('group', '').strip()
        source_type = request.form.get('sourceType', 'file').strip()
        form_layer_id = request.form.get('layer_id', '').strip()
        layer_id = form_layer_id or (effective_layer.id if effective_layer else None)

        # Validate layer_id when layers exist
        if not layer_id and layers:
            flash('Please select a layer for this submission.', 'error')
            return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)

        if group and layer_id and not workgroup_belongs_to_layer(group, layer_id):
            flash('Selected workgroup is not valid for this layer.', 'error')
            return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)

        from services.submission_dedup import (
            compute_content_hash_for_file,
            compute_content_hash_from_bytes,
            find_submission_conflict,
            conflict_message,
        )

        # Process authors (comma-separated)
        authors_list = [a.strip() for a in authors.split(',') if a.strip()]

        # Generate submission ID
        submission_id = ''.join(random.choices(string.ascii_lowercase + string.digits, k=8))

        # Handle based on source type
        if source_type == 'ordinal':
            # Ordinal submission
            ordinal_id = request.form.get('ordinalId', '').strip()
            ordinal_content_url = request.form.get('ordinalContentUrl', '').strip()
            ordinal_content_type = request.form.get('ordinalContentType', '').strip()
            inscription_number = request.form.get('inscriptionNumber', '').strip()
            block_height = request.form.get('blockHeight', '').strip()
            inscription_timestamp = request.form.get('inscriptionTimestamp', '').strip()

            # Validation
            if not title or not authors or not ordinal_id:
                flash('Title, authors, and inscription ID are required', 'error')
                return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)

            if not ordinal_content_url:
                flash('Please preview the ordinal before submitting', 'error')
                return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)

            try:
                ordinal_content_url = validate_ordinals_fetch_url(ordinal_content_url)
            except ValueError:
                flash('Invalid ordinal content URL', 'error')
                return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)

            # Fetch ordinal content and calculate pages/words
            try:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                    'Accept': '*/*',
                    'Connection': 'keep-alive'
                }
                response = requests.get(ordinal_content_url, headers=headers, timeout=30)
                response.raise_for_status()
                content_bytes = response.content
                content_text = content_bytes.decode('utf-8', errors='replace')

                # Calculate pages and words from text
                word_count = len(content_text.split())
                chars_per_page = 3000
                page_count = max(1, (len(content_text) + chars_per_page - 1) // chars_per_page)
            except Exception as e:
                current_app.logger.error(f"Failed to fetch ordinal content for pages/words: {e}")
                page_count = 1
                word_count = 0
                content_bytes = b''

            content_hash = compute_content_hash_from_bytes(
                content_bytes,
                content_type=ordinal_content_type,
            ) if content_bytes else None

            conflict = find_submission_conflict(
                title=title,
                ordinal_id=ordinal_id,
                content_hash=content_hash,
            )
            if conflict:
                flash(conflict_message(conflict[0], conflict[1]), 'error')
                return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)

            # Create submission record with ordinal data
            current_user_info = get_current_user()
            current_app.logger.info(f"📝 CREATING SUBMISSION:")
            current_app.logger.info(f"   current_user_info: {current_user_info}")
            current_app.logger.info(f"   submitted_by will be: {current_user_info['name']}")

            # Get doc_type from form (default to 'draft')
            doc_type = request.form.get('doc_type', 'draft').strip() or 'draft'
            if doc_type not in ['draft', 'rfc']:
                doc_type = 'draft'

            submission = Submission(
                draft_name=submission_id,
                title=title,
                authors=authors_list,
                abstract=abstract,
                group=group,
                layer_id=layer_id,
                submitted_by=current_user_info['name'],
                sourceType='ordinal',
                doc_type=doc_type,
                ordinalId=ordinal_id,
                ordinalContentUrl=ordinal_content_url,
                ordinalContentType=ordinal_content_type,
                inscriptionNumber=int(inscription_number) if inscription_number else None,
                blockHeight=int(block_height) if block_height else None,
                inscriptionTimestamp=datetime.strptime(inscription_timestamp.replace(' UTC', ''), '%Y-%m-%d %H:%M:%S') if inscription_timestamp else None,
                pages=page_count,
                words=word_count,
                content_hash=content_hash,
                prefix_code=(request.form.get('prefix_code') or '').strip().upper() or None,
            )

        else:
            # File upload submission
            file = request.files.get('file')

            # Validation
            if not title or not authors or not file:
                flash('Title, authors, and file are required', 'error')
                return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)

            stored_name, file_path, upload_err = save_submission_upload(file, submission_id)
            if upload_err:
                flash(upload_err, 'error')
                return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)
            filename = stored_name

            # Calculate pages and words
            pages, words = calculate_pages_and_words(file_path, filename)
            content_hash = compute_content_hash_for_file(file_path, filename)

            conflict = find_submission_conflict(
                title=title,
                content_hash=content_hash,
            )
            if conflict:
                flash(conflict_message(conflict[0], conflict[1]), 'error')
                try:
                    os.remove(file_path)
                except OSError:
                    pass
                return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)

            # Create submission record with file data
            submission = Submission(
                draft_name=submission_id,
                title=title,
                authors=authors_list,
                abstract=abstract,
                group=group,
                layer_id=layer_id,
                filename=filename,
                file_path=file_path,
                submitted_by=get_current_user()['name'],
                sourceType='file',
                pages=pages,
                words=words,
                content_hash=content_hash,
                prefix_code=(request.form.get('prefix_code') or '').strip().upper() or None,
            )

        # Save to database
        db.session.add(submission)
        db.session.flush()
        _apply_submission_document_meta(
            submission,
            request.form,
            get_current_user()['id'] if get_current_user() else None,
        )
        # Validate the per-draft prefix override against the chosen layer's
        # current prefix set. An empty/missing value falls back to the layer
        # default, or to the system default ``ML`` if the layer has no
        # configured prefixes yet. A non-empty value must match one of the
        # layer's active prefixes and be a valid 2-uppercase-letter code.
        from services.layer_prefixes import (
            is_valid_prefix_format,
            list_prefixes,
            get_default_prefix,
        )
        raw_prefix = (request.form.get('prefix_code') or '').strip().upper()
        if not raw_prefix:
            # Server-side fallback: if the layer has no prefixes configured,
            # default the draft to the system code ``ML`` so the row is
            # never written with an empty ``prefix_code`` (the client also
            # pre-fills ``ML`` for this branch, but a paranoid server-side
            # default keeps the column non-null even on stale form posts).
            available = list_prefixes(layer_id) if layer_id else []
            if not available:
                raw_prefix = 'ML'
            else:
                default_p = get_default_prefix(layer_id) if layer_id else None
                raw_prefix = (default_p.prefix if default_p else available[0].prefix) or 'ML'
        if raw_prefix:
            if not is_valid_prefix_format(raw_prefix):
                flash('Invalid prefix code. Expected exactly two uppercase letters.', 'error')
                return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)
            available = list_prefixes(layer_id) if layer_id else []
            allowed = {p.prefix for p in available}
            if allowed and raw_prefix not in allowed:
                # If the layer has no prefixes at all the client/server
                # fallback (``ML``) must be accepted without complaint.
                if not (raw_prefix == 'ML' and len(available) == 0):
                    flash('The selected prefix is not available for this layer.', 'error')
                    return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)
            submission.prefix_code = raw_prefix
        db.session.commit()

        # Log the action
        source_desc = f"from ordinal {submission.ordinalId}" if source_type == 'ordinal' else "via file upload"
        add_to_document_history(f"draft-{submission_id}", "submitted", get_current_user()['name'], f"New draft submitted {source_desc}: {title}")

        flash('Draft submitted successfully!', 'success')
        return redirect(url_for('submissions.submission_detail', submission_id=submission.draft_name or submission.id))

    return _format_base_template(title="Submit a Meta-Layer Draft - GovHub", theme=current_theme, user_menu=user_menu, content=submit_template, build_number=BUILD_NUMBER)


@bp.route('/submit/revision/<draft_name>/', methods=['GET', 'POST'])
@require_auth
def submit_revision(draft_name):
    """Submit a new revision of an existing draft"""
    from services.rendering import _format_base_template, generate_user_menu
    from services.identity import get_current_user
    from config import BUILD_NUMBER
    from services.documents import calculate_pages_and_words, load_draft_data, DRAFTS
    from services.workgroup_links import workgroup_belongs_to_layer, workgroup_select_options_html

    user_menu = generate_user_menu()
    current_theme = session.get('theme', get_current_user().get('theme', 'dark') if get_current_user() else 'dark')

    # Find the current draft
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)

    # If not found in DRAFTS, try to find as a submission
    submission = None
    if not draft:
        submission = get_submission_by_ref(draft_name)
        if submission and submission.status == 'approved':
            from services.layer_tags import tags_for_subject
            from models.layer_tag import SUBJECT_SUBMISSION

            draft = {
                'name': submission.id,
                'title': submission.title,
                'authors': ', '.join(submission.authors) if isinstance(submission.authors, list) else submission.authors,
                'abstract': submission.abstract or '',
                'group': submission.group or '',
                'rev': submission.revision_number or '00',
                'ml_number': submission.ml_number,
                'document_category': getattr(submission, 'document_category', None) or 'document',
                'tag_slugs': tags_for_subject(SUBJECT_SUBMISSION, submission.id),
            }
        elif submission:
            flash('Cannot create revision of unapproved submission', 'error')
            return redirect(url_for('submissions.submission_detail', submission_id=submission.id))

    if not draft:
        flash('Draft not found', 'error')
        return redirect(url_for('documents.all_documents'))

    # Inherit project_id from parent draft (revisions belong to same layer)
    parent_sub = get_submission_by_ref(draft_name)
    revision_layer_id = parent_sub.layer_id if parent_sub else (g.layer.id if g.get('layer') else None)

    # Determine display ID (ML-Draft-XXX or internal ID)
    display_id = draft.get('ml_number', draft_name) or draft_name

    # Calculate new revision number
    current_rev = int(draft.get('rev', '00'))
    new_rev = f"{current_rev + 1:02d}"

    if request.method == 'POST':
        # Get form data
        title = request.form.get('title', '').strip()
        authors = request.form.get('authors', '').strip()
        abstract = request.form.get('abstract', '').strip()
        group = request.form.get('group', '').strip()
        what_changed = request.form.get('what_changed', '').strip()
        source_type = request.form.get('sourceType', 'file').strip()

        if group and revision_layer_id and not workgroup_belongs_to_layer(group, revision_layer_id):
            flash('Selected workgroup is not valid for this layer.', 'error')
            return _render_submit_revision_form(
                draft=draft,
                draft_name=draft_name,
                display_id=display_id,
                new_rev=new_rev,
                revision_layer_id=revision_layer_id,
                selected_group=group,
                user_menu=user_menu,
                current_theme=current_theme,
                build_number=BUILD_NUMBER,
            )

        # Process authors
        authors_list = [a.strip() for a in authors.split(',') if a.strip()]

        from services.submission_dedup import (
            compute_content_hash_for_file,
            compute_content_hash_from_bytes,
            find_submission_conflict,
            conflict_message,
        )

        # Generate submission ID
        submission_id = ''.join(random.choices(string.ascii_lowercase + string.digits, k=8))

        # Handle based on source type
        if source_type == 'ordinal':
            # Ordinal submission
            ordinal_id = request.form.get('ordinalId', '').strip()
            ordinal_content_url = request.form.get('ordinalContentUrl', '').strip()
            ordinal_content_type = request.form.get('ordinalContentType', '').strip()
            inscription_number = request.form.get('inscriptionNumber', '').strip()
            block_height = request.form.get('blockHeight', '').strip()
            inscription_timestamp = request.form.get('inscriptionTimestamp', '').strip()

            # Validation
            if not title or not authors or not ordinal_id:
                flash('Title, authors, and inscription ID are required', 'error')
                return redirect(url_for('submissions.submit_revision', draft_name=draft_name))

            if not ordinal_content_url:
                flash('Please preview the ordinal before submitting', 'error')
                return redirect(url_for('submissions.submit_revision', draft_name=draft_name))

            try:
                ordinal_content_url = validate_ordinals_fetch_url(ordinal_content_url)
            except ValueError:
                flash('Invalid ordinal content URL', 'error')
                return redirect(url_for('submissions.submit_revision', draft_name=draft_name))

            # Fetch ordinal content and calculate pages/words
            content_bytes = b''
            try:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                    'Accept': '*/*',
                    'Connection': 'keep-alive'
                }
                response = requests.get(ordinal_content_url, headers=headers, timeout=30)
                response.raise_for_status()
                content_bytes = response.content
                content_text = content_bytes.decode('utf-8', errors='replace')

                word_count = len(content_text.split())
                chars_per_page = 3000
                page_count = max(1, (len(content_text) + chars_per_page - 1) // chars_per_page)
            except Exception as e:
                current_app.logger.error(f"Failed to fetch ordinal content: {e}")
                page_count = 1
                word_count = 0

            content_hash = compute_content_hash_from_bytes(
                content_bytes,
                content_type=ordinal_content_type,
            ) if content_bytes else None

            conflict = find_submission_conflict(
                title=title,
                ordinal_id=ordinal_id,
                content_hash=content_hash,
                exclude_family_parent_id=draft_name,
            )
            if conflict:
                flash(conflict_message(conflict[0], conflict[1]), 'error')
                return redirect(url_for('submissions.submit_revision', draft_name=draft_name))

            # Create revision submission with ordinal data
            submission = Submission(
                draft_name=submission_id,
                title=title,
                authors=authors_list,
                abstract=abstract,
                group=group,
                layer_id=revision_layer_id,
                submitted_by=get_current_user()['name'],
                sourceType='ordinal',
                doc_type='draft',
                ordinalId=ordinal_id,
                ordinalContentUrl=ordinal_content_url,
                ordinalContentType=ordinal_content_type,
                inscriptionNumber=int(inscription_number) if inscription_number else None,
                blockHeight=int(block_height) if block_height else None,
                inscriptionTimestamp=datetime.strptime(inscription_timestamp.replace(' UTC', ''), '%Y-%m-%d %H:%M:%S') if inscription_timestamp else None,
                pages=page_count,
                words=word_count,
                content_hash=content_hash,
                parent_draft_name=draft_name,
                revision_number=new_rev,
                what_changed=what_changed,
                is_revision=True
            )
        else:
            # File upload submission
            file = request.files.get('file')

            # Validation
            if not title or not authors or not file:
                flash('Title, authors, and file are required', 'error')
                return redirect(url_for('submissions.submit_revision', draft_name=draft_name))

            stored_name, file_path, upload_err = save_submission_upload(file, submission_id)
            if upload_err:
                flash(upload_err, 'error')
                return redirect(url_for('submissions.submit_revision', draft_name=draft_name))
            filename = stored_name

            # Calculate pages and words
            pages, words = calculate_pages_and_words(file_path, filename)
            content_hash = compute_content_hash_for_file(file_path, filename)

            conflict = find_submission_conflict(
                title=title,
                content_hash=content_hash,
                exclude_family_parent_id=draft_name,
            )
            if conflict:
                flash(conflict_message(conflict[0], conflict[1]), 'error')
                try:
                    os.remove(file_path)
                except OSError:
                    pass
                return redirect(url_for('submissions.submit_revision', draft_name=draft_name))

            # Create revision submission with file data
            submission = Submission(
                draft_name=submission_id,
                title=title,
                authors=authors_list,
                abstract=abstract,
                group=group,
                layer_id=revision_layer_id,
                filename=filename,
                file_path=file_path,
                submitted_by=get_current_user()['name'],
                sourceType='file',
                pages=pages,
                words=words,
                content_hash=content_hash,
                parent_draft_name=draft_name,
                revision_number=new_rev,
                what_changed=what_changed,
                is_revision=True
            )

        # Save to database
        db.session.add(submission)
        db.session.flush()
        _apply_submission_document_meta(
            submission,
            request.form,
            get_current_user()['id'] if get_current_user() else None,
        )
        db.session.commit()

        # Log the action
        source_desc = f"from ordinal {submission.ordinalId}" if source_type == 'ordinal' else "via file upload"
        change_desc = f" Changes: {what_changed[:100]}" if what_changed else ""
        add_to_document_history(
            draft_name,
            "revision_submitted",
            get_current_user()['name'],
            f"Revision {new_rev} submitted {source_desc}.{change_desc}"
        )

        flash(f'Revision {new_rev} submitted successfully!', 'success')
        return redirect(url_for('submissions.submission_detail', submission_id=submission.draft_name or submission.id))

    return _render_submit_revision_form(
        draft=draft,
        draft_name=draft_name,
        display_id=display_id,
        new_rev=new_rev,
        revision_layer_id=revision_layer_id,
        user_menu=user_menu,
        current_theme=current_theme,
        build_number=BUILD_NUMBER,
    )


@bp.route('/submit/status/')
@require_auth
def submission_status():
    from services.rendering import _format_base_template, generate_user_menu
    from services.identity import get_current_user
    from config import BUILD_NUMBER

    user_menu = generate_user_menu()
    current_theme = session.get('theme', get_current_user().get('theme', 'dark') if get_current_user() else 'dark')

    # Get user's submissions
    user_name = get_current_user()['name']
    submissions = Submission.query.filter_by(submitted_by=user_name).order_by(Submission.submitted_at.desc()).all()

    # Format submissions for template
    submissions_html = ""
    for submission in submissions:
        status_badge = {
            'submitted': 'badge bg-warning text-dark',
            'approved': 'badge bg-success',
            'rejected': 'badge bg-danger',
            'published': 'badge bg-info',
            'inscription_pending': 'badge bg-warning'
        }.get(submission.status, 'badge bg-secondary')

        # Get source type
        source_type = getattr(submission, 'sourceType', 'file')
        source_badge = '<span class="badge bg-info ms-2"><i class="bi bi-coin"></i> Ordinal</span>' if source_type == 'ordinal' else '<span class="badge bg-secondary ms-2"><i class="bi bi-file-earmark"></i> File</span>'

        # Get revision info
        is_revision = getattr(submission, 'is_revision', False)
        revision_number = getattr(submission, 'revision_number', '')
        parent_draft_name = getattr(submission, 'parent_draft_name', '')
        revision_badge = f'<span class="badge bg-success ms-2">Revision {revision_number}</span>' if is_revision and revision_number else ''

        # Get source info (inscription number or filename)
        if source_type == 'ordinal':
            inscription_number = getattr(submission, 'inscriptionNumber', None)
            ordinal_id = getattr(submission, 'ordinalId', None)
            if inscription_number:
                source_info = f'<p class="mb-2"><strong>Inscription:</strong> #{inscription_number}</p>'
            elif ordinal_id:
                shortened_id = shorten_inscription_id(ordinal_id, 8)
                source_info = f'<p class="mb-2"><strong>Inscription:</strong> <a href="https://ordinals.com/inscription/{ordinal_id}" target="_blank" class="text-decoration-none"><code>{shortened_id}</code></a></p>'
            else:
                source_info = ''
        else:
            filename = getattr(submission, 'filename', None)
            source_info = f'<p class="mb-2"><strong>File:</strong> {filename}</p>' if filename else '<p class="mb-2 text-muted"><em>No file</em></p>'

        detail_url = url_for('submissions.submission_detail', submission_id=submission.id)
        draft_url = url_for('documents.draft_detail', draft_name=submission.id)
        submissions_html += f"""
        <div class="submission-item">
            <div class="card mb-3">
                <div class="card-header d-flex justify-content-between align-items-center">
                    <h6 class="mb-0">
                        <a href="{detail_url}" class="text-decoration-none">
                            {submission.title}
                        </a>
                    </h6>
                    <div>
                        <span class="{status_badge}">{submission.status.title()}</span>
                        {source_badge}
                        {revision_badge}
                    </div>
                </div>
                <div class="card-body">
                    <div class="row">
                        <div class="col-md-8">
                            <p class="mb-2"><strong>Authors:</strong> {', '.join(submission.authors)}</p>
                            <p class="mb-2"><strong>Group:</strong> {submission.group or 'None'}</p>
                            <p class="mb-2"><strong>Submitted:</strong> {submission.submitted_at.strftime('%Y-%m-%d %H:%M')}</p>
                            {source_info}
                            {f'<p class="mb-2"><strong>Abstract:</strong> {submission.abstract[:100]}...</p>' if submission.abstract else ''}
                        </div>
                        <div class="col-md-4 text-end">
                            <a href="{detail_url}" class="btn btn-sm btn-primary me-2">View Details</a>
                            <a href="{draft_url}" class="btn btn-sm btn-outline-primary">View Draft</a>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        """

    home_url = url_for('pages.home')
    submit_url = url_for('submissions.submit_draft')
    doc_all_url = url_for('documents.all_documents')
    content = f"""
    <div class="gh-page container mt-4">
        {gh_breadcrumb([('Home', home_url), ('Submit Draft', submit_url), ('My Submissions', None)])}
        {gh_page_header('My Submissions', f'You have {len(submissions)} submission(s)' if submissions else 'No submissions yet', 'fa-inbox', actions_html=f'<a href="{submit_url}" class="btn btn-primary btn-sm">Submit draft</a>')}

        {submissions_html}

        <div class="mt-4">
            <a href="{submit_url}" class="btn btn-primary">Submit Another Draft</a>
            <a href="{home_url}" class="btn btn-secondary ms-2">Back to Home</a>
        </div>
    </div>
    """

    return _format_base_template(title="My Submissions - GovHub", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER)


@bp.route('/api/submissions/<submission_id>/metadata/', methods=['PATCH'])
@require_auth
def patch_submission_metadata(submission_id):
    """Update draft metadata (workgroup assignment, title, abstract)."""
    from models import Workgroup
    from services.workgroup_links import workgroup_belongs_to_layer

    current_user = get_current_user()
    submission = get_submission_by_ref(submission_id)
    if not submission:
        return jsonify({'error': 'Submission not found'}), 404
    if not can_edit_submission_metadata(current_user, submission):
        return jsonify({'error': 'Not allowed to edit this submission'}), 403

    data = request.get_json() or {}
    updated = False

    if 'group' in data:
        group_val = (data.get('group') or '').strip()
        if group_val:
            wg = Workgroup.query.filter_by(acronym=group_val).first()
            if not wg:
                return jsonify({'error': 'Unknown workgroup acronym'}), 400
            if submission.layer_id and not workgroup_belongs_to_layer(
                group_val, submission.layer_id
            ):
                return jsonify({
                    'error': 'Workgroup is not available on this layer',
                }), 400
        submission.group = group_val or None
        updated = True

    if 'title' in data and (data.get('title') or '').strip():
        submission.title = data['title'].strip()
        updated = True

    if 'abstract' in data:
        submission.abstract = (data.get('abstract') or '').strip() or None
        updated = True

    if 'document_category' in data:
        from services.document_categories import normalize_document_category
        submission.document_category = normalize_document_category(data.get('document_category'))
        updated = True

    if 'tag_slugs' in data or 'tags' in data or 'document_tags' in data:
        from flask import current_app
        from services.layer_tags import (
            document_tags_enabled,
            set_submission_tags,
            sync_submission_tags_to_artifact,
        )
        if document_tags_enabled(current_app.config) and submission.layer_id:
            raw = data.get('tag_slugs', data.get('tags', data.get('document_tags')))
            set_submission_tags(submission, raw or [], current_user['id'])
            sync_submission_tags_to_artifact(submission)
            updated = True

    if not updated:
        return jsonify({'error': 'No supported fields to update'}), 400

    db.session.commit()
    from services.layer_tags import tags_for_subject
    from models.layer_tag import SUBJECT_SUBMISSION
    from services.document_categories import document_category_label

    tag_rows = tags_for_subject(SUBJECT_SUBMISSION, submission.id)
    return jsonify({
        'success': True,
        'submission': {
            'id': submission.id,
            'title': submission.title,
            'abstract': submission.abstract,
            'group': submission.group,
            'document_category': submission.document_category,
            'document_category_label': document_category_label(submission.document_category or 'document'),
            'tags': tag_rows,
        },
    })


@bp.route('/submit/status/<submission_id>/')
@require_auth
def submission_detail(submission_id):
    from services.rendering import _format_base_template, generate_user_menu
    from services.identity import get_current_user
    from config import BUILD_NUMBER
    from services.submission_preview_md import markdown_to_safe_preview_html, text_looks_like_markdown

    user_menu = generate_user_menu()
    current_theme = session.get('theme', get_current_user().get('theme', 'dark') if get_current_user() else 'dark')
    current_user = get_current_user()

    submission = get_submission_by_ref(submission_id)
    if not submission:
        return "Submission not found", 404

    # Check if user owns this submission or is admin
    if submission.submitted_by != current_user['name'] and current_user.get('role') not in ['admin', 'editor']:
        return "Access denied", 403

    # Handle content preview based on source type
    file_content = "File preview not available"
    content_preview_html = ""
    source_type = getattr(submission, 'sourceType', 'file')

    if source_type == 'ordinal':
        # Ordinal content - generate preview HTML (HTML inscriptions need script-capable iframe + sniff if MIME wrong)
        import html as html_lib

        ordinal_content_type = getattr(submission, 'ordinalContentType', '') or ''
        ordinal_content_url = getattr(submission, 'ordinalContentUrl', '') or ''

        if ordinal_content_type.startswith('image/'):
            safe_src = html_lib.escape(ordinal_content_url, quote=True)
            content_preview_html = (
                f'<img src="{safe_src}" class="img-fluid" style="max-height: 400px;" alt="Ordinal content">'
            )
        elif looks_like_html_inscription('', ordinal_content_type):
            content_preview_html = format_ordinal_html_iframe_preview(ordinal_content_url)
            file_content = ""
        elif 'text/' in ordinal_content_type or 'application/json' in ordinal_content_type:
            try:
                safe_ordinal_url = validate_ordinals_fetch_url(ordinal_content_url)
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = requests.get(safe_ordinal_url, headers=headers, timeout=10)
                response.raise_for_status()
                text_content = response.text

                if looks_like_html_inscription(text_content, ordinal_content_type):
                    content_preview_html = format_ordinal_html_iframe_preview(ordinal_content_url)
                    file_content = ""
                elif 'text/markdown' in ordinal_content_type or (
                    ordinal_content_type and 'markdown' in ordinal_content_type.lower()
                ):
                    html_content = markdown_to_safe_preview_html(text_content)
                    if html_content:
                        content_preview_html = (
                            f'<div class="border p-3 markdown-body" '
                            f'style="max-height: 400px; overflow-y: auto;">{html_content}</div>'
                        )
                        file_content = ""
                    else:
                        file_content = text_content[:2000] + (
                            "..." if len(text_content) > 2000 else ""
                        )
                elif (
                    'text/plain' in ordinal_content_type
                    or 'text/javascript' in ordinal_content_type
                    or 'application/javascript' in ordinal_content_type
                    or 'application/json' in ordinal_content_type
                ):
                    is_markdown = (
                        'text/plain' in ordinal_content_type
                        and text_looks_like_markdown(text_content)
                    )

                    if is_markdown:
                        html_content = markdown_to_safe_preview_html(text_content)
                        if html_content:
                            content_preview_html = (
                                f'<div class="border p-3 markdown-body" '
                                f'style="max-height: 400px; overflow-y: auto;">{html_content}</div>'
                            )
                            file_content = ""
                        else:
                            file_content = text_content[:2000] + (
                                "..." if len(text_content) > 2000 else ""
                            )
                    else:
                        file_content = text_content[:2000] + (
                            "..." if len(text_content) > 2000 else ""
                        )
                else:
                    # e.g. text/css — show snippet
                    file_content = text_content[:2000] + (
                        "..." if len(text_content) > 2000 else ""
                    )
            except Exception as e:
                current_app.logger.error(f"Error fetching ordinal text content: {e}")
                file_content = "Error loading ordinal text content"
        else:
            file_content = (
                f"Ordinal content type: {ordinal_content_type}\nPreview not available for this content type."
            )

    elif submission.file_path and os.path.exists(submission.file_path):
        # File upload - extract text for preview
        _, ext = os.path.splitext(submission.filename.lower())
        try:
            if ext in ['.txt', '.xml']:
                with open(submission.file_path, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
                    if ext == '.txt' and text_looks_like_markdown(content):
                        html_content = markdown_to_safe_preview_html(content)
                        if html_content:
                            content_preview_html = (
                                f'<div class="border p-3 markdown-body" '
                                f'style="max-height: 400px; overflow-y: auto;">{html_content}</div>'
                            )
                            file_content = ""
                        else:
                            file_content = content[:2000] + "..." if len(content) > 2000 else content
                    else:
                        file_content = content[:2000] + "..." if len(content) > 2000 else content
            elif ext in ('.md', '.markdown'):
                with open(submission.file_path, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
                html_content = markdown_to_safe_preview_html(content)
                if html_content:
                    content_preview_html = (
                        f'<div class="border p-3 markdown-body" '
                        f'style="max-height: 400px; overflow-y: auto;">{html_content}</div>'
                    )
                    file_content = ""
                else:
                    file_content = content[:2000] + "..." if len(content) > 2000 else content
            elif ext == '.docx':
                from docx import Document
                doc = Document(submission.file_path)
                content = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content += cell.text + "\n"
                if content.strip():
                    file_content = content[:2000] + "..." if len(content) > 2000 else content
                else:
                    file_content = "DOCX file appears to be empty or contains no extractable text."
            elif ext == '.pdf':
                file_size = os.path.getsize(submission.file_path)
                file_size_kb = file_size / 1024
                view_url = url_for('submissions.view_submission', submission_id=submission.id)
                download_url = url_for('submissions.download_submission', submission_id=submission.id)
                content_preview_html = f'''
                <div class="pdf-viewer-container">
                    <div class="alert alert-info mb-3">
                        <i class="bi bi-file-pdf"></i> PDF Document ({file_size_kb:.1f} KB) -
                        <a href="{download_url}" class="alert-link">Download PDF</a> for best viewing experience
                    </div>
                    <iframe src="{view_url}"
                            type="application/pdf"
                            style="width: 100%; height: 600px; border: 1px solid var(--card-border);"
                            title="PDF Preview">
                        <p>Your browser does not support PDF preview.
                           <a href="{download_url}">Download the PDF</a> to view it.</p>
                    </iframe>
                </div>
                '''
                file_content = ""
            elif ext == '.doc':
                file_size = os.path.getsize(submission.file_path)
                file_size_kb = file_size / 1024
                file_content = f"Legacy DOC file ({file_size_kb:.1f} KB)\nText extraction not supported for legacy .doc format.\nPlease convert to .docx for text preview."
            else:
                file_size = os.path.getsize(submission.file_path)
                file_size_kb = file_size / 1024
                file_content = f"Unsupported file type: {ext} ({file_size_kb:.1f} KB)\nPreview not available."
        except Exception as e:
            file_size = os.path.getsize(submission.file_path)
            file_size_kb = file_size / 1024
            file_content = f"Error extracting text from {ext[1:].upper()} file ({file_size_kb:.1f} KB): {str(e)}"

    # Prepare template variables
    parent_draft_name = getattr(submission, 'parent_draft_name', '')
    inscription_ts = getattr(submission, 'inscriptionTimestamp', None)
    _title = (submission.title or '').strip()
    _draft = (getattr(submission, 'draft_name', None) or '').strip()
    status_page_heading = _title or _draft or 'Submission Status'
    status_doc_title = _title or _draft or f'Submission {submission.id}'

    template_vars = {
        'submission': submission,
        'current_user': current_user,
        'file_content': file_content,
        'content_preview_html': content_preview_html,
        'status_page_heading': status_page_heading,
        'submission_id': submission.id,
        'submission_status': submission.status,
        'submission_status_title': submission.status.title(),
        'submission_title': submission.title or '',
        'submission_abstract': submission.abstract or '',
        'submission_authors': submission.authors,
        'submission_authors_joined': ', '.join(submission.authors) if submission.authors else '',
        'submission_draft_name': getattr(submission, 'draft_name', submission.id) or '',
        'submission_submitted_at': submission.submitted_at.strftime('%Y-%m-%d %H:%M:%S') if submission.submitted_at else '',
        'submission_filename': submission.filename or '',
        'submission_approved_at': submission.approved_at.strftime('%Y-%m-%d %H:%M:%S') if submission.approved_at else '',
        'submission_rejected_at': submission.rejected_at.strftime('%Y-%m-%d %H:%M:%S') if submission.rejected_at else '',
        'is_admin': current_user and (current_user.get('role') in ['admin', 'editor'] or current_user['name'] in ['admin', 'Admin User']),
        'is_approved': submission.status == 'approved',
        'is_rejected': submission.status == 'rejected',
        'is_approved_or_rejected': submission.status in ['approved', 'rejected'],
        'is_submitted': submission.status == 'submitted',
        'source_type': source_type,
        'is_ordinal': source_type == 'ordinal',
        'is_file': source_type == 'file',
        'ordinal_id': getattr(submission, 'ordinalId', ''),
        'ordinal_id_short': shorten_inscription_id(getattr(submission, 'ordinalId', ''), 8),
        'ordinal_content_url': getattr(submission, 'ordinalContentUrl', ''),
        'ordinal_content_type': getattr(submission, 'ordinalContentType', ''),
        'inscription_number': getattr(submission, 'inscriptionNumber', None),
        'block_height': getattr(submission, 'blockHeight', None),
        'inscription_timestamp': inscription_ts.strftime('%Y-%m-%d %H:%M:%S') if inscription_ts else None,
        'ml_number': submission.ml_number,
        'is_revision': _submission_is_revision(submission),
        'parent_draft_name': parent_draft_name,
        'parent_draft_url': url_for('documents.draft_detail', draft_name=parent_draft_name) if parent_draft_name else '',
        'revision_number': (getattr(submission, 'revision_number', '') or '').strip() or '',
        'what_changed_html': revision_notes_to_safe_html(
            getattr(submission, 'what_changed', '') or ''
        ),
        'artifact_id': getattr(submission, 'artifact_id', None),
        'artifact_layer_slug': Layer.query.get(submission.layer_id).slug if submission.layer_id and getattr(submission, 'artifact_id', None) else None,
        'home_url': url_for('pages.home'),
        'submit_url': url_for('submissions.submit_draft'),
        'doc_all_url': url_for('documents.all_documents'),
        'approve_url': url_for('submissions.approve_submission', submission_id=submission.id),
        'reject_url': url_for('submissions.reject_submission', submission_id=submission.id),
        'download_url': url_for('submissions.download_submission', submission_id=submission.id),
    }

    rendered_content = render_template_string(SUBMISSION_STATUS_TEMPLATE, **template_vars)
    return _format_base_template(
        title=f'{status_doc_title} - GovHub',
        theme=current_theme,
        user_menu=user_menu,
        content=rendered_content,
        build_number=BUILD_NUMBER,
    )


# ---------------------------------------------------------------------------
# API and admin routes
# ---------------------------------------------------------------------------

@bp.route('/api/layers/<layer_id>/submissions/', methods=['GET'])
def list_layer_submissions(layer_id):
    """List approved drafts (not RFCs) for a layer - eligible for voting."""
    Layer.query.get_or_404(layer_id)
    submissions = Submission.query.filter(
        Submission.layer_id == layer_id,
        Submission.status == 'approved',
        Submission.doc_type == 'draft'
    ).order_by(Submission.submitted_at.desc()).all()
    return jsonify({
        'submissions': [{
            'id': s.id,
            'public_id': s.public_id,
            'artifact_id': s.artifact_id,
            'title': s.title,
            'draft_name': s.draft_name,
            'ml_number': s.ml_number,
            'group': s.group,
            'status': s.status,
            'submitted_at': s.submitted_at.isoformat() if s.submitted_at else None
        } for s in submissions]
    })


@bp.route('/api/layers/<layer_id>/docs/', methods=['GET'])
def list_layer_docs(layer_id):
    """List all docs (submissions of doc_type='draft') for a layer — any status.

    Used by the layer detail page 'Docs' tab so contributors can find a draft
    they added to a workgroup-less layer. Returns the most recent first; pending
    (status=submitted) drafts are included so the author can locate their own work
    in progress. Drafts marked deleted are filtered out.
    """
    Layer.query.get_or_404(layer_id)
    submissions = Submission.query.filter(
        Submission.layer_id == layer_id,
        Submission.doc_type == 'draft',
    ).order_by(Submission.submitted_at.desc()).limit(200).all()
    return jsonify({
        'docs': [{
            'id': s.id,
            'public_id': s.public_id,
            'title': s.title or s.draft_name or 'Untitled',
            'draft_name': s.draft_name,
            'ml_number': s.ml_number,
            'group': s.group,
            'status': s.status,
            'document_category': s.document_category,
            'submitted_at': s.submitted_at.isoformat() if s.submitted_at else None,
            'submitted_by': s.submitted_by,
        } for s in submissions]
    })


@bp.route('/submit/approve/<submission_id>', methods=['POST'])
@require_role('admin')
def approve_submission(submission_id):
    submission = get_submission_by_ref(submission_id)
    if not submission:
        flash('Submission not found', 'error')
        return redirect(url_for('submissions.admin_submissions'))

    is_revision = getattr(submission, 'is_revision', False)
    revision_num = getattr(submission, 'revision_number', '')
    ml_num = submission.ml_number
    parent_draft_name = getattr(submission, 'parent_draft_name', '')

    if is_revision and revision_num and ml_num:
        existing_revision = Submission.query.filter(
            Submission.ml_number == ml_num,
            Submission.revision_number == revision_num,
            Submission.status == 'approved',
            Submission.id != submission_id
        ).first()
        if existing_revision:
            all_revisions = Submission.query.filter(
                Submission.ml_number == ml_num,
                Submission.status == 'approved',
                Submission.revision_number.isnot(None)
            ).all()
            existing_nums = []
            for rev in all_revisions:
                try:
                    existing_nums.append(int(rev.revision_number))
                except (ValueError, TypeError):
                    pass
            next_num = 1
            while next_num in existing_nums:
                next_num += 1
            submission.revision_number = f"{next_num:02d}"

    if is_revision and parent_draft_name:
        parent_submission = get_submission_by_ref(parent_draft_name)
        if parent_submission and parent_submission.ml_number:
            revision_num = getattr(submission, 'revision_number', '')
            if revision_num:
                existing_revision = Submission.query.filter(
                    Submission.ml_number == parent_submission.ml_number,
                    Submission.revision_number == revision_num,
                    Submission.status == 'approved',
                    Submission.id != submission_id
                ).first()
                if existing_revision:
                    all_revisions = Submission.query.filter(
                        Submission.ml_number == parent_submission.ml_number,
                        Submission.status == 'approved',
                        Submission.revision_number.isnot(None)
                    ).all()
                    existing_nums = []
                    for rev in all_revisions:
                        try:
                            existing_nums.append(int(rev.revision_number))
                        except (ValueError, TypeError):
                            pass
                    next_num = 1
                    while next_num in existing_nums:
                        next_num += 1
                    submission.revision_number = f"{next_num:02d}"
            submission.ml_number = parent_submission.ml_number
        elif not submission.ml_number:
            try:
                doc_type = getattr(submission, 'doc_type', 'draft') or 'draft'
                submission.ml_number = get_next_ml_number(doc_type, layer_prefix=_layer_prefix_for_submission(submission))
            except Exception:
                pass
    elif not submission.ml_number:
        try:
            doc_type = getattr(submission, 'doc_type', 'draft') or 'draft'
            submission.ml_number = get_next_ml_number(doc_type, layer_prefix=_layer_prefix_for_submission(submission))
        except Exception:
            pass

    old_status = submission.status
    submission.status = 'approved'
    submission.approved_at = datetime.utcnow()
    try:
        db.session.commit()
    except Exception as e:
        from flask import current_app
        current_app.logger.error(f"Failed to commit approval: {e}")
        db.session.rollback()
        flash(f'Failed to approve submission: {str(e)}', 'error')
        return redirect(url_for('submissions.admin_submissions'))

    admin_user = get_current_user()
    try:
        from services.submission_notifications import (
            emit_submission_status_notification,
            run_submission_notification_dispatch,
        )

        bundle = emit_submission_status_notification(
            submission,
            actor_user_id=admin_user['id'],
            old_status=old_status,
            new_status='approved',
        )
        if bundle:
            run_submission_notification_dispatch(bundle, admin_user['id'])
    except Exception as ex:
        from flask import current_app

        current_app.logger.warning('submission notification dispatch (approve): %s', ex)
    try:
        from services.workgroup_links import assign_dp_workgroup_for_submission

        if assign_dp_workgroup_for_submission(submission):
            db.session.commit()
    except Exception as ex:
        from flask import current_app

        current_app.logger.warning('DP workgroup document link (approve): %s', ex)
    action_desc = f"Approved revision {submission.revision_number} of {parent_draft_name}" if is_revision else f"Approved submission: {submission.title}"
    add_to_document_history(f"submission-{submission.id}", "approved", admin_user['name'], action_desc)

    flash_msg = f'Revision {submission.id} approved! ML number: {submission.ml_number}' if is_revision else f'Submission {submission.id} approved! ML number: {submission.ml_number}'
    flash(flash_msg, 'success')
    return redirect(url_for('submissions.submission_detail', submission_id=submission_id))


@bp.route('/submit/reject/<submission_id>', methods=['POST'])
@require_role('admin')
def reject_submission(submission_id):
    submission = get_submission_by_ref(submission_id)
    if not submission:
        flash('Submission not found', 'error')
        return redirect(url_for('submissions.admin_submissions'))

    submission.status = 'rejected'
    submission.rejected_at = datetime.utcnow()
    db.session.commit()

    admin_user = get_current_user()
    add_to_document_history(f"submission-{submission.id}", "rejected", admin_user['name'],
                           f"Rejected submission: {submission.title}")

    flash(f'Submission {submission.id} rejected!', 'warning')
    return redirect(url_for('submissions.submission_detail', submission_id=submission_id))


@bp.route('/view/<submission_id>')
@require_auth
def view_submission(submission_id):
    """View a submission file inline (for PDFs and other viewable files)."""
    submission = get_submission_by_ref(submission_id)
    if not submission:
        return "Submission not found", 404

    current_user = get_current_user()
    if submission.submitted_by != current_user['name'] and current_user.get('role') not in ['admin', 'editor']:
        return "Access denied", 403

    if not submission.file_path or not os.path.exists(submission.file_path):
        return "File not found", 404

    return send_file(submission.file_path, as_attachment=False, download_name=submission.filename)


@bp.route('/download/<submission_id>')
@require_auth
def download_submission(submission_id):
    """Download a submission file."""
    submission = get_submission_by_ref(submission_id)
    if not submission:
        return "Submission not found", 404

    current_user = get_current_user()
    if submission.submitted_by != current_user['name'] and current_user.get('role') not in ['admin', 'editor']:
        return "Access denied", 403

    if not submission.file_path or not os.path.exists(submission.file_path):
        return "File not found", 404

    return send_file(submission.file_path, as_attachment=True, download_name=submission.filename)


@bp.route('/admin/submissions/')
@require_role('admin')
def admin_submissions():
    """Admin submission management page."""
    from services.rendering import _format_base_template, generate_user_menu
    from services.identity import get_current_user
    from config import BUILD_NUMBER

    user_menu = generate_user_menu()
    current_theme = get_current_user().get('theme', 'dark')

    status_filter = request.args.get('status', 'submitted')
    page = request.args.get('page', 1, type=int)
    per_page = 10

    query = Submission.query
    if status_filter and status_filter != 'all':
        query = query.filter_by(status=status_filter)

    submissions = query.order_by(Submission.submitted_at.desc()).paginate(
        page=page, per_page=per_page, error_out=False)

    submission_cards = ""
    for submission in submissions.items:
        status_badge = {
            'submitted': 'badge bg-warning text-dark',
            'approved': 'badge bg-success',
            'rejected': 'badge bg-danger',
            'published': 'badge bg-info',
            'inscription_pending': 'badge bg-warning'
        }.get(submission.status, 'badge bg-secondary')

        is_revision = getattr(submission, 'is_revision', False)
        revision_number = getattr(submission, 'revision_number', '')
        parent_draft_name = getattr(submission, 'parent_draft_name', '')
        revision_badge = f'<span class="badge bg-success ms-2">Revision {revision_number}</span>' if is_revision and revision_number else ''

        source_type = getattr(submission, 'sourceType', 'file')
        if source_type == 'ordinal':
            inscription_number = getattr(submission, 'inscriptionNumber', None)
            ordinal_id = getattr(submission, 'ordinalId', None)
            block_height = getattr(submission, 'blockHeight', None)
            if inscription_number:
                source_info = f'<span class="badge bg-info"><i class="bi bi-coin"></i> Ordinal</span> Inscription #{inscription_number}'
                if block_height:
                    source_info += f' (Block {block_height})'
            elif ordinal_id:
                source_info = f'<span class="badge bg-info"><i class="bi bi-coin"></i> Ordinal</span> {ordinal_id[:16]}...'
            else:
                source_info = '<span class="badge bg-info"><i class="bi bi-coin"></i> Ordinal</span>'
        else:
            file_size = "N/A"
            if submission.file_path and os.path.exists(submission.file_path):
                file_size = f"{os.path.getsize(submission.file_path) / 1024:.1f} KB"
            source_info = f'<span class="badge bg-secondary"><i class="bi bi-file-earmark"></i> File</span> {submission.filename} ({file_size})'

        action_buttons = ""
        if submission.status == 'submitted':
            action_buttons = f"""
            <button class="btn btn-success btn-sm me-2" onclick="approveSubmission('{submission.id}')">
                <i class="fas fa-check me-1"></i>Approve
            </button>
            <button class="btn btn-danger btn-sm me-2" onclick="rejectSubmission('{submission.id}')">
                <i class="fas fa-times me-1"></i>Reject
            </button>
            <button class="btn btn-info btn-sm" onclick="publishAsRFC('{submission.id}')">
                <i class="fas fa-star me-1"></i>Publish as RFC
            </button>
            """
        elif submission.status == 'approved':
            action_buttons = f"""
            <button class="btn btn-info btn-sm me-2" onclick="publishAsRFC('{submission.id}')">
                <i class="fas fa-star me-1"></i>Publish as RFC
            </button>
            <button class="btn btn-warning btn-sm" onclick="unapproveSubmission('{submission.id}')">
                <i class="fas fa-undo me-1"></i>Unapprove
            </button>
            """

        revision_context = ""
        if is_revision and parent_draft_name:
            revision_context = f'<p class="mb-2"><strong>Revision of:</strong> <a href="/doc/draft/{parent_draft_name}/" class="text-decoration-none">{parent_draft_name}</a></p>'

        authors = submission.authors if isinstance(submission.authors, list) else [submission.authors] if submission.authors else []
        submission_cards += f"""
        <div class="card mb-3">
            <div class="card-header d-flex justify-content-between align-items-center">
                <h6 class="mb-0">
                    <a href="/doc/draft/{submission.id}/" class="text-decoration-none">
                        {submission.title}
                    </a>
                </h6>
                <div>
                    <span class="{status_badge}">{submission.status.title()}</span>
                    {revision_badge}
                </div>
            </div>
            <div class="card-body">
                <div class="row">
                    <div class="col-md-8">
                        <p class="mb-2"><strong>Authors:</strong> {', '.join(authors)}</p>
                        <p class="mb-2"><strong>Group:</strong> {submission.group or 'None'}</p>
                        <p class="mb-2"><strong>Submitted:</strong> {submission.submitted_at.strftime('%Y-%m-%d %H:%M')} by {submission.submitted_by}</p>
                        <p class="mb-2"><strong>Source:</strong> {source_info}</p>
                        {revision_context}
                        {f'<p class="mb-2"><strong>Abstract:</strong> {submission.abstract[:200]}...</p>' if submission.abstract else ''}
                    </div>
                    <div class="col-md-4">
                        <div class="d-grid gap-2">
                            <a href="/doc/draft/{submission.id}/" class="btn btn-outline-primary btn-sm">
                                <i class="fas fa-eye me-1"></i>View Draft
                            </a>
                            {action_buttons}
                        </div>
                    </div>
                </div>
            </div>
        </div>
        """

    status_options = f"""
    <option value="all" {'selected' if status_filter == 'all' else ''}>All Submissions</option>
    <option value="submitted" {'selected' if status_filter == 'submitted' else ''}>Pending Review</option>
    <option value="approved" {'selected' if status_filter == 'approved' else ''}>Approved</option>
    <option value="rejected" {'selected' if status_filter == 'rejected' else ''}>Rejected</option>
    <option value="published" {'selected' if status_filter == 'published' else ''}>Published</option>
    """

    content = f"""
    <div class="gh-page container mt-4 gh-admin-page">
        {gh_page_header('Submission Management', 'Review and moderate draft submissions', 'fa-file-alt', actions_html=f'<select class="form-select form-select-sm" onchange="changeStatusFilter(this.value)">{status_options}</select>', breadcrumb_html=gh_breadcrumb([('Admin Dashboard', '/admin/'), ('Submission Management', None)]))}

        <div class="row mb-4">
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-warning">{Submission.query.filter_by(status='submitted').count()}</h4>
                        <p class="mb-0 small">Pending Review</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-success">{Submission.query.filter_by(status='approved').count()}</h4>
                        <p class="mb-0 small">Approved</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-danger">{Submission.query.filter_by(status='rejected').count()}</h4>
                        <p class="mb-0 small">Rejected</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card text-center">
                    <div class="card-body">
                        <h4 class="text-info">{Submission.query.filter_by(status='published').count()}</h4>
                        <p class="mb-0 small">Published as RFC</p>
                    </div>
                </div>
            </div>
        </div>

        <div id="submissions-container">
            {submission_cards}
        </div>

        {f'''
        <nav aria-label="Submission pagination" class="mt-4">
            <ul class="pagination justify-content-center">
                {f'<li class="page-item {"disabled" if not submissions.has_prev else ""}"><a class="page-link" href="?page={submissions.prev_num}&status={status_filter}">Previous</a></li>' if submissions.has_prev else ''}
                {''.join([f'<li class="page-item {"active" if i == submissions.page else ""}"><a class="page-link" href="?page={i}&status={status_filter}">{i}</a></li>' for i in (submissions.iter_pages() or [])])}
                {f'<li class="page-item {"disabled" if not submissions.has_next else ""}"><a class="page-link" href="?page={submissions.next_num}&status={status_filter}">Next</a></li>' if submissions.has_next else ''}
            </ul>
        </nav>
        ''' if submissions.pages and submissions.pages > 1 else ''}
    </div>

    <script>
        function changeStatusFilter(status) {{
            window.location.href = '?status=' + status;
        }}

        async function approveSubmission(submissionId) {{
            const ok = await GhDialog.confirm({{
                title: 'Approve submission',
                message: 'Approve this draft submission? It will be marked as approved and ready for publication.',
                variant: 'warning',
                confirmLabel: 'Approve',
            }});
            if (ok) {{
                updateSubmissionStatus(submissionId, 'approved');
            }}
        }}

        async function rejectSubmission(submissionId) {{
            const reason = await GhDialog.prompt({{
                title: 'Reject submission',
                message: 'Reason for rejection (optional):',
                inputType: 'text',
                placeholder: 'Optional rejection reason',
                confirmLabel: 'Reject',
                required: false,
            }});
            updateSubmissionStatus(submissionId, 'rejected', reason || null);
        }}

        async function unapproveSubmission(submissionId) {{
            const ok = await GhDialog.confirm({{
                title: 'Remove approval',
                message: 'Remove approval for this submission? It will return to the submitted state.',
                variant: 'warning',
                confirmLabel: 'Remove approval',
            }});
            if (ok) {{
                updateSubmissionStatus(submissionId, 'submitted');
            }}
        }}

        async function publishAsRFC(submissionId) {{
            const rfcNumber = await GhDialog.prompt({{
                title: 'Publish as RFC',
                message: 'Enter the RFC number for this submission:',
                inputType: 'text',
                placeholder: 'e.g. 9999',
                confirmLabel: 'Continue',
                required: true,
            }});
            if (!rfcNumber) return;
            const confirmed = await GhDialog.confirm({{
                title: 'Confirm publication',
                message: 'Publish as RFC ' + rfcNumber + '?',
                variant: 'warning',
                confirmLabel: 'Publish',
            }});
            if (confirmed) {{
                updateSubmissionStatus(submissionId, 'published', null, rfcNumber);
            }}
        }}

        async function updateSubmissionStatus(submissionId, status, reason = null, rfcNumber = null) {{
            const data = {{ status: status }};
            if (reason) data.reason = reason;
            if (rfcNumber) data.rfc_number = rfcNumber;

            try {{
                const response = await fetch('/admin/submissions/' + submissionId + '/status', {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify(data),
                }});
                const result = await response.json();
                if (result.success) {{
                    location.reload();
                }} else {{
                    await GhDialog.alert({{ title: 'Update failed', message: ('Error: ' + (result.message || 'Unknown error')), variant: 'danger' }});
                }}
            }} catch (error) {{
                console.error('Error:', error);
                await GhDialog.alert({{ title: 'Update failed', message: 'Error updating submission status', variant: 'danger' }});
            }}
        }}
    </script>
    """

    return _format_base_template(
        title="Submission Management - GovHub",
        theme=current_theme,
        user_menu=user_menu,
        content=content, build_number=BUILD_NUMBER)


@bp.route('/admin/submissions/<submission_id>/status', methods=['POST'])
@require_role('admin')
def update_submission_status(submission_id):
    """Update submission status (approve, reject, publish, unapprove)."""
    data = request.get_json() or {}
    new_status = data.get('status', '')
    reason = data.get('reason', '')
    rfc_number = data.get('rfc_number', '')

    if new_status not in ['submitted', 'approved', 'rejected', 'published']:
        return jsonify({'success': False, 'message': 'Invalid status'}), 400

    submission = get_submission_by_ref(submission_id)
    if not submission:
        return jsonify({'success': False, 'message': 'Submission not found'}), 404

    old_status = submission.status
    submission.status = new_status

    if new_status == 'approved':
        is_revision = getattr(submission, 'is_revision', False)
        revision_num = getattr(submission, 'revision_number', '')
        ml_num = submission.ml_number

        if is_revision and revision_num and ml_num:
            existing_revision = Submission.query.filter(
                Submission.ml_number == ml_num,
                Submission.revision_number == revision_num,
                Submission.status == 'approved',
                Submission.id != submission_id
            ).first()

            if existing_revision:
                all_revisions = Submission.query.filter(
                    Submission.ml_number == ml_num,
                    Submission.status == 'approved',
                    Submission.revision_number.isnot(None)
                ).all()
                existing_nums = []
                for rev in all_revisions:
                    try:
                        existing_nums.append(int(rev.revision_number))
                    except (ValueError, TypeError):
                        pass
                next_num = 1
                while next_num in existing_nums:
                    next_num += 1
                submission.revision_number = f"{next_num:02d}"
                current_app.logger.warning(f"Duplicate revision detected for {ml_num}, auto-assigned {submission.revision_number}")

    if new_status == 'approved' and not submission.ml_number:
        is_revision = getattr(submission, 'is_revision', False)
        parent_draft_name = getattr(submission, 'parent_draft_name', '')

        if is_revision and parent_draft_name:
            parent_submission = get_submission_by_ref(parent_draft_name)
            if parent_submission and parent_submission.ml_number:
                revision_num = getattr(submission, 'revision_number', '')
                if revision_num:
                    existing_revision = Submission.query.filter(
                        Submission.ml_number == parent_submission.ml_number,
                        Submission.revision_number == revision_num,
                        Submission.status == 'approved',
                        Submission.id != submission_id
                    ).first()
                    if existing_revision:
                        all_revisions = Submission.query.filter(
                            Submission.ml_number == parent_submission.ml_number,
                            Submission.status == 'approved',
                            Submission.revision_number.isnot(None)
                        ).all()
                        existing_nums = []
                        for rev in all_revisions:
                            try:
                                existing_nums.append(int(rev.revision_number))
                            except (ValueError, TypeError):
                                pass
                        next_num = 1
                        while next_num in existing_nums:
                            next_num += 1
                        submission.revision_number = f"{next_num:02d}"
                submission.ml_number = parent_submission.ml_number
                submission.approved_at = datetime.utcnow()
            else:
                try:
                    doc_type = getattr(submission, 'doc_type', 'draft') or 'draft'
                    ml_number = get_next_ml_number(doc_type, layer_prefix=_layer_prefix_for_submission(submission))
                    submission.ml_number = ml_number
                    submission.approved_at = datetime.utcnow()
                except Exception as e:
                    current_app.logger.error(f"Failed to assign ML number: {e}")
        else:
            try:
                doc_type = getattr(submission, 'doc_type', 'draft') or 'draft'
                ml_number = get_next_ml_number(doc_type, layer_prefix=_layer_prefix_for_submission(submission))
                submission.ml_number = ml_number
                submission.approved_at = datetime.utcnow()
            except Exception as e:
                current_app.logger.error(f"Failed to assign ML number: {e}")

    if new_status == 'rejected' and reason:
        submission.rejected_at = datetime.utcnow()

    if new_status == 'published':
        if rfc_number:
            try:
                submission.rfc_number = int(rfc_number)
            except (ValueError, TypeError):
                return jsonify({'success': False, 'message': 'Invalid RFC number'}), 400
        if submission.ml_number and '-' in submission.ml_number:
            # ml_number may be either 'ML-Draft-001' or 'CL-Draft-001' (per-layer
            # prefix). Preserve the leading 2-letter prefix when promoting the
            # draft to an RFC so layer-prefixed drafts keep their prefix.
            parts = submission.ml_number.split('-')
            if len(parts) >= 3 and parts[1].upper() == 'DRAFT' and parts[2].isdigit():
                prefix_token = parts[0]
                submission.ml_number = f"{prefix_token}-RFC-{parts[2]}"
        submission.doc_type = 'rfc'

    db.session.commit()

    try:
        from services.submission_notifications import (
            emit_submission_status_notification,
            run_submission_notification_dispatch,
        )

        au = get_current_user()
        if au and new_status in ('approved', 'published'):
            bundle = emit_submission_status_notification(
                submission,
                actor_user_id=au['id'],
                old_status=old_status,
                new_status=new_status,
                rfc_number=submission.rfc_number if new_status == 'published' else None,
            )
            if bundle:
                run_submission_notification_dispatch(bundle, au['id'])
    except Exception as ex:
        current_app.logger.warning('submission notification dispatch (status api): %s', ex)

    admin_user = get_current_user()
    action_details = f"Changed status from {old_status} to {new_status}"
    if reason:
        action_details += f" - Reason: {reason}"
    if rfc_number:
        action_details += f" - Published as RFC {rfc_number}"

    add_to_document_history(f"submission-{submission.id}", "status_changed",
                           admin_user['name'], action_details)

    return jsonify({'success': True, 'message': f'Status updated to {new_status}'})
