"""Documents routes: /doc/active, /doc/all, /doc/draft/*, comments, follow, history, revisions, /test/."""
import json
import os
import re
import html as html_mod
from datetime import datetime, timedelta
from urllib.parse import quote

from flask import Blueprint, request, redirect, url_for, flash, session, Response, current_app, jsonify, send_file

from sqlalchemy import or_

from extensions import db
from models import (
    Comment, DocumentHistory, Submission, Artifact, ArtifactRelation,
    Layer,
)
from services.identity import get_current_user, require_auth
from services.events import emit_event
from services.document_follow_notifications import dispatch_document_followers
from services.event_subscriptions import matrix_from_subscription_post, replace_draft_subscriptions_matrix
from services.submissions import get_submission_by_ref, add_to_document_history, can_edit_submission_metadata
from services.ordinals import (
    process_ordinal_markdown,
    shorten_inscription_id,
    looks_like_html_inscription,
    format_ordinal_html_iframe_preview,
)
from services.submission_preview_md import markdown_to_safe_preview_html, text_looks_like_markdown
from services.documents import (
    load_draft_data,
    DRAFTS,
    build_comment_tree,
    render_comment_tree,
    toggle_comment_like,
    get_comment_likes,
    is_comment_liked,
    render_draft_subscription_form_html,
    add_comment_reply,
    EDIT_DELETE_TIME_MINUTES,
    sort_documents_by_ml_number_desc,
    submission_file_pages_words,
    revision_notes_to_safe_html,
)
from services.draft_reader import (
    build_draft_context,
    draft_display_id,
    load_draft_document_body,
)
from services.directory_ui import gh_page_header, gh_living_module, gh_breadcrumb, gh_filter_row, gh_directory_toolbar
from services.workgroup_links import build_document_workgroup_index, resolve_document_workgroup_meta

bp = Blueprint('documents', __name__, url_prefix='')


def _get_drafts():
    """Get DRAFTS list (cached in services.documents)."""
    return DRAFTS


def _document_workgroup_table_row(draft: dict) -> str:
    """HTML table row for linked workgroup, or empty when none."""
    meta = resolve_document_workgroup_meta(
        name=draft.get('name'),
        group=draft.get('group'),
        title=draft.get('title'),
    )
    label = meta.get('workgroup_name') or meta.get('group')
    if not label:
        return ''
    href = meta.get('workgroup_href')
    if href:
        cell = (
            f'<a href="{html_mod.escape(href, quote=True)}" class="text-decoration-none" '
            f'style="color: var(--accent-color) !important;">{html_mod.escape(label)}</a>'
        )
    else:
        cell = html_mod.escape(label)
    return (
        '<tr><td style="color: var(--text-secondary) !important;"><strong>Workgroup:</strong></td>'
        f'<td style="color: var(--text-primary) !important;">{cell}</td></tr>'
    )


def _submission_uses_display_ordinal(submission):
    if not submission:
        return False
    src = (getattr(submission, 'displayBodySource', None) or 'file').strip().lower()
    return src == 'ordinal' and bool(getattr(submission, 'displayOrdinalContentUrl', None))


def _can_manage_submission_display_body(user, submission):
    if not user or not submission:
        return False
    role = user.get('role')
    if role in ('admin', 'editor'):
        return True
    if (submission.sourceType or 'file').strip().lower() != 'file':
        return False
    sub_by = (submission.submitted_by or '').strip()
    uname = (user.get('name') or '').strip()
    return bool(sub_by and uname and sub_by == uname)


def _load_ordinal_document_body(ordinal_content_url, ordinal_content_type, draft):
    """
    Fetch and render inscription body (markdown / HTML iframe / image).
    Returns (document_content, render_document_as_html, calculated_pages, calculated_words).
    """
    from flask import current_app

    document_content = ''
    render_document_as_html = False
    calculated_pages = draft.get('pages', 1)
    calculated_words = draft.get('words', 0)
    octype = ordinal_content_type or ''

    if ordinal_content_url and ('text/' in octype or 'application/json' in octype):
        try:
            import requests
            from services.url_safety import validate_ordinals_fetch_url

            safe_url = validate_ordinals_fetch_url(ordinal_content_url)
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(safe_url, headers=headers, timeout=10)
            if response.status_code == 200:
                raw_content = response.text
                words = len(raw_content.split())
                calculated_pages = max(1, (words + 499) // 500)
                calculated_words = words
                draft['pages'] = calculated_pages
                draft['words'] = calculated_words

                is_markdown = False
                markdown_patterns = [
                    r'^#{1,6}\s+.+$',
                    r'\*\*.+\*\*',
                    r'\*.+\*',
                    r'^\s*[-*+]\s+',
                    r'^\s*\d+\.\s+',
                    r'\[.+\]\(.+\)',
                    r'!\[.*\]\(.+\)',
                ]
                for pattern in markdown_patterns:
                    if re.search(pattern, raw_content, re.MULTILINE):
                        is_markdown = True
                        break

                if looks_like_html_inscription(raw_content, ordinal_content_type):
                    document_content = format_ordinal_html_iframe_preview(ordinal_content_url)
                    render_document_as_html = True
                elif is_markdown:
                    document_content = process_ordinal_markdown(raw_content)
                    render_document_as_html = True
                else:
                    document_content = raw_content
        except Exception as e:
            current_app.logger.warning(f'Failed to fetch ordinal content for display: {e}')
            document_content = f'Error loading ordinal content: {str(e)}'
    elif ordinal_content_url and octype.startswith('image/'):
        document_content = (
            f'<img src="{html_mod.escape(ordinal_content_url, quote=True)}" '
            f'class="img-fluid" style="max-width: 100%;" alt="Ordinal image content">'
        )
        render_document_as_html = True
    else:
        document_content = (
            f'Ordinal content type: {octype}\nPreview not available for this content type.'
        )

    return document_content, render_document_as_html, calculated_pages, calculated_words


@bp.route('/doc/active/')
def active_documents():
    """Show active documents (alias for all documents)."""
    return all_documents()


def _build_all_documents_catalog():
    """Build JSON-serializable rows for the /doc/all/ directory (client search & sort)."""
    drafts = _get_drafts()
    wg_index = build_document_workgroup_index()
    all_docs = []
    for draft in drafts:
        wg_meta = resolve_document_workgroup_meta(
            name=draft.get('name'),
            group=draft.get('group'),
            title=draft.get('title'),
            index=wg_index,
        )
        all_docs.append({
            'name': draft.get('name'),
            'title': draft.get('title') or '',
            'authors': draft.get('authors') if isinstance(draft.get('authors'), list) else [],
            'group': wg_meta.get('group'),
            'workgroup_name': wg_meta.get('workgroup_name'),
            'workgroup_href': wg_meta.get('workgroup_href'),
            'status': draft.get('status'),
            'rev': draft.get('rev', '00'),
            'pages': draft.get('pages') or 1,
            'words': draft.get('words') or 0,
            'date': draft.get('date') or '',
            'abstract': draft.get('abstract') or '',
            'ml_number': draft.get('ml_number'),
            'is_revision': bool(draft.get('is_revision')),
            'revision_number': draft.get('revision_number') or '',
            'submitted_at': draft.get('date') or '',
        })

    approved_submissions = Submission.query.filter(
        Submission.status.in_(['approved', 'published']),
        Submission.is_revision == False,
    ).all()

    parent_refs = set()
    for submission in approved_submissions:
        parent_refs.add(submission.id)
        if submission.draft_name:
            parent_refs.add(submission.draft_name)

    latest_revisions = {}
    if parent_refs:
        revisions = Submission.query.filter(
            Submission.parent_draft_name.in_(parent_refs),
            Submission.is_revision == True,
            Submission.status.in_(['approved', 'published']),
        ).order_by(Submission.submitted_at.desc()).all()
        for revision in revisions:
            parent = getattr(revision, 'parent_draft_name', None)
            if parent and parent not in latest_revisions:
                latest_revisions[parent] = revision

    for submission in approved_submissions:
        latest_revision = latest_revisions.get(submission.id)
        if not latest_revision and submission.draft_name:
            latest_revision = latest_revisions.get(submission.draft_name)
        display_submission = latest_revision if latest_revision else submission
        is_revision = getattr(display_submission, 'is_revision', False)
        revision_number = getattr(display_submission, 'revision_number', '') or ''
        pages, words = submission_file_pages_words(display_submission)
        submitted_at = (
            display_submission.submitted_at.isoformat()
            if display_submission.submitted_at else ''
        )
        wg_meta = resolve_document_workgroup_meta(
            name=display_submission.id,
            group=display_submission.group,
            title=display_submission.title,
            index=wg_index,
        )
        all_docs.append({
            'name': display_submission.id,
            'title': display_submission.title or '',
            'authors': (
                display_submission.authors
                if isinstance(display_submission.authors, list)
                else [display_submission.authors] if display_submission.authors else []
            ),
            'group': wg_meta.get('group'),
            'workgroup_name': wg_meta.get('workgroup_name'),
            'workgroup_href': wg_meta.get('workgroup_href'),
            'status': display_submission.status,
            'rev': revision_number if is_revision else '00',
            'pages': pages or 1,
            'words': words or 0,
            'date': display_submission.submitted_at.strftime('%Y-%m-%d') if display_submission.submitted_at else '',
            'abstract': display_submission.abstract or '',
            'ml_number': display_submission.ml_number,
            'is_revision': is_revision,
            'revision_number': revision_number,
            'submitted_at': submitted_at,
        })

    return sort_documents_by_ml_number_desc(all_docs)


@bp.route('/doc/all/')
def all_documents():
    from flask import url_for
    from services.rendering import _format_base_template, generate_user_menu
    from config import BUILD_NUMBER
    from services.page_heroes import render_page_hero_html

    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    catalog = _build_all_documents_catalog()
    docs_json = json.dumps(catalog)
    total_docs = len(catalog)
    submit_url = url_for('submissions.submit_draft')

    doc_view_actions = (
        '<div class="btn-group" role="group" aria-label="View mode">'
        '<button type="button" class="btn btn-outline-secondary btn-sm active" id="doc-view-cards" onclick="setDocView(\'cards\')">Cards</button>'
        '<button type="button" class="btn btn-outline-secondary btn-sm" id="doc-view-list" onclick="setDocView(\'list\')">List</button>'
        '</div>'
    )
    submit_draft_col = (
        '<div class="col-md-auto ms-md-auto d-flex align-items-end">'
        f'<a href="{submit_url}" class="btn btn-primary">Submit Draft</a>'
        '</div>'
    )

    content = f"""
    <div class="gh-page container doc-all-page mt-4">
        {render_page_hero_html('docs_drafts')}
        {gh_page_header(
            'Docs & Drafts',
            f'{total_docs} documents in the directory',
            'fa-file-alt',
            actions_html=doc_view_actions,
        )}
        {gh_filter_row(
            gh_directory_toolbar(
                search_placeholder='Search documents…',
                search_col='col-md-5',
                sort_col='col-md-3',
                extra_cols=submit_draft_col,
                sort_options=(
                    ('recent', 'ML number (newest first)'),
                    ('name-asc', 'A–Z'),
                    ('name-desc', 'Z–A'),
                ),
            )
        )}
        <p class="text-muted small mb-3" id="doc-all-count"></p>
        <div id="doc-all-container"></div>
    </div>
    <script>
    const allDocItems = {docs_json};
    let docViewMode = 'cards';

    function setDocView(mode) {{
        docViewMode = mode === 'list' ? 'list' : 'cards';
        document.getElementById('doc-view-cards').classList.toggle('active', docViewMode === 'cards');
        document.getElementById('doc-view-list').classList.toggle('active', docViewMode === 'list');
        renderDocDirectory();
    }}

    function docHref(name) {{
        return encodeURIComponent(String(name || ''));
    }}

    function docWorkgroupHtml(d) {{
        const label = d.workgroup_name || d.group;
        const inner = label
            ? (d.workgroup_href
                ? '<a href="' + GhDirectory.esc(d.workgroup_href) + '">' + GhDirectory.esc(label) + '</a>'
                : GhDirectory.esc(label))
            : '&nbsp;';
        return '<br>Workgroup: ' + inner;
    }}

    function renderDocCards(docs) {{
        return docs.map(function(d) {{
            const displayId = GhDirectory.esc(d.ml_number || d.name || '');
            const href = docHref(d.name);
            const revBadge = (d.is_revision && d.revision_number)
                ? '<span class="badge bg-success ms-2">Revision ' + GhDirectory.esc(d.revision_number) + '</span>' : '';
            const authors = Array.isArray(d.authors) ? d.authors.join(', ') : (d.authors || 'N/A');
            const words = d.words || 0;
            return '<div class="col-md-6 document-card"><div class="card"><div class="card-body">'
                + '<h5 class="card-title document-title"><a href="/doc/draft/' + href + '/">' + displayId + '</a>' + revBadge + '</h5>'
                + '<p class="card-text">' + GhDirectory.esc(d.title || '') + '</p>'
                + '<div class="document-meta"><span class="badge bg-secondary status-badge">' + GhDirectory.esc(d.status || '') + '</span>'
                + '<span>Rev: ' + GhDirectory.esc(d.rev || '00') + '</span>'
                + '<span>' + (d.pages || 1) + ' pages</span>'
                + '<span>' + words + ' words</span></div>'
                + '<div class="mt-2"><small class="text-muted">Authors: ' + GhDirectory.esc(authors)
                + docWorkgroupHtml(d)
                + '<br>Date: ' + GhDirectory.esc(d.date || '') + '</small></div>'
                + '<div class="mt-2 doc-card-actions">'
                + '<a href="/doc/draft/' + href + '/read/" class="btn btn-sm btn-primary">Read</a>'
                + '<a href="/doc/draft/' + href + '/comments/" class="btn btn-sm btn-outline-primary">Comments</a>'
                + '</div></div></div></div>';
        }}).join('');
    }}

    function renderDocList(docs) {{
        const rows = docs.map(function(d) {{
            const displayId = GhDirectory.esc(d.ml_number || d.name || '');
            const href = docHref(d.name);
            return '<tr>'
                + '<td><a href="/doc/draft/' + href + '/">' + displayId + '</a></td>'
                + '<td class="doc-all-title-cell" title="' + GhDirectory.esc(d.title || '') + '">' + GhDirectory.esc(d.title || '') + '</td>'
                + '<td><span class="badge bg-secondary">' + GhDirectory.esc(d.status || '') + '</span></td>'
                + '<td>' + GhDirectory.esc(d.revision_number || d.rev || '00') + '</td>'
                + '<td>' + (d.pages || 1) + '</td>'
                + '<td>' + GhDirectory.esc(d.date || '') + '</td>'
                + '<td class="text-nowrap"><a href="/doc/draft/' + href + '/read/" class="btn btn-sm btn-primary me-1">Read</a>'
                + '<a href="/doc/draft/' + href + '/comments/" class="btn btn-sm btn-outline-primary">Comments</a></td></tr>';
        }}).join('');
        return '<style>.doc-all-list-table th,.doc-all-list-table td{{white-space:nowrap}}.doc-all-list-table .doc-all-title-cell{{max-width:22rem;overflow:hidden;text-overflow:ellipsis}}</style>'
            + '<div class="table-responsive"><table class="table table-hover align-middle doc-all-list-table"><thead><tr>'
            + '<th>ID</th><th>Title</th><th>Status</th><th>Rev</th><th>Pages</th><th>Date</th><th></th></tr></thead><tbody>'
            + (rows || '<tr><td colspan="7" class="text-muted">No documents match your search.</td></tr>')
            + '</tbody></table></div>';
    }}

    function renderDocDirectory() {{
        const items = GhDirectory.filterAndSort(allDocItems, {{
            searchTerm: GhDirectory.getSearchValue('search-input'),
            sort: GhDirectory.getSortValue('sort-filter'),
            searchFields: ['title', 'ml_number', 'name', 'abstract', 'group', 'workgroup_name', 'authors'],
            nameKey: 'title',
            dateKeys: ['submitted_at', 'date'],
            recentSort: 'ml_number',
        }});
        const countEl = document.getElementById('doc-all-count');
        if (countEl) {{
            countEl.textContent = items.length === allDocItems.length
                ? ('Showing ' + items.length + ' documents')
                : ('Showing ' + items.length + ' of ' + allDocItems.length + ' documents');
        }}
        const container = document.getElementById('doc-all-container');
        if (!container) return;
        if (!items.length) {{
            container.innerHTML = '<div class="alert alert-info">No documents match your search.</div>';
            return;
        }}
        if (docViewMode === 'list') {{
            container.innerHTML = renderDocList(items);
        }} else {{
            container.innerHTML = '<div class="row g-3">' + renderDocCards(items) + '</div>';
        }}
    }}

    GhDirectory.bindControls('search-input', 'sort-filter', renderDocDirectory);
    renderDocDirectory();
    </script>
    """

    return _format_base_template(title="Documents - MLGH", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER)


@bp.route('/doc/draft/<path:draft_name>.txt')
def draft_text(draft_name):
    """Serve draft content as plain text."""
    DRAFTS = _get_drafts()
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)
    submission = None
    if not draft:
        submission = get_submission_by_ref(draft_name)
        if submission:
            draft = {
                'name': submission.draft_name or submission.id,
                'title': submission.title,
                'authors': submission.authors,
                'abstract': submission.abstract or 'Abstract not available for this draft.',
                'status': submission.status,
                'group': submission.group,
                'date': submission.submitted_at.strftime('%Y-%m-%d') if submission.submitted_at else '',
            }

    if not draft:
        return "Document not found", 404

    document_content = "Document content not available."

    if submission and submission.file_path and os.path.exists(submission.file_path):
        _, ext = os.path.splitext(submission.filename.lower())
        try:
            if ext in ['.txt', '.xml']:
                with open(submission.file_path, 'r', encoding='utf-8', errors='replace') as f:
                    document_content = f.read()
            elif ext == '.docx':
                from docx import Document
                doc = Document(submission.file_path)
                content_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content_parts.append(cell.text)
                document_content = '\n\n'.join(content_parts)
            elif ext == '.pdf':
                from PyPDF2 import PdfReader
                reader = PdfReader(submission.file_path)
                content_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(text)
                document_content = '\n\n'.join(content_parts)
                document_content = re.sub(r'\n+', '\n', document_content)
                document_content = re.sub(r' +', ' ', document_content)
            else:
                document_content = f"Document content cannot be displayed for {ext.upper()} files. Please download to view."
        except Exception as e:
            document_content = f"Error loading document content: {str(e)}"

    elif draft and 'name' in draft:
        document_content = f"""INTERNET-DRAFT                                               {', '.join(draft.get('authors', []))}
Intended status: Informational                            Meta-Layer Initiative
Expires: {draft.get('date', 'TBD')}                                      {draft.get('date', 'TBD')}


{draft.get('title', 'Document Title')}


Abstract

{draft.get('abstract', 'Abstract not available.')}


1. Introduction

This document describes {draft.get('title', 'the subject matter')}.

The content of this draft is currently being developed and will be available
in the full document once published.

2. Status of This Memo

This Internet-Draft is submitted in full conformance with the provisions
of BCP 78 and BCP 79.

Meta-Layer Drafts are working documents of the Meta-Layer Task Force
(MLGH). These documents represent proposals and specifications for the
Meta-Layer ecosystem. The list of current Meta-Layer Drafts is available
in the MLGH datatracker.

Internet-Drafts are draft documents valid for a maximum of six months and
may be updated, replaced, or obsoleted by other documents at any time. It is
inappropriate to use Internet-Drafts as reference material or to cite them
other than as "work in progress."

This Internet-Draft will expire on {draft.get('date', 'TBD')}.


3. References

[MLGH] MLGH Datatracker, https://rfc.themetalayer.org/

Authors' Addresses

{chr(10).join([f'{author} <email@example.com>' for author in draft.get('authors', [])])}

Meta-Layer Initiative
"""

    return Response(document_content, mimetype='text/plain; charset=utf-8')


@bp.route('/doc/draft/<path:draft_name>/read/')
def draft_reader(draft_name):
    """Full-page reader for a single draft (minimal chrome, document-focused)."""
    from html import escape as html_escape
    from services.rendering import _format_base_template, generate_user_menu
    from config import BUILD_NUMBER

    draft, submission = build_draft_context(draft_name)
    if not draft:
        return 'Document not found', 404

    display_id = draft_display_id(draft)
    title_escaped = html_escape(str(draft.get('title') or ''))
    doc_href = quote(str(draft.get('name') or draft_name), safe='')
    pdf_height = 'calc(100vh - 11rem)'
    document_content, render_html, pages, words = load_draft_document_body(
        draft,
        submission,
        draft_name,
        pdf_iframe_height=pdf_height,
    )

    if render_html:
        body_block = (
            f'<div class="draft-reader-body prose" id="dp-reader-selectable-body">'
            f'{document_content}</div>'
        )
    else:
        body_block = (
            f'<pre class="draft-reader-body draft-reader-pre" id="dp-reader-selectable-body">'
            f'{html_escape(document_content)}</pre>'
        )

    from services.dp_proposal_reader import render_dp_proposal_reader_assets
    from services.read_navigation import draft_reader_back_href

    back_href = html_mod.escape(draft_reader_back_href(request.args.get('return_to')))

    dp_proposal_assets = render_dp_proposal_reader_assets(
        submission,
        draft_name,
        render_html=render_html,
        document_content=document_content if isinstance(document_content, str) else '',
    )

    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')

    content = f'''
    <style>
      .draft-reader-page {{
        --reader-content-max: 52rem;
        --reader-gutter: 15px;
        max-width: 100%;
        padding: 0 var(--reader-gutter) 1rem;
      }}
      @media (min-width: 768px) {{
        .draft-reader-page {{ --reader-gutter: 24px; }}
      }}
      @media (min-width: 1200px) {{
        .draft-reader-page {{ --reader-gutter: 40px; }}
      }}
      .draft-reader-toolbar {{
        position: sticky;
        top: 0;
        z-index: 10;
        background: var(--navbar-bg);
        color: var(--text-primary);
        border-bottom: 1px solid var(--border-color);
        padding: 0.65rem var(--reader-gutter);
        margin: 0 calc(-1 * var(--reader-gutter)) 1rem;
        backdrop-filter: blur(10px);
        -webkit-backdrop-filter: blur(10px);
      }}
      .draft-reader-toolbar-inner {{
        max-width: var(--reader-content-max);
        margin: 0 auto;
        width: 100%;
        min-width: 0;
        display: flex;
        flex-direction: row;
        align-items: center;
        gap: 0.75rem;
        padding: 0.25rem 0;
      }}
      .draft-reader-nav {{
        display: flex;
        flex-shrink: 0;
        align-items: center;
        gap: 0.5rem;
      }}
      .draft-reader-meta {{
        flex: 1;
        min-width: 0;
        display: flex;
        align-items: center;
        gap: 0.35rem;
        overflow: hidden;
        white-space: nowrap;
        color: var(--text-secondary);
        font-size: 0.875rem;
        line-height: 1.4;
      }}
      .draft-reader-meta strong,
      .draft-reader-stats {{
        flex-shrink: 0;
        color: var(--text-primary);
      }}
      .draft-reader-meta .draft-reader-stats {{
        color: var(--text-secondary);
      }}
      .draft-reader-sep {{
        flex-shrink: 0;
        color: var(--text-muted);
      }}
      .draft-reader-title {{
        overflow: hidden;
        text-overflow: ellipsis;
        white-space: nowrap;
        min-width: 0;
      }}
      .draft-reader-toolbar .btn-outline-secondary {{
        color: var(--text-secondary);
        border-color: var(--border-color);
        background: transparent;
      }}
      .draft-reader-toolbar .btn-outline-secondary:hover,
      .draft-reader-toolbar .btn-outline-secondary:focus {{
        color: var(--text-primary);
        border-color: var(--border-hover);
        background: var(--bg-tertiary);
      }}
      .draft-reader-body {{ max-width: var(--reader-content-max); margin: 0 auto; }}
      .draft-reader-pre {{
        white-space: pre-wrap; word-wrap: break-word;
        font-family: var(--bs-font-monospace); font-size: 0.95rem;
        color: var(--text-primary);
      }}
      .draft-reader-body .pdf-viewer-container iframe {{
        width: 100% !important;
        min-height: {pdf_height};
      }}
    </style>
    <div class="draft-reader-page">
      <div class="draft-reader-toolbar">
        <div class="draft-reader-toolbar-inner">
          <div class="draft-reader-nav">
            <a href="{back_href}" class="btn btn-sm btn-outline-secondary" id="draftReaderBack">&larr; Back</a>
            <a href="/doc/draft/{doc_href}/" class="btn btn-sm btn-outline-secondary">Record</a>
          </div>
          <div class="draft-reader-meta">
            <strong>{html_escape(display_id)}</strong>
            <span class="draft-reader-sep">·</span>
            <span class="draft-reader-title" title="{title_escaped}">{title_escaped}</span>
            <span class="draft-reader-sep">·</span>
            <span class="draft-reader-stats">{int(pages)} pg</span>
          </div>
        </div>
      </div>
      {body_block}
    </div>
    {dp_proposal_assets}
    '''

    return _format_base_template(
        title=f'{display_id} — Reader',
        theme=current_theme,
        user_menu=user_menu,
        content=content,
        build_number=BUILD_NUMBER,
    )


@bp.route('/doc/draft/<path:draft_name>/')
def draft_detail(draft_name):
    from services.rendering import _format_base_template, generate_user_menu
    from config import BUILD_NUMBER

    drafts = _get_drafts()
    draft = next((d for d in drafts if d['name'] == draft_name), None)

    submission = None
    if not draft:
        submission = get_submission_by_ref(draft_name)
        if submission:
            source_type = getattr(submission, 'sourceType', 'file')
            pages_count = 1
            words_count = 0
            ordinal_content_url = getattr(submission, 'ordinalContentUrl', None)
            ordinal_content_type = getattr(submission, 'ordinalContentType', '')

            if source_type == 'ordinal':
                if ordinal_content_url and ('text/' in ordinal_content_type or 'application/json' in ordinal_content_type):
                    try:
                        import requests
                        from services.url_safety import validate_ordinals_fetch_url

                        safe_url = validate_ordinals_fetch_url(ordinal_content_url)
                        headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                        }
                        response = requests.get(safe_url, headers=headers, timeout=10)
                        if response.status_code == 200:
                            text_content = response.text
                            words_count = len(text_content.split())
                            pages_count = max(1, (words_count + 499) // 500)
                    except Exception as e:
                        current_app.logger.warning(f"Failed to fetch ordinal content for word/page count: {e}")
                        pass
            else:
                pages_count, words_count = submission_file_pages_words(submission)

            dbs = getattr(submission, 'displayBodySource', None) or 'file'
            displaying_linked = (
                dbs.strip().lower() == 'ordinal' and bool(getattr(submission, 'displayOrdinalContentUrl', None))
            )
            draft = {
                'name': submission.draft_name or submission.id,
                'title': submission.title,
                'authors': submission.authors,
                'abstract': submission.abstract or 'Abstract not available for this draft.',
                'status': submission.status,
                'group': submission.group,
                'date': submission.submitted_at.strftime('%Y-%m-%d') if submission.submitted_at else '',
                'rev': '00',
                'pages': pages_count,
                'words': words_count,
                'stream': 'mltf',
                'ml_number': submission.ml_number,
                'sourceType': source_type,
                'ordinalId': getattr(submission, 'ordinalId', None),
                'inscriptionNumber': getattr(submission, 'inscriptionNumber', None),
                'blockHeight': getattr(submission, 'blockHeight', None),
                'inscriptionTimestamp': getattr(submission, 'inscriptionTimestamp', None),
                'ordinalContentType': ordinal_content_type,
                'is_revision': getattr(submission, 'is_revision', False),
                'revision_number': getattr(submission, 'revision_number', ''),
                'parent_draft_name': getattr(submission, 'parent_draft_name', ''),
                'displayBodySource': dbs,
                'displayOrdinalId': getattr(submission, 'displayOrdinalId', None),
                'displayOrdinalContentUrl': getattr(submission, 'displayOrdinalContentUrl', None),
                'displayOrdinalContentType': getattr(submission, 'displayOrdinalContentType', None),
                'displayingLinkedOrdinal': displaying_linked,
            }

    if not draft:
        return "Document not found", 404

    if not submission:
        submission = get_submission_by_ref(draft.get('name')) or get_submission_by_ref(draft_name)
        if submission:
            dbs = getattr(submission, 'displayBodySource', None) or 'file'
            displaying_linked = (
                dbs.strip().lower() == 'ordinal'
                and bool(getattr(submission, 'displayOrdinalContentUrl', None))
            )
            draft['displayBodySource'] = dbs
            draft['displayOrdinalId'] = getattr(submission, 'displayOrdinalId', None)
            draft['displayOrdinalContentUrl'] = getattr(submission, 'displayOrdinalContentUrl', None)
            draft['displayOrdinalContentType'] = getattr(submission, 'displayOrdinalContentType', None)
            draft['displayingLinkedOrdinal'] = displaying_linked

    document_content = "Document content not available."
    calculated_pages = draft.get('pages', 1)
    calculated_words = draft.get('words', 0)
    # True when body is HTML (markdown → HTML or PDF iframe); use prose styling, not <pre>-like monospace.
    render_document_as_html = False

    if submission and _submission_uses_display_ordinal(submission):
        document_content, render_document_as_html, calculated_pages, calculated_words = _load_ordinal_document_body(
            submission.displayOrdinalContentUrl,
            submission.displayOrdinalContentType or '',
            draft,
        )
    elif submission and draft.get('sourceType') == 'ordinal':
        document_content, render_document_as_html, calculated_pages, calculated_words = _load_ordinal_document_body(
            getattr(submission, 'ordinalContentUrl', None),
            getattr(submission, 'ordinalContentType', '') or '',
            draft,
        )

    elif submission and submission.file_path and os.path.exists(submission.file_path):
        _, ext = os.path.splitext(submission.filename.lower())
        try:
            if ext in ['.txt', '.xml', '.md', '.markdown']:
                with open(submission.file_path, 'r', encoding='utf-8', errors='replace') as f:
                    raw_text = f.read()
                words = len(raw_text.split())
                calculated_pages = max(1, (words + 499) // 500)
                calculated_words = words
                if ext in ('.md', '.markdown'):
                    md_html = markdown_to_safe_preview_html(raw_text)
                    if md_html:
                        document_content = md_html
                        render_document_as_html = True
                    else:
                        document_content = raw_text
                elif ext == '.txt':
                    if text_looks_like_markdown(raw_text):
                        md_html = markdown_to_safe_preview_html(raw_text)
                        if md_html:
                            document_content = md_html
                            render_document_as_html = True
                        else:
                            document_content = raw_text
                    else:
                        document_content = raw_text
                else:
                    document_content = raw_text
            elif ext == '.docx':
                from docx import Document
                doc = Document(submission.file_path)
                content_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content_parts.append(cell.text)
                document_content = '\n\n'.join(content_parts)
                words = len(document_content.split())
                calculated_pages = max(1, (words + 499) // 500)
                calculated_words = words
            elif ext == '.pdf':
                from PyPDF2 import PdfReader
                reader = PdfReader(submission.file_path)
                calculated_pages = len(reader.pages) if reader.pages else 1
                calculated_words = calculated_pages * 275
                file_size = os.path.getsize(submission.file_path)
                file_size_kb = file_size / 1024
                render_document_as_html = True
                document_content = f'''
<div class="pdf-viewer-container">
    <div class="alert alert-info mb-3">
        <i class="bi bi-file-pdf"></i> PDF Document ({calculated_pages} pages, ~{calculated_words} words, {file_size_kb:.1f} KB)
    </div>
    <iframe src="/view/{draft_name}"
            type="application/pdf"
            style="width: 100%; height: 800px; border: 1px solid var(--card-border); border-radius: 4px;"
            title="PDF Document Viewer">
        <p>Your browser does not support PDF preview.
           <a href="/download/{draft_name}">Download the PDF</a> to view it.</p>
    </iframe>
</div>
'''
            else:
                document_content = f"Document content cannot be displayed for {ext.upper()} files. Please download to view."
        except Exception as e:
            document_content = f"Error loading document content: {str(e)}"
        if submission:
            draft['pages'] = calculated_pages
            draft['words'] = calculated_words

    elif draft and 'name' in draft:
        document_content = f"""INTERNET-DRAFT                                               {', '.join(draft.get('authors', []))}
Intended status: Informational                            Meta-Layer Initiative
Expires: {draft.get('date', 'TBD')}                                      {draft.get('date', 'TBD')}


{draft.get('title', 'Document Title')}


Abstract

{draft.get('abstract', 'Abstract not available.')}


1. Introduction

This document describes {draft.get('title', 'the subject matter')}.

The content of this draft is currently being developed and will be available
in the full document once published.

2. Status of This Memo

This Internet-Draft is submitted in full conformance with the provisions
of BCP 78 and BCP 79.

Meta-Layer Drafts are working documents of the Meta-Layer Task Force
(MLGH). These documents represent proposals and specifications for the
Meta-Layer ecosystem. The list of current Meta-Layer Drafts is available
in the MLGH datatracker.

Internet-Drafts are draft documents valid for a maximum of six months and
may be updated, replaced, or obsoleted by other documents at any time. It is
inappropriate to use Internet-Drafts as reference material or to cite them
other than as "work in progress."

This Internet-Draft will expire on {draft.get('date', 'TBD')}.


3. References

[MLGH] MLGH Datatracker, https://rfc.themetalayer.org/

Authors' Addresses

{chr(10).join([f'{author} <email@example.com>' for author in draft.get('authors', [])])}

Meta-Layer Initiative
"""

    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')
    current_user = get_current_user()

    _sub = submission
    if not _sub and draft:
        _sub = get_submission_by_ref(draft.get('name')) or get_submission_by_ref(draft.get('ml_number')) or get_submission_by_ref(draft_name)
    artifact_id = getattr(_sub, 'artifact_id', None) if _sub else None
    layer_slug = None
    supports = []
    opposes = []
    support_oppose_card_html = ''
    if artifact_id:
        artifact = Artifact.query.get(artifact_id)
        if artifact and artifact.layer_id:
            layer = Layer.query.get(artifact.layer_id)
            layer_slug = layer.slug if layer else None
            incoming = ArtifactRelation.query.filter(
                ArtifactRelation.to_object_type == 'artifact',
                ArtifactRelation.to_object_id == artifact_id,
            ).all()
            supports = [r for r in incoming if r.relation_type == 'supports']
            opposes = [r for r in incoming if r.relation_type == 'opposes']

            def _so_row(r):
                a = Artifact.query.get(r.from_object_id)
                t = (a.title or a.id[:8]) if a else r.from_object_id[:8]
                return f'<li class="list-group-item list-group-item-action"><a href="/layers/{layer_slug}/artifacts/{r.from_object_id}/" class="text-decoration-none">{html_mod.escape(str(t)[:50])}</a></li>'
            supports_li = ''.join(_so_row(r) for r in supports) if supports else '<li class="list-group-item text-muted small">No support yet</li>'
            opposes_li = ''.join(_so_row(r) for r in opposes) if opposes else '<li class="list-group-item text-muted small">No opposition yet</li>'
            if current_user:
                add_btns = f'''
                    <div class="d-flex gap-2 mt-2">
                        <button type="button" class="btn btn-outline-success btn-sm flex-grow-1" data-bs-toggle="modal" data-bs-target="#addSupportModal"><i class="fas fa-thumbs-up me-1"></i>Add support</button>
                        <button type="button" class="btn btn-outline-danger btn-sm flex-grow-1" data-bs-toggle="modal" data-bs-target="#addOppositionModal"><i class="fas fa-thumbs-down me-1"></i>Add opposition</button>
                    </div>
                    <div class="modal fade" id="addSupportModal" tabindex="-1"><div class="modal-dialog"><div class="modal-content">
                        <div class="modal-header"><h5 class="modal-title">Add support</h5><button type="button" class="btn-close" data-bs-dismiss="modal"></button></div>
                        <div class="modal-body">
                            <div class="mb-2"><label class="form-label">Title</label><input type="text" class="form-control" id="support-title" placeholder="Support for this proposal"></div>
                            <div class="mb-2"><label class="form-label">Summary (optional)</label><textarea class="form-control" id="support-summary" rows="2"></textarea></div>
                            <div id="support-alert" class="alert d-none"></div>
                        </div>
                        <div class="modal-footer"><button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button><button type="button" class="btn btn-success" id="support-submit-btn">Add support</button></div>
                    </div></div></div>
                    <div class="modal fade" id="addOppositionModal" tabindex="-1"><div class="modal-dialog"><div class="modal-content">
                        <div class="modal-header"><h5 class="modal-title">Add opposition</h5><button type="button" class="btn-close" data-bs-dismiss="modal"></button></div>
                        <div class="modal-body">
                            <div class="mb-2"><label class="form-label">Title</label><input type="text" class="form-control" id="opposition-title" placeholder="Opposition to this proposal"></div>
                            <div class="mb-2"><label class="form-label">Summary (optional)</label><textarea class="form-control" id="opposition-summary" rows="2"></textarea></div>
                            <div id="opposition-alert" class="alert d-none"></div>
                        </div>
                        <div class="modal-footer"><button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button><button type="button" class="btn btn-danger" id="opposition-submit-btn">Add opposition</button></div>
                    </div></div></div>
                    <script>
                    (function(){{const aid='{artifact_id}';
                    document.getElementById('support-submit-btn').addEventListener('click',async function(){{const btn=this;btn.disabled=true;const alert=document.getElementById('support-alert');alert.classList.add('d-none');try{{const r=await fetch('/api/artifacts/'+aid+'/support/',{{method:'POST',headers:{{'Content-Type':'application/json'}},body:JSON.stringify({{title:document.getElementById('support-title').value,summary:document.getElementById('support-summary').value}}),credentials:'same-origin'}});const d=await r.json();if(r.ok)location.reload();else{{alert.textContent=d.error||'Failed';alert.className='alert alert-danger';alert.classList.remove('d-none');}}}}catch(e){{alert.textContent=e.message;alert.className='alert alert-danger';alert.classList.remove('d-none');}}btn.disabled=false;}});
                    document.getElementById('opposition-submit-btn').addEventListener('click',async function(){{const btn=this;btn.disabled=true;const alert=document.getElementById('opposition-alert');alert.classList.add('d-none');try{{const r=await fetch('/api/artifacts/'+aid+'/opposition/',{{method:'POST',headers:{{'Content-Type':'application/json'}},body:JSON.stringify({{title:document.getElementById('opposition-title').value,summary:document.getElementById('opposition-summary').value}}),credentials:'same-origin'}});const d=await r.json();if(r.ok)location.reload();else{{alert.textContent=d.error||'Failed';alert.className='alert alert-danger';alert.classList.remove('d-none');}}}}catch(e){{alert.textContent=e.message;alert.className='alert alert-danger';alert.classList.remove('d-none');}}btn.disabled=false;}});
                    }})();
                    </script>
'''
            else:
                add_btns = '<p class="small text-muted mt-2 mb-0"><a href="/login/">Sign in</a> to add support or opposition.</p>'
            support_oppose_card_html = f'''
                <div class="card mt-3">
                    <div class="card-header"><h5>Support &amp; Opposition</h5></div>
                    <div class="card-body">
                        <div class="row g-2">
                            <div class="col-6"><h6 class="text-success small">Support ({len(supports)})</h6><ul class="list-group list-group-flush small">{supports_li}</ul></div>
                            <div class="col-6"><h6 class="text-danger small">Opposition ({len(opposes)})</h6><ul class="list-group list-group-flush small">{opposes_li}</ul></div>
                        </div>
                        {add_btns}
                        <a href="/layers/{layer_slug}/artifacts/{artifact_id}/" class="btn btn-outline-secondary btn-sm mt-2 w-100">View full artifact</a>
                    </div>
                </div>
'''
    if not layer_slug and _sub and getattr(_sub, 'layer_id', None):
        layer = Layer.query.get(_sub.layer_id)
        layer_slug = layer.slug if layer else None

    def _artifact_card_and_modal_html(draft, sub, aid, lslug, user):
        if not lslug or not user:
            return ''
        if not sub and not aid:
            return ''
        has_artifact = bool(aid)
        sub_id = getattr(sub, 'id', None) if sub else None
        artifact_types = ['proposal', 'evidence', 'insight', 'reflection', 'translation', 'implementation', 'decision', 'monument', 'bridge', 'submission']
        aid_js = f"'{aid}'" if aid else 'null'
        sub_id_js = f"'{sub_id}'" if sub_id else 'null'
        return f'''
        <div class="card mt-3">
            <div class="card-header"><h5>Artifact</h5></div>
            <div class="card-body">
                {f'<a href="/layers/{lslug}/artifacts/{aid}/" class="btn btn-outline-secondary btn-sm mb-2 w-100">View full artifact</a>' if aid else ''}
                <button type="button" class="btn btn-outline-primary btn-sm w-100" id="artifact-modal-btn" data-artifact-id={aid_js} data-submission-id={sub_id_js}>
                    <i class="fas fa-edit me-1"></i>{'Edit Artifact' if has_artifact else 'Create Artifact'}
                </button>
            </div>
        </div>
        <div class="modal fade" id="artifactModal" tabindex="-1">
            <div class="modal-dialog modal-lg">
                <div class="modal-content">
                    <div class="modal-header">
                        <h5 class="modal-title" id="artifactModalTitle">Edit Artifact</h5>
                        <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
                    </div>
                    <div class="modal-body">
                        <div id="artifact-alert" class="alert d-none mb-2"></div>
                        <div class="mb-2"><label class="form-label">Type</label><select class="form-select" id="artifact-type"><option value="">—</option>{''.join(f'<option value="{t}">{t}</option>' for t in artifact_types)}</select></div>
                        <div class="mb-2"><label class="form-label">Subtype</label><input type="text" class="form-control" id="artifact-subtype" placeholder="e.g. governance proposal"></div>
                        <div class="mb-2"><label class="form-label">Title</label><input type="text" class="form-control" id="artifact-title" placeholder="Artifact title"></div>
                        <div class="mb-2"><label class="form-label">Summary</label><textarea class="form-control" id="artifact-summary" rows="2" placeholder="Brief summary"></textarea></div>
                        <div class="mb-2"><label class="form-label">Body</label><textarea class="form-control" id="artifact-body" rows="4" placeholder="Full content"></textarea></div>
                        <div class="mb-2"><label class="form-label">URI</label><input type="text" class="form-control" id="artifact-uri" placeholder="https://..."></div>
                        <div class="mb-2"><label class="form-label">Status</label><select class="form-select" id="artifact-status"><option value="draft">draft</option><option value="published">published</option><option value="archived">archived</option></select></div>
                        <div class="row"><div class="col-6"><label class="form-label">Source language</label><input type="text" class="form-control" id="artifact-source-lang" placeholder="en"></div><div class="col-6"><label class="form-label">Current language</label><input type="text" class="form-control" id="artifact-current-lang" placeholder="en"></div></div>
                        <div class="mb-2 border-top pt-2 mt-2" id="kl-contribution-wrap" style="display:none;">
                            <label class="form-label">Contribution type <span class="text-muted">(optional)</span></label>
                            <select class="form-select" id="kl-contribution-type"><option value="">— Not set</option></select>
                            <p class="small text-muted mb-0">Helps others understand how to engage with this contribution.</p>
                        </div>
                        <div class="mb-2" id="kl-scaffold-wrap" style="display:none;"></div>
                    </div>
                    <div class="modal-footer">
                        <button type="button" class="btn btn-secondary" data-bs-dismiss="modal">Cancel</button>
                        <button type="button" class="btn btn-primary" id="artifact-save-btn">Save</button>
                    </div>
                </div>
            </div>
        </div>
        <script>
        (function(){{
            const btn = document.getElementById('artifact-modal-btn');
            if (!btn) return;
            const aid = btn.dataset.artifactId;
            const subId = btn.dataset.submissionId;
            const modal = new bootstrap.Modal(document.getElementById('artifactModal'));
            const modalTitle = document.getElementById('artifactModalTitle');
            const saveBtn = document.getElementById('artifact-save-btn');
            const alertEl = document.getElementById('artifact-alert');
            const fields = ['artifact_type','artifact_subtype','title','summary','body','uri','status','source_language','current_language'];
            const ids = {{artifact_type:'artifact-type',artifact_subtype:'artifact-subtype',title:'artifact-title',summary:'artifact-summary',body:'artifact-body',uri:'artifact-uri',status:'artifact-status',source_language:'artifact-source-lang',current_language:'artifact-current-lang'}};
            const KL_SCAFFOLD = {{
                inquiry: [{{k:'what_is_unclear',l:'What is unclear?',t:'ta'}},{{k:'status',l:'Status',t:'sel',o:['open','closed']}}],
                principle: [{{k:'why_matters',l:'Why does this matter?',t:'ta'}}],
                model: [{{k:'key_assumptions',l:'Key assumptions',t:'ta'}}],
                claim: [{{k:'why_believe',l:'Why do you believe this?',t:'ta'}}],
                decision: [{{k:'what_resolves',l:'What does this resolve?',t:'ta'}},{{k:'status',l:'Status',t:'sel',o:['draft','final']}}],
                gloss: [{{k:'definition',l:'Definition',t:'ta'}}],
                scenario: [{{k:'actors_context',l:'Actors / context',t:'ta'}}]
            }};
            let klSchema = null;
            async function ensureKlSchema() {{
                if (klSchema) return klSchema;
                try {{
                    const r = await fetch('/api/knowledge-layer/schema/', {{credentials:'same-origin'}});
                    klSchema = await r.json();
                }} catch (e) {{ klSchema = null; }}
                return klSchema;
            }}
            function showAlert(msg,type){{
                alertEl.textContent=msg; alertEl.className='alert alert-'+type; alertEl.classList.remove('d-none');
            }}
            function rebuildKlContribution() {{
                const wrap = document.getElementById('kl-contribution-wrap');
                const sel = document.getElementById('kl-contribution-type');
                const atEl = document.getElementById('artifact-type');
                if (!wrap || !sel || !atEl) return;
                if (!klSchema || !klSchema.feature_flags || !klSchema.feature_flags.knowledge_contribution_type_enabled) {{
                    wrap.style.display = 'none';
                    return;
                }}
                wrap.style.display = 'block';
                const at = (atEl.value || '').trim();
                const spec = klSchema.artifact_types && klSchema.artifact_types[at];
                const prev = sel.value;
                sel.innerHTML = '<option value="">— Not set</option>';
                if (spec && spec.allowed) {{
                    spec.allowed.forEach(function(v) {{ sel.add(new Option(v, v)); }});
                    if (prev && [...sel.options].some(function(o) {{ return o.value === prev; }})) sel.value = prev;
                }}
            }}
            function renderKlScaffold() {{
                const sw = document.getElementById('kl-scaffold-wrap');
                const kls = document.getElementById('kl-contribution-type');
                if (!sw || !kls) return;
                if (!klSchema || !klSchema.feature_flags || !klSchema.feature_flags.knowledge_scaffold_enabled) {{
                    sw.style.display = 'none';
                    sw.innerHTML = '';
                    return;
                }}
                const form = kls.value;
                const rows = form && KL_SCAFFOLD[form];
                if (!rows) {{ sw.style.display = 'none'; sw.innerHTML = ''; return; }}
                sw.style.display = 'block';
                const data = window.__klScaffoldData || {{}};
                let html = '<div class="border rounded p-2 bg-light"><div class="small fw-bold mb-2">Optional details</div>';
                rows.forEach(function(row) {{
                    const id = 'kl-sc-' + row.k;
                    const v = data[row.k] != null ? String(data[row.k]) : '';
                    if (row.t === 'sel') {{
                        html += '<div class="mb-2"><label class="form-label small">' + row.l + '</label><select class="form-select form-select-sm" id="'+id+'" data-kl-scaffold="'+row.k+'"><option value=""></option>';
                        (row.o || []).forEach(function(o) {{ html += '<option value="'+o+'"'+(v===o?' selected':'')+'>'+o+'</option>'; }});
                        html += '</select></div>';
                    }} else {{
                        html += '<div class="mb-2"><label class="form-label small">'+row.l+'</label><textarea class="form-control form-control-sm" id="'+id+'" rows="2" data-kl-scaffold="'+row.k+'"></textarea></div>';
                    }}
                }});
                html += '</div>';
                sw.innerHTML = html;
                rows.forEach(function(row) {{
                    const el = document.getElementById('kl-sc-' + row.k);
                    if (el && row.t !== 'sel' && data[row.k] != null) el.value = data[row.k];
                }});
            }}
            function collectKlScaffold(form) {{
                if (!form || !KL_SCAFFOLD[form]) return null;
                const out = {{}};
                document.querySelectorAll('[data-kl-scaffold]').forEach(function(el) {{
                    const k = el.getAttribute('data-kl-scaffold');
                    if (el.tagName === 'SELECT') {{
                        if (el.value) out[k] = el.value;
                    }} else {{
                        const t = el.value.trim();
                        if (t) out[k] = t;
                    }}
                }});
                return Object.keys(out).length ? out : null;
            }}
            function getPayload(){{
                const p={{}};
                for (const f of fields){{ const el=document.getElementById(ids[f]); if(el) p[f]=el.value===''?null:el.value; }}
                if (klSchema && klSchema.feature_flags && klSchema.feature_flags.knowledge_contribution_type_enabled) {{
                    const kls = document.getElementById('kl-contribution-type');
                    const vf = kls && kls.value ? kls.value : null;
                    p.knowledge_form = vf;
                    if (klSchema.feature_flags.knowledge_scaffold_enabled && vf) {{
                        const sc = collectKlScaffold(vf);
                        p.knowledge_scaffold = sc;
                    }}
                }}
                return p;
            }}
            function setFields(art){{
                for (const f of fields){{ const el=document.getElementById(ids[f]); if(el&&art[f]!==undefined) el.value=art[f]||''; }}
                window.__klScaffoldData = art.knowledge_scaffold || null;
                rebuildKlContribution();
                const kls = document.getElementById('kl-contribution-type');
                var kf = art.knowledge_form;
                if (kf === 'conviction') kf = 'claim';
                if (kls && kf && [...kls.options].some(function(o){{return o.value===kf;}})) kls.value = kf;
                else if (kls) kls.value = '';
                renderKlScaffold();
            }}
            document.getElementById('artifact-type').addEventListener('change', function() {{
                rebuildKlContribution();
                document.getElementById('kl-contribution-type').value = '';
                window.__klScaffoldData = null;
                renderKlScaffold();
            }});
            document.addEventListener('change', function(e) {{
                if (e.target && e.target.id === 'kl-contribution-type') renderKlScaffold();
            }});
            btn.addEventListener('click', async function(){{
                await ensureKlSchema();
                modalTitle.textContent = aid ? 'Edit Artifact' : 'Create Artifact';
                if (aid) {{
                    try {{
                        const r = await fetch('/api/artifacts/'+aid+'/', {{credentials:'same-origin'}});
                        const d = await r.json();
                        if (r.ok) setFields(d); else showAlert(d.error||'Failed to load','danger');
                    }} catch(e) {{ showAlert(e.message,'danger'); }}
                }} else {{
                    setFields({{}});
                    rebuildKlContribution();
                    renderKlScaffold();
                }}
                modal.show();
            }});
            saveBtn.addEventListener('click', async function(){{
                saveBtn.disabled=true; alertEl.classList.add('d-none');
                try {{
                    let d;
                    if (aid) {{
                        const r = await fetch('/api/artifacts/'+aid+'/', {{method:'PATCH',headers:{{'Content-Type':'application/json'}},body:JSON.stringify(getPayload()),credentials:'same-origin'}});
                        d = await r.json();
                        if (r.ok) {{ location.reload(); saveBtn.disabled=false; return; }}
                    }} else {{
                        if (!subId) {{ showAlert('No submission','danger'); saveBtn.disabled=false; return; }}
                        const r0 = await fetch('/api/submissions/'+subId+'/ensure-artifact/', {{method:'POST',headers:{{'Content-Type':'application/json'}},body:JSON.stringify({{}}),credentials:'same-origin'}});
                        d = await r0.json();
                        if (r0.ok && d.artifact_id) {{
                            const r1 = await fetch('/api/artifacts/'+d.artifact_id+'/', {{method:'PATCH',headers:{{'Content-Type':'application/json'}},body:JSON.stringify(getPayload()),credentials:'same-origin'}});
                            d = await r1.json();
                            if (r1.ok) {{ location.reload(); saveBtn.disabled=false; return; }}
                        }}
                    }}
                    showAlert(d.error||'Failed','danger');
                }} catch(e) {{ showAlert(e.message,'danger'); }}
                saveBtn.disabled=false;
            }});
        }})();
        </script>
        '''

    if draft.get('status') == 'approved' and draft.get('ml_number'):
        display_id = draft.get('ml_number')
    else:
        display_id = draft['name']
    is_revision = draft.get('is_revision', False)
    revision_number = draft.get('revision_number', '')
    revision_badge = f'<span class="badge bg-success ms-2">Revision {revision_number}</span>' if is_revision and revision_number else ''

    linked_display_rows = ''
    if draft.get('displayingLinkedOrdinal'):
        doid = draft.get('displayOrdinalId') or ''
        dot = (draft.get('displayOrdinalContentType') or '').replace('<', '')
        if doid:
            esc_id = html_mod.escape(doid, quote=True)
            link_html = (
                f'<a href="https://ordinals.com/inscription/{esc_id}" target="_blank" '
                f'class="text-decoration-none" style="color: var(--accent-color) !important;">'
                f'<code style="font-family: monospace; font-size: 0.85em;">{shorten_inscription_id(doid, 8)}</code></a>'
            )
        else:
            link_html = '<span class="text-muted">(no inscription id)</span>'
        linked_display_rows = (
            '<tr><td colspan="2" style="padding-top: 10px;"><hr style="border-color: var(--border-color);"></td></tr>'
            '<tr><td style="color: var(--text-secondary) !important;"><strong>Displayed body:</strong></td>'
            '<td style="color: var(--text-primary) !important;">'
            '<span class="badge bg-info me-2"><i class="bi bi-coin"></i> Linked ordinal</span>'
            f'{link_html}'
            '</td></tr>'
        )
        if dot:
            linked_display_rows += (
                '<tr><td style="color: var(--text-secondary) !important;"><strong>Display content type:</strong></td>'
                f'<td style="color: var(--text-primary) !important;">{html_mod.escape(dot)}</td></tr>'
            )

    display_body_card_html = ''
    if (
        _sub
        and current_user
        and _can_manage_submission_display_body(current_user, _sub)
        and (_sub.sourceType or 'file').strip().lower() == 'file'
    ):
        is_linked = bool(draft.get('displayingLinkedOrdinal'))
        pref_oid = (getattr(_sub, 'displayOrdinalId', None) or getattr(_sub, 'ordinalId', None) or '') or ''
        pref_url = (
            (getattr(_sub, 'displayOrdinalContentUrl', None) or '')
            if is_linked
            else (getattr(_sub, 'ordinalContentUrl', None) or getattr(_sub, 'displayOrdinalContentUrl', None) or '')
        )
        pref_ct = (
            (getattr(_sub, 'displayOrdinalContentType', None) or '')
            if is_linked
            else (
                getattr(_sub, 'ordinalContentType', None) or getattr(_sub, 'displayOrdinalContentType', None) or ''
            )
        )
        esc_dn = html_mod.escape(draft_name, quote=True)
        display_body_card_html = f'''
                        <div class="border-top pt-2 mt-2">
                            <h6 class="text-muted mb-2">Reader display</h6>
                            <p class="small text-muted mb-2">Show inscription text in the reader while keeping the uploaded file for download and history.</p>
                            <form method="post" action="/doc/draft/{esc_dn}/display-body/" class="mb-2">
                                <input type="hidden" name="display_body" value="file">
                                <button type="submit" class="btn btn-outline-secondary btn-sm w-100"{' disabled' if not is_linked else ''}>Use uploaded file for body</button>
                            </form>
                            <form method="post" action="/doc/draft/{esc_dn}/display-body/">
                                <input type="hidden" name="display_body" value="ordinal">
                                <div class="mb-2">
                                    <label class="form-label small mb-0">Inscription id (optional)</label>
                                    <input class="form-control form-control-sm" name="display_ordinal_id"
                                           value="{html_mod.escape(pref_oid, quote=True)}"
                                           placeholder="abc…i0">
                                </div>
                                <div class="mb-2">
                                    <label class="form-label small mb-0">Content URL</label>
                                    <input class="form-control form-control-sm" name="display_ordinal_content_url" required
                                           value="{html_mod.escape(pref_url, quote=True)}"
                                           placeholder="https://…">
                                </div>
                                <div class="mb-2">
                                    <label class="form-label small mb-0">Content type</label>
                                    <input class="form-control form-control-sm" name="display_ordinal_content_type"
                                           value="{html_mod.escape(pref_ct, quote=True)}"
                                           placeholder="text/plain;charset=utf-8">
                                </div>
                                <button type="submit" class="btn btn-info btn-sm w-100">Use ordinal for body</button>
                            </form>
                        </div>'''

    if draft.get('sourceType') == 'ordinal' or render_document_as_html:
        content_style = "font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; font-size: 1em; line-height: 1.6;"
    else:
        content_style = "font-family: 'Courier New', monospace; font-size: 0.9em; line-height: 1.4; white-space: pre-wrap;"

    workgroup_metadata_edit_html = ''
    if _sub and current_user and can_edit_submission_metadata(current_user, _sub):
        from models import Workgroup
        wg_query = Workgroup.query
        if _sub.layer_id:
            wg_query = wg_query.filter_by(layer_id=_sub.layer_id)
        workgroups = wg_query.order_by(Workgroup.name.asc()).all()
        current_group = (_sub.group or '').strip()
        wg_opts = '<option value="">— None —</option>'
        for wg in workgroups:
            if not wg.acronym:
                continue
            sel = ' selected' if wg.acronym == current_group else ''
            wg_opts += f'<option value="{html_mod.escape(wg.acronym)}"{sel}>{html_mod.escape(wg.name or wg.acronym)}</option>'
        sub_id_esc = html_mod.escape(_sub.id, quote=True)
        workgroup_metadata_edit_html = f'''
                <div class="card mt-3">
                    <div class="card-header"><h6 class="mb-0">Workgroup assignment</h6></div>
                    <div class="card-body">
                        <p class="text-muted small mb-2">Assign this draft to a workgroup (shown under Assigned documents on the workgroup page).</p>
                        <div class="row g-2 align-items-end">
                            <div class="col-md-8">
                                <select class="form-select form-select-sm" id="doc-workgroup-select">{wg_opts}</select>
                            </div>
                            <div class="col-md-4">
                                <button type="button" class="btn btn-sm btn-primary w-100" id="doc-workgroup-save">Save</button>
                            </div>
                        </div>
                        <div id="doc-workgroup-alert" class="alert d-none mt-2 mb-0" role="alert"></div>
                    </div>
                </div>
                <script>
                (function() {{
                    const btn = document.getElementById('doc-workgroup-save');
                    const sel = document.getElementById('doc-workgroup-select');
                    const alertEl = document.getElementById('doc-workgroup-alert');
                    if (!btn || !sel) return;
                    btn.addEventListener('click', async function() {{
                        btn.disabled = true;
                        if (alertEl) alertEl.classList.add('d-none');
                        try {{
                            const r = await fetch('/api/submissions/{sub_id_esc}/metadata/', {{
                                method: 'PATCH',
                                headers: {{ 'Content-Type': 'application/json' }},
                                credentials: 'same-origin',
                                body: JSON.stringify({{ group: sel.value }})
                            }});
                            const d = await r.json();
                            if (r.ok) {{
                                if (alertEl) {{
                                    alertEl.textContent = 'Workgroup updated.';
                                    alertEl.className = 'alert alert-success mt-2 mb-0';
                                    alertEl.classList.remove('d-none');
                                }}
                                setTimeout(function() {{ location.reload(); }}, 600);
                            }} else if (alertEl) {{
                                alertEl.textContent = d.error || 'Update failed';
                                alertEl.className = 'alert alert-danger mt-2 mb-0';
                                alertEl.classList.remove('d-none');
                            }}
                        }} catch (e) {{
                            if (alertEl) {{
                                alertEl.textContent = e.message || 'Update failed';
                                alertEl.className = 'alert alert-danger mt-2 mb-0';
                                alertEl.classList.remove('d-none');
                            }}
                        }}
                        btn.disabled = false;
                    }});
                }})();
                </script>'''

    content = f"""
    <div class="gh-page container mt-4">
        {gh_page_header(str(display_id), draft['title'], 'fa-file-alt', actions_html=f'<a href="/doc/draft/{quote(str(draft.get("name") or draft_name), safe="")}/read/" class="btn btn-primary btn-sm"><i class="bi bi-book"></i> Read full page</a>')}
        <div class="gh-detail-layout mt-3">
            <div class="gh-detail-main">
                <div class="card">
                    <div class="card-header">
                        <h5>Document Information</h5>
                    </div>
                    <div class="card-body">
                        <table class="table" style="color: var(--text-primary) !important;">
                            <tr><td style="color: var(--text-secondary) !important;"><strong>ID:</strong></td><td style="color: var(--text-primary) !important;">{display_id}</td></tr>
                            <tr><td style="color: var(--text-secondary) !important;"><strong>Title:</strong></td><td style="color: var(--text-primary) !important;">{draft['title']}</td></tr>
                            <tr><td style="color: var(--text-secondary) !important;"><strong>Status:</strong></td><td style="color: var(--text-primary) !important;"><span class="badge bg-secondary">{draft['status']}</span></td></tr>
                            <tr><td style="color: var(--text-secondary) !important;"><strong>Authors:</strong></td><td style="color: var(--text-primary) !important;">{', '.join(draft['authors'])}</td></tr>
                            {_document_workgroup_table_row(draft)}
                            <tr><td style="color: var(--text-secondary) !important;"><strong>Date:</strong></td><td style="color: var(--text-primary) !important;">{draft['date']}</td></tr>
                            {f'<tr><td colspan="2" style="padding-top: 15px;"><hr style="border-color: var(--border-color);"></td></tr><tr><td style="color: var(--text-secondary) !important;"><strong>Source:</strong></td><td style="color: var(--text-primary) !important;"><span class="badge bg-info"><i class="bi bi-coin"></i> Bitcoin Ordinal</span></td></tr>' if draft.get('sourceType') == 'ordinal' else f'<tr><td style="color: var(--text-secondary) !important;"><strong>Revision:</strong></td><td style="color: var(--text-primary) !important;">{draft["rev"]}</td></tr><tr><td style="color: var(--text-secondary) !important;"><strong>Pages:</strong></td><td style="color: var(--text-primary) !important;">{draft["pages"]}</td></tr><tr><td style="color: var(--text-secondary) !important;"><strong>Words:</strong></td><td style="color: var(--text-primary) !important;">{draft["words"]}</td></tr>'}
                            {f'<tr><td style="color: var(--text-secondary) !important;"><strong>Inscription #:</strong></td><td style="color: var(--text-primary) !important;">{draft["inscriptionNumber"]}</td></tr>' if draft.get('sourceType') == 'ordinal' and draft.get('inscriptionNumber') else ''}
                            {f'<tr><td style="color: var(--text-secondary) !important;"><strong>Block Height:</strong></td><td style="color: var(--text-primary) !important;">{draft["blockHeight"]}</td></tr>' if draft.get('sourceType') == 'ordinal' and draft.get('blockHeight') else ''}
                            {f'<tr><td style="color: var(--text-secondary) !important;"><strong>Timestamp:</strong></td><td style="color: var(--text-primary) !important;">{draft["inscriptionTimestamp"].strftime("%Y-%m-%d %H:%M UTC") if draft.get("inscriptionTimestamp") else "N/A"}</td></tr>' if draft.get('sourceType') == 'ordinal' else ''}
                            {f'<tr><td style="color: var(--text-secondary) !important;"><strong>Content Type:</strong></td><td style="color: var(--text-primary) !important;">{draft["ordinalContentType"]}</td></tr>' if draft.get('sourceType') == 'ordinal' and draft.get('ordinalContentType') else ''}
                            {f'<tr><td style="color: var(--text-secondary) !important;"><strong>Inscription ID:</strong></td><td style="color: var(--text-primary) !important;"><a href="https://ordinals.com/inscription/{draft["ordinalId"]}" target="_blank" class="text-decoration-none" style="color: var(--accent-color) !important;"><code style="font-family: monospace; font-size: 0.85em;">{shorten_inscription_id(draft["ordinalId"], 8)}</code></a></td></tr>' if draft.get('sourceType') == 'ordinal' and draft.get('ordinalId') else ''}
                            {linked_display_rows}
                        </table>
                    </div>
                </div>

                {workgroup_metadata_edit_html}

                <div class="card mt-3">
                    <div class="card-header">
                        <h5>Abstract</h5>
                    </div>
                    <div class="card-body">
                        <p>{draft.get('abstract', 'Abstract not available for this draft.')}</p>
                    </div>
                </div>

                <div class="card mt-3">
                    <div class="card-header d-flex justify-content-between align-items-center">
                        <h5 class="mb-0">Document Content</h5>
                        <div>
                            {'' if draft.get('sourceType') == 'ordinal' else f'''
                            <a href="/download/{draft['name']}" class="btn btn-sm btn-outline-primary" target="_blank">
                                <i class="fas fa-download me-1"></i>Download
                            </a>
                            <a href="/doc/draft/{draft['name']}.txt" class="btn btn-sm btn-outline-secondary" target="_blank">
                                <i class="fas fa-external-link-alt me-1"></i>View TXT
                            </a>
                            '''}
                        </div>
                    </div>
                    <div class="card-body">
                        <div class="document-content" style="{content_style} background-color: var(--input-bg) !important; color: var(--text-primary) !important; padding: 20px; border-radius: 8px; max-height: 800px; overflow-y: auto; border: 1px solid var(--input-border);">
{document_content}
                        </div>
                    </div>
                </div>
            </div>

            <div class="gh-detail-sidebar">
                {_artifact_card_and_modal_html(draft, _sub, artifact_id, layer_slug, current_user)}
                {support_oppose_card_html}
                <div class="card">
                    <div class="card-header">
                        <h5>Actions</h5>
                    </div>
                    <div class="card-body">
                        {f'<a href="/doc/draft/{draft["name"]}/comments/" class="btn btn-primary w-100 mb-2">View Comments ({Comment.query.filter_by(draft_name=draft_name).count()})</a>' if draft.get('status') == 'approved' else ''}
                        <a href="/doc/draft/{draft['name']}/history/" class="btn btn-secondary w-100 mb-2">View History</a>
                        <a href="/doc/draft/{draft['name']}/revisions/" class="btn btn-info w-100 mb-2">View Revisions</a>

                        {display_body_card_html}
                        {f'<a href="/submit/revision/{draft["name"]}/" class="btn btn-success w-100 mb-2"><i class="fas fa-plus me-1"></i>Submit New Revision</a>' if current_user and draft.get('status') == 'approved' else ''}
                        {'' if draft.get('sourceType') == 'ordinal' else f'<a href="/download/{draft["name"]}" class="btn btn-outline-primary w-100 mb-2">Download Document</a>'}
                        {render_draft_subscription_form_html(draft_name, current_user) if current_user and draft.get('status') == 'approved' else ''}
                        {'' if not current_user else ''}
                    </div>
                </div>

                <div class="card mt-3" id="votes-card" style="display: none;">
                    <div class="card-header">
                        <h5>Votes</h5>
                    </div>
                    <div class="card-body" id="votes-container">
                        <div class="spinner-border spinner-border-sm text-primary"></div>
                    </div>
                </div>

                {f'''<div class="card mt-3">
                    <div class="card-header">
                        <h5>Quick Comment</h5>
                    </div>
                    <div class="card-body">
                        <form method="POST" action="/doc/draft/{draft['name']}/comments/">
                            <div class="mb-3">
                                <textarea class="form-control" name="comment" rows="3" placeholder="Add a quick comment..." required></textarea>
                    </div>
                            <button type="submit" class="btn btn-success btn-sm w-100">Post Comment</button>
                        </form>
        </div>
    </div>''' if draft.get('status') == 'approved' else ''}

                <div class="card mt-3">
                    <div class="card-header">
                        <h5>Related Documents</h5>
                    </div>
                <div class="card-body">
                        <p>Related documents would appear here in the real datatracker.</p>
                    </div>
                    </div>
                </div>
            </div>
        </div>

        <script>
        const submissionId = '{draft["name"] if submission else draft_name}';

        async function loadDraftVotes() {{
            try {{
                const allProjects = await fetch('/api/layers/').then(r => r.json());
                let votes = [];

                for (const proj of allProjects.layers || []) {{
                    const res = await fetch(`/api/layers/${{proj.id}}/votes/`);
                    const data = await res.json();
                    const matchingVotes = (data.votes || []).filter(v => v.submission_id === submissionId);
                    votes.push(...matchingVotes);
                }}

                if (votes.length === 0) {{
                    return;
                }}

                document.getElementById('votes-card').style.display = 'block';
                const container = document.getElementById('votes-container');

                let html = '';
                for (const v of votes) {{
                    const statusBadge = v.status === 'active' ? '<span class="badge bg-success">Active</span>' : v.status === 'closed' ? '<span class="badge bg-secondary">Closed</span>' : '<span class="badge bg-info">Scheduled</span>';
                    const resultBadge = v.result ? '<span class="badge bg-' + (v.result === 'passed' ? 'success' : v.result === 'failed' ? 'danger' : 'warning') + ' ms-1">' + v.result + '</span>' : '';
                    html += '<div class="mb-3">';
                    html += '<h6><a href="/votes/' + v.public_id + '/">' + v.title + '</a> ' + statusBadge + resultBadge + '</h6>';
                    html += '<p class="small mb-1">' + (v.description || '') + '</p>';
                    html += '<p class="small text-muted mb-0">Ends: ' + new Date(v.end_at).toLocaleString() + '</p>';
                    html += '</div>';
                }}

                container.innerHTML = html;
            }} catch (e) {{
                console.error('Error loading votes:', e);
            }}
        }}

        loadDraftVotes();
        </script>
        """

    content = content.replace('{document_content}', document_content)

    if draft.get('status') == 'approved' and draft.get('ml_number'):
        title_id = draft.get('ml_number')
    else:
        title_id = draft['name']

    return _format_base_template(
        title=f"{title_id} - MLGH",
        theme=current_theme,
        user_menu=user_menu,
        content=content,
        build_number=BUILD_NUMBER,
    )


@bp.route('/doc/draft/<path:draft_name>/display-body/', methods=['POST'])
def draft_display_body(draft_name):
    """Switch draft reader body between uploaded file and a linked ordinal (file-backed submissions only)."""
    user = get_current_user()
    if not user:
        flash('You must be signed in to change display settings.', 'error')
        return redirect(url_for('documents.draft_detail', draft_name=draft_name))
    sub = get_submission_by_ref(draft_name)
    if not sub or not _can_manage_submission_display_body(user, sub):
        flash('You cannot change display settings for this draft.', 'error')
        return redirect(url_for('documents.draft_detail', draft_name=draft_name))

    mode = (request.form.get('display_body') or 'file').strip().lower()
    if mode != 'ordinal':
        sub.displayBodySource = 'file'
        sub.displayOrdinalId = None
        sub.displayOrdinalContentUrl = None
        sub.displayOrdinalContentType = None
        sub.displaySwitchedAt = datetime.utcnow()
        sub.displaySwitchedBy = user.get('name')
        details = 'Body display reset to uploaded file.'
        flash('Reader now uses the uploaded file.', 'success')
    else:
        url = (request.form.get('display_ordinal_content_url') or '').strip()
        if not url:
            flash('Content URL is required to show an ordinal body.', 'error')
            return redirect(url_for('documents.draft_detail', draft_name=draft_name))
        try:
            from services.url_safety import validate_ordinals_fetch_url
            url = validate_ordinals_fetch_url(url)
        except ValueError:
            flash('Only ordinals.com content URLs are allowed.', 'error')
            return redirect(url_for('documents.draft_detail', draft_name=draft_name))
        oid = (request.form.get('display_ordinal_id') or '').strip() or None
        ct = (request.form.get('display_ordinal_content_type') or '').strip() or None
        sub.displayBodySource = 'ordinal'
        sub.displayOrdinalContentUrl = url
        sub.displayOrdinalId = oid
        sub.displayOrdinalContentType = ct
        sub.displaySwitchedAt = datetime.utcnow()
        sub.displaySwitchedBy = user.get('name')
        details = f'Body display set to linked ordinal (id={oid or "n/a"})'
        flash('Reader now shows the linked ordinal body; the uploaded file is unchanged.', 'success')

    db.session.add(
        DocumentHistory(
            draft_name=draft_name,
            action='display_body_changed',
            user=user.get('name') or 'unknown',
            details=details,
        )
    )
    db.session.commit()
    return redirect(url_for('documents.draft_detail', draft_name=draft_name))


@bp.route('/doc/draft/<path:draft_name>/comments/', methods=['GET', 'POST'])
@require_auth
def draft_comments(draft_name):
    from services.rendering import _format_base_template, generate_user_menu
    from config import BUILD_NUMBER

    DRAFTS = _get_drafts()
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)

    submission = None
    if not draft:
        submission = get_submission_by_ref(draft_name)
        if submission:
            draft = {
                'name': submission.id,
                'title': submission.title,
                'authors': submission.authors,
                'status': submission.status,
                'group': submission.group,
                'date': submission.submitted_at.strftime('%Y-%m-%d') if submission.submitted_at else '',
                'ml_number': submission.ml_number
            }

    if not draft:
        return "Document not found", 404

    display_id = draft.get('ml_number') or draft_name

    user_menu = generate_user_menu()
    current_theme = session.get('theme', get_current_user().get('theme', 'dark') if get_current_user() else 'dark')
    current_user = get_current_user()

    if request.method == 'POST':
        action = request.form.get('action', 'comment')

        if action == 'comment':
            comment_text = request.form.get('comment', '').strip()
            if comment_text:
                new_comment = Comment(
                    draft_name=draft_name,
                    text=comment_text,
                    author=current_user['name']
                )
                db.session.add(new_comment)
                db.session.flush()
                sub = (
                    Submission.query.filter(
                        or_(
                            Submission.draft_name == draft_name,
                            Submission.parent_draft_name == draft_name,
                        ),
                        Submission.status == 'approved',
                    )
                    .order_by(Submission.submitted_at.desc().nullslast())
                    .first()
                )
                layer_id = getattr(sub, 'layer_id', None) if sub else None
                evt = emit_event(
                    'draft_comment_added',
                    actor_type='user',
                    actor_id=current_user['id'],
                    subject_type='comment',
                    subject_id=str(new_comment.id),
                    layer_id=layer_id,
                    payload={'draft_name': draft_name, 'preview': comment_text[:200]},
                )
                db.session.commit()
                add_to_document_history(draft_name, 'Comment added', current_user['name'], f'Added comment: {comment_text[:50]}...')
                dispatch_document_followers(
                    draft_name=draft_name,
                    event_type='draft_comment_added',
                    event_log=evt,
                    actor_user_id=current_user['id'],
                    title=f'New comment on {draft_name}',
                    body=comment_text[:500],
                    link_path=f'/doc/draft/{draft_name}/comments/',
                )
                flash('Comment added successfully!', 'success')
                return redirect(url_for('documents.draft_comments', draft_name=draft_name))
            else:
                flash('Please enter a comment.', 'error')

        elif action == 'like':
            comment_id = request.form.get('comment_id')
            if comment_id:
                liked = toggle_comment_like(draft_name, comment_id, current_user['name'])
                action_text = 'liked' if liked else 'unliked'
                flash(f'Comment {action_text}!', 'success')
                return redirect(url_for('documents.draft_comments', draft_name=draft_name))
            else:
                flash('Invalid comment ID.', 'error')

        elif action == 'reply':
            parent_comment_id = request.form.get('parent_comment_id')
            reply_text = request.form.get('reply_text', '').strip()
            if reply_text and parent_comment_id:
                reply = add_comment_reply(draft_name, parent_comment_id, reply_text, current_user)
                sub = (
                    Submission.query.filter(
                        or_(
                            Submission.draft_name == draft_name,
                            Submission.parent_draft_name == draft_name,
                        ),
                        Submission.status == 'approved',
                    )
                    .order_by(Submission.submitted_at.desc().nullslast())
                    .first()
                )
                layer_id = getattr(sub, 'layer_id', None) if sub else None
                evt = emit_event(
                    'draft_comment_added',
                    actor_type='user',
                    actor_id=current_user['id'],
                    subject_type='comment',
                    subject_id=str(reply.id),
                    layer_id=layer_id,
                    payload={'draft_name': draft_name, 'preview': reply_text[:200], 'is_reply': True},
                )
                db.session.commit()
                dispatch_document_followers(
                    draft_name=draft_name,
                    event_type='draft_comment_added',
                    event_log=evt,
                    actor_user_id=current_user['id'],
                    title=f'New reply on {draft_name}',
                    body=reply_text[:500],
                    link_path=f'/doc/draft/{draft_name}/comments/',
                )
                flash('Reply added successfully!', 'success')
                return redirect(url_for('documents.draft_comments', draft_name=draft_name))
            else:
                flash('Please enter a reply.', 'error')

        elif action == 'edit':
            comment_id = request.form.get('comment_id')
            new_text = request.form.get('new_text', '').strip()
            if comment_id and new_text:
                comment = Comment.query.filter_by(id=comment_id).first()
                if comment and comment.author == current_user['name']:
                    time_diff = datetime.utcnow() - comment.timestamp
                    time_limit = timedelta(minutes=EDIT_DELETE_TIME_MINUTES)
                    if time_diff <= time_limit and not comment.is_deleted:
                        if not comment.original_text:
                            comment.original_text = comment.text
                        comment.text = new_text
                        comment.edited_at = datetime.utcnow()
                        db.session.commit()
                        flash('Comment updated successfully!', 'success')
                    else:
                        flash('Edit time limit has expired.', 'error')
                else:
                    flash('You can only edit your own comments.', 'error')
                return redirect(url_for('documents.draft_comments', draft_name=draft_name))
            else:
                flash('Invalid comment or empty text.', 'error')

        elif action == 'delete':
            comment_id = request.form.get('comment_id')
            if comment_id:
                comment = Comment.query.filter_by(id=comment_id).first()
                if comment and comment.author == current_user['name']:
                    time_diff = datetime.utcnow() - comment.timestamp
                    time_limit = timedelta(minutes=EDIT_DELETE_TIME_MINUTES)
                    if time_diff <= time_limit and not comment.is_deleted:
                        comment.is_deleted = True
                        comment.text = '[Deleted]'
                        db.session.commit()
                        flash('Comment deleted successfully!', 'success')
                    else:
                        flash('Delete time limit has expired.', 'error')
                else:
                    flash('You can only delete your own comments.', 'error')
                return redirect(url_for('documents.draft_comments', draft_name=draft_name))
            else:
                flash('Invalid comment ID.', 'error')

    all_comments = build_comment_tree(draft_name)
    comments_html = render_comment_tree(all_comments, draft_name, get_current_user, get_comment_likes, is_comment_liked)

    content = f"""
    <div class="gh-page container mt-4">
        {gh_breadcrumb([('Home', '/'), ('Documents', '/doc/all/'), (display_id, f'/doc/draft/{draft_name}/'), ('Comments', None)])}
        {gh_page_header(f'Comments — {display_id}', draft['title'], 'fa-comments')}
        <div class="mb-4">
            <a href="/doc/draft/{draft_name}/" class="btn btn-secondary me-2">
                <i class="fas fa-arrow-left me-1"></i>Back to Draft
            </a>
            <a href="/doc/draft/{draft_name}/history/" class="btn btn-outline-secondary me-2">History</a>
            <a href="/doc/draft/{draft_name}/revisions/" class="btn btn-outline-secondary">Revisions</a>
        </div>

        <div class="row">
            <div class="col-md-8">
                <h3>Comments ({len(all_comments)})</h3>
                <div id="flash-messages"></div>
                {comments_html}

                <div class="card mt-4">
                    <div class="card-header">
                        <h5>Add a Comment</h5>
                    </div>
                    <div class="card-body">
                        <form method="POST">
                            <div class="mb-3">
                                <label for="comment" class="form-label">Your Comment</label>
                                <textarea class="form-control" id="comment" name="comment" rows="4" placeholder="Enter your comment here..." required></textarea>
                            </div>
                            <button type="submit" class="btn btn-primary">Submit Comment</button>
                        </form>
                    </div>
                </div>
            </div>

            <div class="col-md-4">
                <div class="card">
                    <div class="card-header">
                        <h5>Document Info</h5>
                    </div>
                    <div class="card-body">
                        <p><strong>Title:</strong> {draft['title']}</p>
                        <p><strong>Authors:</strong> {', '.join(draft['authors'])}</p>
                        <p><strong>Status:</strong> <span class="badge bg-secondary">{draft['status']}</span></p>
                        <p><strong>Last Updated:</strong> {draft['date']}</p>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <script>
        function toggleLike(commentId) {{
            const form = document.createElement('form');
            form.method = 'POST';
            form.style.display = 'none';
            const actionInput = document.createElement('input');
            actionInput.type = 'hidden';
            actionInput.name = 'action';
            actionInput.value = 'like';
            const commentIdInput = document.createElement('input');
            commentIdInput.type = 'hidden';
            commentIdInput.name = 'comment_id';
            commentIdInput.value = commentId;
            form.appendChild(actionInput);
            form.appendChild(commentIdInput);
            document.body.appendChild(form);
            form.submit();
        }}

        function toggleReply(commentId) {{
            const replyForm = document.getElementById('reply-form-' + commentId);
            if (replyForm) {{
                if (replyForm.style.display === 'none' || replyForm.style.display === '') {{
                replyForm.style.display = 'block';
            }} else {{
                replyForm.style.display = 'none';
            }}
            }}
        }}

        function editComment(commentId) {{
            const commentCard = document.getElementById('comment-' + commentId);
            if (!commentCard) return;
            const commentText = commentCard.querySelector('p.mb-2');
            if (!commentText) return;
            const currentText = commentText.textContent.trim();
            const editForm = document.createElement('form');
            editForm.method = 'POST';
            editForm.style.marginTop = '10px';
            const textarea = document.createElement('textarea');
            textarea.className = 'form-control';
            textarea.name = 'new_text';
            textarea.rows = 3;
            textarea.value = currentText;
            textarea.required = true;
            const actionInput = document.createElement('input');
            actionInput.type = 'hidden';
            actionInput.name = 'action';
            actionInput.value = 'edit';
            const commentIdInput = document.createElement('input');
            commentIdInput.type = 'hidden';
            commentIdInput.name = 'comment_id';
            commentIdInput.value = commentId;
            const buttonDiv = document.createElement('div');
            buttonDiv.className = 'd-flex gap-2 mt-2';
            const saveBtn = document.createElement('button');
            saveBtn.type = 'submit';
            saveBtn.className = 'btn btn-sm btn-primary';
            saveBtn.textContent = 'Save';
            const cancelBtn = document.createElement('button');
            cancelBtn.type = 'button';
            cancelBtn.className = 'btn btn-sm btn-secondary';
            cancelBtn.textContent = 'Cancel';
            cancelBtn.onclick = function() {{
                commentText.style.display = 'block';
                editForm.remove();
            }};
            buttonDiv.appendChild(saveBtn);
            buttonDiv.appendChild(cancelBtn);
            editForm.appendChild(actionInput);
            editForm.appendChild(commentIdInput);
            editForm.appendChild(textarea);
            editForm.appendChild(buttonDiv);
            commentText.style.display = 'none';
            commentText.parentNode.insertBefore(editForm, commentText.nextSibling);
        }}

        function deleteComment(commentId) {{
            if (!confirm('Are you sure you want to delete this comment? This action cannot be undone.')) {{
                return;
            }}
            const form = document.createElement('form');
            form.method = 'POST';
            form.style.display = 'none';
            const actionInput = document.createElement('input');
            actionInput.type = 'hidden';
            actionInput.name = 'action';
            actionInput.value = 'delete';
            const commentIdInput = document.createElement('input');
            commentIdInput.type = 'hidden';
            commentIdInput.name = 'comment_id';
            commentIdInput.value = commentId;
            form.appendChild(actionInput);
            form.appendChild(commentIdInput);
            document.body.appendChild(form);
            form.submit();
        }}
    </script>
"""

    return _format_base_template(title=f"Comments - {draft_name}", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER)


@bp.route('/doc/draft/<path:draft_name>/history/')
def draft_history(draft_name):
    from services.rendering import _format_base_template, generate_user_menu
    from config import BUILD_NUMBER

    DRAFTS = _get_drafts()
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)

    submission = None
    if not draft:
        submission = get_submission_by_ref(draft_name)
        if submission:
            draft = {
                'name': submission.id,
                'title': submission.title,
                'authors': submission.authors,
                'status': submission.status,
                'group': submission.group,
                'date': submission.submitted_at.strftime('%Y-%m-%d') if submission.submitted_at else '',
                'ml_number': submission.ml_number,
            }

    if not draft:
        return "Document not found", 404

    display_id = draft.get('ml_number', draft_name) or draft_name

    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')

    history = DocumentHistory.query.filter_by(draft_name=draft_name).order_by(DocumentHistory.timestamp.desc()).all()

    history_html = ""
    if history:
        for entry in history:
            history_html += f"""
            <div class="card mb-3">
                        <div class="card-body">
                    <div class="d-flex justify-content-between align-items-start mb-2">
                        <span class="badge bg-primary">{entry.action}</span>
                        <small class="text-muted">{entry.timestamp.strftime('%Y-%m-%d %H:%M')}</small>
                            </div>
                    <p class="mb-1"><strong>User:</strong> {entry.user}</p>
                    <p class="mb-0">{entry.details}</p>
                        </div>
                    </div>
            """
    else:
        history_html = """
        <div class="alert alert-info">
            <i class="fas fa-info-circle me-2"></i>
            No history available for this draft.
        </div>
        """

    content = f"""
    <div class="gh-page container mt-4">
        {gh_breadcrumb([('Home', '/'), ('Documents', '/doc/all/'), (display_id, f'/doc/draft/{draft_name}/'), ('History', None)])}
        {gh_page_header(f'History — {display_id}', draft['title'], 'fa-history')}
        <div class="mb-4">
            <a href="/doc/draft/{draft_name}/" class="btn btn-secondary me-2">
                <i class="fas fa-arrow-left me-1"></i>Back to Draft
            </a>
            <a href="/doc/draft/{draft_name}/comments/" class="btn btn-outline-secondary me-2">Comments</a>
            <a href="/doc/draft/{draft_name}/revisions/" class="btn btn-outline-secondary">Revisions</a>
        </div>

                {history_html}
            </div>
    """

    return _format_base_template(title=f"History - {display_id}", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER)


def _redirect_after_subscription(draft_name: str):
    """Same-origin relative path only (avoid open redirects)."""
    next_raw = (request.form.get('next') or '').strip()
    if next_raw.startswith('/') and not next_raw.startswith('//'):
        return redirect(next_raw)
    return redirect(url_for('documents.draft_detail', draft_name=draft_name))


@bp.route('/doc/draft/<path:draft_name>/subscriptions/', methods=['POST'])
def update_draft_subscriptions(draft_name):
    """Save per-event subscription matrix or clear all (draft detail + notifications hub)."""
    current_user = get_current_user()
    if not current_user:
        flash('You must be logged in to manage subscriptions.', 'error')
        return redirect(url_for('documents.draft_detail', draft_name=draft_name))

    if request.form.get('clear_all'):
        replace_draft_subscriptions_matrix(current_user['id'], draft_name, {})
        db.session.commit()
        flash('Removed all notification subscriptions for this draft.', 'success')
        return _redirect_after_subscription(draft_name)

    matrix = matrix_from_subscription_post(request.form)
    replace_draft_subscriptions_matrix(current_user['id'], draft_name, matrix)
    db.session.commit()
    if not matrix:
        flash('No channels enabled — this draft has no notification subscriptions.', 'info')
    else:
        flash('Subscription settings saved.', 'success')
    return _redirect_after_subscription(draft_name)


@bp.route('/doc/draft/<path:draft_name>/revisions/')
def draft_revisions(draft_name):
    from services.rendering import _format_base_template, generate_user_menu
    from config import BUILD_NUMBER

    current_user = get_current_user()
    DRAFTS = _get_drafts()
    draft = next((d for d in DRAFTS if d['name'] == draft_name), None)

    submission = None
    original_submission_id = None
    if not draft:
        submission = get_submission_by_ref(draft_name)
        if submission:
            if getattr(submission, 'is_revision', False) and getattr(submission, 'parent_draft_name', ''):
                original_submission_id = submission.parent_draft_name
            else:
                original_submission_id = submission.id

            sub_p, sub_w = submission_file_pages_words(submission)
            draft = {
                'name': submission.id,
                'title': submission.title,
                'authors': submission.authors,
                'status': submission.status,
                'group': submission.group,
                'date': submission.submitted_at.strftime('%Y-%m-%d') if submission.submitted_at else '',
                'rev': getattr(submission, 'revision_number', '00') or '00',
                'pages': sub_p,
                'words': sub_w,
                'is_revision': getattr(submission, 'is_revision', False),
                'parent_draft_name': getattr(submission, 'parent_draft_name', ''),
                'original_submission_id': original_submission_id,
                'ml_number': submission.ml_number,
            }

    if not draft:
        return "Document not found", 404

    display_id = draft.get('ml_number', draft_name) or draft_name

    if submission:
        cp, cw = submission_file_pages_words(submission)
        draft['pages'] = cp
        draft['words'] = cw

    user_menu = generate_user_menu()
    current_theme = session.get('theme', 'dark')

    original_id = draft.get('original_submission_id', draft['name'])
    original_submission = get_submission_by_ref(original_id)
    orig_pages, orig_words = (
        submission_file_pages_words(original_submission) if original_submission else (1, 0)
    )

    all_revisions = Submission.query.filter(
        Submission.parent_draft_name == original_id,
        Submission.is_revision == True,
        Submission.status.in_(['approved', 'published'])
    ).order_by(Submission.revision_number.desc()).all()

    revisions_list_html = ""
    for rev in all_revisions:
        status_badge_class = {
            'submitted': 'bg-warning text-dark',
            'approved': 'bg-success',
            'rejected': 'bg-danger',
            'published': 'bg-info'
        }.get(rev.status, 'bg-secondary')

        wc_block = revision_notes_to_safe_html(getattr(rev, 'what_changed', '') or '')
        what_changed_html = (
            f'<div class="what-changed-notes"><p class="mb-2"><strong>What changed:</strong></p>{wc_block}</div>'
            if wc_block else ''
        )

        is_current = (rev.id == draft['name'])
        current_badge = '<span class="badge bg-primary ms-2">Current</span>' if is_current else ''

        rev_pages, rev_words = submission_file_pages_words(rev)

        revisions_list_html += f"""
        <div class="card mb-3">
            <div class="card-header d-flex justify-content-between align-items-center">
                <h6 class="mb-0">
                    <a href="/doc/draft/{rev.id}/" class="text-decoration-none">Revision {rev.revision_number}</a>
                    {current_badge}
                </h6>
                <span class="badge {status_badge_class}">{rev.status.title()}</span>
            </div>
            <div class="card-body">
                <p class="mb-2"><strong>Published:</strong> {rev.approved_at.strftime('%Y-%m-%d') if rev.approved_at and rev.status == 'approved' else (rev.submitted_at.strftime('%Y-%m-%d') if rev.submitted_at else 'N/A')}</p>
                <p class="mb-2"><strong>Pages:</strong> {rev_pages} | <strong>Words:</strong> {rev_words}</p>
                {what_changed_html}
            </div>
        </div>
        """

    revisions_html = f"""
                <h4>Revision History</h4>
                {revisions_list_html if revisions_list_html else '<p class="text-muted">No revisions yet.</p>'}

                <div class="card mt-3">
                    <div class="card-header d-flex justify-content-between align-items-center">
                        <h6 class="mb-0">
                            <a href="/doc/draft/{original_id}/" class="text-decoration-none">Original Version (Rev 00)</a>
                        </h6>
                        <span class="badge bg-success">Approved</span>
                    </div>
                    <div class="card-body">
                        <p class="mb-2"><strong>Published:</strong> {original_submission.approved_at.strftime('%Y-%m-%d') if original_submission and original_submission.approved_at else (original_submission.submitted_at.strftime('%Y-%m-%d') if original_submission and original_submission.submitted_at else draft['date'])}</p>
                        <p class="mb-0"><strong>Pages:</strong> {orig_pages} | <strong>Words:</strong> {orig_words}</p>
                    </div>
                </div>

    <div class="alert alert-info mt-3">
        <i class="fas fa-info-circle me-2"></i>
        Detailed revision history and diff viewing would be implemented in a full datatracker system.
            </div>
    """

    content = f"""
    <div class="gh-page container mt-4">
        {gh_breadcrumb([('Home', '/'), ('Documents', '/doc/all/'), (display_id, f'/doc/draft/{draft_name}/'), ('Revisions', None)])}
        {gh_page_header(f'Revisions — {display_id}', draft['title'], 'fa-code-branch')}
        <div class="mb-4">
            <a href="/doc/draft/{draft_name}/" class="btn btn-secondary me-2">
                <i class="fas fa-arrow-left me-1"></i>Back to Draft
            </a>
            {f'<a href="/submit/revision/{draft_name}/" class="btn btn-success me-2"><i class="fas fa-plus me-1"></i>Submit New Revision</a>' if current_user and draft.get('status') == 'approved' else ''}
            <a href="/doc/draft/{draft_name}/comments/" class="btn btn-outline-secondary me-2">Comments</a>
            <a href="/doc/draft/{draft_name}/history/" class="btn btn-outline-secondary">History</a>
        </div>

        {revisions_html}
    </div>
    """

    return _format_base_template(title=f"Revisions - {display_id}", theme=current_theme, user_menu=user_menu, content=content, build_number=BUILD_NUMBER)


# ============================================================================
# Test/draft pages and annotations
# ============================================================================

def _safe_draft_path(filename):
    """Resolve filename under drafts/; return None if path escapes."""
    drafts_dir = os.path.abspath(os.path.join(current_app.root_path, 'drafts'))
    if not filename or ".." in filename or filename.startswith("/"):
        return None
    path = os.path.abspath(os.path.normpath(os.path.join(drafts_dir, filename)))
    if not path.startswith(drafts_dir):
        return None
    return path


@bp.route('/test/', methods=['GET'])
@bp.route('/test/<path:filename>', methods=['GET'])
def draft_page(filename=None):
    """Serve drafts/<filename> as text/html. /test/ serves digitalartifacts.htm."""
    name = filename or "digitalartifacts.htm"
    path = _safe_draft_path(name)
    if not path or not os.path.isfile(path):
        return jsonify({'error': 'Draft page not found'}), 404
    return send_file(path, mimetype='text/html; charset=utf-8')


